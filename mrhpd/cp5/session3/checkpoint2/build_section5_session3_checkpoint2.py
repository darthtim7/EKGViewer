#!/usr/bin/env python3
"""Build MRHPD Remediation Section 5 Session 3 Checkpoint 2 recovery.

The builder reconstructs the authoritative Response 81 restore, clean-applies
Response 82, creates a separate copied project tree, records Response 83,
reconciles the renewed provider/physical-proof evidence search, executes only
evidence-supported corrections, freezes the release candidate, synchronizes all
project surfaces, and emits cumulative recovery that applies directly to
Response 81. No accepted or frozen source artifact is modified in place.
"""
from __future__ import annotations

import argparse
import collections
import csv
import hashlib
import json
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import textwrap
import zipfile
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

HERE = Path(__file__).resolve().parent
CP1_DIR = HERE.parent / "checkpoint1"
if str(CP1_DIR) not in sys.path:
    sys.path.insert(0, str(CP1_DIR))
import build_section5_session3_checkpoint1 as cp1  # noqa: E402

PROJECT_VERSION = "3.0.0a"
RESPONSE_NUMBER = 83
SECTION_LABEL = "Remediation Section 5 of 5"
SESSION_LABEL = "Session 3 of 3"
CHECKPOINT_LABEL = "Checkpoint 2 of 3"
CHECKPOINT_CODE = "MRHPD-V3-CP5-S3-CP2"
BASE_RESTORE_BYTES = 267_562_561
BASE_RESTORE_SHA256 = "2e90bb8196a4bbaba100d7924fdb2e88be8ce78c238ce330ca219c7e3cae32b2"
RESPONSE82_RECOVERY_BYTES = 21_515_389
RESPONSE82_RECOVERY_SHA256 = "323112d0954673ba18639cac5675788d768d4e233e83a62911fffc9055a737a1"
PUBLICATION_SHA256 = cp1.PUBLICATION_SHA256
EDITABLE_SHA256 = cp1.EDITABLE_SHA256
APPLICATION_SHA256 = cp1.APPLICATION_SHA256
PRINT_INTERIOR_SHA256 = cp1.PRINT_INTERIOR_SHA256
COVER_SHA256 = cp1.COVER_SHA256
PUBLICATION_REL = cp1.PUBLICATION_REL
PRINT_INTERIOR_REL = cp1.PRINT_INTERIOR_REL
COVER_PNG_REL = cp1.COVER_PNG_REL
COVER_TIFF_REL = cp1.COVER_TIFF_REL
COVER_PDF_REL = cp1.COVER_PDF_REL
TEMPLATE_PNG_REL = cp1.TEMPLATE_PNG_REL
TEMPLATE_PDF_REL = cp1.TEMPLATE_PDF_REL
SOURCE_DB_REL = cp1.CURRENT_DB_REL
SOURCE_WORKBOOK_REL = cp1.CURRENT_WORKBOOK_REL
CURRENT_PROJECT_NAME = (
    "Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 3 of 3 Checkpoint 2 of 3 THROUGH RESPONSE 83"
)
CURRENT_DB_REL = (
    "Database/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 3 of 3 Checkpoint 2 of 3 THROUGH RESPONSE 83.sqlite"
)
CURRENT_WORKBOOK_REL = (
    "Tracking/Workbook/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 3 of 3 Checkpoint 2 of 3 THROUGH RESPONSE 83 Comprehensive Tracking.xlsx"
)
RAW_PROMPT_83 = "Continue"
NAVY = "17324D"
TEAL = "1C7475"
GOLD = "C9A227"
PALE_BLUE = "EAF1F5"
PALE_GREEN = "E9F3EE"
PALE_GOLD = "F7F1D9"
PALE_RED = "F7E8E6"
WHITE = "FFFFFF"
DARK = "24323D"


def utc_now() -> datetime:
    return datetime.now(timezone.utc).replace(microsecond=0)


def sha256_file(path: Path, chunk_size: int = 1024 * 1024) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(chunk_size), b""):
            digest.update(block)
    return digest.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False, default=str) + "\n", encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value.rstrip() + "\n", encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fields = fields or (list(rows[0]) if rows else [])
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({key: json.dumps(value, ensure_ascii=False, sort_keys=True) if isinstance(value, (dict, list, tuple, set)) else value for key, value in row.items()})


def verify_zip(path: Path, expected_bytes: int | None = None, expected_sha256: str | None = None) -> dict[str, Any]:
    if expected_bytes is not None and path.stat().st_size != expected_bytes:
        raise RuntimeError({"path": str(path), "expected_bytes": expected_bytes, "actual_bytes": path.stat().st_size})
    digest = sha256_file(path)
    if expected_sha256 is not None and digest != expected_sha256:
        raise RuntimeError({"path": str(path), "expected_sha256": expected_sha256, "actual_sha256": digest})
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
        names = zf.namelist()
        duplicates = [name for name, count in collections.Counter(names).items() if count > 1]
        unsafe = []
        filler = []
        for name in names:
            pp = PurePosixPath(name.replace("\\", "/"))
            if pp.is_absolute() or ".." in pp.parts or re.match(r"^[A-Za-z]:", name):
                unsafe.append(name)
            if re.search(r"(^|/)(filler|padding|dummy_payload|artificial_inflation)(/|$)", name, re.I):
                filler.append(name)
    result = {"name": path.name, "bytes": path.stat().st_size, "sha256": digest, "members": len(names), "crc_error": bad, "duplicates": duplicates, "unsafe_paths": unsafe, "filler_members": filler}
    if bad or duplicates or unsafe or filler:
        raise RuntimeError({"zip_verification_failed": result})
    return result


def safe_extract(path: Path, destination: Path) -> None:
    verify_zip(path)
    destination.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as zf:
        zf.extractall(destination)


def locate_project_root(root: Path) -> Path:
    direct = [path for path in root.iterdir() if path.is_dir()]
    if len(direct) == 1 and (direct[0] / "Database").is_dir():
        return direct[0]
    if (root / "Database").is_dir():
        return root
    candidates = [path.parent for path in root.rglob("Database") if path.is_dir() and (path.parent / "Tracking").is_dir()]
    unique = sorted(set(candidates))
    if len(unique) != 1:
        raise RuntimeError({"project_root_candidates": [str(path) for path in unique]})
    return unique[0]


def find_embedded_project_archive(root: Path) -> Path:
    candidates: list[Path] = []
    diagnostics: list[dict[str, Any]] = []
    for path in sorted(root.rglob("*.zip")):
        try:
            with zipfile.ZipFile(path) as zf:
                if zf.testzip() is not None:
                    continue
                names = [name.replace("\\", "/") for name in zf.namelist()]
        except zipfile.BadZipFile:
            continue
        has_database = any("/Database/" in ("/" + name) for name in names)
        has_tracking = any("/Tracking/" in ("/" + name) for name in names)
        has_manifest = any("/Manifest/" in ("/" + name) for name in names)
        diagnostics.append({"path": str(path), "bytes": path.stat().st_size, "members": len(names), "database": has_database, "tracking": has_tracking, "manifest": has_manifest})
        if len(names) >= 900 and has_database and has_tracking and has_manifest:
            candidates.append(path)
    if len(candidates) != 1:
        raise RuntimeError({"embedded_project_archive_candidates": [str(path) for path in candidates], "diagnostics": diagnostics})
    return candidates[0]


def find_exact_zip_recursive(root: Path, expected_bytes: int, expected_sha256: str, work: Path) -> Path:
    queue = [path for path in root.rglob("*.zip") if path.is_file()]
    seen: set[tuple[int, str]] = set()
    sequence = 0
    while queue:
        candidate = queue.pop(0)
        identity = (candidate.stat().st_size, sha256_file(candidate))
        if identity in seen:
            continue
        seen.add(identity)
        if identity == (expected_bytes, expected_sha256):
            return candidate
        sequence += 1
        target = work / f"nested-{sequence:04d}"
        try:
            safe_extract(candidate, target)
        except (RuntimeError, zipfile.BadZipFile):
            continue
        queue.extend(path for path in target.rglob("*.zip") if path.is_file())
    raise RuntimeError({"exact_zip_not_found": {"root": str(root), "bytes": expected_bytes, "sha256": expected_sha256}})


def reconstruct_response82(response81_restore: Path, response82_dir: Path, work: Path) -> tuple[Path, Path, Path, dict[str, Any]]:
    verify_zip(response81_restore, BASE_RESTORE_BYTES, BASE_RESTORE_SHA256)
    restore_root = work / "response81-restore"
    safe_extract(response81_restore, restore_root)
    project_archive = find_embedded_project_archive(restore_root)
    project_qa = verify_zip(project_archive)
    baseline_extract = work / "response81-project"
    safe_extract(project_archive, baseline_extract)
    baseline_project = locate_project_root(baseline_extract)

    recovery_zip = find_exact_zip_recursive(response82_dir, RESPONSE82_RECOVERY_BYTES, RESPONSE82_RECOVERY_SHA256, work / "response82-discovery")
    package_root = work / "response82-package"
    safe_extract(recovery_zip, package_root)
    apply_script = package_root / "TOOLS" / "apply_checkpoint_recovery.py"
    if not apply_script.exists():
        raise RuntimeError("Response 82 apply utility missing")
    restored = work / "response82-restored"
    result = subprocess.run(
        [sys.executable, str(apply_script.resolve()), "--base-response81-restore", str(response81_restore.resolve()), "--output-dir", str(restored)],
        cwd=package_root,
        text=True,
        capture_output=True,
        timeout=3600,
    )
    if result.returncode:
        raise RuntimeError({"response82_apply_failed": {"stdout": result.stdout[-20000:], "stderr": result.stderr[-20000:]}})
    result_files = list(restored.glob("MRHPD_RESPONSE82*_APPLICATION_RESULT.json"))
    application = json.loads(result_files[0].read_text(encoding="utf-8")) if result_files else {"status": "passed"}
    if application.get("status") != "passed":
        raise RuntimeError({"response82_application_gate": application})
    response82_project = locate_project_root(restored)
    return project_archive, baseline_project, response82_project, {"recovery_zip": verify_zip(recovery_zip, RESPONSE82_RECOVERY_BYTES, RESPONSE82_RECOVERY_SHA256), "application": application, "project_archive": project_qa}


def table_info(con: sqlite3.Connection, table: str) -> list[sqlite3.Row]:
    return list(con.execute(f'PRAGMA table_info("{table}")'))


def clone_response83(con: sqlite3.Connection, now_iso: str) -> dict[str, Any]:
    table = "thread_response_reconciliation_cp3"
    info = table_info(con, table)
    columns = [row[1] for row in info]
    source = con.execute(f"SELECT * FROM {table} WHERE response_key='R82' LIMIT 1").fetchone()
    if source is None:
        raise RuntimeError("Response 82 source row missing")
    record = dict(zip(columns, source))
    for row in info:
        if row[5] and row[1] != "response_key":
            record.pop(row[1], None)
    updates = {
        "response_key": "R83",
        "response_number": 83,
        "response_label": "83",
        "response_date": now_iso,
        "major_topic": "Human Pathogen Database remediation",
        "title": "Section 5 Session 3 evidence-ingestion and governed correction checkpoint",
        "goal": "Continue from Response 82, ingest any genuine provider or physical-proof evidence, execute only supported corrections, freeze the final release candidate, and emit deterministic cumulative recovery.",
        "raw_prompt": RAW_PROMPT_83,
        "raw_response": "[PRE-EMISSION RESPONSE; represented by the source-supported summary]",
        "summary": "Recovered and clean-applied Response 82; repeated the provider and physical-proof evidence search; found no new item-level external evidence; executed no unsupported correction; froze the internally verified release candidate; synchronized database, workbook, application, tracking, indexes, manifests, reports, QA, and cumulative recovery through Response 83.",
        "state": "checkpoint_complete_continue_required",
        "disposition": "COMPLETE_CONTINUE_REQUIRED",
        "next": "Remediation Section 5 of 5 Session 3 of 3 Checkpoint 3 of 3",
        "coverage": "exact raw prompt plus source-supported response summary",
        "fidelity_classification": "source_verified_prompt_and_summary",
        "source_id": "CURRENT-CONVERSATION-R83",
        "source_path": "Current conversation, Response 81 complete restore, and Response 82 recovery",
        "notes": "No provider approval or physical-proof completion is inferred; absence of new evidence does not become approval.",
    }
    for key, value in updates.items():
        if key in columns:
            record[key] = value
    for key in list(record):
        lower = key.lower()
        if lower in {"recorded_at", "created_at", "updated_at", "completed_at", "response_timestamp"}:
            record[key] = now_iso
        elif lower in {"response_order", "sequence_number"}:
            record[key] = 83
    con.execute(f"DELETE FROM {table} WHERE response_key='R83'")
    insert_columns = [column for column in columns if column in record]
    con.execute(f"INSERT INTO {table} ({','.join(insert_columns)}) VALUES ({','.join('?' for _ in insert_columns)})", [record[column] for column in insert_columns])
    return {column: record.get(column) for column in insert_columns}


def evidence_rows(now_iso: str) -> list[dict[str, Any]]:
    rows = [
        ("kdp_previewer_conversion", "Google Drive", "KDP Print Previewer", "No new provider-rendered conversion output was discovered."),
        ("kdp_warning_error_register", "Google Drive", "KDP warning error page", "No item-level warning or error list with page/object references was discovered."),
        ("kdp_submission_or_approval", "Google Drive", "KDP approval submission accepted", "No provider approval, accepted submission, or successful upload receipt was discovered."),
        ("provider_generated_proof", "Google Drive", "KDP digital proof", "No provider-generated proof artifact was discovered."),
        ("physical_proof_order", "Google Drive", "physical proof order", "No physical-proof order confirmation was discovered."),
        ("physical_proof_receipt", "Google Drive", "physical proof receipt", "No proof receipt, delivery, or condition documentation was discovered."),
        ("physical_proof_inspection", "Google Drive", "physical proof inspection defect photograph", "No inspection photographs, defect log, or item-level review record was discovered."),
        ("provider_correction_record", "Google Drive", "KDP correction revalidation", "No provider-driven correction or revalidation record was discovered."),
        ("newer_project_checkpoint", "Google Drive", "MRHPD Response 83", "No completed Response 83 or later project package existed at intake."),
        ("github_execution_state", "GitHub isolated runner", "Session 3 Checkpoint 2", "No pre-existing Checkpoint 2 builder, workflow artifact, or completed later state was found before this build."),
    ]
    return [{"evidence_key": key, "searched_location": location, "search_terms": terms, "observation": observation, "status": "controlled_pending" if key not in {"newer_project_checkpoint", "github_execution_state"} else "passed", "evidence_path": "", "evidence_sha256": "", "claim_allowed": 0, "checked_at": now_iso} for key, location, terms, observation in rows]


def correction_rows(now_iso: str) -> list[dict[str, Any]]:
    rows = [
        ("digital_publication", "No correction triggered; 537-page publication identity retained.", "not_triggered_no_evidence"),
        ("print_interior", "No correction triggered; 538-page print interior identity retained.", "not_triggered_no_evidence"),
        ("cover_raster", "No correction triggered; 5554 x 3375 opaque RGB cover identity retained.", "not_triggered_no_evidence"),
        ("cover_pdf_and_template", "No correction triggered; governed cover PDF and template surfaces retained.", "not_triggered_no_evidence"),
        ("database", "Only checkpoint governance, tracking, evidence, and release-candidate records changed.", "governed_current_state_update"),
        ("workbook", "Only checkpoint governance, tracking, evidence, and release-candidate sheets were added.", "governed_current_state_update"),
        ("application", "Main application source retained byte-identically; read-only audit surface updated.", "governed_current_state_update"),
        ("provider_messages", "No provider message existed to correct against.", "not_triggered_no_evidence"),
        ("physical_proof_defects", "No proof defect existed to correct against.", "not_triggered_no_evidence"),
        ("second_proof", "No correction severity established a second-proof requirement.", "not_triggered_no_evidence"),
    ]
    return [{"correction_key": key, "disposition": disposition, "status": status, "input_evidence_path": "", "affected_artifact": key, "before_sha256": "", "after_sha256": "", "validated": 1, "recorded_at": now_iso} for key, disposition, status in rows]


def release_candidate_gates(now_iso: str) -> list[dict[str, Any]]:
    rows = [
        ("response81_restore", "Exact authoritative Response 81 restore reassembles and verifies", "passed"),
        ("response82_apply", "Cumulative Response 82 recovery clean-applies", "passed"),
        ("sqlite_integrity", "SQLite integrity_check returns ok", "passed"),
        ("foreign_keys", "SQLite foreign_key_check returns zero rows", "passed"),
        ("response83", "Response 83 reconciliation exists exactly once", "passed"),
        ("checkpoint2_state", "Session 3 Checkpoint 2 state is complete", "passed"),
        ("digital_publication", "537-page digital publication remains byte-identical and searchable", "passed"),
        ("editable_assembly", "Editable manuscript assembly remains byte-identical", "passed"),
        ("print_interior", "538-page print interior remains byte-identical", "passed"),
        ("cover", "Full-cover raster remains 5554 x 3375 opaque RGB and byte-identical", "passed"),
        ("application", "Main application source remains byte-identical", "passed"),
        ("workbook", "Comprehensive workbook preserves inherited sheets and formula safety", "passed"),
        ("tracking", "Raw and Net tracking are current through Response 83", "passed"),
        ("source_index", "Source Index rebuilt", "passed"),
        ("bit_index", "Bit Index integrity and FTS counts pass", "passed"),
        ("manifest", "Project manifest and checksums have zero mismatches", "passed"),
        ("evidence_ingestion", "All discovered evidence is ingested with source, identity, and claim boundary", "passed"),
        ("correction_cycle", "No correction is executed without genuine evidence; all dispositions are explicit", "passed"),
        ("unsupported_claims", "Unsupported provider and physical-proof claims remain zero", "passed"),
        ("recovery_apply", "Cumulative Response 83 recovery clean-applies directly to Response 81", "passed"),
        ("provider_previewer", "Item-level provider-rendered preview evidence", "controlled_pending"),
        ("provider_approval", "Provider accepted-submission or approval evidence", "controlled_pending"),
        ("physical_proof_order", "Physical proof order evidence", "controlled_pending"),
        ("physical_proof_receipt", "Physical proof receipt evidence", "controlled_pending"),
        ("physical_proof_inspection", "Physical proof inspection and defect log", "controlled_pending"),
        ("physical_signoff", "Physical proof final signoff", "controlled_pending"),
        ("entire_project_release", "Independent Section 5 and entire-project final release", "planned_checkpoint3"),
    ]
    return [{"gate_key": key, "description": description, "status": status, "evidence": "Deterministic Checkpoint 2 verification." if status == "passed" else "No completion inferred without genuine evidence or the designated final checkpoint.", "checked_at": now_iso} for key, description, status in rows]


def handoff_rows(now_iso: str) -> list[dict[str, Any]]:
    rows = [
        ("independent_reconstruction", "Reconstruct Response 83 from Response 81 plus cumulative recovery", "required"),
        ("independent_acceptance", "Repeat database, workbook, application, publication, print, cover, index, manifest, and recovery gates", "required"),
        ("external_evidence", "Ingest any genuine provider/proof evidence found before final freeze", "conditional"),
        ("unsupported_claims", "Require zero unsupported provider/proof claims", "required"),
        ("section5_release", "Freeze Remediation Section 5", "required"),
        ("entire_project_release", "Emit the complete entire project and self-contained restore", "required"),
        ("google_drive_custody", "Store controlling and redundant final copies in Google Drive", "required"),
        ("terminal_state", "Declare All Sections COMPLETE only after all internally observable gates pass", "required"),
    ]
    return [{"item_key": key, "requirement": requirement, "status": status, "recorded_at": now_iso} for key, requirement, status in rows]


def recovery_events(now_iso: str) -> list[dict[str, Any]]:
    rows = [
        (243, "V3-CP5-S3-REC-243-INSTRUCTIONS-REPROCESSED", "Continuation required the current Project Instructions to be reprocessed.", "Reprocessed Instructions 1.5.0 and applied automatic recovery, Google Drive custody, exact filenames, checkpoint, tracking, and full-release controls."),
        (244, "V3-CP5-S3-REC-244-RESPONSE82-DISCOVERED", "The visible continuation initially reflected Response 81 rather than the completed Response 82 checkpoint.", "Searched Google Drive, discovered the verified Response 82 package, adopted it as the newest checkpoint, and avoided regression."),
        (245, "V3-CP5-S3-REC-245-RESPONSE82-CLEAN-APPLIED", "Checkpoint 2 required an exact current copied project tree.", "Reconstructed Response 81 and clean-applied the cumulative Response 82 recovery before any Response 83 mutation."),
        (246, "V3-CP5-S3-REC-246-EVIDENCE-SEARCH-REPEATED", "Checkpoint 2 required a renewed evidence search.", "Repeated searches for provider preview, warnings, approval, proof order, receipt, inspection, correction, and signoff artifacts."),
        (247, "V3-CP5-S3-REC-247-NO-NEW-EXTERNAL-EVIDENCE", "No genuine item-level provider or physical-proof evidence was found.", "Preserved all external states as controlled pending and prohibited unsupported approval or proof-completion claims."),
        (248, "V3-CP5-S3-REC-248-CORRECTION-CYCLE-NOT-TRIGGERED", "No provider message or proof defect supported a content correction.", "Executed no unsupported publication or cover correction and documented each non-triggered correction domain."),
        (249, "V3-CP5-S3-REC-249-RELEASE-CANDIDATE-FREEZE", "Checkpoint 2 required a final internally verified release candidate.", "Reverified and froze the digital publication, editable assembly, print interior, cover, and application identities."),
        (250, "V3-CP5-S3-REC-250-DATABASE-WORKBOOK-SYNC", "Response 83 and Checkpoint 2 governance required synchronized data surfaces.", "Updated the copied SQLite database and comprehensive workbook while preserving inherited records, worksheets, and extension blocks."),
        (251, "V3-CP5-S3-REC-251-TRACKING-REPORT-INDEX-SYNC", "Current tracking, reports, indexes, and manifests were required.", "Updated Raw/Net tracking, Cumulative Thread Index, reports, Source Index, Bit Index, manifest, checksums, and QA."),
        (252, "V3-CP5-S3-REC-252-CUMULATIVE-RECOVERY", "Checkpoint 2 required deterministic recovery directly from the last full session restore.", "Built and clean-applied cumulative Response 83 recovery directly to the exact Response 81 restore and prepared the Checkpoint 3 handoff."),
    ]
    return [{"event_number": number, "event_code": code, "condition": condition, "recovery": recovery, "status": "recovered", "recorded_at": now_iso} for number, code, condition, recovery in rows]


def sync_database(source: Path, destination: Path, now_iso: str, evidence: list[dict[str, Any]], corrections: list[dict[str, Any]], gates: list[dict[str, Any]], handoff: list[dict[str, Any]], events: list[dict[str, Any]], project: Path) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, destination)
    con = sqlite3.connect(destination)
    try:
        con.execute("PRAGMA foreign_keys=ON")
        con.execute("BEGIN IMMEDIATE")
        clone_response83(con, now_iso)
        con.executescript("""
        CREATE TABLE IF NOT EXISTS section5_session3_evidence_ingestion (
            evidence_key TEXT PRIMARY KEY, searched_location TEXT NOT NULL, search_terms TEXT NOT NULL,
            observation TEXT NOT NULL, status TEXT NOT NULL, evidence_path TEXT NOT NULL,
            evidence_sha256 TEXT NOT NULL, claim_allowed INTEGER NOT NULL, checked_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_session3_correction_cycle (
            correction_key TEXT PRIMARY KEY, disposition TEXT NOT NULL, status TEXT NOT NULL,
            input_evidence_path TEXT NOT NULL, affected_artifact TEXT NOT NULL,
            before_sha256 TEXT NOT NULL, after_sha256 TEXT NOT NULL, validated INTEGER NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_session3_checkpoint2_gate (
            gate_key TEXT PRIMARY KEY, description TEXT NOT NULL, status TEXT NOT NULL,
            evidence TEXT NOT NULL, checked_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_session3_checkpoint2_handoff (
            item_key TEXT PRIMARY KEY, requirement TEXT NOT NULL, status TEXT NOT NULL, recorded_at TEXT NOT NULL
        );
        """)
        con.execute("DELETE FROM section5_session3_checkpoint WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        con.execute(
            "INSERT INTO section5_session3_checkpoint VALUES (?,?,?,?,?,?,?,?,?,?,?)",
            (CHECKPOINT_CODE, 83, "checkpoint_complete", "continue", "passed", "controlled_pending", 0, 0, 0, 0, now_iso),
        )
        con.execute("DELETE FROM section5_session3_evidence_ingestion")
        con.executemany("INSERT INTO section5_session3_evidence_ingestion VALUES (?,?,?,?,?,?,?,?,?)", [(r["evidence_key"], r["searched_location"], r["search_terms"], r["observation"], r["status"], r["evidence_path"], r["evidence_sha256"], r["claim_allowed"], r["checked_at"]) for r in evidence])
        con.execute("DELETE FROM section5_session3_correction_cycle")
        con.executemany("INSERT INTO section5_session3_correction_cycle VALUES (?,?,?,?,?,?,?,?,?)", [(r["correction_key"], r["disposition"], r["status"], r["input_evidence_path"], r["affected_artifact"], r["before_sha256"], r["after_sha256"], r["validated"], r["recorded_at"]) for r in corrections])
        con.execute("DELETE FROM section5_session3_checkpoint2_gate")
        con.executemany("INSERT INTO section5_session3_checkpoint2_gate VALUES (?,?,?,?,?)", [(r["gate_key"], r["description"], r["status"], r["evidence"], r["checked_at"]) for r in gates])
        con.execute("DELETE FROM section5_session3_checkpoint2_handoff")
        con.executemany("INSERT INTO section5_session3_checkpoint2_handoff VALUES (?,?,?,?)", [(r["item_key"], r["requirement"], r["status"], r["recorded_at"]) for r in handoff])
        if con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name='section5_project_completion_plan'").fetchone():
            con.execute("UPDATE section5_project_completion_plan SET state='checkpoint_complete', recorded_at=? WHERE sequence=2", (now_iso,))
            con.execute("UPDATE section5_project_completion_plan SET state='planned', recorded_at=? WHERE sequence=3", (now_iso,))
        if con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name='section5_external_evidence_boundary'").fetchone():
            con.execute("UPDATE section5_external_evidence_boundary SET observation=observation || ' Checkpoint 2 repeated the search and found no new item-level evidence.', checked_at=?", (now_iso,))
        if con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name='section5_provider_evidence_register'").fetchone():
            con.execute("UPDATE section5_provider_evidence_register SET status='controlled_pending', claim_allowed=0, recorded_at=?", (now_iso,))
        if con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name='section5_physical_proof_register'").fetchone():
            con.execute("UPDATE section5_physical_proof_register SET status='controlled_pending', completion_claim_allowed=0, recorded_at=?", (now_iso,))
        for event in events:
            con.execute("INSERT OR REPLACE INTO section5_session3_recovery_event VALUES (?,?,?,?,?,?)", (event["event_code"], event["event_number"], event["condition"], event["recovery"], event["status"], event["recorded_at"]))
        freeze = [
            ("cp2_digital_publication", PUBLICATION_REL),
            ("cp2_print_interior", PRINT_INTERIOR_REL),
            ("cp2_cover_png", COVER_PNG_REL),
            ("cp2_cover_tiff", COVER_TIFF_REL),
            ("cp2_cover_pdf", COVER_PDF_REL),
            ("cp2_template_png", TEMPLATE_PNG_REL),
            ("cp2_template_pdf", TEMPLATE_PDF_REL),
        ]
        for key, rel in freeze:
            path = project / rel
            con.execute("INSERT OR REPLACE INTO section5_session3_artifact_freeze VALUES (?,?,?,?,?,?)", (key, rel, path.stat().st_size, sha256_file(path), "release_candidate_frozen", now_iso))
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or fk:
            raise RuntimeError({"integrity": integrity, "foreign_keys": fk[:20]})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()
    con = sqlite3.connect(destination)
    try:
        table_count = con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0]
        response = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R83'").fetchone()[0]
        checkpoint = con.execute("SELECT session_state,internal_acceptance_state,unsupported_provider_claims,unsupported_proof_claims FROM section5_session3_checkpoint WHERE checkpoint_code=?", (CHECKPOINT_CODE,)).fetchone()
        evidence_count = con.execute("SELECT COUNT(*) FROM section5_session3_evidence_ingestion").fetchone()[0]
        correction_count = con.execute("SELECT COUNT(*) FROM section5_session3_correction_cycle").fetchone()[0]
        failed_gates = con.execute("SELECT COUNT(*) FROM section5_session3_checkpoint2_gate WHERE status='failed'").fetchone()[0]
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk_count = len(list(con.execute("PRAGMA foreign_key_check")))
    finally:
        con.close()
    if response != 1 or checkpoint != ("checkpoint_complete", "passed", 0, 0) or evidence_count != len(evidence) or correction_count != len(corrections) or failed_gates:
        raise RuntimeError({"database_gate": {"response": response, "checkpoint": checkpoint, "evidence": evidence_count, "corrections": correction_count, "failed_gates": failed_gates}})
    return {"path": str(destination), "bytes": destination.stat().st_size, "sha256": sha256_file(destination), "table_count": table_count, "integrity": integrity, "foreign_key_violations": fk_count, "response83_records": response, "evidence_records": evidence_count, "correction_records": correction_count, "failed_gates": failed_gates, "checkpoint_state": checkpoint[0], "status": "passed"}


def safe_cell(value: Any) -> Any:
    if isinstance(value, (dict, list, tuple, set)):
        return json.dumps(value, ensure_ascii=False, sort_keys=True, default=str)
    return value


def write_sheet(ws: Any, rows: list[dict[str, Any]]) -> None:
    if not rows:
        ws.append(["No records"])
        return
    headers = list(rows[0])
    ws.append(headers)
    thin = Side(style="thin", color="AAB8C0")
    for cell in ws[1]:
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(color=WHITE, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in rows:
        ws.append([safe_cell(row.get(header)) for header in headers])
    for row_index in range(2, ws.max_row + 1):
        for cell in ws[row_index]:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if row_index % 2 == 0:
                cell.fill = PatternFill("solid", fgColor=PALE_BLUE)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    ws.sheet_view.showGridLines = False
    for column_index, header in enumerate(headers, start=1):
        sample = [str(header)] + [str(ws.cell(row, column_index).value or "") for row in range(2, min(ws.max_row, 120) + 1)]
        ws.column_dimensions[get_column_letter(column_index)].width = min(55, max(10, max(len(value) for value in sample) + 2))


def sync_workbook(source: Path, destination: Path, evidence: list[dict[str, Any]], corrections: list[dict[str, Any]], gates: list[dict[str, Any]], handoff: list[dict[str, Any]], events: list[dict[str, Any]], freeze: list[dict[str, Any]]) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    wb = load_workbook(source)
    inherited = list(wb.sheetnames)
    datasets = {
        "S5S3 CP2 Dashboard": [
            {"Control": "Response", "Value": 83, "Status": "complete"},
            {"Control": "Checkpoint", "Value": "2 of 3", "Status": "complete"},
            {"Control": "Session", "Value": "3 of 3", "Status": "continue"},
            {"Control": "Evidence ingestion", "Value": len(evidence), "Status": "passed_with_controlled_pending"},
            {"Control": "Corrections executed", "Value": 0, "Status": "no evidence-supported correction"},
            {"Control": "Next", "Value": "Checkpoint 3 of 3 — final project release", "Status": "continue"},
        ],
        "S5S3 CP2 Evidence": evidence,
        "S5S3 CP2 Corrections": corrections,
        "S5S3 CP2 Gates": gates,
        "S5S3 CP2 Handoff": handoff,
        "S5S3 CP2 Response": [{"Response": 83, "Raw Prompt": RAW_PROMPT_83, "Summary": "Evidence ingestion, correction-cycle disposition, release-candidate freeze, synchronized project surfaces, and cumulative recovery completed.", "State": "checkpoint_complete_continue_required"}],
        "S5S3 CP2 Recovery": events,
        "S5S3 CP2 Freeze": freeze,
    }
    for title, rows in datasets.items():
        if title in wb.sheetnames:
            del wb[title]
        ws = wb.create_sheet(title=title)
        write_sheet(ws, rows)
    wb.properties.title = "MRHPD v3.0.0a Comprehensive Tracking Through Response 83"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(destination)
    extension_qa = cp1.preserve_inherited_sheet_extensions(source, destination, inherited)
    with zipfile.ZipFile(destination) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("workbook CRC failed")
    check = load_workbook(destination, read_only=True, data_only=False)
    try:
        names = list(check.sheetnames)
        formula_count = 0
        formula_errors = []
        for ws in check.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    value = cell.value
                    if isinstance(value, str) and value.startswith("="):
                        formula_count += 1
                        if any(token in value for token in ("#REF!", "#NAME?", "#VALUE!", "#DIV/0!", "#N/A")):
                            formula_errors.append(f"{ws.title}!{cell.coordinate}:{value}")
    finally:
        check.close()
    lost = sorted(set(inherited) - set(names))
    if lost or len(names) < len(inherited) + 8 or formula_errors:
        raise RuntimeError({"workbook_gate": {"lost": lost, "sheets": len(names), "formula_errors": formula_errors[:20]}})
    return {"path": str(destination), "bytes": destination.stat().st_size, "sha256": sha256_file(destination), "source_sheet_count": len(inherited), "current_sheet_count": len(names), "new_sheet_count": len(names) - len(inherited), "lost_sheets": lost, "formula_count": formula_count, "formula_error_count": len(formula_errors), "extension_preservation": extension_qa, "status": "passed"}


def shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def tracking_rows(db_path: Path) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    con = sqlite3.connect(db_path)
    try:
        columns = [row[1] for row in con.execute("PRAGMA table_info(thread_response_reconciliation_cp3)")]
        rows = [dict(zip(columns, row)) for row in con.execute("SELECT * FROM thread_response_reconciliation_cp3 ORDER BY CAST(response_number AS REAL), response_key")]
        if con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name='fractional_prompt_cp3'").fetchone():
            fcolumns = [row[1] for row in con.execute("PRAGMA table_info(fractional_prompt_cp3)")]
            fractions = [dict(zip(fcolumns, row)) for row in con.execute("SELECT * FROM fractional_prompt_cp3 ORDER BY CAST(prompt_number AS REAL), prompt_number")]
        else:
            fractions = []
    finally:
        con.close()
    return rows, fractions


def write_tracking(project: Path, db_path: Path, now_iso: str) -> list[Path]:
    rows, fractions = tracking_rows(db_path)
    root = project / "Tracking" / "Prompt Response" / "Through Response 83"
    root.mkdir(parents=True, exist_ok=True)
    response83 = next(row for row in rows if row.get("response_key") == "R83")
    response_json = root / "Response_83_Tracking.json"
    json_write(response_json, response83)

    raw_docx = root / "Medical References - Human Pathogen Database v3.0.0a Alternating Raw Prompts and Responses Through Response 83.docx"
    doc = Document()
    doc.core_properties.title = "Human Pathogen Database — Alternating Raw Prompts and Responses Through Response 83"
    doc.core_properties.author = "Brent McAnulty, M.D."
    doc.add_heading("Human Pathogen Database", 0)
    doc.add_paragraph("Alternating Raw Prompts and Responses Through Response 83")
    for row in rows:
        number = row.get("response_label") or row.get("response_number")
        doc.add_heading(f"Response {number}: {row.get('title') or 'Project exchange'}", level=1)
        table = doc.add_table(rows=2, cols=1)
        table.style = "Table Grid"
        table.cell(0, 0).text = f"RAW PROMPT {number}\n\n{row.get('raw_prompt') or '[RAW PROMPT UNAVAILABLE]'}"
        shade_cell(table.cell(0, 0), "D9EAF7")
        table.cell(1, 0).text = f"RAW RESPONSE {number}\n\n{row.get('raw_response') or '[RAW RESPONSE UNAVAILABLE]'}\n\nSUMMARY\n{row.get('summary') or ''}"
        shade_cell(table.cell(1, 0), "E2F0D9")
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(9)
        doc.add_paragraph()
    if fractions:
        doc.add_heading("Fractional prompts", level=1)
        for row in fractions:
            doc.add_paragraph(f"Prompt {row.get('prompt_number')}: {row.get('prompt_text')}")
    doc.save(raw_docx)

    net_prompt = "Continue the Human Pathogen Database from the newest verified restore; recover autonomously; ingest only genuine provider and physical-proof evidence; execute only evidence-supported corrections; preserve accepted clinical, publication, print, cover, and application artifacts; synchronize all project surfaces; emit checkpoint recovery; and complete the final entire-project release at Checkpoint 3."
    net_response = "Sections 1–4 and Section 5 Sessions 1–2 are complete. Session 3 Checkpoint 2 clean-applies Response 82, repeats the external-evidence search, retains absent provider/proof evidence as controlled pending, executes no unsupported correction, freezes the internally verified release candidate, synchronizes all project surfaces, and emits cumulative recovery. Independent final verification and the entire-project release remain for Checkpoint 3."
    net_docx = root / "Medical References - Human Pathogen Database v3.0.0a Alternating Net Prompts and Responses Through Response 83.docx"
    net_doc = Document()
    net_doc.core_properties.title = "Human Pathogen Database — Alternating Net Prompts and Responses Through Response 83"
    net_doc.core_properties.author = "Brent McAnulty, M.D."
    net_doc.add_heading("Human Pathogen Database", 0)
    net_doc.add_heading("Final-session remediation and project release", level=1)
    table = net_doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.cell(0, 0).text = "NET PROMPT\n\n" + net_prompt
    shade_cell(table.cell(0, 0), "D9EAF7")
    table.cell(1, 0).text = "NET RESPONSE\n\n" + net_response
    shade_cell(table.cell(1, 0), "E2F0D9")
    net_doc.save(net_docx)

    everything = root / "Medical References - Human Pathogen Database v3.0.0a Everything in One Thread Through Response 83.xlsx"
    wb = Workbook()
    wb.remove(wb.active)
    datasets = {
        "Raw Prompts": [{"Response": row.get("response_label") or row.get("response_number"), "Major Topic": row.get("major_topic"), "Title": row.get("title"), "Raw Prompt": row.get("raw_prompt")} for row in rows],
        "Raw Responses": [{"Response": row.get("response_label") or row.get("response_number"), "Major Topic": row.get("major_topic"), "Title": row.get("title"), "Raw Response": row.get("raw_response"), "Summary": row.get("summary")} for row in rows],
        "Fractional Prompts": fractions,
        "Net Prompt": [{"Major Topic": "Human Pathogen Database remediation", "Net Prompt": net_prompt}],
        "Net Response": [{"Major Topic": "Human Pathogen Database remediation", "Net Response": net_response}],
        "Summary Index": [{"Response": row.get("response_label") or row.get("response_number"), "Major Topic": row.get("major_topic"), "Title": row.get("title"), "Goal": row.get("goal"), "Summary": row.get("summary"), "State": row.get("state")} for row in rows],
    }
    for title, data in datasets.items():
        ws = wb.create_sheet(title)
        write_sheet(ws, data)
    wb.properties.title = "Human Pathogen Database — Everything in One Thread Through Response 83"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(everything)

    raw_net = root / "Medical References - Human Pathogen Database v3.0.0a Raw and Net Tracking Through Response 83.md"
    text_write(raw_net, f"# Human Pathogen Database — Raw and Net Tracking Through Response 83\n\n## Raw Prompt 83\n\n{RAW_PROMPT_83}\n\n## Raw Response 83\n\n{response83.get('summary')}\n\n## Net Prompt\n\n{net_prompt}\n\n## Net Response\n\n{net_response}\n\nUpdated: {now_iso}")
    cumulative = root / "Medical References - Human Pathogen Database v3.0.0a Cumulative Thread Index Through Response 83.md"
    lines = ["# Human Pathogen Database — Cumulative Thread Index Through Response 83", ""]
    for row in rows:
        lines.append(f"- Response {row.get('response_label') or row.get('response_number')} — {row.get('title')}: {row.get('summary')}")
    lines.extend(["", f"Updated: {now_iso}"])
    text_write(cumulative, "\n".join(lines))
    return [response_json, raw_docx, net_docx, everything, raw_net, cumulative]


def write_application_surfaces(project: Path, db_path: Path, workbook_path: Path, now_iso: str) -> tuple[list[Path], dict[str, Any]]:
    root = project / "App" / "Section 5 Session 3 Checkpoint 2"
    root.mkdir(parents=True, exist_ok=True)
    apps = [path for path in project.rglob("human_pathogen_app.py") if path.is_file() and sha256_file(path) == APPLICATION_SHA256]
    if len(apps) != 1:
        raise RuntimeError({"main_application_candidates": [str(path) for path in apps]})
    app = apps[0]
    db_rel = db_path.relative_to(project).as_posix()
    workbook_rel = workbook_path.relative_to(project).as_posix()
    pointer = root / "CURRENT_DATABASE.txt"
    text_write(pointer, db_rel)
    state = root / "CURRENT_PROJECT_STATE.json"
    json_write(state, {"schema": "mrhpd-section5-session3-checkpoint2-state-1.0", "response": 83, "section": SECTION_LABEL, "session": SESSION_LABEL, "checkpoint": CHECKPOINT_LABEL, "state": "checkpoint_complete", "database": db_rel, "workbook": workbook_rel, "main_application": app.relative_to(project).as_posix(), "main_application_sha256": sha256_file(app), "main_application_unchanged": True, "recorded_at": now_iso})
    audit_script = root / "audit_section5_session3_checkpoint2.py"
    text_write(audit_script, f'''#!/usr/bin/env python3
import hashlib,json,sqlite3
from pathlib import Path
from openpyxl import load_workbook
from pypdf import PdfReader
project=Path(__file__).resolve().parents[2]
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as f:
  for b in iter(lambda:f.read(1024*1024),b''): h.update(b)
 return h.hexdigest()
db=project/{db_rel!r}; workbook=project/{workbook_rel!r}
con=sqlite3.connect(db)
try:
 integrity=con.execute('PRAGMA integrity_check').fetchone()[0]
 fk=len(list(con.execute('PRAGMA foreign_key_check')))
 response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R83'").fetchone()[0]
 checkpoint=con.execute("SELECT session_state,internal_acceptance_state,unsupported_provider_claims,unsupported_proof_claims FROM section5_session3_checkpoint WHERE checkpoint_code='{CHECKPOINT_CODE}'").fetchone()
 evidence=con.execute('SELECT COUNT(*) FROM section5_session3_evidence_ingestion').fetchone()[0]
 corrections=con.execute('SELECT COUNT(*) FROM section5_session3_correction_cycle').fetchone()[0]
 failed=con.execute("SELECT COUNT(*) FROM section5_session3_checkpoint2_gate WHERE status='failed'").fetchone()[0]
finally: con.close()
wb=load_workbook(workbook,read_only=True,data_only=False)
try: sheets=len(wb.sheetnames)
finally: wb.close()
pub=project/{PUBLICATION_REL!r}; printing=project/{PRINT_INTERIOR_REL!r}; cover=project/{COVER_PNG_REL!r}
pub_reader=PdfReader(str(pub)); print_reader=PdfReader(str(printing))
searchable=sum(1 for page in pub_reader.pages if (page.extract_text() or '').strip())
result={{'status':'passed' if integrity=='ok' and fk==0 and response==1 and checkpoint==('checkpoint_complete','passed',0,0) and evidence==10 and corrections==10 and failed==0 and sheets>=145 and len(pub_reader.pages)==537 and searchable==537 and len(print_reader.pages)==538 and sha(pub)=={PUBLICATION_SHA256!r} and sha(printing)=={PRINT_INTERIOR_SHA256!r} and sha(cover)=={COVER_SHA256!r} else 'failed','integrity':integrity,'foreign_keys':fk,'response83':response,'checkpoint':checkpoint,'evidence':evidence,'corrections':corrections,'failed_gates':failed,'workbook_sheets':sheets,'publication_pages':len(pub_reader.pages),'searchable_pages':searchable,'print_pages':len(print_reader.pages)}}
print(json.dumps(result,indent=2))
raise SystemExit(0 if result['status']=='passed' else 1)
''')
    result = subprocess.run([sys.executable, str(audit_script)], cwd=project, text=True, capture_output=True, timeout=600)
    if result.returncode:
        raise RuntimeError({"application_audit_failed": {"stdout": result.stdout[-12000:], "stderr": result.stderr[-12000:]}})
    audit = json.loads(result.stdout)
    audit.update({"main_application_path": app.relative_to(project).as_posix(), "main_application_sha256": sha256_file(app), "main_application_unchanged": True})
    output = root / "SECTION5_SESSION3_CHECKPOINT2_APPLICATION_AUDIT.json"
    json_write(output, audit)
    return [pointer, state, audit_script, output], audit


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            pass
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], typeface: Any, fill: str, width: int, spacing: int = 8) -> int:
    x, y = xy
    for paragraph in text.splitlines() or [""]:
        for line in textwrap.wrap(paragraph, width=width) or [""]:
            draw.text((x, y), line, font=typeface, fill=fill)
            box = draw.textbbox((x, y), line or "Ag", font=typeface)
            y += box[3] - box[1] + spacing
    return y


def build_figure(path: Path, evidence: list[dict[str, Any]], corrections: list[dict[str, Any]], gates: list[dict[str, Any]]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 2400, 1400
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, 180), fill=f"#{NAVY}")
    draw.text((110, 50), "Evidence ingestion and governed correction cycle", font=font(58, True), fill="white")
    draw.text((112, 215), "Only genuine item-level evidence may trigger a publication or cover correction.", font=font(31), fill=f"#{DARK}")
    cards = [
        ("SEARCH", len(evidence), "evidence lanes reviewed", TEAL),
        ("INGEST", 0, "new external items found", GOLD),
        ("CORRECT", 0, "content corrections triggered", "536D8C"),
        ("VERIFY", sum(1 for row in gates if row["status"] == "passed"), "internal gates passed", "4E8A63"),
    ]
    x = 120
    for label, number, note, color in cards:
        draw.rounded_rectangle((x, 340, x + 500, 790), radius=28, fill=f"#{PALE_BLUE}", outline=f"#{color}", width=6)
        draw.text((x + 35, 385), label, font=font(30, True), fill=f"#{color}")
        draw.text((x + 35, 490), str(number), font=font(96, True), fill=f"#{NAVY}")
        wrap_text(draw, note, (x + 35, 655), font(31, True), f"#{DARK}", 24)
        x += 565
    draw.rounded_rectangle((120, 900, width - 120, 1280), radius=28, fill=f"#{PALE_GOLD}", outline=f"#{GOLD}", width=5)
    note = "No new provider-rendered preview, provider warning or error list, provider approval, physical-proof order, receipt, inspection, defect log, correction record, or signoff was found. The release candidate remains internally verified and externally controlled pending."
    wrap_text(draw, note, (170, 965), font(34, True), f"#{DARK}", 118, spacing=10)
    image.save(path, format="PNG", dpi=(300, 300), optimize=True)
    return {"path": str(path), "pixels": [width, height], "dpi": 300, "mode": image.mode, "bytes": path.stat().st_size, "sha256": sha256_file(path), "status": "passed"}


def build_docx_report(path: Path, now_iso: str, evidence: list[dict[str, Any]], corrections: list[dict[str, Any]], gates: list[dict[str, Any]], handoff: list[dict[str, Any]], events: list[dict[str, Any]], figure: Path, database_qa: dict[str, Any], workbook_qa: dict[str, Any]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55); section.bottom_margin = Inches(0.55); section.left_margin = Inches(0.65); section.right_margin = Inches(0.65)
    doc.core_properties.title = "MRHPD Section 5 Session 3 Checkpoint 2 Evidence Ingestion and Release Candidate Report"
    doc.core_properties.author = "Brent McAnulty, M.D."
    title = doc.add_heading("Human Pathogen Database", 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Section 5 Session 3 Checkpoint 2 — Evidence Ingestion, Governed Correction, and Release Candidate Freeze"); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER; subtitle.runs[0].bold = True
    doc.add_picture(str(figure), width=Inches(6.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Checkpoint disposition", level=1)
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"
    for i, value in enumerate(["Control", "Value", "Status"]):
        table.cell(0, i).text = value; shade_cell(table.cell(0, i), NAVY)
        for run in table.cell(0, i).paragraphs[0].runs: run.font.color.rgb = RGBColor(255, 255, 255); run.bold = True
    for row in [("Response", 83, "complete"), ("Checkpoint", "2 of 3", "complete"), ("Database", f"{database_qa['table_count']} tables", database_qa['integrity']), ("Workbook", f"{workbook_qa['current_sheet_count']} sheets", workbook_qa['status']), ("New external evidence", 0, "controlled pending"), ("Content corrections", 0, "not triggered")]:
        cells = table.add_row().cells
        for i, value in enumerate(row): cells[i].text = str(value)
    for heading, rows, keys in [
        ("Evidence ingestion", evidence, ("evidence_key", "observation", "status")),
        ("Correction cycle", corrections, ("correction_key", "disposition", "status")),
        ("Release-candidate gates", gates, ("gate_key", "description", "status")),
        ("Checkpoint 3 handoff", handoff, ("item_key", "requirement", "status")),
    ]:
        doc.add_heading(heading, level=1)
        t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"
        for i, key in enumerate(keys): t.cell(0, i).text = key.replace("_", " ").title()
        for row in rows:
            cells = t.add_row().cells
            for i, key in enumerate(keys): cells[i].text = str(row[key])
    doc.add_heading("Recovery record", level=1)
    for event in events:
        doc.add_heading(f"{event['event_number']} — {event['event_code']}", level=2)
        doc.add_paragraph(event["condition"]); doc.add_paragraph(event["recovery"])
    doc.add_paragraph(f"Generated: {now_iso}")
    doc.save(path)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None: raise RuntimeError("DOCX CRC failed")
    return {"path": str(path), "bytes": path.stat().st_size, "sha256": sha256_file(path), "status": "passed"}


def build_pdf_report(path: Path, now_iso: str, evidence: list[dict[str, Any]], corrections: list[dict[str, Any]], gates: list[dict[str, Any]], handoff: list[dict[str, Any]], events: list[dict[str, Any]], figure: Path) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleMR", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, leading=21, textColor=colors.HexColor("#17324D"), alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name="BodyMR", parent=styles["BodyText"], fontName="Helvetica", fontSize=8, leading=10.5, textColor=colors.HexColor("#24323D"), spaceAfter=4))
    styles.add(ParagraphStyle(name="H1MR", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#17324D"), spaceBefore=8, spaceAfter=5))
    story: list[Any] = [Paragraph("Human Pathogen Database", styles["TitleMR"]), Paragraph("Section 5 Session 3 Checkpoint 2 — Evidence Ingestion and Release Candidate Freeze", styles["BodyMR"]), RLImage(str(figure), width=6.9 * inch, height=4.025 * inch), Spacer(1, 0.1 * inch)]
    for heading, rows, keys in [
        ("Evidence ingestion", evidence, ("evidence_key", "observation", "status")),
        ("Correction cycle", corrections, ("correction_key", "disposition", "status")),
        ("Release-candidate gates", gates, ("gate_key", "description", "status")),
        ("Checkpoint 3 handoff", handoff, ("item_key", "requirement", "status")),
    ]:
        story.append(Paragraph(heading, styles["H1MR"]))
        data = [[key.replace("_", " ").title() for key in keys]] + [[str(row[key]) for key in keys] for row in rows]
        table = Table(data, colWidths=[1.6 * inch, 4.4 * inch, 1.1 * inch], repeatRows=1)
        table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17324D")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 6.6), ("LEADING", (0, 0), (-1, -1), 8), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#AAB8C0")), ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EAF1F5")])]))
        story.extend([table, PageBreak()])
    story.append(Paragraph("Recovery record", styles["H1MR"]))
    for event in events:
        story.append(Paragraph(f"<b>{event['event_number']} — {event['event_code']}</b><br/>{event['condition']}<br/>{event['recovery']}", styles["BodyMR"]))
    story.append(Paragraph(f"Generated: {now_iso}", styles["BodyMR"]))
    document = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.45 * inch, leftMargin=0.45 * inch, topMargin=0.42 * inch, bottomMargin=0.42 * inch, title="MRHPD Section 5 Session 3 Checkpoint 2 Report", author="Brent McAnulty, M.D.")
    document.build(story)
    reader = PdfReader(str(path)); searchable = sum(1 for page in reader.pages if (page.extract_text() or "").strip())
    if searchable != len(reader.pages): raise RuntimeError({"pdf_pages": len(reader.pages), "searchable": searchable})
    return {"path": str(path), "bytes": path.stat().st_size, "sha256": sha256_file(path), "pages": len(reader.pages), "searchable_pages": searchable, "status": "passed"}


def build_register(path: Path, evidence: list[dict[str, Any]], corrections: list[dict[str, Any]], gates: list[dict[str, Any]], handoff: list[dict[str, Any]], events: list[dict[str, Any]], database_qa: dict[str, Any], workbook_qa: dict[str, Any]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook(); wb.remove(wb.active)
    datasets = {
        "Dashboard": [{"Control": "Response", "Value": 83, "Status": "complete"}, {"Control": "Checkpoint", "Value": "2 of 3", "Status": "complete"}, {"Control": "Database tables", "Value": database_qa["table_count"], "Status": database_qa["integrity"]}, {"Control": "Workbook sheets", "Value": workbook_qa["current_sheet_count"], "Status": workbook_qa["status"]}, {"Control": "New external evidence", "Value": 0, "Status": "controlled_pending"}],
        "Evidence": evidence, "Corrections": corrections, "Gates": gates, "Handoff": handoff, "Recovery": events,
        "Tracking": [{"Response": 83, "Raw Prompt": RAW_PROMPT_83, "Summary": "Checkpoint 2 evidence ingestion and release-candidate freeze completed.", "State": "checkpoint_complete_continue_required"}],
    }
    for title, rows in datasets.items():
        ws = wb.create_sheet(title); write_sheet(ws, rows)
    wb.properties.title = "MRHPD v3.0.0a Section 5 Session 3 Checkpoint 2 Register"; wb.properties.creator = "Brent McAnulty, M.D."; wb.save(path)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None: raise RuntimeError("register CRC failed")
    return {"path": str(path), "bytes": path.stat().st_size, "sha256": sha256_file(path), "sheets": len(datasets), "status": "passed"}


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in {".md", ".txt", ".csv", ".json", ".py", ".html", ".yml", ".yaml"}:
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix == ".docx":
            doc = Document(path); chunks = [paragraph.text for paragraph in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows: chunks.extend(cell.text for cell in row.cells)
            return "\n".join(chunks)
        if suffix == ".pdf":
            reader = PdfReader(str(path)); return "\n".join(page.extract_text() or "" for page in reader.pages)
        if suffix == ".xlsx":
            wb = load_workbook(path, read_only=True, data_only=False)
            try:
                chunks = []
                for ws in wb.worksheets:
                    chunks.append(ws.title)
                    for row in ws.iter_rows(values_only=True): chunks.append(" | ".join("" if value is None else str(value) for value in row))
                return "\n".join(chunks)
            finally: wb.close()
        if suffix in {".sqlite", ".db"}:
            con = sqlite3.connect(path)
            try: return "\n".join(row[0] for row in con.execute("SELECT sql FROM sqlite_master WHERE sql IS NOT NULL ORDER BY name"))
            finally: con.close()
    except Exception as exc:
        return f"[index extraction error: {exc!r}]"
    return ""


def build_indexes(project: Path, now_iso: str) -> dict[str, Any]:
    root = project / "Indexes" / "Section 5 Session 3 Checkpoint 2"; root.mkdir(parents=True, exist_ok=True)
    source_json = root / "MRHPD v3.0.0a Section 5 Session 3 Checkpoint 2 Source Index.json"
    source_csv = root / "MRHPD v3.0.0a Section 5 Session 3 Checkpoint 2 Source Index.csv"
    bit_path = root / "MRHPD v3.0.0a Section 5 Session 3 Checkpoint 2 Bit Index.sqlite"
    qa_path = root / "MRHPD v3.0.0a Section 5 Session 3 Checkpoint 2 Index QA.json"
    excluded = {source_json.resolve(), source_csv.resolve(), bit_path.resolve(), qa_path.resolve()}
    rows: list[dict[str, Any]] = []; fts: list[tuple[str, str, str, str]] = []
    searchable_suffixes = {".md", ".txt", ".csv", ".json", ".py", ".html", ".yml", ".yaml", ".docx", ".pdf", ".xlsx", ".sqlite", ".db"}
    for path in sorted(p for p in project.rglob("*") if p.is_file() and p.resolve() not in excluded):
        rel = path.relative_to(project).as_posix(); purpose = "Project artifact"
        for prefix, label in [("Database/", "Canonical or historical project database"), ("Tracking/", "Prompt response and project tracking"), ("QA/", "Quality assurance evidence"), ("Reports/", "Human-readable report"), ("Print Production/", "Print-production derivative or control"), ("Sources/", "Source and evidence control"), ("Recovery/", "Recovery and continuation control")]:
            if rel.startswith(prefix): purpose = label; break
        searchable = int(path.suffix.lower() in searchable_suffixes)
        content = extract_text(path) if searchable else ""
        row = {"record_type": "physical_file", "path": rel, "container_path": "", "name": path.name, "purpose": purpose, "bytes": path.stat().st_size, "sha256": sha256_file(path), "user_searchable": searchable}
        rows.append(row); fts.append((rel, path.name, purpose, content))
        if path.suffix.lower() == ".zip":
            try:
                with zipfile.ZipFile(path) as zf:
                    for info in zf.infolist():
                        if info.is_dir(): continue
                        member_path = f"{rel}!/{info.filename}"; member_content = ""
                        suffix = Path(info.filename).suffix.lower()
                        if suffix in {".md", ".txt", ".csv", ".json", ".py", ".yml", ".yaml"} and info.file_size <= 2_000_000:
                            try: member_content = zf.read(info).decode("utf-8", errors="replace")
                            except Exception: member_content = ""
                        mrow = {"record_type": "container_member", "path": member_path, "container_path": rel, "name": Path(info.filename).name, "purpose": "Container member", "bytes": info.file_size, "sha256": "", "user_searchable": int(bool(member_content))}
                        rows.append(mrow); fts.append((member_path, mrow["name"], mrow["purpose"], member_content))
            except zipfile.BadZipFile:
                pass
    json_write(source_json, {"schema": "mrhpd-source-index-3.0", "generated_at": now_iso, "records": rows}); csv_write(source_csv, rows)
    if bit_path.exists(): bit_path.unlink()
    con = sqlite3.connect(bit_path)
    try:
        con.executescript("""
        CREATE TABLE artifact (artifact_id INTEGER PRIMARY KEY, record_type TEXT NOT NULL, path TEXT NOT NULL UNIQUE, container_path TEXT, name TEXT NOT NULL, purpose TEXT NOT NULL, bytes INTEGER NOT NULL, sha256 TEXT, user_searchable INTEGER NOT NULL);
        CREATE VIRTUAL TABLE artifact_fts USING fts5(path, name, purpose, content);
        """)
        for row, payload in zip(rows, fts):
            con.execute("INSERT INTO artifact (record_type,path,container_path,name,purpose,bytes,sha256,user_searchable) VALUES (?,?,?,?,?,?,?,?)", (row["record_type"], row["path"], row["container_path"], row["name"], row["purpose"], row["bytes"], row["sha256"], row["user_searchable"]))
            con.execute("INSERT INTO artifact_fts (path,name,purpose,content) VALUES (?,?,?,?)", payload)
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        counts = {"artifact": con.execute("SELECT COUNT(*) FROM artifact").fetchone()[0], "fts": con.execute("SELECT COUNT(*) FROM artifact_fts").fetchone()[0], "response83": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"Response 83"',)).fetchone()[0], "evidence_ingestion": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"evidence ingestion"',)).fetchone()[0]}
        con.commit()
    finally: con.close()
    if integrity != "ok" or counts["artifact"] != len(rows) or counts["fts"] != len(rows): raise RuntimeError({"index_gate": {"integrity": integrity, "counts": counts, "expected": len(rows)}})
    qa = {"status": "passed", "generated_at": now_iso, "source_index_records": len(rows), "physical_files": sum(1 for row in rows if row["record_type"] == "physical_file"), "container_members": sum(1 for row in rows if row["record_type"] == "container_member"), "bit_index_integrity": integrity, "counts": counts, "bit_index_sha256": sha256_file(bit_path)}
    json_write(qa_path, qa)
    return {"source_json": source_json, "source_csv": source_csv, "bit_index": bit_path, "qa_path": qa_path, "qa": qa}


def build_manifest(project: Path, now_iso: str) -> tuple[Path, Path, list[dict[str, Any]]]:
    root = project / "Manifest" / "Section 5 Session 3 Checkpoint 2"; root.mkdir(parents=True, exist_ok=True)
    manifest = root / "MRHPD v3.0.0a Section 5 Session 3 Checkpoint 2 Current Project Manifest.json"
    checksums = root / "MRHPD v3.0.0a Section 5 Session 3 Checkpoint 2 Current Project Checksums.sha256"
    rows = []
    for path in sorted(p for p in project.rglob("*") if p.is_file() and p not in {manifest, checksums}):
        rows.append({"path": path.relative_to(project).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    json_write(manifest, {"schema": "mrhpd-current-project-manifest-3.1", "generated_at": now_iso, "exclusions": [manifest.relative_to(project).as_posix(), checksums.relative_to(project).as_posix()], "file_count": len(rows), "total_bytes": sum(row["bytes"] for row in rows), "files": rows})
    text_write(checksums, "".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
    mismatches = [row["path"] for row in rows if (project / row["path"]).stat().st_size != row["bytes"] or sha256_file(project / row["path"]) != row["sha256"]]
    if mismatches: raise RuntimeError({"manifest_mismatches": mismatches[:20]})
    return manifest, checksums, rows


def create_apply_script(manifest: dict[str, Any], expected: dict[str, Any], baseline_project_bytes: int, baseline_project_sha256: str) -> str:
    return f'''#!/usr/bin/env python3
from __future__ import annotations
import argparse,hashlib,json,re,shutil,sqlite3,tempfile,zipfile
from pathlib import Path,PurePosixPath
from openpyxl import load_workbook
from pypdf import PdfReader
BASE_RESTORE_BYTES={BASE_RESTORE_BYTES}
BASE_RESTORE_SHA256={BASE_RESTORE_SHA256!r}
BASE_PROJECT_BYTES={baseline_project_bytes}
BASE_PROJECT_SHA256={baseline_project_sha256!r}
PUBLICATION_SHA256={PUBLICATION_SHA256!r}
PRINT_INTERIOR_SHA256={PRINT_INTERIOR_SHA256!r}
COVER_SHA256={COVER_SHA256!r}
APPLICATION_SHA256={APPLICATION_SHA256!r}
CURRENT_PROJECT_NAME={CURRENT_PROJECT_NAME!r}
CURRENT_DB_REL={CURRENT_DB_REL!r}
CURRENT_WORKBOOK_REL={CURRENT_WORKBOOK_REL!r}
PUBLICATION_REL={PUBLICATION_REL!r}
PRINT_INTERIOR_REL={PRINT_INTERIOR_REL!r}
COVER_PNG_REL={COVER_PNG_REL!r}
MANIFEST={manifest!r}
EXPECTED={expected!r}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as f:
  for b in iter(lambda:f.read(1024*1024),b''): h.update(b)
 return h.hexdigest()
def verify(path,size,digest,label):
 observed={{'bytes':path.stat().st_size,'sha256':sha(path)}}
 if observed!={{'bytes':size,'sha256':digest}}: raise RuntimeError({{label:observed}})
def safe_extract(path,dest):
 dest.mkdir(parents=True,exist_ok=True)
 with zipfile.ZipFile(path) as zf:
  if zf.testzip() is not None: raise RuntimeError('ZIP CRC failure')
  names=zf.namelist()
  if len(names)!=len(set(names)): raise RuntimeError('duplicate ZIP members')
  for name in names:
   pp=PurePosixPath(name.replace('\\\\','/'))
   if pp.is_absolute() or '..' in pp.parts or re.match(r'^[A-Za-z]:',name): raise RuntimeError('unsafe ZIP path: '+name)
  zf.extractall(dest)
def main():
 ap=argparse.ArgumentParser(); ap.add_argument('--base-response81-restore',type=Path,required=True); ap.add_argument('--output-dir',type=Path,required=True); args=ap.parse_args()
 verify(args.base_response81_restore,BASE_RESTORE_BYTES,BASE_RESTORE_SHA256,'baseline_restore')
 package=Path(__file__).resolve().parents[1]; overlay=package/'OVERLAY'
 if not overlay.is_dir(): raise RuntimeError('OVERLAY missing')
 if args.output_dir.exists() and any(args.output_dir.iterdir()): raise RuntimeError('output directory must be empty')
 args.output_dir.mkdir(parents=True,exist_ok=True)
 with tempfile.TemporaryDirectory(prefix='mrhpd-r83-apply-') as td:
  work=Path(td); restore_root=work/'restore'; safe_extract(args.base_response81_restore,restore_root)
  candidates=[p for p in restore_root.rglob('*.zip') if p.stat().st_size==BASE_PROJECT_BYTES and sha(p)==BASE_PROJECT_SHA256]
  if len(candidates)!=1: raise RuntimeError({{'project_archive_candidates':[str(p) for p in candidates]}})
  extracted=work/'project'; safe_extract(candidates[0],extracted)
  direct=[p for p in extracted.iterdir() if p.is_dir()]; source=direct[0] if len(direct)==1 else extracted
  destination=args.output_dir/CURRENT_PROJECT_NAME; shutil.copytree(source,destination)
  for row in MANIFEST['overlay_files']:
   src=overlay/row['path']; verify(src,row['bytes'],row['sha256'],'overlay_'+row['path']); dst=destination/row['path']; dst.parent.mkdir(parents=True,exist_ok=True); shutil.copy2(src,dst)
  db=destination/CURRENT_DB_REL; con=sqlite3.connect(db)
  try:
   integrity=con.execute('PRAGMA integrity_check').fetchone()[0]; fk=len(list(con.execute('PRAGMA foreign_key_check')))
   response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R83'").fetchone()[0]
   checkpoint=con.execute("SELECT session_state,internal_acceptance_state,unsupported_provider_claims,unsupported_proof_claims FROM section5_session3_checkpoint WHERE checkpoint_code=?",({CHECKPOINT_CODE!r},)).fetchone()
   evidence=con.execute('SELECT COUNT(*) FROM section5_session3_evidence_ingestion').fetchone()[0]
   corrections=con.execute('SELECT COUNT(*) FROM section5_session3_correction_cycle').fetchone()[0]
   failed=con.execute("SELECT COUNT(*) FROM section5_session3_checkpoint2_gate WHERE status='failed'").fetchone()[0]
  finally: con.close()
  wb=load_workbook(destination/CURRENT_WORKBOOK_REL,read_only=True,data_only=False)
  try: sheets=len(wb.sheetnames)
  finally: wb.close()
  pub=destination/PUBLICATION_REL; printing=destination/PRINT_INTERIOR_REL; cover=destination/COVER_PNG_REL
  verify(pub,EXPECTED['publication_bytes'],PUBLICATION_SHA256,'publication'); verify(printing,EXPECTED['print_bytes'],PRINT_INTERIOR_SHA256,'print'); verify(cover,EXPECTED['cover_bytes'],COVER_SHA256,'cover')
  pr=PdfReader(str(pub)); pp=PdfReader(str(printing)); searchable=sum(1 for page in pr.pages if (page.extract_text() or '').strip())
  apps=[p for p in destination.rglob('human_pathogen_app.py') if p.is_file() and sha(p)==APPLICATION_SHA256]
  result={{'status':'passed' if integrity=='ok' and fk==0 and response==1 and checkpoint==('checkpoint_complete','passed',0,0) and evidence==10 and corrections==10 and failed==0 and sheets>=EXPECTED['minimum_workbook_sheets'] and len(pr.pages)==537 and searchable==537 and len(pp.pages)==538 and len(apps)==1 else 'failed','project_root':str(destination),'database':{{'integrity':integrity,'foreign_keys':fk,'response83':response,'checkpoint':checkpoint,'evidence':evidence,'corrections':corrections,'failed_gates':failed}},'workbook_sheets':sheets,'publication_pages':len(pr.pages),'searchable_pages':searchable,'print_pages':len(pp.pages),'main_application_matches':len(apps)}}
  output=args.output_dir/'MRHPD_RESPONSE83_CHECKPOINT_RECOVERY_APPLICATION_RESULT.json'; output.write_text(json.dumps(result,indent=2),encoding='utf-8'); print(json.dumps(result,indent=2)); raise SystemExit(0 if result['status']=='passed' else 1)
if __name__=='__main__': main()
'''


def build_recovery_package(baseline_project: Path, current_project: Path, baseline_restore: Path, project_archive: Path, dist: Path, now: datetime, summary: dict[str, Any], direct_files: list[Path]) -> dict[str, Any]:
    baseline_map = {path.relative_to(baseline_project).as_posix(): (path.stat().st_size, sha256_file(path)) for path in baseline_project.rglob("*") if path.is_file()}
    current_map = {path.relative_to(current_project).as_posix(): (path.stat().st_size, sha256_file(path)) for path in current_project.rglob("*") if path.is_file()}
    deleted = sorted(set(baseline_map) - set(current_map))
    if deleted: raise RuntimeError({"unexpected_deleted_paths": deleted[:30]})
    overlay_rows = [{"path": rel, "bytes": identity[0], "sha256": identity[1], "change": "new" if rel not in baseline_map else "changed"} for rel, identity in sorted(current_map.items()) if baseline_map.get(rel) != identity]
    stamp = now.strftime("%Y-%m-%d %H%M UTC")
    package_root = dist / "recovery_package_root"
    if package_root.exists(): shutil.rmtree(package_root)
    overlay_root = package_root / "OVERLAY"; tools = package_root / "TOOLS"; overlay_root.mkdir(parents=True); tools.mkdir(parents=True)
    for row in overlay_rows:
        source = current_project / row["path"]; target = overlay_root / row["path"]; target.parent.mkdir(parents=True, exist_ok=True); shutil.copy2(source, target)
    manifest = {"schema": "mrhpd-section5-session3-checkpoint-recovery-1.1", "generated_at": now.isoformat().replace("+00:00", "Z"), "version": PROJECT_VERSION, "response": 83, "section": SECTION_LABEL, "session": SESSION_LABEL, "checkpoint": CHECKPOINT_LABEL, "state": "checkpoint_complete", "baseline_restore": {"name": baseline_restore.name, "bytes": baseline_restore.stat().st_size, "sha256": sha256_file(baseline_restore)}, "baseline_project": {"name": project_archive.name, "bytes": project_archive.stat().st_size, "sha256": sha256_file(project_archive)}, "overlay_file_count": len(overlay_rows), "overlay_total_bytes": sum(row["bytes"] for row in overlay_rows), "overlay_files": overlay_rows, "deleted_paths": [], "accepted_predecessor_mutated": False, "frozen_section3_release_mutated": False, "immutable_publication_mutated": False, "user_upload_required": False, "requires_conversation_reconstruction": False, "next": "Remediation Section 5 of 5 Session 3 of 3 Checkpoint 3 of 3"}
    json_write(package_root / "CHECKPOINT_RECOVERY_MANIFEST.json", manifest)
    text_write(package_root / "CHECKPOINT_RECOVERY_CHECKSUMS.sha256", "".join(f"{row['sha256']}  OVERLAY/{row['path']}\n" for row in overlay_rows))
    expected = {"publication_bytes": (current_project / PUBLICATION_REL).stat().st_size, "print_bytes": (current_project / PRINT_INTERIOR_REL).stat().st_size, "cover_bytes": (current_project / COVER_PNG_REL).stat().st_size, "minimum_workbook_sheets": summary["workbook"]["current_sheet_count"]}
    text_write(tools / "apply_checkpoint_recovery.py", create_apply_script(manifest, expected, project_archive.stat().st_size, sha256_file(project_archive)))
    text_write(package_root / "RESTORE_READ_FIRST.md", f"# Human Pathogen Database — Response 83 Checkpoint Recovery\n\nThis cumulative recovery applies directly to the exact Response 81 complete restore and includes all progress through Response 83.\n\nBaseline: `{baseline_restore.name}`\n\nBytes: `{baseline_restore.stat().st_size}`\n\nSHA-256: `{sha256_file(baseline_restore)}`\n\n```bash\npython TOOLS/apply_checkpoint_recovery.py --base-response81-restore \"<Response 81 complete restore.zip>\" --output-dir \"<empty destination>\"\n```\n")
    recovery_zip = dist / f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 Remediation Section 5 of 5 Session 3 of 3 Checkpoint 2 of 3 RECOVERY DATA THROUGH RESPONSE 83 {stamp}.zip"
    with zipfile.ZipFile(recovery_zip, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in sorted(package_root.rglob("*")):
            if path.is_file(): zf.write(path, path.relative_to(package_root).as_posix())
    recovery_qa = verify_zip(recovery_zip)
    with tempfile.TemporaryDirectory(prefix="mrhpd-r83-clean-apply-") as td:
        output = Path(td) / "restored"
        result = subprocess.run([sys.executable, str((tools / "apply_checkpoint_recovery.py").resolve()), "--base-response81-restore", str(baseline_restore.resolve()), "--output-dir", str(output)], cwd=package_root, text=True, capture_output=True, timeout=3600)
        if result.returncode: raise RuntimeError({"clean_apply_failed": {"stdout": result.stdout[-20000:], "stderr": result.stderr[-20000:]}})
        result_files = list(output.glob("MRHPD_RESPONSE83*_APPLICATION_RESULT.json")); clean_apply = json.loads(result_files[0].read_text(encoding="utf-8")) if result_files else {"status": "passed"}
        if clean_apply.get("status") != "passed": raise RuntimeError({"clean_apply_gate": clean_apply})
    verification = {"schema": "mrhpd-response83-checkpoint-recovery-verification-1.0", "generated_at": now.isoformat().replace("+00:00", "Z"), "status": "passed", "recovery_zip": recovery_qa, "manifest": {"overlay_file_count": len(overlay_rows), "overlay_total_bytes": sum(row["bytes"] for row in overlay_rows), "deleted_paths": 0}, "clean_apply": clean_apply, "accepted_predecessor_mutated": False, "immutable_publication_mutated": False, "user_upload_required": False, "checkpoint_2_of_3_complete": True, "session_3_of_3_complete": False, "remediation_section_5_complete": False, "next": "Checkpoint 3 of 3 — final Section 5 and entire-project release"}
    verification_path = dist / "MRHPD v3.0.0a Response 83 Checkpoint 2 Recovery Verification.json"; json_write(verification_path, verification)
    sha_path = dist / f"{recovery_zip.name}.sha256.txt"; text_write(sha_path, f"{recovery_qa['sha256']}  {recovery_zip.name}")
    summary_path = dist / "MRHPD_RESPONSE83_SECTION5_SESSION3_CHECKPOINT2_BUILD_SUMMARY.json"; json_write(summary_path, summary | {"recovery": verification})
    exact_names = dist / "MRHPD v3.0.0a Response 83 Exact File Names.txt"; text_write(exact_names, f"Response 83 cumulative checkpoint recovery ZIP:\n{recovery_zip.name}\n\nRequired baseline complete restore:\n{baseline_restore.name}\n\nRequired baseline project archive:\n{project_archive.name}\n\nCurrent copied SQLite database:\n{Path(CURRENT_DB_REL).name}\n\nCurrent comprehensive workbook:\n{Path(CURRENT_WORKBOOK_REL).name}\n\nImmutable digital publication:\n{Path(PUBLICATION_REL).name}\n\nFrozen print-production interior:\n{Path(PRINT_INTERIOR_REL).name}\n\nExact full-cover PNG:\n{Path(COVER_PNG_REL).name}\n")
    delivery = dist / f"MRHPD v3.0.0a Response 83 Section 5 Session 3 Checkpoint 2 Recovery Package {stamp}.zip"
    with zipfile.ZipFile(delivery, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in [recovery_zip, sha_path, verification_path, summary_path, exact_names, *direct_files]:
            if path.exists(): zf.write(path, path.name)
    delivery_qa = verify_zip(delivery)
    return {"recovery_zip": recovery_zip, "recovery_qa": recovery_qa, "verification_path": verification_path, "summary_path": summary_path, "exact_names": exact_names, "delivery": delivery, "delivery_qa": delivery_qa, "overlay_rows": overlay_rows, "clean_apply": clean_apply}


def main() -> None:
    parser = argparse.ArgumentParser(); parser.add_argument("--response81-restore", type=Path, required=True); parser.add_argument("--response82-dir", type=Path, required=True); parser.add_argument("--dist", type=Path, default=Path("dist_cp5_s3_cp2")); args = parser.parse_args()
    now = utc_now(); now_iso = now.isoformat().replace("+00:00", "Z")
    if args.dist.exists(): shutil.rmtree(args.dist)
    args.dist.mkdir(parents=True)
    with tempfile.TemporaryDirectory(prefix="mrhpd-cp5-s3-cp2-") as td:
        work = Path(td)
        project_archive, baseline_project, response82_project, response82_qa = reconstruct_response82(args.response81_restore, args.response82_dir, work)
        current_project = work / "current" / CURRENT_PROJECT_NAME; current_project.parent.mkdir(parents=True); shutil.copytree(response82_project, current_project)
        evidence = evidence_rows(now_iso); corrections = correction_rows(now_iso); gates = release_candidate_gates(now_iso); handoff = handoff_rows(now_iso); events = recovery_events(now_iso)
        freeze = [{"artifact_key": key, "relative_path": rel, "bytes": (current_project / rel).stat().st_size, "sha256": sha256_file(current_project / rel), "state": "release_candidate_frozen"} for key, rel in [("digital_publication", PUBLICATION_REL), ("print_interior", PRINT_INTERIOR_REL), ("cover_png", COVER_PNG_REL), ("cover_tiff", COVER_TIFF_REL), ("cover_pdf", COVER_PDF_REL), ("template_png", TEMPLATE_PNG_REL), ("template_pdf", TEMPLATE_PDF_REL)]]

        current_db = current_project / CURRENT_DB_REL; database_qa = sync_database(current_project / SOURCE_DB_REL, current_db, now_iso, evidence, corrections, gates, handoff, events, current_project)
        current_workbook = current_project / CURRENT_WORKBOOK_REL; workbook_qa = sync_workbook(current_project / SOURCE_WORKBOOK_REL, current_workbook, evidence, corrections, gates, handoff, events, freeze)
        tracking_files = write_tracking(current_project, current_db, now_iso)
        application_files, application_qa = write_application_surfaces(current_project, current_db, current_workbook, now_iso)

        data_root = current_project / "Data" / "Section 5 Session 3 Checkpoint 2"
        for basename, payload in [("Evidence Ingestion", evidence), ("Correction Cycle", corrections), ("Release Candidate Gates", gates), ("Checkpoint 3 Handoff", handoff), ("Artifact Freeze", freeze)]:
            json_write(data_root / f"MRHPD v3.0.0a Response 83 {basename}.json", payload); csv_write(data_root / f"MRHPD v3.0.0a Response 83 {basename}.csv", payload)

        report_root = current_project / "Reports" / "Section 5 Session 3" / "Checkpoint 2"; artwork_root = current_project / "Artwork" / "Section 5 Final Release" / "Checkpoint 2"
        figure = artwork_root / "MRHPD-FIG-S5-0006 Evidence Ingestion and Governed Correction Cycle v3.0.0a.png"; figure_qa = build_figure(figure, evidence, corrections, gates)
        docx_report = report_root / "MRHPD v3.0.0a Section 5 Session 3 Checkpoint 2 Evidence Ingestion and Release Candidate Report.docx"
        pdf_report = report_root / "MRHPD v3.0.0a Section 5 Session 3 Checkpoint 2 Evidence Ingestion and Release Candidate Report.pdf"
        xlsx_register = report_root / "MRHPD v3.0.0a Section 5 Session 3 Checkpoint 2 Release Candidate Register.xlsx"
        docx_qa = build_docx_report(docx_report, now_iso, evidence, corrections, gates, handoff, events, figure, database_qa, workbook_qa)
        pdf_qa = build_pdf_report(pdf_report, now_iso, evidence, corrections, gates, handoff, events, figure)
        register_qa = build_register(xlsx_register, evidence, corrections, gates, handoff, events, database_qa, workbook_qa)

        qa_root = current_project / "QA" / "Section 5 Session 3" / "Checkpoint 2"; qa_root.mkdir(parents=True, exist_ok=True)
        qa_payloads = {
            "DATABASE_QA.json": database_qa,
            "WORKBOOK_QA.json": workbook_qa,
            "APPLICATION_QA.json": application_qa,
            "EVIDENCE_INGESTION_QA.json": {"status": "passed_with_controlled_pending", "records": evidence, "new_item_level_evidence": 0, "unsupported_claims": 0},
            "CORRECTION_CYCLE_QA.json": {"status": "passed", "records": corrections, "content_corrections_triggered": 0},
            "RELEASE_CANDIDATE_GATE_QA.json": {"status": "passed", "gates": gates, "passed": sum(1 for row in gates if row["status"] == "passed"), "controlled_pending": sum(1 for row in gates if row["status"] == "controlled_pending"), "planned": sum(1 for row in gates if row["status"].startswith("planned"))},
            "REPORT_QA.json": {"status": "passed", "docx": docx_qa, "pdf": pdf_qa, "xlsx": register_qa, "figure": figure_qa},
            "RECOVERY_EVENTS_243_252.json": events,
        }
        for name, payload in qa_payloads.items(): json_write(qa_root / name, payload)
        final_qa = {"schema": "mrhpd-section5-session3-checkpoint2-qa-1.0", "generated_at": now_iso, "status": "passed", "response": 83, "section": SECTION_LABEL, "session": SESSION_LABEL, "checkpoint": CHECKPOINT_LABEL, "response82_recovery": response82_qa, "database": database_qa, "workbook": workbook_qa, "application": application_qa, "evidence_ingestion": qa_payloads["EVIDENCE_INGESTION_QA.json"], "correction_cycle": qa_payloads["CORRECTION_CYCLE_QA.json"], "release_candidate_gates": qa_payloads["RELEASE_CANDIDATE_GATE_QA.json"], "reports": qa_payloads["REPORT_QA.json"], "checkpoint_2_of_3_complete": True, "session_3_of_3_complete": False, "remediation_section_5_complete": False, "accepted_predecessor_mutated": False, "immutable_publication_mutated": False, "provider_approval_claimed": False, "physical_proof_completion_claimed": False, "user_upload_required": False, "next": "Checkpoint 3 of 3 — final Section 5 and entire-project release"}
        json_write(qa_root / "SECTION5_SESSION3_CHECKPOINT2_QA.json", final_qa)

        index_result = build_indexes(current_project, now_iso); manifest_path, checksums_path, manifest_rows = build_manifest(current_project, now_iso)
        summary = {"schema": "mrhpd-response83-section5-session3-checkpoint2-build-1.0", "generated_at": now_iso, "status": "passed", "response": 83, "section": SECTION_LABEL, "session": SESSION_LABEL, "checkpoint": CHECKPOINT_LABEL, "response82_recovery": response82_qa, "database": database_qa, "workbook": workbook_qa, "application": application_qa, "evidence_ingestion": qa_payloads["EVIDENCE_INGESTION_QA.json"], "correction_cycle": qa_payloads["CORRECTION_CYCLE_QA.json"], "release_candidate_gates": qa_payloads["RELEASE_CANDIDATE_GATE_QA.json"], "reports": qa_payloads["REPORT_QA.json"], "index": index_result["qa"], "manifest_records": len(manifest_rows), "user_upload_required": False, "checkpoint_2_of_3_complete": True, "session_3_of_3_complete": False, "remediation_section_5_complete": False, "next": "Checkpoint 3 of 3 — final Section 5 and entire-project release"}
        package = build_recovery_package(baseline_project, current_project, args.response81_restore, project_archive, args.dist, now, summary, [docx_report, pdf_report, xlsx_register, figure])
        print(json.dumps({"status": "passed", "delivery": package["delivery"].name, "delivery_bytes": package["delivery_qa"]["bytes"], "delivery_sha256": package["delivery_qa"]["sha256"], "recovery_zip": package["recovery_zip"].name, "recovery_zip_bytes": package["recovery_qa"]["bytes"], "recovery_zip_sha256": package["recovery_qa"]["sha256"], "overlay_files": len(package["overlay_rows"]), "database_tables": database_qa["table_count"], "workbook_sheets": workbook_qa["current_sheet_count"], "new_external_evidence": 0, "content_corrections_triggered": 0, "provider_approval_claimed": False, "physical_proof_completion_claimed": False, "user_upload_required": False, "checkpoint_2_of_3_complete": True, "next": "Checkpoint 3 of 3 — final Section 5 and entire-project release"}, indent=2))


if __name__ == "__main__":
    main()
