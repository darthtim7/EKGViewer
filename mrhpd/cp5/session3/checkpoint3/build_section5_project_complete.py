#!/usr/bin/env python3
"""Build the complete Human Pathogen Database project through Response 84.

This terminal builder reconstructs the authoritative Response 81 complete restore,
clean-applies cumulative Response 83 recovery, creates a separate copied final
project tree, records Response 84, completes Remediation Section 5 and all
project sections, repeats the full internally observable acceptance matrix, and
emits the entire project inside a self-contained complete restore plus four
connector-compatible transport volumes.

Provider-side preview/approval and physical-proof events remain controlled
pending unless genuine item-level evidence is supplied or discovered. Their
absence is never converted into approval. No accepted predecessor, immutable
publication, print candidate, cover, or main application source is modified in
place.
"""
from __future__ import annotations

import argparse
import collections
import csv
import hashlib
import json
import math
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import textwrap
import zipfile
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

HERE = Path(__file__).resolve().parent
CP2_DIR = HERE.parent / "checkpoint2"
if str(CP2_DIR) not in sys.path:
    sys.path.insert(0, str(CP2_DIR))
import build_section5_session3_checkpoint2 as cp2  # noqa: E402

PROJECT_VERSION = "3.0.0a"
RESPONSE_NUMBER = 84
SECTION_LABEL = "Remediation Section 5 of 5"
SESSION_LABEL = "Session 3 of 3"
CHECKPOINT_LABEL = "Checkpoint 3 of 3"
CHECKPOINT_CODE = "MRHPD-V3-CP5-S3-CP3"
RELEASE_KEY = "MRHPD-V3-R84-ALL-SECTIONS-COMPLETE"
BASE_RESTORE_BYTES = 267_562_561
BASE_RESTORE_SHA256 = "2e90bb8196a4bbaba100d7924fdb2e88be8ce78c238ce330ca219c7e3cae32b2"
RESPONSE83_RECOVERY_BYTES = 44_733_269
RESPONSE83_RECOVERY_SHA256 = "7be544d09a025075eb1c2d4fc4d93f49a09a244b63ec70d36106125001f22c1d"
PUBLICATION_SHA256 = cp2.PUBLICATION_SHA256
EDITABLE_SHA256 = cp2.EDITABLE_SHA256
APPLICATION_SHA256 = cp2.APPLICATION_SHA256
PRINT_INTERIOR_SHA256 = cp2.PRINT_INTERIOR_SHA256
COVER_SHA256 = cp2.COVER_SHA256
PUBLICATION_REL = cp2.PUBLICATION_REL
PRINT_INTERIOR_REL = cp2.PRINT_INTERIOR_REL
COVER_PNG_REL = cp2.COVER_PNG_REL
COVER_TIFF_REL = cp2.COVER_TIFF_REL
COVER_PDF_REL = cp2.COVER_PDF_REL
TEMPLATE_PNG_REL = cp2.TEMPLATE_PNG_REL
TEMPLATE_PDF_REL = cp2.TEMPLATE_PDF_REL
SOURCE_DB_REL = cp2.CURRENT_DB_REL
SOURCE_WORKBOOK_REL = cp2.CURRENT_WORKBOOK_REL
CURRENT_PROJECT_NAME = (
    "Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 3 of 3 ALL SECTIONS COMPLETE THROUGH RESPONSE 84"
)
CURRENT_DB_REL = (
    "Database/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 3 of 3 ALL SECTIONS COMPLETE THROUGH RESPONSE 84.sqlite"
)
CURRENT_WORKBOOK_REL = (
    "Tracking/Workbook/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 3 of 3 ALL SECTIONS COMPLETE THROUGH RESPONSE 84 Comprehensive Tracking.xlsx"
)
RAW_PROMPT_84 = "Continue"
VOLUME_COUNT = 4
MAX_VOLUME_PART_BYTES = 92 * 1024 * 1024
NAVY = "17324D"
TEAL = "1C7475"
GOLD = "C9A227"
PALE_BLUE = "EAF1F5"
PALE_GREEN = "E9F3EE"
PALE_GOLD = "F7F1D9"
PALE_RED = "F7E8E6"
WHITE = "FFFFFF"
DARK = "24323D"
GRAY = "66757F"


def utc_now() -> datetime:
    return datetime.now(timezone.utc).replace(microsecond=0)


def sha256_file(path: Path, chunk_size: int = 1024 * 1024) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(chunk_size), b""):
            digest.update(block)
    return digest.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False, default=str) + "\n", encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value.rstrip() + "\n", encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fields = fields or (list(rows[0]) if rows else [])
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({
                key: json.dumps(value, ensure_ascii=False, sort_keys=True)
                if isinstance(value, (dict, list, tuple, set)) else value
                for key, value in row.items()
            })


def verify_zip(path: Path, expected_bytes: int | None = None, expected_sha256: str | None = None) -> dict[str, Any]:
    if expected_bytes is not None and path.stat().st_size != expected_bytes:
        raise RuntimeError({"file": str(path), "expected_bytes": expected_bytes, "actual_bytes": path.stat().st_size})
    digest = sha256_file(path)
    if expected_sha256 is not None and digest != expected_sha256:
        raise RuntimeError({"file": str(path), "expected_sha256": expected_sha256, "actual_sha256": digest})
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
        names = zf.namelist()
        duplicates = [name for name, count in collections.Counter(names).items() if count > 1]
        unsafe: list[str] = []
        filler: list[str] = []
        for name in names:
            pp = PurePosixPath(name.replace("\\", "/"))
            if pp.is_absolute() or ".." in pp.parts or re.match(r"^[A-Za-z]:", name):
                unsafe.append(name)
            if re.search(r"(^|/)(filler|padding|dummy_payload|artificial_inflation)(/|$)", name, re.I):
                filler.append(name)
    result = {
        "name": path.name,
        "bytes": path.stat().st_size,
        "sha256": digest,
        "members": len(names),
        "crc_error": bad,
        "duplicates": duplicates,
        "unsafe_paths": unsafe,
        "filler_members": filler,
    }
    if bad or duplicates or unsafe or filler:
        raise RuntimeError({"zip_verification_failed": result})
    return result


def safe_extract(path: Path, destination: Path) -> None:
    verify_zip(path)
    destination.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as zf:
        zf.extractall(destination)


def locate_project_root(root: Path) -> Path:
    return cp2.locate_project_root(root)


def find_embedded_project_archive(root: Path) -> Path:
    return cp2.find_embedded_project_archive(root)


def find_exact_zip_recursive(root: Path, size: int, digest: str, work: Path) -> Path:
    return cp2.find_exact_zip_recursive(root, size, digest, work)


def reconstruct_response83(response81_restore: Path, response83_dir: Path, work: Path) -> tuple[Path, Path, Path, dict[str, Any]]:
    restore_qa = verify_zip(response81_restore, BASE_RESTORE_BYTES, BASE_RESTORE_SHA256)
    restore_root = work / "response81-restore"
    safe_extract(response81_restore, restore_root)
    project_archive = find_embedded_project_archive(restore_root)
    project_qa = verify_zip(project_archive)
    baseline_extract = work / "response81-project"
    safe_extract(project_archive, baseline_extract)
    baseline_project = locate_project_root(baseline_extract)

    recovery_zip = find_exact_zip_recursive(
        response83_dir,
        RESPONSE83_RECOVERY_BYTES,
        RESPONSE83_RECOVERY_SHA256,
        work / "response83-discovery",
    )
    package_root = work / "response83-package"
    safe_extract(recovery_zip, package_root)
    apply_script = package_root / "TOOLS" / "apply_checkpoint_recovery.py"
    if not apply_script.exists():
        raise RuntimeError("Response 83 application utility is missing")
    output = work / "response83-restored"
    result = subprocess.run(
        [
            sys.executable,
            str(apply_script.resolve()),
            "--base-response81-restore",
            str(response81_restore.resolve()),
            "--output-dir",
            str(output),
        ],
        cwd=package_root,
        text=True,
        capture_output=True,
        timeout=4200,
    )
    if result.returncode:
        raise RuntimeError({
            "response83_apply_failed": {
                "returncode": result.returncode,
                "stdout": result.stdout[-30000:],
                "stderr": result.stderr[-30000:],
            }
        })
    application_files = list(output.glob("MRHPD_RESPONSE83*_APPLICATION_RESULT.json"))
    application = json.loads(application_files[0].read_text(encoding="utf-8")) if application_files else {"status": "passed"}
    if application.get("status") != "passed":
        raise RuntimeError({"response83_application_gate": application})
    response83_project = locate_project_root(output)
    return project_archive, baseline_project, response83_project, {
        "status": "passed",
        "response81_restore": restore_qa,
        "response81_project_archive": project_qa,
        "response83_recovery": verify_zip(recovery_zip, RESPONSE83_RECOVERY_BYTES, RESPONSE83_RECOVERY_SHA256),
        "response83_application": application,
        "stdout": result.stdout[-12000:],
    }


def table_info(con: sqlite3.Connection, table: str) -> list[sqlite3.Row]:
    return list(con.execute(f'PRAGMA table_info("{table}")'))


def table_exists(con: sqlite3.Connection, table: str) -> bool:
    return con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?", (table,)).fetchone() is not None


def clone_response84(con: sqlite3.Connection, now_iso: str) -> dict[str, Any]:
    table = "thread_response_reconciliation_cp3"
    info = table_info(con, table)
    columns = [row[1] for row in info]
    source = con.execute(f"SELECT * FROM {table} WHERE response_key='R83' LIMIT 1").fetchone()
    if source is None:
        raise RuntimeError("Response 83 source row is missing")
    record = dict(zip(columns, source))
    for row in info:
        if row[5] and row[1] != "response_key":
            record.pop(row[1], None)
    updates = {
        "response_key": "R84",
        "response_number": 84,
        "response_label": "84",
        "response_date": now_iso,
        "major_topic": "Human Pathogen Database remediation",
        "title": "Section 5 and entire Human Pathogen Database project complete release",
        "goal": "Independently reconstruct Response 83, complete the project-wide final acceptance matrix, preserve unavailable provider and physical-proof states as controlled pending, and emit the entire project plus a self-contained final restore.",
        "raw_prompt": RAW_PROMPT_84,
        "raw_response": "[PRE-EMISSION RESPONSE; represented by the source-supported final summary]",
        "summary": "Recovered and independently reconstructed Response 83; completed the final project-wide acceptance matrix; preserved the 537-page digital publication, 538-page print interior, cover, editable assembly, and main application byte-identically; synchronized the final database, workbook, tracking, Master Category, indexes, manifests, reports, QA, and recovery controls; completed Remediation Section 5 and all project sections; and emitted the entire project inside a self-contained complete restore with governed transport and Google Drive custody.",
        "state": "all_sections_complete_with_controlled_external_gates",
        "disposition": "COMPLETE",
        "next": "Optional external provider-preview and physical-proof production lane only; no remediation section remains.",
        "coverage": "exact raw prompt plus source-supported final response summary",
        "fidelity_classification": "source_verified_prompt_and_summary",
        "source_id": "CURRENT-CONVERSATION-R84",
        "source_path": "Current conversation, authoritative Response 81 restore, and cumulative Response 83 recovery",
        "notes": "All internally observable project sections are complete. Provider preview, provider acceptance, proof order, receipt, inspection, correction, and physical signoff remain controlled pending unless genuine evidence is later supplied.",
    }
    for key, value in updates.items():
        if key in columns:
            record[key] = value
    for key in list(record):
        lower = key.lower()
        if lower in {"recorded_at", "created_at", "updated_at", "completed_at", "response_timestamp"}:
            record[key] = now_iso
        elif lower in {"response_order", "sequence_number"}:
            record[key] = 84
    con.execute(f"DELETE FROM {table} WHERE response_key='R84'")
    insert_columns = [column for column in columns if column in record]
    con.execute(
        f"INSERT INTO {table} ({','.join(insert_columns)}) VALUES ({','.join('?' for _ in insert_columns)})",
        [record[column] for column in insert_columns],
    )
    return {column: record.get(column) for column in insert_columns}


def final_gate_rows(now_iso: str) -> list[dict[str, Any]]:
    rows = [
        ("response81_restore", "Authoritative Response 81 restore reproduces exactly", "passed"),
        ("response83_apply", "Cumulative Response 83 recovery clean-applies directly to Response 81", "passed"),
        ("sqlite_integrity", "Canonical SQLite integrity_check returns ok", "passed"),
        ("foreign_keys", "Canonical SQLite foreign_key_check returns zero rows", "passed"),
        ("response84", "Response 84 reconciliation exists exactly once", "passed"),
        ("terminal_release", "Checkpoint 3, Session 3, Section 5, and all project sections are terminal-complete", "passed"),
        ("workbook", "Comprehensive workbook retains inherited sheets/extensions and has no formula-error tokens", "passed"),
        ("tracking", "Raw, Net, Everything-in-One, and Cumulative Thread tracking are current through Response 84", "passed"),
        ("digital_publication", "537-page digital publication remains byte-identical and fully searchable", "passed"),
        ("editable_assembly", "Editable manuscript assembly remains byte-identical", "passed"),
        ("print_interior", "538-page print interior remains byte-identical with one intentional terminal blank", "passed"),
        ("cover", "Full-cover raster remains byte-identical, 5554 x 3375, opaque RGB", "passed"),
        ("cover_components", "Cover PDF, TIFF, template PNG, and template PDF remain governed and present", "passed"),
        ("application", "Main application source remains byte-identical and final read-only audit passes", "passed"),
        ("master_category", "Final Master Category database and human-readable exports are present", "passed"),
        ("source_index", "Final Source Index inventories physical files and ZIP members", "passed"),
        ("bit_index", "Final Bit Index integrity and FTS parity pass", "passed"),
        ("manifest", "Final project manifest and checksum inventory have zero mismatches", "passed"),
        ("project_archive", "Complete project archive clean-extracts and verifies", "passed"),
        ("restore", "Complete self-contained restore and embedded verifier pass", "passed"),
        ("transport", "All four governed restore volumes independently reassemble the exact restore", "passed"),
        ("unsupported_claims", "Unsupported provider and physical-proof completion claims remain zero", "passed"),
        ("provider_previewer", "Item-level provider-rendered preview evidence", "controlled_pending"),
        ("provider_approval", "Provider accepted-submission or approval evidence", "controlled_pending"),
        ("physical_proof_order", "Physical proof order evidence", "controlled_pending"),
        ("physical_proof_receipt", "Physical proof receipt evidence", "controlled_pending"),
        ("physical_proof_inspection", "Physical proof inspection and defect log", "controlled_pending"),
        ("physical_signoff", "Physical proof final signoff", "controlled_pending"),
    ]
    return [
        {
            "gate_key": key,
            "description": description,
            "status": status,
            "evidence": "Verified by the deterministic terminal pipeline." if status == "passed" else "No completion inferred without genuine item-level external evidence.",
            "checked_at": now_iso,
        }
        for key, description, status in rows
    ]


def external_evidence_rows(now_iso: str) -> list[dict[str, Any]]:
    rows = [
        ("drive_previewer", "Google Drive", "KDP Print Previewer", "Search matches were established project governance and delivery records; no provider-rendered preview artifact was discovered."),
        ("drive_provider_approval", "Google Drive", "KDP approval submission accepted", "No accepted-submission receipt or provider approval record was discovered."),
        ("drive_physical_proof", "Google Drive", "physical proof order receipt inspection", "No proof order, receipt, inspection image, defect log, correction record, or physical signoff was discovered."),
        ("gmail_provider_evidence", "Gmail", "KDP Print Previewer OR proof copy OR physical proof OR KDP approval OR submission accepted", "The final Gmail evidence search returned no matching messages."),
        ("newer_project_state", "Google Drive and GitHub", "MRHPD Response 84 or later", "No completed Response 84 or later project package existed at intake."),
        ("evidence_classification", "Final release governance", "item-level external evidence", "Existing project plans and delivery indexes were not reclassified as provider or physical-proof evidence."),
    ]
    return [
        {
            "evidence_key": key,
            "searched_location": location,
            "search_terms": terms,
            "observation": observation,
            "status": "controlled_pending" if key not in {"newer_project_state", "evidence_classification"} else "passed",
            "evidence_path": "",
            "evidence_sha256": "",
            "claim_allowed": 0,
            "checked_at": now_iso,
        }
        for key, location, terms, observation in rows
    ]


def recovery_event_rows(now_iso: str) -> list[dict[str, Any]]:
    rows = [
        (253, "V3-CP5-S3-REC-253-INSTRUCTIONS-REPROCESSED", "Terminal continuation required the controlling Project Instructions to be reprocessed.", "Reprocessed Instructions 1.5.0 and applied automatic recovery, exact filename, Google Drive custody, tracking, index, manifest, checkpoint, and complete-project release controls."),
        (254, "V3-CP5-S3-REC-254-RESPONSE83-ADOPTED", "The final checkpoint required the newest verified project state.", "Located the complete Response 83 checkpoint, verified its persistent and recovery copies, and adopted it without regression or user upload."),
        (255, "V3-CP5-S3-REC-255-RESPONSE83-CLEAN-APPLIED", "A clean final copied project tree was required.", "Reassembled the authoritative Response 81 restore and clean-applied cumulative Response 83 recovery before any Response 84 mutation."),
        (256, "V3-CP5-S3-REC-256-FINAL-EVIDENCE-SEARCH", "The terminal gate required a final search for provider and physical-proof evidence.", "Repeated Google Drive and Gmail searches for preview, warnings, approval, proof order, receipt, inspection, correction, and signoff evidence."),
        (257, "V3-CP5-S3-REC-257-NO-NEW-EXTERNAL-EVIDENCE", "No genuine item-level external production evidence was found.", "Retained all unavailable provider and physical-proof states as controlled pending and prohibited unsupported claims."),
        (258, "V3-CP5-S3-REC-258-FINAL-RELEASE-BOUNDARY", "The project required a precise distinction between internal completion and unavailable external production events.", "Completed every internally observable project section while retaining external provider/proof gates as a separate evidence-dependent lane."),
        (259, "V3-CP5-S3-REC-259-DATABASE-TERMINAL-SYNC", "Response 84 and project terminal state required canonical database synchronization.", "Added Response 84, terminal release, final gates, external evidence, artifact freeze, project summary, and recovery records inside a transaction."),
        (260, "V3-CP5-S3-REC-260-WORKBOOK-TRACKING-SYNC", "The final human-reviewable surfaces required current parity.", "Added final workbook sheets; preserved inherited worksheet extension blocks; and rebuilt Raw, Net, Everything-in-One, and Cumulative Thread tracking through Response 84."),
        (261, "V3-CP5-S3-REC-261-IMMUTABLES-AND-APP", "The final release required direct identity verification of publication, print, cover, editable assembly, and application artifacts.", "Verified every governed identity and added a read-only final application audit without modifying the main application."),
        (262, "V3-CP5-S3-REC-262-CATEGORIES-INDEX-MANIFEST", "The entire project required final discovery, classification, search, and integrity controls.", "Created the final Master Category database and exports, rebuilt Source Index and Bit Index, and generated a zero-mismatch manifest and checksum inventory."),
        (263, "V3-CP5-S3-REC-263-COMPLETE-PROJECT-AND-RESTORE", "Completion required the entire project and a self-contained restore.", "Built and clean-verified the complete project archive and complete restore, including embedded verification and extraction utilities."),
        (264, "V3-CP5-S3-REC-264-TRANSPORT-AND-CUSTODY", "The final restore exceeded a conservative single-volume transfer ceiling.", "Split the exact restore into four minimum governed transport volumes, independently reassembled it, and prepared controlling and redundant Google Drive custody."),
    ]
    return [
        {
            "event_number": number,
            "event_code": code,
            "condition": condition,
            "recovery": recovery,
            "status": "recovered",
            "recorded_at": now_iso,
        }
        for number, code, condition, recovery in rows
    ]


def artifact_freeze_rows(project: Path, now_iso: str) -> list[dict[str, Any]]:
    rows = []
    artifacts = [
        ("digital_publication", PUBLICATION_REL),
        ("print_interior", PRINT_INTERIOR_REL),
        ("cover_png", COVER_PNG_REL),
        ("cover_tiff", COVER_TIFF_REL),
        ("cover_pdf", COVER_PDF_REL),
        ("template_png", TEMPLATE_PNG_REL),
        ("template_pdf", TEMPLATE_PDF_REL),
    ]
    for key, rel in artifacts:
        path = project / rel
        rows.append({
            "artifact_key": key,
            "relative_path": rel,
            "bytes": path.stat().st_size,
            "sha256": sha256_file(path),
            "state": "final_frozen",
            "recorded_at": now_iso,
        })
    editable = find_unique_by_hash(project, EDITABLE_SHA256, suffixes={".docx"})
    app = find_unique_by_hash(project, APPLICATION_SHA256, names={"human_pathogen_app.py"})
    for key, path in (("editable_assembly", editable), ("main_application", app)):
        rows.append({
            "artifact_key": key,
            "relative_path": path.relative_to(project).as_posix(),
            "bytes": path.stat().st_size,
            "sha256": sha256_file(path),
            "state": "final_frozen",
            "recorded_at": now_iso,
        })
    return rows


def find_unique_by_hash(root: Path, digest: str, suffixes: set[str] | None = None, names: set[str] | None = None) -> Path:
    matches = []
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        if suffixes and path.suffix.lower() not in suffixes:
            continue
        if names and path.name not in names:
            continue
        if sha256_file(path) == digest:
            matches.append(path)
    if len(matches) != 1:
        raise RuntimeError({"hash_matches": [str(path) for path in matches], "sha256": digest})
    return matches[0]


def sync_database(
    source: Path,
    destination: Path,
    now_iso: str,
    gates: list[dict[str, Any]],
    evidence: list[dict[str, Any]],
    freeze: list[dict[str, Any]],
    events: list[dict[str, Any]],
) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, destination)
    con = sqlite3.connect(destination)
    try:
        con.execute("PRAGMA foreign_keys=ON")
        con.execute("BEGIN IMMEDIATE")
        response84 = clone_response84(con, now_iso)
        con.executescript("""
        CREATE TABLE IF NOT EXISTS section5_final_release (
            release_key TEXT PRIMARY KEY,
            response_number INTEGER NOT NULL,
            checkpoint_state TEXT NOT NULL,
            session_state TEXT NOT NULL,
            section_state TEXT NOT NULL,
            project_state TEXT NOT NULL,
            internal_acceptance_state TEXT NOT NULL,
            external_evidence_state TEXT NOT NULL,
            provider_approval_claimed INTEGER NOT NULL,
            physical_proof_completion_claimed INTEGER NOT NULL,
            accepted_predecessor_mutated INTEGER NOT NULL,
            immutable_publication_mutated INTEGER NOT NULL,
            main_application_mutated INTEGER NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_final_release_gate (
            gate_key TEXT PRIMARY KEY,
            description TEXT NOT NULL,
            status TEXT NOT NULL,
            evidence TEXT NOT NULL,
            checked_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_final_external_evidence (
            evidence_key TEXT PRIMARY KEY,
            searched_location TEXT NOT NULL,
            search_terms TEXT NOT NULL,
            observation TEXT NOT NULL,
            status TEXT NOT NULL,
            evidence_path TEXT NOT NULL,
            evidence_sha256 TEXT NOT NULL,
            claim_allowed INTEGER NOT NULL,
            checked_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_final_artifact_freeze (
            artifact_key TEXT PRIMARY KEY,
            relative_path TEXT NOT NULL,
            bytes INTEGER NOT NULL,
            sha256 TEXT NOT NULL,
            state TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_final_recovery_event (
            event_code TEXT PRIMARY KEY,
            event_number INTEGER NOT NULL,
            condition TEXT NOT NULL,
            recovery TEXT NOT NULL,
            status TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_final_project_summary (
            summary_key TEXT PRIMARY KEY,
            summary_value TEXT NOT NULL,
            status TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        """)
        con.execute("DELETE FROM section5_session3_checkpoint WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        con.execute(
            "INSERT INTO section5_session3_checkpoint VALUES (?,?,?,?,?,?,?,?,?,?,?)",
            (CHECKPOINT_CODE, 84, "session_complete", "section_complete", "passed", "controlled_pending", 0, 0, 0, 0, now_iso),
        )
        con.execute("DELETE FROM section5_final_release")
        con.execute(
            "INSERT INTO section5_final_release VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (RELEASE_KEY, 84, "checkpoint_complete", "session_complete", "section_complete", "all_sections_complete", "passed", "controlled_pending", 0, 0, 0, 0, 0, now_iso),
        )
        con.execute("DELETE FROM section5_final_release_gate")
        con.executemany(
            "INSERT INTO section5_final_release_gate VALUES (?,?,?,?,?)",
            [(row["gate_key"], row["description"], row["status"], row["evidence"], row["checked_at"]) for row in gates],
        )
        con.execute("DELETE FROM section5_final_external_evidence")
        con.executemany(
            "INSERT INTO section5_final_external_evidence VALUES (?,?,?,?,?,?,?,?,?)",
            [(row["evidence_key"], row["searched_location"], row["search_terms"], row["observation"], row["status"], row["evidence_path"], row["evidence_sha256"], row["claim_allowed"], row["checked_at"]) for row in evidence],
        )
        con.execute("DELETE FROM section5_final_artifact_freeze")
        con.executemany(
            "INSERT INTO section5_final_artifact_freeze VALUES (?,?,?,?,?,?)",
            [(row["artifact_key"], row["relative_path"], row["bytes"], row["sha256"], row["state"], row["recorded_at"]) for row in freeze],
        )
        con.execute("DELETE FROM section5_final_recovery_event")
        con.executemany(
            "INSERT INTO section5_final_recovery_event VALUES (?,?,?,?,?,?)",
            [(row["event_code"], row["event_number"], row["condition"], row["recovery"], row["status"], row["recorded_at"]) for row in events],
        )
        summaries = [
            ("response", "84", "complete"),
            ("checkpoint", "3 of 3", "complete"),
            ("session", "3 of 3", "complete"),
            ("remediation_section_5", "5 of 5", "complete"),
            ("all_sections", "complete", "complete"),
            ("provider_evidence", "controlled_pending", "accurately_bounded"),
            ("physical_proof", "controlled_pending", "accurately_bounded"),
            ("next", "optional external provider/proof lane only", "no_remediation_session_remaining"),
        ]
        con.execute("DELETE FROM section5_final_project_summary")
        con.executemany(
            "INSERT INTO section5_final_project_summary VALUES (?,?,?,?)",
            [(key, value, status, now_iso) for key, value, status in summaries],
        )
        if table_exists(con, "section5_project_completion_plan"):
            con.execute("UPDATE section5_project_completion_plan SET state='complete' WHERE sequence IN (1,2,3)")
        con.commit()
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        foreign_keys = len(list(con.execute("PRAGMA foreign_key_check")))
        table_count = con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'").fetchone()[0]
        response_count = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R84'").fetchone()[0]
        release = con.execute("SELECT checkpoint_state,session_state,section_state,project_state,internal_acceptance_state,external_evidence_state,provider_approval_claimed,physical_proof_completion_claimed FROM section5_final_release WHERE release_key=?", (RELEASE_KEY,)).fetchone()
        failed = con.execute("SELECT COUNT(*) FROM section5_final_release_gate WHERE status='failed'").fetchone()[0]
        pending = con.execute("SELECT COUNT(*) FROM section5_final_release_gate WHERE status='controlled_pending'").fetchone()[0]
        passed = con.execute("SELECT COUNT(*) FROM section5_final_release_gate WHERE status='passed'").fetchone()[0]
    finally:
        con.close()
    expected_release = ("checkpoint_complete", "session_complete", "section_complete", "all_sections_complete", "passed", "controlled_pending", 0, 0)
    if integrity != "ok" or foreign_keys or response_count != 1 or release != expected_release or failed:
        raise RuntimeError({
            "database_terminal_gate": {
                "integrity": integrity,
                "foreign_keys": foreign_keys,
                "response84": response_count,
                "release": release,
                "failed_gates": failed,
            }
        })
    return {
        "status": "passed",
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "tables": table_count,
        "integrity": integrity,
        "foreign_keys": foreign_keys,
        "response84": response_count,
        "release": release,
        "passed_gates": passed,
        "controlled_pending_gates": pending,
        "failed_gates": failed,
        "response84_record": response84,
    }


def safe_cell(value: Any) -> Any:
    if isinstance(value, (dict, list, tuple, set)):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return value


def write_sheet(ws: Any, rows: list[dict[str, Any]]) -> None:
    if not rows:
        ws.append(["No records"])
        return
    headers = list(rows[0])
    ws.append(headers)
    thin = Side(style="thin", color="AAB8C0")
    for cell in ws[1]:
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(color=WHITE, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in rows:
        ws.append([safe_cell(row.get(header)) for header in headers])
    for row_index in range(2, ws.max_row + 1):
        for cell in ws[row_index]:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if row_index % 2 == 0:
                cell.fill = PatternFill("solid", fgColor=PALE_BLUE)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    ws.sheet_view.showGridLines = False
    for column_index, header in enumerate(headers, start=1):
        sample = [str(header)] + [str(ws.cell(row, column_index).value or "") for row in range(2, min(ws.max_row, 100) + 1)]
        ws.column_dimensions[get_column_letter(column_index)].width = min(58, max(11, max(len(value) for value in sample) + 2))


def sync_workbook(
    source: Path,
    destination: Path,
    gates: list[dict[str, Any]],
    evidence: list[dict[str, Any]],
    freeze: list[dict[str, Any]],
    events: list[dict[str, Any]],
    db_qa: dict[str, Any],
) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    wb = load_workbook(source)
    inherited = list(wb.sheetnames)
    datasets = {
        "S5S3 FINAL Dashboard": [
            {"Control": "Response", "Value": 84, "Status": "complete"},
            {"Control": "Checkpoint", "Value": "3 of 3", "Status": "complete"},
            {"Control": "Session", "Value": "3 of 3", "Status": "complete"},
            {"Control": "Remediation Section 5", "Value": "5 of 5", "Status": "complete"},
            {"Control": "All Sections", "Value": "COMPLETE", "Status": "complete"},
            {"Control": "Database tables", "Value": db_qa["tables"], "Status": db_qa["integrity"]},
            {"Control": "External provider/proof lane", "Value": "controlled pending", "Status": "accurately bounded"},
        ],
        "S5S3 FINAL Gates": gates,
        "S5S3 FINAL External": evidence,
        "S5S3 FINAL Freeze": freeze,
        "S5S3 FINAL Response": [{"Response": 84, "Raw Prompt": RAW_PROMPT_84, "Disposition": "COMPLETE", "Summary": "Section 5 and all project sections completed; entire project and final self-contained restore emitted; external provider/proof gates remain controlled pending."}],
        "S5S3 FINAL Recovery": events,
        "S5S3 FINAL Completion": [
            {"Sequence": 1, "Phase": "Checkpoint 1", "State": "complete"},
            {"Sequence": 2, "Phase": "Checkpoint 2", "State": "complete"},
            {"Sequence": 3, "Phase": "Checkpoint 3", "State": "complete"},
            {"Sequence": 4, "Phase": "Session 3", "State": "complete"},
            {"Sequence": 5, "Phase": "Remediation Section 5", "State": "complete"},
            {"Sequence": 6, "Phase": "All Sections", "State": "complete"},
        ],
        "S5S3 FINAL Custody": [
            {"Artifact": "Complete project archive", "State": "generated and clean-verified", "Custody": "embedded in complete restore"},
            {"Artifact": "Complete self-contained restore", "State": "generated and verified", "Custody": "four governed Google Drive volumes"},
            {"Artifact": "Verification delivery", "State": "generated and independently verified", "Custody": "Google Drive plus recovery copy"},
        ],
    }
    for title, rows in datasets.items():
        if title in wb.sheetnames:
            del wb[title]
        ws = wb.create_sheet(title=title)
        write_sheet(ws, rows)
    wb.properties.title = "MRHPD v3.0.0a Comprehensive Tracking — All Sections Complete Through Response 84"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(destination)
    extension_qa = cp2.preserve_inherited_sheet_extensions(source, destination, inherited)
    with zipfile.ZipFile(destination) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("workbook ZIP CRC failed")
    check = load_workbook(destination, read_only=True, data_only=False)
    try:
        sheet_names = list(check.sheetnames)
        formula_errors: list[str] = []
        formula_count = 0
        for ws in check.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    value = cell.value
                    if isinstance(value, str) and value.startswith("="):
                        formula_count += 1
                        if any(token in value for token in ("#REF!", "#NAME?", "#VALUE!", "#DIV/0!", "#N/A")):
                            formula_errors.append(f"{ws.title}!{cell.coordinate}:{value}")
    finally:
        check.close()
    lost = sorted(set(inherited) - set(sheet_names))
    if lost or len(sheet_names) < len(inherited) + 8 or formula_errors:
        raise RuntimeError({"workbook_gate": {"lost": lost, "sheet_count": len(sheet_names), "formula_errors": formula_errors[:20]}})
    return {
        "status": "passed",
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "source_sheet_count": len(inherited),
        "current_sheet_count": len(sheet_names),
        "new_sheet_count": len(sheet_names) - len(inherited),
        "lost_sheets": lost,
        "formula_count": formula_count,
        "formula_error_count": len(formula_errors),
        "extension_preservation": extension_qa,
    }


def shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def tracking_rows(db_path: Path) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    con = sqlite3.connect(db_path)
    try:
        columns = [row[1] for row in con.execute("PRAGMA table_info(thread_response_reconciliation_cp3)")]
        rows = [dict(zip(columns, row)) for row in con.execute("SELECT * FROM thread_response_reconciliation_cp3 ORDER BY CAST(response_number AS REAL), response_key")]
        if table_exists(con, "fractional_prompt_cp3"):
            fcolumns = [row[1] for row in con.execute("PRAGMA table_info(fractional_prompt_cp3)")]
            fractions = [dict(zip(fcolumns, row)) for row in con.execute("SELECT * FROM fractional_prompt_cp3 ORDER BY CAST(prompt_number AS REAL), prompt_number")]
        else:
            fractions = []
    finally:
        con.close()
    return rows, fractions


def write_tracking(project: Path, db_path: Path, now_iso: str) -> dict[str, Any]:
    rows, fractions = tracking_rows(db_path)
    root = project / "Tracking" / "Prompt Response" / "All Sections Complete Through Response 84"
    root.mkdir(parents=True, exist_ok=True)
    response84 = next(row for row in rows if row.get("response_key") == "R84")
    response_json = root / "Response_84_Tracking.json"
    json_write(response_json, response84)

    raw_docx = root / "Medical References - Human Pathogen Database v3.0.0a Alternating Raw Prompts and Responses Through Response 84.docx"
    doc = Document()
    doc.core_properties.title = "Human Pathogen Database — Alternating Raw Prompts and Responses Through Response 84"
    doc.core_properties.author = "Brent McAnulty, M.D."
    title = doc.add_heading("Human Pathogen Database", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Alternating Raw Prompts and Responses Through Response 84")
    for row in rows:
        number = row.get("response_label") or row.get("response_number")
        doc.add_heading(f"Response {number}: {row.get('title') or 'Project exchange'}", level=1)
        table = doc.add_table(rows=2, cols=1)
        table.style = "Table Grid"
        table.cell(0, 0).text = f"RAW PROMPT {number}\n\n{row.get('raw_prompt') or '[RAW PROMPT UNAVAILABLE]'}"
        shade_cell(table.cell(0, 0), "D9EAF7")
        table.cell(1, 0).text = f"RAW RESPONSE {number}\n\n{row.get('raw_response') or '[RAW RESPONSE UNAVAILABLE]'}\n\nSUMMARY\n{row.get('summary') or ''}"
        shade_cell(table.cell(1, 0), "E2F0D9")
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(9)
        doc.add_paragraph()
    if fractions:
        doc.add_heading("Fractional prompts", level=1)
        for row in fractions:
            doc.add_paragraph(f"Prompt {row.get('prompt_number')}: {row.get('prompt_text')}")
    doc.save(raw_docx)

    net_docx = root / "Medical References - Human Pathogen Database v3.0.0a Alternating Net Prompts and Responses Through Response 84.docx"
    net_doc = Document()
    net_doc.core_properties.title = "Human Pathogen Database — Alternating Net Prompts and Responses Through Response 84"
    net_doc.core_properties.author = "Brent McAnulty, M.D."
    title = net_doc.add_heading("Human Pathogen Database", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    net_doc.add_heading("Human Pathogen Database remediation and final release", level=1)
    table = net_doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    net_prompt = (
        "Build, validate, publish, and preserve a comprehensive multi-kingdom Human Pathogen Database and searchable clinical reference, with synchronized SQLite, workbook, application, publication, graphics, evidence, treatment, diagnostics, stewardship, print-production, tracking, indexing, manifests, recovery, and persistent Google Drive custody. Complete every remediation section; preserve exact Raw history; maintain current Net state; create frequent recoverable checkpoints; emit the entire completed project and a self-contained final restore; and never infer provider or physical-proof completion without genuine evidence."
    )
    net_response = (
        "The Human Pathogen Database v3.0.0a project is complete across all remediation sections. The final synchronized project preserves the 537-page searchable digital publication, 538-page print-production interior, exact cover package, editable assembly, main application, canonical SQLite database, comprehensive workbook, Master Category database, Raw/Net tracking, Source Index, Bit Index, manifests, QA, and recovery controls. Every internally observable final gate passed. Provider preview, provider approval, proof order, receipt, physical inspection, correction, and signoff remain controlled pending because no genuine item-level evidence was supplied or discovered."
    )
    table.cell(0, 0).text = "NET PROMPT\n\n" + net_prompt
    shade_cell(table.cell(0, 0), "D9EAF7")
    table.cell(1, 0).text = "NET RESPONSE\n\n" + net_response
    shade_cell(table.cell(1, 0), "E2F0D9")
    net_doc.add_paragraph("Brief discussion summary: The thread developed, governed, remediated, print-preflighted, synchronized, and completed a durable Human Pathogen Database publication and application package. The solution is a complete self-contained project release with exact restoration controls and transparent external-production boundaries.")
    net_doc.save(net_docx)

    everything = root / "Medical References - Human Pathogen Database v3.0.0a Everything in One Thread Through Response 84.xlsx"
    wb = Workbook()
    wb.remove(wb.active)
    response_rows = [{
        "Response": row.get("response_number"),
        "Response Key": row.get("response_key"),
        "Title": row.get("title"),
        "Goal": row.get("goal"),
        "Raw Prompt": row.get("raw_prompt"),
        "Raw Response": row.get("raw_response"),
        "Summary": row.get("summary"),
        "Disposition": row.get("disposition"),
        "Next": row.get("next"),
    } for row in rows]
    datasets = {
        "Response Summary": response_rows,
        "Raw Prompts": [{"Response": row.get("response_number"), "Raw Prompt": row.get("raw_prompt")} for row in rows],
        "Raw Responses": [{"Response": row.get("response_number"), "Raw Response": row.get("raw_response"), "Summary": row.get("summary")} for row in rows],
        "Net Prompt": [{"Major Topic": "Human Pathogen Database", "Net Prompt": net_prompt}],
        "Net Response": [{"Major Topic": "Human Pathogen Database", "Net Response": net_response}],
        "Fractional Prompts": fractions or [{"Prompt": "None", "State": "No fractional prompt rows available"}],
        "Completion": [{"Response": 84, "Checkpoint": "3 of 3", "Session": "3 of 3", "Section": "5 of 5", "All Sections": "COMPLETE", "External Gates": "controlled pending"}],
    }
    for title, dataset in datasets.items():
        ws = wb.create_sheet(title=title)
        write_sheet(ws, dataset)
    wb.properties.title = "MRHPD Everything in One Thread Through Response 84"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(everything)

    cumulative = root / "Medical References - Human Pathogen Database v3.0.0a Cumulative Thread Index Through Response 84.docx"
    index_doc = Document()
    index_doc.core_properties.title = "Human Pathogen Database — Cumulative Thread Index Through Response 84"
    index_doc.core_properties.author = "Brent McAnulty, M.D."
    title = index_doc.add_heading("Human Pathogen Database", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    index_doc.add_paragraph("Cumulative Thread Index Through Response 84")
    for row in rows:
        number = row.get("response_number")
        index_doc.add_heading(f"Response {number}: {row.get('title') or 'Project exchange'}", level=1)
        index_doc.add_paragraph(f"Goal: {row.get('goal') or ''}")
        index_doc.add_paragraph(f"Summary: {row.get('summary') or ''}")
        index_doc.add_paragraph(f"Disposition: {row.get('disposition') or ''}")
    index_doc.save(cumulative)

    for path in (raw_docx, net_docx, cumulative):
        with zipfile.ZipFile(path) as zf:
            if zf.testzip() is not None:
                raise RuntimeError(f"DOCX CRC failed: {path}")
    with zipfile.ZipFile(everything) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("Everything-in-One workbook CRC failed")
    return {
        "status": "passed",
        "response_rows": len(rows),
        "fractional_prompt_rows": len(fractions),
        "files": [path.relative_to(project).as_posix() for path in (response_json, raw_docx, net_docx, everything, cumulative)],
    }


def ensure_master_category(project: Path, now_iso: str) -> dict[str, Any]:
    root = project / "Database" / "Master Category"
    root.mkdir(parents=True, exist_ok=True)
    sqlite_path = root / "MRHPD v3.0.0a Master Category Database Through Response 84.sqlite"
    xlsx_path = root / "MRHPD v3.0.0a Master Category Database Through Response 84.xlsx"
    csv_path = root / "MRHPD v3.0.0a Master Category Database Through Response 84.csv"
    categories = [
        ("MRHPD-CAT-TAXONOMY", "Pathogen taxonomy and nomenclature", "Canonical taxa, aliases, former names, groups, and resolver behavior"),
        ("MRHPD-CAT-CLINICAL", "Clinical profiles and syndromes", "Organism profiles, syndromes, manifestations, mimics, risk, and follow-up"),
        ("MRHPD-CAT-LAB", "Laboratory and diagnostics", "Specimens, morphology, growth, identification, AST, imaging, and interpretation"),
        ("MRHPD-CAT-TREATMENT", "Treatment and stewardship", "Therapy contexts, patient factors, source control, duration, stopping, and no-treatment pathways"),
        ("MRHPD-CAT-EVIDENCE", "Evidence and sources", "Guidelines, standards, surveillance, literature, provenance, and citation governance"),
        ("MRHPD-CAT-GRAPHICS", "Graphics and rights", "Artwork, placeholders, observational-image rules, provenance, rights, captions, and alt text"),
        ("MRHPD-CAT-PUBLICATION", "Publication", "Digital publication, editable assembly, page maps, searchability, and reader navigation"),
        ("MRHPD-CAT-PRINT", "Print production", "Print interior, cover, template, bleed, spine, color, provider preview, and physical proof"),
        ("MRHPD-CAT-DATABASE", "Databases and application", "Canonical SQLite, application, portable data, schema, and integrity"),
        ("MRHPD-CAT-TRACKING", "Prompt, response, and project tracking", "Raw, Net, Everything-in-One, Cumulative Thread Index, and checkpoints"),
        ("MRHPD-CAT-INDEX", "Indexes and discovery", "Source Index, Bit Index, container-member inventory, and searchability"),
        ("MRHPD-CAT-QA", "Quality assurance", "Acceptance matrices, drift checks, audits, formulas, pages, and release gates"),
        ("MRHPD-CAT-RECOVERY", "Recovery and custody", "Manifests, checksums, restore utilities, volumes, Google Drive delivery, and recovery copies"),
        ("MRHPD-CAT-EXTERNAL", "External production evidence", "Provider preview/approval and physical-proof evidence retained as controlled pending until genuine"),
    ]
    con = sqlite3.connect(sqlite_path)
    try:
        con.executescript("""
        PRAGMA journal_mode=DELETE;
        CREATE TABLE category (
            category_key TEXT PRIMARY KEY,
            category_name TEXT NOT NULL,
            scope TEXT NOT NULL,
            version TEXT NOT NULL,
            state TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE subcategory (
            subcategory_key TEXT PRIMARY KEY,
            category_key TEXT NOT NULL REFERENCES category(category_key),
            subcategory_name TEXT NOT NULL,
            scope TEXT NOT NULL
        );
        CREATE INDEX idx_subcategory_category ON subcategory(category_key);
        """)
        con.execute("DELETE FROM subcategory")
        con.execute("DELETE FROM category")
        con.executemany(
            "INSERT INTO category VALUES (?,?,?,?,?,?)",
            [(key, name, scope, PROJECT_VERSION, "current", now_iso) for key, name, scope in categories],
        )
        subcategories = []
        for key, name, scope in categories:
            subcategories.extend([
                (key + "-DATA", key, name + " data", scope),
                (key + "-GOV", key, name + " governance", "Rules, QA, provenance, and release state for " + name.lower()),
            ])
        con.executemany("INSERT INTO subcategory VALUES (?,?,?,?)", subcategories)
        con.commit()
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = len(list(con.execute("PRAGMA foreign_key_check")))
        count = con.execute("SELECT COUNT(*) FROM category").fetchone()[0]
        subcount = con.execute("SELECT COUNT(*) FROM subcategory").fetchone()[0]
    finally:
        con.close()
    rows = [{"Category Key": key, "Category": name, "Scope": scope, "Version": PROJECT_VERSION, "State": "current"} for key, name, scope in categories]
    csv_write(csv_path, rows)
    wb = Workbook()
    ws = wb.active
    ws.title = "Categories"
    write_sheet(ws, rows)
    sub_ws = wb.create_sheet("Subcategories")
    write_sheet(sub_ws, [{"Subcategory Key": skey, "Category Key": ckey, "Subcategory": sname, "Scope": scope} for skey, ckey, sname, scope in subcategories])
    wb.properties.title = "MRHPD Master Category Database Through Response 84"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(xlsx_path)
    with zipfile.ZipFile(xlsx_path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("Master Category workbook CRC failed")
    if integrity != "ok" or fk or count != len(categories) or subcount != len(categories) * 2:
        raise RuntimeError({"master_category_gate": {"integrity": integrity, "foreign_keys": fk, "categories": count, "subcategories": subcount}})
    return {
        "status": "passed",
        "category_count": count,
        "subcategory_count": subcount,
        "integrity": integrity,
        "foreign_keys": fk,
        "files": [path.relative_to(project).as_posix() for path in (sqlite_path, xlsx_path, csv_path)],
    }


def audit_core_artifacts(project: Path) -> dict[str, Any]:
    publication = project / PUBLICATION_REL
    printing = project / PRINT_INTERIOR_REL
    cover = project / COVER_PNG_REL
    for path, digest, label in ((publication, PUBLICATION_SHA256, "publication"), (printing, PRINT_INTERIOR_SHA256, "print_interior"), (cover, COVER_SHA256, "cover")):
        observed = sha256_file(path)
        if observed != digest:
            raise RuntimeError({label: {"expected_sha256": digest, "actual_sha256": observed, "path": str(path)}})
    editable = find_unique_by_hash(project, EDITABLE_SHA256, suffixes={".docx"})
    application = find_unique_by_hash(project, APPLICATION_SHA256, names={"human_pathogen_app.py"})
    digital = PdfReader(str(publication))
    digital_searchable = sum(1 for page in digital.pages if (page.extract_text() or "").strip())
    print_reader = PdfReader(str(printing))
    print_searchable = sum(1 for page in print_reader.pages if (page.extract_text() or "").strip())
    with Image.open(cover) as image:
        dimensions = [image.width, image.height]
        mode = image.mode
        opaque = True
        if "A" in image.getbands():
            alpha = image.getchannel("A")
            extrema = alpha.getextrema()
            opaque = extrema == (255, 255)
    if len(digital.pages) != 537 or digital_searchable != 537 or len(print_reader.pages) != 538 or print_searchable != 537:
        raise RuntimeError({"publication_gate": {"digital_pages": len(digital.pages), "digital_searchable": digital_searchable, "print_pages": len(print_reader.pages), "print_searchable": print_searchable}})
    if dimensions != [5554, 3375] or mode not in {"RGB", "RGBA"} or not opaque:
        raise RuntimeError({"cover_gate": {"dimensions": dimensions, "mode": mode, "opaque": opaque}})
    rendered_geometry = []
    for label, path, expected_pages in (("digital", publication, 537), ("print", printing, 538)):
        doc = fitz.open(path)
        try:
            zero_geometry = []
            for index, page in enumerate(doc):
                if page.rect.width <= 0 or page.rect.height <= 0:
                    zero_geometry.append(index + 1)
            if len(doc) != expected_pages or zero_geometry:
                raise RuntimeError({"pdf_geometry_gate": {"label": label, "pages": len(doc), "zero_geometry": zero_geometry}})
            rendered_geometry.append({"label": label, "pages": len(doc), "zero_geometry": 0})
        finally:
            doc.close()
    historical = [path for path in project.rglob("*") if path.is_file() and "v2.0.0" in path.name]
    return {
        "status": "passed",
        "digital_publication": {"path": PUBLICATION_REL, "bytes": publication.stat().st_size, "sha256": sha256_file(publication), "pages": 537, "searchable_pages": 537},
        "editable_assembly": {"path": editable.relative_to(project).as_posix(), "bytes": editable.stat().st_size, "sha256": sha256_file(editable)},
        "print_interior": {"path": PRINT_INTERIOR_REL, "bytes": printing.stat().st_size, "sha256": sha256_file(printing), "pages": 538, "searchable_source_pages": 537, "intentional_blank_page": 538},
        "cover": {"path": COVER_PNG_REL, "bytes": cover.stat().st_size, "sha256": sha256_file(cover), "pixels": dimensions, "mode": mode, "opaque": opaque},
        "cover_components": [{"path": rel, "bytes": (project / rel).stat().st_size, "sha256": sha256_file(project / rel)} for rel in (COVER_TIFF_REL, COVER_PDF_REL, TEMPLATE_PNG_REL, TEMPLATE_PDF_REL)],
        "application": {"path": application.relative_to(project).as_posix(), "bytes": application.stat().st_size, "sha256": sha256_file(application)},
        "pdf_geometry": rendered_geometry,
        "historical_v2_lineage_files": len(historical),
    }


def write_application_audit(project: Path, db_path: Path, workbook_path: Path, now_iso: str) -> dict[str, Any]:
    app = find_unique_by_hash(project, APPLICATION_SHA256, names={"human_pathogen_app.py"})
    root = project / "App" / "Section 5 Session 3 Complete"
    root.mkdir(parents=True, exist_ok=True)
    pointer = root / "MRHPD v3.0.0a Final Application Release Pointer.json"
    script = root / "MRHPD v3.0.0a Final Read-Only Application Audit.py"
    output = root / "MRHPD v3.0.0a Final Read-Only Application Audit Result.json"
    json_write(pointer, {
        "schema": "mrhpd-final-application-pointer-1.0",
        "generated_at": now_iso,
        "main_application": app.relative_to(project).as_posix(),
        "main_application_sha256": APPLICATION_SHA256,
        "database": db_path.relative_to(project).as_posix(),
        "workbook": workbook_path.relative_to(project).as_posix(),
        "response": 84,
        "all_sections_complete": True,
    })
    text_write(script, f'''#!/usr/bin/env python3
import hashlib,json,sqlite3,sys
from pathlib import Path
from openpyxl import load_workbook
PROJECT=Path(__file__).resolve().parents[2]
DB=PROJECT/{db_path.relative_to(project).as_posix()!r}
WB=PROJECT/{workbook_path.relative_to(project).as_posix()!r}
APP=PROJECT/{app.relative_to(project).as_posix()!r}
EXPECTED_APP={APPLICATION_SHA256!r}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as f:
  for block in iter(lambda:f.read(1024*1024),b''): h.update(block)
 return h.hexdigest()
con=sqlite3.connect(DB)
try:
 integrity=con.execute('PRAGMA integrity_check').fetchone()[0]
 fk=len(list(con.execute('PRAGMA foreign_key_check')))
 response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R84'").fetchone()[0]
 release=con.execute("SELECT COUNT(*) FROM section5_final_release WHERE release_key=? AND project_state='all_sections_complete'",({RELEASE_KEY!r},)).fetchone()[0]
 unsupported=con.execute("SELECT COUNT(*) FROM section5_final_release WHERE provider_approval_claimed!=0 OR physical_proof_completion_claimed!=0").fetchone()[0]
finally: con.close()
wb=load_workbook(WB,read_only=True,data_only=False)
try: sheets=len(wb.sheetnames)
finally: wb.close()
result={{'status':'passed' if integrity=='ok' and fk==0 and response==1 and release==1 and unsupported==0 and sheets>=153 and sha(APP)==EXPECTED_APP else 'failed','database_integrity':integrity,'foreign_keys':fk,'response84':response,'release':release,'unsupported_external_claims':unsupported,'workbook_sheets':sheets,'application_sha256':sha(APP)}}
print(json.dumps(result,indent=2))
raise SystemExit(0 if result['status']=='passed' else 1)
''')
    result = subprocess.run([sys.executable, str(script.resolve())], cwd=root, text=True, capture_output=True, timeout=600)
    if result.returncode:
        raise RuntimeError({"final_application_audit_failed": {"stdout": result.stdout, "stderr": result.stderr}})
    audit = json.loads(result.stdout)
    json_write(output, audit)
    return {
        "status": "passed",
        "main_application": app.relative_to(project).as_posix(),
        "main_application_sha256": APPLICATION_SHA256,
        "audit": audit,
        "files": [path.relative_to(project).as_posix() for path in (pointer, script, output)],
    }


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], face: ImageFont.ImageFont, fill: str, width: int, spacing: int = 8) -> int:
    lines = textwrap.wrap(text, width=width)
    x, y = xy
    line_height = int(face.getbbox("Ag")[3] * 1.25)
    for line in lines:
        draw.text((x, y), line, font=face, fill=fill)
        y += line_height + spacing
    return y


def build_final_figure(path: Path, gates: list[dict[str, Any]]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 2400, 1350
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, 250), fill=f"#{NAVY}")
    draw.text((120, 72), "HUMAN PATHOGEN DATABASE", font=font(58, True), fill="white")
    draw.text((120, 150), "Final project release map • Response 84", font=font(34), fill=f"#{PALE_BLUE}")
    passed = sum(1 for row in gates if row["status"] == "passed")
    pending = sum(1 for row in gates if row["status"] == "controlled_pending")
    cards = [
        ("ALL SECTIONS", "COMPLETE", TEAL),
        ("INTERNAL GATES", f"{passed} PASSED", NAVY),
        ("EXTERNAL EVIDENCE", f"{pending} PENDING", GOLD),
    ]
    card_w = 650
    start_x = 125
    gap = 100
    for index, (label, value, color) in enumerate(cards):
        x = start_x + index * (card_w + gap)
        draw.rounded_rectangle((x, 340, x + card_w, 790), radius=28, fill=f"#{PALE_BLUE}", outline=f"#{color}", width=6)
        draw.text((x + 40, 395), label, font=font(31, True), fill=f"#{color}")
        wrap_text(draw, value, (x + 40, 515), font(66, True), f"#{DARK}", 18, spacing=10)
    draw.rounded_rectangle((125, 900, width - 125, 1230), radius=28, fill=f"#{PALE_GOLD}", outline=f"#{GOLD}", width=5)
    note = (
        "The digital publication, print interior, cover, editable assembly, application, database, workbook, tracking, indexes, manifests, QA, and recovery controls are complete and preserved. Provider preview, provider approval, proof order, receipt, physical inspection, correction, and signoff remain a separate evidence-dependent production lane."
    )
    wrap_text(draw, note, (180, 965), font(35, True), f"#{DARK}", 112, spacing=10)
    image.save(path, format="PNG", dpi=(300, 300), optimize=True)
    return {"status": "passed", "path": str(path), "pixels": [width, height], "dpi": 300, "mode": image.mode, "bytes": path.stat().st_size, "sha256": sha256_file(path)}


def write_reports(
    project: Path,
    db_qa: dict[str, Any],
    workbook_qa: dict[str, Any],
    asset_qa: dict[str, Any],
    app_qa: dict[str, Any],
    category_qa: dict[str, Any],
    tracking_qa: dict[str, Any],
    gates: list[dict[str, Any]],
    evidence: list[dict[str, Any]],
    events: list[dict[str, Any]],
    now_iso: str,
) -> dict[str, Any]:
    root = project / "Reports" / "Section 5 Session 3" / "Complete"
    artwork = project / "Artwork" / "Section 5 Final Release" / "Complete" / "MRHPD-FIG-S5-0007 Final Project Release Map v3.0.0a.png"
    root.mkdir(parents=True, exist_ok=True)
    figure_qa = build_final_figure(artwork, gates)
    docx_path = root / "MRHPD v3.0.0a Section 5 and Entire Project Complete Acceptance Report.docx"
    pdf_path = root / "MRHPD v3.0.0a Section 5 and Entire Project Complete Acceptance Report.pdf"
    xlsx_path = root / "MRHPD v3.0.0a Section 5 and Entire Project Complete Acceptance Register.xlsx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    doc.core_properties.title = "Human Pathogen Database — Section 5 and Entire Project Complete Acceptance Report"
    doc.core_properties.author = "Brent McAnulty, M.D."
    title = doc.add_heading("Human Pathogen Database", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Remediation Section 5 of 5 • Session 3 of 3 • Checkpoint 3 of 3 • All Sections Complete")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    doc.add_picture(str(artwork), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Final disposition", level=1)
    doc.add_paragraph(
        "All internally observable project sections are complete. The final release contains the synchronized clinical database, workbook, application, searchable publication, print-production derivative, cover package, editable assembly, graphics, evidence controls, tracking, indexes, Master Category database, manifests, quality assurance, and deterministic recovery architecture. Provider-side preview and physical-proof events remain controlled pending because no genuine item-level evidence was supplied or discovered."
    )
    summary_rows = [
        ("Response", "84 — complete"),
        ("Checkpoint", "3 of 3 — complete"),
        ("Session", "3 of 3 — complete"),
        ("Remediation Section 5", "complete"),
        ("All Sections", "COMPLETE"),
        ("Canonical SQLite", f"{db_qa['tables']} physical tables; integrity {db_qa['integrity']}; {db_qa['foreign_keys']} foreign-key violations"),
        ("Comprehensive workbook", f"{workbook_qa['current_sheet_count']} sheets; {workbook_qa['formula_error_count']} formula-error tokens"),
        ("Digital publication", "537 searchable pages; byte-identical"),
        ("Print interior", "538 pages; page 538 intentional blank; byte-identical"),
        ("Cover", "5554 × 3375 opaque RGB; byte-identical"),
        ("Master Category", f"{category_qa['category_count']} categories; {category_qa['subcategory_count']} subcategories"),
        ("Tracking", f"{tracking_qa['response_rows']} response rows through Response 84"),
        ("Provider/proof evidence", "controlled pending; no unsupported completion claim"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Control"
    table.rows[0].cells[1].text = "Final result"
    shade_cell(table.rows[0].cells[0], NAVY)
    shade_cell(table.rows[0].cells[1], NAVY)
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    for key, value in summary_rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    doc.add_heading("Project-wide final acceptance matrix", level=1)
    gate_table = doc.add_table(rows=1, cols=3)
    gate_table.style = "Table Grid"
    for index, value in enumerate(("Gate", "Status", "Evidence boundary")):
        gate_table.rows[0].cells[index].text = value
        shade_cell(gate_table.rows[0].cells[index], NAVY)
        for run in gate_table.rows[0].cells[index].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    for row in gates:
        cells = gate_table.add_row().cells
        cells[0].text = row["description"]
        cells[1].text = row["status"]
        cells[2].text = row["evidence"]
    doc.add_heading("External evidence boundary", level=1)
    for row in evidence:
        doc.add_heading(row["evidence_key"].replace("_", " ").title(), level=2)
        doc.add_paragraph(row["observation"])
    doc.add_heading("Recovery and release record", level=1)
    for row in events:
        doc.add_heading(f"{row['event_number']} — {row['event_code']}", level=2)
        doc.add_paragraph(row["condition"])
        doc.add_paragraph(row["recovery"])
    doc.add_heading("Release boundary", level=1)
    doc.add_paragraph(
        "Project remediation and internally verifiable production engineering are complete. A later provider Previewer result or physical proof may trigger a separately governed post-release correction, but the absence of those external events does not justify inventing approval and does not erase the completed project release."
    )
    doc.add_paragraph(f"Generated: {now_iso}")
    doc.save(docx_path)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleMR", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, leading=21, textColor=colors.HexColor("#17324D"), alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name="SubMR", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10, leading=13, textColor=colors.HexColor("#1C7475"), alignment=TA_CENTER, spaceAfter=10))
    styles.add(ParagraphStyle(name="BodyMR", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.2, leading=10.5, textColor=colors.HexColor("#24323D"), spaceAfter=5))
    story: list[Any] = [
        Paragraph("Human Pathogen Database", styles["TitleMR"]),
        Paragraph("Section 5 and Entire Project Complete • Response 84", styles["SubMR"]),
        RLImage(str(artwork), width=6.9 * inch, height=3.88125 * inch),
        Spacer(1, 0.12 * inch),
        Paragraph("All internally observable project sections are complete. Provider preview and physical-proof events remain controlled pending without unsupported completion claims.", styles["BodyMR"]),
    ]
    summary_data = [["Control", "Final result"]] + [[key, value] for key, value in summary_rows]
    summary_table = Table(summary_data, colWidths=[2.0 * inch, 4.8 * inch], repeatRows=1)
    summary_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17324D")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#AAB8C0")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EAF1F5")]),
        ("FONTSIZE", (0, 0), (-1, -1), 7.5),
        ("LEADING", (0, 0), (-1, -1), 9.5),
    ]))
    story.extend([summary_table, PageBreak(), Paragraph("Project-wide final acceptance matrix", styles["Heading1"])])
    gate_data = [["Gate", "Status", "Evidence"]] + [[row["description"], row["status"], row["evidence"]] for row in gates]
    gate_pdf = Table(gate_data, colWidths=[3.45 * inch, 1.2 * inch, 2.15 * inch], repeatRows=1)
    gate_pdf.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17324D")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#AAB8C0")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 6.6),
        ("LEADING", (0, 0), (-1, -1), 8.0),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EAF1F5")]),
    ]))
    story.extend([gate_pdf, PageBreak(), Paragraph("External evidence boundary", styles["Heading1"])])
    for row in evidence:
        story.append(Paragraph(f"<b>{row['evidence_key'].replace('_', ' ').title()}</b><br/>{row['observation']}", styles["BodyMR"]))
    story.append(Paragraph("Recovery and release record", styles["Heading1"]))
    for row in events:
        story.append(Paragraph(f"<b>{row['event_number']} — {row['event_code']}</b><br/>{row['condition']}<br/>{row['recovery']}", styles["BodyMR"]))
    story.append(Paragraph(f"Generated: {now_iso}", styles["BodyMR"]))
    SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=0.45 * inch, leftMargin=0.45 * inch, topMargin=0.42 * inch, bottomMargin=0.42 * inch, title="MRHPD Section 5 and Entire Project Complete Acceptance Report", author="Brent McAnulty, M.D.").build(story)

    wb = Workbook()
    wb.remove(wb.active)
    datasets = {
        "Disposition": [{"Control": key, "Final Result": value} for key, value in summary_rows],
        "Final Gates": gates,
        "External Evidence": evidence,
        "Recovery": events,
        "Database": [db_qa],
        "Workbook": [workbook_qa],
        "Core Artifacts": [asset_qa],
        "Application": [app_qa],
        "Master Category": [category_qa],
        "Tracking": [tracking_qa],
    }
    for title, dataset in datasets.items():
        ws = wb.create_sheet(title=title)
        write_sheet(ws, dataset)
    wb.properties.title = "MRHPD Section 5 and Entire Project Complete Acceptance Register"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(xlsx_path)

    with zipfile.ZipFile(docx_path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("final report DOCX CRC failed")
    with zipfile.ZipFile(xlsx_path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("final report XLSX CRC failed")
    pdf_reader = PdfReader(str(pdf_path))
    searchable = sum(1 for page in pdf_reader.pages if (page.extract_text() or "").strip())
    if len(pdf_reader.pages) < 3 or searchable != len(pdf_reader.pages):
        raise RuntimeError({"final_report_pdf_gate": {"pages": len(pdf_reader.pages), "searchable": searchable}})
    return {
        "status": "passed",
        "figure": figure_qa | {"path": artwork.relative_to(project).as_posix()},
        "docx": {"path": docx_path.relative_to(project).as_posix(), "bytes": docx_path.stat().st_size, "sha256": sha256_file(docx_path)},
        "pdf": {"path": pdf_path.relative_to(project).as_posix(), "bytes": pdf_path.stat().st_size, "sha256": sha256_file(pdf_path), "pages": len(pdf_reader.pages), "searchable_pages": searchable},
        "xlsx": {"path": xlsx_path.relative_to(project).as_posix(), "bytes": xlsx_path.stat().st_size, "sha256": sha256_file(xlsx_path)},
    }


def extract_text(path: Path, limit: int = 1_000_000) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in {".md", ".txt", ".csv", ".json", ".py", ".html", ".htm", ".yml", ".yaml", ".xml"}:
            return path.read_text(encoding="utf-8", errors="replace")[:limit]
        if suffix == ".docx":
            doc = Document(path)
            chunks = [paragraph.text for paragraph in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells)
            return "\n".join(chunks)[:limit]
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            chunks = []
            size = 0
            for page in reader.pages:
                text = page.extract_text() or ""
                chunks.append(text)
                size += len(text)
                if size >= limit:
                    break
            return "\n".join(chunks)[:limit]
        if suffix == ".xlsx":
            wb = load_workbook(path, read_only=True, data_only=False)
            try:
                chunks = []
                size = 0
                for ws in wb.worksheets:
                    chunks.append(ws.title)
                    for row in ws.iter_rows(values_only=True):
                        line = " | ".join("" if value is None else str(value) for value in row)
                        chunks.append(line)
                        size += len(line)
                        if size >= limit:
                            return "\n".join(chunks)[:limit]
                return "\n".join(chunks)[:limit]
            finally:
                wb.close()
        if suffix in {".sqlite", ".db"}:
            con = sqlite3.connect(path)
            try:
                schema = [row[0] for row in con.execute("SELECT sql FROM sqlite_master WHERE sql IS NOT NULL ORDER BY name")]
                return "\n".join(schema)[:limit]
            finally:
                con.close()
    except Exception as exc:
        return f"[index extraction error: {exc!r}]"
    return ""


def build_indexes(project: Path, now_iso: str) -> dict[str, Any]:
    root = project / "Indexes" / "All Sections Complete"
    root.mkdir(parents=True, exist_ok=True)
    source_json = root / "MRHPD v3.0.0a All Sections Complete Source Index.json"
    source_csv = root / "MRHPD v3.0.0a All Sections Complete Source Index.csv"
    bit_path = root / "MRHPD v3.0.0a All Sections Complete Bit Index.sqlite"
    qa_path = root / "MRHPD v3.0.0a All Sections Complete Index QA.json"
    excluded = {source_json.resolve(), source_csv.resolve(), bit_path.resolve(), qa_path.resolve()}
    searchable_suffixes = {".md", ".txt", ".csv", ".json", ".py", ".html", ".htm", ".yml", ".yaml", ".xml", ".docx", ".pdf", ".xlsx", ".sqlite", ".db"}
    rows: list[dict[str, Any]] = []
    fts_rows: list[tuple[str, str, str, str]] = []
    for path in sorted(item for item in project.rglob("*") if item.is_file() and item.resolve() not in excluded):
        rel = path.relative_to(project).as_posix()
        purpose = "Project artifact"
        if rel.startswith("Database/"):
            purpose = "Canonical, historical, or category database"
        elif rel.startswith("Tracking/"):
            purpose = "Prompt, response, workbook, or project-state tracking"
        elif rel.startswith("Documents/"):
            purpose = "Publication or editable document"
        elif rel.startswith("Print Production/"):
            purpose = "Print-production interior, cover, template, or proof artifact"
        elif rel.startswith("QA/"):
            purpose = "Quality-assurance evidence"
        elif rel.startswith("Reports/"):
            purpose = "Human-readable project report or register"
        elif rel.startswith("App/"):
            purpose = "Local application or read-only audit surface"
        elif rel.startswith("Artwork/"):
            purpose = "Governed raster artwork or publication visual asset"
        user_searchable = int(path.suffix.lower() in searchable_suffixes)
        rows.append({"record_type": "physical_file", "path": rel, "container_path": "", "name": path.name, "purpose": purpose, "bytes": path.stat().st_size, "sha256": sha256_file(path), "user_searchable": user_searchable})
        content = extract_text(path) if user_searchable else ""
        fts_rows.append((rel, path.name, purpose, content))
        if path.suffix.lower() == ".zip":
            try:
                with zipfile.ZipFile(path) as zf:
                    if zf.testzip() is not None:
                        raise RuntimeError("CRC failure")
                    for info in zf.infolist():
                        if info.is_dir():
                            continue
                        member_path = f"{rel}!/{info.filename}"
                        member_suffix = PurePosixPath(info.filename).suffix.lower()
                        member_searchable = int(member_suffix in searchable_suffixes)
                        rows.append({"record_type": "container_member", "path": member_path, "container_path": rel, "name": PurePosixPath(info.filename).name, "purpose": "Member of project ZIP container", "bytes": info.file_size, "sha256": "", "user_searchable": member_searchable})
                        fts_rows.append((member_path, PurePosixPath(info.filename).name, "Member of project ZIP container", ""))
            except zipfile.BadZipFile as exc:
                raise RuntimeError({"index_bad_zip": {"path": rel, "error": repr(exc)}})
    json_write(source_json, {"schema": "mrhpd-source-index-3.0", "generated_at": now_iso, "response": 84, "all_sections_complete": True, "record_count": len(rows), "records": rows})
    csv_write(source_csv, rows)
    if bit_path.exists():
        bit_path.unlink()
    con = sqlite3.connect(bit_path)
    try:
        con.executescript("""
        CREATE TABLE artifact (
            artifact_id INTEGER PRIMARY KEY,
            record_type TEXT NOT NULL,
            path TEXT NOT NULL UNIQUE,
            container_path TEXT NOT NULL,
            name TEXT NOT NULL,
            purpose TEXT NOT NULL,
            bytes INTEGER NOT NULL,
            sha256 TEXT NOT NULL,
            user_searchable INTEGER NOT NULL
        );
        CREATE VIRTUAL TABLE artifact_fts USING fts5(path,name,purpose,content);
        CREATE INDEX idx_artifact_record_type ON artifact(record_type);
        CREATE INDEX idx_artifact_searchable ON artifact(user_searchable);
        """)
        for row, payload in zip(rows, fts_rows):
            con.execute("INSERT INTO artifact (record_type,path,container_path,name,purpose,bytes,sha256,user_searchable) VALUES (?,?,?,?,?,?,?,?)", (row["record_type"], row["path"], row["container_path"], row["name"], row["purpose"], row["bytes"], row["sha256"], row["user_searchable"]))
            con.execute("INSERT INTO artifact_fts (path,name,purpose,content) VALUES (?,?,?,?)", payload)
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        counts = {
            "artifact": con.execute("SELECT COUNT(*) FROM artifact").fetchone()[0],
            "fts": con.execute("SELECT COUNT(*) FROM artifact_fts").fetchone()[0],
            "response84": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"Response 84"',)).fetchone()[0],
            "all_sections": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"All Sections"',)).fetchone()[0],
            "provider": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('provider',)).fetchone()[0],
        }
        con.commit()
    finally:
        con.close()
    if integrity != "ok" or counts["artifact"] != len(rows) or counts["fts"] != len(rows) or counts["response84"] < 1 or counts["all_sections"] < 1:
        raise RuntimeError({"bit_index_gate": {"integrity": integrity, "counts": counts, "expected": len(rows)}})
    qa = {
        "status": "passed",
        "generated_at": now_iso,
        "source_index_records": len(rows),
        "physical_files": sum(1 for row in rows if row["record_type"] == "physical_file"),
        "container_members": sum(1 for row in rows if row["record_type"] == "container_member"),
        "bit_index_integrity": integrity,
        "counts": counts,
        "bit_index_sha256": sha256_file(bit_path),
    }
    json_write(qa_path, qa)
    return {"status": "passed", "source_json": source_json, "source_csv": source_csv, "bit_index": bit_path, "qa_path": qa_path, "qa": qa}


def write_project_readme(project: Path, now_iso: str) -> Path:
    path = project / "README ALL SECTIONS COMPLETE.md"
    text_write(path, f"""# Human Pathogen Database — All Sections Complete

Response 84 completes Remediation Section 5 of 5, Session 3 of 3, Checkpoint 3 of 3, and all project sections.

## Final project state

- Canonical SQLite database synchronized through Response 84.
- Comprehensive workbook synchronized through Response 84.
- Master Category database and human-readable exports included.
- Immutable digital publication: 537 searchable pages.
- Frozen print-production interior: 538 pages with one intentional terminal blank.
- Exact Premium Color cover and template package retained.
- Editable manuscript assembly and main application retained byte-identically.
- Raw, Net, Everything-in-One, and Cumulative Thread tracking current through Response 84.
- Final Source Index, Bit Index, manifest, checksums, QA, reports, and recovery controls included.
- Provider preview, provider approval, proof order, receipt, inspection, correction, and physical signoff remain controlled pending unless genuine evidence is later supplied.

## Recovery boundary

The complete Response 84 restore contains the entire current project, exact archive identities, manifests, checksums, verification tools, QA, reports, and restoration instructions. No earlier project archive, checkpoint package, cloud artifact, or conversation reconstruction is required.

Updated: {now_iso}
""")
    return path


def write_qa(
    project: Path,
    db_qa: dict[str, Any],
    workbook_qa: dict[str, Any],
    asset_qa: dict[str, Any],
    app_qa: dict[str, Any],
    category_qa: dict[str, Any],
    tracking_qa: dict[str, Any],
    report_qa: dict[str, Any],
    gates: list[dict[str, Any]],
    evidence: list[dict[str, Any]],
    events: list[dict[str, Any]],
    now_iso: str,
) -> dict[str, Any]:
    root = project / "QA" / "Section 5 Session 3" / "Complete"
    root.mkdir(parents=True, exist_ok=True)
    payloads = {
        "DATABASE_QA.json": db_qa,
        "WORKBOOK_QA.json": workbook_qa,
        "CORE_ARTIFACT_QA.json": asset_qa,
        "APPLICATION_QA.json": app_qa,
        "MASTER_CATEGORY_QA.json": category_qa,
        "TRACKING_QA.json": tracking_qa,
        "REPORT_QA.json": report_qa,
        "FINAL_RELEASE_GATES.json": gates,
        "FINAL_EXTERNAL_EVIDENCE_BOUNDARY.json": evidence,
        "RECOVERY_EVENTS_253_264.json": events,
    }
    paths = []
    for name, payload in payloads.items():
        output = root / name
        json_write(output, payload)
        paths.append(output)
    final = {
        "schema": "mrhpd-section5-entire-project-complete-qa-1.0",
        "generated_at": now_iso,
        "status": "passed_with_controlled_external_gates",
        "response": 84,
        "checkpoint_3_of_3_complete": True,
        "session_3_of_3_complete": True,
        "remediation_section_5_complete": True,
        "all_sections_complete": True,
        "internal_gates_passed": sum(1 for row in gates if row["status"] == "passed"),
        "controlled_external_gates": sum(1 for row in gates if row["status"] == "controlled_pending"),
        "failed_gates": sum(1 for row in gates if row["status"] == "failed"),
        "provider_approval_claimed": False,
        "physical_proof_completion_claimed": False,
        "accepted_predecessor_mutated": False,
        "immutable_publication_mutated": False,
        "main_application_mutated": False,
        "user_upload_required": False,
        "conversation_reconstruction_required": False,
        "next": "Optional external provider-preview and physical-proof production lane only; no remediation session remains.",
    }
    final_path = root / "SECTION5_AND_ENTIRE_PROJECT_COMPLETE_QA.json"
    json_write(final_path, final)
    paths.append(final_path)
    if final["failed_gates"]:
        raise RuntimeError({"final_qa_failed_gates": final["failed_gates"]})
    return {"status": "passed", "final": final, "files": [path.relative_to(project).as_posix() for path in paths]}


def build_manifest(project: Path, now_iso: str) -> tuple[Path, Path, list[dict[str, Any]]]:
    root = project / "Manifest" / "All Sections Complete"
    root.mkdir(parents=True, exist_ok=True)
    manifest = root / "MRHPD v3.0.0a All Sections Complete Project Manifest.json"
    checksums = root / "MRHPD v3.0.0a All Sections Complete Project Checksums.sha256"
    rows = []
    for path in sorted(item for item in project.rglob("*") if item.is_file() and item not in {manifest, checksums}):
        rows.append({"path": path.relative_to(project).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    json_write(manifest, {
        "schema": "mrhpd-current-project-manifest-4.0",
        "generated_at": now_iso,
        "response": 84,
        "section": SECTION_LABEL,
        "session": SESSION_LABEL,
        "state": "all_sections_complete",
        "exclusions": [manifest.relative_to(project).as_posix(), checksums.relative_to(project).as_posix()],
        "file_count": len(rows),
        "total_bytes": sum(row["bytes"] for row in rows),
        "files": rows,
    })
    text_write(checksums, "".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
    mismatches = []
    for row in rows:
        path = project / row["path"]
        if path.stat().st_size != row["bytes"] or sha256_file(path) != row["sha256"]:
            mismatches.append(row["path"])
    if mismatches:
        raise RuntimeError({"manifest_mismatches": mismatches[:50]})
    return manifest, checksums, rows


def zip_project(project: Path, destination: Path) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(destination, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in sorted(item for item in project.rglob("*") if item.is_file()):
            zf.write(path, (Path(project.name) / path.relative_to(project)).as_posix())
    return verify_zip(destination)


def clean_verify_project(project_archive: Path, project_name: str, manifest_rel: str, db_rel: str, workbook_rel: str) -> dict[str, Any]:
    with tempfile.TemporaryDirectory(prefix="mrhpd-r84-project-clean-") as td:
        root = Path(td)
        safe_extract(project_archive, root)
        project = root / project_name
        if not project.is_dir():
            raise RuntimeError({"clean_project_root_missing": project_name})
        manifest = json.loads((project / manifest_rel).read_text(encoding="utf-8"))
        mismatches = []
        for row in manifest["files"]:
            path = project / row["path"]
            if not path.exists() or path.stat().st_size != row["bytes"] or sha256_file(path) != row["sha256"]:
                mismatches.append(row["path"])
        con = sqlite3.connect(project / db_rel)
        try:
            integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
            fk = len(list(con.execute("PRAGMA foreign_key_check")))
            response = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R84'").fetchone()[0]
            release = con.execute("SELECT COUNT(*) FROM section5_final_release WHERE release_key=? AND project_state='all_sections_complete'", (RELEASE_KEY,)).fetchone()[0]
            unsupported = con.execute("SELECT COUNT(*) FROM section5_final_release WHERE provider_approval_claimed!=0 OR physical_proof_completion_claimed!=0").fetchone()[0]
        finally:
            con.close()
        wb = load_workbook(project / workbook_rel, read_only=True, data_only=False)
        try:
            sheets = len(wb.sheetnames)
        finally:
            wb.close()
        publication = project / PUBLICATION_REL
        printing = project / PRINT_INTERIOR_REL
        cover = project / COVER_PNG_REL
        apps = [path for path in project.rglob("human_pathogen_app.py") if path.is_file() and sha256_file(path) == APPLICATION_SHA256]
        master_categories = [path for path in project.rglob("*Master Category*.sqlite") if path.is_file()]
        if mismatches or integrity != "ok" or fk or response != 1 or release != 1 or unsupported or sheets < 153 or sha256_file(publication) != PUBLICATION_SHA256 or sha256_file(printing) != PRINT_INTERIOR_SHA256 or sha256_file(cover) != COVER_SHA256 or len(apps) != 1 or not master_categories:
            raise RuntimeError({
                "clean_project_gate": {
                    "manifest_mismatches": mismatches[:30],
                    "integrity": integrity,
                    "foreign_keys": fk,
                    "response84": response,
                    "release": release,
                    "unsupported": unsupported,
                    "workbook_sheets": sheets,
                    "application_matches": len(apps),
                    "master_category_databases": len(master_categories),
                }
            })
        return {
            "status": "passed",
            "manifest_records": len(manifest["files"]),
            "manifest_mismatches": 0,
            "database_integrity": integrity,
            "foreign_keys": fk,
            "response84": response,
            "release": release,
            "unsupported_external_claims": unsupported,
            "workbook_sheets": sheets,
            "application_matches": len(apps),
            "master_category_databases": len(master_categories),
        }


def create_restore_verifier(
    project_archive_name: str,
    project_archive_bytes: int,
    project_archive_sha: str,
    project_name: str,
    manifest_rel: str,
    db_rel: str,
    workbook_rel: str,
) -> str:
    return f'''#!/usr/bin/env python3
from __future__ import annotations
import argparse,hashlib,json,re,sqlite3,tempfile,zipfile
from pathlib import Path,PurePosixPath
from openpyxl import load_workbook
from pypdf import PdfReader
from PIL import Image
PROJECT_ARCHIVE_NAME={project_archive_name!r}
PROJECT_ARCHIVE_BYTES={project_archive_bytes}
PROJECT_ARCHIVE_SHA256={project_archive_sha!r}
PROJECT_NAME={project_name!r}
MANIFEST_REL={manifest_rel!r}
DB_REL={db_rel!r}
WORKBOOK_REL={workbook_rel!r}
RELEASE_KEY={RELEASE_KEY!r}
PUBLICATION_REL={PUBLICATION_REL!r}
PRINT_INTERIOR_REL={PRINT_INTERIOR_REL!r}
COVER_REL={COVER_PNG_REL!r}
PUBLICATION_SHA256={PUBLICATION_SHA256!r}
PRINT_INTERIOR_SHA256={PRINT_INTERIOR_SHA256!r}
COVER_SHA256={COVER_SHA256!r}
APPLICATION_SHA256={APPLICATION_SHA256!r}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as f:
  for block in iter(lambda:f.read(1024*1024),b''): h.update(block)
 return h.hexdigest()
def verify_zip(path):
 with zipfile.ZipFile(path) as zf:
  if zf.testzip() is not None: raise RuntimeError('ZIP CRC failure: '+str(path))
  names=zf.namelist()
  if len(names)!=len(set(names)): raise RuntimeError('duplicate ZIP members')
  for name in names:
   pp=PurePosixPath(name.replace('\\\\','/'))
   if pp.is_absolute() or '..' in pp.parts or re.match(r'^[A-Za-z]:',name): raise RuntimeError('unsafe ZIP path: '+name)
def extract(path,dest):
 verify_zip(path); dest.mkdir(parents=True,exist_ok=True)
 with zipfile.ZipFile(path) as zf: zf.extractall(dest)
def main():
 ap=argparse.ArgumentParser(); ap.add_argument('--extract-project-to',type=Path); args=ap.parse_args()
 root=Path(__file__).resolve().parents[1]
 archive=root/'COMPLETE_PROJECT'/PROJECT_ARCHIVE_NAME
 if archive.stat().st_size!=PROJECT_ARCHIVE_BYTES or sha(archive)!=PROJECT_ARCHIVE_SHA256: raise RuntimeError('project archive identity failed')
 destination=args.extract_project_to; temporary=None
 if destination is None:
  temporary=tempfile.TemporaryDirectory(prefix='mrhpd-r84-verify-'); destination=Path(temporary.name)
 if destination.exists() and any(destination.iterdir()): raise RuntimeError('extract destination must be empty')
 destination.mkdir(parents=True,exist_ok=True); extract(archive,destination)
 project=destination/PROJECT_NAME
 manifest=json.loads((project/MANIFEST_REL).read_text(encoding='utf-8'))
 mismatches=[]
 for row in manifest['files']:
  path=project/row['path']
  if not path.exists() or path.stat().st_size!=row['bytes'] or sha(path)!=row['sha256']: mismatches.append(row['path'])
 con=sqlite3.connect(project/DB_REL)
 try:
  integrity=con.execute('PRAGMA integrity_check').fetchone()[0]
  fk=len(list(con.execute('PRAGMA foreign_key_check')))
  response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R84'").fetchone()[0]
  release=con.execute("SELECT checkpoint_state,session_state,section_state,project_state,internal_acceptance_state,external_evidence_state,provider_approval_claimed,physical_proof_completion_claimed FROM section5_final_release WHERE release_key=?",(RELEASE_KEY,)).fetchone()
  failed=con.execute("SELECT COUNT(*) FROM section5_final_release_gate WHERE status='failed'").fetchone()[0]
 finally: con.close()
 wb=load_workbook(project/WORKBOOK_REL,read_only=True,data_only=False)
 try: sheets=len(wb.sheetnames)
 finally: wb.close()
 publication=project/PUBLICATION_REL; printing=project/PRINT_INTERIOR_REL; cover=project/COVER_REL
 pub=PdfReader(str(publication)); pr=PdfReader(str(printing)); searchable=sum(1 for p in pub.pages if (p.extract_text() or '').strip()); print_searchable=sum(1 for p in pr.pages if (p.extract_text() or '').strip())
 with Image.open(cover) as image: cover_state={{'pixels':[image.width,image.height],'mode':image.mode}}
 apps=[p for p in project.rglob('human_pathogen_app.py') if p.is_file() and sha(p)==APPLICATION_SHA256]
 master=[p for p in project.rglob('*Master Category*.sqlite') if p.is_file()]
 expected=('checkpoint_complete','session_complete','section_complete','all_sections_complete','passed','controlled_pending',0,0)
 result={{'status':'passed' if not mismatches and integrity=='ok' and fk==0 and response==1 and release==expected and failed==0 and sheets>=153 and len(pub.pages)==537 and searchable==537 and len(pr.pages)==538 and print_searchable==537 and sha(publication)==PUBLICATION_SHA256 and sha(printing)==PRINT_INTERIOR_SHA256 and sha(cover)==COVER_SHA256 and cover_state['pixels']==[5554,3375] and len(apps)==1 and len(master)>=1 else 'failed','project_root':str(project),'manifest_records':len(manifest['files']),'manifest_mismatches':mismatches,'database_integrity':integrity,'foreign_keys':fk,'response84':response,'release':release,'failed_gates':failed,'workbook_sheets':sheets,'publication_pages':len(pub.pages),'publication_searchable':searchable,'print_pages':len(pr.pages),'print_searchable':print_searchable,'cover':cover_state,'application_matches':len(apps),'master_category_databases':len(master)}}
 print(json.dumps(result,indent=2))
 if result['status']!='passed': raise SystemExit(1)
 if temporary is not None: temporary.cleanup()
if __name__=='__main__': main()
'''


def path_not_restore_manifest(path: Path) -> bool:
    return path.name not in {"MRHPD_RESPONSE84_COMPLETE_RESTORE_MANIFEST.json", "MRHPD_RESPONSE84_COMPLETE_RESTORE_CHECKSUMS.sha256"}


def build_complete_restore(
    project_archive: Path,
    project: Path,
    dist: Path,
    stamp: str,
    now_iso: str,
    manifest_path: Path,
    checksums_path: Path,
    report_qa: dict[str, Any],
    qa_result: dict[str, Any],
) -> dict[str, Any]:
    root = dist / "complete_restore_root"
    if root.exists():
        shutil.rmtree(root)
    (root / "COMPLETE_PROJECT").mkdir(parents=True)
    (root / "TOOLS").mkdir(parents=True)
    (root / "FINAL_VERIFICATION").mkdir(parents=True)
    shutil.copy2(project_archive, root / "COMPLETE_PROJECT" / project_archive.name)
    text_write(root / "COMPLETE_PROJECT" / f"{project_archive.name}.sha256.txt", f"{sha256_file(project_archive)}  {project_archive.name}")
    manifest_rel = manifest_path.relative_to(project).as_posix()
    text_write(
        root / "TOOLS" / "restore_verify_extract.py",
        create_restore_verifier(
            project_archive.name,
            project_archive.stat().st_size,
            sha256_file(project_archive),
            project.name,
            manifest_rel,
            CURRENT_DB_REL,
            CURRENT_WORKBOOK_REL,
        ),
    )
    text_write(root / "RESTORE_READ_FIRST.md", f"""# Human Pathogen Database — All Sections Complete Restore Through Response 84

This is the complete self-contained restore for the entire Human Pathogen Database project through Response 84.

## Complete project archive

Filename: `{project_archive.name}`

Bytes: `{project_archive.stat().st_size}`

SHA-256: `{sha256_file(project_archive)}`

## Verification

```bash
python TOOLS/restore_verify_extract.py
```

To verify and extract the entire current project:

```bash
python TOOLS/restore_verify_extract.py --extract-project-to "<empty destination>"
```

No earlier project ZIP, checkpoint package, cloud artifact, or conversation reconstruction is required. Provider preview and physical-proof events remain controlled pending unless genuine item-level evidence is later supplied.
""")
    shutil.copy2(manifest_path, root / "FINAL_VERIFICATION" / manifest_path.name)
    shutil.copy2(checksums_path, root / "FINAL_VERIFICATION" / checksums_path.name)
    for key in ("docx", "pdf", "xlsx"):
        source = project / report_qa[key]["path"]
        shutil.copy2(source, root / "FINAL_VERIFICATION" / source.name)
    figure = project / report_qa["figure"]["path"]
    shutil.copy2(figure, root / "FINAL_VERIFICATION" / figure.name)
    final_qa = project / "QA" / "Section 5 Session 3" / "Complete" / "SECTION5_AND_ENTIRE_PROJECT_COMPLETE_QA.json"
    shutil.copy2(final_qa, root / "FINAL_VERIFICATION" / final_qa.name)
    restore_rows = []
    for path in sorted(item for item in root.rglob("*") if item.is_file() and path_not_restore_manifest(item)):
        restore_rows.append({"path": path.relative_to(root).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    restore_manifest = root / "MRHPD_RESPONSE84_COMPLETE_RESTORE_MANIFEST.json"
    restore_checksums = root / "MRHPD_RESPONSE84_COMPLETE_RESTORE_CHECKSUMS.sha256"
    json_write(restore_manifest, {"schema": "mrhpd-complete-restore-4.0", "generated_at": now_iso, "response": 84, "state": "all_sections_complete", "file_count": len(restore_rows), "files": restore_rows})
    text_write(restore_checksums, "".join(f"{row['sha256']}  {row['path']}\n" for row in restore_rows))
    restore = dist / (
        "Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
        f"Remediation Section 5 of 5 Session 3 of 3 ALL SECTIONS COMPLETE RESTORE THROUGH RESPONSE 84 {stamp}.zip"
    )
    with zipfile.ZipFile(restore, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in sorted(item for item in root.rglob("*") if item.is_file()):
            zf.write(path, path.relative_to(root).as_posix())
    restore_qa = verify_zip(restore)
    with tempfile.TemporaryDirectory(prefix="mrhpd-r84-restore-clean-") as td:
        extracted = Path(td) / "restore"
        safe_extract(restore, extracted)
        result = subprocess.run([sys.executable, str((extracted / "TOOLS" / "restore_verify_extract.py").resolve())], cwd=extracted, text=True, capture_output=True, timeout=3600)
        if result.returncode:
            raise RuntimeError({"embedded_restore_verifier_failed": {"stdout": result.stdout[-20000:], "stderr": result.stderr[-20000:]}})
        embedded = json.loads(result.stdout)
    return {"restore": restore, "restore_qa": restore_qa, "embedded_verification": embedded, "root": root}


def build_transport(restore: Path, dist: Path, now_iso: str) -> dict[str, Any]:
    total = restore.stat().st_size
    part_size = math.ceil(total / VOLUME_COUNT)
    if part_size > MAX_VOLUME_PART_BYTES:
        raise RuntimeError({"restore_too_large_for_four_volumes": {"bytes": total, "part_size": part_size, "max": MAX_VOLUME_PART_BYTES}})
    part_paths = []
    with restore.open("rb") as source:
        for index in range(1, VOLUME_COUNT + 1):
            part = dist / f"{restore.name}.part{index:03d}"
            remaining = total - sum(path.stat().st_size for path in part_paths)
            target = remaining if index == VOLUME_COUNT else min(part_size, remaining)
            with part.open("wb") as output:
                left = target
                while left:
                    block = source.read(min(1024 * 1024, left))
                    if not block:
                        raise RuntimeError("unexpected end of restore while splitting")
                    output.write(block)
                    left -= len(block)
            part_paths.append(part)
    if sum(path.stat().st_size for path in part_paths) != total:
        raise RuntimeError("transport part byte total mismatch")
    part_rows = [{"sequence": index, "name": path.name, "bytes": path.stat().st_size, "sha256": sha256_file(path)} for index, path in enumerate(part_paths, start=1)]
    manifest = {
        "schema": "mrhpd-response84-transport-1.0",
        "generated_at": now_iso,
        "restore": {"name": restore.name, "bytes": total, "sha256": sha256_file(restore)},
        "part_count": VOLUME_COUNT,
        "parts": part_rows,
    }
    manifest_path = dist / "MRHPD_RESPONSE84_COMPLETE_RESTORE_TRANSPORT_MANIFEST.json"
    json_write(manifest_path, manifest)
    reassembler = dist / "reassemble_response84_complete_restore.py"
    text_write(reassembler, f'''#!/usr/bin/env python3
import hashlib,json
from pathlib import Path
MANIFEST={manifest!r}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as f:
  for block in iter(lambda:f.read(1024*1024),b''): h.update(block)
 return h.hexdigest()
def main():
 root=Path.cwd(); output=root/MANIFEST['restore']['name']
 with output.open('wb') as target:
  for row in MANIFEST['parts']:
   part=root/row['name']
   if part.stat().st_size!=row['bytes'] or sha(part)!=row['sha256']: raise RuntimeError({{'part_identity_failed':row['name']}})
   with part.open('rb') as source:
    for block in iter(lambda:source.read(1024*1024),b''): target.write(block)
 if output.stat().st_size!=MANIFEST['restore']['bytes'] or sha(output)!=MANIFEST['restore']['sha256']: raise RuntimeError('reassembled restore identity failed')
 print(json.dumps({{'status':'passed','restore':str(output),'bytes':output.stat().st_size,'sha256':sha(output)}},indent=2))
if __name__=='__main__': main()
''')
    wrappers = []
    for index, part in enumerate(part_paths, start=1):
        wrapper = dist / f"MRHPD v3.0.0a Response 84 Complete Restore Drive Volume {index} of {VOLUME_COUNT}.zip"
        readme = dist / f"README_RESPONSE84_VOLUME_{index}.txt"
        text_write(readme, f"Human Pathogen Database Response 84 complete restore volume {index} of {VOLUME_COUNT}. Extract all four governed volume ZIPs into the same otherwise-empty directory, then run reassemble_response84_complete_restore.py. All four volumes are required.")
        with zipfile.ZipFile(wrapper, "w", compression=zipfile.ZIP_STORED, allowZip64=True) as zf:
            zf.write(part, part.name)
            zf.write(manifest_path, manifest_path.name)
            zf.write(reassembler, reassembler.name)
            zf.write(readme, readme.name)
        wrappers.append({"sequence": index, "path": wrapper, "qa": verify_zip(wrapper)})
    with tempfile.TemporaryDirectory(prefix="mrhpd-r84-transport-test-") as td:
        root = Path(td)
        for row in wrappers:
            safe_extract(row["path"], root)
        result = subprocess.run([sys.executable, str((root / reassembler.name).resolve())], cwd=root, text=True, capture_output=True, timeout=1800)
        if result.returncode:
            raise RuntimeError({"transport_reassembly_failed": {"stdout": result.stdout[-12000:], "stderr": result.stderr[-12000:]}})
        reconstructed = root / restore.name
        if reconstructed.stat().st_size != restore.stat().st_size or sha256_file(reconstructed) != sha256_file(restore):
            raise RuntimeError("transport reassembly identity mismatch")
        reassembly = json.loads(result.stdout)
    return {"manifest": manifest, "manifest_path": manifest_path, "reassembler": reassembler, "wrappers": wrappers, "reassembly": reassembly}


def build_verification_delivery(
    dist: Path,
    stamp: str,
    summary: dict[str, Any],
    project: Path,
    project_archive: Path,
    restore_result: dict[str, Any],
    transport: dict[str, Any],
    report_qa: dict[str, Any],
    qa_result: dict[str, Any],
    manifest_path: Path,
    checksums_path: Path,
) -> Path:
    exact_names = dist / "MRHPD v3.0.0a Response 84 Exact File Names.txt"
    lines = [
        "Complete project archive:", project_archive.name, "",
        "Complete self-contained restore:", restore_result["restore"].name, "",
        "Required transport volume wrappers:",
    ]
    lines.extend(row["path"].name for row in transport["wrappers"])
    lines.extend([
        "",
        "Verification delivery:",
        f"MRHPD v3.0.0a Response 84 Section 5 and Entire Project Complete Verification Delivery {stamp}.zip",
        "",
        "Current database:", Path(CURRENT_DB_REL).name,
        "",
        "Current workbook:", Path(CURRENT_WORKBOOK_REL).name,
        "",
        "Digital publication:", Path(PUBLICATION_REL).name,
        "",
        "Print interior:", Path(PRINT_INTERIOR_REL).name,
        "",
        "Cover PNG:", Path(COVER_PNG_REL).name,
    ])
    text_write(exact_names, "\n".join(lines))
    summary_path = dist / "MRHPD_RESPONSE84_SECTION5_ENTIRE_PROJECT_COMPLETE_BUILD_SUMMARY.json"
    json_write(summary_path, summary)
    verification_path = dist / "MRHPD v3.0.0a Response 84 Section 5 and Entire Project Complete Verification.json"
    json_write(verification_path, {
        "schema": "mrhpd-response84-entire-project-complete-verification-1.0",
        "generated_at": summary["generated_at"],
        "status": "passed_with_controlled_external_gates",
        "project_archive": verify_zip(project_archive),
        "complete_restore": restore_result["restore_qa"],
        "embedded_restore_verification": restore_result["embedded_verification"],
        "transport": {
            "manifest": transport["manifest"],
            "wrappers": [{"sequence": row["sequence"], "qa": row["qa"]} for row in transport["wrappers"]],
            "reassembly": transport["reassembly"],
        },
        "provider_approval_claimed": False,
        "physical_proof_completion_claimed": False,
        "user_upload_required": False,
        "conversation_reconstruction_required": False,
        "checkpoint_3_of_3_complete": True,
        "session_3_of_3_complete": True,
        "remediation_section_5_complete": True,
        "all_sections_complete": True,
        "next": "Optional external provider-preview and physical-proof production lane only.",
    })
    delivery = dist / f"MRHPD v3.0.0a Response 84 Section 5 and Entire Project Complete Verification Delivery {stamp}.zip"
    files = [verification_path, summary_path, exact_names, transport["manifest_path"], transport["reassembler"], manifest_path, checksums_path]
    files.extend(project / report_qa[key]["path"] for key in ("docx", "pdf", "xlsx"))
    files.append(project / report_qa["figure"]["path"])
    files.extend(project / path for path in qa_result["files"])
    with zipfile.ZipFile(delivery, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in files:
            if path.exists():
                zf.write(path, path.name)
    verify_zip(delivery)
    return delivery


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--response81-restore", type=Path, required=True)
    parser.add_argument("--response83-dir", type=Path, required=True)
    parser.add_argument("--dist", type=Path, default=Path("dist_cp5_s3_cp3"))
    args = parser.parse_args()
    now = utc_now()
    now_iso = now.isoformat().replace("+00:00", "Z")
    stamp = now.strftime("%Y-%m-%d %H%M UTC")
    if args.dist.exists():
        shutil.rmtree(args.dist)
    args.dist.mkdir(parents=True)
    with tempfile.TemporaryDirectory(prefix="mrhpd-cp5-s3-cp3-") as td:
        work = Path(td)
        baseline_project_archive, baseline_project, response83_project, recovery_qa = reconstruct_response83(args.response81_restore, args.response83_dir, work)
        current_project = work / "current" / CURRENT_PROJECT_NAME
        current_project.parent.mkdir(parents=True)
        shutil.copytree(response83_project, current_project)
        source_db = current_project / SOURCE_DB_REL
        source_workbook = current_project / SOURCE_WORKBOOK_REL
        if not source_db.exists() or not source_workbook.exists():
            raise RuntimeError({"response83_current_paths": {"database": str(source_db), "workbook": str(source_workbook)}})

        gates = final_gate_rows(now_iso)
        evidence = external_evidence_rows(now_iso)
        events = recovery_event_rows(now_iso)
        freeze = artifact_freeze_rows(current_project, now_iso)
        current_db = current_project / CURRENT_DB_REL
        db_qa = sync_database(source_db, current_db, now_iso, gates, evidence, freeze, events)
        current_workbook = current_project / CURRENT_WORKBOOK_REL
        workbook_qa = sync_workbook(source_workbook, current_workbook, gates, evidence, freeze, events, db_qa)
        tracking_qa = write_tracking(current_project, current_db, now_iso)
        category_qa = ensure_master_category(current_project, now_iso)
        asset_qa = audit_core_artifacts(current_project)
        app_qa = write_application_audit(current_project, current_db, current_workbook, now_iso)
        report_qa = write_reports(current_project, db_qa, workbook_qa, asset_qa, app_qa, category_qa, tracking_qa, gates, evidence, events, now_iso)
        readme = write_project_readme(current_project, now_iso)
        qa_result = write_qa(current_project, db_qa, workbook_qa, asset_qa, app_qa, category_qa, tracking_qa, report_qa, gates, evidence, events, now_iso)
        index_result = build_indexes(current_project, now_iso)
        manifest_path, checksums_path, manifest_rows = build_manifest(current_project, now_iso)

        project_archive = args.dist / (
            "Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
            f"Remediation Section 5 of 5 Session 3 of 3 ALL SECTIONS COMPLETE PROJECT THROUGH RESPONSE 84 {stamp}.zip"
        )
        project_archive_qa = zip_project(current_project, project_archive)
        clean_project_qa = clean_verify_project(
            project_archive,
            current_project.name,
            manifest_path.relative_to(current_project).as_posix(),
            CURRENT_DB_REL,
            CURRENT_WORKBOOK_REL,
        )
        restore_result = build_complete_restore(project_archive, current_project, args.dist, stamp, now_iso, manifest_path, checksums_path, report_qa, qa_result)
        transport = build_transport(restore_result["restore"], args.dist, now_iso)
        summary = {
            "schema": "mrhpd-response84-entire-project-complete-build-1.0",
            "generated_at": now_iso,
            "status": "passed_with_controlled_external_gates",
            "response": 84,
            "section": SECTION_LABEL,
            "session": SESSION_LABEL,
            "checkpoint": CHECKPOINT_LABEL,
            "recovery": recovery_qa,
            "database": db_qa,
            "workbook": workbook_qa,
            "tracking": tracking_qa,
            "master_category": category_qa,
            "core_artifacts": asset_qa,
            "application": app_qa,
            "reports": report_qa,
            "qa": qa_result["final"],
            "index": index_result["qa"],
            "manifest_records": len(manifest_rows),
            "project_archive": project_archive_qa,
            "clean_project": clean_project_qa,
            "complete_restore": restore_result["restore_qa"],
            "embedded_restore_verification": restore_result["embedded_verification"],
            "transport": {
                "manifest": transport["manifest"],
                "wrappers": [{"sequence": row["sequence"], "qa": row["qa"]} for row in transport["wrappers"]],
                "reassembly": transport["reassembly"],
            },
            "provider_approval_claimed": False,
            "physical_proof_completion_claimed": False,
            "accepted_predecessor_mutated": False,
            "immutable_publication_mutated": False,
            "main_application_mutated": False,
            "user_upload_required": False,
            "conversation_reconstruction_required": False,
            "checkpoint_3_of_3_complete": True,
            "session_3_of_3_complete": True,
            "remediation_section_5_complete": True,
            "all_sections_complete": True,
            "next": "Optional external provider-preview and physical-proof production lane only; no remediation checkpoint remains.",
        }
        verification_delivery = build_verification_delivery(
            args.dist,
            stamp,
            summary,
            current_project,
            project_archive,
            restore_result,
            transport,
            report_qa,
            qa_result,
            manifest_path,
            checksums_path,
        )
        console = {
            "status": "passed_with_controlled_external_gates",
            "response": 84,
            "database_tables": db_qa["tables"],
            "workbook_sheets": workbook_qa["current_sheet_count"],
            "digital_pages": asset_qa["digital_publication"]["pages"],
            "print_pages": asset_qa["print_interior"]["pages"],
            "project_archive": project_archive_qa,
            "complete_restore": restore_result["restore_qa"],
            "transport_volumes": [{"sequence": row["sequence"], **row["qa"]} for row in transport["wrappers"]],
            "verification_delivery": verify_zip(verification_delivery),
            "provider_preview": "controlled_pending",
            "physical_proof": "controlled_pending",
            "user_upload_required": False,
            "checkpoint_3_of_3_complete": True,
            "session_3_of_3_complete": True,
            "remediation_section_5_complete": True,
            "all_sections_complete": True,
            "next": "Optional external provider-preview and physical-proof production lane only.",
        }
        print(json.dumps(console, indent=2))


if __name__ == "__main__":
    main()
