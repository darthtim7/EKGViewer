#!/usr/bin/env python3
from __future__ import annotations

import argparse
import collections
import csv
import hashlib
import json
import math
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import textwrap
import zipfile
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any, Iterable

import fitz
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont, ImageOps, ImageStat
from pypdf import PdfReader

HERE = Path(__file__).resolve().parent
CP1_DIR = HERE.parent / "checkpoint1"
if str(CP1_DIR) not in sys.path:
    sys.path.insert(0, str(CP1_DIR))
if str(HERE) not in sys.path:
    sys.path.insert(0, str(HERE))

import build_section5_checkpoint1 as cp1  # noqa: E402
import section5_checkpoint2_reporting as reporting  # noqa: E402

PROJECT_VERSION = "3.0.0a"
RESPONSE_NUMBER = 76
SECTION_LABEL = "Remediation Section 5 of 5"
SESSION_LABEL = "Session 1 of 3"
CHECKPOINT_LABEL = "Checkpoint 2 of 3"
CHECKPOINT_CODE = "MRHPD-V3-CP5-S1-CP2"
BASE_RESTORE_BYTES = cp1.BASE_RESTORE_BYTES
BASE_RESTORE_SHA256 = cp1.BASE_RESTORE_SHA256
BASE_PROJECT_BYTES = cp1.BASE_PROJECT_BYTES
BASE_PROJECT_SHA256 = cp1.BASE_PROJECT_SHA256
PUBLICATION_SHA256 = cp1.PUBLICATION_SHA256
EDITABLE_SHA256 = cp1.EDITABLE_SHA256
APPLICATION_SHA256 = cp1.APPLICATION_SHA256
PUBLICATION_REL = cp1.PUBLICATION_REL
COVER_HASHES = cp1.COVER_HASHES
CP1_DB_REL = cp1.CURRENT_DB_REL
CP1_WORKBOOK_REL = cp1.CURRENT_WORKBOOK_REL
CP1_RECOVERY_BYTES = 9_408_775
CP1_RECOVERY_SHA256 = "c841aaa58df4836857ffc26260cb30f0d24c6449b6db0efd6b2ea69e2f078bdc"
CP1_DELIVERY_BYTES = 8_915_858
CP1_DELIVERY_SHA256 = "595e728b0b63bce8b13e8586ddcce55f87b8904752f65db39d1e7339c5b1487f"
CURRENT_DB_REL = (
    "Database/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 1 of 3 Checkpoint 2 of 3 THROUGH RESPONSE 76.sqlite"
)
CURRENT_WORKBOOK_REL = (
    "Tracking/Workbook/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 1 of 3 Checkpoint 2 of 3 THROUGH RESPONSE 76 Comprehensive Tracking.xlsx"
)
PRINT_ROOT_REL = "Print Production/KDP Premium Color Response 76"
PRINT_INTERIOR_REL = (
    PRINT_ROOT_REL + "/Interior/Medical References - Human Pathogen Database v3.0.0a "
    "KDP Premium Color 8.5 x 11 Print Interior 538 Pages Response 76.pdf"
)
COVER_ROOT_REL = PRINT_ROOT_REL + "/Cover"
FINAL_COVER_PNG_REL = COVER_ROOT_REL + "/MRHPD v3.0.0a KDP Premium Color 538-Page Full Cover Response 76 300ppi RGB.png"
FINAL_COVER_TIFF_REL = COVER_ROOT_REL + "/MRHPD v3.0.0a KDP Premium Color 538-Page Full Cover Response 76 300ppi RGB LZW.tif"
FINAL_COVER_PDF_REL = COVER_ROOT_REL + "/MRHPD v3.0.0a KDP Premium Color 538-Page Full Cover Response 76.pdf"
TEMPLATE_PNG_REL = COVER_ROOT_REL + "/MRHPD v3.0.0a KDP Premium Color 538-Page Exact Cover Template Response 76.png"
TEMPLATE_PDF_REL = COVER_ROOT_REL + "/MRHPD v3.0.0a KDP Premium Color 538-Page Exact Cover Template Response 76.pdf"
PROOF_PNG_REL = COVER_ROOT_REL + "/MRHPD v3.0.0a KDP Premium Color 538-Page Full Cover Proof Response 76.png"

NAVY = "17324D"
TEAL = "1C7475"
GOLD = "C9A227"
PALE_BLUE = "EAF1F5"
WHITE = "FFFFFF"


def utc_now() -> datetime:
    return datetime.now(timezone.utc).replace(microsecond=0)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False), encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value, encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fields is None:
        fields = list(rows[0]) if rows else []
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def verify_zip(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
        names = zf.namelist()
        unsafe = []
        filler = []
        for name in names:
            pp = PurePosixPath(name.replace("\\", "/"))
            if pp.is_absolute() or ".." in pp.parts or re.match(r"^[A-Za-z]:", name):
                unsafe.append(name)
            lower = name.lower()
            if any(token in lower for token in ("filler", "padding", "dummy_payload", "artificial_inflation")):
                filler.append(name)
    result = {
        "path": str(path),
        "name": path.name,
        "bytes": path.stat().st_size,
        "sha256": sha256_file(path),
        "members": len(names),
        "crc_error": bad,
        "duplicates": len(names) - len(set(names)),
        "unsafe_paths": unsafe,
        "filler_members": filler,
    }
    if bad or result["duplicates"] or unsafe or filler:
        raise RuntimeError({"zip_verification_failed": result})
    return result


def safe_extract(path: Path, destination: Path) -> None:
    destination.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError(f"ZIP CRC failure: {path}")
        names = zf.namelist()
        if len(names) != len(set(names)):
            raise RuntimeError(f"duplicate ZIP members: {path}")
        for name in names:
            pp = PurePosixPath(name.replace("\\", "/"))
            if pp.is_absolute() or ".." in pp.parts or re.match(r"^[A-Za-z]:", name):
                raise RuntimeError(f"unsafe ZIP path: {name}")
        zf.extractall(destination)


def find_exact_zip_recursive(root: Path, size: int, digest: str, work: Path) -> Path:
    queue = [p for p in root.rglob("*.zip") if p.is_file()]
    seen: set[str] = set()
    sequence = 0
    while queue:
        candidate = queue.pop(0)
        identity = f"{candidate.stat().st_size}:{sha256_file(candidate)}"
        if identity in seen:
            continue
        seen.add(identity)
        if candidate.stat().st_size == size and sha256_file(candidate) == digest:
            return candidate
        sequence += 1
        target = work / f"nested-{sequence:04d}"
        try:
            safe_extract(candidate, target)
        except (zipfile.BadZipFile, RuntimeError):
            continue
        queue.extend(p for p in target.rglob("*.zip") if p.is_file())
    raise RuntimeError({"exact_zip_not_found": {"root": str(root), "bytes": size, "sha256": digest}})


def restore_checkpoint1(
    volume1_dir: Path,
    volume2_dir: Path,
    checkpoint1_dir: Path,
    work: Path,
) -> tuple[Path, Path, Path, Path, dict[str, Any]]:
    restore, project_archive, baseline_project = cp1.reconstruct_baseline(volume1_dir, volume2_dir, work / "baseline")
    recovery_zip = find_exact_zip_recursive(checkpoint1_dir, CP1_RECOVERY_BYTES, CP1_RECOVERY_SHA256, work / "cp1-discovery")
    package_root = work / "cp1-package"
    safe_extract(recovery_zip, package_root)
    apply_script = package_root / "TOOLS" / "apply_checkpoint_recovery.py"
    if not apply_script.exists():
        raise RuntimeError("Checkpoint 1 apply utility is missing")
    restored_root = work / "cp1-restored"
    result = subprocess.run(
        [sys.executable, str(apply_script), "--base-response72-restore", str(restore), "--output-dir", str(restored_root)],
        cwd=package_root,
        text=True,
        capture_output=True,
        timeout=1200,
    )
    if result.returncode:
        raise RuntimeError({"checkpoint1_apply_failed": {"stdout": result.stdout[-12000:], "stderr": result.stderr[-12000:]}})
    result_json = restored_root / "MRHPD_RESPONSE75_CHECKPOINT_RECOVERY_APPLICATION_RESULT.json"
    if not result_json.exists():
        raise RuntimeError("Checkpoint 1 application result is missing")
    application = json.loads(result_json.read_text(encoding="utf-8"))
    if application.get("status") != "passed":
        raise RuntimeError({"checkpoint1_application_gate": application})
    project_candidates = [p for p in restored_root.iterdir() if p.is_dir()]
    if len(project_candidates) != 1:
        raise RuntimeError({"checkpoint1_project_candidates": [str(p) for p in project_candidates]})
    checkpoint1_project = project_candidates[0]
    return restore, project_archive, baseline_project, checkpoint1_project, application


def selection_record(now_iso: str) -> dict[str, Any]:
    pages = 538
    spine = pages * 0.002347
    cover_width = 2 * 8.5 + spine + 2 * 0.125
    cover_height = 11.0 + 2 * 0.125
    return {
        "selection_code": "MRHPD-KDP-PB-PREMIUM-COLOR-WHITE-8.5X11-538-MATTE-R76",
        "provider": "Amazon KDP",
        "binding": "Paperback perfect bound",
        "interior_type": "Premium Color",
        "paper": "White paper",
        "cover_finish": "Matte",
        "reading_direction": "Left-to-right",
        "trim_width_in": 8.5,
        "trim_height_in": 11.0,
        "interior_bleed": "No bleed",
        "cover_bleed_in": 0.125,
        "production_page_count": pages,
        "spine_factor_in_per_page": 0.002347,
        "spine_width_in": round(spine, 6),
        "cover_width_in": round(cover_width, 6),
        "cover_height_in": round(cover_height, 6),
        "cover_width_px_300dpi": math.ceil(cover_width * 300),
        "cover_height_px_300dpi": math.ceil(cover_height * 300),
        "maximum_pages_at_trim": 590,
        "selection_status": "locked_initial_production_master",
        "selection_rationale": (
            "Premium color is the controlling initial production master because the reference contains numerous color-coded teaching figures and the project prioritizes sharp image reproduction and color clarity over minimum unit cost. The 538-page production count remains within the current KDP 8.5 × 11-inch premium-color capacity. Standard color may be generated later as a cost-sensitive derivative after the premium-color master passes provider preview and physical-proof review."
        ),
        "official_sources": [
            "https://kdp.amazon.com/en_US/help/topic/G201857950",
            "https://kdp.amazon.com/en_US/help/topic/GVBQ3CMEQW3W2VL6/",
            "https://kdp.amazon.com/en_US/help/topic/G201834180",
            "https://kdp.amazon.com/en_US/help/topic/G201953020",
            "https://kdp.amazon.com/en_US/cover-calculator",
        ],
        "recorded_at": now_iso,
    }


def _rect_union(rects: list[fitz.Rect], page_rect: fitz.Rect) -> fitz.Rect:
    valid = []
    for rect in rects:
        r = fitz.Rect(rect) & page_rect
        if not r.is_empty and r.width > 0.01 and r.height > 0.01:
            valid.append(r)
    if not valid:
        return fitz.Rect(page_rect)
    union = fitz.Rect(valid[0])
    for rect in valid[1:]:
        union |= rect
    padding = 2.0
    union = fitz.Rect(union.x0 - padding, union.y0 - padding, union.x1 + padding, union.y1 + padding) & page_rect
    return union


def page_content_bbox(page: fitz.Page) -> fitz.Rect:
    rects: list[fitz.Rect] = []
    try:
        for row in page.get_bboxlog():
            if len(row) >= 2:
                rects.append(fitz.Rect(row[1]))
    except Exception:
        pass
    for block in page.get_text("blocks"):
        if len(block) >= 5 and str(block[4]).strip():
            rects.append(fitz.Rect(block[:4]))
    return _rect_union(rects, page.rect)


def normalize_words(text: str) -> collections.Counter[str]:
    words = re.findall(r"[\w]+(?:[-’'][\w]+)*", text.casefold(), flags=re.UNICODE)
    return collections.Counter(words)


def image_ppi_audit(page: fitz.Page) -> list[dict[str, Any]]:
    rows = []
    for info in page.get_image_info(xrefs=True):
        bbox = fitz.Rect(info.get("bbox", (0, 0, 0, 0)))
        width = int(info.get("width", 0) or 0)
        height = int(info.get("height", 0) or 0)
        if bbox.width <= 0 or bbox.height <= 0 or width <= 0 or height <= 0:
            continue
        x_ppi = width / (bbox.width / 72.0)
        y_ppi = height / (bbox.height / 72.0)
        rows.append({"xref": info.get("xref"), "width_px": width, "height_px": height, "bbox": [bbox.x0, bbox.y0, bbox.x1, bbox.y1], "x_ppi": round(x_ppi, 2), "y_ppi": round(y_ppi, 2), "minimum_ppi": round(min(x_ppi, y_ppi), 2)})
    return rows


def generate_print_interior(source: Path, output: Path, qa_root: Path, selection: dict[str, Any]) -> tuple[dict[str, Any], list[dict[str, Any]], list[Path]]:
    output.parent.mkdir(parents=True, exist_ok=True)
    qa_root.mkdir(parents=True, exist_ok=True)
    source_doc = fitz.open(source)
    if source_doc.page_count != 537 or sha256_file(source) != PUBLICATION_SHA256:
        raise RuntimeError("immutable publication identity failed before print derivation")
    output_doc = fitz.open()
    transform_rows: list[dict[str, Any]] = []
    source_inside_failures = 0
    for index, source_page in enumerate(source_doc):
        source_text_blocks = [fitz.Rect(b[:4]) for b in source_page.get_text("blocks") if len(b) >= 5 and str(b[4]).strip()]
        if source_text_blocks:
            source_text_bbox = _rect_union(source_text_blocks, source_page.rect)
            source_inside = source_text_bbox.x0 / 72.0 if (index + 1) % 2 == 1 else (source_page.rect.width - source_text_bbox.x1) / 72.0
            if source_inside < 0.75 - 0.002:
                source_inside_failures += 1
        clip = page_content_bbox(source_page)
        rotate = 90 if source_page.rect.width > source_page.rect.height else 0
        page_number = index + 1
        left = 54.0 if page_number % 2 == 1 else 18.0
        right = 18.0 if page_number % 2 == 1 else 54.0
        safe = fitz.Rect(left, 18.0, 612.0 - right, 774.0)
        rotated_width = clip.height if rotate else clip.width
        rotated_height = clip.width if rotate else clip.height
        scale = min(1.0, safe.width / rotated_width, safe.height / rotated_height)
        target_width = rotated_width * scale
        target_height = rotated_height * scale
        x0 = safe.x0 + (safe.width - target_width) / 2.0
        y0 = safe.y0 + (safe.height - target_height) / 2.0
        target = fitz.Rect(x0, y0, x0 + target_width, y0 + target_height)
        out_page = output_doc.new_page(width=612.0, height=792.0)
        out_page.show_pdf_page(target, source_doc, index, keep_proportion=True, rotate=rotate, clip=clip)
        transformed = rotate != 0 or abs(scale - 1.0) > 0.0005 or abs(target.x0 - clip.x0) > 0.5 or abs(target.y0 - clip.y0) > 0.5
        transform_rows.append({
            "page": page_number,
            "source_width_pt": round(source_page.rect.width, 3),
            "source_height_pt": round(source_page.rect.height, 3),
            "source_orientation": "landscape" if rotate else "portrait",
            "rotation_degrees": rotate,
            "clip_x0": round(clip.x0, 3),
            "clip_y0": round(clip.y0, 3),
            "clip_x1": round(clip.x1, 3),
            "clip_y1": round(clip.y1, 3),
            "scale": round(scale, 8),
            "target_x0": round(target.x0, 3),
            "target_y0": round(target.y0, 3),
            "target_x1": round(target.x1, 3),
            "target_y1": round(target.y1, 3),
            "transformed": transformed,
            "transform_class": "rotated_and_scaled" if rotate and scale < 0.9995 else ("rotated" if rotate else ("scaled" if scale < 0.9995 else "translated")),
        })
    output_doc.new_page(width=612.0, height=792.0)
    try:
        toc = source_doc.get_toc(simple=True)
        if toc:
            output_doc.set_toc(toc)
    except Exception:
        pass
    metadata = dict(source_doc.metadata or {})
    metadata.update({
        "title": "Human Pathogen Database v3.0.0a — KDP Premium Color Print Interior — Response 76",
        "author": "Brent McAnulty, M.D.",
        "subject": "538-page print-production derivative; 537-page digital publication remains immutable",
        "keywords": "Human Pathogen Database; KDP; premium color; print interior; 538 pages",
    })
    output_doc.set_metadata(metadata)
    output_doc.save(output, garbage=4, deflate=True, clean=True)
    output_doc.close()

    output_qa = fitz.open(output)
    if output_qa.page_count != 538:
        raise RuntimeError({"print_interior_page_count": output_qa.page_count})
    page_rows: list[dict[str, Any]] = []
    searchable = 0
    output_inside_failures = 0
    output_outside_failures = 0
    output_top_bottom_failures = 0
    text_mismatch_pages = 0
    unembedded_fonts: set[str] = set()
    low_ppi_images: list[dict[str, Any]] = []
    transparency_objects = 0
    proof_pages = [1, 2, 3, 4, 15, 16, 57, 96, 152, 268, 400, 500, 537, 538]
    proof_paths: list[Path] = []
    for index in range(538):
        page = output_qa[index]
        page_number = index + 1
        text = page.get_text("text")
        if index < 537 and text.strip():
            searchable += 1
        source_counter = normalize_words(source_doc[index].get_text("text")) if index < 537 else collections.Counter()
        output_counter = normalize_words(text)
        text_equal = source_counter == output_counter if index < 537 else not output_counter
        if not text_equal:
            text_mismatch_pages += 1
        blocks = [fitz.Rect(b[:4]) for b in page.get_text("blocks") if len(b) >= 5 and str(b[4]).strip()]
        if blocks:
            bbox = _rect_union(blocks, page.rect)
            inside = bbox.x0 / 72.0 if page_number % 2 == 1 else (page.rect.width - bbox.x1) / 72.0
            outside = (page.rect.width - bbox.x1) / 72.0 if page_number % 2 == 1 else bbox.x0 / 72.0
            top = bbox.y0 / 72.0
            bottom = (page.rect.height - bbox.y1) / 72.0
            inside_pass = inside >= 0.75 - 0.004
            outside_pass = outside >= 0.25 - 0.004
            top_bottom_pass = top >= 0.25 - 0.004 and bottom >= 0.25 - 0.004
            output_inside_failures += int(not inside_pass)
            output_outside_failures += int(not outside_pass)
            output_top_bottom_failures += int(not top_bottom_pass)
        else:
            bbox = None
            inside = outside = top = bottom = None
            inside_pass = outside_pass = top_bottom_pass = page_number == 538
        for font in page.get_fonts(full=True):
            if font and int(font[0]) <= 0:
                unembedded_fonts.add(str(font[3] if len(font) > 3 else font))
        image_rows = image_ppi_audit(page)
        for image_row in image_rows:
            if image_row["minimum_ppi"] < 299.0:
                low_ppi_images.append({"page": page_number, **image_row})
        try:
            for drawing in page.get_drawings():
                if float(drawing.get("fill_opacity", 1.0)) < 0.999 or float(drawing.get("stroke_opacity", 1.0)) < 0.999:
                    transparency_objects += 1
        except Exception:
            pass
        row = {
            **transform_rows[index] if index < 537 else {
                "page": 538,
                "source_width_pt": None,
                "source_height_pt": None,
                "source_orientation": "intentional_blank",
                "rotation_degrees": 0,
                "clip_x0": None,
                "clip_y0": None,
                "clip_x1": None,
                "clip_y1": None,
                "scale": None,
                "target_x0": None,
                "target_y0": None,
                "target_x1": None,
                "target_y1": None,
                "transformed": False,
                "transform_class": "intentional_blank",
            },
            "searchable": bool(text.strip()),
            "text_equal_to_source": text_equal,
            "output_text_bbox": None if bbox is None else [round(v, 3) for v in (bbox.x0, bbox.y0, bbox.x1, bbox.y1)],
            "inside_text_margin_in": None if inside is None else round(inside, 4),
            "outside_text_margin_in": None if outside is None else round(outside, 4),
            "top_text_margin_in": None if top is None else round(top, 4),
            "bottom_text_margin_in": None if bottom is None else round(bottom, 4),
            "inside_pass": inside_pass,
            "outside_pass": outside_pass,
            "top_bottom_pass": top_bottom_pass,
            "raster_images": len(image_rows),
            "minimum_image_ppi": min((r["minimum_ppi"] for r in image_rows), default=None),
        }
        page_rows.append(row)
        if page_number in proof_pages:
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            proof_path = qa_root / "Rendered Proofs" / f"MRHPD v3.0.0a Response 76 Print Interior Page {page_number:03d}.png"
            proof_path.parent.mkdir(parents=True, exist_ok=True)
            pix.save(proof_path)
            proof_paths.append(proof_path)
    output_qa.close()
    source_doc.close()
    transformed_pages = sum(1 for row in transform_rows if row["transformed"])
    translated_pages = sum(1 for row in transform_rows if row["transform_class"] == "translated")
    scaled_pages = sum(1 for row in transform_rows if "scaled" in row["transform_class"])
    summary = {
        "status": "passed",
        "source_path": str(source),
        "source_bytes": source.stat().st_size,
        "source_sha256": sha256_file(source),
        "source_page_count": 537,
        "source_inside_failures": source_inside_failures,
        "output_path": str(output),
        "output_name": output.name,
        "output_bytes": output.stat().st_size,
        "output_sha256": sha256_file(output),
        "output_page_count": 538,
        "searchable_pages": searchable,
        "intentional_blank_page": 538,
        "transformed_pages": transformed_pages,
        "translated_pages": translated_pages,
        "scaled_pages": scaled_pages,
        "output_inside_failures": output_inside_failures,
        "output_outside_failures": output_outside_failures,
        "output_top_bottom_failures": output_top_bottom_failures,
        "text_mismatch_pages": text_mismatch_pages,
        "unembedded_font_count": len(unembedded_fonts),
        "unembedded_fonts": sorted(unembedded_fonts),
        "low_ppi_image_count": len(low_ppi_images),
        "low_ppi_images": low_ppi_images[:250],
        "transparency_object_count": transparency_objects,
        "proof_pages": proof_pages,
        "provider_preview_status": "pending_external_kdp_print_previewer",
        "physical_proof_status": "pending_future_checkpoint",
    }
    hard_failures = {
        "searchable_pages": searchable != 537,
        "inside": output_inside_failures != 0,
        "outside": output_outside_failures != 0,
        "top_bottom": output_top_bottom_failures != 0,
        "text_mismatch": text_mismatch_pages != 0,
        "unembedded_fonts": bool(unembedded_fonts),
    }
    if any(hard_failures.values()):
        raise RuntimeError({"print_interior_hard_gate_failed": hard_failures, "summary": summary})
    csv_write(qa_root / "MRHPD v3.0.0a Response 76 Print Interior Page Transform Register.csv", page_rows)
    json_write(qa_root / "MRHPD v3.0.0a Response 76 Print Interior Page Transform Register.json", page_rows)
    json_write(qa_root / "MRHPD v3.0.0a Response 76 Print Interior QA.json", summary)
    return summary, page_rows, proof_paths


def _fit_panel(image: Image.Image, width: int, height: int) -> Image.Image:
    source = image.convert("RGB")
    if source.size == (width, height):
        return source
    target_ratio = width / height
    source_ratio = source.width / source.height
    if abs(source_ratio - target_ratio) < 0.002:
        return source.resize((width, height), Image.Resampling.LANCZOS)
    if source_ratio < target_ratio:
        crop_height = int(round(source.width / target_ratio))
        top = max(0, (source.height - crop_height) // 2)
        source = source.crop((0, top, source.width, top + crop_height))
    else:
        crop_width = int(round(source.height * target_ratio))
        left = max(0, (source.width - crop_width) // 2)
        source = source.crop((left, 0, left + crop_width, source.height))
    return source.resize((width, height), Image.Resampling.LANCZOS)


def _average_edge(image: Image.Image, side: str) -> tuple[int, int, int]:
    if side == "left":
        crop = image.crop((0, 0, min(8, image.width), image.height))
    else:
        crop = image.crop((max(0, image.width - 8), 0, image.width, image.height))
    values = ImageStat.Stat(crop).mean[:3]
    return tuple(int(round(v)) for v in values)


def _edge_extend(trim: Image.Image, left_bleed: int, right_bleed: int, top_bleed: int, bottom_bleed: int) -> Image.Image:
    canvas = Image.new("RGB", (left_bleed + trim.width + right_bleed, top_bleed + trim.height + bottom_bleed), "white")
    canvas.paste(trim, (left_bleed, top_bleed))
    top = trim.crop((0, 0, trim.width, 1)).resize((trim.width, top_bleed))
    bottom = trim.crop((0, trim.height - 1, trim.width, trim.height)).resize((trim.width, bottom_bleed))
    left = trim.crop((0, 0, 1, trim.height)).resize((left_bleed, trim.height))
    right = trim.crop((trim.width - 1, 0, trim.width, trim.height)).resize((right_bleed, trim.height))
    canvas.paste(top, (left_bleed, 0))
    canvas.paste(bottom, (left_bleed, top_bleed + trim.height))
    canvas.paste(left, (0, top_bleed))
    canvas.paste(right, (left_bleed + trim.width, top_bleed))
    corners = {
        (0, 0): trim.getpixel((0, 0)),
        (left_bleed + trim.width, 0): trim.getpixel((trim.width - 1, 0)),
        (0, top_bleed + trim.height): trim.getpixel((0, trim.height - 1)),
        (left_bleed + trim.width, top_bleed + trim.height): trim.getpixel((trim.width - 1, trim.height - 1)),
    }
    for (x, y), color in corners.items():
        w = left_bleed if x == 0 else right_bleed
        h = top_bleed if y == 0 else bottom_bleed
        ImageDraw.Draw(canvas).rectangle((x, y, x + w, y + h), fill=color)
    return canvas


def create_single_image_pdf(image_path: Path, pdf_path: Path, width_in: float, height_in: float) -> None:
    doc = fitz.open()
    page = doc.new_page(width=width_in * 72.0, height=height_in * 72.0)
    page.insert_image(page.rect, filename=str(image_path), keep_proportion=False)
    doc.set_metadata({"title": image_path.stem, "author": "Brent McAnulty, M.D.", "subject": "Human Pathogen Database print-production cover"})
    doc.save(pdf_path, garbage=4, deflate=True, clean=True)
    doc.close()


def _template_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            pass
    return ImageFont.load_default()


def generate_cover(project: Path, selection: dict[str, Any], qa_root: Path) -> tuple[dict[str, Any], list[Path]]:
    front = Image.open(project / "Cover/MRHPD-COVER-0001 Front Cover v3.0.0a Brent McAnulty MD.png")
    back = Image.open(project / "Cover/MRHPD-COVER-0002 Back Cover v3.0.0a Brent McAnulty MD.png")
    spine = Image.open(project / "Cover/MRHPD-COVER-0003 Spine v3.0.0a Brent McAnulty MD.png")
    panel_w, trim_h = 2550, 3300
    spine_w = int(round(selection["spine_width_in"] * 300))
    left_bleed, right_bleed, top_bleed, bottom_bleed = 38, 37, 38, 37
    expected_width = selection["cover_width_px_300dpi"]
    expected_height = selection["cover_height_px_300dpi"]
    if panel_w * 2 + spine_w + left_bleed + right_bleed != expected_width or trim_h + top_bleed + bottom_bleed != expected_height:
        raise RuntimeError({"cover_pixel_arithmetic": {"spine_w": spine_w, "expected": [expected_width, expected_height]}})
    front_panel = _fit_panel(front, panel_w, trim_h)
    back_panel = _fit_panel(back, panel_w, trim_h)
    spine_height_panel = _fit_panel(spine, max(1, spine.width), trim_h)
    if spine_height_panel.width > spine_w:
        scale = min(1.0, spine_w / spine_height_panel.width)
        spine_height_panel = spine_height_panel.resize((int(round(spine_height_panel.width * scale)), int(round(spine_height_panel.height * scale))), Image.Resampling.LANCZOS)
        if spine_height_panel.height < trim_h:
            padded = Image.new("RGB", (spine_height_panel.width, trim_h), _average_edge(spine_height_panel, "left"))
            padded.paste(spine_height_panel, (0, (trim_h - spine_height_panel.height) // 2))
            spine_height_panel = padded
    back_edge = _average_edge(back_panel, "right")
    front_edge = _average_edge(front_panel, "left")
    spine_panel = Image.new("RGB", (spine_w, trim_h), back_edge)
    spine_draw = ImageDraw.Draw(spine_panel)
    for x in range(spine_w):
        ratio = x / max(1, spine_w - 1)
        color = tuple(int(round(back_edge[c] * (1.0 - ratio) + front_edge[c] * ratio)) for c in range(3))
        spine_draw.line((x, 0, x, trim_h), fill=color)
    spine_x = (spine_w - spine_height_panel.width) // 2
    spine_panel.paste(spine_height_panel, (spine_x, 0))
    trim = Image.new("RGB", (panel_w * 2 + spine_w, trim_h), "white")
    trim.paste(back_panel, (0, 0))
    trim.paste(spine_panel, (panel_w, 0))
    trim.paste(front_panel, (panel_w + spine_w, 0))
    barcode_w, barcode_h = 600, 360
    barcode_right = panel_w - 75
    barcode_bottom = trim_h - 75
    barcode_box = (barcode_right - barcode_w, barcode_bottom - barcode_h, barcode_right, barcode_bottom)
    ImageDraw.Draw(trim).rectangle(barcode_box, fill="white")
    canvas = _edge_extend(trim, left_bleed, right_bleed, top_bleed, bottom_bleed)
    final_png = project / FINAL_COVER_PNG_REL
    final_tiff = project / FINAL_COVER_TIFF_REL
    final_pdf = project / FINAL_COVER_PDF_REL
    template_png = project / TEMPLATE_PNG_REL
    template_pdf = project / TEMPLATE_PDF_REL
    proof_png = project / PROOF_PNG_REL
    for path in (final_png, final_tiff, final_pdf, template_png, template_pdf, proof_png):
        path.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(final_png, format="PNG", dpi=(300, 300), optimize=True)
    canvas.save(final_tiff, format="TIFF", dpi=(300, 300), compression="tiff_lzw")
    create_single_image_pdf(final_png, final_pdf, selection["cover_width_in"], selection["cover_height_in"])

    template = canvas.copy().convert("RGB")
    draw = ImageDraw.Draw(template)
    x_back_trim_start = left_bleed
    x_back_trim_end = left_bleed + panel_w
    x_spine_start = x_back_trim_end
    x_spine_end = x_spine_start + spine_w
    x_front_trim_end = x_spine_end + panel_w
    y_trim_start = top_bleed
    y_trim_end = top_bleed + trim_h
    line_width = 6
    for x in (x_back_trim_start, x_back_trim_end, x_spine_start, x_spine_end, x_front_trim_end):
        draw.line((x, 0, x, expected_height), fill=(197, 0, 0), width=line_width)
    for y in (y_trim_start, y_trim_end):
        draw.line((0, y, expected_width, y), fill=(197, 0, 0), width=line_width)
    safe = 75
    draw.rectangle((x_back_trim_start + safe, y_trim_start + safe, x_back_trim_end - safe, y_trim_end - safe), outline=(0, 128, 0), width=6)
    draw.rectangle((x_spine_end + safe, y_trim_start + safe, x_front_trim_end - safe, y_trim_end - safe), outline=(0, 128, 0), width=6)
    spine_safe = int(math.ceil(0.0625 * 300))
    draw.rectangle((x_spine_start + spine_safe, y_trim_start + safe, x_spine_end - spine_safe, y_trim_end - safe), outline=(201, 153, 0), width=6)
    bx0, by0, bx1, by1 = barcode_box
    barcode_canvas = (bx0 + left_bleed, by0 + top_bleed, bx1 + left_bleed, by1 + top_bleed)
    draw.rectangle(barcode_canvas, outline=(128, 0, 128), width=8)
    label_font = _template_font(42, True)
    small_font = _template_font(30, True)
    labels = [
        ("BACK COVER", (x_back_trim_start + 90, y_trim_start + 70), (0, 128, 0)),
        ("SPINE", (x_spine_start + 30, y_trim_start + 70), (201, 153, 0)),
        ("FRONT COVER", (x_spine_end + 90, y_trim_start + 70), (0, 128, 0)),
        ("2 × 1.2 IN BARCODE CLEAR AREA", (barcode_canvas[0] + 20, barcode_canvas[1] + 20), (128, 0, 128)),
        ("RED = TRIM/FOLD   GREEN = 0.25-IN LIVE SAFE AREA   GOLD = SPINE-TEXT SAFE AREA", (120, 85), (197, 0, 0)),
    ]
    for text, position, color in labels:
        draw.text(position, text, font=small_font if len(text) > 25 else label_font, fill=color)
    template.save(template_png, format="PNG", dpi=(300, 300), optimize=True)
    create_single_image_pdf(template_png, template_pdf, selection["cover_width_in"], selection["cover_height_in"])
    proof = template.resize((2400, int(round(2400 * template.height / template.width))), Image.Resampling.LANCZOS)
    proof.save(proof_png, format="PNG", dpi=(300, 300), optimize=True)

    with Image.open(final_png) as check:
        pixel_width, pixel_height = check.size
        alpha_present = check.mode in {"RGBA", "LA", "PA"}
        color_space = check.mode
        barcode_crop = check.crop(barcode_canvas).convert("RGB")
        barcode_extrema = barcode_crop.getextrema()
        barcode_white = all(channel == (255, 255) for channel in barcode_extrema)
        icc_present = bool(check.info.get("icc_profile"))
    cover_pdf = fitz.open(final_pdf)
    pdf_page = cover_pdf[0]
    pdf_width_in = pdf_page.rect.width / 72.0
    pdf_height_in = pdf_page.rect.height / 72.0
    cover_pdf.close()
    template_pdf_doc = fitz.open(template_pdf)
    template_width_in = template_pdf_doc[0].rect.width / 72.0
    template_height_in = template_pdf_doc[0].rect.height / 72.0
    template_pdf_doc.close()
    spine_text_margin_px = spine_x
    summary = {
        "status": "passed",
        "pixel_width": pixel_width,
        "pixel_height": pixel_height,
        "expected_pixel_width": expected_width,
        "expected_pixel_height": expected_height,
        "color_space": color_space,
        "alpha_present": alpha_present,
        "icc_profile_present": icc_present,
        "pdf_width_in": round(pdf_width_in, 6),
        "pdf_height_in": round(pdf_height_in, 6),
        "template_pdf_width_in": round(template_width_in, 6),
        "template_pdf_height_in": round(template_height_in, 6),
        "barcode_policy": "KDP-applied barcode; reserved white 2 × 1.2-inch area on lower-right back cover",
        "barcode_width_in": 2.0,
        "barcode_height_in": 1.2,
        "barcode_pixel_extrema": barcode_extrema,
        "barcode_white": barcode_white,
        "spine_text_safety": f"Original spine artwork centered with {spine_text_margin_px} px ({spine_text_margin_px / 300:.4f} in) horizontal clearance to each new fold; KDP minimum is 0.0625 in.",
        "front_back_scaling": "No nonuniform stretch; each separate panel is aspect-preserving fitted to the 8.5 × 11-inch trim panel after removing legacy outer bleed where present.",
        "legacy_combined_wrap_preserved": True,
        "legacy_combined_wrap_hashes": {rel: digest for rel, digest in COVER_HASHES.items() if "Combined Cover Wrap" in rel},
        "final_png": {"name": final_png.name, "path": final_png.relative_to(project).as_posix(), "bytes": final_png.stat().st_size, "sha256": sha256_file(final_png)},
        "final_tiff": {"name": final_tiff.name, "path": final_tiff.relative_to(project).as_posix(), "bytes": final_tiff.stat().st_size, "sha256": sha256_file(final_tiff)},
        "final_pdf": {"name": final_pdf.name, "path": final_pdf.relative_to(project).as_posix(), "bytes": final_pdf.stat().st_size, "sha256": sha256_file(final_pdf)},
        "template_png": {"name": template_png.name, "path": template_png.relative_to(project).as_posix(), "bytes": template_png.stat().st_size, "sha256": sha256_file(template_png)},
        "template_pdf": {"name": template_pdf.name, "path": template_pdf.relative_to(project).as_posix(), "bytes": template_pdf.stat().st_size, "sha256": sha256_file(template_pdf)},
        "proof_png": {"name": proof_png.name, "path": proof_png.relative_to(project).as_posix(), "bytes": proof_png.stat().st_size, "sha256": sha256_file(proof_png)},
    }
    failures = {
        "dimensions": (pixel_width, pixel_height) != (expected_width, expected_height),
        "alpha": alpha_present,
        "barcode": not barcode_white,
        "pdf_dimensions": abs(pdf_width_in - selection["cover_width_in"]) > 0.002 or abs(pdf_height_in - selection["cover_height_in"]) > 0.002,
        "template_dimensions": abs(template_width_in - selection["cover_width_in"]) > 0.002 or abs(template_height_in - selection["cover_height_in"]) > 0.002,
        "spine_text": spine_text_margin_px < spine_safe,
    }
    if any(failures.values()):
        raise RuntimeError({"cover_hard_gate_failed": failures, "summary": summary})
    json_write(qa_root / "MRHPD v3.0.0a Response 76 Cover QA.json", summary)
    return summary, [final_png, final_tiff, final_pdf, template_png, template_pdf, proof_png]


def preflight_records(selection: dict[str, Any], interior: dict[str, Any], cover: dict[str, Any], now_iso: str) -> list[dict[str, Any]]:
    rows = [
        ("scenario_locked", "KDP premium color, white paper, paperback, matte, 8.5 × 11", selection["selection_code"], "passed"),
        ("page_count", "538", str(interior["output_page_count"]), "passed" if interior["output_page_count"] == 538 else "failed"),
        ("searchability", "537 searchable source pages", str(interior["searchable_pages"]), "passed" if interior["searchable_pages"] == 537 else "failed"),
        ("intentional_blank", "page 538 blank", str(interior["intentional_blank_page"]), "passed" if interior["intentional_blank_page"] == 538 else "failed"),
        ("text_equivalence", "0 mismatch pages", str(interior["text_mismatch_pages"]), "passed" if interior["text_mismatch_pages"] == 0 else "failed"),
        ("inside_gutter", "0 failures at 0.75 in", str(interior["output_inside_failures"]), "passed" if interior["output_inside_failures"] == 0 else "failed"),
        ("outside_margin", "0 failures at 0.25 in", str(interior["output_outside_failures"]), "passed" if interior["output_outside_failures"] == 0 else "failed"),
        ("top_bottom_margin", "0 failures at 0.25 in", str(interior["output_top_bottom_failures"]), "passed" if interior["output_top_bottom_failures"] == 0 else "failed"),
        ("font_embedding", "0 unembedded fonts", str(interior["unembedded_font_count"]), "passed" if interior["unembedded_font_count"] == 0 else "failed"),
        ("raster_resolution", "300 ppi target", f"{interior['low_ppi_image_count']} placements below 299 ppi", "passed" if interior["low_ppi_image_count"] == 0 else "controlled_warning"),
        ("cover_pixels", f"{selection['cover_width_px_300dpi']} × {selection['cover_height_px_300dpi']}", f"{cover['pixel_width']} × {cover['pixel_height']}", "passed"),
        ("cover_page_size", f"{selection['cover_width_in']:.6f} × {selection['cover_height_in']:.3f} in", f"{cover['pdf_width_in']:.6f} × {cover['pdf_height_in']:.3f} in", "passed"),
        ("cover_flattening", "RGB without alpha", f"{cover['color_space']}; alpha={cover['alpha_present']}", "passed" if not cover["alpha_present"] else "failed"),
        ("barcode_clear_area", "2 × 1.2 in white", str(cover["barcode_white"]), "passed" if cover["barcode_white"] else "failed"),
        ("spine_text_safety", ">= 0.0625 in from folds", cover["spine_text_safety"], "passed"),
        ("legacy_wrap", "preserved; not stretched", str(cover["legacy_combined_wrap_preserved"]), "passed"),
        ("provider_previewer", "KDP Print Previewer approval", "pending external provider review", "controlled_pending"),
        ("physical_proof", "physical proof approval", "pending later Section 5 checkpoint", "controlled_pending"),
    ]
    return [{"gate": gate, "expected": expected, "observed": observed, "status": status, "checked_at": now_iso} for gate, expected, observed, status in rows]


def recovery_events(now_iso: str) -> list[dict[str, Any]]:
    return [
        {
            "event_code": "V3-CP5-S1-REC-181-LOCAL-RUNTIME-UNAVAILABLE",
            "condition": "The local container and both Python execution surfaces returned InvalidArgumentError before code startup.",
            "recovery": "Preserved the verified Google Drive checkpoint, used the existing isolated transient runner only for deterministic computation, and retained Google Drive as controlling storage and user delivery.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_code": "V3-CP5-S1-REC-182-CHECKPOINT1-AUTOMATED-RESTORE",
            "condition": "Checkpoint 2 required the exact Response 75 project state.",
            "recovery": "Retrieved the Response 75 recovery package without user involvement, verified its exact nested identity, clean-applied it to the exact Response 72 restore, and resumed from the newest verified state.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_code": "V3-CP5-S1-REC-183-PREMIUM-COLOR-SCENARIO-LOCKED",
            "condition": "Checkpoint 1 retained several unselected print scenarios and an obsolete 0.750-inch provisional spine.",
            "recovery": "Selected KDP premium color, white paper, paperback, matte, 8.5 × 11 inches, 538 pages, and recalculated the 1.262686-inch spine and exact 18.512686 × 11.250-inch cover canvas.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_code": "V3-CP5-S1-REC-184-PRINT-INTERIOR-DERIVED",
            "condition": "The immutable digital publication contained 537 pages and could not serve directly as an even-page print manuscript.",
            "recovery": "Created a separate 538-page print-only derivative with an explicit final blank page while preserving the immutable digital PDF byte-for-byte.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_code": "V3-CP5-S1-REC-185-GUTTER-NORMALIZATION",
            "condition": "The conservative Checkpoint 1 text-block screen identified pages below the selected 0.75-inch inside-gutter target.",
            "recovery": "Applied deterministic page-specific rotation, translation, and only-when-required proportional reduction; compared source and output words page by page and required zero final text-safe failures.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_code": "V3-CP5-S1-REC-186-EXACT-COVER-REGENERATION",
            "condition": "The legacy 18.000-inch combined wrap was too narrow for the selected 538-page production scenario.",
            "recovery": "Regenerated the cover on the exact calculated canvas from the separate front, back, and spine masters without nonuniform stretching; preserved the legacy wrap as historical evidence.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_code": "V3-CP5-S1-REC-187-BARCODE-FOLD-SPINE-SAFETY",
            "condition": "The final cover required explicit fold, live-area, spine-text, and barcode controls.",
            "recovery": "Generated a project-controlled exact template, centered the existing spine artwork inside the wider spine, preserved fold clearance, and reserved a clean 2 × 1.2-inch KDP barcode area.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_code": "V3-CP5-S1-REC-188-PREFLIGHT-AND-CLEAN-APPLY",
            "condition": "The print candidate required synchronized database, workbook, application, tracking, index, manifest, archive, and recovery verification.",
            "recovery": "Rebuilt all governed derivative surfaces, clean-applied the cumulative Response 76 recovery to the exact Response 72 restore, and retained provider preview and physical proof as explicit controlled external gates.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
    ]


def _clone_response_row(con: sqlite3.Connection, source_key: str, response: dict[str, Any]) -> None:
    info = con.execute("PRAGMA table_info(thread_response_reconciliation_cp3)").fetchall()
    columns = [row[1] for row in info]
    pk_columns = {row[1] for row in info if row[5]}
    source = con.execute("SELECT * FROM thread_response_reconciliation_cp3 WHERE response_key=?", (source_key,)).fetchone()
    if source is None:
        raise RuntimeError(f"source response missing: {source_key}")
    values = dict(zip(columns, source))
    values.update(response)
    insert_columns = [column for column in columns if column not in pk_columns]
    con.execute("DELETE FROM thread_response_reconciliation_cp3 WHERE response_key=?", (response["response_key"],))
    con.execute(
        f"INSERT INTO thread_response_reconciliation_cp3 ({','.join(insert_columns)}) VALUES ({','.join('?' for _ in insert_columns)})",
        [values.get(column) for column in insert_columns],
    )


def synchronize_database(
    cp1_db: Path,
    destination: Path,
    *,
    now_iso: str,
    selection: dict[str, Any],
    interior: dict[str, Any],
    page_rows: list[dict[str, Any]],
    cover: dict[str, Any],
    preflight: list[dict[str, Any]],
    events: list[dict[str, Any]],
) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(cp1_db, destination)
    con = sqlite3.connect(destination)
    try:
        con.execute("PRAGMA foreign_keys=ON")
        con.execute("BEGIN IMMEDIATE")
        response = {
            "response_key": "R76",
            "response_number": 76,
            "response_label": "76",
            "response_date": now_iso,
            "major_topic": "Human Pathogen Database remediation",
            "title": "Section 5 premium-color print-interior and cover-production checkpoint",
            "goal": "Continue from the verified Response 75 state, lock the initial print scenario, generate and preflight the print interior and full cover, synchronize all project surfaces, and emit cumulative recovery.",
            "raw_prompt": "Continue",
            "raw_response": "[PRE-EMISSION RESPONSE; final user-visible response is represented by the source-supported summary]",
            "summary": "Locked the KDP premium-color production master, generated and verified the 538-page print-only interior and exact full cover, synchronized database/workbook/application/tracking/index/manifest controls, and emitted cumulative recovery through Response 76.",
            "state": "checkpoint_complete_continue_required",
            "coverage": "exact raw prompt plus source-supported response summary",
            "fidelity_classification": "source_verified_prompt_and_summary",
            "source_id": "CURRENT-CONVERSATION-R76",
            "source_path": "Current conversation, Response 75 recovery, official KDP production requirements, and Response 76 output package",
            "notes": "Checkpoint 2 of 3 complete. Checkpoint 3 will freeze Session 1 and emit the complete self-contained restore.",
            "reconciled_at": now_iso,
        }
        _clone_response_row(con, "R75", response)
        con.executescript("""
        CREATE TABLE IF NOT EXISTS section5_print_selection (
            section5_print_selection_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            selection_code TEXT NOT NULL,
            provider TEXT NOT NULL,
            binding TEXT NOT NULL,
            interior_type TEXT NOT NULL,
            paper TEXT NOT NULL,
            cover_finish TEXT NOT NULL,
            trim_width_in REAL NOT NULL,
            trim_height_in REAL NOT NULL,
            production_page_count INTEGER NOT NULL,
            spine_width_in REAL NOT NULL,
            cover_width_in REAL NOT NULL,
            cover_height_in REAL NOT NULL,
            status TEXT NOT NULL,
            rationale TEXT NOT NULL,
            recorded_at TEXT NOT NULL,
            UNIQUE(checkpoint_code, selection_code)
        );
        CREATE TABLE IF NOT EXISTS section5_print_derivative (
            section5_print_derivative_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            derivative_type TEXT NOT NULL,
            relative_path TEXT NOT NULL,
            bytes INTEGER NOT NULL,
            sha256 TEXT NOT NULL,
            page_count INTEGER,
            pixel_width INTEGER,
            pixel_height INTEGER,
            status TEXT NOT NULL,
            recorded_at TEXT NOT NULL,
            UNIQUE(checkpoint_code, derivative_type, relative_path)
        );
        CREATE TABLE IF NOT EXISTS section5_page_transform (
            section5_page_transform_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            page_number INTEGER NOT NULL,
            rotation_degrees INTEGER NOT NULL,
            scale REAL,
            transform_class TEXT NOT NULL,
            inside_margin_in REAL,
            outside_margin_in REAL,
            top_margin_in REAL,
            bottom_margin_in REAL,
            searchable INTEGER NOT NULL,
            text_equal INTEGER NOT NULL,
            status TEXT NOT NULL,
            UNIQUE(checkpoint_code, page_number)
        );
        CREATE TABLE IF NOT EXISTS section5_cover_template (
            section5_cover_template_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            template_type TEXT NOT NULL,
            relative_path TEXT NOT NULL,
            width_in REAL NOT NULL,
            height_in REAL NOT NULL,
            pixel_width INTEGER,
            pixel_height INTEGER,
            sha256 TEXT NOT NULL,
            status TEXT NOT NULL,
            recorded_at TEXT NOT NULL,
            UNIQUE(checkpoint_code, template_type)
        );
        CREATE TABLE IF NOT EXISTS section5_print_preflight (
            section5_print_preflight_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            gate_key TEXT NOT NULL,
            expected TEXT NOT NULL,
            observed TEXT NOT NULL,
            status TEXT NOT NULL,
            checked_at TEXT NOT NULL,
            UNIQUE(checkpoint_code, gate_key)
        );
        CREATE TABLE IF NOT EXISTS section5_master_category_extension (
            section5_master_category_extension_id INTEGER PRIMARY KEY,
            category_key TEXT NOT NULL UNIQUE,
            parent_key TEXT,
            label TEXT NOT NULL,
            scope TEXT NOT NULL,
            description TEXT NOT NULL,
            status TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        """)
        con.execute("DELETE FROM section5_session1_checkpoint WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        con.execute(
            """INSERT INTO section5_session1_checkpoint (
            checkpoint_code,response_number,section_label,session_label,checkpoint_label,state,
            baseline_restore_sha256,baseline_project_sha256,publication_sha256,digital_page_count,
            production_page_count,provider_selection_status,print_interior_status,cover_status,
            workbook_status,application_status,next_checkpoint,recorded_at
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (
                CHECKPOINT_CODE, 76, SECTION_LABEL, SESSION_LABEL, CHECKPOINT_LABEL, "checkpoint_complete",
                BASE_RESTORE_SHA256, BASE_PROJECT_SHA256, PUBLICATION_SHA256, 537, 538,
                "locked_initial_production_master", "passed_print_candidate", "passed_regenerated_exact_canvas",
                "pending_final_save", "pending_final_audit", "Checkpoint 3 of 3 - Session 1 freeze and complete restore", now_iso,
            ),
        )
        con.execute("DELETE FROM section5_print_selection WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        con.execute(
            """INSERT INTO section5_print_selection (
            checkpoint_code,selection_code,provider,binding,interior_type,paper,cover_finish,trim_width_in,
            trim_height_in,production_page_count,spine_width_in,cover_width_in,cover_height_in,status,rationale,recorded_at
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (
                CHECKPOINT_CODE, selection["selection_code"], selection["provider"], selection["binding"],
                selection["interior_type"], selection["paper"], selection["cover_finish"], selection["trim_width_in"],
                selection["trim_height_in"], selection["production_page_count"], selection["spine_width_in"],
                selection["cover_width_in"], selection["cover_height_in"], selection["selection_status"],
                selection["selection_rationale"], now_iso,
            ),
        )
        con.execute("DELETE FROM section5_print_derivative WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        derivative_rows = [
            ("print_interior", PRINT_INTERIOR_REL, interior["output_bytes"], interior["output_sha256"], 538, None, None, "passed"),
            ("cover_png", cover["final_png"]["path"], cover["final_png"]["bytes"], cover["final_png"]["sha256"], 1, cover["pixel_width"], cover["pixel_height"], "passed"),
            ("cover_tiff", cover["final_tiff"]["path"], cover["final_tiff"]["bytes"], cover["final_tiff"]["sha256"], 1, cover["pixel_width"], cover["pixel_height"], "passed"),
            ("cover_pdf", cover["final_pdf"]["path"], cover["final_pdf"]["bytes"], cover["final_pdf"]["sha256"], 1, None, None, "passed"),
            ("cover_template_png", cover["template_png"]["path"], cover["template_png"]["bytes"], cover["template_png"]["sha256"], 1, cover["pixel_width"], cover["pixel_height"], "passed"),
            ("cover_template_pdf", cover["template_pdf"]["path"], cover["template_pdf"]["bytes"], cover["template_pdf"]["sha256"], 1, None, None, "passed"),
        ]
        con.executemany(
            "INSERT INTO section5_print_derivative (checkpoint_code,derivative_type,relative_path,bytes,sha256,page_count,pixel_width,pixel_height,status,recorded_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
            [(CHECKPOINT_CODE, *row, now_iso) for row in derivative_rows],
        )
        con.execute("DELETE FROM section5_page_transform WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        con.executemany(
            """INSERT INTO section5_page_transform (
            checkpoint_code,page_number,rotation_degrees,scale,transform_class,inside_margin_in,outside_margin_in,
            top_margin_in,bottom_margin_in,searchable,text_equal,status
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?)""",
            [(
                CHECKPOINT_CODE, row["page"], row["rotation_degrees"], row["scale"], row["transform_class"],
                row["inside_text_margin_in"], row["outside_text_margin_in"], row["top_text_margin_in"], row["bottom_text_margin_in"],
                int(row["searchable"]), int(row["text_equal_to_source"]),
                "passed" if row["inside_pass"] and row["outside_pass"] and row["top_bottom_pass"] and row["text_equal_to_source"] else "failed",
            ) for row in page_rows],
        )
        con.execute("DELETE FROM section5_cover_template WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        template_rows = [
            ("final_cover_png", cover["final_png"]["path"], cover["final_png"]["sha256"], cover["pixel_width"], cover["pixel_height"], "passed"),
            ("final_cover_pdf", cover["final_pdf"]["path"], cover["final_pdf"]["sha256"], None, None, "passed"),
            ("exact_template_png", cover["template_png"]["path"], cover["template_png"]["sha256"], cover["pixel_width"], cover["pixel_height"], "passed"),
            ("exact_template_pdf", cover["template_pdf"]["path"], cover["template_pdf"]["sha256"], None, None, "passed"),
        ]
        con.executemany(
            "INSERT INTO section5_cover_template (checkpoint_code,template_type,relative_path,width_in,height_in,pixel_width,pixel_height,sha256,status,recorded_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
            [(CHECKPOINT_CODE, kind, path, selection["cover_width_in"], selection["cover_height_in"], pxw, pxh, digest, status, now_iso) for kind, path, digest, pxw, pxh, status in template_rows],
        )
        con.execute("DELETE FROM section5_print_preflight WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        con.executemany(
            "INSERT INTO section5_print_preflight (checkpoint_code,gate_key,expected,observed,status,checked_at) VALUES (?,?,?,?,?,?)",
            [(CHECKPOINT_CODE, row["gate"], row["expected"], row["observed"], row["status"], row["checked_at"]) for row in preflight],
        )
        for event in events:
            con.execute(
                "INSERT OR REPLACE INTO section5_recovery_event (checkpoint_code,event_code,condition,recovery,status,recorded_at) VALUES (?,?,?,?,?,?)",
                (CHECKPOINT_CODE, event["event_code"], event["condition"], event["recovery"], event["status"], event["recorded_at"]),
            )
        categories = [
            ("MRHPD-CAT-PRINT-PRODUCTION", None, "Print production", "project-wide", "Provider, material, template, proof, and press-production governance.", "current"),
            ("MRHPD-CAT-PRINT-INTERIOR", "MRHPD-CAT-PRINT-PRODUCTION", "Print interior", "publication", "Even-page manuscript, margin, font, image, and page-render controls.", "current"),
            ("MRHPD-CAT-COVER-PRODUCTION", "MRHPD-CAT-PRINT-PRODUCTION", "Cover production", "publication", "Full-wrap geometry, bleed, fold, spine, live-area, barcode, and raster controls.", "current"),
            ("MRHPD-CAT-PROVIDER-PREVIEW", "MRHPD-CAT-PRINT-PRODUCTION", "Provider preview", "external gate", "KDP Print Previewer or equivalent provider conversion and review evidence.", "current"),
            ("MRHPD-CAT-PHYSICAL-PROOF", "MRHPD-CAT-PRINT-PRODUCTION", "Physical proof", "external gate", "Physical copy review, defect logging, correction, and approval evidence.", "current"),
        ]
        for key, parent, label, scope, description, status in categories:
            con.execute(
                "INSERT OR REPLACE INTO section5_master_category_extension (category_key,parent_key,label,scope,description,status,recorded_at) VALUES (?,?,?,?,?,?,?)",
                (key, parent, label, scope, description, status, now_iso),
            )
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or fk:
            raise RuntimeError({"database_integrity": integrity, "foreign_keys": fk[:20]})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()
    con = sqlite3.connect(destination)
    try:
        table_count = con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0]
        response_count = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R76'").fetchone()[0]
        page_count = con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code=?", (CHECKPOINT_CODE,)).fetchone()[0]
        failed_pages = con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code=? AND status!='passed'", (CHECKPOINT_CODE,)).fetchone()[0]
        selected = con.execute("SELECT status FROM section5_print_selection WHERE checkpoint_code=?", (CHECKPOINT_CODE,)).fetchone()
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk_count = len(list(con.execute("PRAGMA foreign_key_check")))
    finally:
        con.close()
    if response_count != 1 or page_count != 538 or failed_pages or selected != ("locked_initial_production_master",):
        raise RuntimeError({"database_current_gate": {"response": response_count, "pages": page_count, "failed": failed_pages, "selected": selected}})
    return {
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "table_count": table_count,
        "integrity": integrity,
        "foreign_key_violations": fk_count,
        "response76_records": response_count,
        "page_transform_records": page_count,
        "failed_page_transforms": failed_pages,
        "selection_status": selected[0],
        "checkpoint_state": "checkpoint_complete",
    }


def _write_sheet(ws: Any, rows: list[dict[str, Any]]) -> None:
    if not rows:
        ws.append(["No records"])
        return
    headers = list(rows[0])
    ws.append(headers)
    for cell in ws[1]:
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(color=WHITE, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin = Side(style="thin", color="AAB8C0")
    for row in rows:
        ws.append([row.get(header) for header in headers])
    for r in range(2, ws.max_row + 1):
        for cell in ws[r]:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if r % 2 == 0:
                cell.fill = PatternFill("solid", fgColor=PALE_BLUE)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    ws.sheet_view.showGridLines = False
    for idx, header in enumerate(headers, start=1):
        sample = [str(header)] + [str(ws.cell(r, idx).value or "") for r in range(2, min(ws.max_row, 120) + 1)]
        ws.column_dimensions[get_column_letter(idx)].width = min(55, max(10, max(len(value) for value in sample) + 2))


def augment_workbook(
    source: Path,
    destination: Path,
    *,
    selection: dict[str, Any],
    interior: dict[str, Any],
    page_rows: list[dict[str, Any]],
    cover: dict[str, Any],
    preflight: list[dict[str, Any]],
    events: list[dict[str, Any]],
    categories: list[dict[str, Any]],
) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    wb = load_workbook(source)
    inherited = list(wb.sheetnames)
    datasets = {
        "S5S1 CP2 Dashboard": [
            {"Control": "Response", "Value": 76, "Status": "current"},
            {"Control": "Checkpoint", "Value": "2 of 3", "Status": "complete"},
            {"Control": "Scenario", "Value": selection["selection_code"], "Status": selection["selection_status"]},
            {"Control": "Print interior", "Value": interior["output_name"], "Status": interior["status"]},
            {"Control": "Cover", "Value": cover["final_png"]["name"], "Status": cover["status"]},
            {"Control": "Next", "Value": "Checkpoint 3 of 3 - Session 1 freeze and complete restore", "Status": "continue"},
        ],
        "S5S1 CP2 Selection": [selection],
        "S5S1 CP2 Interior": [interior],
        "S5S1 CP2 Transforms": page_rows,
        "S5S1 CP2 Cover": [{k: v if not isinstance(v, dict) else json.dumps(v, ensure_ascii=False) for k, v in cover.items()}],
        "S5S1 CP2 Preflight": preflight,
        "S5S1 CP2 Response": [{"Response": 76, "Raw Prompt": "Continue", "Summary": "KDP premium-color scenario locked; print interior and full cover generated and preflighted; checkpoint recovery emitted.", "State": "checkpoint_complete_continue_required"}],
        "S5S1 CP2 Recovery": events,
        "S5S1 CP2 Categories": categories,
    }
    for title, rows in datasets.items():
        if title in wb.sheetnames:
            del wb[title]
        ws = wb.create_sheet(title=title)
        _write_sheet(ws, rows)
    wb.properties.title = "MRHPD v3.0.0a Comprehensive Tracking Through Response 76"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(destination)
    with zipfile.ZipFile(destination) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("comprehensive workbook CRC failed")
    check = load_workbook(destination, read_only=True, data_only=False)
    try:
        sheet_names = list(check.sheetnames)
        formula_errors = []
        formula_count = 0
        for ws in check.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    value = cell.value
                    if isinstance(value, str) and value.startswith("="):
                        formula_count += 1
                        if any(token in value for token in ("#REF!", "#NAME?", "#VALUE!", "#DIV/0!", "#N/A")):
                            formula_errors.append(f"{ws.title}!{cell.coordinate}:{value}")
    finally:
        check.close()
    lost = sorted(set(inherited) - set(sheet_names))
    if lost or len(sheet_names) < 100 or formula_errors:
        raise RuntimeError({"workbook_gate": {"lost": lost, "sheets": len(sheet_names), "formula_errors": formula_errors[:30]}})
    return {
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "source_sheet_count": len(inherited),
        "current_sheet_count": len(sheet_names),
        "new_sheet_count": len(sheet_names) - len(inherited),
        "lost_sheets": lost,
        "formula_count": formula_count,
        "formula_error_count": len(formula_errors),
        "status": "passed",
    }


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def write_tracking_files(project: Path, db_path: Path, now_iso: str) -> list[Path]:
    root = project / "Tracking" / "Prompt Response" / "Through Response 76"
    root.mkdir(parents=True, exist_ok=True)
    con = sqlite3.connect(db_path)
    try:
        info = con.execute("PRAGMA table_info(thread_response_reconciliation_cp3)").fetchall()
        columns = [row[1] for row in info]
        rows = [dict(zip(columns, row)) for row in con.execute("SELECT * FROM thread_response_reconciliation_cp3 ORDER BY response_number, response_key")]
        fraction_info = con.execute("PRAGMA table_info(fractional_prompt_cp3)").fetchall()
        fraction_columns = [row[1] for row in fraction_info]
        fraction_rows = [dict(zip(fraction_columns, row)) for row in con.execute("SELECT * FROM fractional_prompt_cp3 ORDER BY CAST(prompt_number AS REAL), prompt_number")]
    finally:
        con.close()
    response76 = next(row for row in rows if row.get("response_key") == "R76")
    response_path = root / "Response_76_Tracking.json"
    json_write(response_path, response76)
    raw_docx = root / "Medical References - Human Pathogen Database v3.0.0a Alternating Raw Prompts and Responses Through Response 76.docx"
    raw_doc = Document()
    raw_doc.core_properties.title = "Human Pathogen Database — Alternating Raw Prompts and Responses Through Response 76"
    raw_doc.core_properties.author = "Brent McAnulty, M.D."
    raw_doc.add_heading("Human Pathogen Database", 0)
    raw_doc.add_paragraph("Alternating Raw Prompts and Responses Through Response 76")
    for row in rows:
        number = row.get("response_label") or row.get("response_number")
        raw_doc.add_heading(f"Response {number}: {row.get('title') or 'Untitled exchange'}", level=1)
        table = raw_doc.add_table(rows=2, cols=1)
        table.style = "Table Grid"
        table.cell(0, 0).text = f"RAW PROMPT {number}\n\n{row.get('raw_prompt') or '[RAW PROMPT UNAVAILABLE]'}"
        _shade_cell(table.cell(0, 0), "D9EAF7")
        table.cell(1, 0).text = f"RAW RESPONSE {number}\n\n{row.get('raw_response') or row.get('summary') or '[RAW RESPONSE UNAVAILABLE]'}\n\n{row.get('summary') or ''}"
        _shade_cell(table.cell(1, 0), "E2F0D9")
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(9)
        raw_doc.add_paragraph()
    if fraction_rows:
        raw_doc.add_heading("Fractional prompts", level=1)
        for row in fraction_rows:
            raw_doc.add_paragraph(f"Prompt {row.get('prompt_number')}: {row.get('prompt_text')}")
    raw_doc.save(raw_docx)

    net_prompt = (
        "Continue the Human Pathogen Database from the newest verified checkpoint without regression. Use Google Drive as controlling storage and recover files autonomously. Preserve the immutable accepted clinical and publication state. Complete Section 5 print production through a quality-first premium-color master, an even-page print interior, exact cover/spine geometry, provider preview, physical proof, final preflight, tracking, indexes, manifests, recovery, and complete session/section/project restores."
    )
    net_response = (
        "Remediation Sections 1–4 are complete. Section 5 Session 1 Checkpoint 2 locks the KDP premium-color initial production master, creates the 538-page print-only interior, regenerates the full cover from separate components on the exact calculated canvas, completes deterministic preflight, and synchronizes the database, workbook, application, tracking, indexes, manifests, and recovery state. Provider preview, physical proof, final Session 1 freeze, and later Section 5 release controls remain pending."
    )
    net_docx = root / "Medical References - Human Pathogen Database v3.0.0a Alternating Net Prompts and Responses Through Response 76.docx"
    net_doc = Document()
    net_doc.core_properties.title = "Human Pathogen Database — Alternating Net Prompts and Responses Through Response 76"
    net_doc.core_properties.author = "Brent McAnulty, M.D."
    net_doc.add_heading("Human Pathogen Database", 0)
    net_doc.add_heading("Print-production and final-release remediation", level=1)
    table = net_doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.cell(0, 0).text = "NET PROMPT\n\n" + net_prompt
    _shade_cell(table.cell(0, 0), "D9EAF7")
    table.cell(1, 0).text = "NET RESPONSE\n\n" + net_response
    _shade_cell(table.cell(1, 0), "E2F0D9")
    net_doc.save(net_docx)

    everything = root / "Medical References - Human Pathogen Database v3.0.0a Everything in One Thread Through Response 76.xlsx"
    wb = Workbook()
    wb.remove(wb.active)
    raw_prompts = [{"Response": row.get("response_label") or row.get("response_number"), "Major Topic": row.get("major_topic"), "Title": row.get("title"), "Raw Prompt": row.get("raw_prompt")} for row in rows]
    raw_responses = [{"Response": row.get("response_label") or row.get("response_number"), "Major Topic": row.get("major_topic"), "Title": row.get("title"), "Raw Response": row.get("raw_response"), "Summary": row.get("summary")} for row in rows]
    summary_index = [{"Response": row.get("response_label") or row.get("response_number"), "Major Topic": row.get("major_topic"), "Title": row.get("title"), "Goal": row.get("goal"), "Summary": row.get("summary"), "State": row.get("state")} for row in rows]
    for title, data in {
        "Raw Prompts": raw_prompts,
        "Raw Responses": raw_responses,
        "Fractional Prompts": fraction_rows,
        "Net Prompt": [{"Major Topic": "Human Pathogen Database remediation", "Net Prompt": net_prompt}],
        "Net Response": [{"Major Topic": "Human Pathogen Database remediation", "Net Response": net_response}],
        "Summary Index": summary_index,
    }.items():
        ws = wb.create_sheet(title)
        _write_sheet(ws, data)
    wb.properties.title = "Human Pathogen Database — Everything in One Thread Through Response 76"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(everything)

    raw_net_md = root / "Medical References - Human Pathogen Database v3.0.0a Raw and Net Tracking Through Response 76.md"
    text_write(raw_net_md, f"""# Human Pathogen Database — Raw and Net Tracking Through Response 76

## Raw Prompt 76

Continue

## Raw Response 76

{response76.get('summary')}

## Net Prompt

{net_prompt}

## Net Response

{net_response}

Updated: {now_iso}
""")
    cumulative = root / "Medical References - Human Pathogen Database v3.0.0a Cumulative Thread Index Through Response 76.md"
    lines = ["# Human Pathogen Database — Cumulative Thread Index Through Response 76", ""]
    for row in rows:
        lines.append(f"- Response {row.get('response_label') or row.get('response_number')} — {row.get('title')}: {row.get('summary')}")
    lines.extend(["", f"Updated: {now_iso}"])
    text_write(cumulative, "\n".join(lines))
    return [response_path, raw_docx, net_docx, everything, raw_net_md, cumulative]


def write_application_surfaces(project: Path, db_path: Path, workbook_path: Path, print_interior: Path, cover_png: Path, now_iso: str) -> tuple[list[Path], dict[str, Any]]:
    root = project / "App" / "Section 5 Session 1 Checkpoint 2"
    root.mkdir(parents=True, exist_ok=True)
    main_apps = [p for p in project.rglob("human_pathogen_app.py") if p.is_file() and sha256_file(p) == APPLICATION_SHA256]
    if len(main_apps) != 1:
        raise RuntimeError({"main_application_candidates": [str(p) for p in main_apps]})
    app_path = main_apps[0]
    db_rel = db_path.relative_to(project).as_posix()
    workbook_rel = workbook_path.relative_to(project).as_posix()
    interior_rel = print_interior.relative_to(project).as_posix()
    cover_rel = cover_png.relative_to(project).as_posix()
    pointer = root / "CURRENT_DATABASE.txt"
    text_write(pointer, db_rel + "\n")
    state = root / "CURRENT_PROJECT_STATE.json"
    json_write(state, {
        "schema": "mrhpd-section5-current-project-state-1.1",
        "response": 76,
        "section": SECTION_LABEL,
        "session": SESSION_LABEL,
        "checkpoint": CHECKPOINT_LABEL,
        "state": "checkpoint_complete",
        "database": db_rel,
        "workbook": workbook_rel,
        "print_interior": interior_rel,
        "cover": cover_rel,
        "main_application": app_path.relative_to(project).as_posix(),
        "main_application_sha256": sha256_file(app_path),
        "main_application_unchanged": True,
        "recorded_at": now_iso,
    })
    audit_script = root / "audit_section5_checkpoint2.py"
    text_write(audit_script, f'''#!/usr/bin/env python3
import json, sqlite3
from pathlib import Path
from openpyxl import load_workbook
from pypdf import PdfReader
project=Path(__file__).resolve().parents[2]
db=project/{db_rel!r}
workbook=project/{workbook_rel!r}
interior=project/{interior_rel!r}
cover=project/{cover_rel!r}
con=sqlite3.connect(db)
try:
 integrity=con.execute("PRAGMA integrity_check").fetchone()[0]
 fk=len(list(con.execute("PRAGMA foreign_key_check")))
 response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R76'").fetchone()[0]
 selection=con.execute("SELECT status FROM section5_print_selection WHERE checkpoint_code='{CHECKPOINT_CODE}'").fetchone()
 pages=con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code='{CHECKPOINT_CODE}'").fetchone()[0]
 failed=con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code='{CHECKPOINT_CODE}' AND status!='passed'").fetchone()[0]
finally: con.close()
wb=load_workbook(workbook,read_only=True,data_only=False)
try: sheets=len(wb.sheetnames)
finally: wb.close()
pdf=PdfReader(str(interior)); page_count=len(pdf.pages); searchable=sum(1 for p in pdf.pages[:537] if (p.extract_text() or '').strip())
result={{'status':'passed' if integrity=='ok' and fk==0 and response==1 and selection==('locked_initial_production_master',) and pages==538 and failed==0 and sheets>=100 and page_count==538 and searchable==537 and cover.exists() else 'failed','integrity':integrity,'foreign_keys':fk,'response76':response,'selection':selection,'page_records':pages,'failed_page_records':failed,'workbook_sheets':sheets,'print_pages':page_count,'searchable_pages':searchable,'cover_exists':cover.exists()}}
print(json.dumps(result,indent=2))
raise SystemExit(0 if result['status']=='passed' else 1)
''')
    result = subprocess.run([sys.executable, str(audit_script)], cwd=project, text=True, capture_output=True, timeout=300)
    if result.returncode:
        raise RuntimeError({"checkpoint2_application_audit_failed": {"stdout": result.stdout[-8000:], "stderr": result.stderr[-8000:]}})
    audit = json.loads(result.stdout)
    audit.update({"main_application_path": app_path.relative_to(project).as_posix(), "main_application_sha256": sha256_file(app_path), "main_application_unchanged": True})
    output = root / "SECTION5_CHECKPOINT2_APPLICATION_AUDIT.json"
    json_write(output, audit)
    return [pointer, state, audit_script, output], audit


def extract_text_for_index(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in {".md", ".txt", ".csv", ".json", ".py", ".html", ".yml", ".yaml"}:
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix == ".docx":
            doc = Document(path)
            chunks = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells)
            return "\n".join(chunks)
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        if suffix == ".xlsx":
            wb = load_workbook(path, read_only=True, data_only=False)
            try:
                chunks: list[str] = []
                for ws in wb.worksheets:
                    chunks.append(ws.title)
                    for row in ws.iter_rows(values_only=True):
                        chunks.append(" | ".join("" if value is None else str(value) for value in row))
                return "\n".join(chunks)
            finally:
                wb.close()
        if suffix in {".sqlite", ".db"}:
            con = sqlite3.connect(path)
            try:
                schema = [row[0] for row in con.execute("SELECT sql FROM sqlite_master WHERE sql IS NOT NULL ORDER BY name")]
                return "\n".join(schema)
            finally:
                con.close()
    except Exception as exc:
        return f"[index extraction error: {exc!r}]"
    return ""


def build_source_and_bit_indexes(project: Path, now_iso: str) -> dict[str, Any]:
    root = project / "Indexes" / "Section 5 Session 1 Checkpoint 2"
    root.mkdir(parents=True, exist_ok=True)
    source_json = root / "MRHPD v3.0.0a Section 5 Checkpoint 2 Source Index.json"
    source_csv = root / "MRHPD v3.0.0a Section 5 Checkpoint 2 Source Index.csv"
    bit_path = root / "MRHPD v3.0.0a Section 5 Checkpoint 2 Bit Index.sqlite"
    qa_path = root / "MRHPD v3.0.0a Section 5 Checkpoint 2 Index QA.json"
    excluded = {source_json.resolve(), source_csv.resolve(), bit_path.resolve(), qa_path.resolve()}
    rows: list[dict[str, Any]] = []
    fts_payloads: list[tuple[str, str, str, str]] = []
    searchable_suffixes = {".md", ".txt", ".csv", ".json", ".py", ".html", ".yml", ".yaml", ".docx", ".pdf", ".xlsx", ".sqlite", ".db"}
    for path in sorted(p for p in project.rglob("*") if p.is_file() and p.resolve() not in excluded):
        rel = path.relative_to(project).as_posix()
        purpose = "Project artifact"
        if rel.startswith("Database/"):
            purpose = "Canonical or historical project database"
        elif rel.startswith("Documents/"):
            purpose = "Publication or editable assembly"
        elif rel.startswith("Print Production/"):
            purpose = "Section 5 print-production derivative"
        elif rel.startswith("Cover/"):
            purpose = "Cover component or historical cover master"
        elif rel.startswith("Tracking/"):
            purpose = "Prompt, response, summary, and project tracking"
        elif rel.startswith("QA/"):
            purpose = "Validation and acceptance evidence"
        elif rel.startswith("Reports/"):
            purpose = "Human-readable report or register"
        elif rel.startswith("App/"):
            purpose = "Local application or current-state audit surface"
        user_searchable = path.suffix.lower() in searchable_suffixes
        content = extract_text_for_index(path) if user_searchable else ""
        row = {"record_type": "physical_file", "path": rel, "container_path": "", "name": path.name, "purpose": purpose, "bytes": path.stat().st_size, "sha256": sha256_file(path), "user_searchable": int(user_searchable)}
        rows.append(row)
        fts_payloads.append((rel, path.name, purpose, content))
        if path.suffix.lower() == ".zip":
            try:
                with zipfile.ZipFile(path) as zf:
                    if zf.testzip() is not None:
                        continue
                    for info in zf.infolist():
                        if info.is_dir():
                            continue
                        member_path = f"{rel}!/{info.filename}"
                        member_content = ""
                        member_suffix = Path(info.filename).suffix.lower()
                        if member_suffix in {".md", ".txt", ".csv", ".json", ".py", ".html", ".yml", ".yaml"} and info.file_size <= 5_000_000:
                            try:
                                member_content = zf.read(info).decode("utf-8", errors="replace")
                            except Exception:
                                member_content = ""
                        member_row = {"record_type": "container_member", "path": member_path, "container_path": rel, "name": Path(info.filename).name, "purpose": "Member of project container", "bytes": info.file_size, "sha256": "", "user_searchable": int(bool(member_content))}
                        rows.append(member_row)
                        fts_payloads.append((member_path, member_row["name"], member_row["purpose"], member_content))
            except zipfile.BadZipFile:
                pass
    json_write(source_json, {"schema": "mrhpd-source-index-2.0", "generated_at": now_iso, "record_count": len(rows), "records": rows})
    csv_write(source_csv, rows)
    if bit_path.exists():
        bit_path.unlink()
    con = sqlite3.connect(bit_path)
    try:
        con.executescript("""
        CREATE TABLE artifact (
            artifact_id INTEGER PRIMARY KEY,
            record_type TEXT NOT NULL,
            path TEXT NOT NULL UNIQUE,
            container_path TEXT,
            name TEXT NOT NULL,
            purpose TEXT NOT NULL,
            bytes INTEGER NOT NULL,
            sha256 TEXT,
            user_searchable INTEGER NOT NULL
        );
        CREATE VIRTUAL TABLE artifact_fts USING fts5(path, name, purpose, content);
        """)
        for row, payload in zip(rows, fts_payloads):
            con.execute("INSERT INTO artifact (record_type,path,container_path,name,purpose,bytes,sha256,user_searchable) VALUES (?,?,?,?,?,?,?,?)", (row["record_type"], row["path"], row["container_path"], row["name"], row["purpose"], row["bytes"], row["sha256"], row["user_searchable"]))
            con.execute("INSERT INTO artifact_fts (path,name,purpose,content) VALUES (?,?,?,?)", payload)
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        counts = {
            "artifact": con.execute("SELECT COUNT(*) FROM artifact").fetchone()[0],
            "fts": con.execute("SELECT COUNT(*) FROM artifact_fts").fetchone()[0],
            "response76": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"Response 76"',)).fetchone()[0],
            "premium_color": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"premium color"',)).fetchone()[0],
            "print_interior": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"print interior"',)).fetchone()[0],
        }
        con.commit()
    finally:
        con.close()
    if integrity != "ok" or counts["artifact"] != len(rows) or counts["fts"] != len(rows):
        raise RuntimeError({"bit_index_gate": {"integrity": integrity, "counts": counts, "expected": len(rows)}})
    qa = {"status": "passed", "generated_at": now_iso, "source_index_records": len(rows), "physical_files": sum(1 for row in rows if row["record_type"] == "physical_file"), "container_members": sum(1 for row in rows if row["record_type"] == "container_member"), "bit_index_integrity": integrity, "counts": counts, "bit_index_sha256": sha256_file(bit_path)}
    json_write(qa_path, qa)
    return {"source_json": source_json, "source_csv": source_csv, "bit_index": bit_path, "qa_path": qa_path, "qa": qa}


def build_project_manifest(project: Path, now_iso: str) -> tuple[Path, Path, list[dict[str, Any]]]:
    root = project / "Manifest" / "Section 5 Session 1 Checkpoint 2"
    root.mkdir(parents=True, exist_ok=True)
    manifest = root / "MRHPD v3.0.0a Section 5 Checkpoint 2 Current Project Manifest.json"
    checksums = root / "MRHPD v3.0.0a Section 5 Checkpoint 2 Current Project Checksums.sha256"
    rows = []
    for path in sorted(p for p in project.rglob("*") if p.is_file() and p not in {manifest, checksums}):
        rows.append({"path": path.relative_to(project).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    json_write(manifest, {"schema": "mrhpd-current-project-manifest-2.0", "generated_at": now_iso, "exclusions": [manifest.relative_to(project).as_posix(), checksums.relative_to(project).as_posix()], "file_count": len(rows), "total_bytes": sum(row["bytes"] for row in rows), "files": rows})
    text_write(checksums, "".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
    for row in rows:
        path = project / row["path"]
        if path.stat().st_size != row["bytes"] or sha256_file(path) != row["sha256"]:
            raise RuntimeError({"manifest_mismatch": row["path"]})
    return manifest, checksums, rows


def create_apply_script(manifest: dict[str, Any], expected: dict[str, Any]) -> str:
    return f'''#!/usr/bin/env python3
from __future__ import annotations
import argparse, hashlib, json, re, shutil, sqlite3, sys, tempfile, zipfile
from pathlib import Path, PurePosixPath
from openpyxl import load_workbook
from pypdf import PdfReader
BASE_RESTORE_BYTES={BASE_RESTORE_BYTES}
BASE_RESTORE_SHA256={BASE_RESTORE_SHA256!r}
BASE_PROJECT_BYTES={BASE_PROJECT_BYTES}
BASE_PROJECT_SHA256={BASE_PROJECT_SHA256!r}
PUBLICATION_SHA256={PUBLICATION_SHA256!r}
APPLICATION_SHA256={APPLICATION_SHA256!r}
CURRENT_DB_REL={CURRENT_DB_REL!r}
CURRENT_WORKBOOK_REL={CURRENT_WORKBOOK_REL!r}
PUBLICATION_REL={PUBLICATION_REL!r}
PRINT_INTERIOR_REL={PRINT_INTERIOR_REL!r}
FINAL_COVER_PNG_REL={FINAL_COVER_PNG_REL!r}
MANIFEST={manifest!r}
EXPECTED={expected!r}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as f:
  for b in iter(lambda:f.read(1024*1024),b''): h.update(b)
 return h.hexdigest()
def safe_extract(path,dest):
 dest.mkdir(parents=True,exist_ok=True)
 with zipfile.ZipFile(path) as zf:
  if zf.testzip() is not None: raise RuntimeError('ZIP CRC failure: '+str(path))
  names=zf.namelist()
  if len(names)!=len(set(names)): raise RuntimeError('duplicate ZIP members')
  for name in names:
   pp=PurePosixPath(name.replace('\\\\','/'))
   if pp.is_absolute() or '..' in pp.parts or re.match(r'^[A-Za-z]:',name): raise RuntimeError('unsafe ZIP path: '+name)
  zf.extractall(dest)
def verify(path,size,digest,label):
 observed={{'bytes':path.stat().st_size,'sha256':sha(path)}}
 if observed!={{'bytes':size,'sha256':digest}}: raise RuntimeError({{label:observed}})
def main():
 ap=argparse.ArgumentParser()
 ap.add_argument('--base-response72-restore',type=Path,required=True)
 ap.add_argument('--output-dir',type=Path,required=True)
 args=ap.parse_args()
 verify(args.base_response72_restore,BASE_RESTORE_BYTES,BASE_RESTORE_SHA256,'baseline_restore')
 package=Path(__file__).resolve().parents[1]
 overlay=package/'OVERLAY'
 if not overlay.is_dir(): raise RuntimeError('OVERLAY missing')
 if args.output_dir.exists() and any(args.output_dir.iterdir()): raise RuntimeError('output directory must be empty')
 args.output_dir.mkdir(parents=True,exist_ok=True)
 with tempfile.TemporaryDirectory(prefix='mrhpd-r76-apply-') as td:
  work=Path(td); restore_root=work/'restore'; safe_extract(args.base_response72_restore,restore_root)
  candidates=[p for p in restore_root.rglob('*.zip') if p.stat().st_size==BASE_PROJECT_BYTES and sha(p)==BASE_PROJECT_SHA256]
  if len(candidates)!=1: raise RuntimeError({{'project_archive_candidates':[str(p) for p in candidates]}})
  extracted=work/'project'; safe_extract(candidates[0],extracted)
  roots=[p for p in extracted.iterdir() if p.is_dir()]
  source=roots[0] if len(roots)==1 else extracted
  destination=args.output_dir/source.name
  shutil.copytree(source,destination)
  for row in MANIFEST['overlay_files']:
   src=overlay/row['path']; verify(src,row['bytes'],row['sha256'],'overlay_'+row['path'])
   dst=destination/row['path']; dst.parent.mkdir(parents=True,exist_ok=True); shutil.copy2(src,dst)
  db=destination/CURRENT_DB_REL
  con=sqlite3.connect(db)
  try:
   integrity=con.execute('PRAGMA integrity_check').fetchone()[0]
   fk=len(list(con.execute('PRAGMA foreign_key_check')))
   response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R76'").fetchone()[0]
   selection=con.execute("SELECT status FROM section5_print_selection WHERE checkpoint_code='MRHPD-V3-CP5-S1-CP2'").fetchone()
   pages=con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code='MRHPD-V3-CP5-S1-CP2'").fetchone()[0]
   failed=con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code='MRHPD-V3-CP5-S1-CP2' AND status!='passed'").fetchone()[0]
  finally: con.close()
  wb=load_workbook(destination/CURRENT_WORKBOOK_REL,read_only=True,data_only=False)
  try: sheets=len(wb.sheetnames)
  finally: wb.close()
  interior=destination/PRINT_INTERIOR_REL
  cover=destination/FINAL_COVER_PNG_REL
  verify(interior,EXPECTED['interior_bytes'],EXPECTED['interior_sha256'],'interior')
  verify(cover,EXPECTED['cover_bytes'],EXPECTED['cover_sha256'],'cover')
  reader=PdfReader(str(interior)); print_pages=len(reader.pages); searchable=sum(1 for page in reader.pages[:537] if (page.extract_text() or '').strip())
  publication=destination/PUBLICATION_REL
  if sha(publication)!=PUBLICATION_SHA256: raise RuntimeError('immutable publication changed')
  apps=[p for p in destination.rglob('human_pathogen_app.py') if p.is_file() and sha(p)==APPLICATION_SHA256]
  result={{'status':'passed' if integrity=='ok' and fk==0 and response==1 and selection==('locked_initial_production_master',) and pages==538 and failed==0 and sheets>=100 and print_pages==538 and searchable==537 and len(apps)==1 else 'failed','project_root':str(destination),'database':{{'integrity':integrity,'foreign_keys':fk,'response76':response,'selection':selection,'page_records':pages,'failed_pages':failed}},'workbook_sheets':sheets,'print_pages':print_pages,'searchable_pages':searchable,'publication_sha256':sha(publication),'main_application_matches':len(apps),'interior_sha256':sha(interior),'cover_sha256':sha(cover)}}
  output=args.output_dir/'MRHPD_RESPONSE76_CHECKPOINT_RECOVERY_APPLICATION_RESULT.json'; output.write_text(json.dumps(result,indent=2),encoding='utf-8')
  print(json.dumps(result,indent=2))
  raise SystemExit(0 if result['status']=='passed' else 1)
if __name__=='__main__': main()
'''


def build_recovery_package(
    *,
    baseline_project: Path,
    current_project: Path,
    baseline_restore: Path,
    project_archive: Path,
    dist: Path,
    now: datetime,
    summary: dict[str, Any],
    direct_files: list[Path],
) -> dict[str, Any]:
    baseline_map = {p.relative_to(baseline_project).as_posix(): (p.stat().st_size, sha256_file(p)) for p in baseline_project.rglob("*") if p.is_file()}
    current_map = {p.relative_to(current_project).as_posix(): (p.stat().st_size, sha256_file(p)) for p in current_project.rglob("*") if p.is_file()}
    deleted = sorted(set(baseline_map) - set(current_map))
    if deleted:
        raise RuntimeError({"unexpected_deleted_paths": deleted})
    overlay_rows = []
    for rel, identity in sorted(current_map.items()):
        if baseline_map.get(rel) != identity:
            overlay_rows.append({"path": rel, "bytes": identity[0], "sha256": identity[1], "change": "new" if rel not in baseline_map else "changed"})
    stamp = now.strftime("%Y-%m-%d %H%M UTC")
    package_root = dist / "recovery_package_root"
    if package_root.exists():
        shutil.rmtree(package_root)
    overlay_root = package_root / "OVERLAY"
    tools = package_root / "TOOLS"
    overlay_root.mkdir(parents=True)
    tools.mkdir(parents=True)
    for row in overlay_rows:
        source = current_project / row["path"]
        target = overlay_root / row["path"]
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, target)
    manifest = {
        "schema": "mrhpd-section5-checkpoint-recovery-1.1",
        "generated_at": now.isoformat().replace("+00:00", "Z"),
        "version": PROJECT_VERSION,
        "response": RESPONSE_NUMBER,
        "section": SECTION_LABEL,
        "session": SESSION_LABEL,
        "checkpoint": CHECKPOINT_LABEL,
        "state": "checkpoint_complete",
        "baseline_restore": {"name": baseline_restore.name, "bytes": baseline_restore.stat().st_size, "sha256": sha256_file(baseline_restore)},
        "baseline_project": {"name": project_archive.name, "bytes": project_archive.stat().st_size, "sha256": sha256_file(project_archive)},
        "overlay_file_count": len(overlay_rows),
        "overlay_total_bytes": sum(row["bytes"] for row in overlay_rows),
        "overlay_files": overlay_rows,
        "deleted_paths": [],
        "accepted_predecessor_mutated": False,
        "frozen_section3_release_mutated": False,
        "immutable_publication_mutated": False,
        "user_upload_required": False,
        "requires_conversation_reconstruction": False,
        "next": "Remediation Section 5 of 5 Session 1 of 3 Checkpoint 3 of 3",
    }
    json_write(package_root / "CHECKPOINT_RECOVERY_MANIFEST.json", manifest)
    text_write(package_root / "CHECKPOINT_RECOVERY_CHECKSUMS.sha256", "".join(f"{row['sha256']}  OVERLAY/{row['path']}\n" for row in overlay_rows))
    expected = {
        "interior_bytes": (current_project / PRINT_INTERIOR_REL).stat().st_size,
        "interior_sha256": sha256_file(current_project / PRINT_INTERIOR_REL),
        "cover_bytes": (current_project / FINAL_COVER_PNG_REL).stat().st_size,
        "cover_sha256": sha256_file(current_project / FINAL_COVER_PNG_REL),
    }
    text_write(tools / "apply_checkpoint_recovery.py", create_apply_script(manifest, expected))
    text_write(package_root / "RESTORE_READ_FIRST.md", f"""# Human Pathogen Database — Response 76 Checkpoint Recovery

This cumulative intermediate recovery applies directly to the exact Response 72 complete restore and includes all current progress through Response 76. Response 75 does not need to be applied separately.

## Required baseline

Filename: `{baseline_restore.name}`

Bytes: `{baseline_restore.stat().st_size}`

SHA-256: `{sha256_file(baseline_restore)}`

## Automated apply

```bash
python TOOLS/apply_checkpoint_recovery.py \
  --base-response72-restore "<Response 72 complete restore.zip>" \
  --output-dir "<empty destination>"
```

The utility verifies the exact baseline, every overlay file, the Response 76 database and workbook, the 538-page print interior, the regenerated exact cover, the immutable 537-page publication, and the unchanged main application.
""")
    recovery_zip = dist / (
        f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 "
        f"Remediation Section 5 of 5 Session 1 of 3 Checkpoint 2 of 3 RECOVERY DATA THROUGH RESPONSE 76 {stamp}.zip"
    )
    with zipfile.ZipFile(recovery_zip, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in sorted(package_root.rglob("*")):
            if path.is_file():
                zf.write(path, path.relative_to(package_root).as_posix())
    recovery_qa = verify_zip(recovery_zip)
    with tempfile.TemporaryDirectory(prefix="mrhpd-r76-clean-apply-") as td:
        output = Path(td) / "restored"
        result = subprocess.run([sys.executable, str(tools / "apply_checkpoint_recovery.py"), "--base-response72-restore", str(baseline_restore), "--output-dir", str(output)], cwd=package_root, text=True, capture_output=True, timeout=1800)
        if result.returncode:
            raise RuntimeError({"clean_apply_failed": {"stdout": result.stdout[-16000:], "stderr": result.stderr[-16000:]}})
        clean_apply = json.loads((output / "MRHPD_RESPONSE76_CHECKPOINT_RECOVERY_APPLICATION_RESULT.json").read_text(encoding="utf-8"))
    verification = {
        "schema": "mrhpd-response76-recovery-verification-1.0",
        "generated_at": now.isoformat().replace("+00:00", "Z"),
        "status": "passed",
        "recovery_zip": recovery_qa,
        "manifest": {"overlay_file_count": len(overlay_rows), "overlay_total_bytes": sum(row["bytes"] for row in overlay_rows), "deleted_paths": 0},
        "clean_apply": clean_apply,
        "accepted_predecessor_mutated": False,
        "frozen_section3_release_mutated": False,
        "immutable_publication_mutated": False,
        "user_upload_required": False,
        "checkpoint_2_of_3_complete": True,
        "session_1_of_3_complete": False,
        "remediation_section_5_complete": False,
        "next": "Checkpoint 3 of 3 - Session 1 freeze and complete restore",
    }
    verification_path = dist / "MRHPD v3.0.0a Response 76 Checkpoint 2 Recovery Verification.json"
    json_write(verification_path, verification)
    sha_path = dist / f"{recovery_zip.name}.sha256.txt"
    text_write(sha_path, f"{recovery_qa['sha256']}  {recovery_zip.name}\n")
    summary_path = dist / "MRHPD_RESPONSE76_SECTION5_CHECKPOINT2_BUILD_SUMMARY.json"
    json_write(summary_path, summary | {"recovery": verification})
    exact_names = dist / "MRHPD v3.0.0a Response 76 Exact File Names.txt"
    text_write(exact_names, f"""Response 76 cumulative checkpoint recovery ZIP:
{recovery_zip.name}

Required baseline complete restore:
{baseline_restore.name}

Required baseline project archive embedded in that restore:
{project_archive.name}

Current copied SQLite database:
{Path(CURRENT_DB_REL).name}

Current comprehensive workbook:
{Path(CURRENT_WORKBOOK_REL).name}

Print-production interior:
{Path(PRINT_INTERIOR_REL).name}

Final full-cover PNG:
{Path(FINAL_COVER_PNG_REL).name}

Final full-cover PDF:
{Path(FINAL_COVER_PDF_REL).name}

Exact template PNG:
{Path(TEMPLATE_PNG_REL).name}

Exact template PDF:
{Path(TEMPLATE_PDF_REL).name}
""")
    delivery = dist / f"MRHPD v3.0.0a Response 76 Section 5 Session 1 Checkpoint 2 Recovery Package {stamp}.zip"
    with zipfile.ZipFile(delivery, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in [recovery_zip, sha_path, verification_path, summary_path, exact_names, *direct_files]:
            if path.exists():
                zf.write(path, path.name)
    delivery_qa = verify_zip(delivery)
    return {
        "recovery_zip": recovery_zip,
        "recovery_qa": recovery_qa,
        "verification_path": verification_path,
        "summary_path": summary_path,
        "exact_names": exact_names,
        "delivery": delivery,
        "delivery_qa": delivery_qa,
        "overlay_rows": overlay_rows,
        "clean_apply": clean_apply,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--volume1-dir", type=Path, required=True)
    parser.add_argument("--volume2-dir", type=Path, required=True)
    parser.add_argument("--checkpoint1-dir", type=Path, required=True)
    parser.add_argument("--dist", type=Path, default=Path("dist_cp5_s1_cp2"))
    args = parser.parse_args()
    now = utc_now()
    now_iso = now.isoformat().replace("+00:00", "Z")
    if args.dist.exists():
        shutil.rmtree(args.dist)
    args.dist.mkdir(parents=True)
    with tempfile.TemporaryDirectory(prefix="mrhpd-cp5-s1-cp2-") as td:
        work = Path(td)
        restore, project_archive, baseline_project, checkpoint1_project, checkpoint1_application = restore_checkpoint1(args.volume1_dir, args.volume2_dir, args.checkpoint1_dir, work)
        current_project = work / "current_project" / checkpoint1_project.name
        current_project.parent.mkdir(parents=True)
        shutil.copytree(checkpoint1_project, current_project)
        selection = selection_record(now_iso)
        qa_root = current_project / "QA" / "Section 5 Session 1" / "Checkpoint 2"
        print_interior = current_project / PRINT_INTERIOR_REL
        interior_summary, page_rows, proof_paths = generate_print_interior(current_project / PUBLICATION_REL, print_interior, qa_root / "Interior", selection)
        cover_summary, cover_paths = generate_cover(current_project, selection, qa_root / "Cover")
        preflight = preflight_records(selection, interior_summary, cover_summary, now_iso)
        events = recovery_events(now_iso)
        current_db = current_project / CURRENT_DB_REL
        database_qa = synchronize_database(checkpoint1_project / CP1_DB_REL, current_db, now_iso=now_iso, selection=selection, interior=interior_summary, page_rows=page_rows, cover=cover_summary, preflight=preflight, events=events)
        categories = [
            {"Category Key": "MRHPD-CAT-PRINT-PRODUCTION", "Parent": "", "Label": "Print production", "Scope": "project-wide", "Status": "current"},
            {"Category Key": "MRHPD-CAT-PRINT-INTERIOR", "Parent": "MRHPD-CAT-PRINT-PRODUCTION", "Label": "Print interior", "Scope": "publication", "Status": "current"},
            {"Category Key": "MRHPD-CAT-COVER-PRODUCTION", "Parent": "MRHPD-CAT-PRINT-PRODUCTION", "Label": "Cover production", "Scope": "publication", "Status": "current"},
            {"Category Key": "MRHPD-CAT-PROVIDER-PREVIEW", "Parent": "MRHPD-CAT-PRINT-PRODUCTION", "Label": "Provider preview", "Scope": "external gate", "Status": "current"},
            {"Category Key": "MRHPD-CAT-PHYSICAL-PROOF", "Parent": "MRHPD-CAT-PRINT-PRODUCTION", "Label": "Physical proof", "Scope": "external gate", "Status": "current"},
        ]
        current_workbook = current_project / CURRENT_WORKBOOK_REL
        workbook_qa = augment_workbook(checkpoint1_project / CP1_WORKBOOK_REL, current_workbook, selection=selection, interior=interior_summary, page_rows=page_rows, cover=cover_summary, preflight=preflight, events=events, categories=categories)
        tracking_files = write_tracking_files(current_project, current_db, now_iso)
        application_files, application_qa = write_application_surfaces(current_project, current_db, current_workbook, print_interior, current_project / FINAL_COVER_PNG_REL, now_iso)
        source_control_root = current_project / "Sources" / "Print Production" / "Response 76"
        source_control_path = source_control_root / "MRHPD v3.0.0a Response 76 KDP Official Requirement Control.json"
        json_write(source_control_path, {"generated_at": now_iso, "selection": selection, "requirements": [
            {"requirement": "Premium color intended for image-heavy books where sharpness and clarity are pivotal", "url": "https://kdp.amazon.com/en_US/help/topic/G201953020"},
            {"requirement": "538 pages fits the 8.5 × 11-inch premium-color maximum", "url": "https://kdp.amazon.com/en_US/help/topic/G201834180"},
            {"requirement": "Cover bleed, color spine factor, live-area and spine-text safety", "url": "https://kdp.amazon.com/en_US/help/topic/G201857950"},
            {"requirement": "501–700-page no-bleed gutter and outside-margin controls", "url": "https://kdp.amazon.com/en_US/help/topic/GVBQ3CMEQW3W2VL6/"},
            {"requirement": "Image, font, barcode, and cover production controls", "url": "https://kdp.amazon.com/en_US/help/topic/G201953020"},
        ], "evidence_boundary": "Official-provider production requirements and current deterministic calculations; provider Print Previewer and physical proof remain controlling external gates."})
        data_root = current_project / "Data" / "Section 5 Session 1 Checkpoint 2"
        selection_json = data_root / "MRHPD v3.0.0a Response 76 Print Selection.json"
        selection_csv = data_root / "MRHPD v3.0.0a Response 76 Print Selection.csv"
        preflight_json = data_root / "MRHPD v3.0.0a Response 76 Print Preflight.json"
        preflight_csv = data_root / "MRHPD v3.0.0a Response 76 Print Preflight.csv"
        categories_json = data_root / "MRHPD v3.0.0a Response 76 Master Category Extension.json"
        categories_csv = data_root / "MRHPD v3.0.0a Response 76 Master Category Extension.csv"
        json_write(selection_json, selection); csv_write(selection_csv, [selection])
        json_write(preflight_json, preflight); csv_write(preflight_csv, preflight)
        json_write(categories_json, categories); csv_write(categories_csv, categories)
        report_root = current_project / "Reports" / "Section 5 Session 1" / "Checkpoint 2"
        artwork_root = current_project / "Artwork" / "Section 5 Print Production" / "Checkpoint 2"
        selection_figure = artwork_root / "MRHPD-FIG-S5-0003 Locked Premium Color Production Scenario v3.0.0a.png"
        margin_figure = artwork_root / "MRHPD-FIG-S5-0004 Print Interior Margin Normalization v3.0.0a.png"
        selection_figure_qa = reporting.build_selection_figure(selection_figure, selection)
        margin_figure_qa = reporting.build_margin_figure(margin_figure, interior_summary)
        docx_report = report_root / "MRHPD v3.0.0a Section 5 Session 1 Checkpoint 2 Print Production Candidate Report.docx"
        pdf_report = report_root / "MRHPD v3.0.0a Section 5 Session 1 Checkpoint 2 Print Production Candidate Report.pdf"
        xlsx_register = report_root / "MRHPD v3.0.0a Section 5 Session 1 Checkpoint 2 Print Production Register.xlsx"
        direct_figures = [selection_figure, margin_figure, current_project / PROOF_PNG_REL]
        docx_qa = reporting.build_docx_report(docx_report, generated_at=now_iso, selection=selection, interior=interior_summary, cover=cover_summary, preflight=preflight, recovery_events=events, figure_paths=direct_figures)
        pdf_qa = reporting.build_pdf_report(pdf_report, generated_at=now_iso, selection=selection, interior=interior_summary, cover=cover_summary, preflight=preflight, recovery_events=events, figure_paths=direct_figures)
        register_qa = reporting.build_register(xlsx_register, selection=selection, interior_summary=interior_summary, page_rows=page_rows, cover=cover_summary, preflight=preflight, recovery_events=events, tracking_rows=[{"Response": 76, "Raw Prompt": "Continue", "Summary": "Premium-color scenario, 538-page interior, exact cover, and preflight completed."}])
        report_render_qa = reporting.render_report_qa(pdf_report, report_root / "Rendered Report QA")
        qa_root.mkdir(parents=True, exist_ok=True)
        qa_files = {
            "DATABASE_QA.json": database_qa,
            "WORKBOOK_QA.json": workbook_qa,
            "APPLICATION_QA.json": application_qa,
            "PRINT_INTERIOR_QA.json": interior_summary,
            "COVER_QA.json": cover_summary,
            "PREFLIGHT_QA.json": {"status": "passed_with_controlled_external_gates", "records": preflight, "hard_failures": sum(1 for row in preflight if row["status"] == "failed"), "controlled_pending": sum(1 for row in preflight if row["status"] == "controlled_pending"), "controlled_warnings": sum(1 for row in preflight if row["status"] == "controlled_warning")},
            "REPORT_QA.json": {"status": "passed", "docx": docx_qa, "pdf": pdf_qa, "xlsx": register_qa, "render": report_render_qa, "figures": [selection_figure_qa, margin_figure_qa]},
            "RECOVERY_EVENTS_181_188.json": events,
        }
        for name, payload in qa_files.items():
            json_write(qa_root / name, payload)
        final_qa = {
            "schema": "mrhpd-section5-checkpoint2-qa-1.0",
            "generated_at": now_iso,
            "status": "passed",
            "response": 76,
            "section": SECTION_LABEL,
            "session": SESSION_LABEL,
            "checkpoint": CHECKPOINT_LABEL,
            "selection": selection,
            "database": database_qa,
            "workbook": workbook_qa,
            "application": application_qa,
            "print_interior": interior_summary,
            "cover": cover_summary,
            "preflight": qa_files["PREFLIGHT_QA.json"],
            "reports": qa_files["REPORT_QA.json"],
            "checkpoint_2_of_3_complete": True,
            "session_1_of_3_complete": False,
            "remediation_section_5_complete": False,
            "accepted_predecessor_mutated": False,
            "frozen_section3_release_mutated": False,
            "immutable_publication_mutated": False,
            "main_application_mutated": False,
            "provider_previewer": "pending_external",
            "physical_proof": "pending_future_checkpoint",
            "next": "Checkpoint 3 of 3 - Session 1 freeze and complete restore",
        }
        final_qa_path = qa_root / "SECTION5_CHECKPOINT2_QA.json"
        json_write(final_qa_path, final_qa)
        index_result = build_source_and_bit_indexes(current_project, now_iso)
        manifest_path, checksums_path, manifest_rows = build_project_manifest(current_project, now_iso)
        summary = {
            "schema": "mrhpd-response76-section5-checkpoint2-build-1.0",
            "generated_at": now_iso,
            "status": "passed",
            "response": 76,
            "section": SECTION_LABEL,
            "session": SESSION_LABEL,
            "checkpoint": CHECKPOINT_LABEL,
            "checkpoint1_application": checkpoint1_application,
            "selection": selection,
            "database": database_qa,
            "workbook": workbook_qa,
            "application": application_qa,
            "print_interior": interior_summary,
            "cover": cover_summary,
            "preflight": qa_files["PREFLIGHT_QA.json"],
            "reports": qa_files["REPORT_QA.json"],
            "index": index_result["qa"],
            "manifest_records": len(manifest_rows),
            "user_upload_required": False,
            "checkpoint_2_of_3_complete": True,
            "session_1_of_3_complete": False,
            "remediation_section_5_complete": False,
            "next": "Checkpoint 3 of 3 - Session 1 freeze and complete restore",
        }
        direct_files = [docx_report, pdf_report, xlsx_register, selection_figure, margin_figure, current_project / PROOF_PNG_REL]
        package = build_recovery_package(baseline_project=baseline_project, current_project=current_project, baseline_restore=restore, project_archive=project_archive, dist=args.dist, now=now, summary=summary, direct_files=direct_files)
        console = {
            "status": "passed",
            "delivery": package["delivery"].name,
            "delivery_bytes": package["delivery_qa"]["bytes"],
            "delivery_sha256": package["delivery_qa"]["sha256"],
            "recovery_zip": package["recovery_zip"].name,
            "recovery_zip_bytes": package["recovery_qa"]["bytes"],
            "recovery_zip_sha256": package["recovery_qa"]["sha256"],
            "overlay_files": len(package["overlay_rows"]),
            "database_tables": database_qa["table_count"],
            "workbook_sheets": workbook_qa["current_sheet_count"],
            "print_pages": interior_summary["output_page_count"],
            "searchable_pages": interior_summary["searchable_pages"],
            "cover_pixels": [cover_summary["pixel_width"], cover_summary["pixel_height"]],
            "preflight_hard_failures": qa_files["PREFLIGHT_QA.json"]["hard_failures"],
            "user_upload_required": False,
            "checkpoint_2_of_3_complete": True,
            "next": "Checkpoint 3 of 3 - Session 1 freeze and complete restore",
        }
        print(json.dumps(console, indent=2))


if __name__ == "__main__":
    main()
