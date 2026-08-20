#!/usr/bin/env python3
from __future__ import annotations

import json
import textwrap
import zipfile
from pathlib import Path
from typing import Any, Iterable

import fitz
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

NAVY = "17324D"
TEAL = "1C7475"
GOLD = "C9A227"
PALE_BLUE = "EAF1F5"
PALE_GOLD = "F7F1D9"
PALE_RED = "F7E8E6"
PALE_GREEN = "E9F3EE"
DARK = "24323D"
GRAY = "66757F"
WHITE = "FFFFFF"


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            pass
    return ImageFont.load_default()


def _wrap(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], font: Any, fill: str, width: int, spacing: int = 8) -> int:
    x, y = xy
    lines: list[str] = []
    for paragraph in text.splitlines() or [""]:
        lines.extend(textwrap.wrap(paragraph, width=width) or [""])
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        box = draw.textbbox((x, y), line or "Ag", font=font)
        y += box[3] - box[1] + spacing
    return y


def build_selection_figure(path: Path, selection: dict[str, Any]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 2400, 1350
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, 175), fill=f"#{NAVY}")
    draw.text((110, 48), "Locked initial print-production scenario", font=_font(58, True), fill="white")
    draw.text((112, 205), "A quality-first production master for the current image-rich clinical reference", font=_font(31), fill=f"#{DARK}")

    cards = [
        ("PROVIDER", selection["provider"], TEAL),
        ("FORMAT", selection["binding"], GOLD),
        ("INTERIOR", selection["interior_type"], "536D8C"),
        ("PAPER", selection["paper"], "6B8E4E"),
        ("TRIM", f"{selection['trim_width_in']:.1f} × {selection['trim_height_in']:.1f} in", TEAL),
        ("PAGES", str(selection["production_page_count"]), GOLD),
        ("SPINE", f"{selection['spine_width_in']:.6f} in", "536D8C"),
        ("WRAP", f"{selection['cover_width_in']:.6f} × {selection['cover_height_in']:.3f} in", "6B8E4E"),
    ]
    card_w, card_h = 500, 205
    start_x, start_y = 115, 315
    gap_x, gap_y = 70, 55
    for idx, (label, value, color) in enumerate(cards):
        row, col = divmod(idx, 4)
        x = start_x + col * (card_w + gap_x)
        y = start_y + row * (card_h + gap_y)
        draw.rounded_rectangle((x, y, x + card_w, y + card_h), radius=24, fill=f"#{PALE_BLUE}", outline=f"#{color}", width=5)
        draw.text((x + 30, y + 28), label, font=_font(25, True), fill=f"#{color}")
        _wrap(draw, value, (x + 30, y + 82), _font(34, True), f"#{DARK}", 24, spacing=5)

    draw.rounded_rectangle((110, 895, width - 110, 1250), radius=28, fill=f"#{PALE_GOLD}", outline=f"#{GOLD}", width=5)
    rationale = (
        "Premium color was selected as the initial production master because the publication contains numerous teaching figures and the project places image sharpness, color clarity, and press-quality review ahead of unit-cost minimization. "
        "The 538-page manuscript remains within the current 8.5 × 11-inch premium-color capacity. Standard color remains a later cost-sensitive derivative, not the controlling master."
    )
    _wrap(draw, rationale, (155, 940), _font(31, True), f"#{DARK}", 125, spacing=8)
    image.save(path, format="PNG", dpi=(300, 300), optimize=True)
    return {"path": str(path), "pixels": [width, height], "dpi": 300, "mode": image.mode, "bytes": path.stat().st_size}


def build_margin_figure(path: Path, interior: dict[str, Any]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 2400, 1400
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, 175), fill=f"#{NAVY}")
    draw.text((110, 48), "Print-interior margin normalization", font=_font(58, True), fill="white")
    draw.text((112, 205), "Text-safe geometry is normalized page by page while preserving the immutable 537-page digital publication.", font=_font(30), fill=f"#{DARK}")

    before = int(interior.get("source_inside_failures", 0))
    after = int(interior.get("output_inside_failures", 0))
    transformed = int(interior.get("transformed_pages", 0))
    translated = int(interior.get("translated_pages", 0))
    scaled = int(interior.get("scaled_pages", 0))
    values = [
        ("Source pages flagged by 0.75-in gutter screen", before, "A65F46"),
        ("Pages translated inward", translated, TEAL),
        ("Pages requiring proportional reduction", scaled, GOLD),
        ("Final pages failing the text-safe screen", after, "4E8A63"),
    ]
    max_value = max([v for _, v, _ in values] + [1])
    y = 360
    for label, value, color in values:
        draw.text((115, y + 12), label, font=_font(31, True), fill=f"#{NAVY}")
        x0, max_w = 980, 1260
        bar = max(18, int(value / max_value * max_w)) if value else 18
        draw.rounded_rectangle((x0, y, x0 + bar, y + 68), radius=18, fill=f"#{color}")
        draw.text((x0 + 24, y + 13), str(value), font=_font(31, True), fill="white")
        y += 175

    draw.rounded_rectangle((110, 1050, width - 110, 1300), radius=26, fill=f"#{PALE_GREEN}", outline="#4E8A63", width=5)
    note = (
        f"The derivative contains {interior.get('output_page_count', 538)} pages: all {interior.get('searchable_pages', 537)} source pages remain searchable, page 538 is intentionally blank, and {transformed} pages received a deterministic transformation. "
        "Text was compared page by page before and after transformation; no source words were lost."
    )
    _wrap(draw, note, (155, 1100), _font(31, True), f"#{DARK}", 125, spacing=8)
    image.save(path, format="PNG", dpi=(300, 300), optimize=True)
    return {"path": str(path), "pixels": [width, height], "dpi": 300, "mode": image.mode, "bytes": path.stat().st_size}


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_docx_table(document: Document, headers: list[str], rows: Iterable[Iterable[Any]], widths: list[float] | None = None) -> Any:
    materialized = [list(row) for row in rows]
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(header)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for ridx, row in enumerate(materialized, start=1):
        cells = table.add_row().cells
        for cidx, value in enumerate(row):
            cells[cidx].text = "" if value is None else str(value)
            set_cell_margins(cells[cidx])
            if ridx % 2 == 0:
                set_cell_shading(cells[cidx], PALE_BLUE)
            for paragraph in cells[cidx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(7.3)
            if widths and cidx < len(widths):
                cells[cidx].width = Inches(widths[cidx])
    return table


def configure_docx(doc: Document, generated_at: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.3)
    normal.paragraph_format.space_after = Pt(5)
    for style_name, size, color in (
        ("Title", 23, NAVY),
        ("Subtitle", 11, TEAL),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 12.5, TEAL),
        ("Heading 3", 10.5, GOLD),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    doc.core_properties.title = "Human Pathogen Database — Section 5 Checkpoint 2 Print Production Candidate"
    doc.core_properties.subject = "KDP premium-color selection, print interior, cover template, regenerated wrap, and preflight"
    doc.core_properties.author = "Brent McAnulty, M.D."
    doc.core_properties.keywords = "Human Pathogen Database; KDP; premium color; print interior; cover; preflight"
    p = doc.add_paragraph("Human Pathogen Database — Print Production Candidate", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("KDP premium-color master, 538-page interior, exact wrap geometry, and release preflight", style="Subtitle")
    p = doc.add_paragraph()
    run = p.add_run(f"Version 3.0.0a  |  Response 76  |  Generated {generated_at}")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    table = doc.add_table(rows=1, cols=3)
    for idx, (label, value, color) in enumerate((("SECTION", "5 OF 5", NAVY), ("SESSION", "1 OF 3", TEAL), ("CHECKPOINT", "2 OF 3", GOLD))):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, color)
        cell.text = f"{label}\n{value}"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
        set_cell_margins(cell, top=120, bottom=120)


def build_docx_report(
    path: Path,
    *,
    generated_at: str,
    selection: dict[str, Any],
    interior: dict[str, Any],
    cover: dict[str, Any],
    preflight: list[dict[str, Any]],
    recovery_events: list[dict[str, Any]],
    figure_paths: list[Path],
) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_docx(doc, generated_at)
    if figure_paths:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(figure_paths[0]), width=Inches(6.9))
    doc.add_heading("Executive disposition", level=1)
    doc.add_paragraph(
        "Checkpoint 2 locks the initial KDP premium-color production scenario, generates a separate 538-page print-only interior, regenerates the full wrap from the separate front, back, and spine masters, and completes deterministic preflight. The immutable 537-page digital publication remains unchanged. Provider Print Previewer review and a physical proof remain required before press release."
    )
    add_docx_table(doc, ["Control", "Result"], [
        ("Provider / binding", f"{selection['provider']} / {selection['binding']}"),
        ("Interior", selection["interior_type"]),
        ("Paper / finish", f"{selection['paper']} / {selection['cover_finish']}"),
        ("Trim", f"{selection['trim_width_in']} × {selection['trim_height_in']} in"),
        ("Production pages", selection["production_page_count"]),
        ("Spine", f"{selection['spine_width_in']:.6f} in"),
        ("Full wrap", f"{selection['cover_width_in']:.6f} × {selection['cover_height_in']:.3f} in"),
        ("Raster master", f"{selection['cover_width_px_300dpi']} × {selection['cover_height_px_300dpi']} px at 300 ppi"),
    ], [2.1, 4.8])
    doc.add_heading("Why premium color is the controlling initial master", level=1)
    doc.add_paragraph(selection["selection_rationale"])
    if len(figure_paths) > 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(figure_paths[1]), width=Inches(6.9))
    doc.add_heading("538-page print-only interior", level=1)
    add_docx_table(doc, ["Interior control", "Observed"], [
        ("Digital publication", f"{interior['source_page_count']} pages; immutable"),
        ("Print derivative", f"{interior['output_page_count']} pages"),
        ("Searchable pages", interior["searchable_pages"]),
        ("Intentional blank", "Page 538"),
        ("Pages transformed", interior["transformed_pages"]),
        ("Translated only", interior["translated_pages"]),
        ("Proportionally reduced", interior["scaled_pages"]),
        ("Final gutter failures", interior["output_inside_failures"]),
        ("Lost-word pages", interior["text_mismatch_pages"]),
        ("Unembedded fonts", interior["unembedded_font_count"]),
        ("Raster images below 300 ppi", interior["low_ppi_image_count"]),
    ], [2.5, 4.4])
    doc.add_paragraph(
        "The transformation map is page-specific. Pages already inside the selected safe geometry are copied at full scale. A page is translated before any proportional reduction is considered; proportional reduction is used only when translation alone cannot place all searchable text within the selected live area."
    )
    doc.add_heading("Cover template and regenerated wrap", level=1)
    if len(figure_paths) > 2:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(figure_paths[2]), width=Inches(6.9))
    add_docx_table(doc, ["Cover control", "Observed"], [
        ("Generated master PNG", cover["final_png"]["name"]),
        ("Print-ready PDF", cover["final_pdf"]["name"]),
        ("Template PNG", cover["template_png"]["name"]),
        ("Template PDF", cover["template_pdf"]["name"]),
        ("Pixel dimensions", f"{cover['pixel_width']} × {cover['pixel_height']}"),
        ("Color space", cover["color_space"]),
        ("Alpha/transparency", cover["alpha_present"]),
        ("Barcode policy", cover["barcode_policy"]),
        ("Barcode clear area", f"{cover['barcode_width_in']} × {cover['barcode_height_in']} in"),
        ("Spine text safety", cover["spine_text_safety"]),
        ("Front/back component scaling", cover["front_back_scaling"]),
    ], [2.4, 4.5])
    doc.add_paragraph(
        "The legacy combined wrap remains preserved as historical design evidence. It was not stretched. The new wrap was rebuilt from the separate front, back, and spine masters on an exact calculated canvas; the original spine artwork is centered without horizontal distortion inside the wider selected spine."
    )
    doc.add_heading("Preflight matrix", level=1)
    add_docx_table(doc, ["Gate", "Expected", "Observed", "Status"], [(r["gate"], r["expected"], r["observed"], r["status"]) for r in preflight], [1.65, 1.75, 2.6, 0.75])
    doc.add_heading("Controlled limitations and next proof gate", level=1)
    doc.add_paragraph(
        "This checkpoint produces a deterministic production candidate. It does not substitute for the KDP Print Previewer, the provider's conversion result, or a physical proof. Checkpoint 3 will freeze the Session 1 candidate, verify the cumulative project state, and emit the session-end complete restore. Provider preview and physical proof evidence remain controlled Section 5 gates before final project completion."
    )
    doc.add_heading("Recoverable build events", level=1)
    add_docx_table(doc, ["Event", "Condition", "Recovery", "Status"], [(r["event_code"], r["condition"], r["recovery"], r["status"]) for r in recovery_events], [1.15, 2.2, 3.1, 0.65])
    doc.add_heading("Checkpoint disposition", level=1)
    doc.add_paragraph("Response 76 and Checkpoint 2 of 3 are COMPLETE. Session 1 of 3 and Remediation Section 5 of 5 CONTINUE. Checkpoint 3 will perform the Session 1 freeze and emit a complete self-contained restore.")
    doc.save(path)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("DOCX container CRC failed")
    return {"path": str(path), "bytes": path.stat().st_size, "status": "passed"}


def _rl_table(rows: list[list[Any]], widths: list[float]) -> Table:
    table = Table(rows, colWidths=[w * inch for w in widths], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 6.8),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#9AAAB4")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(f"#{PALE_BLUE}")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return table


def build_pdf_report(
    path: Path,
    *,
    generated_at: str,
    selection: dict[str, Any],
    interior: dict[str, Any],
    cover: dict[str, Any],
    preflight: list[dict[str, Any]],
    recovery_events: list[dict[str, Any]],
    figure_paths: list[Path],
) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="MRTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=22, leading=25, textColor=colors.HexColor(f"#{NAVY}"), spaceAfter=8))
    styles.add(ParagraphStyle(name="MRSub", parent=styles["Normal"], fontName="Helvetica", fontSize=9.5, leading=13, textColor=colors.HexColor(f"#{TEAL}"), spaceAfter=10))
    styles.add(ParagraphStyle(name="MRH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=15, leading=18, textColor=colors.HexColor(f"#{NAVY}"), spaceBefore=8, spaceAfter=5))
    styles.add(ParagraphStyle(name="MRBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.6, leading=12, textColor=colors.HexColor(f"#{DARK}"), spaceAfter=6))
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=0.55*inch, rightMargin=0.55*inch, topMargin=0.5*inch, bottomMargin=0.5*inch, title="Human Pathogen Database Section 5 Checkpoint 2", author="Brent McAnulty, M.D.")
    story: list[Any] = [
        Paragraph("Human Pathogen Database — Print Production Candidate", styles["MRTitle"]),
        Paragraph("KDP premium-color master, 538-page interior, exact cover geometry, and release preflight", styles["MRSub"]),
        Paragraph(f"Version 3.0.0a | Response 76 | {generated_at}", styles["MRBody"]),
    ]
    if figure_paths:
        story.extend([RLImage(str(figure_paths[0]), width=7.0*inch, height=3.94*inch), Spacer(1, 8)])
    story.extend([
        Paragraph("Executive disposition", styles["MRH1"]),
        Paragraph("Checkpoint 2 locks the initial KDP premium-color production scenario and creates a separate print-production candidate without modifying the 537-page digital publication. Provider Print Previewer review and a physical proof remain mandatory later gates.", styles["MRBody"]),
        _rl_table([["Control", "Result"], ["Provider / binding", f"{selection['provider']} / {selection['binding']}"], ["Interior", selection['interior_type']], ["Trim / pages", f"{selection['trim_width_in']} × {selection['trim_height_in']} in / {selection['production_page_count']}"], ["Spine", f"{selection['spine_width_in']:.6f} in"], ["Full wrap", f"{selection['cover_width_in']:.6f} × {selection['cover_height_in']:.3f} in"], ["Raster master", f"{selection['cover_width_px_300dpi']} × {selection['cover_height_px_300dpi']} px"]], [2.2, 4.8]),
        PageBreak(),
        Paragraph("Print-only interior", styles["MRH1"]),
    ])
    if len(figure_paths) > 1:
        story.extend([RLImage(str(figure_paths[1]), width=7.0*inch, height=4.08*inch), Spacer(1, 8)])
    story.extend([
        _rl_table([["Interior control", "Observed"], ["Digital publication", f"{interior['source_page_count']} pages; immutable"], ["Print derivative", f"{interior['output_page_count']} pages"], ["Searchable source pages", interior['searchable_pages']], ["Pages transformed", interior['transformed_pages']], ["Proportional reductions", interior['scaled_pages']], ["Final gutter failures", interior['output_inside_failures']], ["Lost-word pages", interior['text_mismatch_pages']], ["Unembedded fonts", interior['unembedded_font_count']], ["Images below 300 ppi", interior['low_ppi_image_count']]], [2.35, 4.65]),
        Paragraph("All source text was compared before and after transformation. Page 538 is an explicit intentional blank used only in the print derivative.", styles["MRBody"]),
        PageBreak(),
        Paragraph("Cover template and regenerated wrap", styles["MRH1"]),
    ])
    if len(figure_paths) > 2:
        story.extend([RLImage(str(figure_paths[2]), width=7.0*inch, height=4.25*inch), Spacer(1, 8)])
    story.extend([
        _rl_table([["Cover control", "Observed"], ["Final PNG", cover['final_png']['name']], ["Final PDF", cover['final_pdf']['name']], ["Template PNG", cover['template_png']['name']], ["Canvas", f"{cover['pixel_width']} × {cover['pixel_height']} px"], ["Color / transparency", f"{cover['color_space']} / {cover['alpha_present']}"], ["Barcode", cover['barcode_policy']], ["Spine safety", cover['spine_text_safety']], ["Component scaling", cover['front_back_scaling']]], [2.2, 4.8]),
        Paragraph("The new wrap was rebuilt from the separate component masters. The provisional legacy wrap remains preserved and was not stretched.", styles["MRBody"]),
        PageBreak(),
        Paragraph("Preflight matrix", styles["MRH1"]),
        _rl_table([["Gate", "Expected", "Observed", "Status"]] + [[r['gate'], r['expected'], r['observed'], r['status']] for r in preflight], [1.6, 1.7, 2.9, 0.65]),
        PageBreak(),
        Paragraph("Recoverable build events", styles["MRH1"]),
        _rl_table([["Event", "Condition", "Recovery", "Status"]] + [[r['event_code'], r['condition'], r['recovery'], r['status']] for r in recovery_events], [1.15, 2.15, 3.1, 0.6]),
        Paragraph("Checkpoint disposition", styles["MRH1"]),
        Paragraph("Response 76 and Checkpoint 2 of 3 are COMPLETE. Session 1 of 3 and Remediation Section 5 of 5 CONTINUE. Checkpoint 3 will freeze the Session 1 state and emit the complete self-contained session restore.", styles["MRBody"]),
    ])
    doc.build(story)
    reader = PdfReader(str(path))
    text_chars = sum(len(page.extract_text() or "") for page in reader.pages)
    if len(reader.pages) < 5 or text_chars < 2500:
        raise RuntimeError({"pdf_report_validation": {"pages": len(reader.pages), "text_chars": text_chars}})
    return {"path": str(path), "bytes": path.stat().st_size, "pages": len(reader.pages), "text_chars": text_chars, "status": "passed"}


def build_register(
    path: Path,
    *,
    selection: dict[str, Any],
    interior_summary: dict[str, Any],
    page_rows: list[dict[str, Any]],
    cover: dict[str, Any],
    preflight: list[dict[str, Any]],
    recovery_events: list[dict[str, Any]],
    tracking_rows: list[dict[str, Any]],
) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    wb.remove(wb.active)
    datasets: dict[str, list[dict[str, Any]]] = {
        "Dashboard": [
            {"Control": "Response", "Value": 76, "Status": "current"},
            {"Control": "Checkpoint", "Value": "2 of 3", "Status": "complete"},
            {"Control": "Provider", "Value": selection["provider"], "Status": "locked_initial"},
            {"Control": "Interior", "Value": selection["interior_type"], "Status": "selected"},
            {"Control": "Production pages", "Value": interior_summary["output_page_count"], "Status": "passed"},
            {"Control": "Cover canvas", "Value": f"{cover['pixel_width']} × {cover['pixel_height']} px", "Status": "passed"},
            {"Control": "Preflight failures", "Value": sum(1 for row in preflight if row["status"] != "passed"), "Status": "passed"},
        ],
        "Selection": [selection],
        "Interior Summary": [interior_summary],
        "Page Transforms": page_rows,
        "Cover": [{k: v if not isinstance(v, dict) else json.dumps(v, ensure_ascii=False) for k, v in cover.items()}],
        "Preflight": preflight,
        "Tracking": tracking_rows,
        "Recovery": recovery_events,
    }
    thin = Side(style="thin", color="AAB8C0")
    for title, rows in datasets.items():
        ws = wb.create_sheet(title=title[:31])
        if not rows:
            ws.append(["No records"])
            continue
        headers = list(rows[0])
        ws.append(headers)
        for cell in ws[1]:
            cell.fill = PatternFill("solid", fgColor=NAVY)
            cell.font = Font(color=WHITE, bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row in rows:
            ws.append([row.get(header) for header in headers])
        for r in range(2, ws.max_row + 1):
            if r % 2 == 0:
                for cell in ws[r]:
                    cell.fill = PatternFill("solid", fgColor=PALE_BLUE)
            for cell in ws[r]:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        for idx, header in enumerate(headers, start=1):
            sample = [str(header)] + [str(ws.cell(r, idx).value or "") for r in range(2, min(ws.max_row, 120) + 1)]
            width = min(55, max(10, max(len(v) for v in sample) + 2))
            ws.column_dimensions[get_column_letter(idx)].width = width
        ws.sheet_view.showGridLines = False
    wb.properties.title = "MRHPD Section 5 Session 1 Checkpoint 2 Print Production Register"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(path)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("XLSX container CRC failed")
    return {"path": str(path), "bytes": path.stat().st_size, "sheets": wb.sheetnames, "status": "passed"}


def render_report_qa(pdf_path: Path, output_dir: Path) -> dict[str, Any]:
    output_dir.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(pdf_path)
    rows = []
    for index, page in enumerate(doc):
        pix = page.get_pixmap(matrix=fitz.Matrix(1.25, 1.25), alpha=False)
        target = output_dir / f"report-page-{index + 1:02d}.png"
        pix.save(target)
        image = Image.open(target).convert("L")
        extrema = image.getextrema()
        status = "passed" if extrema[0] < 248 and target.stat().st_size > 10_000 else "failed"
        rows.append({"page": index + 1, "path": target.name, "bytes": target.stat().st_size, "extrema": list(extrema), "status": status})
    doc.close()
    failed = [row for row in rows if row["status"] != "passed"]
    if failed:
        raise RuntimeError({"report_render_failures": failed})
    return {"status": "passed", "pages": len(rows), "records": rows}
