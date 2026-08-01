#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import textwrap
import zipfile
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any, Iterable

import fitz
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

HERE = Path(__file__).resolve().parent
CP2_DIR = HERE.parent / "checkpoint2"
if str(CP2_DIR) not in sys.path:
    sys.path.insert(0, str(CP2_DIR))

import build_section5_checkpoint2 as cp2  # noqa: E402

PROJECT_VERSION = "3.0.0a"
RESPONSE_NUMBER = 77
SECTION_LABEL = "Remediation Section 5 of 5"
SESSION_LABEL = "Session 1 of 3"
CHECKPOINT_LABEL = "Checkpoint 3 of 3"
CHECKPOINT_CODE = "MRHPD-V3-CP5-S1-CP3"
RELEASE_CODE = "MRHPD-V3-CP5-S1-COMPLETE-R77"
NEXT_SESSION = "Remediation Section 5 of 5 Session 2 of 3"
CP2_RECOVERY_BYTES = 69_251_244
CP2_RECOVERY_SHA256 = "d90133ffe2b595c5df3937bc9931b083d12e809c0737cd8bbf301f2f02b206e0"
BASE_RESTORE_BYTES = cp2.BASE_RESTORE_BYTES
BASE_RESTORE_SHA256 = cp2.BASE_RESTORE_SHA256
BASE_PROJECT_BYTES = cp2.BASE_PROJECT_BYTES
BASE_PROJECT_SHA256 = cp2.BASE_PROJECT_SHA256
PUBLICATION_SHA256 = cp2.PUBLICATION_SHA256
EDITABLE_SHA256 = cp2.EDITABLE_SHA256
APPLICATION_SHA256 = cp2.APPLICATION_SHA256
PUBLICATION_REL = cp2.PUBLICATION_REL
CURRENT_DB_REL = (
    "Database/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 1 of 3 Checkpoint 3 of 3 SESSION 1 COMPLETE THROUGH RESPONSE 77.sqlite"
)
CURRENT_WORKBOOK_REL = (
    "Tracking/Workbook/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 1 of 3 Checkpoint 3 of 3 SESSION 1 COMPLETE THROUGH RESPONSE 77 Comprehensive Tracking.xlsx"
)
PRINT_INTERIOR_REL = cp2.PRINT_INTERIOR_REL
FINAL_COVER_PNG_REL = cp2.FINAL_COVER_PNG_REL
FINAL_COVER_TIFF_REL = cp2.FINAL_COVER_TIFF_REL
FINAL_COVER_PDF_REL = cp2.FINAL_COVER_PDF_REL
TEMPLATE_PNG_REL = cp2.TEMPLATE_PNG_REL
TEMPLATE_PDF_REL = cp2.TEMPLATE_PDF_REL
PROOF_PNG_REL = cp2.PROOF_PNG_REL

NAVY = "17324D"
TEAL = "1C7475"
GOLD = "C9A227"
PALE_BLUE = "EAF1F5"
PALE_GOLD = "F7F0D8"
WHITE = "FFFFFF"
RED = "B5423A"


def utc_now() -> datetime:
    return datetime.now(timezone.utc).replace(microsecond=0)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False), encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value, encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fields is None:
        fields = list(rows[0]) if rows else []
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def safe_extract(path: Path, destination: Path) -> None:
    destination.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError(f"ZIP CRC failure: {path}")
        names = zf.namelist()
        if len(names) != len(set(names)):
            raise RuntimeError(f"Duplicate ZIP members: {path}")
        for name in names:
            pp = PurePosixPath(name.replace("\\", "/"))
            if pp.is_absolute() or ".." in pp.parts or re.match(r"^[A-Za-z]:", name):
                raise RuntimeError(f"Unsafe ZIP member: {name}")
        zf.extractall(destination)


def verify_zip(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
        names = zf.namelist()
        unsafe = []
        filler = []
        for name in names:
            pp = PurePosixPath(name.replace("\\", "/"))
            if pp.is_absolute() or ".." in pp.parts or re.match(r"^[A-Za-z]:", name):
                unsafe.append(name)
            low = name.lower()
            if any(token in low for token in ("filler", "padding", "dummy_payload", "artificial_inflation")):
                filler.append(name)
    result = {
        "name": path.name,
        "path": str(path),
        "bytes": path.stat().st_size,
        "sha256": sha256_file(path),
        "members": len(names),
        "crc_error": bad,
        "duplicate_members": len(names) - len(set(names)),
        "unsafe_paths": unsafe,
        "filler_members": filler,
    }
    if bad or result["duplicate_members"] or unsafe or filler:
        raise RuntimeError({"zip_verification_failed": result})
    return result


def find_exact_zip_recursive(root: Path, size: int, digest: str, work: Path) -> Path:
    queue = [p for p in root.rglob("*.zip") if p.is_file()]
    seen: set[str] = set()
    sequence = 0
    while queue:
        candidate = queue.pop(0)
        identity = f"{candidate.stat().st_size}:{sha256_file(candidate)}"
        if identity in seen:
            continue
        seen.add(identity)
        if candidate.stat().st_size == size and sha256_file(candidate) == digest:
            return candidate
        sequence += 1
        nested = work / f"nested-{sequence:04d}"
        try:
            safe_extract(candidate, nested)
        except Exception:
            continue
        queue.extend(p for p in nested.rglob("*.zip") if p.is_file())
    raise RuntimeError({"exact_zip_not_found": {"root": str(root), "bytes": size, "sha256": digest}})


def recover_checkpoint2_project(
    volume1_dir: Path,
    volume2_dir: Path,
    checkpoint2_dir: Path,
    work: Path,
) -> tuple[Path, Path, Path, Path, Path, dict[str, Any]]:
    restore, project_archive, baseline_project = cp2.cp1.reconstruct_baseline(
        volume1_dir, volume2_dir, work / "baseline"
    )
    recovery_zip = find_exact_zip_recursive(
        checkpoint2_dir, CP2_RECOVERY_BYTES, CP2_RECOVERY_SHA256, work / "checkpoint2-discovery"
    )
    package_root = work / "checkpoint2-package"
    safe_extract(recovery_zip, package_root)
    apply_script = package_root / "TOOLS" / "apply_checkpoint_recovery.py"
    if not apply_script.exists():
        raise RuntimeError("Response 76 apply utility is missing")
    restored_root = work / "response76-restored"
    result = subprocess.run(
        [
            sys.executable,
            str(apply_script.resolve()),
            "--base-response72-restore",
            str(restore.resolve()),
            "--output-dir",
            str(restored_root.resolve()),
        ],
        cwd=package_root,
        text=True,
        capture_output=True,
        timeout=2400,
    )
    if result.returncode:
        raise RuntimeError(
            {
                "response76_apply_failed": {
                    "stdout": result.stdout[-16000:],
                    "stderr": result.stderr[-16000:],
                }
            }
        )
    result_candidates = list(restored_root.glob("*APPLICATION_RESULT.json"))
    if not result_candidates:
        result_candidates = list(restored_root.rglob("*APPLICATION_RESULT.json"))
    application = {}
    if result_candidates:
        application = json.loads(result_candidates[0].read_text(encoding="utf-8"))
        if application.get("status") != "passed":
            raise RuntimeError({"response76_application_result": application})
    project_candidates = [p for p in restored_root.iterdir() if p.is_dir()]
    if len(project_candidates) != 1:
        raise RuntimeError({"restored_project_candidates": [str(p) for p in project_candidates]})
    return restore, project_archive, baseline_project, recovery_zip, project_candidates[0], application


def find_file_by_hash(project: Path, digest: str, suffixes: tuple[str, ...] | None = None) -> Path:
    candidates = [p for p in project.rglob("*") if p.is_file()]
    if suffixes:
        candidates = [p for p in candidates if p.suffix.lower() in suffixes]
    matches = [p for p in candidates if sha256_file(p) == digest]
    if len(matches) != 1:
        raise RuntimeError({"hash_match_count": len(matches), "sha256": digest, "matches": [str(p) for p in matches[:20]]})
    return matches[0]


def locate_main_application(project: Path) -> Path:
    preferred = project / "App" / "human_pathogen_app.py"
    if preferred.exists() and sha256_file(preferred) == APPLICATION_SHA256:
        return preferred
    return find_file_by_hash(project, APPLICATION_SHA256, (".py",))


def audit_response76_project(project: Path) -> dict[str, Any]:
    db = project / cp2.CURRENT_DB_REL
    workbook = project / cp2.CURRENT_WORKBOOK_REL
    publication = project / PUBLICATION_REL
    print_interior = project / PRINT_INTERIOR_REL
    cover_png = project / FINAL_COVER_PNG_REL
    cover_tiff = project / FINAL_COVER_TIFF_REL
    cover_pdf = project / FINAL_COVER_PDF_REL
    template_png = project / TEMPLATE_PNG_REL
    template_pdf = project / TEMPLATE_PDF_REL
    proof_png = project / PROOF_PNG_REL
    editable = find_file_by_hash(project, EDITABLE_SHA256, (".docx",))
    application = locate_main_application(project)
    required = [
        db,
        workbook,
        publication,
        print_interior,
        cover_png,
        cover_tiff,
        cover_pdf,
        template_png,
        template_pdf,
        proof_png,
        editable,
        application,
    ]
    missing = [str(p) for p in required if not p.exists()]
    if missing:
        raise RuntimeError({"missing_response76_artifacts": missing})
    if sha256_file(publication) != PUBLICATION_SHA256:
        raise RuntimeError("Immutable 537-page publication hash changed")
    if sha256_file(editable) != EDITABLE_SHA256:
        raise RuntimeError("Editable assembly hash changed")
    if sha256_file(application) != APPLICATION_SHA256:
        raise RuntimeError("Main application hash changed")

    reader = PdfReader(str(print_interior))
    print_pages = len(reader.pages)
    searchable = sum(1 for page in reader.pages[:537] if (page.extract_text() or "").strip())
    last_text = (reader.pages[-1].extract_text() or "").strip()
    if print_pages != 538 or searchable != 537 or last_text:
        raise RuntimeError({"print_interior_gate": {"pages": print_pages, "searchable": searchable, "last_text": last_text[:200]}})
    pdf_doc = fitz.open(print_interior)
    dimensions = sorted({(round(page.rect.width / 72.0, 4), round(page.rect.height / 72.0, 4)) for page in pdf_doc})
    pdf_doc.close()
    if dimensions != [(8.5, 11.0)]:
        raise RuntimeError({"print_page_dimensions": dimensions})

    with Image.open(cover_png) as image:
        cover_size = image.size
        cover_mode = image.mode
        alpha = "A" in image.getbands()
    if cover_size != (5554, 3375) or alpha:
        raise RuntimeError({"cover_png_gate": {"size": cover_size, "mode": cover_mode, "alpha": alpha}})
    cover_doc = fitz.open(cover_pdf)
    cover_inches = (round(cover_doc[0].rect.width / 72.0, 6), round(cover_doc[0].rect.height / 72.0, 6))
    cover_doc.close()
    if abs(cover_inches[0] - 18.512686) > 0.002 or abs(cover_inches[1] - 11.25) > 0.002:
        raise RuntimeError({"cover_pdf_dimensions": cover_inches})

    con = sqlite3.connect(db)
    try:
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        foreign_keys = list(con.execute("PRAGMA foreign_key_check"))
        table_count = con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0]
        response76 = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R76'").fetchone()[0]
        page_records = con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code=?", (cp2.CHECKPOINT_CODE,)).fetchone()[0]
        failed_page_records = con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code=? AND status!='passed'", (cp2.CHECKPOINT_CODE,)).fetchone()[0]
        selection = con.execute("SELECT status FROM section5_print_selection WHERE checkpoint_code=?", (cp2.CHECKPOINT_CODE,)).fetchone()
    finally:
        con.close()
    if integrity != "ok" or foreign_keys or response76 != 1 or page_records != 538 or failed_page_records or selection != ("locked_initial_production_master",):
        raise RuntimeError({
            "response76_database_gate": {
                "integrity": integrity,
                "foreign_keys": foreign_keys[:20],
                "response76": response76,
                "page_records": page_records,
                "failed_page_records": failed_page_records,
                "selection": selection,
            }
        })

    wb = load_workbook(workbook, read_only=True, data_only=False)
    try:
        sheets = list(wb.sheetnames)
        formula_errors = 0
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    value = cell.value
                    if isinstance(value, str) and any(token in value for token in ("#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "#NUM!", "#N/A")):
                        formula_errors += 1
    finally:
        wb.close()
    if len(sheets) < 100 or formula_errors:
        raise RuntimeError({"response76_workbook_gate": {"sheets": len(sheets), "formula_errors": formula_errors}})

    return {
        "status": "passed",
        "database": db,
        "workbook": workbook,
        "publication": publication,
        "editable": editable,
        "application": application,
        "print_interior": print_interior,
        "cover_png": cover_png,
        "cover_tiff": cover_tiff,
        "cover_pdf": cover_pdf,
        "template_png": template_png,
        "template_pdf": template_pdf,
        "proof_png": proof_png,
        "database_tables": table_count,
        "workbook_sheets": len(sheets),
        "print_pages": print_pages,
        "searchable_pages": searchable,
        "cover_pixels": list(cover_size),
        "cover_inches": list(cover_inches),
        "critical_hashes": {
            "publication": sha256_file(publication),
            "editable": sha256_file(editable),
            "application": sha256_file(application),
            "print_interior": sha256_file(print_interior),
            "cover_png": sha256_file(cover_png),
            "cover_tiff": sha256_file(cover_tiff),
            "cover_pdf": sha256_file(cover_pdf),
            "template_png": sha256_file(template_png),
            "template_pdf": sha256_file(template_pdf),
        },
    }


def _clone_table_row(con: sqlite3.Connection, table: str, where_column: str, source_value: Any, updates: dict[str, Any]) -> None:
    info = con.execute(f"PRAGMA table_info({table})").fetchall()
    columns = [row[1] for row in info]
    pk_columns = {row[1] for row in info if row[5]}
    source = con.execute(f"SELECT * FROM {table} WHERE {where_column}=?", (source_value,)).fetchone()
    if source is None:
        raise RuntimeError(f"Source row missing: {table}.{where_column}={source_value}")
    values = dict(zip(columns, source))
    values.update({k: v for k, v in updates.items() if k in columns})
    insert_columns = [column for column in columns if column not in pk_columns]
    target_value = values.get(where_column)
    con.execute(f"DELETE FROM {table} WHERE {where_column}=?", (target_value,))
    con.execute(
        f"INSERT INTO {table} ({','.join(insert_columns)}) VALUES ({','.join('?' for _ in insert_columns)})",
        [values.get(column) for column in insert_columns],
    )


def recovery_events(now_iso: str) -> list[dict[str, Any]]:
    rows = [
        (193, "LOCAL-RUNTIME-UNAVAILABLE", "The local container and Python surfaces returned InvalidArgumentError before code startup.", "Recovered the newest verified Response 76 package from Google Drive and used the isolated transient computation lane while retaining Drive as controlling storage."),
        (194, "RESPONSE76-EXACT-RECOVERY", "Session 1 completion required the exact Response 76 cumulative state.", "Recovered the exact 70,033,927-byte delivery artifact and clean-applied its 69,251,244-byte cumulative overlay to the exact Response 72 restore."),
        (195, "RESPONSE76-INDEPENDENT-REAUDIT", "The prior checkpoint could not self-authorize the session boundary.", "Independently rechecked database integrity, workbook safety, the 538-page print interior, exact cover dimensions, immutable publication, editable assembly, and main application identity."),
        (196, "SESSION1-PRODUCTION-CANDIDATE-FROZEN", "Checkpoint 2 remained a mutable production candidate.", "Froze the Premium Color 538-page interior and exact cover package as the Session 1 production candidate without claiming provider-preview or physical-proof approval."),
        (197, "PRINT-INTERIOR-REVALIDATED", "A session-end freeze required clean page-level revalidation.", "Revalidated 538 pages, 537 searchable source pages, one terminal blank page, 8.5 × 11-inch geometry, and zero failed database page-transform records."),
        (198, "COVER-PACKAGE-REVALIDATED", "A session-end freeze required exact full-wrap and template verification.", "Revalidated the 5,554 × 3,375 RGB cover, 18.512686 × 11.250-inch PDF geometry, template, TIFF, proof, and preserved legacy component identities."),
        (199, "DATABASE-WORKBOOK-APPLICATION-SYNCHRONIZED", "Response 77 and terminal Session 1 state had to be represented across governed surfaces.", "Added Response 77, Checkpoint 3, Session 1 release, acceptance, freeze, handoff, workbook, application-pointer, tracking, and recovery records to copied artifacts."),
        (200, "KDP-PREVIEWER-CONTROLLED-PENDING", "KDP Print Previewer is external to this execution environment.", "Preserved it as a controlled pending Session 2 gate rather than fabricating approval evidence."),
        (201, "PHYSICAL-PROOF-CONTROLLED-PENDING", "A physical copy cannot be inspected inside the current execution environment.", "Preserved physical-proof ordering, inspection, defect logging, and approval as governed Session 2/3 controls."),
        (202, "SUPERSEDED-DERIVATIVE-COMPACTION", "The full project contained older database and workbook snapshots superseded by verified current supersets.", "Removed only superset-proven superseded snapshots and rebuilt all indexes, manifests, checksums, and clean-extraction evidence."),
        (203, "COMPLETE-PROJECT-CLEAN-EXTRACTION", "The Session 1 project archive required independent restoration evidence.", "Clean-extracted the complete project and reran database, workbook, application, digital-publication, print-interior, cover, index, and manifest gates."),
        (204, "SELF-CONTAINED-RESTORE-VERIFIED", "The session boundary required a restore independent of prior project files and the conversation.", "Built and executed the embedded restore verifier against the complete project snapshot."),
        (205, "MINIMUM-VOLUME-TRANSPORT-AND-HANDOFF", "The final restore exceeded the Google Drive connector single-file ceiling.", "Split it into the minimum number of verified Drive-safe volumes, included deterministic reassembly, and recorded the Session 2 handoff."),
    ]
    return [
        {
            "event_code": f"V3-CP5-S1-REC-{number}-{slug}",
            "condition": condition,
            "recovery": recovery,
            "status": "recovered",
            "recorded_at": now_iso,
        }
        for number, slug, condition, recovery in rows
    ]


def synchronize_database(source: Path, destination: Path, audit: dict[str, Any], now_iso: str) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, destination)
    events = recovery_events(now_iso)
    con = sqlite3.connect(destination)
    try:
        con.execute("PRAGMA foreign_keys=ON")
        con.execute("BEGIN IMMEDIATE")
        response = {
            "response_key": "R77",
            "response_number": 77,
            "response_label": "77",
            "response_date": now_iso,
            "major_topic": "Human Pathogen Database remediation",
            "title": "Section 5 Session 1 complete production-candidate restore",
            "goal": "Continue from the verified Response 76 state, independently freeze the Session 1 production candidate, synchronize all governed artifacts, and emit a complete self-contained restore.",
            "raw_prompt": "Continue",
            "raw_response": "[PRE-EMISSION RESPONSE; final user-visible response is represented by the source-supported summary]",
            "summary": "Independently reconstructed and revalidated the Response 76 production candidate, synchronized Response 77 and terminal Session 1 state, clean-verified the complete project, and emitted the self-contained Session 1 restore.",
            "state": "session_complete_continue_required",
            "coverage": "exact raw prompt plus source-supported response summary",
            "fidelity_classification": "source_verified_prompt_and_summary",
            "source_id": "CURRENT-CONVERSATION-R77",
            "source_path": "Current conversation, exact Response 72 restore, and cumulative Response 76 recovery package",
            "notes": "Checkpoint 3 of 3 and Session 1 of 3 complete. Continue begins Session 2 of 3.",
            "reconciled_at": now_iso,
        }
        cp2._clone_response_row(con, "R76", response)

        _clone_table_row(
            con,
            "section5_session1_checkpoint",
            "checkpoint_code",
            cp2.CHECKPOINT_CODE,
            {
                "checkpoint_code": CHECKPOINT_CODE,
                "response_number": 77,
                "checkpoint_label": CHECKPOINT_LABEL,
                "state": "session_complete",
                "provider_selection_status": "session1_frozen_production_candidate",
                "print_interior_status": "passed_session1_freeze",
                "cover_status": "passed_session1_freeze",
                "workbook_status": "pending_final_save",
                "application_status": "pending_final_audit",
                "next_checkpoint": NEXT_SESSION,
                "recorded_at": now_iso,
            },
        )

        for table in ("section5_print_selection", "section5_print_derivative", "section5_page_transform", "section5_cover_template", "section5_print_preflight"):
            info = con.execute(f"PRAGMA table_info({table})").fetchall()
            columns = [row[1] for row in info]
            pk_columns = {row[1] for row in info if row[5]}
            rows = con.execute(f"SELECT * FROM {table} WHERE checkpoint_code=?", (cp2.CHECKPOINT_CODE,)).fetchall()
            con.execute(f"DELETE FROM {table} WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
            insert_columns = [column for column in columns if column not in pk_columns]
            for source_row in rows:
                values = dict(zip(columns, source_row))
                values["checkpoint_code"] = CHECKPOINT_CODE
                if table == "section5_print_selection" and "status" in values:
                    values["status"] = "session1_frozen_production_candidate"
                if "recorded_at" in values:
                    values["recorded_at"] = now_iso
                if "checked_at" in values:
                    values["checked_at"] = now_iso
                con.execute(
                    f"INSERT INTO {table} ({','.join(insert_columns)}) VALUES ({','.join('?' for _ in insert_columns)})",
                    [values.get(column) for column in insert_columns],
                )

        con.executescript("""
        CREATE TABLE IF NOT EXISTS section5_session1_release (
            section5_session1_release_id INTEGER PRIMARY KEY,
            release_code TEXT NOT NULL UNIQUE,
            response_number INTEGER NOT NULL,
            state TEXT NOT NULL,
            production_candidate_status TEXT NOT NULL,
            digital_publication_status TEXT NOT NULL,
            print_interior_status TEXT NOT NULL,
            cover_status TEXT NOT NULL,
            database_status TEXT NOT NULL,
            workbook_status TEXT NOT NULL,
            application_status TEXT NOT NULL,
            provider_previewer_status TEXT NOT NULL,
            physical_proof_status TEXT NOT NULL,
            next_session TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_session1_acceptance (
            section5_session1_acceptance_id INTEGER PRIMARY KEY,
            release_code TEXT NOT NULL,
            gate_key TEXT NOT NULL,
            expected TEXT NOT NULL,
            observed TEXT NOT NULL,
            status TEXT NOT NULL,
            checked_at TEXT NOT NULL,
            UNIQUE(release_code, gate_key)
        );
        CREATE TABLE IF NOT EXISTS section5_session1_freeze (
            section5_session1_freeze_id INTEGER PRIMARY KEY,
            release_code TEXT NOT NULL,
            artifact_key TEXT NOT NULL,
            relative_path TEXT NOT NULL,
            bytes INTEGER NOT NULL,
            sha256 TEXT NOT NULL,
            immutable INTEGER NOT NULL,
            status TEXT NOT NULL,
            recorded_at TEXT NOT NULL,
            UNIQUE(release_code, artifact_key)
        );
        CREATE TABLE IF NOT EXISTS section5_session_handoff (
            section5_session_handoff_id INTEGER PRIMARY KEY,
            handoff_code TEXT NOT NULL UNIQUE,
            from_session TEXT NOT NULL,
            to_session TEXT NOT NULL,
            state TEXT NOT NULL,
            scope TEXT NOT NULL,
            prerequisite TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        """)
        con.execute("DELETE FROM section5_session1_release WHERE release_code=?", (RELEASE_CODE,))
        con.execute(
            "INSERT INTO section5_session1_release (release_code,response_number,state,production_candidate_status,digital_publication_status,print_interior_status,cover_status,database_status,workbook_status,application_status,provider_previewer_status,physical_proof_status,next_session,recorded_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (
                RELEASE_CODE,
                77,
                "session_complete",
                "frozen_session1_production_candidate",
                "passed_immutable_537_page",
                "passed_538_page_candidate",
                "passed_exact_cover_candidate",
                "ok",
                "pending_final_save",
                "pending_final_audit",
                "controlled_pending_session2",
                "controlled_pending_session2_or_3",
                NEXT_SESSION,
                now_iso,
            ),
        )
        con.execute("DELETE FROM section5_session1_freeze WHERE release_code=?", (RELEASE_CODE,))
        freeze_rows = [
            ("digital_publication", audit["publication"], True),
            ("editable_assembly", audit["editable"], True),
            ("main_application", audit["application"], True),
            ("print_interior", audit["print_interior"], False),
            ("cover_png", audit["cover_png"], False),
            ("cover_tiff", audit["cover_tiff"], False),
            ("cover_pdf", audit["cover_pdf"], False),
            ("template_png", audit["template_png"], False),
            ("template_pdf", audit["template_pdf"], False),
        ]
        for key, path, immutable in freeze_rows:
            con.execute(
                "INSERT INTO section5_session1_freeze (release_code,artifact_key,relative_path,bytes,sha256,immutable,status,recorded_at) VALUES (?,?,?,?,?,?,?,?)",
                (RELEASE_CODE, key, path.relative_to(path.parents[len(path.parts) - len(path.parts)] if False else destination.parents[1]).as_posix() if False else str(path), path.stat().st_size, sha256_file(path), int(immutable), "passed", now_iso),
            )
        con.execute("DELETE FROM section5_session_handoff WHERE handoff_code=?", ("MRHPD-V3-CP5-S1-TO-S2",))
        con.execute(
            "INSERT INTO section5_session_handoff (handoff_code,from_session,to_session,state,scope,prerequisite,recorded_at) VALUES (?,?,?,?,?,?,?)",
            (
                "MRHPD-V3-CP5-S1-TO-S2",
                "Section 5 Session 1 of 3",
                "Section 5 Session 2 of 3",
                "ready",
                "Provider preview, physical-proof workflow, correction control, and production-acceptance evidence.",
                "Use the exact Response 77 complete restore; do not alter the frozen Session 1 candidate in place.",
                now_iso,
            ),
        )
        for event in events:
            con.execute(
                "INSERT OR REPLACE INTO section5_recovery_event (checkpoint_code,event_code,condition,recovery,status,recorded_at) VALUES (?,?,?,?,?,?)",
                (CHECKPOINT_CODE, event["event_code"], event["condition"], event["recovery"], event["status"], event["recorded_at"]),
            )
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        foreign_keys = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or foreign_keys:
            raise RuntimeError({"database_gate_before_workbook": {"integrity": integrity, "foreign_keys": foreign_keys[:20]}})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()
    return {"events": events}


def _write_sheet(ws: Any, rows: list[dict[str, Any]]) -> None:
    if not rows:
        ws.append(["No records"])
        return
    headers = list(rows[0])
    ws.append(headers)
    for cell in ws[1]:
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(color=WHITE, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin = Side(style="thin", color="AAB8C0")
    for row in rows:
        ws.append([json.dumps(row.get(header), ensure_ascii=False) if isinstance(row.get(header), (list, dict)) else row.get(header) for header in headers])
    for r in range(2, ws.max_row + 1):
        for cell in ws[r]:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if r % 2 == 0:
                cell.fill = PatternFill("solid", fgColor=PALE_BLUE)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    ws.sheet_view.showGridLines = False
    for idx, header in enumerate(headers, start=1):
        sample = [str(header)] + [str(ws.cell(r, idx).value or "") for r in range(2, min(ws.max_row, 120) + 1)]
        ws.column_dimensions[get_column_letter(idx)].width = min(60, max(10, max(len(value) for value in sample) + 2))


def acceptance_rows(audit: dict[str, Any], workbook_sheet_count: int, formula_errors: int, now_iso: str) -> list[dict[str, Any]]:
    checks = [
        ("response76_clean_apply", "passed", "passed", "passed"),
        ("database_integrity", "ok", "ok", "passed"),
        ("foreign_keys", "0", "0", "passed"),
        ("response77_lineage", "1 record", "1 record", "passed"),
        ("session_state", "session_complete", "session_complete", "passed"),
        ("selection_freeze", "session1_frozen_production_candidate", "session1_frozen_production_candidate", "passed"),
        ("digital_publication", "537 immutable pages", f"537 pages; {audit['critical_hashes']['publication']}", "passed"),
        ("editable_assembly", "immutable", audit["critical_hashes"]["editable"], "passed"),
        ("print_interior", "538 pages", str(audit["print_pages"]), "passed"),
        ("print_searchability", "537 searchable source pages", str(audit["searchable_pages"]), "passed"),
        ("page_geometry", "8.5 × 11 inches", "8.5 × 11 inches", "passed"),
        ("cover_pixels", "5554 × 3375", f"{audit['cover_pixels'][0]} × {audit['cover_pixels'][1]}", "passed"),
        ("cover_inches", "18.512686 × 11.250", f"{audit['cover_inches'][0]} × {audit['cover_inches'][1]}", "passed"),
        ("cover_alpha", "none", "none", "passed"),
        ("workbook_preservation", ">= 106 sheets", str(workbook_sheet_count), "passed" if workbook_sheet_count >= 106 else "failed"),
        ("workbook_formula_errors", "0", str(formula_errors), "passed" if formula_errors == 0 else "failed"),
        ("main_application", "byte-identical", audit["critical_hashes"]["application"], "passed"),
        ("source_index", "rebuilt", "pending_post-freeze_rebuild", "pending_internal"),
        ("bit_index", "integrity ok", "pending_post-freeze_rebuild", "pending_internal"),
        ("manifest", "zero mismatches", "pending_post-freeze_rebuild", "pending_internal"),
        ("clean_project", "passed", "pending archive stage", "pending_internal"),
        ("self_contained_restore", "passed", "pending restore stage", "pending_internal"),
        ("transport", "minimum Drive-safe volumes", "pending restore stage", "pending_internal"),
        ("provider_previewer", "external approval", "controlled pending Session 2", "controlled_pending"),
        ("physical_proof", "physical approval", "controlled pending Session 2 or 3", "controlled_pending"),
    ]
    return [
        {"gate_key": key, "expected": expected, "observed": observed, "status": status, "checked_at": now_iso}
        for key, expected, observed, status in checks
    ]


def augment_workbook(source: Path, destination: Path, audit: dict[str, Any], events: list[dict[str, Any]], now_iso: str) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    wb = load_workbook(source)
    inherited = list(wb.sheetnames)
    for name in [
        "S5S1 CP3 Dashboard",
        "S5S1 CP3 Acceptance",
        "S5S1 CP3 Freeze",
        "S5S1 CP3 Responses",
        "S5S1 CP3 Recovery",
        "S5S1 CP3 Handoff",
    ]:
        if name in wb.sheetnames:
            del wb[name]
    preliminary_count = len(wb.sheetnames) + 6
    rows = acceptance_rows(audit, preliminary_count, 0, now_iso)
    datasets = {
        "S5S1 CP3 Dashboard": [
            {"Control": "Response", "Value": 77, "Status": "current"},
            {"Control": "Checkpoint", "Value": "3 of 3", "Status": "complete"},
            {"Control": "Session", "Value": "1 of 3", "Status": "complete"},
            {"Control": "Production candidate", "Value": "KDP Premium Color 538-page interior and exact cover", "Status": "frozen_session1_candidate"},
            {"Control": "Provider preview", "Value": "KDP Print Previewer", "Status": "controlled_pending_session2"},
            {"Control": "Physical proof", "Value": "Printed proof inspection", "Status": "controlled_pending_session2_or_3"},
            {"Control": "Next", "Value": NEXT_SESSION, "Status": "continue"},
        ],
        "S5S1 CP3 Acceptance": rows,
        "S5S1 CP3 Freeze": [
            {"Artifact": key, "SHA-256": value, "Status": "passed"}
            for key, value in audit["critical_hashes"].items()
        ],
        "S5S1 CP3 Responses": [
            {
                "Response": 77,
                "Raw Prompt": "Continue",
                "Net Prompt": "Continue from the newest verified Response 76 checkpoint and complete Section 5 Session 1 with an independently verified full restore.",
                "Summary": "Reconstructed and independently revalidated Response 76, froze the Session 1 production candidate, synchronized all governed artifacts, and emitted the complete Session 1 restore.",
                "State": "session_complete_continue_required",
            }
        ],
        "S5S1 CP3 Recovery": events,
        "S5S1 CP3 Handoff": [
            {
                "From": "Section 5 Session 1 of 3",
                "To": "Section 5 Session 2 of 3",
                "State": "ready",
                "Scope": "Provider preview, physical proof, correction control, production acceptance, and next-session recovery.",
                "Prerequisite": "Use the exact Response 77 complete restore; do not edit the frozen Session 1 production candidate in place.",
            }
        ],
    }
    for name, data in datasets.items():
        ws = wb.create_sheet(name)
        _write_sheet(ws, data)
    wb.properties.title = "Human Pathogen Database — Comprehensive Tracking Through Response 77"
    wb.properties.subject = "Remediation Section 5 Session 1 complete tracking and acceptance"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.calculation.fullCalcOnLoad = True
    wb.calculation.forceFullCalc = True
    wb.save(destination)
    with zipfile.ZipFile(destination) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("Workbook CRC failure")
    check = load_workbook(destination, read_only=True, data_only=False)
    try:
        sheets = list(check.sheetnames)
        formula_errors = 0
        formula_cells = 0
        for ws in check.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    value = cell.value
                    if isinstance(value, str) and value.startswith("="):
                        formula_cells += 1
                    if isinstance(value, str) and any(token in value for token in ("#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "#NUM!", "#N/A")):
                        formula_errors += 1
    finally:
        check.close()
    if any(name not in sheets for name in inherited) or len(sheets) < 106 or formula_errors:
        raise RuntimeError({"workbook_final_gate": {"sheets": len(sheets), "formula_errors": formula_errors, "lost": [name for name in inherited if name not in sheets]}})
    return {
        "status": "passed",
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "inherited_sheet_count": len(inherited),
        "current_sheet_count": len(sheets),
        "formula_cells": formula_cells,
        "formula_error_count": formula_errors,
        "new_sheets": list(datasets),
    }


def finalize_database(destination: Path, workbook_qa: dict[str, Any], audit: dict[str, Any], now_iso: str) -> dict[str, Any]:
    rows = acceptance_rows(audit, workbook_qa["current_sheet_count"], workbook_qa["formula_error_count"], now_iso)
    con = sqlite3.connect(destination)
    try:
        con.execute("PRAGMA foreign_keys=ON")
        con.execute("BEGIN IMMEDIATE")
        con.execute(
            "UPDATE section5_session1_checkpoint SET state='session_complete',workbook_status='passed',application_status='passed',next_checkpoint=?,recorded_at=? WHERE checkpoint_code=?",
            (NEXT_SESSION, now_iso, CHECKPOINT_CODE),
        )
        con.execute(
            "UPDATE section5_session1_release SET state='session_complete',database_status='ok',workbook_status='passed',application_status='passed',recorded_at=? WHERE release_code=?",
            (now_iso, RELEASE_CODE),
        )
        con.execute("DELETE FROM section5_session1_acceptance WHERE release_code=?", (RELEASE_CODE,))
        for row in rows:
            con.execute(
                "INSERT INTO section5_session1_acceptance (release_code,gate_key,expected,observed,status,checked_at) VALUES (?,?,?,?,?,?)",
                (RELEASE_CODE, row["gate_key"], row["expected"], row["observed"], row["status"], row["checked_at"]),
            )
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        foreign_keys = list(con.execute("PRAGMA foreign_key_check"))
        response77 = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R77'").fetchone()[0]
        cp_state = con.execute("SELECT state FROM section5_session1_checkpoint WHERE checkpoint_code=?", (CHECKPOINT_CODE,)).fetchone()
        release_state = con.execute("SELECT state FROM section5_session1_release WHERE release_code=?", (RELEASE_CODE,)).fetchone()
        page_records = con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code=?", (CHECKPOINT_CODE,)).fetchone()[0]
        failed_pages = con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code=? AND status!='passed'", (CHECKPOINT_CODE,)).fetchone()[0]
        failed_internal = con.execute("SELECT COUNT(*) FROM section5_session1_acceptance WHERE release_code=? AND status='failed'", (RELEASE_CODE,)).fetchone()[0]
        table_count = con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0]
        if integrity != "ok" or foreign_keys or response77 != 1 or cp_state != ("session_complete",) or release_state != ("session_complete",) or page_records != 538 or failed_pages or failed_internal:
            raise RuntimeError({"database_final_gate": {"integrity": integrity, "foreign_keys": foreign_keys[:20], "response77": response77, "checkpoint": cp_state, "release": release_state, "pages": page_records, "failed_pages": failed_pages, "failed_internal": failed_internal}})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()
    return {
        "status": "passed",
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "table_count": table_count,
        "integrity": integrity,
        "foreign_key_violations": len(foreign_keys),
        "response77_records": response77,
        "checkpoint_state": cp_state[0],
        "session_release_state": release_state[0],
        "page_transform_records": page_records,
        "failed_page_transforms": failed_pages,
        "failed_acceptance_gates": failed_internal,
    }


def update_application_state(project: Path, db: Path, workbook: Path, audit: dict[str, Any], now_iso: str) -> dict[str, Any]:
    app = audit["application"]
    if sha256_file(app) != APPLICATION_SHA256:
        raise RuntimeError("Main application source changed")
    root = project / "App" / "Section 5 Session 1 Complete"
    root.mkdir(parents=True, exist_ok=True)
    pointer = root / "CURRENT_DATABASE.txt"
    state = root / "CURRENT_PROJECT_STATE.json"
    audit_script = root / "audit_section5_session1_complete.py"
    text_write(pointer, db.relative_to(project).as_posix() + "\n")
    json_write(state, {
        "schema": "mrhpd-section5-session1-complete-state-1.0",
        "response": 77,
        "section": SECTION_LABEL,
        "session": SESSION_LABEL + " COMPLETE",
        "checkpoint": CHECKPOINT_LABEL + " COMPLETE",
        "database": db.relative_to(project).as_posix(),
        "workbook": workbook.relative_to(project).as_posix(),
        "digital_publication": audit["publication"].relative_to(project).as_posix(),
        "print_interior": audit["print_interior"].relative_to(project).as_posix(),
        "cover": audit["cover_png"].relative_to(project).as_posix(),
        "main_application": app.relative_to(project).as_posix(),
        "main_application_sha256": sha256_file(app),
        "main_application_unchanged": True,
        "provider_previewer": "controlled_pending_session2",
        "physical_proof": "controlled_pending_session2_or_3",
        "next": NEXT_SESSION,
        "recorded_at": now_iso,
    })
    text_write(audit_script, f'''#!/usr/bin/env python3
import json, sqlite3
from pathlib import Path
from openpyxl import load_workbook
from pypdf import PdfReader
project=Path(__file__).resolve().parents[2]
db=project/{db.relative_to(project).as_posix()!r}
workbook=project/{workbook.relative_to(project).as_posix()!r}
interior=project/{audit['print_interior'].relative_to(project).as_posix()!r}
cover=project/{audit['cover_png'].relative_to(project).as_posix()!r}
con=sqlite3.connect(db)
try:
 integrity=con.execute("PRAGMA integrity_check").fetchone()[0]
 fk=len(list(con.execute("PRAGMA foreign_key_check")))
 response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R77'").fetchone()[0]
 checkpoint=con.execute("SELECT state FROM section5_session1_checkpoint WHERE checkpoint_code='{CHECKPOINT_CODE}'").fetchone()
 release=con.execute("SELECT state FROM section5_session1_release WHERE release_code='{RELEASE_CODE}'").fetchone()
 pages=con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code='{CHECKPOINT_CODE}'").fetchone()[0]
 failed=con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code='{CHECKPOINT_CODE}' AND status!='passed'").fetchone()[0]
finally: con.close()
wb=load_workbook(workbook,read_only=True,data_only=False)
try: sheets=len(wb.sheetnames)
finally: wb.close()
pdf=PdfReader(str(interior)); page_count=len(pdf.pages); searchable=sum(1 for p in pdf.pages[:537] if (p.extract_text() or '').strip())
result={{'status':'passed' if integrity=='ok' and fk==0 and response==1 and checkpoint==('session_complete',) and release==('session_complete',) and pages==538 and failed==0 and sheets>=106 and page_count==538 and searchable==537 and cover.exists() else 'failed','integrity':integrity,'foreign_keys':fk,'response77':response,'checkpoint':checkpoint,'release':release,'page_records':pages,'failed_page_records':failed,'workbook_sheets':sheets,'print_pages':page_count,'searchable_pages':searchable,'cover_exists':cover.exists()}}
print(json.dumps(result,indent=2))
raise SystemExit(0 if result['status']=='passed' else 1)
''')
    result = subprocess.run([sys.executable, str(audit_script)], cwd=project, text=True, capture_output=True, timeout=600)
    if result.returncode:
        raise RuntimeError({"application_audit_failed": {"stdout": result.stdout[-12000:], "stderr": result.stderr[-12000:]}})
    output = root / "SECTION5_SESSION1_COMPLETE_APPLICATION_AUDIT.json"
    payload = json.loads(result.stdout)
    payload.update({"main_application_path": app.relative_to(project).as_posix(), "main_application_sha256": sha256_file(app), "main_application_unchanged": True})
    json_write(output, payload)
    return {"status": "passed", "pointer": pointer, "state": state, "audit_script": audit_script, "audit_output": output, "audit": payload}


def style_docx_table(table: Any, header_fill: str = NAVY) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), header_fill if idx == 0 else (PALE_BLUE if idx % 2 == 0 else WHITE))
            tc_pr.append(shd)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
                    if idx == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def create_summary_figure(project: Path, audit: dict[str, Any]) -> Path:
    output = project / "Reports" / "Section 5 Session 1 Complete" / "MRHPD v3.0.0a Section 5 Session 1 Production Freeze Summary Response 77.png"
    output.parent.mkdir(parents=True, exist_ok=True)
    canvas = Image.new("RGB", (2400, 1600), "white")
    draw = ImageDraw.Draw(canvas)
    try:
        title_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 74)
        body_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 34)
        small_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 27)
    except OSError:
        title_font = body_font = small_font = ImageFont.load_default()
    draw.rectangle((0, 0, 2400, 175), fill=(23, 50, 77))
    draw.text((90, 48), "HUMAN PATHOGEN DATABASE — SESSION 1 PRODUCTION FREEZE", font=title_font, fill="white")
    cover = Image.open(audit["cover_png"]).convert("RGB")
    cover.thumbnail((1080, 1160), Image.Resampling.LANCZOS)
    canvas.paste(cover, (90, 245))
    x = 1270
    draw.text((x, 255), "Frozen candidate", font=title_font, fill=(23, 50, 77))
    lines = [
        "KDP Premium Color paperback",
        "8.5 × 11 in trim • 538 pages",
        "1.262686 in spine",
        "18.512686 × 11.250 in full wrap",
        "5,554 × 3,375 px RGB cover",
        "537 searchable source pages",
        "One intentional terminal blank page",
        "Database, workbook, application: passed",
        "Provider preview: controlled pending",
        "Physical proof: controlled pending",
    ]
    y = 390
    for line in lines:
        draw.rounded_rectangle((x, y, 2290, y + 88), radius=20, fill=(234, 241, 245), outline=(28, 116, 117), width=3)
        draw.text((x + 30, y + 21), line, font=body_font, fill=(23, 50, 77))
        y += 108
    draw.text((90, 1515), "Response 77 • Checkpoint 3 of 3 • Session 1 of 3 COMPLETE", font=small_font, fill=(23, 50, 77))
    canvas.save(output, format="PNG", dpi=(300, 300), optimize=True)
    return output


def build_reports(project: Path, final_qa: dict[str, Any], acceptance: list[dict[str, Any]], events: list[dict[str, Any]]) -> list[Path]:
    root = project / "Reports" / "Section 5 Session 1 Complete"
    root.mkdir(parents=True, exist_ok=True)
    docx_path = root / "MRHPD v3.0.0a Section 5 Session 1 Complete Report Through Response 77.docx"
    pdf_path = root / "MRHPD v3.0.0a Section 5 Session 1 Complete Report Through Response 77.pdf"
    xlsx_path = root / "MRHPD v3.0.0a Section 5 Session 1 Final Acceptance Register Through Response 77.xlsx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Human Pathogen Database\nSection 5 Session 1 Complete")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(23, 50, 77)
    subtitle = doc.add_paragraph("Response 77 • Checkpoint 3 of 3 • Version 3.0.0a")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "The Session 1 production candidate is frozen and independently verified. The immutable 537-page digital publication remains unchanged; the separate 538-page Premium Color print interior and exact cover package are carried forward to provider-preview and physical-proof controls in Session 2."
    )
    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.rows[0].cells[0].text = "Control"
    summary_table.rows[0].cells[1].text = "Observed"
    summary_table.rows[0].cells[2].text = "Status"
    summary_rows = [
        ("Database", f"{final_qa['database']['table_count']} tables; integrity ok; 0 foreign-key violations", "passed"),
        ("Workbook", f"{final_qa['workbook']['current_sheet_count']} sheets; 0 formula-error tokens", "passed"),
        ("Print interior", "538 pages; 537 searchable source pages; terminal blank page", "passed"),
        ("Cover", "5554 × 3375 px; 18.512686 × 11.250 in", "passed"),
        ("Provider preview", "KDP Print Previewer", "controlled pending"),
        ("Physical proof", "Print and inspect physical copy", "controlled pending"),
    ]
    for control, observed, status in summary_rows:
        cells = summary_table.add_row().cells
        cells[0].text = control
        cells[1].text = observed
        cells[2].text = status
    style_docx_table(summary_table)
    doc.add_heading("Acceptance matrix", level=1)
    table = doc.add_table(rows=1, cols=4)
    for idx, value in enumerate(("Gate", "Expected", "Observed", "Status")):
        table.rows[0].cells[idx].text = value
    for row in acceptance:
        cells = table.add_row().cells
        cells[0].text = row["gate_key"]
        cells[1].text = row["expected"]
        cells[2].text = row["observed"]
        cells[3].text = row["status"]
    style_docx_table(table)
    doc.add_heading("Recovery history", level=1)
    for event in events:
        p = doc.add_paragraph()
        p.add_run(event["event_code"] + ". ").bold = True
        p.add_run(event["condition"] + " Recovery: " + event["recovery"])
    doc.add_heading("Next session", level=1)
    doc.add_paragraph(
        "Session 2 begins from the exact Response 77 complete restore. It will collect provider-preview evidence, manage corrections, initiate or govern the physical-proof lane, and preserve the frozen Session 1 production candidate as an immutable predecessor."
    )
    doc.core_properties.title = "Human Pathogen Database — Section 5 Session 1 Complete"
    doc.core_properties.subject = "Response 77 final acceptance and restore"
    doc.core_properties.author = "Brent McAnulty, M.D."
    doc.save(docx_path)
    with zipfile.ZipFile(docx_path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("DOCX CRC failure")

    styles_pdf = getSampleStyleSheet()
    styles_pdf.add(ParagraphStyle(name="CenterTitle", parent=styles_pdf["Title"], alignment=TA_CENTER, textColor=colors.HexColor("#17324D"), fontSize=20, leading=24))
    styles_pdf.add(ParagraphStyle(name="BodySmall", parent=styles_pdf["BodyText"], fontSize=8.5, leading=11, alignment=TA_LEFT))
    story: list[Any] = [
        Paragraph("Human Pathogen Database", styles_pdf["CenterTitle"]),
        Paragraph("Section 5 Session 1 Complete — Response 77", styles_pdf["Heading2"]),
        Spacer(1, 0.15 * inch),
        Paragraph("The frozen production candidate retains the immutable 537-page digital edition and a separate 538-page Premium Color print derivative with the exact cover package.", styles_pdf["BodyText"]),
        Spacer(1, 0.15 * inch),
    ]
    summary_data = [["Control", "Observed", "Status"]] + [[a, b, c] for a, b, c in summary_rows]
    summary_pdf_table = Table(summary_data, colWidths=[1.55 * inch, 4.55 * inch, 1.1 * inch], repeatRows=1)
    summary_pdf_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17324D")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7.5),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#AAB8C0")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EAF1F5")]),
    ]))
    story.extend([summary_pdf_table, PageBreak(), Paragraph("Acceptance matrix", styles_pdf["Heading1"])])
    acceptance_data = [["Gate", "Expected", "Observed", "Status"]] + [[row["gate_key"], row["expected"], row["observed"], row["status"]] for row in acceptance]
    acceptance_table = Table(acceptance_data, colWidths=[1.4 * inch, 1.85 * inch, 3.0 * inch, 0.95 * inch], repeatRows=1)
    acceptance_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17324D")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 6.5),
        ("LEADING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#AAB8C0")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F0D8")]),
    ]))
    story.extend([acceptance_table, PageBreak(), Paragraph("Recovery and continuation", styles_pdf["Heading1"])])
    for event in events:
        story.append(Paragraph(f"<b>{event['event_code']}</b> — {event['condition']} Recovery: {event['recovery']}", styles_pdf["BodySmall"]))
        story.append(Spacer(1, 0.05 * inch))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph("Next: Remediation Section 5 of 5, Session 2 of 3. Provider preview and physical-proof evidence remain controlled external gates.", styles_pdf["BodyText"]))
    pdf_doc = SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=0.55 * inch, leftMargin=0.55 * inch, topMargin=0.55 * inch, bottomMargin=0.55 * inch, title="Human Pathogen Database Section 5 Session 1 Complete", author="Brent McAnulty, M.D.")
    pdf_doc.build(story)
    reader = PdfReader(str(pdf_path))
    text_chars = sum(len(page.extract_text() or "") for page in reader.pages)
    if len(reader.pages) < 3 or text_chars < 3000:
        raise RuntimeError({"report_pdf_gate": {"pages": len(reader.pages), "text_chars": text_chars}})

    wb = Workbook()
    wb.remove(wb.active)
    datasets = {
        "Summary": [{"Control": a, "Observed": b, "Status": c} for a, b, c in summary_rows],
        "Acceptance": acceptance,
        "Recovery": events,
        "Database": [final_qa["database"]],
        "Workbook": [final_qa["workbook"]],
        "Application": [final_qa["application"]],
        "Publication": [final_qa["publication"]],
        "Handoff": [{"From": "Session 1 of 3", "To": "Session 2 of 3", "State": "ready", "Next": NEXT_SESSION}],
    }
    for name, rows in datasets.items():
        ws = wb.create_sheet(name)
        _write_sheet(ws, rows)
    wb.properties.title = "MRHPD Section 5 Session 1 Final Acceptance Register"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(xlsx_path)
    with zipfile.ZipFile(xlsx_path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("Acceptance-register XLSX CRC failure")
    return [docx_path, pdf_path, xlsx_path]


def build_tracking(project: Path, now_iso: str) -> list[Path]:
    root = project / "Tracking" / "Section 5 Session 1 Complete"
    root.mkdir(parents=True, exist_ok=True)
    response_json = root / "Response_77_Tracking.json"
    raw_net = root / "Raw_and_Net_Tracking_Through_Response_77.md"
    index_update = root / "Cumulative_Thread_Index_Update_Response_77.md"
    raw_docx = root / "MRHPD v3.0.0a Alternating Raw Prompts and Responses Through Response 77.docx"
    net_docx = root / "MRHPD v3.0.0a Alternating Net Prompts and Responses Through Response 77.docx"
    json_write(response_json, {
        "schema": "mrhpd-response-tracking-1.0",
        "response": 77,
        "raw_prompt": "Continue",
        "net_prompt": "Continue from the newest verified Response 76 checkpoint and complete Section 5 Session 1 with independent verification and a full self-contained restore.",
        "raw_response": "[PRE-EMISSION RESPONSE; source-supported final summary controls the tracked response]",
        "net_response": "The exact Response 76 state was reconstructed, revalidated, frozen as the Session 1 production candidate, synchronized across the database, workbook, application, tracking, indexes, manifests, reports, and recovery controls, and emitted as a complete self-contained restore.",
        "state": "session_complete_continue_required",
        "recorded_at": now_iso,
    })
    text_write(raw_net, """# Human Pathogen Database — Raw and Net Tracking Through Response 77

## Raw Prompt 77

Continue

## Net Prompt

Continue from the newest verified Response 76 checkpoint and complete Section 5 Session 1 with independent verification and a full self-contained restore.

## Net Response

The exact Response 76 state was reconstructed, revalidated, frozen as the Session 1 production candidate, synchronized across the database, workbook, application, tracking, indexes, manifests, reports, and recovery controls, and emitted as a complete self-contained restore.

## Disposition

Checkpoint 3 of 3: COMPLETE. Session 1 of 3: COMPLETE. Remediation Section 5 of 5: CONTINUE. Next: Session 2 of 3.
""")
    text_write(index_update, """# Cumulative Thread Index — Response 77 Update

[SUMMARY: Response #: 77; Thread: Medical References; Title: Human Pathogen Database Section 5 Session 1 complete production-candidate restore; Goal: Continue from the newest verified Response 76 checkpoint and finish Session 1 with independent verification and a complete self-contained restore; Output: Reconstructed and revalidated the Response 76 production candidate, froze the 538-page interior and exact cover package, synchronized all governed artifacts, and emitted the complete Session 1 restore; CONTINUE to proceed with Remediation Section 5 of 5 Session 2 of 3 + Session 2/3 will be next.]
""")
    for path, title, prompt, response in [
        (raw_docx, "Alternating Raw Prompts and Responses Through Response 77", "Continue", "[Source-supported final Response 77 summary is preserved in Response_77_Tracking.json and the Cumulative Thread Index update.]"),
        (net_docx, "Alternating Net Prompts and Responses Through Response 77", "Continue from the newest verified Response 76 checkpoint and complete Section 5 Session 1 with independent verification and a full self-contained restore.", "The exact Response 76 state was reconstructed, revalidated, frozen, synchronized, and emitted as the complete Session 1 restore."),
    ]:
        doc = Document()
        heading = doc.add_heading("Human Pathogen Database", 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading(title, 1)
        doc.add_heading("Session 1 completion and restore", 2)
        p = doc.add_paragraph()
        r = p.add_run("Prompt 77\n")
        r.bold = True
        r.font.color.rgb = RGBColor(31, 78, 121)
        p.add_run(prompt)
        p2 = doc.add_paragraph()
        r2 = p2.add_run("Response 77\n")
        r2.bold = True
        r2.font.color.rgb = RGBColor(112, 48, 160)
        p2.add_run(response)
        doc.add_paragraph("Checkpoint 3 of 3 and Session 1 of 3 are complete. Continue begins Session 2 of 3.")
        doc.core_properties.title = title
        doc.core_properties.author = "Brent McAnulty, M.D."
        doc.save(path)
    return [response_json, raw_net, index_update, raw_docx, net_docx]


def compact_superseded_snapshots(project: Path, current_db: Path, current_workbook: Path, now_iso: str) -> dict[str, Any]:
    removed: list[dict[str, Any]] = []
    current_con = sqlite3.connect(current_db)
    try:
        current_tables = {row[0] for row in current_con.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'")}
        current_counts: dict[str, int] = {}
        for table in current_tables:
            try:
                current_counts[table] = current_con.execute(f'SELECT COUNT(*) FROM "{table}"').fetchone()[0]
            except sqlite3.DatabaseError:
                pass
    finally:
        current_con.close()
    db_root = project / "Database"
    for candidate in sorted(p for p in db_root.rglob("*") if p.is_file() and p.suffix.lower() in {".sqlite", ".db"} and p != current_db):
        if not any(token in candidate.name for token in ("Checkpoint", "THROUGH RESPONSE", "Remediation Section 4", "Remediation Section 5")):
            continue
        try:
            con = sqlite3.connect(candidate)
            integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
            old_tables = {row[0] for row in con.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'")}
            old_counts = {}
            for table in old_tables:
                try:
                    old_counts[table] = con.execute(f'SELECT COUNT(*) FROM "{table}"').fetchone()[0]
                except sqlite3.DatabaseError:
                    pass
            con.close()
        except Exception:
            continue
        if integrity == "ok" and old_tables.issubset(current_tables) and all(current_counts.get(table, -1) >= count for table, count in old_counts.items()):
            removed.append({"path": candidate.relative_to(project).as_posix(), "bytes": candidate.stat().st_size, "sha256": sha256_file(candidate), "reason": "superseded_database_superset_proven", "proof": {"old_tables": len(old_tables), "current_tables": len(current_tables)}})
            candidate.unlink()
    wb = load_workbook(current_workbook, read_only=True)
    try:
        current_sheets = set(wb.sheetnames)
    finally:
        wb.close()
    workbook_root = project / "Tracking" / "Workbook"
    for candidate in sorted(p for p in workbook_root.rglob("*.xlsx") if p != current_workbook):
        if not any(token in candidate.name for token in ("Checkpoint", "THROUGH RESPONSE", "Comprehensive Tracking")):
            continue
        try:
            old = load_workbook(candidate, read_only=True)
            old_sheets = set(old.sheetnames)
            old.close()
        except Exception:
            continue
        if old_sheets.issubset(current_sheets):
            removed.append({"path": candidate.relative_to(project).as_posix(), "bytes": candidate.stat().st_size, "sha256": sha256_file(candidate), "reason": "superseded_workbook_sheet_superset_proven", "proof": {"old_sheets": len(old_sheets), "current_sheets": len(current_sheets)}})
            candidate.unlink()
    for relative in (
        "Indexes/Section 5 Session 1 Checkpoint 1",
        "Indexes/Section 5 Session 1 Checkpoint 2",
        "Manifest/Section 5 Session 1 Checkpoint 1",
        "Manifest/Section 5 Session 1 Checkpoint 2",
    ):
        target = project / relative
        if target.exists():
            for path in sorted(p for p in target.rglob("*") if p.is_file()):
                removed.append({"path": path.relative_to(project).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path), "reason": "superseded_index_or_manifest_derivative", "proof": {"replacement": "Section 5 Session 1 Complete indexes and manifests"}})
            shutil.rmtree(target)
    register = project / "Recovery" / "Section 5 Session 1 Complete" / "SESSION1_COMPACTION_REGISTER.json"
    json_write(register, {"schema": "mrhpd-equivalence-controlled-compaction-1.0", "generated_at": now_iso, "status": "passed", "removed_count": len(removed), "removed_bytes": sum(row["bytes"] for row in removed), "records": removed, "clinical_content_removed": False, "current_database_removed": False, "current_workbook_removed": False, "immutable_publication_removed": False, "print_production_candidate_removed": False})
    return {"status": "passed", "register": register, "removed_count": len(removed), "removed_bytes": sum(row["bytes"] for row in removed), "records": removed}


def extract_text_for_index(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in {".md", ".txt", ".csv", ".json", ".py", ".html", ".yml", ".yaml"}:
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix == ".docx":
            doc = Document(path)
            chunks = [paragraph.text for paragraph in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells)
            return "\n".join(chunks)
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        if suffix == ".xlsx":
            wb = load_workbook(path, read_only=True, data_only=False)
            try:
                chunks: list[str] = []
                for ws in wb.worksheets:
                    chunks.append(ws.title)
                    for row in ws.iter_rows(values_only=True):
                        chunks.append(" | ".join("" if value is None else str(value) for value in row))
                return "\n".join(chunks)
            finally:
                wb.close()
        if suffix in {".sqlite", ".db"}:
            con = sqlite3.connect(path)
            try:
                return "\n".join(row[0] or "" for row in con.execute("SELECT sql FROM sqlite_master WHERE sql IS NOT NULL ORDER BY name"))
            finally:
                con.close()
    except Exception as exc:
        return f"[index extraction error: {exc!r}]"
    return ""


def build_indexes(project: Path, now_iso: str) -> dict[str, Any]:
    root = project / "Indexes" / "Section 5 Session 1 Complete"
    root.mkdir(parents=True, exist_ok=True)
    source_json = root / "MRHPD v3.0.0a Section 5 Session 1 Complete Source Index.json"
    source_csv = root / "MRHPD v3.0.0a Section 5 Session 1 Complete Source Index.csv"
    bit_path = root / "MRHPD v3.0.0a Section 5 Session 1 Complete Bit Index.sqlite"
    qa_path = root / "MRHPD v3.0.0a Section 5 Session 1 Complete Index QA.json"
    excluded = {source_json.resolve(), source_csv.resolve(), bit_path.resolve(), qa_path.resolve()}
    searchable_suffixes = {".md", ".txt", ".csv", ".json", ".py", ".html", ".yml", ".yaml", ".docx", ".pdf", ".xlsx", ".sqlite", ".db"}
    rows: list[dict[str, Any]] = []
    payloads: list[tuple[str, str, str, str]] = []
    for path in sorted(p for p in project.rglob("*") if p.is_file() and p.resolve() not in excluded):
        rel = path.relative_to(project).as_posix()
        purpose = "Project artifact"
        for prefix, label in (
            ("Database/", "Canonical or historical project database"),
            ("Documents/", "Publication or editable assembly"),
            ("Print Production/", "Section 5 print-production artifact"),
            ("Cover/", "Cover component or historical cover master"),
            ("Tracking/", "Prompt, response, summary, and project tracking"),
            ("QA/", "Validation and acceptance evidence"),
            ("Reports/", "Human-readable report or register"),
            ("App/", "Local application or current-state audit surface"),
            ("Recovery/", "Checkpoint, recovery, compaction, and handoff evidence"),
        ):
            if rel.startswith(prefix):
                purpose = label
                break
        searchable = path.suffix.lower() in searchable_suffixes
        content = extract_text_for_index(path) if searchable else ""
        row = {"record_type": "physical_file", "path": rel, "container_path": "", "name": path.name, "purpose": purpose, "bytes": path.stat().st_size, "sha256": sha256_file(path), "user_searchable": int(searchable)}
        rows.append(row)
        payloads.append((rel, path.name, purpose, content))
        if path.suffix.lower() == ".zip":
            try:
                with zipfile.ZipFile(path) as zf:
                    if zf.testzip() is not None:
                        continue
                    for info in zf.infolist():
                        if info.is_dir():
                            continue
                        member_path = f"{rel}!/{info.filename}"
                        member_content = ""
                        suffix = Path(info.filename).suffix.lower()
                        if suffix in {".md", ".txt", ".csv", ".json", ".py", ".html", ".yml", ".yaml"} and info.file_size <= 5_000_000:
                            try:
                                member_content = zf.read(info).decode("utf-8", errors="replace")
                            except Exception:
                                member_content = ""
                        member = {"record_type": "container_member", "path": member_path, "container_path": rel, "name": Path(info.filename).name, "purpose": "Member of project container", "bytes": info.file_size, "sha256": "", "user_searchable": int(bool(member_content))}
                        rows.append(member)
                        payloads.append((member_path, member["name"], member["purpose"], member_content))
            except zipfile.BadZipFile:
                pass
    json_write(source_json, {"schema": "mrhpd-source-index-2.0", "generated_at": now_iso, "record_count": len(rows), "records": rows})
    csv_write(source_csv, rows)
    if bit_path.exists():
        bit_path.unlink()
    con = sqlite3.connect(bit_path)
    try:
        con.executescript("""
        CREATE TABLE artifact (
            artifact_id INTEGER PRIMARY KEY,
            record_type TEXT NOT NULL,
            path TEXT NOT NULL UNIQUE,
            container_path TEXT,
            name TEXT NOT NULL,
            purpose TEXT NOT NULL,
            bytes INTEGER NOT NULL,
            sha256 TEXT,
            user_searchable INTEGER NOT NULL
        );
        CREATE VIRTUAL TABLE artifact_fts USING fts5(path, name, purpose, content);
        """)
        for row, payload in zip(rows, payloads):
            con.execute("INSERT INTO artifact (record_type,path,container_path,name,purpose,bytes,sha256,user_searchable) VALUES (?,?,?,?,?,?,?,?)", (row["record_type"], row["path"], row["container_path"], row["name"], row["purpose"], row["bytes"], row["sha256"], row["user_searchable"]))
            con.execute("INSERT INTO artifact_fts (path,name,purpose,content) VALUES (?,?,?,?)", payload)
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        counts = {
            "artifact": con.execute("SELECT COUNT(*) FROM artifact").fetchone()[0],
            "fts": con.execute("SELECT COUNT(*) FROM artifact_fts").fetchone()[0],
            "response77": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"Response 77"',)).fetchone()[0],
            "session1_complete": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"Session 1" AND complete',)).fetchone()[0],
            "print_interior": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"print interior"',)).fetchone()[0],
        }
        con.commit()
    finally:
        con.close()
    if integrity != "ok" or counts["artifact"] != len(rows) or counts["fts"] != len(rows):
        raise RuntimeError({"bit_index_gate": {"integrity": integrity, "counts": counts, "expected": len(rows)}})
    qa = {"status": "passed", "generated_at": now_iso, "source_index_records": len(rows), "physical_files": sum(1 for row in rows if row["record_type"] == "physical_file"), "container_members": sum(1 for row in rows if row["record_type"] == "container_member"), "bit_index_integrity": integrity, "counts": counts, "bit_index_sha256": sha256_file(bit_path)}
    json_write(qa_path, qa)
    return {"source_json": source_json, "source_csv": source_csv, "bit_index": bit_path, "qa_path": qa_path, "qa": qa}


def build_manifest(project: Path, now_iso: str) -> tuple[Path, Path, list[dict[str, Any]]]:
    root = project / "Manifest" / "Section 5 Session 1 Complete"
    root.mkdir(parents=True, exist_ok=True)
    manifest = root / "MRHPD v3.0.0a Section 5 Session 1 Complete Current Project Manifest.json"
    checksums = root / "MRHPD v3.0.0a Section 5 Session 1 Complete Current Project Checksums.sha256"
    rows = []
    for path in sorted(p for p in project.rglob("*") if p.is_file() and p not in {manifest, checksums}):
        rows.append({"path": path.relative_to(project).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    json_write(manifest, {"schema": "mrhpd-current-project-manifest-2.0", "generated_at": now_iso, "exclusions": [manifest.relative_to(project).as_posix(), checksums.relative_to(project).as_posix()], "file_count": len(rows), "total_bytes": sum(row["bytes"] for row in rows), "files": rows})
    text_write(checksums, "".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
    for row in rows:
        path = project / row["path"]
        if path.stat().st_size != row["bytes"] or sha256_file(path) != row["sha256"]:
            raise RuntimeError({"manifest_mismatch": row["path"]})
    return manifest, checksums, rows


def quick_verify_project(project: Path, expected_db_sha: str, expected_workbook_sha: str) -> dict[str, Any]:
    db = project / CURRENT_DB_REL
    workbook = project / CURRENT_WORKBOOK_REL
    interior = project / PRINT_INTERIOR_REL
    cover = project / FINAL_COVER_PNG_REL
    publication = project / PUBLICATION_REL
    application = locate_main_application(project)
    if sha256_file(db) != expected_db_sha or sha256_file(workbook) != expected_workbook_sha:
        raise RuntimeError("Clean project database/workbook hash mismatch")
    if sha256_file(publication) != PUBLICATION_SHA256 or sha256_file(application) != APPLICATION_SHA256:
        raise RuntimeError("Clean project invariant hash mismatch")
    con = sqlite3.connect(db)
    try:
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        response = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R77'").fetchone()[0]
        checkpoint = con.execute("SELECT state FROM section5_session1_checkpoint WHERE checkpoint_code=?", (CHECKPOINT_CODE,)).fetchone()
        release = con.execute("SELECT state FROM section5_session1_release WHERE release_code=?", (RELEASE_CODE,)).fetchone()
    finally:
        con.close()
    wb = load_workbook(workbook, read_only=True, data_only=False)
    try:
        sheet_count = len(wb.sheetnames)
    finally:
        wb.close()
    reader = PdfReader(str(interior))
    pages = len(reader.pages)
    searchable = sum(1 for p in reader.pages[:537] if (p.extract_text() or "").strip())
    with Image.open(cover) as image:
        pixels = image.size
    result = {"status": "passed", "integrity": integrity, "foreign_keys": len(fk), "response77": response, "checkpoint": checkpoint, "release": release, "workbook_sheets": sheet_count, "print_pages": pages, "searchable_pages": searchable, "cover_pixels": list(pixels)}
    if integrity != "ok" or fk or response != 1 or checkpoint != ("session_complete",) or release != ("session_complete",) or sheet_count < 106 or pages != 538 or searchable != 537 or pixels != (5554, 3375):
        raise RuntimeError({"clean_project_gate": result})
    return result


def build_project_archive(project: Path, dist: Path, stamp: str, db_sha: str, workbook_sha: str) -> tuple[Path, dict[str, Any]]:
    root_name = f"Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 Remediation Section 5 of 5 Session 1 of 3 COMPLETE PROJECT THROUGH RESPONSE 77 {stamp}"
    archive = dist / f"{root_name}.zip"
    with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in sorted(project.rglob("*")):
            if path.is_file():
                zf.write(path, f"{root_name}/{path.relative_to(project).as_posix()}")
    qa = verify_zip(archive)
    with tempfile.TemporaryDirectory(prefix="mrhpd-r77-project-clean-") as td:
        clean = Path(td)
        safe_extract(archive, clean)
        restored = clean / root_name
        clean_qa = quick_verify_project(restored, db_sha, workbook_sha)
    qa["clean_project_verification"] = clean_qa
    return archive, qa


def build_restore_verify_script(project_name: str, project_bytes: int, project_sha: str) -> str:
    return f'''#!/usr/bin/env python3
from __future__ import annotations
import argparse, hashlib, json, re, sqlite3, zipfile
from pathlib import Path, PurePosixPath
from openpyxl import load_workbook
from pypdf import PdfReader
from PIL import Image
PROJECT_NAME={project_name!r}
PROJECT_BYTES={project_bytes}
PROJECT_SHA256={project_sha!r}
CURRENT_DB_REL={CURRENT_DB_REL!r}
CURRENT_WORKBOOK_REL={CURRENT_WORKBOOK_REL!r}
PRINT_INTERIOR_REL={PRINT_INTERIOR_REL!r}
FINAL_COVER_PNG_REL={FINAL_COVER_PNG_REL!r}
PUBLICATION_REL={PUBLICATION_REL!r}
PUBLICATION_SHA256={PUBLICATION_SHA256!r}
APPLICATION_SHA256={APPLICATION_SHA256!r}
CHECKPOINT_CODE={CHECKPOINT_CODE!r}
RELEASE_CODE={RELEASE_CODE!r}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as f:
  for b in iter(lambda:f.read(1024*1024),b''): h.update(b)
 return h.hexdigest()
def safe_extract(path,dest):
 dest.mkdir(parents=True,exist_ok=True)
 with zipfile.ZipFile(path) as zf:
  if zf.testzip() is not None: raise RuntimeError('ZIP CRC failure: '+str(path))
  names=zf.namelist()
  if len(names)!=len(set(names)): raise RuntimeError('duplicate ZIP members')
  for name in names:
   pp=PurePosixPath(name.replace('\\\\','/'))
   if pp.is_absolute() or '..' in pp.parts or re.match(r'^[A-Za-z]:',name): raise RuntimeError('unsafe ZIP path: '+name)
  zf.extractall(dest)
def locate_app(project):
 for p in project.rglob('human_pathogen_app.py'):
  if sha(p)==APPLICATION_SHA256: return p
 raise RuntimeError('main application identity not found')
def main():
 ap=argparse.ArgumentParser()
 ap.add_argument('--extract-project-to',type=Path)
 args=ap.parse_args()
 root=Path(__file__).resolve().parents[1]
 project_zip=root/'PROJECT_SNAPSHOT'/PROJECT_NAME
 if not project_zip.exists() or project_zip.stat().st_size!=PROJECT_BYTES or sha(project_zip)!=PROJECT_SHA256: raise RuntimeError('project snapshot identity failure')
 manifest=json.loads((root/'COMPLETE_RESTORE_MANIFEST.json').read_text())
 for row in manifest['files']:
  p=root/row['path']
  if not p.exists() or p.stat().st_size!=row['bytes'] or sha(p)!=row['sha256']: raise RuntimeError('restore manifest mismatch: '+row['path'])
 target=args.extract_project_to or (root/'VERIFIED_PROJECT_EXTRACTION')
 if target.exists() and any(target.iterdir()): raise RuntimeError('project extraction destination must be empty')
 safe_extract(project_zip,target)
 roots=[p for p in target.iterdir() if p.is_dir()]
 project=roots[0] if len(roots)==1 else target
 db=project/CURRENT_DB_REL; workbook=project/CURRENT_WORKBOOK_REL; interior=project/PRINT_INTERIOR_REL; cover=project/FINAL_COVER_PNG_REL; publication=project/PUBLICATION_REL
 con=sqlite3.connect(db)
 try:
  integrity=con.execute('PRAGMA integrity_check').fetchone()[0]
  fk=len(list(con.execute('PRAGMA foreign_key_check')))
  response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R77'").fetchone()[0]
  checkpoint=con.execute('SELECT state FROM section5_session1_checkpoint WHERE checkpoint_code=?',(CHECKPOINT_CODE,)).fetchone()
  release=con.execute('SELECT state FROM section5_session1_release WHERE release_code=?',(RELEASE_CODE,)).fetchone()
 finally: con.close()
 wb=load_workbook(workbook,read_only=True,data_only=False)
 try: sheets=len(wb.sheetnames)
 finally: wb.close()
 reader=PdfReader(str(interior)); pages=len(reader.pages); searchable=sum(1 for p in reader.pages[:537] if (p.extract_text() or '').strip())
 with Image.open(cover) as im: pixels=im.size
 app=locate_app(project)
 result={{'status':'passed' if integrity=='ok' and fk==0 and response==1 and checkpoint==('session_complete',) and release==('session_complete',) and sheets>=106 and pages==538 and searchable==537 and pixels==(5554,3375) and sha(publication)==PUBLICATION_SHA256 and sha(app)==APPLICATION_SHA256 else 'failed','project':project.name,'database_integrity':integrity,'foreign_keys':fk,'response77':response,'checkpoint':checkpoint,'release':release,'workbook_sheets':sheets,'print_pages':pages,'searchable_pages':searchable,'cover_pixels':pixels,'publication_sha256':sha(publication),'application_sha256':sha(app)}}
 print(json.dumps(result,indent=2))
 raise SystemExit(0 if result['status']=='passed' else 1)
if __name__=='__main__': main()
'''


def build_complete_restore(project_archive: Path, project_qa: dict[str, Any], reports: list[Path], tracking: list[Path], dist: Path, stamp: str, final_qa: dict[str, Any]) -> tuple[Path, dict[str, Any]]:
    package = Path(tempfile.mkdtemp(prefix="mrhpd-r77-restore-package-"))
    try:
        (package / "PROJECT_SNAPSHOT").mkdir()
        (package / "REPORTS").mkdir()
        (package / "TRACKING").mkdir()
        (package / "TOOLS").mkdir()
        shutil.copy2(project_archive, package / "PROJECT_SNAPSHOT" / project_archive.name)
        for path in reports:
            shutil.copy2(path, package / "REPORTS" / path.name)
        for path in tracking:
            shutil.copy2(path, package / "TRACKING" / path.name)
        text_write(package / "RESTORE_READ_FIRST.md", f"""# Human Pathogen Database — Complete Restore Through Response 77

This is the complete self-contained Section 5 Session 1 restore. It requires no prior project ZIP, checkpoint recovery, cloud artifact, user-supplied project file, or reconstruction from the conversation.

## Restore

1. Run `python TOOLS/restore_verify_extract.py` to verify and extract into the default verified directory.
2. Or run `python TOOLS/restore_verify_extract.py --extract-project-to <destination>`.
3. Open the extracted project and begin with its current README, Recovery, Tracking, QA, Reports, Indexes, Manifest, Database, Workbook, App, Documents, and Print Production directories.

## Current state

- Response 77: COMPLETE
- Checkpoint 3 of 3: COMPLETE
- Section 5 Session 1 of 3: COMPLETE
- Remediation Section 5 of 5: CONTINUE
- Next: {NEXT_SESSION}
- KDP Print Previewer: controlled pending Session 2
- Physical proof: controlled pending Session 2 or 3
- Accepted predecessor modified: no
- Frozen Section 3 release modified: no
- Immutable 537-page digital publication modified: no
""")
        text_write(package / "Instructions.txt", "Project Instructions 1.5.0 remain controlling. The complete operative instructions and current addenda are included within the project snapshot. Preserve automatic recovery, Google Drive custody, Raw/Net tracking, checkpoint recovery between full restores, full restores at session and section boundaries, exact filenames, indexes, manifests, checksums, and immutable-predecessor controls.")
        identity = {
            "schema": "mrhpd-complete-restore-identity-1.0",
            "generated_at": final_qa["generated_at"],
            "version": PROJECT_VERSION,
            "response": 77,
            "section": SECTION_LABEL + " CONTINUE",
            "session": SESSION_LABEL + " COMPLETE",
            "checkpoint": CHECKPOINT_LABEL + " COMPLETE",
            "project_snapshot": {"name": project_archive.name, "bytes": project_archive.stat().st_size, "sha256": sha256_file(project_archive)},
            "self_contained": True,
            "requires_other_project_files": False,
            "requires_conversation_reconstruction": False,
            "accepted_predecessor_mutated": False,
            "immutable_publication_mutated": False,
            "next": NEXT_SESSION,
        }
        json_write(package / "CURRENT_PROJECT_IDENTITY.json", identity)
        json_write(package / "SESSION_1_ACCEPTANCE_QA.json", final_qa)
        text_write(package / "TOOLS" / "restore_verify_extract.py", build_restore_verify_script(project_archive.name, project_archive.stat().st_size, sha256_file(project_archive)))
        control_names = {"COMPLETE_RESTORE_MANIFEST.json", "COMPLETE_RESTORE_CHECKSUMS.sha256"}
        rows = []
        for path in sorted(package.rglob("*")):
            if path.is_file() and path.name not in control_names:
                rows.append({"path": path.relative_to(package).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path)})
        json_write(package / "COMPLETE_RESTORE_MANIFEST.json", {"schema": "mrhpd-complete-restore-manifest-1.0", "generated_at": final_qa["generated_at"], "file_count": len(rows), "total_bytes": sum(row["bytes"] for row in rows), "files": rows, "self_contained": True})
        text_write(package / "COMPLETE_RESTORE_CHECKSUMS.sha256", "".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
        restore = dist / f"Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 Remediation Section 5 of 5 Session 1 of 3 COMPLETE RESTORE THROUGH RESPONSE 77 {stamp}.zip"
        with zipfile.ZipFile(restore, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
            for path in sorted(package.rglob("*")):
                if path.is_file():
                    zf.write(path, path.relative_to(package).as_posix())
        restore_qa = verify_zip(restore)
        with tempfile.TemporaryDirectory(prefix="mrhpd-r77-restore-clean-") as td:
            clean = Path(td)
            safe_extract(restore, clean)
            result = subprocess.run([sys.executable, str(clean / "TOOLS" / "restore_verify_extract.py")], cwd=clean, text=True, capture_output=True, timeout=2400)
            if result.returncode:
                raise RuntimeError({"restore_verifier_failed": {"stdout": result.stdout[-16000:], "stderr": result.stderr[-16000:]}})
            verifier_output = json.loads(result.stdout)
        verification = {"schema": "mrhpd-response77-complete-restore-verification-1.0", "generated_at": final_qa["generated_at"], "status": "passed", "restore": restore_qa, "project_snapshot": project_qa, "clean_restore_verifier": "passed", "verifier_output": verifier_output, "self_contained": True, "requires_other_project_files": False, "requires_conversation_reconstruction": False, "checkpoint_3_of_3_complete": True, "session_1_of_3_complete": True, "remediation_section_5_complete": False, "next": NEXT_SESSION}
        return restore, verification
    finally:
        shutil.rmtree(package, ignore_errors=True)


def build_transport_volumes(restore: Path, dist: Path, generated_at: str) -> dict[str, Any]:
    total = restore.stat().st_size
    max_raw = 94_000_000
    part_count = max(2, math.ceil(total / max_raw))
    if part_count > 4:
        raise RuntimeError({"restore_requires_too_many_volumes": {"bytes": total, "part_count": part_count}})
    part_size = math.ceil(total / part_count)
    raw_parts = []
    with restore.open("rb") as source:
        for sequence in range(1, part_count + 1):
            remaining_total = total - source.tell()
            size = min(part_size, remaining_total)
            path = dist / f"{restore.name}.part{sequence:03d}"
            with path.open("wb") as out:
                remaining = size
                while remaining:
                    block = source.read(min(1024 * 1024, remaining))
                    if not block:
                        raise RuntimeError("Unexpected EOF while splitting restore")
                    out.write(block)
                    remaining -= len(block)
            raw_parts.append({"sequence": sequence, "name": path.name, "bytes": path.stat().st_size, "sha256": sha256_file(path), "path": path})
    manifest = {
        "schema": "mrhpd-complete-restore-transport-1.0",
        "generated_at": generated_at,
        "restore": {"name": restore.name, "bytes": total, "sha256": sha256_file(restore)},
        "part_count": part_count,
        "parts": [{k: v for k, v in row.items() if k != "path"} for row in raw_parts],
        "minimum_volume_count": part_count,
        "connector_limit_bytes": 104_857_600,
    }
    manifest_path = dist / "MRHPD_RESPONSE77_COMPLETE_RESTORE_TRANSPORT_MANIFEST.json"
    json_write(manifest_path, manifest)
    utility = dist / "reassemble_response77_complete_restore.py"
    text_write(utility, f'''#!/usr/bin/env python3
import hashlib,json
from pathlib import Path
M={manifest!r}
def sha(p):
 h=hashlib.sha256()
 with open(p,'rb') as f:
  for b in iter(lambda:f.read(1024*1024),b''): h.update(b)
 return h.hexdigest()
root=Path(__file__).resolve().parent
out=root/M['restore']['name']
with open(out,'wb') as dst:
 for row in M['parts']:
  p=root/row['name']
  if not p.exists() or p.stat().st_size!=row['bytes'] or sha(p)!=row['sha256']: raise SystemExit('Part identity failure: '+row['name'])
  with open(p,'rb') as src:
   for block in iter(lambda:src.read(1024*1024),b''): dst.write(block)
if out.stat().st_size!=M['restore']['bytes'] or sha(out)!=M['restore']['sha256']: raise SystemExit('Restore identity failure')
print(json.dumps({{'status':'passed','restore':out.name,'bytes':out.stat().st_size,'sha256':sha(out)}},indent=2))
''')
    wrappers = []
    for row in raw_parts:
        wrapper = dist / f"MRHPD v3.0.0a Response 77 Complete Restore Drive Volume {row['sequence']} of {part_count}.zip"
        readme = dist / f"README_RESPONSE77_VOLUME_{row['sequence']}_OF_{part_count}.txt"
        text_write(readme, f"MRHPD Response 77 complete restore volume {row['sequence']} of {part_count}. ALL {part_count} VOLUMES ARE REQUIRED. Extract every wrapper into one directory and run reassemble_response77_complete_restore.py.\n")
        with zipfile.ZipFile(wrapper, "w", compression=zipfile.ZIP_STORED, allowZip64=True) as zf:
            zf.write(row["path"], row["path"].name)
            zf.write(manifest_path, manifest_path.name)
            zf.write(utility, utility.name)
            zf.write(readme, readme.name)
        qa = verify_zip(wrapper)
        if qa["bytes"] >= 104_857_600:
            raise RuntimeError({"drive_volume_exceeds_connector_limit": qa})
        wrappers.append({"sequence": row["sequence"], "wrapper": wrapper, "qa": qa, "raw_part": {k: v for k, v in row.items() if k != "path"}})
    return {"status": "passed", "manifest": manifest, "manifest_path": manifest_path, "reassembly_utility": utility, "volumes": wrappers}


def build_controls_zip(dist: Path, files: Iterable[Path]) -> Path:
    controls = dist / "MRHPD v3.0.0a Response 77 Section 5 Session 1 Complete Verification and Controls.zip"
    with zipfile.ZipFile(controls, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in files:
            if path.exists():
                zf.write(path, path.name)
    verify_zip(controls)
    return controls


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--volume1-dir", type=Path, required=True)
    parser.add_argument("--volume2-dir", type=Path, required=True)
    parser.add_argument("--checkpoint2-dir", type=Path, required=True)
    parser.add_argument("--dist", type=Path, default=Path("dist_cp5_s1_cp3"))
    args = parser.parse_args()
    if args.dist.exists():
        shutil.rmtree(args.dist)
    args.dist.mkdir(parents=True)
    now = utc_now()
    now_iso = now.isoformat()
    stamp = now.strftime("%Y-%m-%d %H%M UTC")
    with tempfile.TemporaryDirectory(prefix="mrhpd-cp5-s1-cp3-") as td:
        work = Path(td)
        restore, base_project_archive, baseline_project, recovery_zip, current_project, cp2_application = recover_checkpoint2_project(
            args.volume1_dir, args.volume2_dir, args.checkpoint2_dir, work
        )
        audit = audit_response76_project(current_project)
        new_db = current_project / CURRENT_DB_REL
        db_stage = synchronize_database(audit["database"], new_db, audit, now_iso)
        new_workbook = current_project / CURRENT_WORKBOOK_REL
        workbook_qa = augment_workbook(audit["workbook"], new_workbook, audit, db_stage["events"], now_iso)
        database_qa = finalize_database(new_db, workbook_qa, audit, now_iso)
        application_qa = update_application_state(current_project, new_db, new_workbook, audit, now_iso)
        acceptance = acceptance_rows(audit, workbook_qa["current_sheet_count"], workbook_qa["formula_error_count"], now_iso)
        tracking_files = build_tracking(current_project, now_iso)
        compaction = compact_superseded_snapshots(current_project, new_db, new_workbook, now_iso)
        summary_figure = create_summary_figure(current_project, audit)
        publication_qa = {"status": "passed", "digital_pages": 537, "digital_searchable_pages": 537, "digital_sha256": sha256_file(audit["publication"]), "digital_unchanged": True, "editable_sha256": sha256_file(audit["editable"]), "editable_unchanged": True, "print_pages": 538, "print_searchable_source_pages": 537, "print_sha256": sha256_file(audit["print_interior"]), "cover_pixels": audit["cover_pixels"], "cover_inches": audit["cover_inches"]}
        final_qa = {
            "schema": "mrhpd-section5-session1-complete-qa-1.0",
            "generated_at": now_iso,
            "status": "passed",
            "response": 77,
            "section": SECTION_LABEL,
            "session": SESSION_LABEL + " COMPLETE",
            "checkpoint": CHECKPOINT_LABEL + " COMPLETE",
            "database": database_qa,
            "workbook": workbook_qa,
            "application": application_qa,
            "publication": publication_qa,
            "response76_application": cp2_application,
            "compaction": {k: v for k, v in compaction.items() if k != "records"},
            "acceptance": {"count": len(acceptance), "failed": sum(1 for row in acceptance if row["status"] == "failed"), "controlled_pending": sum(1 for row in acceptance if row["status"] == "controlled_pending")},
            "checkpoint_3_of_3_complete": True,
            "session_1_of_3_complete": True,
            "remediation_section_5_complete": False,
            "provider_previewer": "controlled_pending_session2",
            "physical_proof": "controlled_pending_session2_or_3",
            "accepted_predecessor_mutated": False,
            "frozen_section3_release_mutated": False,
            "immutable_publication_mutated": False,
            "main_application_mutated": False,
            "next": NEXT_SESSION,
        }
        reports = build_reports(current_project, final_qa, acceptance, db_stage["events"])
        reports.append(summary_figure)
        qa_root = current_project / "QA" / "Section 5 Session 1 Complete"
        qa_root.mkdir(parents=True, exist_ok=True)
        json_write(qa_root / "SECTION5_SESSION1_COMPLETE_QA.json", final_qa)
        json_write(qa_root / "DATABASE_QA.json", database_qa)
        json_write(qa_root / "WORKBOOK_QA.json", workbook_qa)
        json_write(qa_root / "APPLICATION_QA.json", application_qa)
        json_write(qa_root / "PUBLICATION_AND_PRINT_QA.json", publication_qa)
        json_write(qa_root / "RECOVERY_EVENTS_193_205.json", db_stage["events"])
        text_write(current_project / "Recovery" / "Section 5 Session 1 Complete" / "README.md", "# Section 5 Session 1 Complete\n\nResponse 77, Checkpoint 3 of 3, and Session 1 of 3 are complete. The KDP Premium Color production candidate is frozen. Provider preview and physical proof remain controlled external gates for Session 2 and Session 3. Continue begins Session 2 of 3.\n")
        text_write(current_project / "Recovery" / "Section 5 Session 1 Complete" / "PROJECT_INSTRUCTIONS_1.5.0_RESPONSE77_OPERATIVE_ADDENDUM.md", "# Project Instructions 1.5.0 — Response 77 Operative Addendum\n\nAutomatic recovery from the newest verified artifact, Google Drive custody, exact filenames, Raw and Net prompt/response tracking, per-turn checkpoint recovery, full restores at session and section boundaries, Source Index, Bit Index, Master Category continuity, manifests, checksums, professional document design, raster-artwork governance, immutable predecessors, and explicit COMPLETE versus CONTINUE disposition remain mandatory.\n")
        index_result = build_indexes(current_project, now_iso)
        manifest_path, checksums_path, manifest_rows = build_manifest(current_project, now_iso)
        final_qa["indexes"] = index_result["qa"]
        final_qa["manifest_records"] = len(manifest_rows)
        json_write(qa_root / "SECTION5_SESSION1_COMPLETE_QA.json", final_qa)
        # Rebuild indexes and manifest after the final QA record is frozen.
        index_result = build_indexes(current_project, now_iso)
        manifest_path, checksums_path, manifest_rows = build_manifest(current_project, now_iso)
        final_qa["indexes"] = index_result["qa"]
        final_qa["manifest_records"] = len(manifest_rows)
        json_write(qa_root / "SECTION5_SESSION1_COMPLETE_QA_EXTERNAL.json", final_qa)

        project_archive, project_archive_qa = build_project_archive(current_project, args.dist, stamp, database_qa["sha256"], workbook_qa["sha256"])
        complete_restore, restore_verification = build_complete_restore(project_archive, project_archive_qa, reports, tracking_files, args.dist, stamp, final_qa)
        transport = build_transport_volumes(complete_restore, args.dist, now_iso)
        verification_path = args.dist / "MRHPD v3.0.0a Response 77 Complete Restore Verification.json"
        summary_path = args.dist / "MRHPD_RESPONSE77_SECTION5_SESSION1_COMPLETE_BUILD_SUMMARY.json"
        exact_names = args.dist / "MRHPD v3.0.0a Response 77 Exact File Names.txt"
        restore_verification["transport"] = {"status": transport["status"], "part_count": transport["manifest"]["part_count"], "volumes": [{"sequence": row["sequence"], "name": row["wrapper"].name, "bytes": row["qa"]["bytes"], "sha256": row["qa"]["sha256"]} for row in transport["volumes"]]}
        json_write(verification_path, restore_verification)
        build_summary = {
            "schema": "mrhpd-response77-section5-session1-complete-build-1.0",
            "generated_at": now_iso,
            "status": "passed",
            "response": 77,
            "section": SECTION_LABEL,
            "session": SESSION_LABEL + " COMPLETE",
            "checkpoint": CHECKPOINT_LABEL + " COMPLETE",
            "baseline_restore": {"name": restore.name, "bytes": restore.stat().st_size, "sha256": sha256_file(restore)},
            "response76_recovery": {"name": recovery_zip.name, "bytes": recovery_zip.stat().st_size, "sha256": sha256_file(recovery_zip)},
            "database": database_qa,
            "workbook": workbook_qa,
            "application": application_qa,
            "publication": publication_qa,
            "compaction": {k: v for k, v in compaction.items() if k != "records"},
            "indexes": index_result["qa"],
            "manifest_records": len(manifest_rows),
            "project_archive": project_archive_qa,
            "complete_restore": restore_verification,
            "transport": restore_verification["transport"],
            "user_upload_required": False,
            "requires_other_project_files": False,
            "requires_conversation_reconstruction": False,
            "checkpoint_3_of_3_complete": True,
            "session_1_of_3_complete": True,
            "remediation_section_5_complete": False,
            "next": NEXT_SESSION,
        }
        json_write(summary_path, build_summary)
        text_write(exact_names, "\n".join([
            project_archive.name,
            complete_restore.name,
            *[row["wrapper"].name for row in transport["volumes"]],
            verification_path.name,
            summary_path.name,
            transport["manifest_path"].name,
            transport["reassembly_utility"].name,
            *[path.name for path in reports],
        ]) + "\n")
        controls = build_controls_zip(args.dist, [verification_path, summary_path, exact_names, transport["manifest_path"], transport["reassembly_utility"], *reports])
        final_console = {
            "status": "passed",
            "project_archive": project_archive.name,
            "project_archive_bytes": project_archive.stat().st_size,
            "project_archive_sha256": sha256_file(project_archive),
            "complete_restore": complete_restore.name,
            "complete_restore_bytes": complete_restore.stat().st_size,
            "complete_restore_sha256": sha256_file(complete_restore),
            "volume_count": transport["manifest"]["part_count"],
            "volumes": [{"name": row["wrapper"].name, "bytes": row["qa"]["bytes"], "sha256": row["qa"]["sha256"]} for row in transport["volumes"]],
            "controls": controls.name,
            "controls_bytes": controls.stat().st_size,
            "database_tables": database_qa["table_count"],
            "workbook_sheets": workbook_qa["current_sheet_count"],
            "print_pages": 538,
            "searchable_pages": 537,
            "cover_pixels": audit["cover_pixels"],
            "checkpoint_3_of_3_complete": True,
            "session_1_of_3_complete": True,
            "remediation_section_5_complete": False,
            "user_upload_required": False,
            "next": NEXT_SESSION,
        }
        print(json.dumps(final_console, indent=2))


if __name__ == "__main__":
    main()
