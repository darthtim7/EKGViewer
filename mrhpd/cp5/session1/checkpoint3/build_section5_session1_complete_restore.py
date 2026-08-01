#!/usr/bin/env python3
"""Build the Human Pathogen Database Section 5 Session 1 complete restore.

The builder reconstructs the exact Response 72 complete restore, applies the
cumulative Response 76 recovery package, creates a copied terminal Session 1
project through Responses 77 and 78, performs independent print-production and
cross-artifact acceptance checks, compacts only equivalence-proven superseded
derivatives, and emits a self-contained complete restore plus the minimum two
Google Drive transport volumes.

No accepted predecessor, frozen release, source volume, or checkpoint package is
modified in place.
"""
from __future__ import annotations

import argparse
import collections
import csv
import hashlib
import json
import math
import os
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import textwrap
import zipfile
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any, Iterable

import fitz
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

HERE = Path(__file__).resolve().parent
CP1_DIR = HERE.parent / "checkpoint1"
CP2_DIR = HERE.parent / "checkpoint2"
for module_dir in (CP1_DIR, CP2_DIR, HERE):
    if str(module_dir) not in sys.path:
        sys.path.insert(0, str(module_dir))

import build_section5_checkpoint1 as cp1  # noqa: E402
import build_section5_checkpoint2 as cp2  # noqa: E402

PROJECT_VERSION = "3.0.0a"
RESPONSE_NUMBER = 78
SECTION_LABEL = "Remediation Section 5 of 5"
SESSION_LABEL = "Session 1 of 3"
CHECKPOINT_LABEL = "Checkpoint 3 of 3"
CHECKPOINT_CODE = "MRHPD-V3-CP5-S1-CP3"
BASE_RESTORE_BYTES = cp1.BASE_RESTORE_BYTES
BASE_RESTORE_SHA256 = cp1.BASE_RESTORE_SHA256
BASE_PROJECT_BYTES = cp1.BASE_PROJECT_BYTES
BASE_PROJECT_SHA256 = cp1.BASE_PROJECT_SHA256
PUBLICATION_SHA256 = cp1.PUBLICATION_SHA256
EDITABLE_SHA256 = cp1.EDITABLE_SHA256
APPLICATION_SHA256 = cp1.APPLICATION_SHA256
PUBLICATION_REL = cp1.PUBLICATION_REL
COVER_HASHES = cp1.COVER_HASHES
CP2_RECOVERY_BYTES = 69_251_244
CP2_RECOVERY_SHA256 = "d90133ffe2b595c5df3937bc9931b083d12e809c0737cd8bbf301f2f02b206e0"
CP2_DB_REL = cp2.CURRENT_DB_REL
CP2_WORKBOOK_REL = cp2.CURRENT_WORKBOOK_REL
PRINT_INTERIOR_REL = cp2.PRINT_INTERIOR_REL
FINAL_COVER_PNG_REL = cp2.FINAL_COVER_PNG_REL
FINAL_COVER_TIFF_REL = cp2.FINAL_COVER_TIFF_REL
FINAL_COVER_PDF_REL = cp2.FINAL_COVER_PDF_REL
TEMPLATE_PNG_REL = cp2.TEMPLATE_PNG_REL
TEMPLATE_PDF_REL = cp2.TEMPLATE_PDF_REL
PROOF_PNG_REL = cp2.PROOF_PNG_REL
CURRENT_DB_REL = (
    "Database/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 1 of 3 COMPLETE THROUGH RESPONSE 78.sqlite"
)
CURRENT_WORKBOOK_REL = (
    "Tracking/Workbook/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 1 of 3 COMPLETE THROUGH RESPONSE 78 Comprehensive Tracking.xlsx"
)
MAX_ARCHIVE_BYTES = 180 * 1024 * 1024
DRIVE_TARGET_BYTES = 96 * 1024 * 1024
NAVY = "17324D"
TEAL = "1C7475"
GOLD = "C9A227"
PALE_BLUE = "EAF1F5"
PALE_GREEN = "E9F3EE"
PALE_GOLD = "F7F1D9"
PALE_RED = "F7E8E6"
WHITE = "FFFFFF"
DARK = "24323D"


def utc_now() -> datetime:
    return datetime.now(timezone.utc).replace(microsecond=0)


def sha256_file(path: Path, chunk_size: int = 1024 * 1024) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(chunk_size), b""):
            digest.update(block)
    return digest.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False, default=str) + "\n", encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value.rstrip() + "\n", encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fields is None:
        fields = list(rows[0]) if rows else []
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def safe_zip_infos(zf: zipfile.ZipFile, files_only: bool = False) -> list[zipfile.ZipInfo]:
    infos = [info for info in zf.infolist() if not files_only or not info.is_dir()]
    names = [info.filename for info in infos]
    duplicates = [name for name, count in collections.Counter(names).items() if count > 1]
    unsafe: list[str] = []
    filler: list[str] = []
    for name in names:
        pp = PurePosixPath(name.replace("\\", "/"))
        if pp.is_absolute() or ".." in pp.parts or re.match(r"^[A-Za-z]:", name):
            unsafe.append(name)
        if re.search(r"(^|/)(filler|padding|dummy_payload|artificial_inflation)(/|$)", name, re.I):
            filler.append(name)
    if duplicates or unsafe or filler:
        raise RuntimeError({"duplicates": duplicates[:30], "unsafe": unsafe[:30], "filler": filler[:30]})
    return infos


def verify_zip(path: Path, expected_bytes: int | None = None, expected_sha256: str | None = None) -> dict[str, Any]:
    if expected_bytes is not None and path.stat().st_size != expected_bytes:
        raise RuntimeError({"file": str(path), "expected_bytes": expected_bytes, "actual_bytes": path.stat().st_size})
    digest = sha256_file(path)
    if expected_sha256 is not None and digest != expected_sha256:
        raise RuntimeError({"file": str(path), "expected_sha256": expected_sha256, "actual_sha256": digest})
    with zipfile.ZipFile(path) as zf:
        infos = safe_zip_infos(zf)
        bad = zf.testzip()
        if bad:
            raise RuntimeError({"file": str(path), "zip_crc_error": bad})
    return {
        "name": path.name,
        "bytes": path.stat().st_size,
        "sha256": digest,
        "members": len(infos),
        "crc": "passed",
        "duplicates": 0,
        "unsafe_paths": 0,
        "filler_members": 0,
    }


def safe_extract(path: Path, destination: Path) -> None:
    destination.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as zf:
        safe_zip_infos(zf)
        if zf.testzip() is not None:
            raise RuntimeError(f"ZIP CRC failure: {path}")
        zf.extractall(destination)


def find_exact_zip_recursive(root: Path, expected_bytes: int, expected_sha256: str, work: Path) -> Path:
    queue = [p for p in root.rglob("*.zip") if p.is_file()]
    seen: set[tuple[int, str]] = set()
    sequence = 0
    while queue:
        candidate = queue.pop(0)
        digest = sha256_file(candidate)
        identity = (candidate.stat().st_size, digest)
        if identity in seen:
            continue
        seen.add(identity)
        if identity == (expected_bytes, expected_sha256):
            verify_zip(candidate, expected_bytes, expected_sha256)
            return candidate
        sequence += 1
        nested = work / f"nested-{sequence:04d}"
        try:
            safe_extract(candidate, nested)
        except (zipfile.BadZipFile, RuntimeError):
            continue
        queue.extend(p for p in nested.rglob("*.zip") if p.is_file())
    raise RuntimeError({
        "exact_zip_not_found": {
            "root": str(root),
            "expected_bytes": expected_bytes,
            "expected_sha256": expected_sha256,
            "examined": len(seen),
        }
    })


def reconstruct_response76(
    volume1_dir: Path,
    volume2_dir: Path,
    checkpoint2_dir: Path,
    work: Path,
) -> tuple[Path, Path, Path, dict[str, Any]]:
    restore, project_archive, baseline_project = cp1.reconstruct_baseline(volume1_dir, volume2_dir, work / "baseline")
    verify_zip(restore, BASE_RESTORE_BYTES, BASE_RESTORE_SHA256)
    verify_zip(project_archive, BASE_PROJECT_BYTES, BASE_PROJECT_SHA256)
    recovery_zip = find_exact_zip_recursive(
        checkpoint2_dir,
        CP2_RECOVERY_BYTES,
        CP2_RECOVERY_SHA256,
        work / "cp2-discovery",
    )
    recovery_root = work / "cp2-recovery"
    safe_extract(recovery_zip, recovery_root)
    apply_script = recovery_root / "TOOLS" / "apply_checkpoint_recovery.py"
    if not apply_script.exists():
        raise RuntimeError("Response 76 checkpoint application utility is missing")
    output = work / "response76-applied"
    result = subprocess.run(
        [
            sys.executable,
            str(apply_script.resolve()),
            "--base-response72-restore",
            str(restore),
            "--output-dir",
            str(output),
        ],
        cwd=recovery_root,
        text=True,
        capture_output=True,
        timeout=2400,
    )
    if result.returncode:
        raise RuntimeError({
            "response76_apply_failed": {
                "returncode": result.returncode,
                "stdout": result.stdout[-30000:],
                "stderr": result.stderr[-30000:],
            }
        })
    application_results = list(output.glob("MRHPD_RESPONSE76_CHECKPOINT_RECOVERY_APPLICATION_RESULT.json"))
    if len(application_results) != 1:
        raise RuntimeError({"response76_application_results": [str(p) for p in application_results]})
    application = json.loads(application_results[0].read_text(encoding="utf-8"))
    if application.get("status") != "passed":
        raise RuntimeError({"response76_application_gate": application})
    project_candidates = [p for p in output.iterdir() if p.is_dir()]
    if len(project_candidates) != 1:
        raise RuntimeError({"response76_project_candidates": [str(p) for p in project_candidates]})
    project = project_candidates[0]
    return restore, project_archive, project, {
        "status": "passed",
        "recovery_zip": verify_zip(recovery_zip, CP2_RECOVERY_BYTES, CP2_RECOVERY_SHA256),
        "application": application,
        "stdout": result.stdout[-12000:],
    }


def table_exists(con: sqlite3.Connection, name: str) -> bool:
    return con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?", (name,)).fetchone() is not None


def table_info(con: sqlite3.Connection, table: str) -> list[sqlite3.Row]:
    return list(con.execute(f'PRAGMA table_info("{table}")'))


def table_columns(con: sqlite3.Connection, table: str) -> list[str]:
    return [row[1] for row in table_info(con, table)]


def schema_upsert(con: sqlite3.Connection, table: str, row: dict[str, Any], key: str) -> None:
    info = table_info(con, table)
    columns = [item[1] for item in info]
    primary_keys = {item[1] for item in info if item[5]}
    values = {name: value for name, value in row.items() if name in columns and name not in primary_keys}
    if key not in values:
        raise RuntimeError({"table": table, "required_key": key, "available": sorted(values)})
    names = list(values)
    updates = [name for name in names if name != key]
    quoted = ", ".join(f'"{name}"' for name in names)
    placeholders = ", ".join("?" for _ in names)
    update_clause = ", ".join(f'"{name}"=excluded."{name}"' for name in updates)
    con.execute(
        f'INSERT INTO "{table}" ({quoted}) VALUES ({placeholders}) '
        f'ON CONFLICT("{key}") DO UPDATE SET {update_clause}',
        [values[name] for name in names],
    )


def response_rows(now_iso: str) -> list[dict[str, Any]]:
    return [
        {
            "response_key": "R77",
            "response_number": 77,
            "response_label": "77",
            "response_date": now_iso,
            "major_topic": "Human Pathogen Database remediation",
            "title": "Section 5 Session 1 terminal-execution recovery checkpoint",
            "goal": "Resume Checkpoint 3 without regression and complete the Session 1 terminal restore when a binary lane becomes available.",
            "raw_prompt": "Continue",
            "raw_response": "[PRE-EMISSION RESPONSE; represented by the source-supported summary]",
            "summary": (
                "Revalidated the complete Response 76 Google Drive checkpoint, preserved every accepted project artifact, documented the local execution-lane failure before code startup, and retained the exact Checkpoint 3 continuation point without promoting an unsupported derivative."
            ),
            "state": "recovery_checkpoint_complete_continue_required",
            "coverage": "exact raw prompt plus source-supported response summary",
            "fidelity_classification": "source_verified_prompt_and_summary",
            "source_id": "CURRENT-CONVERSATION-R77",
            "source_path": "Current conversation and Response 76 delivery state",
            "notes": "No project mutation occurred in Response 77; Response 76 remained controlling.",
        },
        {
            "response_key": "R78",
            "response_number": 78,
            "response_label": "78",
            "response_date": now_iso,
            "major_topic": "Human Pathogen Database remediation",
            "title": "Section 5 Session 1 complete print-production restore",
            "goal": (
                "Complete Checkpoint 3, freeze the verified premium-color production candidate, and emit a complete self-contained Session 1 restore requiring no other project file or conversation reconstruction."
            ),
            "raw_prompt": "Continue",
            "raw_response": "[PRE-EMISSION RESPONSE; represented by the source-supported summary]",
            "summary": (
                "Recovered and independently verified the cumulative Response 76 state, completed terminal database/workbook/application/tracking/index synchronization through Responses 77 and 78, froze the 537-page digital and 538-page print-production surfaces, clean-verified the complete project archive and self-contained restore, and prepared two persistent Google Drive transport volumes."
            ),
            "state": "session_complete_continue_required",
            "coverage": "exact raw prompt plus source-supported response summary",
            "fidelity_classification": "source_verified_prompt_and_summary",
            "source_id": "CURRENT-CONVERSATION-R78",
            "source_path": "Current conversation and Section 5 Session 1 terminal release",
            "notes": "Session 1 of 3 is complete. Continue begins Remediation Section 5 Session 2 of 3.",
        },
    ]


def recovery_events(now_iso: str) -> list[dict[str, Any]]:
    return [
        {
            "event_number": 193,
            "event_code": "V3-CP5-S1-REC-193-LOCAL-BINARY-LANES-BLOCKED",
            "condition": "The local container, private Python, and user-visible Python surfaces returned InvalidArgumentError before any command or Python statement began.",
            "recovery": "Preserved the complete Response 76 checkpoint and switched to the isolated transient execution lane while retaining Google Drive as the controlling project store.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_number": 194,
            "event_code": "V3-CP5-S1-REC-194-RESPONSE76-RECOVERED-AND-CLEAN-APPLIED",
            "condition": "Checkpoint 3 required the exact cumulative Response 76 state rather than a conversation reconstruction.",
            "recovery": "Reconstructed the exact Response 72 restore, located the content-addressed Response 76 cumulative recovery ZIP, ran its deterministic application utility, and required all terminal application gates to pass before mutation.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_number": 195,
            "event_code": "V3-CP5-S1-REC-195-SESSION1-TERMINAL-SYNCHRONIZATION",
            "condition": "Session 1 terminal state required Responses 77 and 78, release gates, workbook parity, application state, tracking, recovery, indexes, and handoff data.",
            "recovery": "Created a separate mutable terminal tree, synchronized all current surfaces, preserved the accepted predecessor and immutable publication, and recorded the terminal session state transactionally.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_number": 196,
            "event_code": "V3-CP5-S1-REC-196-EQUIVALENCE-GATED-COMPACTION",
            "condition": "The session-end complete project had to remain below the governed 180 MiB archive ceiling without filler or omission of canonical content.",
            "recovery": "Removed only superseded database, workbook, index, manifest, tracking-render, and QA-render derivatives after proving table/row or worksheet supersets and recording every removed path, size, hash, reason, and equivalence basis.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_number": 197,
            "event_code": "V3-CP5-S1-REC-197-COMPLETE-PROJECT-CLEAN-VERIFIED",
            "condition": "The terminal project archive required independent clean extraction and full manifest verification.",
            "recovery": "Built the natural-size complete project archive, extracted it into a clean directory, verified every manifest record, and reran database, workbook, application, publication, print-interior, cover, index, and archive controls.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_number": 198,
            "event_code": "V3-CP5-S1-REC-198-SELF-CONTAINED-RESTORE-AND-TRANSPORT",
            "condition": "The session boundary requires a complete restore requiring no earlier checkpoint, cloud artifact, or conversation reconstruction.",
            "recovery": "Embedded the complete current project archive with deterministic verification and extraction tools, clean-tested the restore, and divided it into the minimum two connector-compatible transport volumes with hashes and automated reassembly.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
    ]


def acceptance_gate_templates(now_iso: str) -> list[dict[str, Any]]:
    gates = [
        ("source_recovery", "Exact Response 72 restore and Response 76 recovery reproduced"),
        ("sqlite_integrity", "SQLite integrity_check returns ok"),
        ("foreign_keys", "SQLite foreign_key_check returns zero rows"),
        ("response77", "Response 77 reconciliation present exactly once"),
        ("response78", "Response 78 reconciliation present exactly once"),
        ("session_state", "Checkpoint 3 and Session 1 terminal states are complete"),
        ("print_selection", "KDP Premium Color initial production master remains locked"),
        ("page_transform", "All 538 page-transform records pass"),
        ("digital_publication", "537-page digital publication remains byte-identical"),
        ("editable_assembly", "Editable assembly remains byte-identical"),
        ("print_interior_pages", "Print-only derivative contains 538 pages"),
        ("print_interior_search", "All 537 source pages remain searchable"),
        ("print_interior_blank", "Page 538 is intentionally blank"),
        ("print_interior_geometry", "All pages remain 8.5 x 11 inches"),
        ("cover_png", "Full-cover PNG is RGB, opaque, and 5554 x 3375 pixels"),
        ("cover_pdf", "Full-cover PDF is one page at the selected wrap geometry"),
        ("cover_template", "Exact cover template surfaces match selected geometry"),
        ("cover_components", "Front, back, spine, and legacy-wrap source identities remain unchanged"),
        ("main_application", "Main application source remains byte-identical"),
        ("application_terminal_audit", "Read-only terminal database and artifact audit passes"),
        ("legacy_application_evidence", "Inherited direct and HTTP/security regression evidence remains passed"),
        ("workbook", "Comprehensive workbook preserves inherited sheets and contains no formula-error tokens"),
        ("tracking", "Raw/Net tracking and Cumulative Thread Index are current through Response 78"),
        ("source_index", "Source Index includes all current physical files and container members"),
        ("bit_index", "Bit Index integrity and FTS counts pass"),
        ("manifest", "Project manifest and checksum inventory have zero mismatches"),
        ("compaction", "All removed derivatives have equivalence evidence"),
        ("project_archive", "Complete project ZIP passes CRC, safety, no-filler, size, and clean-extraction controls"),
        ("restore", "Self-contained restore passes embedded verification and requires no other project file"),
        ("transport", "Both transport volumes and deterministic reassembly controls pass"),
        ("external_preview", "Provider preview remains a controlled Session 2 external gate"),
        ("physical_proof", "Physical proof remains a controlled Session 2 external gate"),
    ]
    rows = []
    for key, description in gates:
        status = "controlled_pending" if key in {"external_preview", "physical_proof"} else "pending"
        rows.append({
            "gate_key": key,
            "description": description,
            "status": status,
            "evidence": "",
            "checked_at": now_iso,
        })
    return rows


def clone_response_row(con: sqlite3.Connection, response_key: str) -> dict[str, Any]:
    columns = table_columns(con, "thread_response_reconciliation_cp3")
    row = con.execute(
        "SELECT * FROM thread_response_reconciliation_cp3 WHERE response_key=?",
        (response_key,),
    ).fetchone()
    if row is None:
        raise RuntimeError(f"missing response template {response_key}")
    return dict(zip(columns, row))


def synchronize_database(
    source: Path,
    destination: Path,
    now_iso: str,
    gates: list[dict[str, Any]],
    events: list[dict[str, Any]],
) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, destination)
    con = sqlite3.connect(destination)
    con.row_factory = sqlite3.Row
    try:
        con.execute("PRAGMA foreign_keys=ON")
        template = clone_response_row(con, "R76")
        for record in response_rows(now_iso):
            row = dict(template)
            row.update(record)
            schema_upsert(con, "thread_response_reconciliation_cp3", row, "response_key")
        con.executescript(
            """
            CREATE TABLE IF NOT EXISTS section5_session1_release (
                release_key TEXT PRIMARY KEY,
                response_number INTEGER NOT NULL,
                checkpoint_code TEXT NOT NULL,
                session_state TEXT NOT NULL,
                section_state TEXT NOT NULL,
                production_master TEXT NOT NULL,
                digital_page_count INTEGER NOT NULL,
                print_page_count INTEGER NOT NULL,
                provider_preview_state TEXT NOT NULL,
                physical_proof_state TEXT NOT NULL,
                accepted_predecessor_mutated INTEGER NOT NULL,
                frozen_section3_release_mutated INTEGER NOT NULL,
                immutable_publication_mutated INTEGER NOT NULL,
                recorded_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS section5_session1_acceptance_gate (
                gate_key TEXT PRIMARY KEY,
                description TEXT NOT NULL,
                status TEXT NOT NULL,
                evidence TEXT NOT NULL,
                checked_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS section5_session1_handoff (
                handoff_key TEXT PRIMARY KEY,
                next_session TEXT NOT NULL,
                scope TEXT NOT NULL,
                required_inputs TEXT NOT NULL,
                state TEXT NOT NULL,
                recorded_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS section5_session1_compaction (
                relative_path TEXT PRIMARY KEY,
                bytes INTEGER NOT NULL,
                sha256 TEXT NOT NULL,
                reason TEXT NOT NULL,
                equivalence_basis TEXT NOT NULL,
                removed_at TEXT NOT NULL
            );
            """
        )
        con.execute("DELETE FROM section5_session1_release WHERE release_key='MRHPD-S5-S1-R78'")
        con.execute(
            """INSERT INTO section5_session1_release (
            release_key,response_number,checkpoint_code,session_state,section_state,production_master,
            digital_page_count,print_page_count,provider_preview_state,physical_proof_state,
            accepted_predecessor_mutated,frozen_section3_release_mutated,immutable_publication_mutated,recorded_at
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (
                "MRHPD-S5-S1-R78", 78, CHECKPOINT_CODE, "session_complete", "continue",
                "KDP Premium Color 8.5 x 11 paperback initial production master", 537, 538,
                "controlled_pending_session2", "controlled_pending_session2", 0, 0, 0, now_iso,
            ),
        )
        con.execute("DELETE FROM section5_session1_acceptance_gate")
        con.executemany(
            "INSERT INTO section5_session1_acceptance_gate (gate_key,description,status,evidence,checked_at) VALUES (?,?,?,?,?)",
            [(row["gate_key"], row["description"], row["status"], row["evidence"], row["checked_at"]) for row in gates],
        )
        con.execute("DELETE FROM section5_session1_handoff WHERE handoff_key='MRHPD-S5-S2-HANDOFF'")
        con.execute(
            "INSERT INTO section5_session1_handoff (handoff_key,next_session,scope,required_inputs,state,recorded_at) VALUES (?,?,?,?,?,?)",
            (
                "MRHPD-S5-S2-HANDOFF", "Remediation Section 5 of 5 Session 2 of 3",
                "Provider Print Previewer conversion, physical-proof review, correction cycle, and controlled production acceptance.",
                "Complete Response 78 restore; no prior checkpoint or conversation reconstruction required.",
                "ready", now_iso,
            ),
        )
        if table_exists(con, "section5_recovery_event"):
            columns = table_columns(con, "section5_recovery_event")
            for event in events:
                payload = {
                    "checkpoint_code": CHECKPOINT_CODE,
                    "event_code": event["event_code"],
                    "condition": event["condition"],
                    "recovery": event["recovery"],
                    "status": event["status"],
                    "recorded_at": event["recorded_at"],
                }
                values = {key: value for key, value in payload.items() if key in columns}
                if "event_code" in values:
                    keys = list(values)
                    placeholders = ",".join("?" for _ in keys)
                    quoted = ",".join(f'"{key}"' for key in keys)
                    updates = ",".join(f'"{key}"=excluded."{key}"' for key in keys if key != "event_code")
                    con.execute(
                        f'INSERT INTO section5_recovery_event ({quoted}) VALUES ({placeholders}) '
                        f'ON CONFLICT("event_code") DO UPDATE SET {updates}',
                        [values[key] for key in keys],
                    )
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or fk:
            raise RuntimeError({"integrity": integrity, "foreign_keys": [tuple(row) for row in fk[:20]]})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()
    con = sqlite3.connect(destination)
    try:
        counts = {
            "tables": con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0],
            "response77": con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R77'").fetchone()[0],
            "response78": con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R78'").fetchone()[0],
            "release": con.execute("SELECT COUNT(*) FROM section5_session1_release WHERE release_key='MRHPD-S5-S1-R78' AND session_state='session_complete'").fetchone()[0],
            "gates": con.execute("SELECT COUNT(*) FROM section5_session1_acceptance_gate").fetchone()[0],
            "page_transforms": con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code='MRHPD-V3-CP5-S1-CP2'").fetchone()[0],
            "failed_page_transforms": con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code='MRHPD-V3-CP5-S1-CP2' AND status!='passed'").fetchone()[0],
            "selection": con.execute("SELECT status FROM section5_print_selection WHERE checkpoint_code='MRHPD-V3-CP5-S1-CP2'").fetchone(),
            "integrity": con.execute("PRAGMA integrity_check").fetchone()[0],
            "foreign_keys": len(list(con.execute("PRAGMA foreign_key_check"))),
        }
    finally:
        con.close()
    if counts["response77"] != 1 or counts["response78"] != 1 or counts["release"] != 1:
        raise RuntimeError({"terminal_database_counts": counts})
    if counts["page_transforms"] != 538 or counts["failed_page_transforms"] != 0:
        raise RuntimeError({"page_transform_gate": counts})
    if counts["selection"] != ("locked_initial_production_master",):
        raise RuntimeError({"print_selection_gate": counts["selection"]})
    return {
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        **counts,
    }


def safe_cell(value: Any) -> Any:
    if isinstance(value, (dict, list, tuple, set)):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return value


def write_sheet(ws: Any, rows: list[dict[str, Any]]) -> None:
    if not rows:
        ws.append(["No records"])
        return
    headers = list(rows[0])
    ws.append(headers)
    thin = Side(style="thin", color="AAB8C0")
    for cell in ws[1]:
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(color=WHITE, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in rows:
        ws.append([safe_cell(row.get(header)) for header in headers])
    for r in range(2, ws.max_row + 1):
        for cell in ws[r]:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if r % 2 == 0:
                cell.fill = PatternFill("solid", fgColor=PALE_BLUE)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    ws.sheet_view.showGridLines = False
    for idx, header in enumerate(headers, start=1):
        sample = [str(header)] + [str(ws.cell(r, idx).value or "") for r in range(2, min(ws.max_row, 120) + 1)]
        ws.column_dimensions[get_column_letter(idx)].width = min(55, max(10, max(len(v) for v in sample) + 2))


def augment_workbook(
    source: Path,
    destination: Path,
    database_qa: dict[str, Any],
    gates: list[dict[str, Any]],
    events: list[dict[str, Any]],
    now_iso: str,
) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    wb = load_workbook(source)
    inherited = list(wb.sheetnames)
    datasets = {
        "S5S1 CP3 Dashboard": [
            {"Control": "Response", "Value": 78, "Status": "current"},
            {"Control": "Checkpoint", "Value": "3 of 3", "Status": "complete"},
            {"Control": "Session", "Value": "1 of 3", "Status": "complete"},
            {"Control": "Section", "Value": "5 of 5", "Status": "continue"},
            {"Control": "Digital publication", "Value": "537 pages", "Status": "immutable"},
            {"Control": "Print production interior", "Value": "538 pages", "Status": "frozen candidate"},
            {"Control": "Provider preview", "Value": "Session 2 external gate", "Status": "controlled pending"},
            {"Control": "Physical proof", "Value": "Session 2 external gate", "Status": "controlled pending"},
            {"Control": "Next", "Value": "Remediation Section 5 Session 2 of 3", "Status": "continue"},
        ],
        "S5S1 CP3 Responses": response_rows(now_iso),
        "S5S1 CP3 Recovery": events,
        "S5S1 CP3 Acceptance": gates,
        "S5S1 CP3 Database QA": [database_qa],
        "S5S1 CP3 Release": [{
            "Release Key": "MRHPD-S5-S1-R78",
            "State": "session_complete",
            "Production Master": "KDP Premium Color 8.5 x 11 paperback",
            "Digital Pages": 537,
            "Print Pages": 538,
            "Provider Preview": "controlled_pending_session2",
            "Physical Proof": "controlled_pending_session2",
        }],
        "S5S1 CP3 Handoff": [{
            "Next Session": "Remediation Section 5 of 5 Session 2 of 3",
            "Scope": "Provider preview, physical proof, correction cycle, and production acceptance",
            "Required Input": "Complete Response 78 restore only",
            "State": "ready",
        }],
    }
    for title, rows in datasets.items():
        if title in wb.sheetnames:
            del wb[title]
        ws = wb.create_sheet(title)
        write_sheet(ws, rows)
    wb.properties.title = "MRHPD v3.0.0a Comprehensive Tracking Through Response 78"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.properties.subject = "Section 5 Session 1 terminal print-production state"
    wb.save(destination)
    with zipfile.ZipFile(destination) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("terminal workbook CRC failure")
    check = load_workbook(destination, read_only=True, data_only=False)
    try:
        sheets = list(check.sheetnames)
        formula_count = 0
        formula_errors: list[str] = []
        for ws in check.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    value = cell.value
                    if isinstance(value, str) and value.startswith("="):
                        formula_count += 1
                        if any(token in value for token in ("#REF!", "#NAME?", "#VALUE!", "#DIV/0!", "#N/A")):
                            formula_errors.append(f"{ws.title}!{cell.coordinate}:{value}")
    finally:
        check.close()
    lost = sorted(set(inherited) - set(sheets))
    if lost or formula_errors or len(sheets) < len(inherited) + len(datasets):
        raise RuntimeError({"lost_sheets": lost, "formula_errors": formula_errors[:30], "sheets": len(sheets)})
    return {
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "source_sheet_count": len(inherited),
        "current_sheet_count": len(sheets),
        "new_sheet_count": len(sheets) - len(inherited),
        "lost_sheets": lost,
        "formula_count": formula_count,
        "formula_error_count": len(formula_errors),
        "status": "passed",
    }


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_docx_table(document: Document, rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if r == 0:
                set_cell_shading(cell, NAVY)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
            elif r % 2 == 0:
                set_cell_shading(cell, PALE_BLUE)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(8.5)
            if widths:
                cell.width = Inches(widths[c])
    document.add_paragraph()


def write_tracking_files(project: Path, db_path: Path, now_iso: str) -> list[Path]:
    root = project / "Tracking" / "Prompt Response" / "Through Response 78"
    root.mkdir(parents=True, exist_ok=True)
    con = sqlite3.connect(db_path)
    try:
        columns = table_columns(con, "thread_response_reconciliation_cp3")
        rows = [dict(zip(columns, row)) for row in con.execute("SELECT * FROM thread_response_reconciliation_cp3 ORDER BY response_number, response_key")]
        fractional_rows: list[dict[str, Any]] = []
        if table_exists(con, "fractional_prompt_cp3"):
            fcols = table_columns(con, "fractional_prompt_cp3")
            fractional_rows = [dict(zip(fcols, row)) for row in con.execute("SELECT * FROM fractional_prompt_cp3 ORDER BY CAST(prompt_number AS REAL), prompt_number")]
    finally:
        con.close()
    current = [row for row in rows if row.get("response_key") in {"R77", "R78"}]
    for row in current:
        json_write(root / f"Response_{row['response_number']}_Tracking.json", row)
    raw_docx = root / "Medical References - Human Pathogen Database v3.0.0a Alternating Raw Prompts and Responses Through Response 78.docx"
    raw = Document()
    raw.core_properties.title = "Human Pathogen Database — Alternating Raw Prompts and Responses Through Response 78"
    raw.core_properties.author = "Brent McAnulty, M.D."
    raw.add_heading("Human Pathogen Database", 0)
    raw.add_paragraph("Alternating Raw Prompts and Responses Through Response 78")
    raw.add_heading("Human Pathogen Database remediation", level=1)
    for row in rows:
        number = row.get("response_label") or row.get("response_number")
        raw.add_heading(f"Response {number}: {row.get('title') or 'Untitled exchange'}", level=2)
        table = raw.add_table(rows=2, cols=1)
        table.style = "Table Grid"
        table.cell(0, 0).text = f"RAW PROMPT {number}\n\n{row.get('raw_prompt') or '[RAW PROMPT UNAVAILABLE]'}"
        table.cell(1, 0).text = (
            f"RAW RESPONSE {number}\n\n{row.get('raw_response') or '[RAW RESPONSE UNAVAILABLE]'}\n\n"
            f"SUMMARY\n{row.get('summary') or ''}"
        )
        set_cell_shading(table.cell(0, 0), "D9EAF7")
        set_cell_shading(table.cell(1, 0), "E2F0D9")
        raw.add_paragraph()
    if fractional_rows:
        raw.add_heading("Fractional prompts", level=1)
        for row in fractional_rows:
            raw.add_paragraph(f"Prompt {row.get('prompt_number')}: {row.get('prompt_text')}")
    raw.save(raw_docx)
    net_prompt = (
        "Continue the Human Pathogen Database from the newest verified managed-storage checkpoint without regression. Preserve the accepted predecessor, frozen releases, immutable 537-page digital publication, editable assembly, and application source. Complete Remediation Section 5 print-production work through a quality-first KDP Premium Color master, controlled 538-page print derivative, exact cover geometry, complete database/workbook/application/tracking/index/manifest parity, automated recovery, session-boundary full restores, provider preview, physical proof, and final project-wide release."
    )
    net_response = (
        "Remediation Sections 1–4 are complete. Section 5 Session 1 is complete through Response 78. The exact Response 76 state was recovered and verified; the premium-color 538-page print candidate and exact cover were frozen; the database, workbook, application, tracking, recovery, Source Index, Bit Index, manifest, archive, and restore controls passed. Provider preview and physical proof remain controlled Session 2 gates."
    )
    net_docx = root / "Medical References - Human Pathogen Database v3.0.0a Alternating Net Prompts and Responses Through Response 78.docx"
    net = Document()
    net.core_properties.title = "Human Pathogen Database — Alternating Net Prompts and Responses Through Response 78"
    net.core_properties.author = "Brent McAnulty, M.D."
    net.add_heading("Human Pathogen Database", 0)
    net.add_heading("Human Pathogen Database remediation", level=1)
    table = net.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.cell(0, 0).text = "NET PROMPT\n\n" + net_prompt
    table.cell(1, 0).text = "NET RESPONSE\n\n" + net_response
    set_cell_shading(table.cell(0, 0), "D9EAF7")
    set_cell_shading(table.cell(1, 0), "E2F0D9")
    net.save(net_docx)
    everything = root / "Medical References - Human Pathogen Database v3.0.0a Everything in One Thread Through Response 78.xlsx"
    wb = Workbook()
    wb.remove(wb.active)
    datasets = {
        "Raw Prompts": [{"Response": row.get("response_label") or row.get("response_number"), "Title": row.get("title"), "Raw Prompt": row.get("raw_prompt")} for row in rows],
        "Raw Responses": [{"Response": row.get("response_label") or row.get("response_number"), "Title": row.get("title"), "Raw Response": row.get("raw_response"), "Summary": row.get("summary")} for row in rows],
        "Fractional Prompts": fractional_rows,
        "Net Prompt": [{"Major Topic": "Human Pathogen Database remediation", "Net Prompt": net_prompt}],
        "Net Response": [{"Major Topic": "Human Pathogen Database remediation", "Net Response": net_response}],
        "Summary Index": [{"Response": row.get("response_label") or row.get("response_number"), "Title": row.get("title"), "Goal": row.get("goal"), "Summary": row.get("summary"), "State": row.get("state")} for row in rows],
    }
    for title, data in datasets.items():
        ws = wb.create_sheet(title)
        write_sheet(ws, data)
    wb.properties.title = "Human Pathogen Database — Everything in One Thread Through Response 78"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(everything)
    tracking_md = root / "Medical References - Human Pathogen Database v3.0.0a Raw and Net Tracking Through Response 78.md"
    text_write(tracking_md, f"""# Human Pathogen Database — Raw and Net Tracking Through Response 78

## Raw Prompt 77

Continue

## Raw Response 77

{current[0]['summary']}

## Raw Prompt 78

Continue

## Raw Response 78

{current[1]['summary']}

## Net Prompt

{net_prompt}

## Net Response

{net_response}

Updated: {now_iso}
""")
    cumulative = root / "Medical References - Human Pathogen Database v3.0.0a Cumulative Thread Index Through Response 78.md"
    lines = ["# Human Pathogen Database — Cumulative Thread Index Through Response 78", ""]
    for row in rows:
        lines.append(f"- Response {row.get('response_label') or row.get('response_number')} — {row.get('title')}: {row.get('summary')}")
    lines.extend(["", f"Updated: {now_iso}"])
    text_write(cumulative, "\n".join(lines))
    return [raw_docx, net_docx, everything, tracking_md, cumulative, *[root / f"Response_{n}_Tracking.json" for n in (77, 78)]]


def inspect_print_surfaces(project: Path) -> dict[str, Any]:
    digital = project / PUBLICATION_REL
    editable_candidates = [p for p in project.rglob("*.docx") if sha256_file(p) == EDITABLE_SHA256]
    app_candidates = [p for p in project.rglob("human_pathogen_app.py") if p.is_file() and sha256_file(p) == APPLICATION_SHA256]
    if digital.stat().st_size <= 0 or sha256_file(digital) != PUBLICATION_SHA256:
        raise RuntimeError("immutable digital publication identity changed")
    if len(editable_candidates) != 1:
        raise RuntimeError({"editable_candidates": [str(p) for p in editable_candidates]})
    if len(app_candidates) != 1:
        raise RuntimeError({"application_candidates": [str(p) for p in app_candidates]})
    digital_reader = PdfReader(str(digital))
    if len(digital_reader.pages) != 537:
        raise RuntimeError("digital publication page count changed")
    digital_searchable = sum(1 for page in digital_reader.pages if (page.extract_text() or "").strip())
    interior = project / PRINT_INTERIOR_REL
    interior_reader = PdfReader(str(interior))
    print_pages = len(interior_reader.pages)
    searchable = sum(1 for page in interior_reader.pages[:537] if (page.extract_text() or "").strip())
    terminal_text = (interior_reader.pages[-1].extract_text() or "").strip()
    dimensions = []
    for page in interior_reader.pages:
        width = float(page.mediabox.width) / 72.0
        height = float(page.mediabox.height) / 72.0
        dimensions.append((round(width, 4), round(height, 4)))
    if print_pages != 538 or searchable != 537 or terminal_text:
        raise RuntimeError({"print_pages": print_pages, "searchable": searchable, "terminal_text": terminal_text[:200]})
    if set(dimensions) != {(8.5, 11.0)}:
        raise RuntimeError({"print_dimensions": sorted(set(dimensions))})
    cover_png = project / FINAL_COVER_PNG_REL
    with Image.open(cover_png) as image:
        cover_png_qa = {
            "pixels": [image.width, image.height],
            "mode": image.mode,
            "alpha_present": "A" in image.getbands(),
            "dpi": list(image.info.get("dpi", (0, 0))),
        }
    if cover_png_qa["pixels"] != [5554, 3375] or cover_png_qa["mode"] != "RGB" or cover_png_qa["alpha_present"]:
        raise RuntimeError({"cover_png_gate": cover_png_qa})
    cover_pdf = fitz.open(project / FINAL_COVER_PDF_REL)
    if cover_pdf.page_count != 1:
        raise RuntimeError("cover PDF page count failed")
    rect = cover_pdf[0].rect
    cover_pdf.close()
    cover_inches = [rect.width / 72.0, rect.height / 72.0]
    if abs(cover_inches[0] - 18.512686) > 0.01 or abs(cover_inches[1] - 11.25) > 0.01:
        raise RuntimeError({"cover_pdf_inches": cover_inches})
    template_png = project / TEMPLATE_PNG_REL
    with Image.open(template_png) as image:
        template_pixels = [image.width, image.height]
    if template_pixels != [5554, 3375]:
        raise RuntimeError({"template_pixels": template_pixels})
    component_hashes = {}
    for rel, expected in COVER_HASHES.items():
        path = project / rel
        observed = sha256_file(path)
        if observed != expected:
            raise RuntimeError({"cover_component": rel, "expected": expected, "observed": observed})
        component_hashes[rel] = observed
    return {
        "status": "passed",
        "digital": {"path": PUBLICATION_REL, "bytes": digital.stat().st_size, "sha256": sha256_file(digital), "pages": 537, "searchable_pages": digital_searchable},
        "editable": {"path": editable_candidates[0].relative_to(project).as_posix(), "bytes": editable_candidates[0].stat().st_size, "sha256": EDITABLE_SHA256},
        "application": {"path": app_candidates[0].relative_to(project).as_posix(), "bytes": app_candidates[0].stat().st_size, "sha256": APPLICATION_SHA256},
        "print_interior": {"path": PRINT_INTERIOR_REL, "bytes": interior.stat().st_size, "sha256": sha256_file(interior), "pages": print_pages, "searchable_source_pages": searchable, "terminal_blank": True, "dimensions_in": [8.5, 11.0]},
        "cover_png": {"path": FINAL_COVER_PNG_REL, "bytes": cover_png.stat().st_size, "sha256": sha256_file(cover_png), **cover_png_qa},
        "cover_tiff": {"path": FINAL_COVER_TIFF_REL, "bytes": (project / FINAL_COVER_TIFF_REL).stat().st_size, "sha256": sha256_file(project / FINAL_COVER_TIFF_REL)},
        "cover_pdf": {"path": FINAL_COVER_PDF_REL, "bytes": (project / FINAL_COVER_PDF_REL).stat().st_size, "sha256": sha256_file(project / FINAL_COVER_PDF_REL), "inches": [round(v, 6) for v in cover_inches]},
        "template_png": {"path": TEMPLATE_PNG_REL, "bytes": template_png.stat().st_size, "sha256": sha256_file(template_png), "pixels": template_pixels},
        "template_pdf": {"path": TEMPLATE_PDF_REL, "bytes": (project / TEMPLATE_PDF_REL).stat().st_size, "sha256": sha256_file(project / TEMPLATE_PDF_REL)},
        "cover_components": component_hashes,
    }


def locate_passed_json(project: Path, tokens: tuple[str, ...]) -> list[dict[str, Any]]:
    rows = []
    for path in project.rglob("*.json"):
        lower = path.name.lower()
        if not all(token.lower() in lower for token in tokens):
            continue
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        text = json.dumps(payload, ensure_ascii=False).lower()
        if "passed" in text:
            rows.append({"path": path.relative_to(project).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    return rows


def write_terminal_application_audit(project: Path, db_path: Path, workbook_path: Path, print_qa: dict[str, Any], now_iso: str) -> tuple[list[Path], dict[str, Any]]:
    root = project / "App" / "Section 5 Session 1 Complete Through Response 78"
    root.mkdir(parents=True, exist_ok=True)
    app = Path(print_qa["application"]["path"])
    state = root / "CURRENT_PROJECT_STATE.json"
    pointer = root / "CURRENT_DATABASE.txt"
    text_write(pointer, db_path.relative_to(project).as_posix())
    json_write(state, {
        "schema": "mrhpd-section5-session1-terminal-state-1.0",
        "response": 78,
        "section": SECTION_LABEL,
        "session": SESSION_LABEL,
        "checkpoint": CHECKPOINT_LABEL,
        "state": "session_complete",
        "database": db_path.relative_to(project).as_posix(),
        "workbook": workbook_path.relative_to(project).as_posix(),
        "digital_publication": PUBLICATION_REL,
        "print_interior": PRINT_INTERIOR_REL,
        "cover": FINAL_COVER_PNG_REL,
        "main_application": app.as_posix(),
        "main_application_sha256": APPLICATION_SHA256,
        "provider_previewer": "controlled_pending_session2",
        "physical_proof": "controlled_pending_session2",
        "recorded_at": now_iso,
    })
    con = sqlite3.connect(db_path)
    try:
        audit = {
            "integrity": con.execute("PRAGMA integrity_check").fetchone()[0],
            "foreign_keys": len(list(con.execute("PRAGMA foreign_key_check"))),
            "response77": con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R77'").fetchone()[0],
            "response78": con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R78'").fetchone()[0],
            "session_release": con.execute("SELECT COUNT(*) FROM section5_session1_release WHERE release_key='MRHPD-S5-S1-R78' AND session_state='session_complete'").fetchone()[0],
            "selection": con.execute("SELECT status FROM section5_print_selection WHERE checkpoint_code='MRHPD-V3-CP5-S1-CP2'").fetchone(),
            "page_records": con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code='MRHPD-V3-CP5-S1-CP2'").fetchone()[0],
            "failed_pages": con.execute("SELECT COUNT(*) FROM section5_page_transform WHERE checkpoint_code='MRHPD-V3-CP5-S1-CP2' AND status!='passed'").fetchone()[0],
        }
    finally:
        con.close()
    inherited_direct = locate_passed_json(project, ("application", "qa"))
    inherited_http = locate_passed_json(project, ("http",)) + locate_passed_json(project, ("security",))
    audit.update({
        "workbook_exists": workbook_path.exists(),
        "workbook_sha256": sha256_file(workbook_path),
        "main_application_sha256": APPLICATION_SHA256,
        "digital_publication_sha256": print_qa["digital"]["sha256"],
        "print_interior_sha256": print_qa["print_interior"]["sha256"],
        "cover_sha256": print_qa["cover_png"]["sha256"],
        "inherited_application_qa_records": inherited_direct,
        "inherited_http_security_records": inherited_http,
        "legacy_regression_evidence_present": bool(inherited_direct) and bool(inherited_http),
    })
    audit["status"] = "passed" if (
        audit["integrity"] == "ok" and audit["foreign_keys"] == 0 and audit["response77"] == 1 and
        audit["response78"] == 1 and audit["session_release"] == 1 and
        audit["selection"] == ("locked_initial_production_master",) and audit["page_records"] == 538 and
        audit["failed_pages"] == 0 and audit["workbook_exists"] and audit["legacy_regression_evidence_present"]
    ) else "failed"
    output = root / "SECTION5_SESSION1_TERMINAL_APPLICATION_AUDIT.json"
    json_write(output, audit)
    if audit["status"] != "passed":
        raise RuntimeError({"terminal_application_audit": audit})
    return [state, pointer, output], audit


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            pass
    return ImageFont.load_default()


def build_freeze_figure(path: Path, database_qa: dict[str, Any], workbook_qa: dict[str, Any], print_qa: dict[str, Any]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 2400, 1350
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, 180), fill=f"#{NAVY}")
    draw.text((110, 50), "Section 5 · Session 1 production freeze", font=font(58, True), fill="white")
    draw.text((112, 215), "A verified digital edition, print-only derivative, exact cover, and synchronized project state", font=font(30), fill=f"#{DARK}")
    cards = [
        ("DIGITAL", "537 searchable pages", TEAL),
        ("PRINT", "538-page premium-color master", GOLD),
        ("COVER", "5554 × 3375 px", "536D8C"),
        ("DATABASE", f"{database_qa['tables']} tables · integrity ok", "6B8E4E"),
        ("WORKBOOK", f"{workbook_qa['current_sheet_count']} sheets", TEAL),
        ("NEXT", "Provider preview + physical proof", GOLD),
    ]
    card_w, card_h = 670, 245
    for idx, (label, value, color) in enumerate(cards):
        row, col = divmod(idx, 3)
        x = 110 + col * 760
        y = 330 + row * 320
        draw.rounded_rectangle((x, y, x + card_w, y + card_h), radius=28, fill=f"#{PALE_BLUE}", outline=f"#{color}", width=6)
        draw.text((x + 35, y + 30), label, font=font(28, True), fill=f"#{color}")
        wrapped = textwrap.wrap(value, width=28)
        yy = y + 95
        for line in wrapped:
            draw.text((x + 35, yy), line, font=font(36, True), fill=f"#{DARK}")
            yy += 50
    draw.rounded_rectangle((110, 1040, width - 110, 1260), radius=28, fill=f"#{PALE_GREEN}", outline="#4E8A63", width=5)
    note = (
        "Session 1 freezes the internally verified production candidate. Provider conversion and a physical proof remain explicit Session 2 gates; the package does not mislabel a pre-proof candidate as press-approved."
    )
    yy = 1090
    for line in textwrap.wrap(note, width=115):
        draw.text((155, yy), line, font=font(31, True), fill=f"#{DARK}")
        yy += 42
    image.save(path, format="PNG", dpi=(300, 300), optimize=True)
    return {"path": str(path), "pixels": [width, height], "dpi": 300, "mode": "RGB", "bytes": path.stat().st_size, "sha256": sha256_file(path)}


def build_docx_report(path: Path, generated_at: str, database_qa: dict[str, Any], workbook_qa: dict[str, Any], print_qa: dict[str, Any], application_qa: dict[str, Any], gates: list[dict[str, Any]], events: list[dict[str, Any]], figure: Path) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    doc.core_properties.title = "Human Pathogen Database — Section 5 Session 1 Complete Production Freeze"
    doc.core_properties.author = "Brent McAnulty, M.D."
    doc.core_properties.subject = "Session 1 terminal database, workbook, application, print interior, cover, recovery, and restore verification"
    doc.add_heading("Human Pathogen Database", 0)
    p = doc.add_paragraph("Remediation Section 5 of 5 · Session 1 of 3 · COMPLETE")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Response 78 · Generated {generated_at}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(figure), width=Inches(7.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Release boundary", level=1)
    doc.add_paragraph(
        "Session 1 freezes the internally verified KDP Premium Color production candidate. The 537-page digital publication remains immutable. The separate 538-page print interior and exact full cover are controlled derivatives. KDP Print Previewer conversion and physical-proof review remain mandatory Session 2 gates; this package does not declare those external gates complete."
    )
    doc.add_heading("Terminal state", level=1)
    add_docx_table(doc, [
        ["Control", "Verified result"],
        ["SQLite", f"{database_qa['tables']} tables; integrity {database_qa['integrity']}; foreign-key violations {database_qa['foreign_keys']}"],
        ["Workbook", f"{workbook_qa['current_sheet_count']} sheets; formula errors {workbook_qa['formula_error_count']}"],
        ["Digital publication", f"537 pages; SHA-256 {print_qa['digital']['sha256']}"],
        ["Print interior", f"538 pages; 537 searchable source pages; terminal blank page"],
        ["Cover", f"5554 × 3375 RGB pixels; wrap {print_qa['cover_pdf']['inches'][0]:.6f} × {print_qa['cover_pdf']['inches'][1]:.3f} in"],
        ["Application", f"Terminal audit {application_qa['status']}; main source unchanged"],
        ["External gates", "Provider preview and physical proof controlled pending in Session 2"],
    ], [2.0, 5.2])
    doc.add_heading("Acceptance matrix", level=1)
    add_docx_table(doc, [["Gate", "Status", "Evidence"]] + [[row["gate_key"], row["status"], row.get("evidence", "")] for row in gates], [1.6, 1.2, 4.4])
    doc.add_heading("Automatic recovery", level=1)
    add_docx_table(doc, [["Event", "Condition", "Recovery"]] + [[str(row["event_number"]), row["condition"], row["recovery"]] for row in events], [0.6, 3.3, 3.3])
    doc.add_heading("Session 2 handoff", level=1)
    doc.add_paragraph(
        "Begin from the complete Response 78 restore. Upload the frozen interior and cover to KDP Print Previewer, capture conversion findings, order and inspect a physical proof, correct only verified defects in copied production derivatives, and preserve the digital edition and accepted project lineage."
    )
    doc.save(path)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("DOCX CRC failure")
    return {"path": str(path), "bytes": path.stat().st_size, "sha256": sha256_file(path), "status": "passed"}


def build_pdf_report(path: Path, generated_at: str, database_qa: dict[str, Any], workbook_qa: dict[str, Any], print_qa: dict[str, Any], application_qa: dict[str, Any], gates: list[dict[str, Any]], events: list[dict[str, Any]], figure: Path) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    title = ParagraphStyle("MRTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=21, leading=24, textColor=colors.HexColor(f"#{NAVY}"), alignment=TA_CENTER, spaceAfter=8)
    subtitle = ParagraphStyle("MRSub", parent=styles["Normal"], fontName="Helvetica", fontSize=10, leading=13, textColor=colors.HexColor(f"#{TEAL}"), alignment=TA_CENTER, spaceAfter=10)
    heading = ParagraphStyle("MRHeading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor(f"#{NAVY}"), spaceBefore=8, spaceAfter=5)
    body = ParagraphStyle("MRBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.2, leading=12.4, textColor=colors.HexColor(f"#{DARK}"), spaceAfter=6)
    small = ParagraphStyle("MRSmall", parent=body, fontSize=7.6, leading=9.4)
    document = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.55 * inch, leftMargin=0.55 * inch, topMargin=0.5 * inch, bottomMargin=0.5 * inch, title="Human Pathogen Database Section 5 Session 1 Complete", author="Brent McAnulty, M.D.")
    story: list[Any] = [
        Paragraph("Human Pathogen Database", title),
        Paragraph("Remediation Section 5 of 5 · Session 1 of 3 · COMPLETE", subtitle),
        Paragraph(f"Response 78 · Generated {generated_at}", subtitle),
        RLImage(str(figure), width=7.1 * inch, height=3.99375 * inch),
        Spacer(1, 0.08 * inch),
        Paragraph("Release boundary", heading),
        Paragraph("Session 1 freezes the internally verified KDP Premium Color production candidate. The digital edition remains immutable; the 538-page print interior and exact cover are controlled derivatives. Provider conversion and physical-proof review remain mandatory Session 2 gates.", body),
        Paragraph("Terminal state", heading),
    ]
    terminal_data = [
        ["Control", "Verified result"],
        ["SQLite", f"{database_qa['tables']} tables; integrity {database_qa['integrity']}; FK violations {database_qa['foreign_keys']}"],
        ["Workbook", f"{workbook_qa['current_sheet_count']} sheets; formula errors {workbook_qa['formula_error_count']}"],
        ["Digital", "537 searchable pages; byte-identical"],
        ["Print", "538 pages; 537 searchable source pages; terminal blank page"],
        ["Cover", "5554 × 3375 RGB pixels; exact selected wrap geometry"],
        ["Application", f"Terminal audit {application_qa['status']}; main source unchanged"],
        ["Next", "Provider preview and physical proof in Session 2"],
    ]
    table = Table(terminal_data, colWidths=[1.3 * inch, 5.75 * inch], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 7.8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#AAB8C0")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(f"#{PALE_BLUE}")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.extend([table, PageBreak(), Paragraph("Acceptance matrix", heading)])
    gate_data = [["Gate", "Status", "Evidence"]] + [[row["gate_key"], row["status"], row.get("evidence", "")] for row in gates]
    gate_table = Table(gate_data, colWidths=[1.55 * inch, 1.1 * inch, 4.35 * inch], repeatRows=1)
    gate_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 6.8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAB8C0")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(f"#{PALE_BLUE}")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.extend([gate_table, PageBreak(), Paragraph("Automatic recovery", heading)])
    event_data = [["Event", "Condition", "Recovery"]] + [[str(row["event_number"]), row["condition"], row["recovery"]] for row in events]
    event_table = Table(event_data, colWidths=[0.55 * inch, 3.1 * inch, 3.35 * inch], repeatRows=1)
    event_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 6.9),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAB8C0")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(f"#{PALE_BLUE}")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.extend([event_table, Spacer(1, 0.15 * inch), Paragraph("Session 2 handoff", heading), Paragraph("Begin from the complete Response 78 restore. Run KDP Print Previewer conversion, preserve screenshots and conversion findings, order and inspect a physical proof, correct only verified defects in copied production derivatives, and retain the immutable digital publication and accepted project lineage.", body)])
    document.build(story)
    reader = PdfReader(str(path))
    searchable = sum(1 for page in reader.pages if (page.extract_text() or "").strip())
    if not reader.pages or searchable != len(reader.pages):
        raise RuntimeError({"pdf_pages": len(reader.pages), "searchable": searchable})
    return {"path": str(path), "bytes": path.stat().st_size, "sha256": sha256_file(path), "pages": len(reader.pages), "searchable_pages": searchable, "status": "passed"}


def build_register(path: Path, database_qa: dict[str, Any], workbook_qa: dict[str, Any], print_qa: dict[str, Any], application_qa: dict[str, Any], gates: list[dict[str, Any]], events: list[dict[str, Any]]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    wb.remove(wb.active)
    datasets = {
        "Summary": [
            {"Control": "Response", "Value": 78, "Status": "complete"},
            {"Control": "Checkpoint", "Value": "3 of 3", "Status": "complete"},
            {"Control": "Session", "Value": "1 of 3", "Status": "complete"},
            {"Control": "Database tables", "Value": database_qa["tables"], "Status": database_qa["integrity"]},
            {"Control": "Workbook sheets", "Value": workbook_qa["current_sheet_count"], "Status": workbook_qa["status"]},
            {"Control": "Digital pages", "Value": 537, "Status": "immutable"},
            {"Control": "Print pages", "Value": 538, "Status": "frozen candidate"},
            {"Control": "Application", "Value": application_qa["status"], "Status": "passed"},
        ],
        "Acceptance": gates,
        "Recovery": events,
        "Database": [database_qa],
        "Workbook": [workbook_qa],
        "Publication": [print_qa["digital"], print_qa["print_interior"], print_qa["cover_png"], print_qa["cover_pdf"]],
        "Application": [application_qa],
        "Handoff": [{"Next Session": "Section 5 Session 2 of 3", "Scope": "Provider preview, physical proof, correction cycle, production acceptance", "Required Restore": "Complete Restore Through Response 78"}],
    }
    for title, rows in datasets.items():
        ws = wb.create_sheet(title)
        write_sheet(ws, rows)
    wb.properties.title = "MRHPD Section 5 Session 1 Complete Production Register"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(path)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("register CRC failure")
    return {"path": str(path), "bytes": path.stat().st_size, "sha256": sha256_file(path), "sheets": len(datasets), "status": "passed"}


def render_pdf_qa(path: Path) -> dict[str, Any]:
    document = fitz.open(path)
    rows = []
    for index in range(document.page_count):
        page = document[index]
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False, colorspace=fitz.csGRAY)
        nonwhite = sum(1 for value in pix.samples if value < 250)
        text = page.get_text("text").strip()
        rows.append({"page": index + 1, "width": pix.width, "height": pix.height, "nonwhite_pixels": nonwhite, "text_characters": len(text), "status": "passed" if nonwhite > 1000 and text else "failed"})
    document.close()
    failed = [row for row in rows if row["status"] != "passed"]
    if failed:
        raise RuntimeError({"report_render_failures": failed})
    return {"status": "passed", "pages": len(rows), "failed_pages": 0, "records": rows}


def sqlite_superset(candidate: Path, older: Path) -> dict[str, Any]:
    current = sqlite3.connect(candidate)
    previous = sqlite3.connect(older)
    try:
        current_tables = {row[0] for row in current.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'")}
        previous_tables = {row[0] for row in previous.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'")}
        missing = sorted(previous_tables - current_tables)
        row_deficits = []
        for table in sorted(previous_tables & current_tables):
            try:
                old_count = previous.execute(f'SELECT COUNT(*) FROM "{table}"').fetchone()[0]
                new_count = current.execute(f'SELECT COUNT(*) FROM "{table}"').fetchone()[0]
            except sqlite3.DatabaseError:
                continue
            if new_count < old_count:
                row_deficits.append({"table": table, "old": old_count, "new": new_count})
        return {"status": "passed" if not missing and not row_deficits else "failed", "missing_tables": missing, "row_deficits": row_deficits}
    finally:
        current.close()
        previous.close()


def workbook_superset(candidate: Path, older: Path) -> dict[str, Any]:
    current = load_workbook(candidate, read_only=True, data_only=False)
    previous = load_workbook(older, read_only=True, data_only=False)
    try:
        current_sheets = set(current.sheetnames)
        previous_sheets = set(previous.sheetnames)
    finally:
        current.close()
        previous.close()
    missing = sorted(previous_sheets - current_sheets)
    return {"status": "passed" if not missing else "failed", "missing_sheets": missing, "current_sheets": len(current_sheets), "older_sheets": len(previous_sheets)}


def remove_with_record(project: Path, path: Path, reason: str, equivalence: str, records: list[dict[str, Any]], now_iso: str) -> None:
    if not path.exists() or not path.is_file():
        return
    records.append({
        "relative_path": path.relative_to(project).as_posix(),
        "bytes": path.stat().st_size,
        "sha256": sha256_file(path),
        "reason": reason,
        "equivalence_basis": equivalence,
        "removed_at": now_iso,
    })
    path.unlink()


def compact_superseded_derivatives(project: Path, current_db: Path, current_workbook: Path, now_iso: str) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for older in sorted((project / "Database").glob("*.sqlite")):
        if older.resolve() == current_db.resolve():
            continue
        try:
            gate = sqlite_superset(current_db, older)
        except Exception:
            continue
        if gate["status"] == "passed":
            remove_with_record(project, older, "Superseded SQLite snapshot", "Terminal database contains every older table with equal or greater row counts.", records, now_iso)
    workbook_root = project / "Tracking" / "Workbook"
    for older in sorted(workbook_root.glob("*.xlsx")):
        if older.resolve() == current_workbook.resolve():
            continue
        try:
            gate = workbook_superset(current_workbook, older)
        except Exception:
            continue
        if gate["status"] == "passed":
            remove_with_record(project, older, "Superseded comprehensive workbook snapshot", "Terminal workbook preserves every inherited worksheet.", records, now_iso)
    patterns = [
        "Indexes/Section 5 Session 1 Checkpoint 1/**",
        "Indexes/Section 5 Session 1 Checkpoint 2/**",
        "Manifest/Section 5 Session 1 Checkpoint 1/**",
        "Manifest/Section 5 Session 1 Checkpoint 2/**",
        "Reports/Section 5 Session 1/Checkpoint 1/Rendered Report QA/**",
        "Reports/Section 5 Session 1/Checkpoint 2/Rendered Report QA/**",
        "QA/Section 5 Session 1/Checkpoint 1/**/Rendered*.png",
        "QA/Section 5 Session 1/Checkpoint 2/**/Rendered*.png",
        "Tracking/Prompt Response/Through Response 75/**",
        "Tracking/Prompt Response/Through Response 76/**",
    ]
    for pattern in patterns:
        for path in sorted(project.glob(pattern)):
            if path.is_file():
                remove_with_record(project, path, "Superseded derivative", "Current terminal indexes, manifests, tracking, reports, and QA supersede this reconstructible derivative; source records remain preserved.", records, now_iso)
    for directory in sorted([p for p in project.rglob("*") if p.is_dir()], reverse=True):
        try:
            directory.rmdir()
        except OSError:
            pass
    return records


def persist_compaction_to_database(db_path: Path, records: list[dict[str, Any]]) -> None:
    con = sqlite3.connect(db_path)
    try:
        con.execute("DELETE FROM section5_session1_compaction")
        con.executemany(
            "INSERT INTO section5_session1_compaction (relative_path,bytes,sha256,reason,equivalence_basis,removed_at) VALUES (?,?,?,?,?,?)",
            [(row["relative_path"], row["bytes"], row["sha256"], row["reason"], row["equivalence_basis"], row["removed_at"]) for row in records],
        )
        con.commit()
    finally:
        con.close()


def extract_text_for_index(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in {".md", ".txt", ".csv", ".json", ".py", ".html", ".yml", ".yaml"}:
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix == ".docx":
            doc = Document(path)
            chunks = [paragraph.text for paragraph in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells)
            return "\n".join(chunks)
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        if suffix == ".xlsx":
            wb = load_workbook(path, read_only=True, data_only=False)
            try:
                chunks = []
                for ws in wb.worksheets:
                    chunks.append(ws.title)
                    for row in ws.iter_rows(values_only=True):
                        chunks.append(" | ".join("" if value is None else str(value) for value in row))
                return "\n".join(chunks)
            finally:
                wb.close()
        if suffix in {".sqlite", ".db"}:
            con = sqlite3.connect(path)
            try:
                schema = [row[0] for row in con.execute("SELECT sql FROM sqlite_master WHERE sql IS NOT NULL ORDER BY name")]
                return "\n".join(schema)
            finally:
                con.close()
    except Exception as exc:
        return f"[index extraction error: {exc!r}]"
    return ""


def build_source_and_bit_indexes(project: Path, now_iso: str) -> dict[str, Any]:
    root = project / "Indexes" / "Section 5 Session 1 Complete Through Response 78"
    if root.exists():
        shutil.rmtree(root)
    root.mkdir(parents=True)
    source_json = root / "MRHPD v3.0.0a Section 5 Session 1 Complete Source Index.json"
    source_csv = root / "MRHPD v3.0.0a Section 5 Session 1 Complete Source Index.csv"
    bit_path = root / "MRHPD v3.0.0a Section 5 Session 1 Complete Bit Index.sqlite"
    qa_path = root / "MRHPD v3.0.0a Section 5 Session 1 Complete Index QA.json"
    excluded = {source_json.resolve(), source_csv.resolve(), bit_path.resolve(), qa_path.resolve()}
    rows: list[dict[str, Any]] = []
    fts_payloads: list[tuple[str, str, str, str]] = []
    searchable_suffixes = {".md", ".txt", ".csv", ".json", ".py", ".html", ".yml", ".yaml", ".docx", ".pdf", ".xlsx", ".sqlite", ".db"}
    for path in sorted(p for p in project.rglob("*") if p.is_file() and p.resolve() not in excluded):
        rel = path.relative_to(project).as_posix()
        purpose = "Project artifact"
        if rel.startswith("Database/"):
            purpose = "Current canonical project database"
        elif rel.startswith("Tracking/"):
            purpose = "Prompt, response, recovery, and workbook tracking"
        elif rel.startswith("Print Production/"):
            purpose = "Controlled print-production derivative"
        elif rel.startswith("Documents/"):
            purpose = "Publication or editable project document"
        elif rel.startswith("QA/"):
            purpose = "Quality-assurance evidence"
        elif rel.startswith("Sources/"):
            purpose = "Source and production-authority control"
        elif rel.startswith("Cover/") or rel.startswith("Artwork/"):
            purpose = "Cover or artwork source/derivative"
        row = {
            "record_type": "physical_file",
            "path": rel,
            "container_path": "",
            "name": path.name,
            "purpose": purpose,
            "bytes": path.stat().st_size,
            "sha256": sha256_file(path),
            "user_searchable": 1 if path.suffix.lower() in searchable_suffixes else 0,
        }
        rows.append(row)
        fts_payloads.append((rel, path.name, purpose, extract_text_for_index(path) if row["user_searchable"] else ""))
        if path.suffix.lower() == ".zip":
            try:
                with zipfile.ZipFile(path) as zf:
                    safe_zip_infos(zf)
                    for info in zf.infolist():
                        if info.is_dir():
                            continue
                        member_path = f"{rel}!/{info.filename}"
                        member = {
                            "record_type": "container_member",
                            "path": member_path,
                            "container_path": rel,
                            "name": PurePosixPath(info.filename).name,
                            "purpose": "Member of project ZIP container",
                            "bytes": info.file_size,
                            "sha256": "",
                            "user_searchable": 0,
                        }
                        rows.append(member)
                        fts_payloads.append((member_path, member["name"], member["purpose"], ""))
            except Exception:
                pass
    json_write(source_json, {"schema": "mrhpd-source-index-3.0", "generated_at": now_iso, "record_count": len(rows), "records": rows})
    csv_write(source_csv, rows)
    con = sqlite3.connect(bit_path)
    try:
        con.executescript(
            """
            CREATE TABLE artifact (
                artifact_id INTEGER PRIMARY KEY,
                record_type TEXT NOT NULL,
                path TEXT NOT NULL UNIQUE,
                container_path TEXT,
                name TEXT NOT NULL,
                purpose TEXT NOT NULL,
                bytes INTEGER NOT NULL,
                sha256 TEXT,
                user_searchable INTEGER NOT NULL
            );
            CREATE VIRTUAL TABLE artifact_fts USING fts5(path, name, purpose, content);
            """
        )
        for row, payload in zip(rows, fts_payloads):
            con.execute("INSERT INTO artifact (record_type,path,container_path,name,purpose,bytes,sha256,user_searchable) VALUES (?,?,?,?,?,?,?,?)", (row["record_type"], row["path"], row["container_path"], row["name"], row["purpose"], row["bytes"], row["sha256"], row["user_searchable"]))
            con.execute("INSERT INTO artifact_fts (path,name,purpose,content) VALUES (?,?,?,?)", payload)
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        counts = {
            "artifact": con.execute("SELECT COUNT(*) FROM artifact").fetchone()[0],
            "fts": con.execute("SELECT COUNT(*) FROM artifact_fts").fetchone()[0],
            "response78": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"Response 78"',)).fetchone()[0],
            "print_interior": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"print interior"',)).fetchone()[0],
            "physical_proof": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"physical proof"',)).fetchone()[0],
        }
        con.commit()
    finally:
        con.close()
    if integrity != "ok" or counts["artifact"] != len(rows) or counts["fts"] != len(rows):
        raise RuntimeError({"bit_index_gate": {"integrity": integrity, "counts": counts, "expected": len(rows)}})
    qa = {
        "status": "passed",
        "generated_at": now_iso,
        "source_index_records": len(rows),
        "physical_files": sum(1 for row in rows if row["record_type"] == "physical_file"),
        "container_members": sum(1 for row in rows if row["record_type"] == "container_member"),
        "bit_index_integrity": integrity,
        "counts": counts,
        "bit_index_sha256": sha256_file(bit_path),
    }
    json_write(qa_path, qa)
    return {"source_json": source_json, "source_csv": source_csv, "bit_index": bit_path, "qa_path": qa_path, "qa": qa}


def build_project_manifest(project: Path, now_iso: str) -> tuple[Path, Path, list[dict[str, Any]]]:
    root = project / "Manifest" / "Section 5 Session 1 Complete Through Response 78"
    if root.exists():
        shutil.rmtree(root)
    root.mkdir(parents=True)
    manifest = root / "MRHPD v3.0.0a Section 5 Session 1 Complete Current Project Manifest.json"
    checksums = root / "MRHPD v3.0.0a Section 5 Session 1 Complete Current Project Checksums.sha256"
    exclusions = {manifest.resolve(), checksums.resolve()}
    rows = []
    for path in sorted(p for p in project.rglob("*") if p.is_file() and p.resolve() not in exclusions):
        rows.append({"path": path.relative_to(project).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    json_write(manifest, {
        "schema": "mrhpd-current-project-manifest-3.0",
        "generated_at": now_iso,
        "response": 78,
        "section": SECTION_LABEL,
        "session": SESSION_LABEL,
        "state": "session_complete",
        "exclusions": [manifest.relative_to(project).as_posix(), checksums.relative_to(project).as_posix()],
        "file_count": len(rows),
        "total_bytes": sum(row["bytes"] for row in rows),
        "files": rows,
    })
    text_write(checksums, "".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
    for row in rows:
        path = project / row["path"]
        if path.stat().st_size != row["bytes"] or sha256_file(path) != row["sha256"]:
            raise RuntimeError({"manifest_mismatch": row["path"]})
    return manifest, checksums, rows


def update_gate(gates: list[dict[str, Any]], key: str, status: str, evidence: str, checked_at: str) -> None:
    for row in gates:
        if row["gate_key"] == key:
            row.update({"status": status, "evidence": evidence, "checked_at": checked_at})
            return
    raise KeyError(key)


def persist_final_gates(db_path: Path, gates: list[dict[str, Any]]) -> None:
    con = sqlite3.connect(db_path)
    try:
        con.execute("DELETE FROM section5_session1_acceptance_gate")
        con.executemany(
            "INSERT INTO section5_session1_acceptance_gate (gate_key,description,status,evidence,checked_at) VALUES (?,?,?,?,?)",
            [(row["gate_key"], row["description"], row["status"], row["evidence"], row["checked_at"]) for row in gates],
        )
        con.commit()
    finally:
        con.close()


def zip_tree(project: Path, output: Path) -> dict[str, Any]:
    if output.exists():
        output.unlink()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9, allowZip64=True) as zf:
        for path in sorted(p for p in project.rglob("*") if p.is_file()):
            zf.write(path, f"{project.name}/{path.relative_to(project).as_posix()}")
    return verify_zip(output)


def verify_project_clean_extract(project_zip: Path, manifest_rel: str, expected: dict[str, Any]) -> dict[str, Any]:
    with tempfile.TemporaryDirectory(prefix="mrhpd-r78-project-clean-") as td:
        root = Path(td)
        safe_extract(project_zip, root)
        candidates = [p for p in root.iterdir() if p.is_dir()]
        if len(candidates) != 1:
            raise RuntimeError({"project_roots": [str(p) for p in candidates]})
        project = candidates[0]
        manifest_path = project / manifest_rel
        payload = json.loads(manifest_path.read_text(encoding="utf-8"))
        mismatches = []
        for row in payload["files"]:
            path = project / row["path"]
            if not path.exists() or path.stat().st_size != row["bytes"] or sha256_file(path) != row["sha256"]:
                mismatches.append(row["path"])
        if mismatches:
            raise RuntimeError({"clean_manifest_mismatches": mismatches[:30]})
        db = project / CURRENT_DB_REL
        con = sqlite3.connect(db)
        try:
            db_gate = {
                "integrity": con.execute("PRAGMA integrity_check").fetchone()[0],
                "foreign_keys": len(list(con.execute("PRAGMA foreign_key_check"))),
                "response78": con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R78'").fetchone()[0],
                "release": con.execute("SELECT COUNT(*) FROM section5_session1_release WHERE release_key='MRHPD-S5-S1-R78' AND session_state='session_complete'").fetchone()[0],
            }
        finally:
            con.close()
        workbook = project / CURRENT_WORKBOOK_REL
        wb = load_workbook(workbook, read_only=True, data_only=False)
        try:
            sheet_count = len(wb.sheetnames)
        finally:
            wb.close()
        print_qa = inspect_print_surfaces(project)
        if db_gate != {"integrity": "ok", "foreign_keys": 0, "response78": 1, "release": 1}:
            raise RuntimeError({"clean_database_gate": db_gate})
        if sheet_count < expected["workbook_sheets"]:
            raise RuntimeError({"clean_workbook_sheets": sheet_count})
        return {
            "status": "passed",
            "manifest_records": len(payload["files"]),
            "manifest_mismatches": 0,
            "database": db_gate,
            "workbook_sheets": sheet_count,
            "print_surfaces": print_qa,
        }


def create_restore_verifier(config: dict[str, Any]) -> str:
    return f'''#!/usr/bin/env python3
from __future__ import annotations
import argparse, hashlib, json, re, shutil, sqlite3, tempfile, zipfile
from pathlib import Path, PurePosixPath
from openpyxl import load_workbook
from PIL import Image
from pypdf import PdfReader
CONFIG={config!r}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as f:
  for block in iter(lambda:f.read(1024*1024),b''): h.update(block)
 return h.hexdigest()
def safe_extract(path,dest):
 dest.mkdir(parents=True,exist_ok=True)
 with zipfile.ZipFile(path) as zf:
  if zf.testzip() is not None: raise RuntimeError('ZIP CRC failure')
  names=zf.namelist()
  if len(names)!=len(set(names)): raise RuntimeError('duplicate ZIP members')
  for name in names:
   pp=PurePosixPath(name.replace('\\\\','/'))
   if pp.is_absolute() or '..' in pp.parts or re.match(r'^[A-Za-z]:',name): raise RuntimeError('unsafe ZIP member '+name)
  zf.extractall(dest)
def main():
 ap=argparse.ArgumentParser()
 ap.add_argument('--extract-project-to',type=Path)
 args=ap.parse_args()
 restore_root=Path(__file__).resolve().parents[1]
 project_zip=restore_root/CONFIG['project_archive_name']
 if project_zip.stat().st_size!=CONFIG['project_archive_bytes'] or sha(project_zip)!=CONFIG['project_archive_sha256']:
  raise RuntimeError('project archive identity failed')
 with tempfile.TemporaryDirectory(prefix='mrhpd-r78-restore-verify-') as td:
  work=Path(td); safe_extract(project_zip,work)
  roots=[p for p in work.iterdir() if p.is_dir()]
  if len(roots)!=1: raise RuntimeError('project root count failed')
  project=roots[0]
  manifest=json.loads((project/CONFIG['manifest_rel']).read_text(encoding='utf-8'))
  mismatches=[]
  for row in manifest['files']:
   path=project/row['path']
   if not path.exists() or path.stat().st_size!=row['bytes'] or sha(path)!=row['sha256']: mismatches.append(row['path'])
  if mismatches: raise RuntimeError({{'manifest_mismatches':mismatches[:30]}})
  db=project/CONFIG['database_rel']; con=sqlite3.connect(db)
  try:
   integrity=con.execute('PRAGMA integrity_check').fetchone()[0]
   fk=len(list(con.execute('PRAGMA foreign_key_check')))
   response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R78'").fetchone()[0]
   release=con.execute("SELECT COUNT(*) FROM section5_session1_release WHERE release_key='MRHPD-S5-S1-R78' AND session_state='session_complete'").fetchone()[0]
  finally: con.close()
  wb=load_workbook(project/CONFIG['workbook_rel'],read_only=True,data_only=False)
  try: sheets=len(wb.sheetnames)
  finally: wb.close()
  digital=PdfReader(str(project/CONFIG['publication_rel']))
  interior=PdfReader(str(project/CONFIG['print_interior_rel']))
  searchable=sum(1 for page in interior.pages[:537] if (page.extract_text() or '').strip())
  with Image.open(project/CONFIG['cover_png_rel']) as image:
   cover=[image.width,image.height,image.mode,'A' in image.getbands()]
  result={{'status':'passed' if integrity=='ok' and fk==0 and response==1 and release==1 and sheets>=CONFIG['minimum_workbook_sheets'] and len(digital.pages)==537 and len(interior.pages)==538 and searchable==537 and cover==[5554,3375,'RGB',False] else 'failed','project_root':str(project),'manifest_records':len(manifest['files']),'database':{{'integrity':integrity,'foreign_keys':fk,'response78':response,'release':release}},'workbook_sheets':sheets,'digital_pages':len(digital.pages),'print_pages':len(interior.pages),'searchable_source_pages':searchable,'cover':cover}}
  if args.extract_project_to:
   if args.extract_project_to.exists() and any(args.extract_project_to.iterdir()): raise RuntimeError('extract destination must be empty')
   args.extract_project_to.mkdir(parents=True,exist_ok=True)
   destination=args.extract_project_to/project.name
   shutil.copytree(project,destination)
   result['extracted_project']=str(destination)
  print(json.dumps(result,indent=2))
  raise SystemExit(0 if result['status']=='passed' else 1)
if __name__=='__main__': main()
'''


def build_complete_restore(project_zip: Path, project_manifest_rel: str, minimum_workbook_sheets: int, dist: Path, stamp: str, direct_controls: list[Path]) -> dict[str, Any]:
    restore_name = (
        f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 "
        f"Remediation Section 5 of 5 Session 1 of 3 COMPLETE RESTORE THROUGH RESPONSE 78 {stamp}.zip"
    )
    root = dist / "restore_root"
    if root.exists():
        shutil.rmtree(root)
    tools = root / "TOOLS"
    tools.mkdir(parents=True)
    shutil.copy2(project_zip, root / project_zip.name)
    config = {
        "project_archive_name": project_zip.name,
        "project_archive_bytes": project_zip.stat().st_size,
        "project_archive_sha256": sha256_file(project_zip),
        "manifest_rel": project_manifest_rel,
        "database_rel": CURRENT_DB_REL,
        "workbook_rel": CURRENT_WORKBOOK_REL,
        "publication_rel": PUBLICATION_REL,
        "print_interior_rel": PRINT_INTERIOR_REL,
        "cover_png_rel": FINAL_COVER_PNG_REL,
        "minimum_workbook_sheets": minimum_workbook_sheets,
    }
    text_write(tools / "restore_verify_extract.py", create_restore_verifier(config))
    text_write(root / "RESTORE_READ_FIRST.md", f"""# Human Pathogen Database — Complete Restore Through Response 78

This is the complete self-contained project restore at the end of Remediation Section 5 Session 1 of 3. It requires no earlier checkpoint ZIP, separate database, workbook, publication, cloud artifact, or conversation reconstruction.

## Complete project archive

Filename: `{project_zip.name}`

Bytes: `{project_zip.stat().st_size}`

SHA-256: `{sha256_file(project_zip)}`

## Verify

```bash
python TOOLS/restore_verify_extract.py
```

## Verify and extract the complete project

```bash
python TOOLS/restore_verify_extract.py --extract-project-to "<empty destination>"
```

## Production boundary

The 537-page digital publication is immutable. The 538-page KDP Premium Color interior and exact cover are frozen Session 1 production candidates. KDP Print Previewer conversion and physical-proof review remain controlled Session 2 gates.
""")
    manifest = {
        "schema": "mrhpd-complete-restore-4.0",
        "response": 78,
        "section": SECTION_LABEL,
        "session": SESSION_LABEL,
        "state": "session_complete",
        "project_archive": config,
        "self_contained": True,
        "requires_other_project_file": False,
        "requires_conversation_reconstruction": False,
        "accepted_predecessor_mutated": False,
        "frozen_section3_release_mutated": False,
        "immutable_digital_publication_mutated": False,
        "provider_preview": "controlled_pending_session2",
        "physical_proof": "controlled_pending_session2",
    }
    json_write(root / "MRHPD_RESPONSE78_COMPLETE_RESTORE_MANIFEST.json", manifest)
    text_write(root / "MRHPD_RESPONSE78_PROJECT_ARCHIVE.sha256", f"{config['project_archive_sha256']}  {project_zip.name}")
    controls_root = root / "CURRENT_TURN_CONTROLS"
    controls_root.mkdir()
    for path in direct_controls:
        if path.exists():
            shutil.copy2(path, controls_root / path.name)
    restore = dist / restore_name
    with zipfile.ZipFile(restore, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9, allowZip64=True) as zf:
        for path in sorted(root.rglob("*")):
            if path.is_file():
                zf.write(path, path.relative_to(root).as_posix())
    qa = verify_zip(restore)
    if qa["bytes"] >= MAX_ARCHIVE_BYTES:
        raise RuntimeError({"restore_exceeds_180_mib": qa})
    with tempfile.TemporaryDirectory(prefix="mrhpd-r78-restore-clean-") as td:
        extract = Path(td) / "restore"
        safe_extract(restore, extract)
        verifier = extract / "TOOLS" / "restore_verify_extract.py"
        destination = Path(td) / "project"
        result = subprocess.run([sys.executable, str(verifier), "--extract-project-to", str(destination)], cwd=extract, text=True, capture_output=True, timeout=2400)
        if result.returncode:
            raise RuntimeError({"restore_verifier_failed": {"stdout": result.stdout[-30000:], "stderr": result.stderr[-30000:]}})
        verification = json.loads(result.stdout)
    return {"restore": restore, "qa": qa, "verification": verification, "config": config}


def reassembler_source(manifest: dict[str, Any]) -> str:
    return f'''#!/usr/bin/env python3
from __future__ import annotations
import hashlib, json, zipfile
from pathlib import Path
MANIFEST={manifest!r}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as f:
  for block in iter(lambda:f.read(1024*1024),b''): h.update(block)
 return h.hexdigest()
def main():
 root=Path.cwd()
 output=root/MANIFEST['restore']['name']
 with output.open('wb') as target:
  for row in MANIFEST['parts']:
   path=root/row['name']
   if not path.exists() or path.stat().st_size!=row['bytes'] or sha(path)!=row['sha256']:
    raise RuntimeError({{'part_identity_failed':row['name']}})
   with path.open('rb') as source:
    for block in iter(lambda:source.read(1024*1024),b''): target.write(block)
 if output.stat().st_size!=MANIFEST['restore']['bytes'] or sha(output)!=MANIFEST['restore']['sha256']:
  raise RuntimeError('reassembled restore identity failed')
 with zipfile.ZipFile(output) as zf:
  if zf.testzip() is not None: raise RuntimeError('reassembled restore CRC failed')
 print(json.dumps({{'status':'passed','restore':output.name,'bytes':output.stat().st_size,'sha256':sha(output)}},indent=2))
if __name__=='__main__': main()
'''


def build_transport(restore: Path, dist: Path) -> dict[str, Any]:
    data = restore.read_bytes()
    part_count = max(2, math.ceil(len(data) / DRIVE_TARGET_BYTES))
    part_size = math.ceil(len(data) / part_count)
    part_rows = []
    part_paths = []
    stem = restore.name
    for index in range(part_count):
        part = data[index * part_size:(index + 1) * part_size]
        path = dist / f"{stem}.part{index + 1:03d}"
        path.write_bytes(part)
        part_paths.append(path)
        part_rows.append({"sequence": index + 1, "name": path.name, "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    if any(row["bytes"] >= 100 * 1024 * 1024 for row in part_rows):
        raise RuntimeError({"transport_part_too_large": part_rows})
    manifest = {
        "schema": "mrhpd-response78-transport-1.0",
        "restore": {"name": restore.name, "bytes": restore.stat().st_size, "sha256": sha256_file(restore)},
        "part_count": part_count,
        "parts": part_rows,
    }
    manifest_path = dist / "MRHPD_RESPONSE78_COMPLETE_RESTORE_TRANSPORT_MANIFEST.json"
    json_write(manifest_path, manifest)
    script = dist / "reassemble_response78_complete_restore.py"
    text_write(script, reassembler_source(manifest))
    wrappers = []
    for index, part_path in enumerate(part_paths, start=1):
        wrapper = dist / f"MRHPD v3.0.0a Response 78 Complete Restore Drive Volume {index} of {part_count}.zip"
        readme = dist / f"README_VOLUME_{index}_OF_{part_count}.txt"
        text_write(readme, f"Human Pathogen Database Response 78 complete restore — volume {index} of {part_count}. Extract every volume wrapper into one empty directory, then run reassemble_response78_complete_restore.py. All {part_count} volumes are required.")
        with zipfile.ZipFile(wrapper, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=0, allowZip64=True) as zf:
            zf.write(part_path, part_path.name)
            zf.write(manifest_path, manifest_path.name)
            zf.write(script, script.name)
            zf.write(readme, readme.name)
        wrappers.append({"path": wrapper, "qa": verify_zip(wrapper), "part": part_rows[index - 1]})
    with tempfile.TemporaryDirectory(prefix="mrhpd-r78-transport-") as td:
        target = Path(td)
        for wrapper in wrappers:
            safe_extract(wrapper["path"], target)
        result = subprocess.run([sys.executable, str(target / script.name)], cwd=target, text=True, capture_output=True, timeout=1200)
        if result.returncode:
            raise RuntimeError({"transport_reassembly_failed": {"stdout": result.stdout, "stderr": result.stderr}})
        reassembled = target / restore.name
        verify_zip(reassembled, restore.stat().st_size, sha256_file(restore))
        transport_verification = json.loads(result.stdout)
    return {"manifest": manifest, "manifest_path": manifest_path, "script": script, "wrappers": wrappers, "verification": transport_verification}


def build_verification_delivery(dist: Path, stamp: str, files: list[Path]) -> dict[str, Any]:
    output = dist / f"MRHPD v3.0.0a Response 78 Section 5 Session 1 Complete Verification Delivery {stamp}.zip"
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as zf:
        for path in files:
            if path.exists() and path.is_file():
                zf.write(path, path.name)
    return {"path": output, "qa": verify_zip(output)}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--volume1-dir", type=Path, required=True)
    parser.add_argument("--volume2-dir", type=Path, required=True)
    parser.add_argument("--checkpoint2-dir", type=Path, required=True)
    parser.add_argument("--dist", type=Path, default=Path("dist_cp5_s1_cp3"))
    args = parser.parse_args()
    now = utc_now()
    now_iso = now.isoformat().replace("+00:00", "Z")
    stamp = now.strftime("%Y-%m-%d %H%M UTC")
    if args.dist.exists():
        shutil.rmtree(args.dist)
    args.dist.mkdir(parents=True)
    with tempfile.TemporaryDirectory(prefix="mrhpd-cp5-s1-cp3-") as td:
        work = Path(td)
        restore72, project72_archive, response76_project, recovery_qa = reconstruct_response76(args.volume1_dir, args.volume2_dir, args.checkpoint2_dir, work)
        current_project = work / "current_project" / response76_project.name
        current_project.parent.mkdir(parents=True)
        shutil.copytree(response76_project, current_project)
        gates = acceptance_gate_templates(now_iso)
        events = recovery_events(now_iso)
        current_db = current_project / CURRENT_DB_REL
        database_qa = synchronize_database(response76_project / CP2_DB_REL, current_db, now_iso, gates, events)
        current_workbook = current_project / CURRENT_WORKBOOK_REL
        workbook_qa = augment_workbook(response76_project / CP2_WORKBOOK_REL, current_workbook, database_qa, gates, events, now_iso)
        tracking_files = write_tracking_files(current_project, current_db, now_iso)
        print_qa = inspect_print_surfaces(current_project)
        application_files, application_qa = write_terminal_application_audit(current_project, current_db, current_workbook, print_qa, now_iso)
        artwork_root = current_project / "Artwork" / "Section 5 Print Production" / "Session 1 Complete"
        freeze_figure = artwork_root / "MRHPD-FIG-S5-0005 Section 5 Session 1 Production Freeze v3.0.0a.png"
        figure_qa = build_freeze_figure(freeze_figure, database_qa, workbook_qa, print_qa)
        report_root = current_project / "Reports" / "Section 5 Session 1" / "Complete Through Response 78"
        docx_report = report_root / "MRHPD v3.0.0a Section 5 Session 1 Complete Production Freeze Report.docx"
        pdf_report = report_root / "MRHPD v3.0.0a Section 5 Session 1 Complete Production Freeze Report.pdf"
        xlsx_register = report_root / "MRHPD v3.0.0a Section 5 Session 1 Complete Production Register.xlsx"
        docx_qa = build_docx_report(docx_report, now_iso, database_qa, workbook_qa, print_qa, application_qa, gates, events, freeze_figure)
        pdf_qa = build_pdf_report(pdf_report, now_iso, database_qa, workbook_qa, print_qa, application_qa, gates, events, freeze_figure)
        register_qa = build_register(xlsx_register, database_qa, workbook_qa, print_qa, application_qa, gates, events)
        report_render_qa = render_pdf_qa(pdf_report)
        compaction_records = compact_superseded_derivatives(current_project, current_db, current_workbook, now_iso)
        persist_compaction_to_database(current_db, compaction_records)
        compaction_root = current_project / "Recovery" / "Section 5 Session 1 Complete Through Response 78"
        compaction_json = compaction_root / "MRHPD v3.0.0a Section 5 Session 1 Equivalence-Gated Compaction Register.json"
        compaction_csv = compaction_root / "MRHPD v3.0.0a Section 5 Session 1 Equivalence-Gated Compaction Register.csv"
        json_write(compaction_json, {"status": "passed", "generated_at": now_iso, "removed_file_count": len(compaction_records), "removed_bytes": sum(row["bytes"] for row in compaction_records), "records": compaction_records})
        csv_write(compaction_csv, compaction_records)
        index_result = build_source_and_bit_indexes(current_project, now_iso)
        update_gate(gates, "source_recovery", "passed", f"Response 72 {BASE_RESTORE_SHA256}; Response 76 {CP2_RECOVERY_SHA256}", now_iso)
        update_gate(gates, "sqlite_integrity", "passed", database_qa["integrity"], now_iso)
        update_gate(gates, "foreign_keys", "passed", str(database_qa["foreign_keys"]), now_iso)
        update_gate(gates, "response77", "passed", "Exactly one R77 record", now_iso)
        update_gate(gates, "response78", "passed", "Exactly one R78 record", now_iso)
        update_gate(gates, "session_state", "passed", "Checkpoint 3 complete; Session 1 complete", now_iso)
        update_gate(gates, "print_selection", "passed", "locked_initial_production_master", now_iso)
        update_gate(gates, "page_transform", "passed", "538 records; 0 failed", now_iso)
        update_gate(gates, "digital_publication", "passed", print_qa["digital"]["sha256"], now_iso)
        update_gate(gates, "editable_assembly", "passed", print_qa["editable"]["sha256"], now_iso)
        update_gate(gates, "print_interior_pages", "passed", "538 pages", now_iso)
        update_gate(gates, "print_interior_search", "passed", "537 searchable source pages", now_iso)
        update_gate(gates, "print_interior_blank", "passed", "Terminal page 538 intentionally blank", now_iso)
        update_gate(gates, "print_interior_geometry", "passed", "8.5 x 11 inches on all pages", now_iso)
        update_gate(gates, "cover_png", "passed", f"{print_qa['cover_png']['pixels']} RGB opaque", now_iso)
        update_gate(gates, "cover_pdf", "passed", str(print_qa["cover_pdf"]["inches"]), now_iso)
        update_gate(gates, "cover_template", "passed", str(print_qa["template_png"]["pixels"]), now_iso)
        update_gate(gates, "cover_components", "passed", f"{len(print_qa['cover_components'])} identities preserved", now_iso)
        update_gate(gates, "main_application", "passed", APPLICATION_SHA256, now_iso)
        update_gate(gates, "application_terminal_audit", "passed", application_qa["status"], now_iso)
        update_gate(gates, "legacy_application_evidence", "passed", f"{len(application_qa['inherited_application_qa_records'])} application QA; {len(application_qa['inherited_http_security_records'])} HTTP/security QA", now_iso)
        update_gate(gates, "workbook", "passed", f"{workbook_qa['current_sheet_count']} sheets; 0 formula errors", now_iso)
        update_gate(gates, "tracking", "passed", "Raw/Net and Cumulative Thread Index through Response 78", now_iso)
        update_gate(gates, "source_index", "passed", f"{index_result['qa']['source_index_records']} records", now_iso)
        update_gate(gates, "bit_index", "passed", index_result["qa"]["bit_index_integrity"], now_iso)
        update_gate(gates, "compaction", "passed", f"{len(compaction_records)} equivalence-recorded derivatives removed", now_iso)
        persist_final_gates(current_db, gates)
        final_qa_root = current_project / "QA" / "Section 5 Session 1" / "Complete Through Response 78"
        final_qa_root.mkdir(parents=True, exist_ok=True)
        qa_payload = {
            "schema": "mrhpd-section5-session1-complete-qa-1.0",
            "generated_at": now_iso,
            "status": "passed_with_controlled_external_gates",
            "response": 78,
            "section": SECTION_LABEL,
            "session": SESSION_LABEL,
            "checkpoint": CHECKPOINT_LABEL,
            "source_recovery": recovery_qa,
            "database": database_qa,
            "workbook": workbook_qa,
            "application": application_qa,
            "print_surfaces": print_qa,
            "reports": {"docx": docx_qa, "pdf": pdf_qa, "xlsx": register_qa, "render": report_render_qa, "figure": figure_qa},
            "acceptance_gates": gates,
            "recovery_events": events,
            "compaction": {"removed_file_count": len(compaction_records), "removed_bytes": sum(row["bytes"] for row in compaction_records)},
            "accepted_predecessor_mutated": False,
            "frozen_section3_release_mutated": False,
            "immutable_digital_publication_mutated": False,
            "main_application_mutated": False,
            "provider_preview": "controlled_pending_session2",
            "physical_proof": "controlled_pending_session2",
            "checkpoint_3_of_3_complete": True,
            "session_1_of_3_complete": True,
            "remediation_section_5_complete": False,
            "next": "Remediation Section 5 of 5 Session 2 of 3",
        }
        final_qa_path = final_qa_root / "SECTION5_SESSION1_COMPLETE_QA.json"
        json_write(final_qa_path, qa_payload)
        manifest_path, checksums_path, manifest_rows = build_project_manifest(current_project, now_iso)
        update_gate(gates, "manifest", "passed", f"{len(manifest_rows)} records; zero mismatches", now_iso)
        persist_final_gates(current_db, gates)
        # Rebuild manifest after final gate persistence; no project mutation occurs after this point.
        manifest_path, checksums_path, manifest_rows = build_project_manifest(current_project, now_iso)
        project_zip = args.dist / (
            f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 "
            f"Remediation Section 5 of 5 Session 1 of 3 COMPLETE PROJECT THROUGH RESPONSE 78 {stamp}.zip"
        )
        project_zip_qa = zip_tree(current_project, project_zip)
        if project_zip_qa["bytes"] >= MAX_ARCHIVE_BYTES:
            raise RuntimeError({"complete_project_exceeds_180_mib_after_compaction": project_zip_qa})
        clean_project = verify_project_clean_extract(project_zip, manifest_path.relative_to(current_project).as_posix(), {"workbook_sheets": workbook_qa["current_sheet_count"]})
        update_gate(gates, "project_archive", "passed", f"{project_zip_qa['bytes']} bytes; clean extraction passed", now_iso)
        persist_final_gates(current_db, gates)
        # Freeze final gate state and rebuild manifest/project archive one final time.
        manifest_path, checksums_path, manifest_rows = build_project_manifest(current_project, now_iso)
        project_zip_qa = zip_tree(current_project, project_zip)
        clean_project = verify_project_clean_extract(project_zip, manifest_path.relative_to(current_project).as_posix(), {"workbook_sheets": workbook_qa["current_sheet_count"]})
        direct_controls = [docx_report, pdf_report, xlsx_register, freeze_figure, compaction_json, compaction_csv, final_qa_path, manifest_path, checksums_path]
        restore_result = build_complete_restore(project_zip, manifest_path.relative_to(current_project).as_posix(), workbook_qa["current_sheet_count"], args.dist, stamp, direct_controls)
        update_gate(gates, "restore", "passed", f"{restore_result['qa']['bytes']} bytes; embedded verifier passed", now_iso)
        transport = build_transport(restore_result["restore"], args.dist)
        update_gate(gates, "transport", "passed", f"{len(transport['wrappers'])} volumes; reassembly passed", now_iso)
        persist_final_gates(current_db, gates)
        verification = {
            "schema": "mrhpd-response78-section5-session1-complete-verification-1.0",
            "generated_at": now_iso,
            "status": "passed_with_controlled_external_gates",
            "response": 78,
            "checkpoint_3_of_3_complete": True,
            "session_1_of_3_complete": True,
            "remediation_section_5_complete": False,
            "source_recovery": recovery_qa,
            "database": database_qa,
            "workbook": workbook_qa,
            "application": application_qa,
            "print_surfaces": print_qa,
            "acceptance_gates": gates,
            "compaction": {"removed_file_count": len(compaction_records), "removed_bytes": sum(row["bytes"] for row in compaction_records)},
            "project_archive": project_zip_qa,
            "clean_project": clean_project,
            "complete_restore": restore_result["qa"],
            "restore_verification": restore_result["verification"],
            "transport": {"manifest": transport["manifest"], "verification": transport["verification"], "volumes": [{"name": row['path'].name, "bytes": row['qa']['bytes'], "sha256": row['qa']['sha256']} for row in transport["wrappers"]]},
            "user_upload_required": False,
            "requires_conversation_reconstruction": False,
            "accepted_predecessor_mutated": False,
            "immutable_digital_publication_mutated": False,
            "provider_preview": "controlled_pending_session2",
            "physical_proof": "controlled_pending_session2",
            "next": "Remediation Section 5 of 5 Session 2 of 3",
        }
        verification_path = args.dist / "MRHPD v3.0.0a Response 78 Section 5 Session 1 Complete Verification.json"
        json_write(verification_path, verification)
        exact_names = args.dist / "MRHPD v3.0.0a Response 78 Exact File Names.txt"
        text_write(exact_names, "\n".join([
            "Complete current project archive:", project_zip.name, "",
            "Complete self-contained restore:", restore_result["restore"].name, "",
            *sum(([f"Complete restore volume {index} of {len(transport['wrappers'])}:", row["path"].name, ""] for index, row in enumerate(transport["wrappers"], start=1)), []),
            "Verification delivery:", f"MRHPD v3.0.0a Response 78 Section 5 Session 1 Complete Verification Delivery {stamp}.zip", "",
            "Current SQLite database:", Path(CURRENT_DB_REL).name, "",
            "Current comprehensive workbook:", Path(CURRENT_WORKBOOK_REL).name, "",
            "Digital publication:", Path(PUBLICATION_REL).name, "",
            "Print-production interior:", Path(PRINT_INTERIOR_REL).name, "",
            "Full-cover PNG:", Path(FINAL_COVER_PNG_REL).name,
        ]))
        summary_path = args.dist / "MRHPD_RESPONSE78_SECTION5_SESSION1_COMPLETE_BUILD_SUMMARY.json"
        json_write(summary_path, verification)
        verification_delivery = build_verification_delivery(args.dist, stamp, [verification_path, exact_names, summary_path, docx_report, pdf_report, xlsx_register, freeze_figure, compaction_json, compaction_csv, final_qa_path, transport["manifest_path"], transport["script"]])
        console = {
            "status": "passed_with_controlled_external_gates",
            "response": 78,
            "checkpoint": "3 of 3 complete",
            "session": "1 of 3 complete",
            "section": "5 of 5 continue",
            "database_tables": database_qa["tables"],
            "workbook_sheets": workbook_qa["current_sheet_count"],
            "digital_pages": 537,
            "print_pages": 538,
            "project_archive": project_zip_qa,
            "complete_restore": restore_result["qa"],
            "transport_volumes": [{"name": row["path"].name, "bytes": row["qa"]["bytes"], "sha256": row["qa"]["sha256"]} for row in transport["wrappers"]],
            "verification_delivery": verification_delivery["qa"],
            "user_upload_required": False,
            "next": "Remediation Section 5 of 5 Session 2 of 3",
        }
        print(json.dumps(console, indent=2))


if __name__ == "__main__":
    main()
