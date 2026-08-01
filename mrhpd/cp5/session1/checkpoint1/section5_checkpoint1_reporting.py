#!/usr/bin/env python3
from __future__ import annotations

import csv
import json
import math
import textwrap
import zipfile
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

NAVY = "17324D"
TEAL = "1C7475"
GOLD = "C9A227"
PALE_BLUE = "EAF1F5"
PALE_GOLD = "F7F1D9"
PALE_RED = "F7E8E6"
DARK = "24323D"
WHITE = "FFFFFF"
GRAY = "66757F"


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False), encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fieldnames: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fieldnames is None:
        fieldnames = list(rows[0]) if rows else []
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow(row)


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            pass
    return ImageFont.load_default()


def _draw_wrapped(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font: Any, fill: str, width_chars: int, spacing: int = 8) -> int:
    x, y = xy
    lines: list[str] = []
    for paragraph in text.splitlines() or [""]:
        lines.extend(textwrap.wrap(paragraph, width=width_chars) or [""])
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        bbox = draw.textbbox((x, y), line or "Ag", font=font)
        y += bbox[3] - bbox[1] + spacing
    return y


def build_spine_scenario_figure(path: Path, scenarios: list[dict[str, Any]]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 2400, 1350
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = _font(60, True)
    subtitle_font = _font(31, False)
    label_font = _font(34, True)
    body_font = _font(29, False)
    small_font = _font(24, False)
    draw.rectangle((0, 0, width, 170), fill=f"#{NAVY}")
    draw.text((110, 45), "KDP spine scenarios for the 537-page interior", font=title_font, fill="white")
    draw.text((112, 190), "KDP rounds the production count to 538 pages. Widths below are calculated scenarios, not a selected printer template.", font=subtitle_font, fill=f"#{DARK}")
    numeric = [row for row in scenarios if isinstance(row.get("spine_width_in"), (int, float))]
    max_spine = max((float(row["spine_width_in"]) for row in numeric), default=1.0)
    y = 315
    bar_left = 700
    bar_max_width = 1320
    palette = ["#1C7475", "#C9A227", "#536D8C", "#A65F46"]
    for index, row in enumerate(numeric):
        color = palette[index % len(palette)]
        name = str(row["scenario_name"])
        spine = float(row["spine_width_in"])
        cover = float(row["cover_width_in"])
        pixels = int(row["cover_width_px_300dpi"])
        draw.text((115, y + 12), name, font=label_font, fill=f"#{NAVY}")
        bar_width = max(20, int(spine / max_spine * bar_max_width))
        draw.rounded_rectangle((bar_left, y, bar_left + bar_width, y + 72), radius=20, fill=color)
        draw.text((bar_left + 24, y + 17), f"{spine:.6f} in", font=body_font, fill="white")
        draw.text((bar_left, y + 88), f"Full wrap: {cover:.6f} × 11.250 in  |  minimum raster: {pixels} × 3375 px at 300 ppi", font=small_font, fill=f"#{DARK}")
        y += 205
    draw.rounded_rectangle((110, height - 235, width - 110, height - 75), radius=24, fill=f"#{PALE_GOLD}", outline=f"#{GOLD}", width=4)
    _draw_wrapped(
        draw,
        (145, height - 208),
        "The existing 18.000 × 11.250-inch wrap with a 0.750-inch spine is retained only as a provisional legacy reference. It is too narrow for every current 538-page KDP paper scenario and must not be stretched into compliance.",
        _font(30, True),
        f"#{DARK}",
        125,
        spacing=7,
    )
    image.save(path, format="PNG", dpi=(300, 300), optimize=True)
    return {"path": str(path), "pixels": [width, height], "dpi": 300, "mode": image.mode, "bytes": path.stat().st_size}


def build_cover_width_figure(path: Path, current_width_in: float, scenarios: list[dict[str, Any]]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 2400, 1250
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, 165), fill=f"#{NAVY}")
    draw.text((110, 44), "Existing cover wrap versus calculated production widths", font=_font(58, True), fill="white")
    draw.text((110, 195), "Visual comparison at a common scale. Final dimensions remain provider- and material-specific.", font=_font(31), fill=f"#{DARK}")
    values = [("Existing provisional wrap", current_width_in, "#A65F46")]
    for idx, row in enumerate([r for r in scenarios if isinstance(r.get("cover_width_in"), (int, float))]):
        values.append((str(row["scenario_name"]), float(row["cover_width_in"]), ["#1C7475", "#C9A227", "#536D8C", "#6B8E4E"][idx % 4]))
    max_width = max(v for _, v, _ in values)
    y = 320
    left, available = 570, 1650
    for name, value, color in values:
        draw.text((105, y + 18), name, font=_font(31, True), fill=f"#{NAVY}")
        bw = int(value / max_width * available)
        draw.rounded_rectangle((left, y, left + bw, y + 78), radius=18, fill=color)
        draw.text((left + 24, y + 19), f"{value:.6f} in", font=_font(30, True), fill="white")
        y += 150
    minimum = min(v for name, v, _ in values if name != "Existing provisional wrap")
    shortage = minimum - current_width_in
    draw.rounded_rectangle((105, height - 175, width - 105, height - 55), radius=22, fill=f"#{PALE_RED}", outline="#A65F46", width=4)
    draw.text((140, height - 138), f"Minimum calculated shortfall: {shortage:.6f} inches. Regenerate from separate front, back, and spine assets after the provider scenario is selected.", font=_font(29, True), fill=f"#{DARK}")
    image.save(path, format="PNG", dpi=(300, 300), optimize=True)
    return {"path": str(path), "pixels": [width, height], "dpi": 300, "mode": image.mode, "bytes": path.stat().st_size}


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_docx_table(document: Document, headers: list[str], rows: Iterable[Iterable[Any]], widths: list[float] | None = None) -> Any:
    materialized = [list(row) for row in rows]
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(header)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for r_idx, row in enumerate(materialized, start=1):
        cells = table.add_row().cells
        for c_idx, value in enumerate(row):
            cells[c_idx].text = "" if value is None else str(value)
            set_cell_margins(cells[c_idx])
            if r_idx % 2 == 0:
                set_cell_shading(cells[c_idx], PALE_BLUE)
            for paragraph in cells[c_idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(7.4)
            cells[c_idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths and c_idx < len(widths):
                cells[c_idx].width = Inches(widths[c_idx])
    return table


def configure_docx(document: Document, title: str, subtitle: str, generated_at: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    for style_name, size, color, before, after in (
        ("Title", 24, NAVY, 0, 8),
        ("Subtitle", 11, TEAL, 0, 10),
        ("Heading 1", 16, NAVY, 12, 5),
        ("Heading 2", 12.5, TEAL, 9, 4),
        ("Heading 3", 10.5, GOLD, 7, 3),
    ):
        style = document.styles[style_name]
        style.font.name = "Aptos Display" if style_name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    document.core_properties.title = title
    document.core_properties.subject = subtitle
    document.core_properties.author = "Brent McAnulty, M.D."
    document.core_properties.keywords = "Human Pathogen Database; print production; spine; preflight; recovery"
    p = document.add_paragraph(title, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = document.add_paragraph(subtitle, style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"Version 3.0.0a  |  Response 75  |  Generated {generated_at}")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (label, value) in enumerate((
        ("SECTION", "5 OF 5"),
        ("SESSION", "1 OF 3"),
        ("CHECKPOINT", "1 OF 3"),
    )):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, [NAVY, TEAL, GOLD][idx])
        cell.text = f"{label}\n{value}"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
        set_cell_margins(cell, top=120, bottom=120)


def build_docx_report(
    path: Path,
    *,
    generated_at: str,
    baseline: dict[str, Any],
    provider_specs: list[dict[str, Any]],
    scenarios: list[dict[str, Any]],
    interior: dict[str, Any],
    cover: dict[str, Any],
    risks: list[dict[str, Any]],
    recovery_events: list[dict[str, Any]],
    tracking: list[dict[str, Any]],
    figure_paths: list[Path],
) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    title = "Human Pathogen Database — Section 5 Print-Production Intake"
    subtitle = "Provider requirements, even-page production count, spine scenarios, interior screening, provisional-cover disposition, and deterministic recovery"
    configure_docx(doc, title, subtitle, generated_at)

    doc.add_heading("Checkpoint disposition", level=1)
    p = doc.add_paragraph()
    p.add_run("Checkpoint 1 of 3 is complete. ").bold = True
    p.add_run("The immutable Response 72 restore was verified, a copied Section 5 working state was created, and print-production controls were synchronized. No press-ready cover or print interior is declared at this checkpoint.")
    add_docx_table(doc, ["Control", "Result"], [
        ("Response 72 restore", f"{baseline['restore_bytes']:,} bytes; SHA-256 {baseline['restore_sha256']}"),
        ("Response 72 project archive", f"{baseline['project_bytes']:,} bytes; SHA-256 {baseline['project_sha256']}"),
        ("Current digital interior", "537 searchable pages; immutable source retained"),
        ("Production page count", "538 pages after even-page normalization; print-only derivative pending"),
        ("Existing cover wrap", "Retained as provisional legacy reference; not press-ready"),
        ("User upload required", "No"),
    ])

    doc.add_heading("Production boundary", level=1)
    doc.add_paragraph(
        "Section 5 separates the immutable 537-page digital publication from provider-specific print derivatives. The digital PDF and editable assembly remain authoritative clinical-content artifacts. A print interior may add one controlled blank page and may require margin reflow, but it must not silently change clinical content, locators, evidence, or accessibility."
    )

    doc.add_heading("Official provider requirements", level=1)
    doc.add_paragraph(
        "The provider register stores each requirement with an official source URL, verification date, evidence boundary, and implementation status. KDP arithmetic is deterministic once paper/color is selected. IngramSpark dimensions remain template-controlled until trim, binding, paper, and page count are submitted to its cover-template generator."
    )
    add_docx_table(
        doc,
        ["Provider", "Requirement", "Current value", "Status", "Official source"],
        [
            (row["provider"], row["requirement"], row["value"], row["status"], row["source_url"])
            for row in provider_specs
        ],
        widths=[0.65, 1.75, 1.55, 0.9, 2.35],
    )

    doc.add_heading("Spine and wrap scenarios", level=1)
    doc.add_paragraph(
        "The current 18.000 × 11.250-inch wrap assumes a 0.750-inch spine. Every calculated 538-page KDP scenario requires a materially wider spine and cover. The correct operation is regeneration from separate front, back, and spine components after selecting the provider scenario—not horizontal stretching of the existing wrap."
    )
    add_docx_table(
        doc,
        ["Scenario", "Pages", "Spine in", "Wrap W × H in", "300-ppi pixels", "Disposition"],
        [
            (
                row["scenario_name"], row["production_page_count"],
                "template required" if row.get("spine_width_in") is None else f"{row['spine_width_in']:.6f}",
                "template required" if row.get("cover_width_in") is None else f"{row['cover_width_in']:.6f} × {row['cover_height_in']:.3f}",
                "template required" if row.get("cover_width_px_300dpi") is None else f"{row['cover_width_px_300dpi']} × {row['cover_height_px_300dpi']}",
                row["status"],
            )
            for row in scenarios
        ],
        widths=[1.45, 0.48, 0.8, 1.1, 1.1, 1.55],
    )
    for figure in figure_paths:
        doc.add_paragraph()
        doc.add_picture(str(figure), width=Inches(7.05))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Interior preflight intake", level=1)
    add_docx_table(doc, ["Interior control", "Observed state", "Disposition"], [
        ("Digital page count", interior["page_count"], "Preserve immutable digital edition"),
        ("KDP/Ingram even production count", interior["production_page_count"], "Create controlled print-only blank page later"),
        ("Searchable pages", interior["searchable_pages"], "Passed"),
        ("Nominal trim", interior["nominal_trim"], "8.5 × 11 in target retained"),
        ("Portrait/landscape page boxes", f"{interior['portrait_pages']} / {interior['landscape_pages']}", "Review landscape content at provider preview"),
        ("Pages below automated 0.75-in inside text-block screen", interior["inside_margin_screen_failures"], "Automated screening only; margin reflow decision required"),
        ("Pages below automated outside-margin screen", interior["outside_margin_screen_failures"], "Automated screening only; inspect provider preview"),
        ("Publication SHA-256", interior["sha256"], "Immutable"),
    ])
    doc.add_paragraph(
        "The text-block margin screen is conservative and not a substitute for a provider previewer or manual page review. It identifies pages requiring attention; it does not alter the publication or certify final print compliance."
    )

    doc.add_heading("Current cover asset audit", level=1)
    add_docx_table(doc, ["Asset", "Observed", "Status"], [
        (row["relative_path"], row["observed"], row["status"])
        for row in cover["assets"]
    ])
    doc.add_paragraph(
        "The separate front, back, and spine assets are retained as reusable design inputs. The combined wrap is a provisional legacy proof. Final cover production remains blocked on provider selection, paper/color choice, template acquisition, final print-interior page count, barcode policy, and proof approval."
    )

    doc.add_heading("Controlled forward-work register", level=1)
    add_docx_table(doc, ["Risk/control", "Current state", "Required closure"], [
        (row["risk"], row["status"], row["closure_requirement"]) for row in risks
    ])

    doc.add_heading("Tracking and recovery", level=1)
    add_docx_table(doc, ["Record", "Disposition"], [(row["record"], row["disposition"]) for row in tracking])
    add_docx_table(doc, ["Event", "Condition", "Recovery"], [
        (row["event_code"], row["condition"], row["recovery"]) for row in recovery_events
    ])

    doc.add_heading("Next checkpoint", level=1)
    doc.add_paragraph(
        "Checkpoint 2 will select and lock the first production scenario, generate a controlled 538-page print-interior derivative, perform page-by-page margin and render review, acquire or generate the exact provider cover template, regenerate the wrap from separate components, and rerun database, workbook, application, index, manifest, checksum, and recovery controls."
    )

    doc.save(path)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("DOCX CRC failed")
    return {"path": str(path), "bytes": path.stat().st_size, "title": title}


def _rl_paragraph(text: Any, style: ParagraphStyle) -> Paragraph:
    safe = str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
    return Paragraph(safe, style)


def build_pdf_report(
    path: Path,
    *,
    generated_at: str,
    baseline: dict[str, Any],
    provider_specs: list[dict[str, Any]],
    scenarios: list[dict[str, Any]],
    interior: dict[str, Any],
    cover: dict[str, Any],
    risks: list[dict[str, Any]],
    recovery_events: list[dict[str, Any]],
    figure_paths: list[Path],
) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    title = ParagraphStyle("TitleMR", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=21, leading=24, textColor=colors.HexColor(f"#{NAVY}"), alignment=TA_LEFT, spaceAfter=7)
    subtitle = ParagraphStyle("SubMR", parent=styles["Normal"], fontName="Helvetica", fontSize=9.5, leading=12, textColor=colors.HexColor(f"#{TEAL}"), spaceAfter=9)
    h1 = ParagraphStyle("H1MR", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14, leading=17, textColor=colors.HexColor(f"#{NAVY}"), spaceBefore=8, spaceAfter=5)
    body = ParagraphStyle("BodyMR", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.2, leading=10.5, textColor=colors.HexColor(f"#{DARK}"), spaceAfter=5)
    small = ParagraphStyle("SmallMR", parent=body, fontSize=6.5, leading=8)
    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.55 * inch, leftMargin=0.55 * inch, topMargin=0.55 * inch, bottomMargin=0.55 * inch, title="Human Pathogen Database — Section 5 Print-Production Intake", author="Brent McAnulty, M.D.")
    story: list[Any] = []
    story.append(_rl_paragraph("Human Pathogen Database — Section 5 Print-Production Intake", title))
    story.append(_rl_paragraph("Provider requirements, spine scenarios, interior screening, provisional-cover disposition, and deterministic recovery", subtitle))
    status_table = Table([
        ["SECTION", "SESSION", "CHECKPOINT"],
        ["5 OF 5", "1 OF 3", "1 OF 3"],
    ], colWidths=[2.35 * inch] * 3)
    status_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.white),
        ("BACKGROUND", (0, 1), (0, 1), colors.HexColor(f"#{NAVY}")),
        ("BACKGROUND", (1, 1), (1, 1), colors.HexColor(f"#{TEAL}")),
        ("BACKGROUND", (2, 1), (2, 1), colors.HexColor(f"#{GOLD}")),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.extend([status_table, Spacer(1, 0.12 * inch)])
    story.append(_rl_paragraph(f"Version 3.0.0a | Response 75 | Generated {generated_at}", small))

    story.append(_rl_paragraph("Checkpoint disposition", h1))
    story.append(_rl_paragraph("Checkpoint 1 of 3 is complete. The exact Response 72 restore was verified and the first copied Section 5 print-production state was synchronized. No press-ready cover or print interior is declared at this checkpoint.", body))
    base_data = [
        ["Control", "Result"],
        ["Response 72 restore", f"{baseline['restore_bytes']:,} bytes; {baseline['restore_sha256']}"],
        ["Response 72 project", f"{baseline['project_bytes']:,} bytes; {baseline['project_sha256']}"],
        ["Digital interior", "537 searchable pages; immutable"],
        ["Production page count", "538 pages; print-only derivative pending"],
        ["User upload required", "No"],
    ]
    story.append(_pdf_table(base_data, [1.55 * inch, 5.45 * inch], small))

    story.append(_rl_paragraph("Official provider requirements", h1))
    provider_data = [["Provider", "Requirement", "Value", "Status"]] + [[row["provider"], row["requirement"], row["value"], row["status"]] for row in provider_specs]
    story.append(_pdf_table(provider_data, [0.7 * inch, 2.5 * inch, 2.25 * inch, 1.55 * inch], small))

    story.append(_rl_paragraph("Spine and wrap scenarios", h1))
    scenario_data = [["Scenario", "Pages", "Spine", "Wrap", "Disposition"]]
    for row in scenarios:
        scenario_data.append([
            row["scenario_name"], row["production_page_count"],
            "template" if row.get("spine_width_in") is None else f"{row['spine_width_in']:.6f} in",
            "template" if row.get("cover_width_in") is None else f"{row['cover_width_in']:.6f} × {row['cover_height_in']:.3f} in",
            row["status"],
        ])
    story.append(_pdf_table(scenario_data, [1.6 * inch, 0.45 * inch, 0.8 * inch, 1.45 * inch, 2.7 * inch], small))
    for figure in figure_paths:
        story.extend([Spacer(1, 0.08 * inch), RLImage(str(figure), width=7.0 * inch, height=7.0 * inch * 1350 / 2400)])

    story.append(PageBreak())
    story.append(_rl_paragraph("Interior preflight intake", h1))
    interior_data = [
        ["Control", "Observed", "Disposition"],
        ["Digital pages", interior["page_count"], "Immutable"],
        ["Production pages", interior["production_page_count"], "Controlled print derivative pending"],
        ["Searchable pages", interior["searchable_pages"], "Passed"],
        ["Nominal trim", interior["nominal_trim"], "8.5 × 11 target"],
        ["Portrait / landscape", f"{interior['portrait_pages']} / {interior['landscape_pages']}", "Preview all pages"],
        ["Inside-margin screen", interior["inside_margin_screen_failures"], "Manual/provider review required"],
        ["Outside-margin screen", interior["outside_margin_screen_failures"], "Manual/provider review required"],
    ]
    story.append(_pdf_table(interior_data, [2.0 * inch, 1.5 * inch, 3.5 * inch], small))
    story.append(_rl_paragraph("The automated text-block margin screen is a conservative triage tool, not a final printer certification. Provider preview and page-level visual review remain required.", body))

    story.append(_rl_paragraph("Current cover asset audit", h1))
    cover_data = [["Asset", "Observed", "Status"]] + [[row["relative_path"], row["observed"], row["status"]] for row in cover["assets"]]
    story.append(_pdf_table(cover_data, [3.6 * inch, 1.6 * inch, 1.8 * inch], small))

    story.append(_rl_paragraph("Controlled forward-work register", h1))
    risk_data = [["Risk/control", "State", "Closure requirement"]] + [[row["risk"], row["status"], row["closure_requirement"]] for row in risks]
    story.append(_pdf_table(risk_data, [2.05 * inch, 1.05 * inch, 3.9 * inch], small))

    story.append(_rl_paragraph("Recovery events", h1))
    event_data = [["Event", "Condition", "Recovery"]] + [[row["event_code"], row["condition"], row["recovery"]] for row in recovery_events]
    story.append(_pdf_table(event_data, [1.1 * inch, 2.55 * inch, 3.35 * inch], small))

    story.append(_rl_paragraph("Next checkpoint", h1))
    story.append(_rl_paragraph("Checkpoint 2 will lock the first production scenario, generate the controlled 538-page print interior, obtain the exact provider template, regenerate the cover wrap from separate components, and repeat all document, database, workbook, application, index, manifest, checksum, and recovery gates.", body))
    doc.build(story)
    reader = PdfReader(str(path))
    text_chars = sum(len(page.extract_text() or "") for page in reader.pages)
    if len(reader.pages) < 3 or text_chars < 2500:
        raise RuntimeError({"pdf_report_validation": {"pages": len(reader.pages), "text_chars": text_chars}})
    return {"path": str(path), "bytes": path.stat().st_size, "pages": len(reader.pages), "text_chars": text_chars}


def _pdf_table(data: list[list[Any]], widths: list[float], style: ParagraphStyle) -> Table:
    converted = [[_rl_paragraph(value, style) for value in row] for row in data]
    table = Table(converted, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#82909A")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    for idx in range(1, len(data)):
        if idx % 2 == 0:
            table.setStyle(TableStyle([("BACKGROUND", (0, idx), (-1, idx), colors.HexColor(f"#{PALE_BLUE}"))]))
    return table


def _style_sheet(ws: Any, widths: dict[int, float] | None = None, freeze: str = "A2", autofilter: bool = True) -> None:
    header_fill = PatternFill("solid", fgColor=NAVY)
    header_font = Font(name="Aptos", bold=True, color=WHITE, size=10)
    normal_font = Font(name="Aptos", size=9)
    thin = Side(style="thin", color="AAB4BA")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    ws.row_dimensions[1].height = 32
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.font = normal_font
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if cell.row % 2 == 0:
                cell.fill = PatternFill("solid", fgColor=PALE_BLUE)
    ws.freeze_panes = freeze
    if autofilter and ws.max_column and ws.max_row:
        ws.auto_filter.ref = ws.dimensions
    if widths:
        for column, width in widths.items():
            ws.column_dimensions[get_column_letter(column)].width = width
    else:
        for column in range(1, ws.max_column + 1):
            max_len = 0
            for cell in ws.iter_cols(min_col=column, max_col=column, min_row=1, max_row=min(ws.max_row, 250)):
                for item in cell:
                    max_len = max(max_len, len(str(item.value or "")))
            ws.column_dimensions[get_column_letter(column)].width = min(max(max_len + 2, 10), 55)


def build_checkpoint_register(
    path: Path,
    *,
    summary_rows: list[dict[str, Any]],
    provider_specs: list[dict[str, Any]],
    scenarios: list[dict[str, Any]],
    interior_page_rows: list[dict[str, Any]],
    cover_assets: list[dict[str, Any]],
    risks: list[dict[str, Any]],
    tracking: list[dict[str, Any]],
    recovery_events: list[dict[str, Any]],
) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    wb.remove(wb.active)
    datasets: list[tuple[str, list[dict[str, Any]], dict[int, float] | None]] = [
        ("Summary", summary_rows, {1: 30, 2: 42, 3: 24}),
        ("Provider Specs", provider_specs, {1: 14, 2: 31, 3: 25, 4: 17, 5: 45, 6: 22}),
        ("Spine Scenarios", scenarios, {1: 28, 2: 12, 3: 13, 4: 14, 5: 14, 6: 14, 7: 18, 8: 24}),
        ("Interior Pages", interior_page_rows, {1: 10, 2: 13, 3: 13, 4: 12, 5: 12, 6: 13, 7: 13, 8: 13, 9: 15, 10: 18}),
        ("Cover Assets", cover_assets, {1: 65, 2: 15, 3: 20, 4: 20, 5: 35}),
        ("Risks", risks, {1: 35, 2: 22, 3: 70}),
        ("Tracking", tracking, {1: 35, 2: 80}),
        ("Recovery", recovery_events, {1: 28, 2: 50, 3: 70}),
    ]
    for sheet_name, rows, widths in datasets:
        ws = wb.create_sheet(sheet_name)
        headers = list(rows[0]) if rows else ["No records"]
        ws.append(headers)
        for row in rows:
            ws.append([row.get(header) for header in headers])
        _style_sheet(ws, widths=widths)
    wb.properties.title = "MRHPD Section 5 Session 1 Checkpoint 1 Register"
    wb.properties.subject = "Print-production intake, spine scenarios, preflight, tracking, and recovery"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(path)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("XLSX CRC failed")
    readback = load_workbook(path, read_only=True, data_only=False)
    try:
        sheets = list(readback.sheetnames)
        formula_errors = 0
        error_tokens = ("#REF!", "#DIV/0!", "#NAME?", "#VALUE!", "#N/A")
        for ws in readback.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str) and any(token in cell.value for token in error_tokens):
                        formula_errors += 1
    finally:
        readback.close()
    if formula_errors:
        raise RuntimeError({"workbook_formula_error_tokens": formula_errors})
    return {"path": str(path), "bytes": path.stat().st_size, "sheets": sheets, "formula_error_tokens": formula_errors}


def augment_comprehensive_workbook(
    source: Path,
    destination: Path,
    *,
    generated_at: str,
    provider_specs: list[dict[str, Any]],
    scenarios: list[dict[str, Any]],
    risks: list[dict[str, Any]],
    tracking_rows: list[dict[str, Any]],
    recovery_events: list[dict[str, Any]],
    interior_summary: dict[str, Any],
    cover_summary: dict[str, Any],
) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    wb = load_workbook(source)
    inherited = list(wb.sheetnames)
    new_names = [
        "S5S1 Dashboard",
        "S5S1 Provider Specs",
        "S5S1 Spine Scenarios",
        "S5S1 Interior QA",
        "S5S1 Cover QA",
        "S5S1 Risks",
        "S5S1 Responses",
        "S5S1 Recovery",
    ]
    for name in new_names:
        if name in wb.sheetnames:
            del wb[name]
    ws = wb.create_sheet("S5S1 Dashboard", 0)
    ws.append(["Control", "Current state", "Disposition"])
    dashboard_rows = [
        ("Section", "Remediation Section 5 of 5", "CONTINUE"),
        ("Session", "Session 1 of 3", "CONTINUE"),
        ("Checkpoint", "Checkpoint 1 of 3", "COMPLETE"),
        ("Current response", 75, "Current"),
        ("Digital interior", f"{interior_summary['page_count']} pages", "Immutable"),
        ("Print production count", interior_summary["production_page_count"], "Print-only derivative pending"),
        ("Current combined cover", cover_summary["combined_cover_status"], "Provisional legacy reference"),
        ("Database", "Copied Section 5 database", "Integrity and foreign keys passed"),
        ("Generated", generated_at, "UTC"),
    ]
    for row in dashboard_rows:
        ws.append(row)
    _style_sheet(ws, widths={1: 28, 2: 45, 3: 45})

    sheet_rows = {
        "S5S1 Provider Specs": provider_specs,
        "S5S1 Spine Scenarios": scenarios,
        "S5S1 Interior QA": [interior_summary],
        "S5S1 Cover QA": cover_summary["assets"],
        "S5S1 Risks": risks,
        "S5S1 Responses": tracking_rows,
        "S5S1 Recovery": recovery_events,
    }
    for name, rows in sheet_rows.items():
        ws = wb.create_sheet(name)
        headers = list(rows[0]) if rows else ["No records"]
        ws.append(headers)
        for row in rows:
            ws.append([row.get(header) for header in headers])
        _style_sheet(ws)

    wb.properties.title = "MRHPD v3.0.0a Comprehensive Tracking Through Response 75"
    wb.properties.subject = "Section 5 print-production intake and cumulative tracking"
    wb.save(destination)
    with zipfile.ZipFile(destination) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("Comprehensive workbook CRC failed")
    check = load_workbook(destination, read_only=True, data_only=False)
    try:
        current_sheets = list(check.sheetnames)
        missing = [name for name in inherited if name not in current_sheets]
        formula_errors = 0
        for ws in check.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str) and any(token in cell.value for token in ("#REF!", "#DIV/0!", "#NAME?", "#VALUE!", "#N/A")):
                        formula_errors += 1
    finally:
        check.close()
    if missing or formula_errors:
        raise RuntimeError({"workbook_validation": {"missing_inherited": missing, "formula_errors": formula_errors}})
    return {
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "inherited_sheet_count": len(inherited),
        "current_sheet_count": len(current_sheets),
        "new_sheets": new_names,
        "missing_inherited_sheets": missing,
        "formula_error_tokens": formula_errors,
    }
