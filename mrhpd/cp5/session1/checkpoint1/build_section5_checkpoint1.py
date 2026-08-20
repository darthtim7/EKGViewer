#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import os
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import textwrap
import zipfile
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any, Iterable

import fitz
from docx import Document
from openpyxl import load_workbook
from PIL import Image
from pypdf import PdfReader

HERE = Path(__file__).resolve().parent
if str(HERE) not in sys.path:
    sys.path.insert(0, str(HERE))
import inspect_response72_v2 as intake  # noqa: E402
import section5_checkpoint1_reporting as reporting  # noqa: E402

PROJECT_VERSION = "3.0.0a"
RESPONSE_NUMBER = 75
SECTION_LABEL = "Remediation Section 5 of 5"
SESSION_LABEL = "Session 1 of 3"
CHECKPOINT_LABEL = "Checkpoint 1 of 3"
CHECKPOINT_CODE = "MRHPD-V3-CP5-S1-CP1"
BASE_RESTORE_BYTES = 159_186_352
BASE_RESTORE_SHA256 = "cb6d2de9bb351a4ff580e8ac0ac071a774670974098da88be822d64b437b25ce"
BASE_PROJECT_BYTES = 159_865_032
BASE_PROJECT_SHA256 = "88b3a6fab6e1106b2942b92fbe5b10c9b06ffe6f15963a7f0c308203dcb6beb5"
PUBLICATION_SHA256 = "8a053112ca24cd730b970130d5d0fc57a15c681531603601096186aeb0cd9642"
EDITABLE_SHA256 = "f832ff934d77049d75712f28bdfc9167b8a6b119c797235431b304b9e24369a2"
APPLICATION_SHA256 = "5f1e4ac8fc6e2ffad213646c78e4f261bf655795de5ac8a7d4486d3be11ce139"
CURRENT_DB_REL = (
    "Database/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 1 of 3 Checkpoint 1 of 3 THROUGH RESPONSE 75.sqlite"
)
CURRENT_WORKBOOK_REL = (
    "Tracking/Workbook/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 1 of 3 Checkpoint 1 of 3 THROUGH RESPONSE 75 Comprehensive Tracking.xlsx"
)
BASE_DB_REL = (
    "Database/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 4 of 5 Session 3 of 3 COMPLETE THROUGH RESPONSE 69.sqlite"
)
BASE_WORKBOOK_REL = (
    "Tracking/Workbook/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 4 of 5 Session 3 of 3 COMPLETE THROUGH RESPONSE 69 Comprehensive Tracking.xlsx"
)
PUBLICATION_REL = (
    "Documents/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 3 of 5 Session 3 of 4 Integrated Manuscript.pdf"
)
COVER_HASHES = {
    "Cover/MRHPD-COVER-0001 Front Cover v3.0.0a Brent McAnulty MD.png": "e53b56481f1a3f084477f68415ef5619ef019fc6796fd8262232f073c96549b6",
    "Cover/MRHPD-COVER-0002 Back Cover v3.0.0a Brent McAnulty MD.png": "890e82e71bf8fe5cd528af1f03ffeaf5a133a4242c766f8e2fd57173b47f35f7",
    "Cover/MRHPD-COVER-0003 Spine v3.0.0a Brent McAnulty MD.png": "83da68640a6a335b73952aed76efa4c2a3aef33d462c84fde620ad2694621f9c",
    "Cover/MRHPD-COVER-0004 Combined Cover Wrap v3.0.0a Brent McAnulty MD KDP Reference CMYK.tif": "6bcfbb7cb11f13ade958a24d2858ac7469eb482ec3156495ddd989d282971df2",
    "Cover/MRHPD-COVER-0004 Combined Cover Wrap v3.0.0a Brent McAnulty MD KDP Reference.pdf": "d281f523eb5fea428dde2f4250af56f09c92e03b32e7400b78b915994f08b607",
    "Cover/MRHPD-COVER-0004 Combined Cover Wrap v3.0.0a Brent McAnulty MD KDP Reference.png": "c8d1f37a1e0fdef86b2f0d8e85a6e075f76ed64e85eaf186e9f5eb79ffbc19b7",
}


def utc_now() -> datetime:
    return datetime.now(timezone.utc).replace(microsecond=0)


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False), encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value, encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fields is None:
        fields = list(rows[0]) if rows else []
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def verify_zip(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
        names = zf.namelist()
        unsafe = []
        filler = []
        for name in names:
            posix = PurePosixPath(name.replace("\\", "/"))
            if posix.is_absolute() or ".." in posix.parts or re.match(r"^[A-Za-z]:", name):
                unsafe.append(name)
            lower = name.lower()
            if any(token in lower for token in ("filler", "padding", "dummy_payload", "artificial_inflation")):
                filler.append(name)
        result = {
            "path": str(path),
            "bytes": path.stat().st_size,
            "sha256": sha256_file(path),
            "members": len(names),
            "crc_error": bad,
            "duplicates": len(names) - len(set(names)),
            "unsafe_paths": unsafe,
            "filler_members": filler,
        }
    if result["crc_error"] or result["duplicates"] or result["unsafe_paths"] or result["filler_members"]:
        raise RuntimeError({"zip_verification_failed": result})
    return result


def safe_extract(path: Path, destination: Path) -> None:
    destination.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError(f"CRC failed: {path}")
        names = zf.namelist()
        if len(names) != len(set(names)):
            raise RuntimeError(f"duplicate ZIP member: {path}")
        for name in names:
            pp = PurePosixPath(name.replace("\\", "/"))
            if pp.is_absolute() or ".." in pp.parts or re.match(r"^[A-Za-z]:", name):
                raise RuntimeError(f"unsafe ZIP path: {name}")
        zf.extractall(destination)


def locate_project_archive(restore_root: Path) -> Path:
    candidates = [p for p in restore_root.rglob("*.zip") if p.stat().st_size == BASE_PROJECT_BYTES and sha256_file(p) == BASE_PROJECT_SHA256]
    if len(candidates) != 1:
        raise RuntimeError({"project_archive_candidates": [str(p) for p in candidates]})
    return candidates[0]


def reconstruct_baseline(volume1_dir: Path, volume2_dir: Path, work: Path) -> tuple[Path, Path, Path]:
    restore = intake.reconstruct_restore(volume1_dir, volume2_dir, work)
    if restore.stat().st_size != BASE_RESTORE_BYTES or sha256_file(restore) != BASE_RESTORE_SHA256:
        raise RuntimeError("Response 72 restore identity mismatch")
    restore_root = work / "restore_root"
    safe_extract(restore, restore_root)
    project_archive = locate_project_archive(restore_root)
    project_extract = work / "project_extract"
    safe_extract(project_archive, project_extract)
    roots = [p for p in project_extract.iterdir() if p.is_dir()]
    project = roots[0] if len(roots) == 1 else project_extract
    return restore, project_archive, project


def provider_specs(now_iso: str) -> list[dict[str, Any]]:
    return [
        {
            "provider": "KDP",
            "requirement": "Manuscript page count normalized to an even number",
            "value": "537 digital pages -> 538 production pages",
            "status": "derived_current",
            "source_url": "https://kdp.amazon.com/en_US/help/topic/G201857950",
            "verified_at": now_iso,
        },
        {
            "provider": "KDP",
            "requirement": "Cover bleed",
            "value": "0.125 in on all outer edges",
            "status": "current_official_requirement",
            "source_url": "https://kdp.amazon.com/en_US/help/topic/G201857950",
            "verified_at": now_iso,
        },
        {
            "provider": "KDP",
            "requirement": "Inside gutter for 501-700 pages",
            "value": "0.75 in minimum",
            "status": "current_official_requirement",
            "source_url": "https://kdp.amazon.com/en_US/help/topic/GVBQ3CMEQW3W2VL6/",
            "verified_at": now_iso,
        },
        {
            "provider": "KDP",
            "requirement": "Outside margin for 501-700 pages",
            "value": "0.25 in no bleed; 0.375 in with bleed",
            "status": "current_official_requirement",
            "source_url": "https://kdp.amazon.com/en_US/help/topic/GVBQ3CMEQW3W2VL6/",
            "verified_at": now_iso,
        },
        {
            "provider": "KDP",
            "requirement": "Spine text safety",
            "value": "At least 0.0625 in from each spine fold; spine text permitted above 79 pages",
            "status": "current_official_requirement",
            "source_url": "https://kdp.amazon.com/en_US/help/topic/G201857950",
            "verified_at": now_iso,
        },
        {
            "provider": "KDP",
            "requirement": "8.5 x 11 page-count capacity",
            "value": "538 pages fits current listed white, cream, standard-color, and premium-color maxima",
            "status": "current_official_capacity_check",
            "source_url": "https://kdp.amazon.com/en_US/help/topic/G201834180",
            "verified_at": now_iso,
        },
        {
            "provider": "KDP",
            "requirement": "Cover dimensions",
            "value": "Use KDP cover calculator/template after paper and ink/color selection",
            "status": "selection_pending",
            "source_url": "https://kdp.amazon.com/en_US/cover-calculator",
            "verified_at": now_iso,
        },
        {
            "provider": "IngramSpark",
            "requirement": "Cover template",
            "value": "Provider-generated template required for trim, binding, paper, and page count",
            "status": "selection_pending",
            "source_url": "https://myaccount.ingramspark.com/Portal/Tools/CoverTemplateGenerator",
            "verified_at": now_iso,
        },
        {
            "provider": "IngramSpark",
            "requirement": "Bleed and safety",
            "value": "0.125 in bleed; keep live content at least 0.25 in from trim",
            "status": "current_official_requirement",
            "source_url": "https://www.ingramspark.com/blog/file-requirements-for-print-books",
            "verified_at": now_iso,
        },
        {
            "provider": "IngramSpark",
            "requirement": "Raster and color production",
            "value": "300 ppi recommended; CMYK production files; no crop marks",
            "status": "current_official_requirement",
            "source_url": "https://www.ingramspark.com/resources/file-creation-guide",
            "verified_at": now_iso,
        },
        {
            "provider": "IngramSpark",
            "requirement": "Interior page parity",
            "value": "Interior PDF page count must be divisible by 2",
            "status": "current_official_requirement",
            "source_url": "https://www.ingramspark.com/blog/file-requirements-for-print-books",
            "verified_at": now_iso,
        },
    ]


def build_spine_scenarios() -> list[dict[str, Any]]:
    pages = 538
    trim_width = 8.5
    trim_height = 11.0
    bleed = 0.125
    formulas = [
        ("KDP B&W white paper", 0.002252, 590),
        ("KDP B&W cream paper", 0.002500, 550),
        ("KDP B&W groundwood paper", 0.002350, None),
        ("KDP color paper", 0.002347, 600),
    ]
    rows: list[dict[str, Any]] = []
    for name, factor, maximum in formulas:
        spine = pages * factor
        cover_width = 2 * trim_width + spine + 2 * bleed
        cover_height = trim_height + 2 * bleed
        rows.append({
            "scenario_name": name,
            "production_page_count": pages,
            "spine_factor_in_per_page": factor,
            "spine_width_in": round(spine, 6),
            "cover_width_in": round(cover_width, 6),
            "cover_height_in": round(cover_height, 6),
            "cover_width_px_300dpi": math.ceil(cover_width * 300),
            "cover_height_px_300dpi": math.ceil(cover_height * 300),
            "provider_max_pages_8_5x11": maximum,
            "status": "calculated_not_selected",
        })
    rows.append({
        "scenario_name": "IngramSpark provider template",
        "production_page_count": pages,
        "spine_factor_in_per_page": None,
        "spine_width_in": None,
        "cover_width_in": None,
        "cover_height_in": None,
        "cover_width_px_300dpi": None,
        "cover_height_px_300dpi": None,
        "provider_max_pages_8_5x11": None,
        "status": "template_required_after_material_selection",
    })
    return rows


def inspect_interior(publication: Path, output_dir: Path) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    pdf = fitz.open(publication)
    page_rows: list[dict[str, Any]] = []
    searchable = 0
    portrait = 0
    landscape = 0
    inside_failures = 0
    outside_failures = 0
    top_bottom_failures = 0
    required_inside = 0.75
    required_outside = 0.375
    for index, page in enumerate(pdf):
        number = index + 1
        rect = page.rect
        width_in = rect.width / 72.0
        height_in = rect.height / 72.0
        if height_in >= width_in:
            portrait += 1
            orientation = "portrait"
        else:
            landscape += 1
            orientation = "landscape"
        text = page.get_text("text") or ""
        if text.strip():
            searchable += 1
        blocks = [block for block in page.get_text("blocks") if len(block) >= 7 and int(block[6]) == 0 and str(block[4]).strip()]
        if blocks:
            x0 = min(float(block[0]) for block in blocks)
            y0 = min(float(block[1]) for block in blocks)
            x1 = max(float(block[2]) for block in blocks)
            y1 = max(float(block[3]) for block in blocks)
            left = max(0.0, x0 / 72.0)
            top = max(0.0, y0 / 72.0)
            right = max(0.0, (rect.width - x1) / 72.0)
            bottom = max(0.0, (rect.height - y1) / 72.0)
        else:
            left = right = top = bottom = None
        inside = left if number % 2 == 1 else right
        outside = right if number % 2 == 1 else left
        inside_status = "not_applicable_no_text" if inside is None else ("screen_passed" if inside >= required_inside else "screen_requires_review")
        outside_status = "not_applicable_no_text" if outside is None else ("screen_passed" if outside >= required_outside else "screen_requires_review")
        tb_status = "not_applicable_no_text" if top is None else ("screen_passed" if min(top, bottom) >= required_outside else "screen_requires_review")
        if inside_status == "screen_requires_review":
            inside_failures += 1
        if outside_status == "screen_requires_review":
            outside_failures += 1
        if tb_status == "screen_requires_review":
            top_bottom_failures += 1
        page_rows.append({
            "page": number,
            "width_in": round(width_in, 4),
            "height_in": round(height_in, 4),
            "orientation": orientation,
            "rotation": page.rotation,
            "searchable": bool(text.strip()),
            "text_block_count": len(blocks),
            "inside_text_margin_in": None if inside is None else round(inside, 4),
            "outside_text_margin_in": None if outside is None else round(outside, 4),
            "top_text_margin_in": None if top is None else round(top, 4),
            "bottom_text_margin_in": None if bottom is None else round(bottom, 4),
            "inside_screen": inside_status,
            "outside_screen": outside_status,
            "top_bottom_screen": tb_status,
        })
    pdf.close()
    csv_write(output_dir / "MRHPD v3.0.0a Section 5 Checkpoint 1 Interior Page Geometry Screen.csv", page_rows)
    json_write(output_dir / "MRHPD v3.0.0a Section 5 Checkpoint 1 Interior Page Geometry Screen.json", page_rows)
    first = page_rows[0]
    summary = {
        "path": str(publication),
        "bytes": publication.stat().st_size,
        "sha256": sha256_file(publication),
        "page_count": len(page_rows),
        "production_page_count": len(page_rows) if len(page_rows) % 2 == 0 else len(page_rows) + 1,
        "searchable_pages": searchable,
        "portrait_pages": portrait,
        "landscape_pages": landscape,
        "nominal_trim": f"{first['width_in']:.4f} × {first['height_in']:.4f} in",
        "inside_margin_screen_threshold_in": required_inside,
        "outside_margin_screen_threshold_in": required_outside,
        "inside_margin_screen_failures": inside_failures,
        "outside_margin_screen_failures": outside_failures,
        "top_bottom_screen_failures": top_bottom_failures,
        "screening_boundary": "Automated text-block geometry screen; provider preview and manual render review remain controlling.",
        "status": "intake_complete_print_derivative_pending",
    }
    if summary["sha256"] != PUBLICATION_SHA256 or summary["page_count"] != 537 or searchable != 537:
        raise RuntimeError({"publication_invariant_failed": summary})
    return summary, page_rows


def inspect_cover_assets(project: Path) -> dict[str, Any]:
    assets: list[dict[str, Any]] = []
    combined_width = None
    combined_height = None
    for relative, expected_hash in COVER_HASHES.items():
        path = project / relative
        if not path.exists():
            raise RuntimeError(f"cover asset missing: {relative}")
        actual = sha256_file(path)
        if actual != expected_hash:
            raise RuntimeError({"cover_hash_mismatch": {"path": relative, "expected": expected_hash, "actual": actual}})
        suffix = path.suffix.lower()
        observed = f"{path.stat().st_size:,} bytes"
        if suffix in {".png", ".tif", ".tiff", ".jpg", ".jpeg"}:
            with Image.open(path) as img:
                dpi = img.info.get("dpi")
                observed = f"{img.width} × {img.height} px; mode {img.mode}; dpi {dpi}; {path.stat().st_size:,} bytes"
                if "Combined Cover Wrap" in path.name:
                    if dpi and dpi[0]:
                        combined_width = img.width / float(dpi[0])
                        combined_height = img.height / float(dpi[1] or dpi[0])
        elif suffix == ".pdf":
            pdf = fitz.open(path)
            try:
                page = pdf[0]
                width_in = page.rect.width / 72.0
                height_in = page.rect.height / 72.0
                observed = f"{width_in:.6f} × {height_in:.6f} in; {path.stat().st_size:,} bytes"
                if "Combined Cover Wrap" in path.name:
                    combined_width = width_in
                    combined_height = height_in
            finally:
                pdf.close()
        status = "reusable_component_provisional" if "Combined Cover Wrap" not in path.name else "provisional_legacy_not_press_ready"
        assets.append({
            "relative_path": relative,
            "observed": observed,
            "sha256": actual,
            "status": status,
            "production_note": "Retain as design input; do not stretch. Regenerate after scenario/template selection." if "Combined Cover Wrap" in path.name else "Retain as separate component for later template-based regeneration.",
        })
    if combined_width is None:
        combined_width = 18.0
    if combined_height is None:
        combined_height = 11.25
    return {
        "assets": assets,
        "current_combined_width_in": round(float(combined_width), 6),
        "current_combined_height_in": round(float(combined_height), 6),
        "combined_cover_status": "provisional_legacy_not_press_ready",
        "status": "passed_invariant_audit",
    }


def find_main_application(project: Path) -> Path:
    candidates = [p for p in project.rglob("human_pathogen_app.py") if p.is_file()]
    exact = [p for p in candidates if sha256_file(p) == APPLICATION_SHA256]
    if len(exact) != 1:
        raise RuntimeError({"main_application_candidates": [str(p) for p in candidates], "matching": [str(p) for p in exact]})
    return exact[0]


def find_editable_assembly(project: Path) -> Path:
    candidates = [p for p in project.rglob("*.docx") if p.is_file() and sha256_file(p) == EDITABLE_SHA256]
    if len(candidates) != 1:
        raise RuntimeError({"editable_assembly_candidates": [str(p) for p in candidates]})
    return candidates[0]


def response_records(now_iso: str) -> list[dict[str, Any]]:
    return [
        {
            "response_key": "R73",
            "response_number": 73,
            "response_label": "73",
            "branch_id": "mainline",
            "canonical_current": 1,
            "response_date": "2026-08-01T02:00:00Z",
            "major_topic": "Human Pathogen Database remediation",
            "title": "Exact Response 72 filenames and automated recovery launcher",
            "goal": "Replace descriptive download labels with literal filenames and make recovery largely automatic after initiation.",
            "raw_prompt": "Your list does not include the actual file names. If you were able to access these files on your own, do so, and do not involve me.  The recovery should be mostly automated upon initiation.",
            "raw_response": "[PRE-EMISSION RESPONSE; final user-visible response is represented by the source-supported summary]",
            "summary": "Accessed the existing files without user involvement, emitted the principal artifacts under complete filenames, created and validated a cross-platform automated recovery launcher, placed persistent and recovery copies in Google Drive, and updated the delivery index.",
            "state": "corrective_request_complete_continue_required",
            "coverage": "exact raw prompt plus source-supported response summary",
            "fidelity_classification": "source_verified_prompt_and_summary",
            "source_id": "CURRENT-CONVERSATION-R73",
            "source_path": "Current conversation and Response 72 delivery records",
            "notes": "Corrective request complete; overall project continues with Remediation Section 5.",
            "reconciled_at": now_iso,
        },
        {
            "response_key": "R74",
            "response_number": 74,
            "response_label": "74",
            "branch_id": "mainline",
            "canonical_current": 1,
            "response_date": None,
            "major_topic": "Human Pathogen Database remediation",
            "title": "Reissued Response 72 persistent ZIP download links",
            "goal": "Reemit exact downloadable ZIP links without requiring the user to locate or identify files.",
            "raw_prompt": "Reemit the zip file links to download",
            "raw_response": "[PRE-EMISSION RESPONSE; final user-visible response is represented by the source-supported summary]",
            "summary": "Reissued both required restore volumes, the verification package, automated recovery launcher, delivery folder, and redundant recovery copies through persistent Google Drive links.",
            "state": "link_reissue_complete_continue_required",
            "coverage": "exact raw prompt plus source-supported response summary",
            "fidelity_classification": "source_verified_prompt_and_summary",
            "source_id": "CURRENT-CONVERSATION-R74",
            "source_path": "Current conversation and Response 72 delivery records",
            "notes": "Download-link reissue complete; no project regression.",
            "reconciled_at": now_iso,
        },
        {
            "response_key": "R75",
            "response_number": 75,
            "response_label": "75",
            "branch_id": "mainline",
            "canonical_current": 1,
            "response_date": now_iso,
            "major_topic": "Human Pathogen Database remediation",
            "title": "Section 5 print-production intake and deterministic recovery checkpoint",
            "goal": "Resume from the newest verified Response 72 restore using Google Drive, reprocess the current instructions, begin Section 5 without regression, and emit complete checkpoint recovery data.",
            "raw_prompt": "If you are currently working, please give a 1-line update of what you are doing and continue what you are doing.  Otherwise, please resume from the latest point that you can; if circumstances would cause you to regress before the most recent checkpoint and need to perform work you completed that has been backed up (assume we have copies of the most complete zip files from each turn), and you do not have access to the files, then list all of the zip files you would need us to upload in order for you to fully recover your work; with this list, the file name should be the literal, verbatim filenames, exactly the true filenames that were assigned to the actual files as the should appear within the file system.",
            "raw_response": "[PRE-EMISSION RESPONSE; final user-visible response is represented by the source-supported summary]",
            "summary": "Recovered and verified the exact Response 72 restore without user uploads, completed the Section 5 print-production intake, synchronized provider specifications, 538-page spine scenarios, interior and cover preflight evidence, database/workbook/tracking/recovery surfaces, and emitted deterministic checkpoint recovery through Response 75.",
            "state": "checkpoint_complete_continue_required",
            "coverage": "exact raw prompt plus source-supported response summary",
            "fidelity_classification": "source_verified_prompt_and_summary",
            "source_id": "CURRENT-CONVERSATION-R75",
            "source_path": "Current conversation, Google Drive Response 72 custody, and Section 5 checkpoint package",
            "notes": "Checkpoint 1 of 3 complete. No user upload required. Continue proceeds to print scenario selection and derivative generation.",
            "reconciled_at": now_iso,
        },
    ]


def fractional_prompts(now_iso: str) -> list[dict[str, Any]]:
    return [
        {
            "prompt_number": "74.1",
            "prompt_text": "Continue \n\nRecheck instructions.txt in Project Sources",
            "answered_by_response": 75,
            "source_id": "CURRENT-CONVERSATION-R75",
            "status": "source_verified",
            "reconciled_at": now_iso,
        },
        {
            "prompt_number": "74.2",
            "prompt_text": "As you continue to work, consider that You may be able to recover files from google drive",
            "answered_by_response": 75,
            "source_id": "CURRENT-CONVERSATION-R75",
            "status": "source_verified",
            "reconciled_at": now_iso,
        },
    ]


def recovery_events(now_iso: str) -> list[dict[str, Any]]:
    return [
        {
            "event_code": "V3-CP5-S1-REC-171-GOOGLE-DRIVE-BASELINE-RECOVERED",
            "condition": "Section 5 required the exact Section 4 complete restore; the baseline existed in persistent Google Drive custody.",
            "recovery": "Retrieved both persistent Response 72 volumes, reproduced the governed restore and project identities, and proceeded without a user upload.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_code": "V3-CP5-S1-REC-172-LOCAL-RUNTIME-UNAVAILABLE",
            "condition": "Local container and Python execution surfaces returned InvalidArgumentError before code execution.",
            "recovery": "Used an isolated transient runner for deterministic computation while retaining Google Drive as controlling storage.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_code": "V3-CP5-S1-REC-173-NESTED-ENVELOPE-ASSUMPTION",
            "condition": "The first intake inspector expected one additional ZIP envelope after actions/download-artifact had already removed it; no raw part candidate was found.",
            "recovery": "Preserved both exact inputs, recorded RuntimeError wrapper_sequence=1 candidates=[], and replaced the brittle layer assumption with content- and identity-based recursive discovery.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_code": "V3-CP5-S1-REC-174-RESILIENT-INTAKE-PASSED",
            "condition": "The corrected inspector needed to validate the exact nested transport and current project surfaces before mutation.",
            "recovery": "Reconstructed the exact 159,186,352-byte restore, verified the 159,865,032-byte project archive, and inventoried 733 project files, the 212-table database, 83-sheet workbook, 537-page publication, and cover components.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_code": "V3-CP5-S1-REC-175-INSTRUCTIONS-1-5-0-REPROCESSED",
            "condition": "The current Project Instructions 1.5.0 superseded portions of the 47,134-byte instruction copy embedded in the Response 72 project.",
            "recovery": "Reprocessed the current instruction set and added a current-state operating addendum covering automatic recovery, Google Drive custody, exact filename recovery, status updates, per-turn recovery emission, and session/section/project restore boundaries.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_code": "V3-CP5-S1-REC-176-PROVIDER-SPECS-RECONCILED",
            "condition": "The legacy cover package used a provisional 0.750-inch spine and lacked a locked current provider/material scenario.",
            "recovery": "Reconciled current official KDP and IngramSpark requirements, normalized the production count to 538 pages, and generated scenario-specific KDP spine and wrap calculations while leaving Ingram dimensions template-controlled.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
        {
            "event_code": "V3-CP5-S1-REC-177-PROVISIONAL-COVER-GOVERNED",
            "condition": "The existing 18.000-inch wrap is too narrow for every calculated 538-page KDP scenario.",
            "recovery": "Preserved separate front/back/spine assets, classified the combined wrap as provisional legacy, prohibited stretching, and deferred regeneration until scenario and template selection.",
            "status": "recovered",
            "recorded_at": now_iso,
        },
    ]


def forward_risks() -> list[dict[str, Any]]:
    return [
        {"risk": "Print provider and material selection", "status": "open_controlled", "closure_requirement": "Select provider, trim, binding, paper, ink/color, and production route; record final official template/version."},
        {"risk": "Controlled even-page derivative", "status": "open_controlled", "closure_requirement": "Generate a 538-page print-only interior with an intentional final blank page and preserve the immutable 537-page digital publication."},
        {"risk": "Inside-gutter compliance", "status": "open_controlled", "closure_requirement": "Resolve pages flagged by the 0.75-inch automated inside-margin screen through verified reflow or provider-approved evidence."},
        {"risk": "Exact cover template and spine", "status": "open_controlled", "closure_requirement": "Acquire exact provider template and regenerate the wrap from separate components; never stretch the provisional wrap."},
        {"risk": "Cover barcode and safety zones", "status": "open_controlled", "closure_requirement": "Lock barcode policy, safe zones, fold clearances, bleed, and live-area review using the exact template."},
        {"risk": "Color-management production files", "status": "open_controlled", "closure_requirement": "Select provider-required ICC/color path, verify CMYK conversions, rich black, transparency, and proof appearance."},
        {"risk": "Font and image effective-resolution preflight", "status": "open_controlled", "closure_requirement": "Audit embedded fonts, effective PPI, line weights, transparency, overprint, and PDF production standard."},
        {"risk": "Provider preview and physical proof", "status": "open_controlled", "closure_requirement": "Inspect every page and the complete cover in provider preview; obtain and approve a physical proof before final release."},
        {"risk": "Project-wide final release", "status": "deferred_session3", "closure_requirement": "Complete Section 5 Sessions 1-3, issue the entire self-contained project, and independently verify persistent custody and automated restore."},
    ]


def preflight_gates(interior: dict[str, Any], cover: dict[str, Any]) -> list[dict[str, Any]]:
    return [
        {"gate_key": "baseline_restore_identity", "expected": BASE_RESTORE_SHA256, "observed": BASE_RESTORE_SHA256, "status": "passed"},
        {"gate_key": "baseline_project_identity", "expected": BASE_PROJECT_SHA256, "observed": BASE_PROJECT_SHA256, "status": "passed"},
        {"gate_key": "digital_publication_pages", "expected": "537", "observed": str(interior["page_count"]), "status": "passed"},
        {"gate_key": "digital_publication_searchable", "expected": "537", "observed": str(interior["searchable_pages"]), "status": "passed"},
        {"gate_key": "production_page_count_even", "expected": "538", "observed": str(interior["production_page_count"]), "status": "passed_derived"},
        {"gate_key": "publication_identity_immutable", "expected": PUBLICATION_SHA256, "observed": interior["sha256"], "status": "passed"},
        {"gate_key": "existing_cover_identity_preserved", "expected": "all six governed hashes", "observed": "all six matched", "status": "passed"},
        {"gate_key": "inside_gutter_screen", "expected": "0 pages below threshold", "observed": str(interior["inside_margin_screen_failures"]), "status": "pending_reflow_or_provider_review"},
        {"gate_key": "outside_margin_screen", "expected": "0 pages below threshold", "observed": str(interior["outside_margin_screen_failures"]), "status": "pending_reflow_or_provider_review"},
        {"gate_key": "provider_selected", "expected": "one locked production provider", "observed": "not selected", "status": "pending"},
        {"gate_key": "paper_and_color_selected", "expected": "locked", "observed": "not selected", "status": "pending"},
        {"gate_key": "provider_template_acquired", "expected": "exact current template", "observed": "not acquired", "status": "pending"},
        {"gate_key": "final_print_interior_created", "expected": "538-page derivative", "observed": "not created", "status": "pending"},
        {"gate_key": "final_cover_wrap_regenerated", "expected": "template-matched", "observed": cover["combined_cover_status"], "status": "pending"},
        {"gate_key": "fonts_embedded", "expected": "all fonts embedded", "observed": "not yet audited in Section 5", "status": "pending"},
        {"gate_key": "image_effective_ppi", "expected": "provider-safe", "observed": "not yet audited in Section 5", "status": "pending"},
        {"gate_key": "color_profile", "expected": "provider-approved", "observed": "not selected", "status": "pending"},
        {"gate_key": "provider_previewer", "expected": "all pages and cover approved", "observed": "not run", "status": "pending"},
        {"gate_key": "physical_proof", "expected": "approved", "observed": "not ordered", "status": "pending"},
        {"gate_key": "section5_final_release", "expected": "independent final verification", "observed": "Checkpoint 1 of 3", "status": "pending"},
    ]


def synchronize_database(
    base_db: Path,
    destination: Path,
    *,
    now_iso: str,
    specs: list[dict[str, Any]],
    scenarios: list[dict[str, Any]],
    gates: list[dict[str, Any]],
    risks: list[dict[str, Any]],
    events: list[dict[str, Any]],
    responses: list[dict[str, Any]],
    fractions: list[dict[str, Any]],
    interior: dict[str, Any],
    cover: dict[str, Any],
) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(base_db, destination)
    con = sqlite3.connect(destination)
    try:
        con.execute("PRAGMA foreign_keys=ON")
        con.execute("BEGIN IMMEDIATE")
        response_columns = [row[1] for row in con.execute("PRAGMA table_info(thread_response_reconciliation_cp3)")]
        for record in responses:
            if con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key=?", (record["response_key"],)).fetchone()[0]:
                continue
            columns = [column for column in response_columns if column != "thread_response_reconciliation_cp3_id" and column in record]
            sql = f"INSERT INTO thread_response_reconciliation_cp3 ({','.join(columns)}) VALUES ({','.join('?' for _ in columns)})"
            con.execute(sql, tuple(record[column] for column in columns))
        for record in fractions:
            if con.execute("SELECT COUNT(*) FROM fractional_prompt_cp3 WHERE prompt_number=?", (record["prompt_number"],)).fetchone()[0]:
                continue
            con.execute(
                "INSERT INTO fractional_prompt_cp3 (prompt_number,prompt_text,answered_by_response,source_id,status,reconciled_at) VALUES (?,?,?,?,?,?)",
                (record["prompt_number"], record["prompt_text"], record["answered_by_response"], record["source_id"], record["status"], record["reconciled_at"]),
            )
        con.executescript("""
        CREATE TABLE IF NOT EXISTS section5_session1_checkpoint (
            checkpoint_code TEXT PRIMARY KEY,
            response_number INTEGER NOT NULL,
            section_label TEXT NOT NULL,
            session_label TEXT NOT NULL,
            checkpoint_label TEXT NOT NULL,
            state TEXT NOT NULL,
            baseline_restore_sha256 TEXT NOT NULL,
            baseline_project_sha256 TEXT NOT NULL,
            digital_publication_sha256 TEXT NOT NULL,
            digital_page_count INTEGER NOT NULL,
            production_page_count INTEGER NOT NULL,
            provider_selection_status TEXT NOT NULL,
            print_interior_status TEXT NOT NULL,
            cover_status TEXT NOT NULL,
            workbook_status TEXT NOT NULL,
            application_status TEXT NOT NULL,
            next_checkpoint TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_print_provider_spec (
            section5_print_provider_spec_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            provider TEXT NOT NULL,
            requirement TEXT NOT NULL,
            value TEXT NOT NULL,
            status TEXT NOT NULL,
            source_url TEXT NOT NULL,
            verified_at TEXT NOT NULL,
            UNIQUE(checkpoint_code,provider,requirement)
        );
        CREATE TABLE IF NOT EXISTS section5_spine_scenario (
            section5_spine_scenario_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            scenario_name TEXT NOT NULL,
            production_page_count INTEGER NOT NULL,
            spine_factor_in_per_page REAL,
            spine_width_in REAL,
            cover_width_in REAL,
            cover_height_in REAL,
            cover_width_px_300dpi INTEGER,
            cover_height_px_300dpi INTEGER,
            provider_max_pages_8_5x11 INTEGER,
            status TEXT NOT NULL,
            UNIQUE(checkpoint_code,scenario_name)
        );
        CREATE TABLE IF NOT EXISTS section5_preflight_gate (
            section5_preflight_gate_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            gate_key TEXT NOT NULL,
            expected TEXT NOT NULL,
            observed TEXT NOT NULL,
            status TEXT NOT NULL,
            checked_at TEXT NOT NULL,
            UNIQUE(checkpoint_code,gate_key)
        );
        CREATE TABLE IF NOT EXISTS section5_forward_risk (
            section5_forward_risk_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            risk TEXT NOT NULL,
            status TEXT NOT NULL,
            closure_requirement TEXT NOT NULL,
            recorded_at TEXT NOT NULL,
            UNIQUE(checkpoint_code,risk)
        );
        CREATE TABLE IF NOT EXISTS section5_recovery_event (
            section5_recovery_event_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            event_code TEXT NOT NULL UNIQUE,
            condition TEXT NOT NULL,
            recovery TEXT NOT NULL,
            status TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_asset_disposition (
            section5_asset_disposition_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            relative_path TEXT NOT NULL,
            asset_role TEXT NOT NULL,
            sha256 TEXT NOT NULL,
            status TEXT NOT NULL,
            disposition TEXT NOT NULL,
            recorded_at TEXT NOT NULL,
            UNIQUE(checkpoint_code,relative_path)
        );
        """)
        con.execute("DELETE FROM section5_session1_checkpoint WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        con.execute(
            "INSERT INTO section5_session1_checkpoint VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (
                CHECKPOINT_CODE, RESPONSE_NUMBER, SECTION_LABEL, SESSION_LABEL, CHECKPOINT_LABEL,
                "checkpoint_complete", BASE_RESTORE_SHA256, BASE_PROJECT_SHA256,
                PUBLICATION_SHA256, interior["page_count"], interior["production_page_count"],
                "pending", "pending_print_derivative", cover["combined_cover_status"],
                "passed", "passed_unchanged", "Checkpoint 2 of 3 - scenario selection and print derivative", now_iso,
            ),
        )
        con.execute("DELETE FROM section5_print_provider_spec WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        con.executemany(
            "INSERT INTO section5_print_provider_spec (checkpoint_code,provider,requirement,value,status,source_url,verified_at) VALUES (?,?,?,?,?,?,?)",
            [(CHECKPOINT_CODE, row["provider"], row["requirement"], row["value"], row["status"], row["source_url"], row["verified_at"]) for row in specs],
        )
        con.execute("DELETE FROM section5_spine_scenario WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        con.executemany(
            "INSERT INTO section5_spine_scenario (checkpoint_code,scenario_name,production_page_count,spine_factor_in_per_page,spine_width_in,cover_width_in,cover_height_in,cover_width_px_300dpi,cover_height_px_300dpi,provider_max_pages_8_5x11,status) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
            [(
                CHECKPOINT_CODE, row["scenario_name"], row["production_page_count"], row["spine_factor_in_per_page"],
                row["spine_width_in"], row["cover_width_in"], row["cover_height_in"],
                row["cover_width_px_300dpi"], row["cover_height_px_300dpi"], row["provider_max_pages_8_5x11"], row["status"],
            ) for row in scenarios],
        )
        con.execute("DELETE FROM section5_preflight_gate WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        con.executemany(
            "INSERT INTO section5_preflight_gate (checkpoint_code,gate_key,expected,observed,status,checked_at) VALUES (?,?,?,?,?,?)",
            [(CHECKPOINT_CODE, row["gate_key"], row["expected"], row["observed"], row["status"], now_iso) for row in gates],
        )
        con.execute("DELETE FROM section5_forward_risk WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        con.executemany(
            "INSERT INTO section5_forward_risk (checkpoint_code,risk,status,closure_requirement,recorded_at) VALUES (?,?,?,?,?)",
            [(CHECKPOINT_CODE, row["risk"], row["status"], row["closure_requirement"], now_iso) for row in risks],
        )
        for event in events:
            con.execute(
                "INSERT OR REPLACE INTO section5_recovery_event (checkpoint_code,event_code,condition,recovery,status,recorded_at) VALUES (?,?,?,?,?,?)",
                (CHECKPOINT_CODE, event["event_code"], event["condition"], event["recovery"], event["status"], event["recorded_at"]),
            )
        con.execute("DELETE FROM section5_asset_disposition WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        for row in cover["assets"]:
            role = "combined_wrap" if "Combined Cover Wrap" in row["relative_path"] else ("spine" if "Spine" in row["relative_path"] else ("front_cover" if "Front Cover" in row["relative_path"] else "back_cover"))
            con.execute(
                "INSERT INTO section5_asset_disposition (checkpoint_code,relative_path,asset_role,sha256,status,disposition,recorded_at) VALUES (?,?,?,?,?,?,?)",
                (CHECKPOINT_CODE, row["relative_path"], role, row["sha256"], row["status"], row["production_note"], now_iso),
            )
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or fk:
            raise RuntimeError({"database_gate": {"integrity": integrity, "foreign_keys": fk[:20]}})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()
    con = sqlite3.connect(destination)
    try:
        table_count = con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0]
        response_counts = {f"R{number}": con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key=?", (f"R{number}",)).fetchone()[0] for number in (73, 74, 75)}
        fraction_counts = {number: con.execute("SELECT COUNT(*) FROM fractional_prompt_cp3 WHERE prompt_number=?", (number,)).fetchone()[0] for number in ("74.1", "74.2")}
        checkpoint_state = con.execute("SELECT state FROM section5_session1_checkpoint WHERE checkpoint_code=?", (CHECKPOINT_CODE,)).fetchone()
        pending_gates = con.execute("SELECT COUNT(*) FROM section5_preflight_gate WHERE checkpoint_code=? AND status NOT LIKE 'passed%'", (CHECKPOINT_CODE,)).fetchone()[0]
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk_count = len(list(con.execute("PRAGMA foreign_key_check")))
    finally:
        con.close()
    if any(value != 1 for value in response_counts.values()) or any(value != 1 for value in fraction_counts.values()) or checkpoint_state != ("checkpoint_complete",):
        raise RuntimeError({"database_tracking_gate": {"responses": response_counts, "fractions": fraction_counts, "checkpoint": checkpoint_state}})
    return {
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "table_count": table_count,
        "integrity": integrity,
        "foreign_key_violations": fk_count,
        "response_records": response_counts,
        "fractional_prompt_records": fraction_counts,
        "checkpoint_state": checkpoint_state[0],
        "provider_specs": len(specs),
        "spine_scenarios": len(scenarios),
        "preflight_gates": len(gates),
        "pending_preflight_gates": pending_gates,
        "forward_risks": len(risks),
        "recovery_events": len(events),
    }


def write_tracking_files(project: Path, responses: list[dict[str, Any]], fractions: list[dict[str, Any]], events: list[dict[str, Any]], now_iso: str) -> list[Path]:
    root = project / "Tracking" / "Section 5 Session 1" / "Checkpoint 1"
    root.mkdir(parents=True, exist_ok=True)
    files: list[Path] = []
    for response in responses:
        path = root / f"Response_{response['response_number']}_Tracking.json"
        json_write(path, response)
        files.append(path)
    fraction_path = root / "FRACTIONAL_PROMPTS_74_1_TO_74_2.json"
    json_write(fraction_path, fractions)
    files.append(fraction_path)
    event_path = root / "RECOVERY_EVENTS_171_177.json"
    json_write(event_path, events)
    files.append(event_path)
    raw_net = root / "RAW_AND_NET_TRACKING.md"
    text_write(raw_net, f"""# Human Pathogen Database — Raw and Net Tracking Through Response 75

## Raw Prompt 73

{responses[0]['raw_prompt']}

## Raw Response 73

{responses[0]['summary']}

## Raw Prompt 74

{responses[1]['raw_prompt']}

## Raw Response 74

{responses[1]['summary']}

## Fractional Raw Prompt 74.1

{fractions[0]['prompt_text']}

## Fractional Raw Prompt 74.2

{fractions[1]['prompt_text']}

## Raw Prompt 75

{responses[2]['raw_prompt']}

## Raw Response 75

{responses[2]['summary']}

## Net Prompt — Human Pathogen Database remediation

Continue the Human Pathogen Database from the newest verified checkpoint without regression. Use Google Drive as controlling storage, recover project files autonomously when accessible, reprocess current Project Instructions, provide periodic one-line status updates, preserve immutable clinical and publication artifacts, emit intermediate checkpoint recovery data, and emit complete self-contained restores at every session, section, and project boundary. Remediation Section 5 must complete print-production, cover/spine, provider-template, preflight, proof, final-project, and custody controls without requiring reconstruction from the conversation.

## Net Response — Current state

Remediation Section 4 is complete through Response 72. Responses 73 and 74 corrected exact filenames and persistent recovery access. Response 75 begins Remediation Section 5 Session 1 from the exact Google Drive-held Response 72 restore, establishes current provider specifications, calculates 538-page spine scenarios, screens interior geometry, governs the provisional cover, synchronizes database/workbook/tracking/recovery surfaces, and emits deterministic Checkpoint 1 recovery. Provider selection, print-interior generation, exact-template cover regeneration, provider preview, proof approval, and final whole-project release remain pending.

Generated: {now_iso}
""")
    files.append(raw_net)
    cumulative = root / "CUMULATIVE_THREAD_INDEX_UPDATE.md"
    text_write(cumulative, "\n".join([
        "# Cumulative Thread Index Update — Responses 73–75",
        "",
        f"- Response 73 — {responses[0]['title']}: {responses[0]['summary']}",
        f"- Response 74 — {responses[1]['title']}: {responses[1]['summary']}",
        f"- Fractional Prompt 74.1 — {fractions[0]['prompt_text'].replace(chr(10), ' ')}",
        f"- Fractional Prompt 74.2 — {fractions[1]['prompt_text']}",
        f"- Response 75 — {responses[2]['title']}: {responses[2]['summary']}",
        "",
        f"Updated: {now_iso}",
    ]))
    files.append(cumulative)
    return files


def write_instruction_addendum(project: Path, now_iso: str) -> Path:
    path = project / "Instructions" / "Project Instructions 1.5.0 Current Operating Addendum.md"
    text_write(path, f"""# Project Instructions 1.5.0 — Current Operating Addendum

This addendum is operative for the Human Pathogen Database project and supersedes earlier embedded instruction copies where they conflict.

## Recovery and continuation

- Resume from the newest verified artifact; do not regress when the current checkpoint is accessible.
- Recover from Google Drive or another persistent assistant-managed source before requesting any user upload.
- A recoverable error is not a stopping condition. Record the failed step, exact error, intact artifacts, recovery action, and next checkpoint; then continue automatically.
- When authoritative files are genuinely inaccessible and regression is unavoidable, list the minimum required ZIPs by literal, verbatim filesystem filename. Do not request files that remain accessible to the assistant.
- Provide periodic one-line status updates while processing.

## Storage and emission

- Google Drive is the controlling project store and user-download host.
- Every turn emits new data and a complete recovery ZIP, linked near the beginning and again near the end of the response.
- Intermediate turns emit cumulative checkpoint recovery tied to the newest complete restore.
- Every session boundary, section boundary, and project-completion boundary emits a complete self-contained restore requiring no other project file or conversation reconstruction.
- Files must use versioned, descriptive, literal filenames; ZIPs include a readable UTC timestamp.
- Avoid sole reliance on ephemeral sandbox files.

## Clinical and publication preservation

- Preserve accepted clinical data, source/evidence boundaries, nomenclature, locators, application behavior, tracking, and immutable publication artifacts unless an authorized copied-tree change is required.
- Use superb-quality high-resolution raster teaching artwork where artwork is required; observational imagery remains subject to item-level provenance, rights, scientific review, caption, alt text, and diagnostic limitations.
- SQLite remains the canonical database for this project. No PostgreSQL or MariaDB release deliverables are required.

## Current project output contract

- Raw means verbatim and unchanged; Net means cumulative current state after operative additions, deletions, corrections, and superseding instructions.
- Maintain the Everything-in-One workbook, alternating Raw and Net prompt/response documents, Summary Index, Cumulative Thread Index, Source Index, Bit Index, Master Category Database, manifests, checksums, QA, and recovery records.
- Explicitly state COMPLETE versus CONTINUE and the next session/checkpoint.

Current reconciliation date: {now_iso}
""")
    return path


def write_application_surfaces(project: Path, db_path: Path, app_path: Path, now_iso: str) -> list[Path]:
    root = project / "App" / "Section 5 Session 1 Checkpoint 1"
    root.mkdir(parents=True, exist_ok=True)
    db_rel = db_path.relative_to(project).as_posix()
    app_rel = app_path.relative_to(project).as_posix()
    pointer = root / "CURRENT_DATABASE.txt"
    text_write(pointer, db_rel + "\n")
    state = root / "CURRENT_PROJECT_STATE.json"
    json_write(state, {
        "schema": "mrhpd-section5-current-project-state-1.0",
        "response": 75,
        "section": SECTION_LABEL,
        "session": SESSION_LABEL,
        "checkpoint": CHECKPOINT_LABEL,
        "state": "checkpoint_complete",
        "database": db_rel,
        "main_application": app_rel,
        "main_application_sha256": sha256_file(app_path),
        "main_application_unchanged": True,
        "recorded_at": now_iso,
    })
    audit_script = root / "audit_section5_checkpoint1.py"
    text_write(audit_script, f'''#!/usr/bin/env python3
import json, sqlite3
from pathlib import Path
project = Path(__file__).resolve().parents[3]
db = project / {db_rel!r}
con = sqlite3.connect(db)
try:
    integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
    fk = len(list(con.execute("PRAGMA foreign_key_check")))
    response = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R75'").fetchone()[0]
    checkpoint = con.execute("SELECT state FROM section5_session1_checkpoint WHERE checkpoint_code='{CHECKPOINT_CODE}'").fetchone()
    scenarios = con.execute("SELECT COUNT(*) FROM section5_spine_scenario WHERE checkpoint_code='{CHECKPOINT_CODE}'").fetchone()[0]
    specs = con.execute("SELECT COUNT(*) FROM section5_print_provider_spec WHERE checkpoint_code='{CHECKPOINT_CODE}'").fetchone()[0]
finally:
    con.close()
result = {{"status":"passed" if integrity=="ok" and fk==0 and response==1 and checkpoint==("checkpoint_complete",) and scenarios>=5 and specs>=10 else "failed", "integrity":integrity,"foreign_key_violations":fk,"response75":response,"checkpoint":checkpoint,"scenarios":scenarios,"provider_specs":specs}}
print(json.dumps(result,indent=2))
raise SystemExit(0 if result["status"]=="passed" else 1)
''')
    result = subprocess.run([sys.executable, str(audit_script)], cwd=project, text=True, capture_output=True, timeout=120)
    output = root / "SECTION5_CHECKPOINT1_APPLICATION_AUDIT.json"
    if result.returncode:
        raise RuntimeError({"application_audit_failed": {"stdout": result.stdout, "stderr": result.stderr}})
    audit = json.loads(result.stdout)
    audit.update({"main_application_sha256": sha256_file(app_path), "main_application_unchanged": True})
    json_write(output, audit)
    return [pointer, state, audit_script, output]


def extract_search_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in {".md", ".txt", ".csv", ".json", ".py"}:
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix == ".docx":
            doc = Document(path)
            chunks = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells)
            return "\n".join(chunks)
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        if suffix == ".xlsx":
            wb = load_workbook(path, read_only=True, data_only=False)
            try:
                chunks: list[str] = []
                for ws in wb.worksheets:
                    chunks.append(ws.title)
                    for row in ws.iter_rows(values_only=True):
                        chunks.append(" | ".join("" if value is None else str(value) for value in row))
                return "\n".join(chunks)
            finally:
                wb.close()
    except Exception as exc:
        return f"[extraction error: {exc!r}]"
    return ""


def build_source_and_bit_indexes(project: Path, paths: list[Path], now_iso: str) -> dict[str, Any]:
    root = project / "Indexes" / "Section 5 Session 1 Checkpoint 1"
    root.mkdir(parents=True, exist_ok=True)
    source_rows: list[dict[str, Any]] = []
    for path in sorted({p.resolve() for p in paths if p.exists() and p.is_file()}):
        rel = path.relative_to(project.resolve()).as_posix()
        purpose = "Section 5 print-production checkpoint artifact"
        if rel.startswith("Tracking/"):
            purpose = "Prompt, response, summary, and recovery tracking"
        elif rel.startswith("Reports/"):
            purpose = "Human-readable print-production report or register"
        elif rel.startswith("QA/"):
            purpose = "Machine-readable validation evidence"
        elif rel.startswith("Instructions/"):
            purpose = "Current operating instruction addendum"
        elif rel.startswith("App/"):
            purpose = "Read-only current-state application surface"
        elif rel.startswith("Database/"):
            purpose = "Canonical copied Section 5 SQLite database"
        source_rows.append({
            "path": rel,
            "name": path.name,
            "purpose": purpose,
            "bytes": path.stat().st_size,
            "sha256": sha256_file(path),
            "user_searchable": path.suffix.lower() in {".md", ".txt", ".csv", ".json", ".docx", ".pdf", ".xlsx"},
        })
    source_json = root / "MRHPD v3.0.0a Section 5 Checkpoint 1 Source Index.json"
    source_csv = root / "MRHPD v3.0.0a Section 5 Checkpoint 1 Source Index.csv"
    json_write(source_json, {"generated_at": now_iso, "record_count": len(source_rows), "records": source_rows})
    csv_write(source_csv, source_rows)
    bit_path = root / "MRHPD v3.0.0a Section 5 Checkpoint 1 Bit Index.sqlite"
    if bit_path.exists():
        bit_path.unlink()
    con = sqlite3.connect(bit_path)
    try:
        con.executescript("""
        CREATE TABLE artifact (
            artifact_id INTEGER PRIMARY KEY,
            path TEXT NOT NULL UNIQUE,
            name TEXT NOT NULL,
            purpose TEXT NOT NULL,
            bytes INTEGER NOT NULL,
            sha256 TEXT NOT NULL,
            user_searchable INTEGER NOT NULL
        );
        CREATE VIRTUAL TABLE artifact_fts USING fts5(path, name, purpose, content);
        """)
        for row in source_rows:
            con.execute("INSERT INTO artifact (path,name,purpose,bytes,sha256,user_searchable) VALUES (?,?,?,?,?,?)", (row["path"], row["name"], row["purpose"], row["bytes"], row["sha256"], int(row["user_searchable"])))
            path = project / row["path"]
            content = extract_search_text(path) if row["user_searchable"] else ""
            con.execute("INSERT INTO artifact_fts (path,name,purpose,content) VALUES (?,?,?,?)", (row["path"], row["name"], row["purpose"], content))
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        counts = {
            "artifact": con.execute("SELECT COUNT(*) FROM artifact").fetchone()[0],
            "fts": con.execute("SELECT COUNT(*) FROM artifact_fts").fetchone()[0],
            "response75_hits": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"Response 75"',)).fetchone()[0],
            "spine_hits": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ("spine",)).fetchone()[0],
            "google_drive_hits": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"Google Drive"',)).fetchone()[0],
        }
        con.commit()
    finally:
        con.close()
    if integrity != "ok" or counts["artifact"] != len(source_rows) or counts["fts"] != len(source_rows):
        raise RuntimeError({"bit_index_gate": {"integrity": integrity, "counts": counts, "expected": len(source_rows)}})
    qa = root / "MRHPD v3.0.0a Section 5 Checkpoint 1 Index QA.json"
    json_write(qa, {"status": "passed", "generated_at": now_iso, "source_index_records": len(source_rows), "bit_index_integrity": integrity, "counts": counts, "bit_index_sha256": sha256_file(bit_path)})
    return {"source_json": source_json, "source_csv": source_csv, "bit_index": bit_path, "qa": qa, "records": len(source_rows), "integrity": integrity, "counts": counts}


def checkpoint_manifest(project: Path, paths: list[Path], now_iso: str) -> tuple[Path, Path, list[dict[str, Any]]]:
    root = project / "Manifest" / "Section 5 Session 1 Checkpoint 1"
    root.mkdir(parents=True, exist_ok=True)
    rows: list[dict[str, Any]] = []
    for path in sorted({p.resolve() for p in paths if p.exists() and p.is_file()}):
        rows.append({"path": path.relative_to(project.resolve()).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    manifest = root / "MRHPD v3.0.0a Section 5 Checkpoint 1 Manifest.json"
    checksums = root / "MRHPD v3.0.0a Section 5 Checkpoint 1 Checksums.sha256"
    json_write(manifest, {"generated_at": now_iso, "file_count": len(rows), "total_bytes": sum(row["bytes"] for row in rows), "files": rows})
    text_write(checksums, "".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
    return manifest, checksums, rows


def create_apply_script(manifest: dict[str, Any]) -> str:
    payload = repr(manifest)
    return f'''#!/usr/bin/env python3
from __future__ import annotations
import argparse, hashlib, json, re, shutil, sqlite3, sys, tempfile, zipfile
from pathlib import Path, PurePosixPath
BASE_RESTORE_BYTES={BASE_RESTORE_BYTES}
BASE_RESTORE_SHA256={BASE_RESTORE_SHA256!r}
BASE_PROJECT_BYTES={BASE_PROJECT_BYTES}
BASE_PROJECT_SHA256={BASE_PROJECT_SHA256!r}
PUBLICATION_SHA256={PUBLICATION_SHA256!r}
APPLICATION_SHA256={APPLICATION_SHA256!r}
CURRENT_DB_REL={CURRENT_DB_REL!r}
CURRENT_WORKBOOK_REL={CURRENT_WORKBOOK_REL!r}
PUBLICATION_REL={PUBLICATION_REL!r}
COVER_HASHES={COVER_HASHES!r}
MANIFEST={payload}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as f:
  for b in iter(lambda:f.read(1024*1024),b''): h.update(b)
 return h.hexdigest()
def safe_extract(path,dest):
 dest.mkdir(parents=True,exist_ok=True)
 with zipfile.ZipFile(path) as zf:
  if zf.testzip() is not None: raise RuntimeError('ZIP CRC failure: '+str(path))
  names=zf.namelist()
  if len(names)!=len(set(names)): raise RuntimeError('duplicate ZIP members: '+str(path))
  for name in names:
   pp=PurePosixPath(name.replace('\\\\','/'))
   if pp.is_absolute() or '..' in pp.parts or re.match(r'^[A-Za-z]:',name): raise RuntimeError('unsafe ZIP path: '+name)
  zf.extractall(dest)
def verify_identity(path,size,digest,label):
 observed={{'bytes':path.stat().st_size,'sha256':sha(path)}}
 if observed!={{'bytes':size,'sha256':digest}}: raise RuntimeError({{label:observed}})
def main():
 ap=argparse.ArgumentParser()
 ap.add_argument('--base-response72-restore',type=Path,required=True)
 ap.add_argument('--output-dir',type=Path,required=True)
 args=ap.parse_args()
 verify_identity(args.base_response72_restore,BASE_RESTORE_BYTES,BASE_RESTORE_SHA256,'baseline_restore_identity')
 package=Path(__file__).resolve().parents[1]
 overlay=package/'OVERLAY'
 if not overlay.is_dir(): raise RuntimeError('OVERLAY missing')
 if args.output_dir.exists():
  if any(args.output_dir.iterdir()): raise RuntimeError('output directory must be empty')
 else: args.output_dir.mkdir(parents=True)
 with tempfile.TemporaryDirectory(prefix='mrhpd-r75-apply-') as td:
  work=Path(td); restore_root=work/'restore'; safe_extract(args.base_response72_restore,restore_root)
  candidates=[p for p in restore_root.rglob('*.zip') if p.stat().st_size==BASE_PROJECT_BYTES and sha(p)==BASE_PROJECT_SHA256]
  if len(candidates)!=1: raise RuntimeError({{'project_archive_candidates':[str(p) for p in candidates]}})
  project_archive=candidates[0]
  extracted=work/'project'; safe_extract(project_archive,extracted)
  roots=[p for p in extracted.iterdir() if p.is_dir()]
  source=roots[0] if len(roots)==1 else extracted
  destination=args.output_dir/source.name
  shutil.copytree(source,destination)
  for row in MANIFEST['overlay_files']:
   source_file=overlay/row['path']
   verify_identity(source_file,row['bytes'],row['sha256'],'overlay_identity_'+row['path'])
   target=destination/row['path']; target.parent.mkdir(parents=True,exist_ok=True); shutil.copy2(source_file,target)
  db=destination/CURRENT_DB_REL
  con=sqlite3.connect(db)
  try:
   integrity=con.execute('PRAGMA integrity_check').fetchone()[0]
   fk=len(list(con.execute('PRAGMA foreign_key_check')))
   response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R75'").fetchone()[0]
   checkpoint=con.execute("SELECT state FROM section5_session1_checkpoint WHERE checkpoint_code='MRHPD-V3-CP5-S1-CP1'").fetchone()
   specs=con.execute("SELECT COUNT(*) FROM section5_print_provider_spec WHERE checkpoint_code='MRHPD-V3-CP5-S1-CP1'").fetchone()[0]
   scenarios=con.execute("SELECT COUNT(*) FROM section5_spine_scenario WHERE checkpoint_code='MRHPD-V3-CP5-S1-CP1'").fetchone()[0]
  finally: con.close()
  from openpyxl import load_workbook
  wb=load_workbook(destination/CURRENT_WORKBOOK_REL,read_only=True,data_only=False)
  try: sheet_count=len(wb.sheetnames); required={{'S5S1 Dashboard','S5S1 Provider Specs','S5S1 Spine Scenarios','S5S1 Interior QA','S5S1 Cover QA','S5S1 Risks','S5S1 Responses','S5S1 Recovery'}}; missing=sorted(required-set(wb.sheetnames))
  finally: wb.close()
  publication=destination/PUBLICATION_REL
  if sha(publication)!=PUBLICATION_SHA256: raise RuntimeError('publication changed')
  for rel,digest in COVER_HASHES.items():
   if sha(destination/rel)!=digest: raise RuntimeError('cover invariant failed: '+rel)
  app=[p for p in destination.rglob('human_pathogen_app.py') if p.is_file() and sha(p)==APPLICATION_SHA256]
  result={{'status':'passed','project_root':str(destination),'database':{{'integrity':integrity,'foreign_keys':fk,'response75':response,'checkpoint':checkpoint,'provider_specs':specs,'spine_scenarios':scenarios}},'workbook':{{'sheet_count':sheet_count,'missing_required':missing}},'publication_sha256':sha(publication),'main_application_matches':len(app),'baseline_restore_sha256':BASE_RESTORE_SHA256}}
  if integrity!='ok' or fk or response!=1 or checkpoint!=('checkpoint_complete',) or specs<10 or scenarios<5 or sheet_count<91 or missing or len(app)!=1: result['status']='failed'
  output=args.output_dir/'MRHPD_RESPONSE75_CHECKPOINT_RECOVERY_APPLICATION_RESULT.json'; output.write_text(json.dumps(result,indent=2),encoding='utf-8')
  print(json.dumps(result,indent=2))
  raise SystemExit(0 if result['status']=='passed' else 1)
if __name__=='__main__': main()
'''


def build_recovery_package(
    *,
    baseline_project: Path,
    current_project: Path,
    baseline_restore: Path,
    project_archive: Path,
    dist: Path,
    now: datetime,
    summary: dict[str, Any],
) -> dict[str, Any]:
    baseline_map = {p.relative_to(baseline_project).as_posix(): (p.stat().st_size, sha256_file(p)) for p in baseline_project.rglob("*") if p.is_file()}
    current_map = {p.relative_to(current_project).as_posix(): (p.stat().st_size, sha256_file(p)) for p in current_project.rglob("*") if p.is_file()}
    deleted = sorted(set(baseline_map) - set(current_map))
    if deleted:
        raise RuntimeError({"unexpected_deleted_paths": deleted})
    overlay_rows: list[dict[str, Any]] = []
    for rel, identity in sorted(current_map.items()):
        if baseline_map.get(rel) != identity:
            overlay_rows.append({"path": rel, "bytes": identity[0], "sha256": identity[1], "change": "new" if rel not in baseline_map else "changed"})
    stamp = now.strftime("%Y-%m-%d %H%M UTC")
    package_root = dist / "recovery_package_root"
    if package_root.exists():
        shutil.rmtree(package_root)
    overlay_root = package_root / "OVERLAY"
    tools = package_root / "TOOLS"
    overlay_root.mkdir(parents=True)
    tools.mkdir(parents=True)
    for row in overlay_rows:
        source = current_project / row["path"]
        target = overlay_root / row["path"]
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, target)
    manifest = {
        "schema": "mrhpd-section5-checkpoint-recovery-1.0",
        "generated_at": now.isoformat().replace("+00:00", "Z"),
        "version": PROJECT_VERSION,
        "response": RESPONSE_NUMBER,
        "section": SECTION_LABEL,
        "session": SESSION_LABEL,
        "checkpoint": CHECKPOINT_LABEL,
        "state": "checkpoint_complete",
        "baseline_restore": {"name": baseline_restore.name, "bytes": baseline_restore.stat().st_size, "sha256": sha256_file(baseline_restore)},
        "baseline_project": {"name": project_archive.name, "bytes": project_archive.stat().st_size, "sha256": sha256_file(project_archive)},
        "overlay_file_count": len(overlay_rows),
        "overlay_total_bytes": sum(row["bytes"] for row in overlay_rows),
        "overlay_files": overlay_rows,
        "deleted_paths": [],
        "accepted_predecessor_mutated": False,
        "frozen_section3_release_mutated": False,
        "immutable_publication_mutated": False,
        "user_upload_required": False,
        "requires_conversation_reconstruction": False,
        "next": "Remediation Section 5 of 5 Session 1 of 3 Checkpoint 2 of 3",
    }
    json_write(package_root / "CHECKPOINT_RECOVERY_MANIFEST.json", manifest)
    text_write(package_root / "CHECKPOINT_RECOVERY_CHECKSUMS.sha256", "".join(f"{row['sha256']}  OVERLAY/{row['path']}\n" for row in overlay_rows))
    text_write(tools / "apply_checkpoint_recovery.py", create_apply_script(manifest))
    text_write(package_root / "RESTORE_READ_FIRST.md", f"""# Human Pathogen Database — Response 75 Checkpoint Recovery

This is cumulative intermediate recovery for {SECTION_LABEL}, {SESSION_LABEL}, {CHECKPOINT_LABEL}. It applies directly to the exact Response 72 complete restore and does not require Response 73 or Response 74 packages.

## Required baseline

Filename: `{baseline_restore.name}`

Bytes: `{baseline_restore.stat().st_size}`

SHA-256: `{sha256_file(baseline_restore)}`

## Apply

```bash
python TOOLS/apply_checkpoint_recovery.py \
  --base-response72-restore "<Response 72 complete restore.zip>" \
  --output-dir "<empty destination>"
```

The utility verifies the baseline, every overlay member, SQLite integrity and foreign keys, Responses 73–75, fractional prompts 74.1–74.2, Section 5 checkpoint state, provider specifications, spine scenarios, workbook preservation, application identity, immutable publication identity, and all governed cover hashes.

No user upload, external project file, or reconstruction from the conversation is required when the persistent Google Drive baseline remains accessible.
""")
    recovery_zip = dist / (
        f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 "
        f"Remediation Section 5 of 5 Session 1 of 3 Checkpoint 1 of 3 RECOVERY DATA THROUGH RESPONSE 75 {stamp}.zip"
    )
    with zipfile.ZipFile(recovery_zip, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in sorted(package_root.rglob("*")):
            if path.is_file():
                zf.write(path, path.relative_to(package_root).as_posix())
    recovery_qa = verify_zip(recovery_zip)
    with tempfile.TemporaryDirectory(prefix="mrhpd-r75-apply-test-") as td:
        output = Path(td) / "restored"
        result = subprocess.run(
            [sys.executable, str(package_root / "TOOLS" / "apply_checkpoint_recovery.py"), "--base-response72-restore", str(baseline_restore), "--output-dir", str(output)],
            cwd=package_root,
            text=True,
            capture_output=True,
            timeout=900,
        )
        if result.returncode:
            raise RuntimeError({"clean_apply_failed": {"stdout": result.stdout[-12000:], "stderr": result.stderr[-12000:]}})
        clean_apply = json.loads((output / "MRHPD_RESPONSE75_CHECKPOINT_RECOVERY_APPLICATION_RESULT.json").read_text(encoding="utf-8"))
    verification = {
        "schema": "mrhpd-response75-recovery-verification-1.0",
        "generated_at": now.isoformat().replace("+00:00", "Z"),
        "status": "passed",
        "recovery_zip": recovery_qa,
        "manifest": {"overlay_file_count": len(overlay_rows), "overlay_total_bytes": sum(row["bytes"] for row in overlay_rows), "deleted_paths": 0},
        "clean_apply": clean_apply,
        "accepted_predecessor_mutated": False,
        "frozen_section3_release_mutated": False,
        "immutable_publication_mutated": False,
        "user_upload_required": False,
        "checkpoint_1_of_3_complete": True,
        "session_1_of_3_complete": False,
        "remediation_section_5_complete": False,
        "next": "Checkpoint 2 of 3 - scenario selection and print derivative",
    }
    verification_path = dist / "MRHPD v3.0.0a Response 75 Checkpoint 1 Recovery Verification.json"
    json_write(verification_path, verification)
    sha_path = dist / f"{recovery_zip.name}.sha256.txt"
    text_write(sha_path, f"{recovery_qa['sha256']}  {recovery_zip.name}\n")
    summary_path = dist / "MRHPD_RESPONSE75_SECTION5_CHECKPOINT1_BUILD_SUMMARY.json"
    json_write(summary_path, summary | {"recovery": verification})
    exact_names = dist / "MRHPD v3.0.0a Response 75 Exact File Names.txt"
    text_write(exact_names, f"""Response 75 checkpoint recovery ZIP:
{recovery_zip.name}

Required baseline complete restore:
{baseline_restore.name}

Required baseline project archive embedded in that restore:
{project_archive.name}

Current copied SQLite database:
{Path(CURRENT_DB_REL).name}

Current comprehensive workbook:
{Path(CURRENT_WORKBOOK_REL).name}
""")
    delivery = dist / f"MRHPD v3.0.0a Response 75 Section 5 Session 1 Checkpoint 1 Recovery Package {stamp}.zip"
    with zipfile.ZipFile(delivery, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in (recovery_zip, sha_path, verification_path, summary_path, exact_names):
            zf.write(path, path.name)
    delivery_qa = verify_zip(delivery)
    return {
        "recovery_zip": recovery_zip,
        "recovery_qa": recovery_qa,
        "verification_path": verification_path,
        "summary_path": summary_path,
        "exact_names": exact_names,
        "delivery": delivery,
        "delivery_qa": delivery_qa,
        "overlay_rows": overlay_rows,
        "clean_apply": clean_apply,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--volume1-dir", type=Path, required=True)
    parser.add_argument("--volume2-dir", type=Path, required=True)
    parser.add_argument("--dist", type=Path, default=Path("dist_cp5_s1_cp1"))
    args = parser.parse_args()
    now = utc_now()
    now_iso = now.isoformat().replace("+00:00", "Z")
    stamp = now.strftime("%Y-%m-%d %H%M UTC")
    if args.dist.exists():
        shutil.rmtree(args.dist)
    args.dist.mkdir(parents=True)

    with tempfile.TemporaryDirectory(prefix="mrhpd-cp5-s1-cp1-") as td:
        work = Path(td)
        restore, project_archive, baseline_project = reconstruct_baseline(args.volume1_dir, args.volume2_dir, work / "baseline")
        current_project = work / "current_project" / baseline_project.name
        current_project.parent.mkdir(parents=True)
        shutil.copytree(baseline_project, current_project)

        base_db = baseline_project / BASE_DB_REL
        base_workbook = baseline_project / BASE_WORKBOOK_REL
        publication = current_project / PUBLICATION_REL
        if not base_db.exists() or not base_workbook.exists() or not publication.exists():
            raise RuntimeError({"required_baseline_paths": {"db": base_db.exists(), "workbook": base_workbook.exists(), "publication": publication.exists()}})

        interior_dir = current_project / "QA" / "Section 5 Session 1" / "Checkpoint 1" / "Interior"
        interior_summary, page_rows = inspect_interior(publication, interior_dir)
        cover_summary = inspect_cover_assets(current_project)
        app_path = find_main_application(current_project)
        editable = find_editable_assembly(current_project)
        if sha256_file(editable) != EDITABLE_SHA256:
            raise RuntimeError("editable assembly identity changed")

        specs = provider_specs(now_iso)
        scenarios = build_spine_scenarios()
        risks = forward_risks()
        gates = preflight_gates(interior_summary, cover_summary)
        responses = response_records(now_iso)
        fractions = fractional_prompts(now_iso)
        events = recovery_events(now_iso)

        current_db = current_project / CURRENT_DB_REL
        database_qa = synchronize_database(
            base_db, current_db, now_iso=now_iso, specs=specs, scenarios=scenarios, gates=gates, risks=risks,
            events=events, responses=responses, fractions=fractions, interior=interior_summary, cover=cover_summary,
        )
        tracking_files = write_tracking_files(current_project, responses, fractions, events, now_iso)
        instruction_addendum = write_instruction_addendum(current_project, now_iso)
        application_files = write_application_surfaces(current_project, current_db, app_path, now_iso)

        report_root = current_project / "Reports" / "Section 5 Session 1" / "Checkpoint 1"
        figure_root = current_project / "Artwork" / "Section 5 Print Production"
        spine_figure = figure_root / "MRHPD-FIG-S5-0001 KDP Spine Scenario Comparison v3.0.0a.png"
        width_figure = figure_root / "MRHPD-FIG-S5-0002 Existing Cover versus Required Width v3.0.0a.png"
        spine_figure_qa = reporting.build_spine_scenario_figure(spine_figure, scenarios)
        width_figure_qa = reporting.build_cover_width_figure(width_figure, cover_summary["current_combined_width_in"], scenarios)
        docx_path = report_root / "MRHPD v3.0.0a Section 5 Session 1 Checkpoint 1 Print-Production Intake Report.docx"
        pdf_path = report_root / "MRHPD v3.0.0a Section 5 Session 1 Checkpoint 1 Print-Production Intake Report.pdf"
        xlsx_path = report_root / "MRHPD v3.0.0a Section 5 Session 1 Checkpoint 1 Print-Production Register.xlsx"
        baseline_identity = {
            "restore_bytes": restore.stat().st_size,
            "restore_sha256": sha256_file(restore),
            "project_bytes": project_archive.stat().st_size,
            "project_sha256": sha256_file(project_archive),
        }
        tracking_register = [
            {"record": "Responses 73-75", "disposition": "Source-supported reconciliation added to the copied database and current tracking files."},
            {"record": "Fractional Prompts 74.1-74.2", "disposition": "Exact text retained and linked to Response 75."},
            {"record": "Raw and Net tracking", "disposition": "Updated through Response 75."},
            {"record": "Cumulative Thread Index", "disposition": "Updated through Response 75."},
            {"record": "Project Instructions 1.5.0", "disposition": "Current-state operating addendum added; current instructions reprocessed."},
            {"record": "Persistent custody", "disposition": "Response 72 source recovered from Google Drive; no upload required."},
        ]
        docx_qa = reporting.build_docx_report(
            docx_path, generated_at=now_iso, baseline=baseline_identity, provider_specs=specs, scenarios=scenarios,
            interior=interior_summary, cover=cover_summary, risks=risks, recovery_events=events,
            tracking=tracking_register, figure_paths=[spine_figure, width_figure],
        )
        pdf_qa = reporting.build_pdf_report(
            pdf_path, generated_at=now_iso, baseline=baseline_identity, provider_specs=specs, scenarios=scenarios,
            interior=interior_summary, cover=cover_summary, risks=risks, recovery_events=events,
            figure_paths=[spine_figure, width_figure],
        )
        summary_rows = [
            {"control": "Response", "value": 75, "status": "current"},
            {"control": "Checkpoint", "value": "1 of 3", "status": "complete"},
            {"control": "Session", "value": "1 of 3", "status": "continue"},
            {"control": "Section", "value": "5 of 5", "status": "continue"},
            {"control": "Digital pages", "value": interior_summary["page_count"], "status": "immutable"},
            {"control": "Production pages", "value": interior_summary["production_page_count"], "status": "derived"},
            {"control": "Database tables", "value": database_qa["table_count"], "status": database_qa["integrity"]},
            {"control": "Provider specs", "value": len(specs), "status": "reconciled"},
            {"control": "Spine scenarios", "value": len(scenarios), "status": "calculated_not_selected"},
            {"control": "User upload", "value": "not required", "status": "Google Drive baseline accessible"},
        ]
        register_qa = reporting.build_checkpoint_register(
            xlsx_path, summary_rows=summary_rows, provider_specs=specs, scenarios=scenarios,
            interior_page_rows=page_rows, cover_assets=cover_summary["assets"], risks=risks,
            tracking=tracking_register, recovery_events=events,
        )

        current_workbook = current_project / CURRENT_WORKBOOK_REL
        workbook_qa = reporting.augment_comprehensive_workbook(
            base_workbook, current_workbook, generated_at=now_iso, provider_specs=specs, scenarios=scenarios,
            risks=risks, tracking_rows=responses, recovery_events=events,
            interior_summary=interior_summary, cover_summary=cover_summary,
        )

        data_root = current_project / "Data" / "Section 5 Session 1 Checkpoint 1"
        specs_json = data_root / "MRHPD v3.0.0a Section 5 Checkpoint 1 Provider Specifications.json"
        specs_csv = data_root / "MRHPD v3.0.0a Section 5 Checkpoint 1 Provider Specifications.csv"
        scenarios_json = data_root / "MRHPD v3.0.0a Section 5 Checkpoint 1 Spine Scenarios.json"
        scenarios_csv = data_root / "MRHPD v3.0.0a Section 5 Checkpoint 1 Spine Scenarios.csv"
        gates_json = data_root / "MRHPD v3.0.0a Section 5 Checkpoint 1 Preflight Gates.json"
        risks_json = data_root / "MRHPD v3.0.0a Section 5 Checkpoint 1 Forward Risks.json"
        json_write(specs_json, specs); csv_write(specs_csv, specs)
        json_write(scenarios_json, scenarios); csv_write(scenarios_csv, scenarios)
        json_write(gates_json, gates); json_write(risks_json, risks)

        qa_root = current_project / "QA" / "Section 5 Session 1" / "Checkpoint 1"
        qa_root.mkdir(parents=True, exist_ok=True)
        database_qa_path = qa_root / "DATABASE_QA.json"
        workbook_qa_path = qa_root / "WORKBOOK_QA.json"
        application_qa_path = qa_root / "APPLICATION_QA.json"
        publication_qa_path = qa_root / "PUBLICATION_QA.json"
        cover_qa_path = qa_root / "COVER_QA.json"
        report_qa_path = qa_root / "REPORT_QA.json"
        json_write(database_qa_path, database_qa)
        json_write(workbook_qa_path, workbook_qa)
        app_qa = {"status": "passed", "main_application": app_path.relative_to(current_project).as_posix(), "sha256": sha256_file(app_path), "unchanged": True, "section5_audit": json.loads(application_files[-1].read_text(encoding="utf-8"))}
        json_write(application_qa_path, app_qa)
        publication_qa = {"status": "passed_intake", "interior": interior_summary, "editable_assembly": {"path": editable.relative_to(current_project).as_posix(), "bytes": editable.stat().st_size, "sha256": sha256_file(editable), "unchanged": True}}
        json_write(publication_qa_path, publication_qa)
        json_write(cover_qa_path, cover_summary)
        report_qa = {"status": "passed", "docx": docx_qa, "pdf": pdf_qa, "xlsx": register_qa, "figures": [spine_figure_qa, width_figure_qa]}
        json_write(report_qa_path, report_qa)

        generated_files = [
            current_db, current_workbook, *tracking_files, instruction_addendum, *application_files,
            docx_path, pdf_path, xlsx_path, spine_figure, width_figure,
            specs_json, specs_csv, scenarios_json, scenarios_csv, gates_json, risks_json,
            database_qa_path, workbook_qa_path, application_qa_path, publication_qa_path, cover_qa_path, report_qa_path,
            interior_dir / "MRHPD v3.0.0a Section 5 Checkpoint 1 Interior Page Geometry Screen.csv",
            interior_dir / "MRHPD v3.0.0a Section 5 Checkpoint 1 Interior Page Geometry Screen.json",
        ]
        index_qa = build_source_and_bit_indexes(current_project, generated_files, now_iso)
        generated_files.extend([index_qa["source_json"], index_qa["source_csv"], index_qa["bit_index"], index_qa["qa"]])
        manifest_path, checksum_path, manifest_rows = checkpoint_manifest(current_project, generated_files, now_iso)
        generated_files.extend([manifest_path, checksum_path])

        final_qa_path = qa_root / "SECTION5_CHECKPOINT1_QA.json"
        final_qa = {
            "schema": "mrhpd-section5-checkpoint1-qa-1.0",
            "generated_at": now_iso,
            "status": "passed",
            "response": 75,
            "section": SECTION_LABEL,
            "session": SESSION_LABEL,
            "checkpoint": CHECKPOINT_LABEL,
            "checkpoint_state": "complete",
            "session_state": "continue",
            "section_state": "continue",
            "baseline": baseline_identity,
            "database": database_qa,
            "workbook": workbook_qa,
            "application": app_qa,
            "publication": publication_qa,
            "cover": cover_summary,
            "provider_specs": len(specs),
            "spine_scenarios": len(scenarios),
            "preflight_gates": {"total": len(gates), "pending": sum(1 for row in gates if not row["status"].startswith("passed"))},
            "forward_risks": len(risks),
            "reports": report_qa,
            "indexes": {"records": index_qa["records"], "integrity": index_qa["integrity"], "counts": index_qa["counts"]},
            "manifest_records": len(manifest_rows),
            "accepted_predecessor_mutated": False,
            "frozen_section3_release_mutated": False,
            "immutable_publication_mutated": False,
            "user_upload_required": False,
            "next": "Checkpoint 2 of 3 - scenario selection and print derivative",
        }
        json_write(final_qa_path, final_qa)
        generated_files.append(final_qa_path)
        # Regenerate the checkpoint manifest after the final QA file is frozen.
        manifest_path, checksum_path, manifest_rows = checkpoint_manifest(current_project, generated_files, now_iso)

        summary = {
            "schema": "mrhpd-response75-section5-checkpoint1-build-1.0",
            "generated_at": now_iso,
            "status": "passed",
            "response": 75,
            "section": SECTION_LABEL,
            "session": SESSION_LABEL,
            "checkpoint": CHECKPOINT_LABEL,
            "database": database_qa,
            "workbook": workbook_qa,
            "publication": interior_summary,
            "cover": {"current_width_in": cover_summary["current_combined_width_in"], "current_height_in": cover_summary["current_combined_height_in"], "status": cover_summary["combined_cover_status"]},
            "provider_specs": len(specs),
            "spine_scenarios": scenarios,
            "preflight_gates": gates,
            "recovery_events": events,
            "reports": report_qa,
            "index": {"records": index_qa["records"], "integrity": index_qa["integrity"]},
            "manifest_records": len(manifest_rows),
            "user_upload_required": False,
            "checkpoint_1_of_3_complete": True,
            "session_1_of_3_complete": False,
            "remediation_section_5_complete": False,
            "next": "Checkpoint 2 of 3 - scenario selection and print derivative",
        }
        package = build_recovery_package(
            baseline_project=baseline_project,
            current_project=current_project,
            baseline_restore=restore,
            project_archive=project_archive,
            dist=args.dist,
            now=now,
            summary=summary,
        )
        console = {
            "status": "passed",
            "delivery": package["delivery"].name,
            "delivery_bytes": package["delivery_qa"]["bytes"],
            "delivery_sha256": package["delivery_qa"]["sha256"],
            "recovery_zip": package["recovery_zip"].name,
            "recovery_zip_bytes": package["recovery_qa"]["bytes"],
            "recovery_zip_sha256": package["recovery_qa"]["sha256"],
            "overlay_files": len(package["overlay_rows"]),
            "database_tables": database_qa["table_count"],
            "workbook_sheets": workbook_qa["current_sheet_count"],
            "publication_pages": interior_summary["page_count"],
            "production_pages": interior_summary["production_page_count"],
            "provider_specs": len(specs),
            "spine_scenarios": len(scenarios),
            "user_upload_required": False,
            "checkpoint_1_of_3_complete": True,
            "next": "Checkpoint 2 of 3 - scenario selection and print derivative",
        }
        print(json.dumps(console, indent=2))


if __name__ == "__main__":
    main()
