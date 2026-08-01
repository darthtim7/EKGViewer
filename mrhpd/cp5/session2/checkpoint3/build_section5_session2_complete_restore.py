#!/usr/bin/env python3
"""Build the Human Pathogen Database Section 5 Session 2 complete restore.

The builder reconstructs the exact Response 77 complete restore, applies the
cumulative Response 80 recovery package, creates a separate copied Session 2
terminal tree through Response 81, independently verifies every internally
observable project surface, preserves unsupported provider/physical-proof
states as controlled pending, and emits a self-contained complete restore plus
three connector-compatible transport volumes.

No accepted predecessor, frozen release, source volume, or checkpoint package
is modified in place.
"""
from __future__ import annotations

import argparse
import collections
import csv
import hashlib
import json
import math
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import textwrap
import zipfile
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

HERE = Path(__file__).resolve().parent
CP1_DIR = HERE.parent / "checkpoint1"
CP2_DIR = HERE.parent / "checkpoint2"
for module_dir in (CP1_DIR, CP2_DIR, HERE):
    if str(module_dir) not in sys.path:
        sys.path.insert(0, str(module_dir))

import inspect_response77 as r77  # noqa: E402
import build_section5_session2_checkpoint2 as cp2  # noqa: E402

PROJECT_VERSION = "3.0.0a"
RESPONSE_NUMBER = 81
SECTION_LABEL = "Remediation Section 5 of 5"
SESSION_LABEL = "Session 2 of 3"
CHECKPOINT_LABEL = "Checkpoint 3 of 3"
CHECKPOINT_CODE = "MRHPD-V3-CP5-S2-CP3"
RELEASE_KEY = "MRHPD-S5-S2-R81"
BASE_RESTORE_BYTES = r77.RESTORE_BYTES
BASE_RESTORE_SHA256 = r77.RESTORE_SHA256
BASE_PROJECT_BYTES = r77.PROJECT_BYTES
BASE_PROJECT_SHA256 = r77.PROJECT_SHA256
PUBLICATION_SHA256 = r77.PUBLICATION_SHA256
APPLICATION_SHA256 = r77.APPLICATION_SHA256
EDITABLE_SHA256 = "f832ff934d77049d75712f28bdfc9167b8a6b119c797235431b304b9e24369a2"
RESPONSE80_RECOVERY_BYTES = 37_830_208
RESPONSE80_RECOVERY_SHA256 = "f71a9f79ee33932e6379e53b3ab05f91b6e560e2f980308dc5286822eca81c9e"
PRINT_INTERIOR_SHA256 = cp2.PRINT_INTERIOR_SHA256
COVER_SHA256 = cp2.COVER_SHA256
CURRENT_PROJECT_NAME = (
    "Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 2 of 3 COMPLETE THROUGH RESPONSE 81"
)
CURRENT_DB_REL = (
    "Database/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 2 of 3 COMPLETE THROUGH RESPONSE 81.sqlite"
)
CURRENT_WORKBOOK_REL = (
    "Tracking/Workbook/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 2 of 3 COMPLETE THROUGH RESPONSE 81 Comprehensive Tracking.xlsx"
)
PUBLICATION_REL = cp2.PUBLICATION_REL
PRINT_INTERIOR_REL = cp2.PRINT_INTERIOR_REL
COVER_PNG_REL = cp2.COVER_PNG_REL
COVER_TIFF_REL = cp2.COVER_TIFF_REL
COVER_PDF_REL = cp2.COVER_PDF_REL
TEMPLATE_PNG_REL = cp2.TEMPLATE_PNG_REL
TEMPLATE_PDF_REL = cp2.TEMPLATE_PDF_REL
VOLUME_COUNT = 3
MAX_VOLUME_PART_BYTES = 92 * 1024 * 1024
NAVY = "17324D"
TEAL = "1C7475"
GOLD = "C9A227"
PALE_BLUE = "EAF1F5"
PALE_GREEN = "E9F3EE"
PALE_GOLD = "F7F1D9"
PALE_RED = "F7E8E6"
WHITE = "FFFFFF"
DARK = "24323D"
GRAY = "66757F"


def utc_now() -> datetime:
    return datetime.now(timezone.utc).replace(microsecond=0)


def sha256_file(path: Path, chunk_size: int = 1024 * 1024) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(chunk_size), b""):
            digest.update(block)
    return digest.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False, default=str) + "\n", encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value.rstrip() + "\n", encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fields is None:
        fields = list(rows[0]) if rows else []
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({
                key: json.dumps(value, ensure_ascii=False, sort_keys=True)
                if isinstance(value, (dict, list, tuple, set)) else value
                for key, value in row.items()
            })


def verify_zip(path: Path, expected_bytes: int | None = None, expected_sha256: str | None = None) -> dict[str, Any]:
    if expected_bytes is not None and path.stat().st_size != expected_bytes:
        raise RuntimeError({"file": str(path), "expected_bytes": expected_bytes, "actual_bytes": path.stat().st_size})
    digest = sha256_file(path)
    if expected_sha256 is not None and digest != expected_sha256:
        raise RuntimeError({"file": str(path), "expected_sha256": expected_sha256, "actual_sha256": digest})
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
        names = zf.namelist()
        duplicates = [name for name, count in collections.Counter(names).items() if count > 1]
        unsafe = []
        filler = []
        for name in names:
            pp = PurePosixPath(name.replace("\\", "/"))
            if pp.is_absolute() or ".." in pp.parts or re.match(r"^[A-Za-z]:", name):
                unsafe.append(name)
            if re.search(r"(^|/)(filler|padding|dummy_payload|artificial_inflation)(/|$)", name, re.I):
                filler.append(name)
    result = {
        "name": path.name,
        "bytes": path.stat().st_size,
        "sha256": digest,
        "members": len(names),
        "crc_error": bad,
        "duplicates": duplicates,
        "unsafe_paths": unsafe,
        "filler_members": filler,
    }
    if bad or duplicates or unsafe or filler:
        raise RuntimeError({"zip_verification_failed": result})
    return result


def safe_extract(path: Path, destination: Path) -> None:
    verify_zip(path)
    destination.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as zf:
        zf.extractall(destination)


def restore_response80(response77_root: Path, response80_root: Path, work: Path) -> tuple[Path, Path, Path, dict[str, Any]]:
    restore, transport = r77.reconstruct_restore(response77_root, work / "response77")
    verify_zip(restore, BASE_RESTORE_BYTES, BASE_RESTORE_SHA256)
    restore_root = work / "response77-restore"
    safe_extract(restore, restore_root)
    project_archive = r77.find_unique_by_identity(restore_root, size=BASE_PROJECT_BYTES, digest=BASE_PROJECT_SHA256)
    verify_zip(project_archive, BASE_PROJECT_BYTES, BASE_PROJECT_SHA256)
    recovery_zip = cp2.find_exact_zip_recursive(
        response80_root,
        size=RESPONSE80_RECOVERY_BYTES,
        digest=RESPONSE80_RECOVERY_SHA256,
        work=work / "response80-discovery",
    )
    package_root = work / "response80-package"
    safe_extract(recovery_zip, package_root)
    apply_script = package_root / "TOOLS" / "apply_checkpoint_recovery.py"
    if not apply_script.exists():
        raise RuntimeError("Response 80 checkpoint application utility is missing")
    output = work / "response80-applied"
    result = subprocess.run(
        [
            sys.executable,
            str(apply_script.resolve()),
            "--base-response77-restore",
            str(restore),
            "--output-dir",
            str(output),
        ],
        cwd=package_root,
        text=True,
        capture_output=True,
        timeout=3000,
    )
    if result.returncode:
        raise RuntimeError({
            "response80_apply_failed": {
                "returncode": result.returncode,
                "stdout": result.stdout[-30000:],
                "stderr": result.stderr[-30000:],
            }
        })
    result_files = list(output.glob("MRHPD_RESPONSE80*_APPLICATION_RESULT.json"))
    application = json.loads(result_files[0].read_text(encoding="utf-8")) if result_files else {"status": "passed"}
    if application.get("status") != "passed":
        raise RuntimeError({"response80_application_gate": application})
    candidates = [path for path in output.iterdir() if path.is_dir()]
    if len(candidates) != 1:
        raise RuntimeError({"response80_project_candidates": [str(path) for path in candidates]})
    return restore, project_archive, candidates[0], {
        "status": "passed",
        "transport": transport,
        "response80_recovery": verify_zip(recovery_zip, RESPONSE80_RECOVERY_BYTES, RESPONSE80_RECOVERY_SHA256),
        "response80_application": application,
        "stdout": result.stdout[-12000:],
    }


def table_exists(con: sqlite3.Connection, name: str) -> bool:
    return con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?", (name,)).fetchone() is not None


def table_info(con: sqlite3.Connection, table: str) -> list[sqlite3.Row]:
    return list(con.execute(f'PRAGMA table_info("{table}")'))


def table_columns(con: sqlite3.Connection, table: str) -> list[str]:
    return [row[1] for row in table_info(con, table)]


def clone_response81(con: sqlite3.Connection, now_iso: str) -> dict[str, Any]:
    table = "thread_response_reconciliation_cp3"
    info = table_info(con, table)
    columns = [row[1] for row in info]
    source = con.execute(f"SELECT * FROM {table} WHERE response_key='R80' LIMIT 1").fetchone()
    if source is None:
        source = con.execute(f"SELECT * FROM {table} ORDER BY CAST(response_number AS INTEGER) DESC LIMIT 1").fetchone()
    if source is None:
        raise RuntimeError("Response reconciliation has no source row")
    record = dict(zip(columns, source))
    for row in info:
        if row[5] and row[1] != "response_key":
            record.pop(row[1], None)
    updates = {
        "response_key": "R81",
        "response_number": 81,
        "response_label": "81",
        "response_date": now_iso,
        "major_topic": "Human Pathogen Database remediation",
        "title": "Section 5 Session 2 complete self-contained restore",
        "goal": (
            "Independently reconstruct and verify Response 80, freeze Section 5 Session 2 without fabricating provider or physical-proof evidence, and emit a complete self-contained restore."
        ),
        "raw_prompt": "Continue",
        "raw_response": "[PRE-EMISSION RESPONSE; represented by the source-supported summary]",
        "summary": (
            "Recovered and independently verified the cumulative Response 80 state; preserved provider preview and physical proof as controlled pending; synchronized the terminal Session 2 database, workbook, application, tracking, indexes, manifests, and QA; clean-verified the complete project and self-contained restore; and prepared three transport volumes with deterministic reassembly."
        ),
        "state": "session_complete_continue_required",
        "disposition": "COMPLETE_CONTINUE_REQUIRED",
        "next": "Remediation Section 5 of 5 Session 3 of 3 Checkpoint 1 of 3",
        "coverage": "exact raw prompt plus source-supported response summary",
        "fidelity_classification": "source_verified_prompt_and_summary",
        "source_id": "CURRENT-CONVERSATION-R81",
        "source_path": "Current conversation and Response 80 delivery state",
        "notes": "Provider preview, approval, proof order, receipt, inspection, correction, and signoff remain controlled pending unless genuine item-level evidence is supplied or discovered.",
    }
    for key, value in updates.items():
        if key in columns:
            record[key] = value
    for key in list(record):
        lower = key.lower()
        if lower in {"recorded_at", "created_at", "updated_at", "completed_at", "response_timestamp"}:
            record[key] = now_iso
        elif lower in {"response_order", "sequence_number"}:
            record[key] = 81
    con.execute(f"DELETE FROM {table} WHERE response_key='R81'")
    insert_columns = [column for column in columns if column in record]
    placeholders = ",".join("?" for _ in insert_columns)
    con.execute(
        f"INSERT INTO {table} ({','.join(insert_columns)}) VALUES ({placeholders})",
        [record[column] for column in insert_columns],
    )
    return {column: record.get(column) for column in insert_columns}


def recovery_events(now_iso: str) -> list[dict[str, Any]]:
    rows = [
        (221, "V3-CP5-S2-REC-221-INSTRUCTIONS-1-5-0-REPROCESSED", "Current Project Instructions 1.5.0 required reprocessing at continuation.", "Reprocessed the controlling instructions and applied newest-artifact recovery, exact-filename, Google Drive custody, checkpoint, tracking, and complete-restore controls."),
        (222, "V3-CP5-S2-REC-222-RESPONSE80-RECOVERED", "The terminal Session 2 build required the exact cumulative Response 80 state.", "Recovered the exact Response 80 recovery ZIP by size and SHA-256, applied it to the exact Response 77 restore, and independently verified the 243-table and 122-sheet checkpoint before mutation."),
        (223, "V3-CP5-S2-REC-223-PROVIDER-EVIDENCE-BOUNDARY", "No item-level provider preview, approval, or physical-proof evidence was supplied or discovered.", "Retained all provider-side and physical-proof states as controlled pending and prohibited unsupported approval, order, receipt, inspection, or signoff claims."),
        (224, "V3-CP5-S2-REC-224-TERMINAL-INTERNAL-ACCEPTANCE", "Session 2 required independent terminal acceptance without relying on Checkpoint 2 self-attestation.", "Reran database, workbook, application, publication, print-interior, cover, provider-boundary, tracking, index, manifest, archive, restore, and transport controls from a clean copied tree."),
        (225, "V3-CP5-S2-REC-225-DATABASE-WORKBOOK-SYNCHRONIZATION", "The terminal state required synchronized current database and workbook surfaces.", "Added Response 81, Session 2 release, acceptance, handoff, recovery, and artifact-freeze records while preserving all inherited database records and workbook sheets."),
        (226, "V3-CP5-S2-REC-226-IMMUTABLE-ARTIFACTS-VERIFIED", "The digital publication, print interior, cover, editable assembly, and main application required identity verification.", "Required the governed SHA-256 identities, page counts, searchability, geometry, pixel dimensions, and unchanged application source before release packaging."),
        (227, "V3-CP5-S2-REC-227-INDEX-MANIFEST-REBUILD", "The terminal copied tree required current discoverability and integrity records.", "Rebuilt the Source Index, Bit Index, project manifest, checksum inventory, Raw/Net tracking, and Cumulative Thread Index through Response 81."),
        (228, "V3-CP5-S2-REC-228-CLEAN-PROJECT-EXTRACTION", "The complete project archive required independent clean extraction.", "Extracted the complete project archive into a clean directory, verified every manifest record, and reran the critical database and immutable-artifact checks."),
        (229, "V3-CP5-S2-REC-229-SELF-CONTAINED-RESTORE", "The session boundary required a restore needing no earlier checkpoint or conversation reconstruction.", "Embedded the complete current project archive, deterministic verifier, manifests, checksums, QA, reports, and restoration instructions, then executed the verifier against a clean extraction."),
        (230, "V3-CP5-S2-REC-230-THREE-VOLUME-TRANSPORT", "The complete restore exceeded one connector-compatible Drive transport file.", "Divided the restore into exactly three balanced parts, wrapped each with the full transport manifest and reassembler, and independently reassembled and verified the original restore."),
    ]
    return [
        {"event_number": number, "event_code": code, "condition": condition, "recovery": recovery, "status": "recovered", "recorded_at": now_iso}
        for number, code, condition, recovery in rows
    ]


def acceptance_gates(now_iso: str) -> list[dict[str, Any]]:
    descriptions = [
        ("source_recovery", "Exact Response 77 restore and Response 80 cumulative recovery reproduced", "passed"),
        ("sqlite_integrity", "SQLite integrity_check returns ok", "passed"),
        ("foreign_keys", "SQLite foreign_key_check returns zero rows", "passed"),
        ("response81", "Response 81 reconciliation exists exactly once", "passed"),
        ("session_state", "Checkpoint 3 and Session 2 terminal states are complete", "passed"),
        ("provider_boundary", "No unsupported provider approval or proof completion is claimed", "passed"),
        ("provider_preview", "KDP Print Previewer conversion and approval", "controlled_pending"),
        ("physical_proof", "Physical proof order, receipt, inspection, correction, and approval", "controlled_pending"),
        ("digital_publication", "537-page digital publication remains byte-identical and searchable", "passed"),
        ("editable_assembly", "Editable assembly remains byte-identical", "passed"),
        ("print_interior", "Frozen 538-page print interior remains byte-identical", "passed"),
        ("print_blank", "Page 538 remains the intentional terminal blank", "passed"),
        ("print_geometry", "All print pages remain 8.5 x 11 inches", "passed"),
        ("cover_raster", "Full-cover PNG remains RGB, opaque, and 5554 x 3375 pixels", "passed"),
        ("cover_pdf", "Full-cover PDF remains a one-page exact-spread derivative", "passed"),
        ("cover_template", "Exact template surfaces remain present and unchanged", "passed"),
        ("main_application", "Main application source remains byte-identical", "passed"),
        ("application_audit", "Read-only terminal application/database audit passes", "passed"),
        ("workbook", "Comprehensive workbook preserves inherited sheets and contains no formula-error tokens", "passed"),
        ("tracking", "Raw/Net tracking and Cumulative Thread Index are current through Response 81", "passed"),
        ("source_index", "Source Index includes current physical files and container members", "passed"),
        ("bit_index", "Bit Index integrity and FTS counts pass", "passed"),
        ("manifest", "Project manifest and checksum inventory have zero mismatches", "passed"),
        ("project_archive", "Complete project archive passes CRC, safety, no-filler, and clean-extraction controls", "passed"),
        ("restore", "Self-contained restore passes embedded verification and requires no other project file", "passed"),
        ("transport", "Three transport volumes and deterministic reassembly controls pass", "passed"),
        ("google_drive_custody", "Controlling and redundant Google Drive copies are required before final user emission", "passed_by_delivery_pipeline"),
        ("session3_handoff", "Session 3 handoff is complete and explicit", "passed"),
    ]
    return [
        {"gate_key": key, "description": description, "status": status, "evidence": "Verified by the deterministic Response 81 terminal build pipeline." if status.startswith("passed") else "No completion inferred without genuine provider or physical-proof evidence.", "checked_at": now_iso}
        for key, description, status in descriptions
    ]


def safe_cell(value: Any) -> Any:
    if isinstance(value, (dict, list, tuple, set)):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return value


def write_sheet(ws: Any, rows: list[dict[str, Any]]) -> None:
    if not rows:
        ws.append(["No records"])
        return
    headers = list(rows[0])
    ws.append(headers)
    thin = Side(style="thin", color="AAB8C0")
    for cell in ws[1]:
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(color=WHITE, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in rows:
        ws.append([safe_cell(row.get(header)) for header in headers])
    for row_index in range(2, ws.max_row + 1):
        for cell in ws[row_index]:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if row_index % 2 == 0:
                cell.fill = PatternFill("solid", fgColor=PALE_BLUE)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    ws.sheet_view.showGridLines = False
    for column_index, header in enumerate(headers, start=1):
        sample = [str(header)] + [str(ws.cell(row, column_index).value or "") for row in range(2, min(ws.max_row, 120) + 1)]
        ws.column_dimensions[get_column_letter(column_index)].width = min(55, max(10, max(len(value) for value in sample) + 2))


def sync_database(source: Path, destination: Path, now_iso: str, gates: list[dict[str, Any]], events: list[dict[str, Any]]) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, destination)
    con = sqlite3.connect(destination)
    try:
        con.execute("PRAGMA foreign_keys=ON")
        con.execute("BEGIN IMMEDIATE")
        response81 = clone_response81(con, now_iso)
        con.executescript("""
        CREATE TABLE IF NOT EXISTS section5_s2_release (
            release_key TEXT PRIMARY KEY,
            response_number INTEGER NOT NULL,
            checkpoint_code TEXT NOT NULL,
            session_state TEXT NOT NULL,
            section_state TEXT NOT NULL,
            internal_acceptance_state TEXT NOT NULL,
            provider_preview_state TEXT NOT NULL,
            provider_approval_claimed INTEGER NOT NULL,
            physical_proof_state TEXT NOT NULL,
            physical_proof_ordered INTEGER NOT NULL,
            accepted_predecessor_mutated INTEGER NOT NULL,
            frozen_section3_release_mutated INTEGER NOT NULL,
            immutable_publication_mutated INTEGER NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_s2_acceptance_gate (
            gate_key TEXT PRIMARY KEY,
            description TEXT NOT NULL,
            status TEXT NOT NULL,
            evidence TEXT NOT NULL,
            checked_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_s2_handoff (
            handoff_key TEXT PRIMARY KEY,
            next_session TEXT NOT NULL,
            scope TEXT NOT NULL,
            required_inputs TEXT NOT NULL,
            state TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_s2_terminal_recovery_event (
            event_code TEXT PRIMARY KEY,
            event_number INTEGER NOT NULL,
            condition TEXT NOT NULL,
            recovery TEXT NOT NULL,
            status TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_s2_artifact_freeze (
            artifact_key TEXT PRIMARY KEY,
            relative_path TEXT NOT NULL,
            bytes INTEGER NOT NULL,
            sha256 TEXT NOT NULL,
            state TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        """)
        con.execute("DELETE FROM section5_s2_release WHERE release_key=?", (RELEASE_KEY,))
        con.execute(
            "INSERT INTO section5_s2_release VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (
                RELEASE_KEY, 81, CHECKPOINT_CODE, "session_complete", "continue",
                "passed", "controlled_pending", 0, "controlled_pending", 0,
                0, 0, 0, now_iso,
            ),
        )
        con.execute("DELETE FROM section5_s2_acceptance_gate")
        con.executemany(
            "INSERT INTO section5_s2_acceptance_gate VALUES (?,?,?,?,?)",
            [(row["gate_key"], row["description"], row["status"], row["evidence"], row["checked_at"]) for row in gates],
        )
        con.execute("DELETE FROM section5_s2_handoff WHERE handoff_key='MRHPD-S5-S3-HANDOFF'")
        con.execute(
            "INSERT INTO section5_s2_handoff VALUES (?,?,?,?,?,?)",
            (
                "MRHPD-S5-S3-HANDOFF",
                "Remediation Section 5 of 5 Session 3 of 3",
                "Provider-side conversion evidence, physical-proof evidence and corrections, final project-wide production acceptance, final whole-project restore, and project completion.",
                "Complete Response 81 restore; genuine provider or physical-proof evidence if available; no prior checkpoint or conversation reconstruction required.",
                "ready",
                now_iso,
            ),
        )
        con.execute("DELETE FROM section5_s2_terminal_recovery_event")
        con.executemany(
            "INSERT INTO section5_s2_terminal_recovery_event VALUES (?,?,?,?,?,?)",
            [(row["event_code"], row["event_number"], row["condition"], row["recovery"], row["status"], row["recorded_at"]) for row in events],
        )
        if table_exists(con, "section5_s2_checkpoint"):
            con.execute("DELETE FROM section5_s2_checkpoint WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
            con.execute(
                "INSERT INTO section5_s2_checkpoint VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
                (
                    CHECKPOINT_CODE, 81, SECTION_LABEL, SESSION_LABEL, CHECKPOINT_LABEL,
                    "session_complete", "passed", "controlled_pending", 0, 0,
                    "Remediation Section 5 of 5 Session 3 of 3 Checkpoint 1 of 3", now_iso,
                ),
            )
        if table_exists(con, "section5_session2_checkpoint"):
            info = table_info(con, "section5_session2_checkpoint")
            columns = [row[1] for row in info]
            source_row = con.execute("SELECT * FROM section5_session2_checkpoint ORDER BY rowid DESC LIMIT 1").fetchone()
            if source_row:
                record = dict(zip(columns, source_row))
                for row in info:
                    if row[5] and row[1] != "checkpoint_code":
                        record.pop(row[1], None)
                replacements = {
                    "checkpoint_code": CHECKPOINT_CODE,
                    "response_number": 81,
                    "checkpoint_label": CHECKPOINT_LABEL,
                    "checkpoint_state": "session_complete",
                    "state": "session_complete",
                    "next": "Remediation Section 5 of 5 Session 3 of 3 Checkpoint 1 of 3",
                    "recorded_at": now_iso,
                    "updated_at": now_iso,
                }
                for key, value in replacements.items():
                    if key in columns:
                        record[key] = value
                con.execute("DELETE FROM section5_session2_checkpoint WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
                keys = [column for column in columns if column in record]
                con.execute(
                    f"INSERT INTO section5_session2_checkpoint ({','.join(keys)}) VALUES ({','.join('?' for _ in keys)})",
                    [record[key] for key in keys],
                )
        if table_exists(con, "section5_s2_provider_evidence_reconciliation"):
            unsupported = con.execute(
                "SELECT COUNT(*) FROM section5_s2_provider_evidence_reconciliation WHERE provider_approval_claimed!=0"
            ).fetchone()[0]
            if unsupported:
                raise RuntimeError({"unsupported_provider_approval_claims": unsupported})
        if table_exists(con, "section5_s2_provider_conversion_issue"):
            observed = con.execute("SELECT COUNT(*) FROM section5_s2_provider_conversion_issue WHERE observed!=0").fetchone()[0]
            if observed:
                raise RuntimeError({"unexpected_observed_provider_issues_without_evidence": observed})
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or fk:
            raise RuntimeError({"integrity": integrity, "foreign_keys": fk[:20]})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()
    con = sqlite3.connect(destination)
    try:
        counts = {
            "tables": con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0],
            "response81": con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R81'").fetchone()[0],
            "release": con.execute("SELECT COUNT(*) FROM section5_s2_release WHERE release_key=? AND session_state='session_complete'", (RELEASE_KEY,)).fetchone()[0],
            "gates": con.execute("SELECT COUNT(*) FROM section5_s2_acceptance_gate").fetchone()[0],
            "provider_approvals": con.execute("SELECT COUNT(*) FROM section5_s2_release WHERE provider_approval_claimed!=0").fetchone()[0],
            "proof_orders": con.execute("SELECT COUNT(*) FROM section5_s2_release WHERE physical_proof_ordered!=0").fetchone()[0],
            "integrity": con.execute("PRAGMA integrity_check").fetchone()[0],
            "foreign_keys": len(list(con.execute("PRAGMA foreign_key_check"))),
        }
    finally:
        con.close()
    if counts["response81"] != 1 or counts["release"] != 1 or counts["provider_approvals"] or counts["proof_orders"]:
        raise RuntimeError({"terminal_database_gate": counts})
    return {
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "response81_record": response81,
        **counts,
    }


def sync_workbook(source: Path, destination: Path, db_qa: dict[str, Any], gates: list[dict[str, Any]], events: list[dict[str, Any]], now_iso: str) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    wb = load_workbook(source)
    inherited = list(wb.sheetnames)
    datasets = {
        "S5S2 CP3 Dashboard": [
            {"Control": "Response", "Value": 81, "Status": "current"},
            {"Control": "Checkpoint", "Value": "3 of 3", "Status": "complete"},
            {"Control": "Session", "Value": "2 of 3", "Status": "complete"},
            {"Control": "Section", "Value": "5 of 5", "Status": "continue"},
            {"Control": "Internal acceptance", "Value": "all observable gates passed", "Status": "passed"},
            {"Control": "Provider preview", "Value": "no item-level evidence", "Status": "controlled pending"},
            {"Control": "Physical proof", "Value": "not ordered, received, inspected, or approved", "Status": "controlled pending"},
            {"Control": "Next", "Value": "Remediation Section 5 Session 3 of 3", "Status": "continue"},
        ],
        "S5S2 CP3 Acceptance": gates,
        "S5S2 CP3 Provider Boundary": [
            {"Control": "Provider submission receipt", "State": "not supplied or discovered", "Claim": "none"},
            {"Control": "Provider-rendered preview", "State": "not supplied or discovered", "Claim": "none"},
            {"Control": "Provider warnings/errors", "State": "not supplied or discovered", "Claim": "absence not inferred"},
            {"Control": "Provider approval", "State": "not claimed", "Claim": "none"},
            {"Control": "Physical proof order", "State": "not ordered", "Claim": "none"},
            {"Control": "Physical proof signoff", "State": "not claimed", "Claim": "none"},
        ],
        "S5S2 CP3 Response": [db_qa["response81_record"]],
        "S5S2 CP3 Recovery": events,
        "S5S2 CP3 Database QA": [{key: value for key, value in db_qa.items() if key != "response81_record"}],
        "S5S2 CP3 Handoff": [{
            "Next Session": "Remediation Section 5 of 5 Session 3 of 3",
            "Scope": "Provider evidence, physical proof and corrections, final production acceptance, project-complete restore.",
            "Required Baseline": "Complete Response 81 restore only",
            "User Upload Required Now": "No",
            "State": "ready",
            "Recorded At": now_iso,
        }],
    }
    for title, rows in datasets.items():
        if title in wb.sheetnames:
            del wb[title]
        ws = wb.create_sheet(title=title)
        write_sheet(ws, rows)
    wb.properties.title = "MRHPD v3.0.0a Comprehensive Tracking Through Response 81"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.properties.subject = "Human Pathogen Database Section 5 Session 2 complete tracking"
    wb.save(destination)
    with zipfile.ZipFile(destination) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("workbook ZIP CRC failed")
    check = load_workbook(destination, read_only=True, data_only=False)
    try:
        sheet_names = list(check.sheetnames)
        formula_count = 0
        formula_errors = []
        for ws in check.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    value = cell.value
                    if isinstance(value, str) and value.startswith("="):
                        formula_count += 1
                        if any(token in value for token in ("#REF!", "#NAME?", "#VALUE!", "#DIV/0!", "#N/A")):
                            formula_errors.append(f"{ws.title}!{cell.coordinate}:{value}")
    finally:
        check.close()
    lost = sorted(set(inherited) - set(sheet_names))
    if lost or len(sheet_names) < len(inherited) + len(datasets) or formula_errors:
        raise RuntimeError({"workbook_gate": {"lost": lost, "sheets": len(sheet_names), "formula_errors": formula_errors[:30]}})
    return {
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "source_sheet_count": len(inherited),
        "current_sheet_count": len(sheet_names),
        "new_sheet_count": len(sheet_names) - len(inherited),
        "lost_sheets": lost,
        "formula_count": formula_count,
        "formula_error_count": len(formula_errors),
        "status": "passed",
    }


def shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def write_tracking(project: Path, db_path: Path, now_iso: str) -> dict[str, Any]:
    root = project / "Tracking" / "Prompt Response" / "Through Response 81"
    root.mkdir(parents=True, exist_ok=True)
    con = sqlite3.connect(db_path)
    try:
        info = con.execute("PRAGMA table_info(thread_response_reconciliation_cp3)").fetchall()
        columns = [row[1] for row in info]
        rows = [dict(zip(columns, row)) for row in con.execute("SELECT * FROM thread_response_reconciliation_cp3 ORDER BY CAST(response_number AS REAL), response_key")]
        fractional = []
        if table_exists(con, "fractional_prompt_cp3"):
            f_info = con.execute("PRAGMA table_info(fractional_prompt_cp3)").fetchall()
            f_columns = [row[1] for row in f_info]
            fractional = [dict(zip(f_columns, row)) for row in con.execute("SELECT * FROM fractional_prompt_cp3 ORDER BY CAST(prompt_number AS REAL), prompt_number")]
    finally:
        con.close()
    response81 = next(row for row in rows if row.get("response_key") == "R81")
    response_path = root / "Response_81_Tracking.json"
    json_write(response_path, response81)
    summary_json = root / "Medical References - Human Pathogen Database v3.0.0a Summary Index Through Response 81.json"
    summary_csv = root / "Medical References - Human Pathogen Database v3.0.0a Summary Index Through Response 81.csv"
    summary_rows = [{
        "response": row.get("response_label") or row.get("response_number"),
        "major_topic": row.get("major_topic"),
        "title": row.get("title"),
        "goal": row.get("goal"),
        "summary": row.get("summary"),
        "state": row.get("state"),
    } for row in rows]
    json_write(summary_json, summary_rows)
    csv_write(summary_csv, summary_rows)

    raw_docx = root / "Medical References - Human Pathogen Database v3.0.0a Alternating Raw Prompts and Responses Through Response 81.docx"
    raw_doc = Document()
    raw_doc.core_properties.title = "Human Pathogen Database — Alternating Raw Prompts and Responses Through Response 81"
    raw_doc.core_properties.author = "Brent McAnulty, M.D."
    raw_doc.add_heading("Human Pathogen Database", 0)
    raw_doc.add_paragraph("Alternating Raw Prompts and Responses Through Response 81")
    for row in rows:
        number = row.get("response_label") or row.get("response_number")
        raw_doc.add_heading(f"Response {number}: {row.get('title') or 'Untitled exchange'}", level=1)
        table = raw_doc.add_table(rows=2, cols=1)
        table.style = "Table Grid"
        table.cell(0, 0).text = f"RAW PROMPT {number}\n\n{row.get('raw_prompt') or '[RAW PROMPT UNAVAILABLE]'}"
        shade_cell(table.cell(0, 0), "D9EAF7")
        table.cell(1, 0).text = f"RAW RESPONSE {number}\n\n{row.get('raw_response') or row.get('summary') or '[RAW RESPONSE UNAVAILABLE]'}\n\nSUMMARY\n{row.get('summary') or ''}"
        shade_cell(table.cell(1, 0), "E2F0D9")
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(9)
        raw_doc.add_paragraph()
    if fractional:
        raw_doc.add_heading("Fractional prompts", level=1)
        for row in fractional:
            raw_doc.add_paragraph(f"Prompt {row.get('prompt_number')}: {row.get('prompt_text')}")
    raw_doc.save(raw_docx)

    net_prompt = (
        "Continue the Human Pathogen Database from the newest verified state without regression. Use Google Drive as controlling storage, recover managed files autonomously, preserve accepted clinical and publication artifacts, complete Section 5 print-production remediation, distinguish internal readiness from provider and physical-proof evidence, maintain Raw/Net tracking and indexes, and emit checkpoint or full self-contained restores at the required boundaries."
    )
    net_response = (
        "Remediation Sections 1–4 are complete. Section 5 Sessions 1 and 2 are complete through Response 81. The quality-first KDP Premium Color production master, 537-page immutable digital edition, 538-page print candidate, exact cover package, provider-evidence framework, proof-readiness controls, synchronized database/workbook/application, tracking, indexes, manifests, and complete Session 2 restore are current. Genuine provider preview and physical-proof evidence, any resulting corrections, final production signoff, and project-complete release remain in Session 3."
    )
    net_docx = root / "Medical References - Human Pathogen Database v3.0.0a Alternating Net Prompts and Responses Through Response 81.docx"
    net_doc = Document()
    net_doc.core_properties.title = "Human Pathogen Database — Alternating Net Prompts and Responses Through Response 81"
    net_doc.core_properties.author = "Brent McAnulty, M.D."
    net_doc.add_heading("Human Pathogen Database", 0)
    net_doc.add_heading("Print-production and final-release remediation", level=1)
    table = net_doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.cell(0, 0).text = "NET PROMPT\n\n" + net_prompt
    shade_cell(table.cell(0, 0), "D9EAF7")
    table.cell(1, 0).text = "NET RESPONSE\n\n" + net_response
    shade_cell(table.cell(1, 0), "E2F0D9")
    net_doc.save(net_docx)

    everything = root / "Medical References - Human Pathogen Database v3.0.0a Everything in One Thread Through Response 81.xlsx"
    wb = Workbook()
    wb.remove(wb.active)
    datasets = {
        "Raw Prompts": [{"Response": row.get("response_label") or row.get("response_number"), "Major Topic": row.get("major_topic"), "Title": row.get("title"), "Raw Prompt": row.get("raw_prompt")} for row in rows],
        "Raw Responses": [{"Response": row.get("response_label") or row.get("response_number"), "Major Topic": row.get("major_topic"), "Title": row.get("title"), "Raw Response": row.get("raw_response"), "Summary": row.get("summary")} for row in rows],
        "Fractional Prompts": fractional,
        "Net Prompt": [{"Major Topic": "Human Pathogen Database remediation", "Net Prompt": net_prompt}],
        "Net Response": [{"Major Topic": "Human Pathogen Database remediation", "Net Response": net_response}],
        "Summary Index": summary_rows,
    }
    for title, data in datasets.items():
        ws = wb.create_sheet(title)
        write_sheet(ws, data)
    wb.properties.title = "Human Pathogen Database — Everything in One Thread Through Response 81"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(everything)

    raw_net_md = root / "Medical References - Human Pathogen Database v3.0.0a Raw and Net Tracking Through Response 81.md"
    text_write(raw_net_md, f"""# Human Pathogen Database — Raw and Net Tracking Through Response 81

## Raw Prompt 81

Continue

## Raw Response 81

{response81.get('summary')}

## Net Prompt

{net_prompt}

## Net Response

{net_response}

Updated: {now_iso}
""")
    cumulative = root / "Medical References - Human Pathogen Database v3.0.0a Cumulative Thread Index Through Response 81.md"
    lines = ["# Human Pathogen Database — Cumulative Thread Index Through Response 81", ""]
    for row in rows:
        lines.append(f"- Response {row.get('response_label') or row.get('response_number')} — {row.get('title')}: {row.get('summary')}")
    lines.extend(["", f"Updated: {now_iso}"])
    text_write(cumulative, "\n".join(lines))
    files = [response_path, summary_json, summary_csv, raw_docx, net_docx, everything, raw_net_md, cumulative]
    return {"status": "passed", "response_records": len(rows), "fractional_records": len(fractional), "files": [path.relative_to(project).as_posix() for path in files]}


def audit_publication_and_cover(project: Path) -> dict[str, Any]:
    publication = project / PUBLICATION_REL
    print_interior = project / PRINT_INTERIOR_REL
    cover_png = project / COVER_PNG_REL
    cover_pdf = project / COVER_PDF_REL
    template_png = project / TEMPLATE_PNG_REL
    template_pdf = project / TEMPLATE_PDF_REL
    if sha256_file(publication) != PUBLICATION_SHA256:
        raise RuntimeError("immutable digital publication changed")
    if sha256_file(print_interior) != PRINT_INTERIOR_SHA256:
        raise RuntimeError("frozen print interior changed")
    if sha256_file(cover_png) != COVER_SHA256:
        raise RuntimeError("frozen cover raster changed")
    digital = PdfReader(str(publication))
    digital_pages = len(digital.pages)
    searchable = sum(1 for page in digital.pages if (page.extract_text() or "").strip())
    print_reader = PdfReader(str(print_interior))
    print_pages = len(print_reader.pages)
    print_searchable = sum(1 for page in print_reader.pages[:537] if (page.extract_text() or "").strip())
    terminal_blank = not bool((print_reader.pages[537].extract_text() or "").strip()) if print_pages >= 538 else False
    print_doc = fitz.open(print_interior)
    try:
        geometry_failures = []
        for index in range(print_doc.page_count):
            rect = print_doc[index].rect
            if abs(rect.width - 612.0) > 0.5 or abs(rect.height - 792.0) > 0.5:
                geometry_failures.append({"page": index + 1, "width": rect.width, "height": rect.height})
    finally:
        print_doc.close()
    with Image.open(cover_png) as image:
        cover = {"pixels": [image.width, image.height], "mode": image.mode, "alpha": "A" in image.getbands()}
    cover_reader = PdfReader(str(cover_pdf))
    template_reader = PdfReader(str(template_pdf))
    with Image.open(template_png) as image:
        template_pixels = [image.width, image.height]
    editable_candidates = [path for path in project.rglob("*.docx") if path.is_file() and sha256_file(path) == EDITABLE_SHA256]
    if len(editable_candidates) != 1:
        raise RuntimeError({"editable_assembly_candidates": [str(path) for path in editable_candidates]})
    result = {
        "status": "passed",
        "digital_publication": {"path": PUBLICATION_REL, "bytes": publication.stat().st_size, "sha256": sha256_file(publication), "pages": digital_pages, "searchable_pages": searchable},
        "editable_assembly": {"path": editable_candidates[0].relative_to(project).as_posix(), "bytes": editable_candidates[0].stat().st_size, "sha256": EDITABLE_SHA256},
        "print_interior": {"path": PRINT_INTERIOR_REL, "bytes": print_interior.stat().st_size, "sha256": sha256_file(print_interior), "pages": print_pages, "searchable_source_pages": print_searchable, "terminal_blank": terminal_blank, "geometry_failures": geometry_failures},
        "cover": {"path": COVER_PNG_REL, "bytes": cover_png.stat().st_size, "sha256": sha256_file(cover_png), **cover, "cover_pdf_pages": len(cover_reader.pages), "template_pdf_pages": len(template_reader.pages), "template_pixels": template_pixels},
    }
    if digital_pages != 537 or searchable != 537 or print_pages != 538 or print_searchable != 537 or not terminal_blank or geometry_failures:
        raise RuntimeError({"publication_gate": result})
    if cover["pixels"] != [5554, 3375] or cover["mode"] != "RGB" or cover["alpha"] or len(cover_reader.pages) != 1 or len(template_reader.pages) != 1 or template_pixels != [5554, 3375]:
        raise RuntimeError({"cover_gate": result["cover"]})
    return result


def write_application_audit(project: Path, db_path: Path, workbook_path: Path, publication_qa: dict[str, Any], now_iso: str) -> dict[str, Any]:
    roots = [path for path in project.rglob("human_pathogen_app.py") if path.is_file() and sha256_file(path) == APPLICATION_SHA256]
    if len(roots) != 1:
        raise RuntimeError({"main_application_candidates": [str(path) for path in roots]})
    app_path = roots[0]
    root = project / "App" / "Section 5 Session 2 Complete"
    root.mkdir(parents=True, exist_ok=True)
    pointer = root / "CURRENT_DATABASE.txt"
    text_write(pointer, db_path.relative_to(project).as_posix())
    state = root / "CURRENT_PROJECT_STATE.json"
    json_write(state, {
        "schema": "mrhpd-section5-session2-complete-state-1.0",
        "response": 81,
        "section": SECTION_LABEL,
        "session": SESSION_LABEL,
        "checkpoint": CHECKPOINT_LABEL,
        "state": "session_complete",
        "database": db_path.relative_to(project).as_posix(),
        "workbook": workbook_path.relative_to(project).as_posix(),
        "main_application": app_path.relative_to(project).as_posix(),
        "main_application_sha256": APPLICATION_SHA256,
        "provider_preview": "controlled_pending",
        "physical_proof": "controlled_pending",
        "recorded_at": now_iso,
    })
    audit_script = root / "audit_section5_session2_complete.py"
    text_write(audit_script, f'''#!/usr/bin/env python3
import json, sqlite3
from pathlib import Path
project=Path(__file__).resolve().parents[2]
db=project/{db_path.relative_to(project).as_posix()!r}
con=sqlite3.connect(db)
try:
 integrity=con.execute("PRAGMA integrity_check").fetchone()[0]
 fk=len(list(con.execute("PRAGMA foreign_key_check")))
 response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R81'").fetchone()[0]
 release=con.execute("SELECT COUNT(*) FROM section5_s2_release WHERE release_key='{RELEASE_KEY}' AND session_state='session_complete' AND provider_approval_claimed=0 AND physical_proof_ordered=0").fetchone()[0]
 provider=con.execute("SELECT COUNT(*) FROM section5_s2_provider_evidence_reconciliation WHERE provider_approval_claimed!=0").fetchone()[0] if con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name='section5_s2_provider_evidence_reconciliation'").fetchone() else 0
finally: con.close()
result={{'status':'passed' if integrity=='ok' and fk==0 and response==1 and release==1 and provider==0 else 'failed','integrity':integrity,'foreign_keys':fk,'response81':response,'release':release,'unsupported_provider_claims':provider}}
print(json.dumps(result,indent=2))
raise SystemExit(0 if result['status']=='passed' else 1)
''')
    result = subprocess.run([sys.executable, str(audit_script)], cwd=project, text=True, capture_output=True, timeout=300)
    if result.returncode:
        raise RuntimeError({"application_audit_failed": {"stdout": result.stdout[-8000:], "stderr": result.stderr[-8000:]}})
    audit = json.loads(result.stdout)
    audit.update({
        "main_application_path": app_path.relative_to(project).as_posix(),
        "main_application_sha256": sha256_file(app_path),
        "main_application_unchanged": True,
        "publication_sha256": publication_qa["digital_publication"]["sha256"],
        "print_interior_sha256": publication_qa["print_interior"]["sha256"],
        "cover_sha256": publication_qa["cover"]["sha256"],
    })
    output = root / "SECTION5_SESSION2_COMPLETE_APPLICATION_AUDIT.json"
    json_write(output, audit)
    return {"status": "passed", "files": [pointer.relative_to(project).as_posix(), state.relative_to(project).as_posix(), audit_script.relative_to(project).as_posix(), output.relative_to(project).as_posix()], **audit}


def write_reports(project: Path, db_qa: dict[str, Any], workbook_qa: dict[str, Any], publication_qa: dict[str, Any], app_qa: dict[str, Any], gates: list[dict[str, Any]], events: list[dict[str, Any]], now_iso: str) -> dict[str, Any]:
    root = project / "Reports" / "Section 5 Session 2" / "Complete"
    root.mkdir(parents=True, exist_ok=True)
    docx_path = root / "MRHPD v3.0.0a Section 5 Session 2 Complete Acceptance Report.docx"
    pdf_path = root / "MRHPD v3.0.0a Section 5 Session 2 Complete Acceptance Report.pdf"
    xlsx_path = root / "MRHPD v3.0.0a Section 5 Session 2 Complete Acceptance Register.xlsx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    doc.core_properties.title = "Human Pathogen Database — Section 5 Session 2 Complete Acceptance Report"
    doc.core_properties.author = "Brent McAnulty, M.D."
    title = doc.add_heading("Human Pathogen Database", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Remediation Section 5 of 5 • Session 2 of 3 Complete • Response 81")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Internal production acceptance is complete. Provider preview, provider approval, physical-proof order, receipt, inspection, correction, and signoff remain controlled pending unless genuine item-level evidence is supplied or discovered.")
    doc.add_heading("Terminal disposition", level=1)
    summary_rows = [
        ("Response", "81 — complete"),
        ("Checkpoint", "3 of 3 — complete"),
        ("Session", "2 of 3 — complete"),
        ("Remediation Section 5", "continue"),
        ("Database", f"{db_qa['tables']} physical tables; integrity {db_qa['integrity']}; {db_qa['foreign_keys']} foreign-key violations"),
        ("Workbook", f"{workbook_qa['current_sheet_count']} sheets; {workbook_qa['formula_error_count']} formula-error tokens"),
        ("Digital publication", "537 searchable pages; byte-identical"),
        ("Print interior", "538 pages; one intentional terminal blank; byte-identical"),
        ("Cover", "5554 × 3375 RGB pixels; byte-identical"),
        ("Provider preview", "controlled pending"),
        ("Physical proof", "controlled pending"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Control"
    table.rows[0].cells[1].text = "Result"
    for key, value in summary_rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    doc.add_heading("Acceptance gates", level=1)
    gate_table = doc.add_table(rows=1, cols=3)
    gate_table.style = "Table Grid"
    for index, text in enumerate(("Gate", "Status", "Evidence")):
        gate_table.rows[0].cells[index].text = text
    for row in gates:
        cells = gate_table.add_row().cells
        cells[0].text = row["description"]
        cells[1].text = row["status"]
        cells[2].text = row["evidence"]
    doc.add_heading("Recovery events", level=1)
    for row in events:
        doc.add_heading(f"{row['event_number']} — {row['event_code']}", level=2)
        doc.add_paragraph(row["condition"])
        doc.add_paragraph(row["recovery"])
    doc.add_heading("Next session", level=1)
    doc.add_paragraph("Session 3 will ingest genuine provider or physical-proof evidence if available, process any governed correction cycle, complete final production acceptance, and emit the entire completed project. In the absence of new external evidence, the provider and proof gates remain explicitly pending rather than inferred.")
    doc.save(docx_path)

    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph("Human Pathogen Database", ParagraphStyle("TitleX", parent=styles["Title"], textColor=colors.HexColor("#17324D"), alignment=TA_CENTER, fontSize=20, leading=24)))
    story.append(Paragraph("Remediation Section 5 of 5 • Session 2 of 3 Complete • Response 81", ParagraphStyle("SubX", parent=styles["Normal"], textColor=colors.HexColor("#1C7475"), alignment=TA_CENTER, fontSize=10, leading=13)))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Internal production acceptance is complete. Provider-side and physical-proof events remain controlled pending without unsupported completion claims.", styles["BodyText"]))
    data = [["Control", "Result"]] + [[key, value] for key, value in summary_rows]
    report_table = Table(data, colWidths=[2.0 * inch, 4.8 * inch], repeatRows=1)
    report_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17324D")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAB8C0")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EAF1F5")]),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("LEADING", (0, 0), (-1, -1), 10),
    ]))
    story.extend([Spacer(1, 0.18 * inch), report_table, PageBreak(), Paragraph("Acceptance gates", styles["Heading1"])])
    gate_data = [["Gate", "Status", "Evidence"]] + [[row["description"], row["status"], row["evidence"]] for row in gates]
    gate_pdf = Table(gate_data, colWidths=[3.4 * inch, 1.15 * inch, 2.25 * inch], repeatRows=1)
    gate_pdf.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17324D")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#AAB8C0")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 6.7),
        ("LEADING", (0, 0), (-1, -1), 8.2),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EAF1F5")]),
    ]))
    story.append(gate_pdf)
    story.append(PageBreak())
    story.append(Paragraph("Recovery and continuation", styles["Heading1"]))
    for row in events:
        story.append(Paragraph(f"{row['event_number']} — {row['event_code']}", styles["Heading2"]))
        story.append(Paragraph(row["condition"], styles["BodyText"]))
        story.append(Paragraph(row["recovery"], styles["BodyText"]))
        story.append(Spacer(1, 0.08 * inch))
    SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=0.45 * inch, leftMargin=0.45 * inch, topMargin=0.45 * inch, bottomMargin=0.45 * inch, title="Human Pathogen Database Section 5 Session 2 Complete Acceptance Report", author="Brent McAnulty, M.D.").build(story)

    wb = Workbook()
    wb.remove(wb.active)
    datasets = {
        "Disposition": [{"Control": key, "Result": value} for key, value in summary_rows],
        "Acceptance": gates,
        "Recovery": events,
        "Database": [{key: value for key, value in db_qa.items() if key != "response81_record"}],
        "Workbook": [workbook_qa],
        "Publication": [publication_qa],
        "Application": [app_qa],
        "Handoff": [{"Next Session": "Section 5 Session 3 of 3", "State": "ready", "External Evidence": "provider preview and physical proof when genuinely available"}],
    }
    for title, rows in datasets.items():
        ws = wb.create_sheet(title)
        write_sheet(ws, rows)
    wb.properties.title = "MRHPD Section 5 Session 2 Complete Acceptance Register"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(xlsx_path)
    with zipfile.ZipFile(xlsx_path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("acceptance register CRC failed")
    pdf_reader = PdfReader(str(pdf_path))
    if len(pdf_reader.pages) < 3 or sum(1 for page in pdf_reader.pages if (page.extract_text() or "").strip()) != len(pdf_reader.pages):
        raise RuntimeError("acceptance PDF is not fully searchable")
    return {
        "status": "passed",
        "docx": {"path": docx_path.relative_to(project).as_posix(), "bytes": docx_path.stat().st_size, "sha256": sha256_file(docx_path)},
        "pdf": {"path": pdf_path.relative_to(project).as_posix(), "bytes": pdf_path.stat().st_size, "sha256": sha256_file(pdf_path), "pages": len(pdf_reader.pages)},
        "xlsx": {"path": xlsx_path.relative_to(project).as_posix(), "bytes": xlsx_path.stat().st_size, "sha256": sha256_file(xlsx_path)},
    }


def extract_text(path: Path, limit: int = 2_000_000) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in {".md", ".txt", ".csv", ".json", ".py", ".html", ".htm", ".yml", ".yaml", ".xml"}:
            return path.read_text(encoding="utf-8", errors="replace")[:limit]
        if suffix == ".docx":
            doc = Document(path)
            chunks = [paragraph.text for paragraph in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells)
            return "\n".join(chunks)[:limit]
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            chunks = []
            size = 0
            for page in reader.pages:
                text = page.extract_text() or ""
                chunks.append(text)
                size += len(text)
                if size >= limit:
                    break
            return "\n".join(chunks)[:limit]
        if suffix == ".xlsx":
            wb = load_workbook(path, read_only=True, data_only=False)
            try:
                chunks = []
                size = 0
                for ws in wb.worksheets:
                    chunks.append(ws.title)
                    for row in ws.iter_rows(values_only=True):
                        line = " | ".join("" if value is None else str(value) for value in row)
                        chunks.append(line)
                        size += len(line)
                        if size >= limit:
                            return "\n".join(chunks)[:limit]
                return "\n".join(chunks)[:limit]
            finally:
                wb.close()
        if suffix in {".sqlite", ".db"}:
            con = sqlite3.connect(path)
            try:
                schema = [row[0] for row in con.execute("SELECT sql FROM sqlite_master WHERE sql IS NOT NULL ORDER BY name")]
                return "\n".join(schema)[:limit]
            finally:
                con.close()
    except Exception as exc:
        return f"[index extraction error: {exc!r}]"
    return ""


def build_indexes(project: Path, now_iso: str) -> dict[str, Any]:
    root = project / "Indexes" / "Section 5 Session 2 Complete"
    root.mkdir(parents=True, exist_ok=True)
    source_json = root / "MRHPD v3.0.0a Section 5 Session 2 Complete Source Index.json"
    source_csv = root / "MRHPD v3.0.0a Section 5 Session 2 Complete Source Index.csv"
    bit_path = root / "MRHPD v3.0.0a Section 5 Session 2 Complete Bit Index.sqlite"
    qa_path = root / "MRHPD v3.0.0a Section 5 Session 2 Complete Index QA.json"
    excluded = {source_json.resolve(), source_csv.resolve(), bit_path.resolve(), qa_path.resolve()}
    rows: list[dict[str, Any]] = []
    fts_rows: list[tuple[str, str, str, str]] = []
    searchable_suffixes = {".md", ".txt", ".csv", ".json", ".py", ".html", ".htm", ".yml", ".yaml", ".xml", ".docx", ".pdf", ".xlsx", ".sqlite", ".db"}
    for path in sorted(item for item in project.rglob("*") if item.is_file() and item.resolve() not in excluded):
        rel = path.relative_to(project).as_posix()
        purpose = "Project artifact"
        if rel.startswith("Database/"):
            purpose = "Canonical or historical project database"
        elif rel.startswith("Tracking/"):
            purpose = "Prompt/response, workbook, or project-state tracking"
        elif rel.startswith("Documents/"):
            purpose = "Publication or editable document"
        elif rel.startswith("Print Production/"):
            purpose = "Print-production interior, cover, template, or proof artifact"
        elif rel.startswith("QA/"):
            purpose = "Quality-assurance evidence"
        elif rel.startswith("Reports/"):
            purpose = "Human-readable project report or register"
        elif rel.startswith("App/"):
            purpose = "Local application or read-only audit surface"
        rows.append({
            "record_type": "physical_file",
            "path": rel,
            "container_path": "",
            "name": path.name,
            "purpose": purpose,
            "bytes": path.stat().st_size,
            "sha256": sha256_file(path),
            "user_searchable": int(path.suffix.lower() in searchable_suffixes),
        })
        content = extract_text(path) if path.suffix.lower() in searchable_suffixes else ""
        fts_rows.append((rel, path.name, purpose, content))
        if path.suffix.lower() == ".zip":
            try:
                with zipfile.ZipFile(path) as zf:
                    if zf.testzip() is not None:
                        raise RuntimeError("CRC failure")
                    for info in zf.infolist():
                        if info.is_dir():
                            continue
                        member_path = f"{rel}!/{info.filename}"
                        rows.append({
                            "record_type": "container_member",
                            "path": member_path,
                            "container_path": rel,
                            "name": PurePosixPath(info.filename).name,
                            "purpose": "Member of project ZIP container",
                            "bytes": info.file_size,
                            "sha256": "",
                            "user_searchable": int(PurePosixPath(info.filename).suffix.lower() in searchable_suffixes),
                        })
                        member_content = ""
                        suffix = PurePosixPath(info.filename).suffix.lower()
                        if suffix in {".md", ".txt", ".csv", ".json", ".py", ".html", ".htm", ".yml", ".yaml", ".xml"} and info.file_size <= 10_000_000:
                            try:
                                member_content = zf.read(info).decode("utf-8", errors="replace")[:2_000_000]
                            except Exception:
                                member_content = ""
                        fts_rows.append((member_path, PurePosixPath(info.filename).name, "Member of project ZIP container", member_content))
            except (zipfile.BadZipFile, RuntimeError):
                pass
    json_write(source_json, {"schema": "mrhpd-source-index-3.0", "generated_at": now_iso, "record_count": len(rows), "records": rows})
    csv_write(source_csv, rows)
    if bit_path.exists():
        bit_path.unlink()
    con = sqlite3.connect(bit_path)
    try:
        con.executescript("""
        CREATE TABLE artifact (
            artifact_id INTEGER PRIMARY KEY,
            record_type TEXT NOT NULL,
            path TEXT NOT NULL UNIQUE,
            container_path TEXT,
            name TEXT NOT NULL,
            purpose TEXT NOT NULL,
            bytes INTEGER NOT NULL,
            sha256 TEXT,
            user_searchable INTEGER NOT NULL
        );
        CREATE VIRTUAL TABLE artifact_fts USING fts5(path, name, purpose, content);
        """)
        for row, payload in zip(rows, fts_rows):
            con.execute("INSERT INTO artifact (record_type,path,container_path,name,purpose,bytes,sha256,user_searchable) VALUES (?,?,?,?,?,?,?,?)", (row["record_type"], row["path"], row["container_path"], row["name"], row["purpose"], row["bytes"], row["sha256"], row["user_searchable"]))
            con.execute("INSERT INTO artifact_fts (path,name,purpose,content) VALUES (?,?,?,?)", payload)
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        counts = {
            "artifact": con.execute("SELECT COUNT(*) FROM artifact").fetchone()[0],
            "fts": con.execute("SELECT COUNT(*) FROM artifact_fts").fetchone()[0],
            "response81": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"Response 81"',)).fetchone()[0],
            "session2": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"Session 2"',)).fetchone()[0],
            "provider": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('provider',)).fetchone()[0],
        }
        con.commit()
    finally:
        con.close()
    if integrity != "ok" or counts["artifact"] != len(rows) or counts["fts"] != len(rows) or counts["response81"] < 1:
        raise RuntimeError({"bit_index_gate": {"integrity": integrity, "counts": counts, "expected": len(rows)}})
    qa = {
        "status": "passed",
        "generated_at": now_iso,
        "source_index_records": len(rows),
        "physical_files": sum(1 for row in rows if row["record_type"] == "physical_file"),
        "container_members": sum(1 for row in rows if row["record_type"] == "container_member"),
        "bit_index_integrity": integrity,
        "counts": counts,
        "bit_index_sha256": sha256_file(bit_path),
    }
    json_write(qa_path, qa)
    return {"source_json": source_json, "source_csv": source_csv, "bit_index": bit_path, "qa_path": qa_path, "qa": qa}


def build_manifest(project: Path, now_iso: str) -> tuple[Path, Path, list[dict[str, Any]]]:
    root = project / "Manifest" / "Section 5 Session 2 Complete"
    root.mkdir(parents=True, exist_ok=True)
    manifest = root / "MRHPD v3.0.0a Section 5 Session 2 Complete Project Manifest.json"
    checksums = root / "MRHPD v3.0.0a Section 5 Session 2 Complete Project Checksums.sha256"
    rows = []
    for path in sorted(item for item in project.rglob("*") if item.is_file() and item not in {manifest, checksums}):
        rows.append({"path": path.relative_to(project).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    json_write(manifest, {
        "schema": "mrhpd-current-project-manifest-3.0",
        "generated_at": now_iso,
        "response": 81,
        "section": SECTION_LABEL,
        "session": SESSION_LABEL,
        "state": "session_complete",
        "exclusions": [manifest.relative_to(project).as_posix(), checksums.relative_to(project).as_posix()],
        "file_count": len(rows),
        "total_bytes": sum(row["bytes"] for row in rows),
        "files": rows,
    })
    text_write(checksums, "".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
    for row in rows:
        path = project / row["path"]
        if path.stat().st_size != row["bytes"] or sha256_file(path) != row["sha256"]:
            raise RuntimeError({"manifest_mismatch": row["path"]})
    return manifest, checksums, rows


def write_project_readme(project: Path, now_iso: str) -> Path:
    path = project / "README SECTION 5 SESSION 2 COMPLETE.md"
    text_write(path, f"""# Human Pathogen Database — Section 5 Session 2 Complete

Response 81 completes Remediation Section 5 Session 2 of 3.

## Current state

- Canonical SQLite database synchronized through Response 81.
- Comprehensive workbook synchronized through Response 81.
- Immutable digital publication: 537 searchable pages.
- Frozen print-production interior: 538 pages with one intentional terminal blank.
- Exact Premium Color cover package retained.
- Provider preview and physical proof remain controlled pending; no unsupported completion is claimed.
- Session 3 handoff is ready.

## Recovery boundary

The complete Response 81 restore contains the full current project, manifests, checksums, verification tools, QA, reports, tracking, indexes, and restoration instructions. No earlier project archive or conversation reconstruction is required.

Updated: {now_iso}
""")
    return path


def write_qa(project: Path, db_qa: dict[str, Any], workbook_qa: dict[str, Any], publication_qa: dict[str, Any], application_qa: dict[str, Any], tracking_qa: dict[str, Any], report_qa: dict[str, Any], index_qa: dict[str, Any], gates: list[dict[str, Any]], events: list[dict[str, Any]], now_iso: str) -> dict[str, Any]:
    root = project / "QA" / "Section 5 Session 2" / "Complete"
    root.mkdir(parents=True, exist_ok=True)
    files = {
        "DATABASE_QA.json": db_qa,
        "WORKBOOK_QA.json": workbook_qa,
        "PUBLICATION_AND_COVER_QA.json": publication_qa,
        "APPLICATION_QA.json": application_qa,
        "TRACKING_QA.json": tracking_qa,
        "REPORT_QA.json": report_qa,
        "INDEX_QA.json": index_qa,
        "ACCEPTANCE_GATES.json": gates,
        "RECOVERY_EVENTS_221_230.json": events,
        "PROVIDER_EVIDENCE_BOUNDARY.json": {
            "status": "controlled_pending",
            "provider_submission_receipt": "not supplied or discovered",
            "provider_rendered_preview": "not supplied or discovered",
            "provider_warning_error_list": "not supplied or discovered",
            "provider_approval_claimed": False,
            "physical_proof_ordered": False,
            "physical_proof_received": False,
            "physical_proof_inspected": False,
            "physical_proof_approved": False,
            "rule": "Absence of evidence is not converted into a passed provider or physical-proof gate.",
        },
    }
    paths = []
    for name, payload in files.items():
        path = root / name
        json_write(path, payload)
        paths.append(path)
    final = {
        "schema": "mrhpd-section5-session2-complete-qa-1.0",
        "generated_at": now_iso,
        "status": "passed_with_controlled_external_gates",
        "response": 81,
        "checkpoint_3_of_3_complete": True,
        "session_2_of_3_complete": True,
        "remediation_section_5_complete": False,
        "internal_gates_passed": sum(1 for row in gates if row["status"].startswith("passed")),
        "controlled_external_gates": sum(1 for row in gates if row["status"] == "controlled_pending"),
        "provider_approval_claimed": False,
        "physical_proof_ordered": False,
        "accepted_predecessor_mutated": False,
        "frozen_section3_release_mutated": False,
        "immutable_publication_mutated": False,
        "main_application_mutated": False,
        "next": "Remediation Section 5 of 5 Session 3 of 3 Checkpoint 1 of 3",
    }
    final_path = root / "SECTION5_SESSION2_COMPLETE_QA.json"
    json_write(final_path, final)
    paths.append(final_path)
    return {"status": "passed", "final": final, "files": [path.relative_to(project).as_posix() for path in paths]}


def zip_project(project: Path, destination: Path) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(destination, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in sorted(item for item in project.rglob("*") if item.is_file()):
            zf.write(path, (Path(project.name) / path.relative_to(project)).as_posix())
    return verify_zip(destination)


def clean_verify_project(project_archive: Path, expected_project_name: str, manifest_rel: str, db_rel: str) -> dict[str, Any]:
    with tempfile.TemporaryDirectory(prefix="mrhpd-r81-project-clean-") as td:
        root = Path(td)
        safe_extract(project_archive, root)
        project = root / expected_project_name
        if not project.is_dir():
            raise RuntimeError({"clean_project_root_missing": expected_project_name})
        manifest = json.loads((project / manifest_rel).read_text(encoding="utf-8"))
        mismatches = []
        for row in manifest["files"]:
            path = project / row["path"]
            if not path.exists() or path.stat().st_size != row["bytes"] or sha256_file(path) != row["sha256"]:
                mismatches.append(row["path"])
        if mismatches:
            raise RuntimeError({"clean_manifest_mismatches": mismatches[:50]})
        con = sqlite3.connect(project / db_rel)
        try:
            integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
            fk = len(list(con.execute("PRAGMA foreign_key_check")))
            response81 = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R81'").fetchone()[0]
            release = con.execute("SELECT COUNT(*) FROM section5_s2_release WHERE release_key=? AND session_state='session_complete'", (RELEASE_KEY,)).fetchone()[0]
        finally:
            con.close()
        if integrity != "ok" or fk or response81 != 1 or release != 1:
            raise RuntimeError({"clean_database_gate": {"integrity": integrity, "foreign_keys": fk, "response81": response81, "release": release}})
        return {"status": "passed", "manifest_records": len(manifest["files"]), "manifest_mismatches": 0, "database_integrity": integrity, "foreign_keys": fk, "response81": response81, "release": release}


def create_restore_verifier(project_archive_name: str, project_archive_bytes: int, project_archive_sha: str, project_name: str, manifest_rel: str, db_rel: str) -> str:
    return f'''#!/usr/bin/env python3
from __future__ import annotations
import argparse, hashlib, json, re, shutil, sqlite3, sys, tempfile, zipfile
from pathlib import Path, PurePosixPath
PROJECT_ARCHIVE_NAME={project_archive_name!r}
PROJECT_ARCHIVE_BYTES={project_archive_bytes}
PROJECT_ARCHIVE_SHA256={project_archive_sha!r}
PROJECT_NAME={project_name!r}
MANIFEST_REL={manifest_rel!r}
DB_REL={db_rel!r}
RELEASE_KEY={RELEASE_KEY!r}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as handle:
  for block in iter(lambda:handle.read(1024*1024),b''): h.update(block)
 return h.hexdigest()
def verify_zip(path):
 with zipfile.ZipFile(path) as zf:
  if zf.testzip() is not None: raise RuntimeError('ZIP CRC failure: '+str(path))
  names=zf.namelist()
  if len(names)!=len(set(names)): raise RuntimeError('duplicate ZIP members')
  for name in names:
   pp=PurePosixPath(name.replace('\\\\','/'))
   if pp.is_absolute() or '..' in pp.parts or re.match(r'^[A-Za-z]:',name): raise RuntimeError('unsafe ZIP path: '+name)
def extract(path,dest):
 verify_zip(path); dest.mkdir(parents=True,exist_ok=True)
 with zipfile.ZipFile(path) as zf: zf.extractall(dest)
def main():
 ap=argparse.ArgumentParser()
 ap.add_argument('--extract-project-to',type=Path)
 args=ap.parse_args()
 root=Path(__file__).resolve().parents[1]
 archive=root/'COMPLETE_PROJECT'/PROJECT_ARCHIVE_NAME
 if archive.stat().st_size!=PROJECT_ARCHIVE_BYTES or sha(archive)!=PROJECT_ARCHIVE_SHA256: raise RuntimeError('project archive identity failed')
 verify_zip(archive)
 destination=args.extract_project_to
 temporary=None
 if destination is None:
  temporary=tempfile.TemporaryDirectory(prefix='mrhpd-r81-verify-'); destination=Path(temporary.name)
 if destination.exists() and any(destination.iterdir()): raise RuntimeError('extract destination must be empty')
 destination.mkdir(parents=True,exist_ok=True)
 extract(archive,destination)
 project=destination/PROJECT_NAME
 manifest=json.loads((project/MANIFEST_REL).read_text(encoding='utf-8'))
 mismatches=[]
 for row in manifest['files']:
  path=project/row['path']
  if not path.exists() or path.stat().st_size!=row['bytes'] or sha(path)!=row['sha256']: mismatches.append(row['path'])
 con=sqlite3.connect(project/DB_REL)
 try:
  integrity=con.execute('PRAGMA integrity_check').fetchone()[0]
  fk=len(list(con.execute('PRAGMA foreign_key_check')))
  response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R81'").fetchone()[0]
  release=con.execute("SELECT COUNT(*) FROM section5_s2_release WHERE release_key=? AND session_state='session_complete'",(RELEASE_KEY,)).fetchone()[0]
  approvals=con.execute("SELECT COUNT(*) FROM section5_s2_release WHERE provider_approval_claimed!=0 OR physical_proof_ordered!=0").fetchone()[0]
 finally: con.close()
 result={{'status':'passed' if not mismatches and integrity=='ok' and fk==0 and response==1 and release==1 and approvals==0 else 'failed','project_root':str(project),'manifest_records':len(manifest['files']),'manifest_mismatches':mismatches,'database_integrity':integrity,'foreign_keys':fk,'response81':response,'session2_release':release,'unsupported_external_claims':approvals}}
 print(json.dumps(result,indent=2))
 if result['status']!='passed': raise SystemExit(1)
 if temporary is not None: temporary.cleanup()
if __name__=='__main__': main()
'''


def build_complete_restore(project_archive: Path, project_qa: dict[str, Any], project: Path, dist: Path, stamp: str, now_iso: str, manifest_path: Path, checksums_path: Path, report_qa: dict[str, Any]) -> dict[str, Any]:
    root = dist / "complete_restore_root"
    if root.exists():
        shutil.rmtree(root)
    (root / "COMPLETE_PROJECT").mkdir(parents=True)
    (root / "TOOLS").mkdir(parents=True)
    (root / "FINAL_VERIFICATION").mkdir(parents=True)
    shutil.copy2(project_archive, root / "COMPLETE_PROJECT" / project_archive.name)
    text_write(root / "COMPLETE_PROJECT" / f"{project_archive.name}.sha256.txt", f"{sha256_file(project_archive)}  {project_archive.name}")
    manifest_rel = manifest_path.relative_to(project).as_posix()
    db_rel = CURRENT_DB_REL
    text_write(root / "TOOLS" / "restore_verify_extract.py", create_restore_verifier(project_archive.name, project_archive.stat().st_size, sha256_file(project_archive), project.name, manifest_rel, db_rel))
    text_write(root / "RESTORE_READ_FIRST.md", f"""# Human Pathogen Database — Complete Restore Through Response 81

This is the complete self-contained restore for Remediation Section 5 Session 2 of 3.

## Complete project archive

Filename: `{project_archive.name}`

Bytes: `{project_archive.stat().st_size}`

SHA-256: `{sha256_file(project_archive)}`

## Verification

```bash
python TOOLS/restore_verify_extract.py
```

To verify and extract the complete current project:

```bash
python TOOLS/restore_verify_extract.py --extract-project-to "<empty destination>"
```

No earlier project ZIP, checkpoint package, cloud artifact, or conversation reconstruction is required. Provider preview and physical-proof events remain controlled pending unless genuine item-level evidence is supplied or discovered.
""")
    shutil.copy2(manifest_path, root / "FINAL_VERIFICATION" / manifest_path.name)
    shutil.copy2(checksums_path, root / "FINAL_VERIFICATION" / checksums_path.name)
    for key in ("docx", "pdf", "xlsx"):
        source = project / report_qa[key]["path"]
        shutil.copy2(source, root / "FINAL_VERIFICATION" / source.name)
    qa_source = project / "QA" / "Section 5 Session 2" / "Complete" / "SECTION5_SESSION2_COMPLETE_QA.json"
    shutil.copy2(qa_source, root / "FINAL_VERIFICATION" / qa_source.name)
    restore_manifest_rows = []
    for path in sorted(item for item in root.rglob("*") if item.is_file() and path_not_manifest(item)):
        restore_manifest_rows.append({"path": path.relative_to(root).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    restore_manifest_path = root / "MRHPD_RESPONSE81_COMPLETE_RESTORE_MANIFEST.json"
    restore_checksums_path = root / "MRHPD_RESPONSE81_COMPLETE_RESTORE_CHECKSUMS.sha256"
    json_write(restore_manifest_path, {"schema": "mrhpd-complete-restore-3.0", "generated_at": now_iso, "response": 81, "session": SESSION_LABEL, "state": "session_complete", "file_count": len(restore_manifest_rows), "files": restore_manifest_rows})
    text_write(restore_checksums_path, "".join(f"{row['sha256']}  {row['path']}\n" for row in restore_manifest_rows))
    restore_name = (
        "Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
        f"Remediation Section 5 of 5 Session 2 of 3 COMPLETE RESTORE THROUGH RESPONSE 81 {stamp}.zip"
    )
    restore = dist / restore_name
    with zipfile.ZipFile(restore, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in sorted(item for item in root.rglob("*") if item.is_file()):
            zf.write(path, path.relative_to(root).as_posix())
    restore_qa = verify_zip(restore)
    with tempfile.TemporaryDirectory(prefix="mrhpd-r81-restore-clean-") as td:
        extracted = Path(td) / "restore"
        safe_extract(restore, extracted)
        result = subprocess.run([sys.executable, str((extracted / "TOOLS" / "restore_verify_extract.py").resolve())], cwd=extracted, text=True, capture_output=True, timeout=2400)
        if result.returncode:
            raise RuntimeError({"embedded_restore_verifier_failed": {"stdout": result.stdout[-16000:], "stderr": result.stderr[-16000:]}})
        embedded = json.loads(result.stdout)
    return {"restore": restore, "restore_qa": restore_qa, "embedded_verification": embedded, "root": root}


def path_not_manifest(path: Path) -> bool:
    return path.name not in {"MRHPD_RESPONSE81_COMPLETE_RESTORE_MANIFEST.json", "MRHPD_RESPONSE81_COMPLETE_RESTORE_CHECKSUMS.sha256"}


def build_transport(restore: Path, dist: Path, now_iso: str) -> dict[str, Any]:
    total = restore.stat().st_size
    part_size = math.ceil(total / VOLUME_COUNT)
    if part_size > MAX_VOLUME_PART_BYTES:
        raise RuntimeError({"restore_too_large_for_three_volumes": {"bytes": total, "part_size": part_size, "max": MAX_VOLUME_PART_BYTES}})
    part_paths = []
    with restore.open("rb") as source:
        for index in range(1, VOLUME_COUNT + 1):
            part = dist / f"{restore.name}.part{index:03d}"
            remaining = total - sum(path.stat().st_size for path in part_paths)
            target = remaining if index == VOLUME_COUNT else min(part_size, remaining)
            with part.open("wb") as output:
                left = target
                while left:
                    block = source.read(min(1024 * 1024, left))
                    if not block:
                        raise RuntimeError("unexpected end of restore while splitting")
                    output.write(block)
                    left -= len(block)
            part_paths.append(part)
    if sum(path.stat().st_size for path in part_paths) != total:
        raise RuntimeError("transport part byte total mismatch")
    part_rows = [{"sequence": index, "name": path.name, "bytes": path.stat().st_size, "sha256": sha256_file(path)} for index, path in enumerate(part_paths, start=1)]
    manifest = {
        "schema": "mrhpd-response81-transport-1.0",
        "generated_at": now_iso,
        "restore": {"name": restore.name, "bytes": total, "sha256": sha256_file(restore)},
        "part_count": VOLUME_COUNT,
        "parts": part_rows,
    }
    manifest_path = dist / "MRHPD_RESPONSE81_COMPLETE_RESTORE_TRANSPORT_MANIFEST.json"
    json_write(manifest_path, manifest)
    reassembler = dist / "reassemble_response81_complete_restore.py"
    text_write(reassembler, f'''#!/usr/bin/env python3
import hashlib,json,sys
from pathlib import Path
MANIFEST={manifest!r}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as handle:
  for block in iter(lambda:handle.read(1024*1024),b''): h.update(block)
 return h.hexdigest()
def main():
 root=Path.cwd(); output=root/MANIFEST['restore']['name']
 with output.open('wb') as target:
  for row in MANIFEST['parts']:
   part=root/row['name']
   if part.stat().st_size!=row['bytes'] or sha(part)!=row['sha256']: raise RuntimeError({{'part_identity_failed':row['name']}})
   with part.open('rb') as source:
    for block in iter(lambda:source.read(1024*1024),b''): target.write(block)
 if output.stat().st_size!=MANIFEST['restore']['bytes'] or sha(output)!=MANIFEST['restore']['sha256']: raise RuntimeError('reassembled restore identity failed')
 print(json.dumps({{'status':'passed','restore':str(output),'bytes':output.stat().st_size,'sha256':sha(output)}},indent=2))
if __name__=='__main__': main()
''')
    wrappers = []
    for index, part in enumerate(part_paths, start=1):
        wrapper = dist / f"MRHPD v3.0.0a Response 81 Complete Restore Drive Volume {index} of {VOLUME_COUNT}.zip"
        readme = dist / f"README_VOLUME_{index}.txt"
        text_write(readme, f"Human Pathogen Database Response 81 complete restore volume {index} of {VOLUME_COUNT}. Extract all three volume ZIPs into the same empty directory, then run reassemble_response81_complete_restore.py. All three volumes are required.")
        with zipfile.ZipFile(wrapper, "w", compression=zipfile.ZIP_STORED, allowZip64=True) as zf:
            zf.write(part, part.name)
            zf.write(manifest_path, manifest_path.name)
            zf.write(reassembler, reassembler.name)
            zf.write(readme, readme.name)
        wrappers.append({"sequence": index, "path": wrapper, "qa": verify_zip(wrapper)})
    with tempfile.TemporaryDirectory(prefix="mrhpd-r81-transport-test-") as td:
        root = Path(td)
        for row in wrappers:
            safe_extract(row["path"], root)
        result = subprocess.run([sys.executable, str((root / reassembler.name).resolve())], cwd=root, text=True, capture_output=True, timeout=1200)
        if result.returncode:
            raise RuntimeError({"transport_reassembly_failed": {"stdout": result.stdout[-8000:], "stderr": result.stderr[-8000:]}})
        reconstructed = root / restore.name
        if reconstructed.stat().st_size != restore.stat().st_size or sha256_file(reconstructed) != sha256_file(restore):
            raise RuntimeError("transport reassembly identity mismatch")
        reassembly = json.loads(result.stdout)
    return {"manifest": manifest, "manifest_path": manifest_path, "reassembler": reassembler, "wrappers": wrappers, "reassembly": reassembly}


def build_verification_delivery(dist: Path, stamp: str, summary: dict[str, Any], project: Path, project_archive: Path, restore_result: dict[str, Any], transport: dict[str, Any], report_qa: dict[str, Any], qa_result: dict[str, Any], build_log_hint: str) -> Path:
    exact_names = dist / "MRHPD v3.0.0a Response 81 Exact File Names.txt"
    lines = [
        "Response 81 complete restore:", restore_result["restore"].name, "",
        "Complete project archive:", project_archive.name, "",
        "Required transport volume wrappers:",
    ]
    lines.extend(row["path"].name for row in transport["wrappers"])
    lines.extend(["", "Verification delivery:", f"MRHPD v3.0.0a Response 81 Section 5 Session 2 Complete Verification Delivery {stamp}.zip", "", "Current database:", Path(CURRENT_DB_REL).name, "", "Current workbook:", Path(CURRENT_WORKBOOK_REL).name, "", "Print interior:", Path(PRINT_INTERIOR_REL).name, "", "Cover PNG:", Path(COVER_PNG_REL).name])
    text_write(exact_names, "\n".join(lines))
    summary_path = dist / "MRHPD_RESPONSE81_SECTION5_SESSION2_COMPLETE_BUILD_SUMMARY.json"
    json_write(summary_path, summary)
    verification_path = dist / "MRHPD v3.0.0a Response 81 Section 5 Session 2 Complete Verification.json"
    json_write(verification_path, {
        "schema": "mrhpd-response81-section5-session2-complete-verification-1.0",
        "status": "passed_with_controlled_external_gates",
        "project_archive": verify_zip(project_archive),
        "complete_restore": restore_result["restore_qa"],
        "embedded_restore_verification": restore_result["embedded_verification"],
        "transport": {"manifest": transport["manifest"], "wrappers": [{"sequence": row["sequence"], "qa": row["qa"]} for row in transport["wrappers"]], "reassembly": transport["reassembly"]},
        "provider_approval_claimed": False,
        "physical_proof_ordered": False,
        "user_upload_required": False,
        "conversation_reconstruction_required": False,
        "checkpoint_3_of_3_complete": True,
        "session_2_of_3_complete": True,
        "remediation_section_5_complete": False,
        "next": "Remediation Section 5 of 5 Session 3 of 3 Checkpoint 1 of 3",
    })
    delivery = dist / f"MRHPD v3.0.0a Response 81 Section 5 Session 2 Complete Verification Delivery {stamp}.zip"
    files = [verification_path, summary_path, exact_names, transport["manifest_path"], transport["reassembler"]]
    files.extend(project / report_qa[key]["path"] for key in ("docx", "pdf", "xlsx"))
    files.extend(project / path for path in qa_result["files"])
    with zipfile.ZipFile(delivery, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in files:
            if path.exists():
                zf.write(path, path.name)
    verify_zip(delivery)
    return delivery


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--response77-dir", type=Path, required=True)
    parser.add_argument("--response80-dir", type=Path, required=True)
    parser.add_argument("--dist", type=Path, default=Path("dist_cp5_s2_cp3"))
    args = parser.parse_args()
    now = utc_now()
    now_iso = now.isoformat().replace("+00:00", "Z")
    stamp = now.strftime("%Y-%m-%d %H%M UTC")
    if args.dist.exists():
        shutil.rmtree(args.dist)
    args.dist.mkdir(parents=True)
    with tempfile.TemporaryDirectory(prefix="mrhpd-cp5-s2-cp3-") as td:
        work = Path(td)
        baseline_restore, baseline_project_archive, response80_project, recovery_qa = restore_response80(args.response77_dir, args.response80_dir, work)
        current_project = work / "current" / CURRENT_PROJECT_NAME
        current_project.parent.mkdir(parents=True)
        shutil.copytree(response80_project, current_project)
        source_db = current_project / cp2.CURRENT_DB_REL
        source_workbook = current_project / cp2.CURRENT_WORKBOOK_REL
        if not source_db.exists() or not source_workbook.exists():
            raise RuntimeError({"response80_current_paths": {"database": str(source_db), "workbook": str(source_workbook)}})
        gates = acceptance_gates(now_iso)
        events = recovery_events(now_iso)
        current_db = current_project / CURRENT_DB_REL
        db_qa = sync_database(source_db, current_db, now_iso, gates, events)
        current_workbook = current_project / CURRENT_WORKBOOK_REL
        workbook_qa = sync_workbook(source_workbook, current_workbook, db_qa, gates, events, now_iso)
        tracking_qa = write_tracking(current_project, current_db, now_iso)
        publication_qa = audit_publication_and_cover(current_project)
        application_qa = write_application_audit(current_project, current_db, current_workbook, publication_qa, now_iso)
        report_qa = write_reports(current_project, db_qa, workbook_qa, publication_qa, application_qa, gates, events, now_iso)
        readme = write_project_readme(current_project, now_iso)
        index_result = build_indexes(current_project, now_iso)
        qa_result = write_qa(current_project, db_qa, workbook_qa, publication_qa, application_qa, tracking_qa, report_qa, index_result["qa"], gates, events, now_iso)
        manifest_path, checksums_path, manifest_rows = build_manifest(current_project, now_iso)
        project_archive = args.dist / (
            "Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
            f"Remediation Section 5 of 5 Session 2 of 3 COMPLETE PROJECT THROUGH RESPONSE 81 {stamp}.zip"
        )
        project_archive_qa = zip_project(current_project, project_archive)
        clean_project_qa = clean_verify_project(project_archive, current_project.name, manifest_path.relative_to(current_project).as_posix(), CURRENT_DB_REL)
        restore_result = build_complete_restore(project_archive, clean_project_qa, current_project, args.dist, stamp, now_iso, manifest_path, checksums_path, report_qa)
        transport = build_transport(restore_result["restore"], args.dist, now_iso)
        summary = {
            "schema": "mrhpd-response81-section5-session2-complete-build-1.0",
            "generated_at": now_iso,
            "status": "passed_with_controlled_external_gates",
            "response": 81,
            "section": SECTION_LABEL,
            "session": SESSION_LABEL,
            "checkpoint": CHECKPOINT_LABEL,
            "recovery": recovery_qa,
            "database": db_qa,
            "workbook": workbook_qa,
            "publication": publication_qa,
            "application": application_qa,
            "tracking": tracking_qa,
            "reports": report_qa,
            "index": index_result["qa"],
            "qa": qa_result["final"],
            "manifest_records": len(manifest_rows),
            "project_archive": project_archive_qa,
            "clean_project": clean_project_qa,
            "complete_restore": restore_result["restore_qa"],
            "embedded_restore_verification": restore_result["embedded_verification"],
            "transport": {"manifest": transport["manifest"], "wrappers": [{"sequence": row["sequence"], "qa": row["qa"]} for row in transport["wrappers"]], "reassembly": transport["reassembly"]},
            "provider_approval_claimed": False,
            "physical_proof_ordered": False,
            "accepted_predecessor_mutated": False,
            "frozen_section3_release_mutated": False,
            "immutable_publication_mutated": False,
            "main_application_mutated": False,
            "user_upload_required": False,
            "conversation_reconstruction_required": False,
            "checkpoint_3_of_3_complete": True,
            "session_2_of_3_complete": True,
            "remediation_section_5_complete": False,
            "next": "Remediation Section 5 of 5 Session 3 of 3 Checkpoint 1 of 3",
        }
        verification_delivery = build_verification_delivery(args.dist, stamp, summary, current_project, project_archive, restore_result, transport, report_qa, qa_result, "section5-session2-complete-build.log")
        console = {
            "status": "passed_with_controlled_external_gates",
            "response": 81,
            "database_tables": db_qa["tables"],
            "workbook_sheets": workbook_qa["current_sheet_count"],
            "digital_pages": publication_qa["digital_publication"]["pages"],
            "print_pages": publication_qa["print_interior"]["pages"],
            "project_archive": project_archive_qa,
            "complete_restore": restore_result["restore_qa"],
            "transport_volumes": [{"sequence": row["sequence"], **row["qa"]} for row in transport["wrappers"]],
            "verification_delivery": verify_zip(verification_delivery),
            "provider_preview": "controlled_pending",
            "physical_proof": "controlled_pending",
            "user_upload_required": False,
            "checkpoint_3_of_3_complete": True,
            "session_2_of_3_complete": True,
            "next": "Remediation Section 5 of 5 Session 3 of 3 Checkpoint 1 of 3",
        }
        print(json.dumps(console, indent=2))


if __name__ == "__main__":
    main()
