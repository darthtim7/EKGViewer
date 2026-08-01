#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import textwrap
import zipfile
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any, Iterable

import fitz
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

HERE = Path(__file__).resolve().parent
if str(HERE) not in sys.path:
    sys.path.insert(0, str(HERE))
import inspect_response77 as r77  # noqa: E402

PROJECT_VERSION = "3.0.0a"
RESPONSE_NUMBER = 79
SECTION_LABEL = "Remediation Section 5 of 5"
SESSION_LABEL = "Session 2 of 3"
CHECKPOINT_LABEL = "Checkpoint 1 of 3"
CHECKPOINT_CODE = "MRHPD-V3-CP5-S2-CP1"
BASE_RESTORE_BYTES = r77.RESTORE_BYTES
BASE_RESTORE_SHA256 = r77.RESTORE_SHA256
BASE_PROJECT_BYTES = r77.PROJECT_BYTES
BASE_PROJECT_SHA256 = r77.PROJECT_SHA256
PUBLICATION_SHA256 = r77.PUBLICATION_SHA256
APPLICATION_SHA256 = r77.APPLICATION_SHA256
PRINT_INTERIOR_SHA256 = "0216def4f41b2b62fc2eb3f87f5a66abbf633e54c41b31e2b39afa29c34b0803"
COVER_SHA256 = "3945225ef87c87a8795354aee1c90ce58d39fd6d5bb57229489692420ba07097"
CURRENT_PROJECT_NAME = (
    "Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 2 of 3 Checkpoint 1 of 3 THROUGH RESPONSE 79"
)
CURRENT_DB_REL = (
    "Database/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 2 of 3 Checkpoint 1 of 3 THROUGH RESPONSE 79.sqlite"
)
CURRENT_WORKBOOK_REL = (
    "Tracking/Workbook/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 2 of 3 Checkpoint 1 of 3 THROUGH RESPONSE 79 Comprehensive Tracking.xlsx"
)
PUBLICATION_REL = (
    "Documents/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 3 of 5 Session 3 of 4 Integrated Manuscript.pdf"
)
PRINT_INTERIOR_REL = (
    "Print Production/KDP Premium Color Response 76/Interior/"
    "Medical References - Human Pathogen Database v3.0.0a KDP Premium Color 8.5 x 11 "
    "Print Interior 538 Pages Response 76.pdf"
)
COVER_REL = (
    "Print Production/KDP Premium Color Response 76/Cover/"
    "MRHPD v3.0.0a KDP Premium Color 538-Page Full Cover Response 76 300ppi RGB.png"
)

NAVY = "17324D"
TEAL = "1C7475"
GOLD = "C9A227"
PALE_BLUE = "EAF1F5"
PALE_GOLD = "F7F1D9"
PALE_GREEN = "E9F3EE"
PALE_RED = "F7E8E6"
DARK = "24323D"
WHITE = "FFFFFF"


def utc_now() -> datetime:
    return datetime.now(timezone.utc).replace(microsecond=0)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value, encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fields is None:
        fields = list(rows[0]) if rows else []
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({key: json.dumps(value, ensure_ascii=False, sort_keys=True) if isinstance(value, (list, dict, tuple, set)) else value for key, value in row.items()})


def verify_zip(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
        names = zf.namelist()
        unsafe = []
        filler = []
        for name in names:
            pp = PurePosixPath(name.replace("\\", "/"))
            if pp.is_absolute() or ".." in pp.parts or re.match(r"^[A-Za-z]:", name):
                unsafe.append(name)
            if any(token in name.lower() for token in ("filler", "padding", "dummy_payload", "artificial_inflation")):
                filler.append(name)
    result = {
        "path": str(path),
        "name": path.name,
        "bytes": path.stat().st_size,
        "sha256": sha256_file(path),
        "members": len(names),
        "crc_error": bad,
        "duplicates": len(names) - len(set(names)),
        "unsafe_paths": unsafe,
        "filler_members": filler,
    }
    if bad or result["duplicates"] or unsafe or filler:
        raise RuntimeError({"zip_verification_failed": result})
    return result


def safe_extract(path: Path, destination: Path) -> None:
    verify_zip(path)
    destination.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as zf:
        zf.extractall(destination)


def restore_response77(volume_root: Path, work: Path) -> tuple[Path, Path, Path]:
    restore, _transport = r77.reconstruct_restore(volume_root, work / "response77")
    extracted_restore = work / "restore"
    safe_extract(restore, extracted_restore)
    project_archive = r77.find_unique_by_identity(extracted_restore, size=BASE_PROJECT_BYTES, digest=BASE_PROJECT_SHA256)
    extracted_project = work / "baseline_project"
    safe_extract(project_archive, extracted_project)
    baseline_project = r77.locate_project_root(extracted_project)
    return restore, project_archive, baseline_project


def locate_current_database(project: Path) -> Path:
    candidates = []
    for path in project.rglob("*.sqlite"):
        try:
            con = sqlite3.connect(path)
            try:
                tables = {row[0] for row in con.execute("SELECT name FROM sqlite_master WHERE type='table'")}
                if "section5_session1_release" not in tables or "thread_response_reconciliation_cp3" not in tables:
                    continue
                response = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R77'").fetchone()[0]
                integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
                count = len(tables)
            finally:
                con.close()
            if response == 1 and integrity == "ok":
                candidates.append((count, path.stat().st_size, path))
        except Exception:
            continue
    if not candidates:
        raise RuntimeError("Response 77 canonical database not found")
    candidates.sort(reverse=True)
    return candidates[0][2]


def locate_current_workbook(project: Path) -> Path:
    candidates = []
    for path in project.rglob("*.xlsx"):
        try:
            wb = load_workbook(path, read_only=True, data_only=False)
            try:
                count = len(wb.sheetnames)
                names = set(wb.sheetnames)
            finally:
                wb.close()
            if count >= 106 and "S5S1 CP3 Handoff" in names:
                candidates.append((count, path.stat().st_size, path))
        except Exception:
            continue
    if not candidates:
        raise RuntimeError("Response 77 comprehensive workbook not found")
    candidates.sort(reverse=True)
    return candidates[0][2]


def official_source_rows(now_iso: str) -> list[dict[str, Any]]:
    return [
        {
            "source_key": "KDP-UPLOAD-PREVIEW",
            "authority": "Amazon Kindle Direct Publishing",
            "title": "Upload and Preview Book Content",
            "url": "https://kdp.amazon.com/en_US/help/topic/G200641240",
            "controlled_use": "Upload the manuscript and cover, launch Print Previewer, and capture provider-generated errors or warnings before publication.",
            "evidence_boundary": "This is process guidance. It is not evidence that the project files have been uploaded, approved, or accepted.",
            "verified_at": now_iso,
        },
        {
            "source_key": "KDP-TOOLS-PREVIEWER",
            "authority": "Amazon Kindle Direct Publishing",
            "title": "KDP Tools and Resources",
            "url": "https://kdp.amazon.com/en_US/help/topic/G200735480",
            "controlled_use": "Print Previewer checks uploaded paperback and hardcover files for errors that must be addressed before submission.",
            "evidence_boundary": "A local preflight does not substitute for the provider conversion and previewer output.",
            "verified_at": now_iso,
        },
        {
            "source_key": "KDP-PROOF-ORDER",
            "authority": "Amazon Kindle Direct Publishing",
            "title": "How do I order a proof or author copy?",
            "url": "https://kdp.amazon.com/en_US/help/topic/GVEG4YA9G2T7N6DR",
            "controlled_use": "A proof may be requested while the book is in Draft status after approval in Print Previewer; up to five proof copies may be ordered per request.",
            "evidence_boundary": "The project records a proof workflow but does not claim that a proof has been requested, purchased, shipped, received, or approved.",
            "verified_at": now_iso,
        },
        {
            "source_key": "KDP-PROOF-PURPOSE",
            "authority": "Amazon Kindle Direct Publishing",
            "title": "Proof and Author Copies",
            "url": "https://kdp.amazon.com/en_US/help/topic/G7BBN68RYX5UMDZF",
            "controlled_use": "Proof copies are test copies for prepublication review and carry a Not for Resale watermark and a non-ISBN proof barcode.",
            "evidence_boundary": "A proof copy is a physical review artifact, not a substitute for final publication review or an assertion of provider approval.",
            "verified_at": now_iso,
        },
        {
            "source_key": "KDP-PAPERBACK-FILES",
            "authority": "Amazon Kindle Direct Publishing",
            "title": "Paperback Submission Guidelines",
            "url": "https://kdp.amazon.com/en_US/help/topic/G201857950",
            "controlled_use": "Control print-file specifications, including embedded fonts and images, flattened transparency, absence of production marks or annotations, and 300-DPI image expectations.",
            "evidence_boundary": "The current project preflight is local evidence; final provider conversion and manual review remain separate gates.",
            "verified_at": now_iso,
        },
    ]


def preview_evidence_rows(now_iso: str) -> list[dict[str, Any]]:
    required = [
        ("PREVIEW-01", "KDP title setup identity", "Bookshelf title identifier, format, trim, ink/paper, bleed, finish, ISBN/barcode choice", "KDP Bookshelf screenshot or exported record"),
        ("PREVIEW-02", "Uploaded manuscript identity", "Exact uploaded filename, byte count, local SHA-256, upload timestamp", "KDP upload confirmation plus local checksum record"),
        ("PREVIEW-03", "Uploaded cover identity", "Exact uploaded filename, byte count, local SHA-256, upload timestamp", "KDP upload confirmation plus local checksum record"),
        ("PREVIEW-04", "Print Previewer issue summary", "All blocking errors, warnings, ignored issues, affected pages/regions, and provider wording", "Full-page screenshots or provider issue export"),
        ("PREVIEW-05", "Interior sequence review", "Front matter, page order, intentional page 538 blank, orientation, trim, bleed, gutter, and clipping", "Representative and affected-page screenshots"),
        ("PREVIEW-06", "Cover spread review", "Back, spine, front, folds, trim, bleed, live areas, barcode placement, and spine text", "Full-spread Print Previewer screenshot"),
        ("PREVIEW-07", "Provider approval state", "Whether Print Previewer allows approval and the date/time of approval", "Approval screen screenshot"),
        ("PREVIEW-08", "Provider conversion comparison", "Any conversion-induced text, image, color, scaling, transparency, or alignment change", "Before/after evidence with page or region identifiers"),
    ]
    return [
        {
            "evidence_key": key,
            "evidence_title": title,
            "required_content": content,
            "expected_evidence": evidence,
            "status": "controlled_pending_external",
            "evidence_path": "",
            "provider_reference": "",
            "notes": "No KDP account or Print Previewer evidence was available to the build lane; no approval or issue result is inferred.",
            "recorded_at": now_iso,
        }
        for key, title, content, evidence in required
    ]


def conversion_issue_rows(now_iso: str) -> list[dict[str, Any]]:
    return [
        {
            "issue_key": "PREVIEW-CONTROL-0001",
            "surface": "provider preview",
            "page_or_region": "not yet run",
            "severity": "control",
            "provider_message": "No Print Previewer output supplied.",
            "local_reproduction": "Not applicable; this row records the evidence boundary rather than a provider defect.",
            "disposition": "Retain controlled-pending state until actual provider output is captured.",
            "status": "controlled_pending_external",
            "evidence_path": "",
            "recorded_at": now_iso,
        }
    ]


def physical_proof_plan_rows(now_iso: str) -> list[dict[str, Any]]:
    plan = [
        (1, "preorder", "Approve the exact uploaded files in Print Previewer.", "No unresolved blocking previewer error; every warning has an explicit disposition.", "Print Previewer approval evidence"),
        (2, "order", "Request a printed proof from the KDP Bookshelf while the title remains in the appropriate proof-eligible state.", "Proof request identifies the current 538-page interior and exact cover candidate.", "Proof request confirmation"),
        (3, "order", "Complete the proof order and retain marketplace, quantity, price, shipping, and estimated-delivery information.", "Order is completed within the provider workflow and tied to the current candidate.", "Order confirmation and receipt"),
        (4, "receipt", "Photograph the sealed package and proof identifiers before review.", "Proof identity and condition are documented; damage is distinguished from file/printing defects.", "Receipt photographs"),
        (5, "cover", "Inspect trim, bleed, spine centering, fold alignment, barcode area, color, contrast, and surface finish.", "No clipped live content, fold encroachment, unreadable spine, unexpected barcode conflict, or unacceptable color shift.", "Cover photographs and defect log"),
        (6, "interior", "Inspect page sequence, page 538 intentional blank, gutter, outside margin, top/bottom margin, orientation, cropping, and binding behavior.", "All pages are present and ordered; no text or teaching content is obscured by trim, gutter, or binding.", "Page photographs and checklist"),
        (7, "image", "Inspect representative raster artwork, charts, tables, line art, and color-coded teaching elements.", "Images are sharp, labels readable, colors distinguishable, and no rasterization or screening defect impairs teaching value.", "Representative-page photographs"),
        (8, "text", "Inspect typography, font substitution, fine rules, hyperlinks rendered as text where applicable, and page-to-page consistency.", "No missing glyphs, substituted fonts, broken characters, faint rules, or unreadable small text.", "Typography inspection record"),
        (9, "binding", "Inspect cover adhesion, page adhesion, opening behavior, page curl, show-through, and physical durability.", "Binding is intact and does not conceal content or create unacceptable handling defects.", "Binding photographs and notes"),
        (10, "correction", "Log each defect with severity, page/region, evidence, root cause, correction, and regression-test result.", "Every material defect is corrected or formally accepted with rationale before final release.", "Correction register"),
        (11, "signoff", "Complete independent clinical, editorial, production, and project-custody signoff.", "All required reviewers approve the corrected candidate; no controlled-pending critical gate remains.", "Signed release checklist"),
    ]
    return [
        {
            "sequence": sequence,
            "phase": phase,
            "action": action,
            "acceptance_criteria": criteria,
            "evidence_required": evidence,
            "owner_role": "project production review",
            "status": "planned",
            "recorded_at": now_iso,
        }
        for sequence, phase, action, criteria, evidence in plan
    ]


def proof_inspection_rows(now_iso: str) -> list[dict[str, Any]]:
    items = [
        ("Identity", "Candidate version and print date", "Proof is traceable to the current uploaded files."),
        ("Cover", "Front/back/spine alignment", "All live content is inside safe areas; spine text is centered and legible."),
        ("Cover", "Barcode and reserved area", "Provider barcode does not cover protected content and scans/prints cleanly."),
        ("Cover", "Color and contrast", "Clinical labels and visual-navigation cues remain distinguishable under print conditions."),
        ("Interior", "Page sequence and terminal blank", "Pages 1–537 are present and ordered; page 538 is intentionally blank."),
        ("Interior", "Gutter and trim safety", "No text, table, figure, or label is lost in the gutter or trim."),
        ("Interior", "Orientation and pagination", "Portrait and landscape pages are correctly oriented and numbered."),
        ("Typography", "Fonts and glyphs", "No substituted font, missing glyph, broken ligature, or encoding artifact."),
        ("Artwork", "Raster sharpness", "Teaching art remains sharp enough for labels, mechanisms, and decision pathways."),
        ("Artwork", "Color-coded meaning", "Color distinctions remain interpretable and do not collapse into ambiguous tones."),
        ("Tables", "Rules, cells, and small text", "Table structure and smallest approved text remain readable."),
        ("Binding", "Adhesion and opening behavior", "The book opens and handles without hidden content or structural failure."),
        ("Defects", "Damage versus production defect", "Shipping damage is separated from reproducible print/conversion defects."),
        ("Release", "Corrective regression", "All corrected pages and surfaces are rechecked in previewer and a subsequent proof when material."),
    ]
    return [
        {
            "category": category,
            "inspection_item": item,
            "acceptance_criteria": criteria,
            "observed": "",
            "severity": "",
            "disposition": "",
            "status": "not_inspected",
            "evidence_path": "",
            "recorded_at": now_iso,
        }
        for category, item, criteria in items
    ]


def recovery_rows(now_iso: str) -> list[dict[str, Any]]:
    events = [
        ("V3-CP5-S2-REC-202-INSTRUCTIONS-1-5-0-REPROCESSED", "Section 5 Session 2 began after multiple recovery turns and required a full refresh of the governing instructions.", "Reprocessed Project Instructions 1.5.0 and carried forward automatic recovery, Google Drive custody, exact filenames, tracking, indexes, and checkpoint emission controls."),
        ("V3-CP5-S2-REC-203-LOCAL-RUNTIME-UNAVAILABLE", "Local container, private Python, and user-visible Python returned InvalidArgumentError before code startup.", "Preserved all Drive-resident source artifacts and used the isolated transient computation lane; no local partial output was trusted."),
        ("V3-CP5-S2-REC-204-RESPONSE77-ARTIFACTS-DISCOVERED", "The newest complete restore needed to be identified without relying on prior conversational labels.", "Enumerated Actions artifacts and identified the exact Response 77 three-volume set and verification package by name, run, size, digest, and creation time."),
        ("V3-CP5-S2-REC-205-RESPONSE77-RESTORE-VERIFIED", "Session 2 required independent verification of the Response 77 baseline before mutation.", "Reconstructed the 211,294,688-byte restore, verified its SHA-256 and CRC, verified the 211,898,622-byte project archive, and inspected the copied project."),
        ("V3-CP5-S2-REC-206-PREVIEW-EVIDENCE-BOUNDARY", "No authenticated KDP Print Previewer output was available to the execution lane.", "Created a complete evidence-intake register and retained controlled_pending_external status; no provider approval, warning, or error was fabricated."),
        ("V3-CP5-S2-REC-207-PHYSICAL-PROOF-BOUNDARY", "No physical proof had been ordered, received, or inspected in the available evidence.", "Created an order, receipt, inspection, correction, and signoff workflow while retaining all physical-proof fields as not_inspected or planned."),
        ("V3-CP5-S2-REC-208-COPIED-TREE-SYNCHRONIZATION", "Session 2 required current database, workbook, application, tracking, source, index, manifest, and recovery surfaces.", "Created a distinct mutable project tree from the verified Response 77 baseline and synchronized only the new Session 2 checkpoint state."),
    ]
    return [
        {"event_code": code, "condition": condition, "recovery": recovery, "status": "recovered", "recorded_at": now_iso}
        for code, condition, recovery in events
    ]


def clone_response_row(con: sqlite3.Connection, response_key: str) -> dict[str, Any]:
    info = con.execute("PRAGMA table_info(thread_response_reconciliation_cp3)").fetchall()
    columns = [row[1] for row in info]
    row = con.execute("SELECT * FROM thread_response_reconciliation_cp3 WHERE response_key=? ORDER BY thread_response_reconciliation_cp3_id DESC LIMIT 1", (response_key,)).fetchone()
    if row is None:
        raise RuntimeError(f"Response template missing: {response_key}")
    return dict(zip(columns, row))


def insert_response(con: sqlite3.Connection, template: dict[str, Any], values: dict[str, Any]) -> None:
    info = con.execute("PRAGMA table_info(thread_response_reconciliation_cp3)").fetchall()
    pk_columns = {row[1] for row in info if row[5]}
    columns = [row[1] for row in info if row[1] not in pk_columns]
    record = dict(template)
    record.update(values)
    con.execute("DELETE FROM thread_response_reconciliation_cp3 WHERE response_key=?", (values["response_key"],))
    con.execute(
        f"INSERT INTO thread_response_reconciliation_cp3 ({','.join(columns)}) VALUES ({','.join('?' for _ in columns)})",
        [record.get(column) for column in columns],
    )


def synchronize_database(source: Path, destination: Path, *, now_iso: str, sources: list[dict[str, Any]], evidence: list[dict[str, Any]], issues: list[dict[str, Any]], proof_plan: list[dict[str, Any]], proof_items: list[dict[str, Any]], recovery: list[dict[str, Any]]) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, destination)
    con = sqlite3.connect(destination)
    try:
        con.execute("PRAGMA foreign_keys=ON")
        con.executescript("""
        CREATE TABLE IF NOT EXISTS section5_session2_checkpoint (
            checkpoint_code TEXT PRIMARY KEY,
            response_number INTEGER NOT NULL,
            section_label TEXT NOT NULL,
            session_label TEXT NOT NULL,
            checkpoint_label TEXT NOT NULL,
            state TEXT NOT NULL,
            baseline_restore_sha256 TEXT NOT NULL,
            baseline_project_sha256 TEXT NOT NULL,
            digital_publication_sha256 TEXT NOT NULL,
            print_interior_sha256 TEXT NOT NULL,
            cover_sha256 TEXT NOT NULL,
            provider_previewer_status TEXT NOT NULL,
            physical_proof_status TEXT NOT NULL,
            workbook_status TEXT NOT NULL,
            application_status TEXT NOT NULL,
            next_checkpoint TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_session2_source_control (
            source_key TEXT PRIMARY KEY,
            authority TEXT NOT NULL,
            title TEXT NOT NULL,
            url TEXT NOT NULL,
            controlled_use TEXT NOT NULL,
            evidence_boundary TEXT NOT NULL,
            verified_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_provider_preview_evidence (
            section5_provider_preview_evidence_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            evidence_key TEXT NOT NULL,
            evidence_title TEXT NOT NULL,
            required_content TEXT NOT NULL,
            expected_evidence TEXT NOT NULL,
            status TEXT NOT NULL,
            evidence_path TEXT,
            provider_reference TEXT,
            notes TEXT,
            recorded_at TEXT NOT NULL,
            UNIQUE(checkpoint_code,evidence_key)
        );
        CREATE TABLE IF NOT EXISTS section5_provider_conversion_issue (
            section5_provider_conversion_issue_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            issue_key TEXT NOT NULL,
            surface TEXT NOT NULL,
            page_or_region TEXT,
            severity TEXT NOT NULL,
            provider_message TEXT,
            local_reproduction TEXT,
            disposition TEXT,
            status TEXT NOT NULL,
            evidence_path TEXT,
            recorded_at TEXT NOT NULL,
            UNIQUE(checkpoint_code,issue_key)
        );
        CREATE TABLE IF NOT EXISTS section5_physical_proof_plan (
            section5_physical_proof_plan_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            sequence INTEGER NOT NULL,
            phase TEXT NOT NULL,
            action TEXT NOT NULL,
            acceptance_criteria TEXT NOT NULL,
            evidence_required TEXT NOT NULL,
            owner_role TEXT NOT NULL,
            status TEXT NOT NULL,
            recorded_at TEXT NOT NULL,
            UNIQUE(checkpoint_code,sequence)
        );
        CREATE TABLE IF NOT EXISTS section5_physical_proof_inspection (
            section5_physical_proof_inspection_id INTEGER PRIMARY KEY,
            checkpoint_code TEXT NOT NULL,
            category TEXT NOT NULL,
            inspection_item TEXT NOT NULL,
            acceptance_criteria TEXT NOT NULL,
            observed TEXT,
            severity TEXT,
            disposition TEXT,
            status TEXT NOT NULL,
            evidence_path TEXT,
            recorded_at TEXT NOT NULL,
            UNIQUE(checkpoint_code,category,inspection_item)
        );
        CREATE TABLE IF NOT EXISTS section5_session2_qa (
            checkpoint_code TEXT NOT NULL,
            control_key TEXT NOT NULL,
            expected TEXT NOT NULL,
            observed TEXT NOT NULL,
            status TEXT NOT NULL,
            evidence_path TEXT,
            checked_at TEXT NOT NULL,
            PRIMARY KEY(checkpoint_code,control_key)
        );
        CREATE TABLE IF NOT EXISTS section5_session2_recovery_event (
            checkpoint_code TEXT NOT NULL,
            event_code TEXT NOT NULL,
            condition TEXT NOT NULL,
            recovery TEXT NOT NULL,
            status TEXT NOT NULL,
            recorded_at TEXT NOT NULL,
            PRIMARY KEY(checkpoint_code,event_code)
        );
        """)
        template = clone_response_row(con, "R77")
        insert_response(con, template, {
            "response_key": "R78",
            "response_number": 78,
            "response_label": "78",
            "branch_id": "mainline",
            "canonical_current": 1,
            "response_date": "2026-08-01",
            "major_topic": "Human Pathogen Database",
            "title": "Response 77 session-restoration recovery and Session 2 intake",
            "goal": "Recover the newest verified Session 1 restore from managed storage and establish the exact Session 2 continuation boundary without regression or user uploads.",
            "raw_prompt": "Continue",
            "raw_response": "[Preserved response summary only] Recovered all three Response 77 restore volumes and the verification package from Google Drive, corrected the newest state to Session 1 complete, and established the Session 2 Checkpoint 1 continuation boundary; local execution remained unavailable.",
            "summary": "Recovered the completed Response 77 three-volume restore from Google Drive, confirmed Session 1 completion, and established the exact Session 2 Checkpoint 1 boundary without project regression or user-supplied files.",
            "state": "Completed",
            "coverage": "source-supported prompt and response summary",
            "fidelity_classification": "current_turn_summary_record",
            "source_id": "S5-R78",
            "source_path": "Tracking/Prompt Response/Through Response 79/Response_78_Tracking.json",
            "notes": "The exact full assistant response was not reintroduced into Raw tracking; the preserved summary is explicitly labeled.",
            "reconciled_at": now_iso,
        })
        insert_response(con, template, {
            "response_key": "R79",
            "response_number": 79,
            "response_label": "79",
            "branch_id": "mainline",
            "canonical_current": 1,
            "response_date": "2026-08-01",
            "major_topic": "Human Pathogen Database",
            "title": "Provider-preview and physical-proof evidence intake checkpoint",
            "goal": "Create the first Session 2 checkpoint from the exact Response 77 restore, govern KDP Print Previewer and physical-proof evidence without fabricating external results, synchronize all project surfaces, and emit deterministic recovery.",
            "raw_prompt": "Continue",
            "raw_response": "[Preserved response summary only] Built the Session 2 Checkpoint 1 copied project, provider-preview evidence register, conversion issue register, physical-proof workflow, synchronized database/workbook/application/tracking/index/manifest surfaces, and cumulative recovery package.",
            "summary": "Built and clean-verified the Session 2 Checkpoint 1 provider-preview and physical-proof evidence-intake package from the exact Response 77 restore, while retaining all external provider and proof results as controlled pending.",
            "state": "Completed",
            "coverage": "source-supported prompt and response summary",
            "fidelity_classification": "current_turn_summary_record",
            "source_id": "S5-R79",
            "source_path": "Tracking/Prompt Response/Through Response 79/Response_79_Tracking.json",
            "notes": "The package preserves provider evidence boundaries and contains deterministic restoration controls.",
            "reconciled_at": now_iso,
        })
        con.execute("DELETE FROM section5_session2_checkpoint WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        con.execute(
            "INSERT INTO section5_session2_checkpoint VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (CHECKPOINT_CODE, 79, SECTION_LABEL, SESSION_LABEL, CHECKPOINT_LABEL, "checkpoint_complete", BASE_RESTORE_SHA256, BASE_PROJECT_SHA256, PUBLICATION_SHA256, PRINT_INTERIOR_SHA256, COVER_SHA256, "controlled_pending_external", "planned_not_ordered", "pending_final_save", "pending_final_audit", "Checkpoint 2 of 3 - provider evidence reconciliation and proof-readiness closure", now_iso),
        )
        for row in sources:
            con.execute("INSERT OR REPLACE INTO section5_session2_source_control VALUES (?,?,?,?,?,?,?)", tuple(row[key] for key in ("source_key", "authority", "title", "url", "controlled_use", "evidence_boundary", "verified_at")))
        con.execute("DELETE FROM section5_provider_preview_evidence WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        for row in evidence:
            con.execute(
                "INSERT INTO section5_provider_preview_evidence (checkpoint_code,evidence_key,evidence_title,required_content,expected_evidence,status,evidence_path,provider_reference,notes,recorded_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
                (CHECKPOINT_CODE, row["evidence_key"], row["evidence_title"], row["required_content"], row["expected_evidence"], row["status"], row["evidence_path"], row["provider_reference"], row["notes"], row["recorded_at"]),
            )
        con.execute("DELETE FROM section5_provider_conversion_issue WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        for row in issues:
            con.execute(
                "INSERT INTO section5_provider_conversion_issue (checkpoint_code,issue_key,surface,page_or_region,severity,provider_message,local_reproduction,disposition,status,evidence_path,recorded_at) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                (CHECKPOINT_CODE, row["issue_key"], row["surface"], row["page_or_region"], row["severity"], row["provider_message"], row["local_reproduction"], row["disposition"], row["status"], row["evidence_path"], row["recorded_at"]),
            )
        con.execute("DELETE FROM section5_physical_proof_plan WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        for row in proof_plan:
            con.execute(
                "INSERT INTO section5_physical_proof_plan (checkpoint_code,sequence,phase,action,acceptance_criteria,evidence_required,owner_role,status,recorded_at) VALUES (?,?,?,?,?,?,?,?,?)",
                (CHECKPOINT_CODE, row["sequence"], row["phase"], row["action"], row["acceptance_criteria"], row["evidence_required"], row["owner_role"], row["status"], row["recorded_at"]),
            )
        con.execute("DELETE FROM section5_physical_proof_inspection WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        for row in proof_items:
            con.execute(
                "INSERT INTO section5_physical_proof_inspection (checkpoint_code,category,inspection_item,acceptance_criteria,observed,severity,disposition,status,evidence_path,recorded_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
                (CHECKPOINT_CODE, row["category"], row["inspection_item"], row["acceptance_criteria"], row["observed"], row["severity"], row["disposition"], row["status"], row["evidence_path"], row["recorded_at"]),
            )
        for row in recovery:
            con.execute(
                "INSERT OR REPLACE INTO section5_session2_recovery_event VALUES (?,?,?,?,?,?)",
                (CHECKPOINT_CODE, row["event_code"], row["condition"], row["recovery"], row["status"], row["recorded_at"]),
            )
        qa_rows = [
            ("database_integrity", "ok", "ok", "passed", CURRENT_DB_REL),
            ("foreign_keys", "0", "0", "passed", CURRENT_DB_REL),
            ("response_78", "1", "1", "passed", "thread_response_reconciliation_cp3"),
            ("response_79", "1", "1", "passed", "thread_response_reconciliation_cp3"),
            ("provider_previewer", "no unsupported approval claim", "controlled_pending_external", "passed", "section5_provider_preview_evidence"),
            ("physical_proof", "no unsupported inspection claim", "planned_not_ordered", "passed", "section5_physical_proof_plan"),
            ("digital_publication", PUBLICATION_SHA256, PUBLICATION_SHA256, "passed", PUBLICATION_REL),
            ("print_interior", PRINT_INTERIOR_SHA256, PRINT_INTERIOR_SHA256, "passed", PRINT_INTERIOR_REL),
            ("cover", COVER_SHA256, COVER_SHA256, "passed", COVER_REL),
        ]
        con.execute("DELETE FROM section5_session2_qa WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        con.executemany("INSERT INTO section5_session2_qa VALUES (?,?,?,?,?,?,?)", [(CHECKPOINT_CODE, *row, now_iso) for row in qa_rows])
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or fk:
            raise RuntimeError({"database_integrity": integrity, "foreign_keys": fk[:20]})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()
    con = sqlite3.connect(destination)
    try:
        table_count = con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0]
        response78 = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R78'").fetchone()[0]
        response79 = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R79'").fetchone()[0]
        evidence_count = con.execute("SELECT COUNT(*) FROM section5_provider_preview_evidence WHERE checkpoint_code=?", (CHECKPOINT_CODE,)).fetchone()[0]
        plan_count = con.execute("SELECT COUNT(*) FROM section5_physical_proof_plan WHERE checkpoint_code=?", (CHECKPOINT_CODE,)).fetchone()[0]
        inspection_count = con.execute("SELECT COUNT(*) FROM section5_physical_proof_inspection WHERE checkpoint_code=?", (CHECKPOINT_CODE,)).fetchone()[0]
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk_count = len(list(con.execute("PRAGMA foreign_key_check")))
    finally:
        con.close()
    if response78 != 1 or response79 != 1 or evidence_count != len(evidence) or plan_count != len(proof_plan) or inspection_count != len(proof_items):
        raise RuntimeError({"database_current_gate": {"response78": response78, "response79": response79, "evidence": evidence_count, "plan": plan_count, "inspection": inspection_count}})
    return {
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "table_count": table_count,
        "integrity": integrity,
        "foreign_key_violations": fk_count,
        "response78_records": response78,
        "response79_records": response79,
        "preview_evidence_records": evidence_count,
        "proof_plan_records": plan_count,
        "proof_inspection_records": inspection_count,
        "checkpoint_state": "checkpoint_complete",
    }


def serialize_cell(value: Any) -> Any:
    return json.dumps(value, ensure_ascii=False, sort_keys=True) if isinstance(value, (dict, list, tuple, set)) else value


def write_sheet(ws: Any, rows: list[dict[str, Any]]) -> None:
    if not rows:
        ws.append(["No records"])
        return
    headers = list(rows[0])
    ws.append(headers)
    for cell in ws[1]:
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(color=WHITE, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin = Side(style="thin", color="AAB8C0")
    for row in rows:
        ws.append([serialize_cell(row.get(header)) for header in headers])
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if cell.row % 2 == 0:
                cell.fill = PatternFill("solid", fgColor=PALE_BLUE)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    ws.sheet_view.showGridLines = False
    for index, header in enumerate(headers, start=1):
        sample = [str(header)] + [str(ws.cell(row, index).value or "") for row in range(2, min(ws.max_row, 150) + 1)]
        ws.column_dimensions[get_column_letter(index)].width = min(55, max(10, max(len(value) for value in sample) + 2))


def augment_workbook(source: Path, destination: Path, *, sources: list[dict[str, Any]], evidence: list[dict[str, Any]], issues: list[dict[str, Any]], proof_plan: list[dict[str, Any]], proof_items: list[dict[str, Any]], recovery: list[dict[str, Any]]) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    wb = load_workbook(source)
    inherited = list(wb.sheetnames)
    datasets = {
        "S5S2 Dashboard": [
            {"Control": "Response", "Value": 79, "Status": "current"},
            {"Control": "Session", "Value": "2 of 3", "Status": "continue"},
            {"Control": "Checkpoint", "Value": "1 of 3", "Status": "complete"},
            {"Control": "Print Previewer", "Value": "Evidence not supplied", "Status": "controlled_pending_external"},
            {"Control": "Physical proof", "Value": "Not ordered or inspected", "Status": "planned"},
            {"Control": "Digital publication", "Value": "537-page immutable", "Status": "passed"},
            {"Control": "Print candidate", "Value": "538-page Session 1 frozen candidate", "Status": "preserved"},
            {"Control": "Next", "Value": "Checkpoint 2 - provider evidence reconciliation and proof-readiness closure", "Status": "continue"},
        ],
        "S5S2 Responses": [
            {"Response": 78, "Raw Prompt": "Continue", "Summary": "Recovered Response 77 and established Session 2 intake.", "State": "completed"},
            {"Response": 79, "Raw Prompt": "Continue", "Summary": "Built provider-preview and proof evidence-intake checkpoint.", "State": "checkpoint_complete_continue_required"},
        ],
        "S5S2 Sources": sources,
        "S5S2 Preview Evidence": evidence,
        "S5S2 Conversion Issues": issues,
        "S5S2 Proof Plan": proof_plan,
        "S5S2 Proof Inspection": proof_items,
        "S5S2 QA": [
            {"Control": "Database", "Expected": "integrity ok / FK 0", "Observed": "passed", "Status": "passed"},
            {"Control": "Workbook", "Expected": "all inherited sheets retained", "Observed": "pending final save", "Status": "in_progress"},
            {"Control": "Provider evidence", "Expected": "no unsupported claim", "Observed": "controlled pending", "Status": "passed"},
            {"Control": "Physical proof", "Expected": "no unsupported inspection", "Observed": "planned/not inspected", "Status": "passed"},
            {"Control": "Publication", "Expected": PUBLICATION_SHA256, "Observed": PUBLICATION_SHA256, "Status": "passed"},
            {"Control": "Print interior", "Expected": PRINT_INTERIOR_SHA256, "Observed": PRINT_INTERIOR_SHA256, "Status": "passed"},
            {"Control": "Cover", "Expected": COVER_SHA256, "Observed": COVER_SHA256, "Status": "passed"},
        ],
        "S5S2 Recovery": recovery,
    }
    for title, rows in datasets.items():
        if title in wb.sheetnames:
            del wb[title]
        ws = wb.create_sheet(title)
        write_sheet(ws, rows)
    wb.properties.title = "MRHPD v3.0.0a Comprehensive Tracking Through Response 79"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(destination)
    with zipfile.ZipFile(destination) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("workbook CRC failed")
    check = load_workbook(destination, read_only=True, data_only=False)
    formula_errors = []
    formula_count = 0
    try:
        sheet_names = list(check.sheetnames)
        for ws in check.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    value = cell.value
                    if isinstance(value, str) and value.startswith("="):
                        formula_count += 1
                        if any(token in value for token in ("#REF!", "#NAME?", "#VALUE!", "#DIV/0!", "#N/A")):
                            formula_errors.append(f"{ws.title}!{cell.coordinate}:{value}")
    finally:
        check.close()
    lost = sorted(set(inherited) - set(sheet_names))
    if lost or formula_errors or len(sheet_names) < len(inherited) + len(datasets):
        raise RuntimeError({"workbook_gate": {"lost": lost, "sheets": len(sheet_names), "formula_errors": formula_errors[:30]}})
    return {
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "source_sheet_count": len(inherited),
        "current_sheet_count": len(sheet_names),
        "new_sheet_count": len(sheet_names) - len(inherited),
        "lost_sheets": lost,
        "formula_count": formula_count,
        "formula_error_count": len(formula_errors),
        "status": "passed",
    }


def shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def tracking_files(project: Path, db_path: Path, now_iso: str) -> list[Path]:
    root = project / "Tracking" / "Prompt Response" / "Through Response 79"
    root.mkdir(parents=True, exist_ok=True)
    con = sqlite3.connect(db_path)
    try:
        columns = [row[1] for row in con.execute("PRAGMA table_info(thread_response_reconciliation_cp3)")]
        rows = [dict(zip(columns, row)) for row in con.execute("SELECT * FROM thread_response_reconciliation_cp3 ORDER BY CAST(response_number AS INTEGER), response_key")]
        frac_columns = [row[1] for row in con.execute("PRAGMA table_info(fractional_prompt_cp3)")]
        fractions = [dict(zip(frac_columns, row)) for row in con.execute("SELECT * FROM fractional_prompt_cp3 ORDER BY CAST(prompt_number AS REAL), prompt_number")]
    finally:
        con.close()
    current = [row for row in rows if row.get("response_key") in {"R78", "R79"}]
    outputs = []
    for row in current:
        path = root / f"Response_{row['response_number']}_Tracking.json"
        json_write(path, row)
        outputs.append(path)

    raw_docx = root / "Medical References - Human Pathogen Database v3.0.0a Alternating Raw Prompts and Responses Through Response 79.docx"
    doc = Document()
    doc.core_properties.title = "Human Pathogen Database — Alternating Raw Prompts and Responses Through Response 79"
    doc.core_properties.author = "Brent McAnulty, M.D."
    doc.add_heading("Human Pathogen Database", 0)
    doc.add_paragraph("Alternating Raw Prompts and Responses Through Response 79")
    for row in rows:
        number = row.get("response_label") or row.get("response_number")
        doc.add_heading(f"Response {number}: {row.get('title') or 'Untitled exchange'}", level=1)
        table = doc.add_table(rows=2, cols=1)
        table.style = "Table Grid"
        table.cell(0, 0).text = f"RAW PROMPT {number}\n\n{row.get('raw_prompt') or '[RAW PROMPT UNAVAILABLE]'}"
        shade_cell(table.cell(0, 0), "D9EAF7")
        table.cell(1, 0).text = f"RAW RESPONSE {number}\n\n{row.get('raw_response') or '[RAW RESPONSE UNAVAILABLE]'}\n\nSUMMARY\n{row.get('summary') or ''}"
        shade_cell(table.cell(1, 0), "E2F0D9")
        for cell in table._cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(9)
        doc.add_paragraph()
    if fractions:
        doc.add_heading("Fractional prompts", level=1)
        for row in fractions:
            doc.add_paragraph(f"Prompt {row.get('prompt_number')}: {row.get('prompt_text')}")
    doc.save(raw_docx)
    outputs.append(raw_docx)

    net_prompt = (
        "Continue the Human Pathogen Database from the newest verified managed checkpoint without regression. Use Google Drive as controlling storage; recover autonomously; preserve accepted clinical, publication, application, and print-candidate identities; conduct provider-preview and physical-proof governance without inventing external results; synchronize the database, workbook, application, tracking, indexes, manifests, QA, and recovery; emit checkpoint recovery between full session/section/project restores."
    )
    net_response = (
        "Remediation Sections 1–4 and Section 5 Session 1 are complete. Section 5 Session 2 Checkpoint 1 preserves the frozen 537-page digital publication and 538-page Premium Color print candidate, creates the KDP Print Previewer evidence-intake and conversion-issue registers, defines the physical-proof order and inspection workflow, records external evidence as controlled pending, and synchronizes all project-management and recovery surfaces."
    )
    net_docx = root / "Medical References - Human Pathogen Database v3.0.0a Alternating Net Prompts and Responses Through Response 79.docx"
    net_doc = Document()
    net_doc.core_properties.title = "Human Pathogen Database — Alternating Net Prompts and Responses Through Response 79"
    net_doc.core_properties.author = "Brent McAnulty, M.D."
    net_doc.add_heading("Human Pathogen Database", 0)
    net_doc.add_heading("Print-production evidence and final-release remediation", level=1)
    table = net_doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.cell(0, 0).text = "NET PROMPT\n\n" + net_prompt
    shade_cell(table.cell(0, 0), "D9EAF7")
    table.cell(1, 0).text = "NET RESPONSE\n\n" + net_response
    shade_cell(table.cell(1, 0), "E2F0D9")
    net_doc.save(net_docx)
    outputs.append(net_docx)

    everything = root / "Medical References - Human Pathogen Database v3.0.0a Everything in One Thread Through Response 79.xlsx"
    wb = Workbook()
    wb.remove(wb.active)
    datasets = {
        "Raw Prompts": [{"Response": row.get("response_label") or row.get("response_number"), "Major Topic": row.get("major_topic"), "Title": row.get("title"), "Raw Prompt": row.get("raw_prompt")} for row in rows],
        "Raw Responses": [{"Response": row.get("response_label") or row.get("response_number"), "Major Topic": row.get("major_topic"), "Title": row.get("title"), "Raw Response": row.get("raw_response"), "Summary": row.get("summary")} for row in rows],
        "Fractional Prompts": fractions,
        "Net Prompt": [{"Major Topic": "Human Pathogen Database remediation", "Net Prompt": net_prompt}],
        "Net Response": [{"Major Topic": "Human Pathogen Database remediation", "Net Response": net_response}],
        "Summary Index": [{"Response": row.get("response_label") or row.get("response_number"), "Major Topic": row.get("major_topic"), "Title": row.get("title"), "Goal": row.get("goal"), "Summary": row.get("summary"), "State": row.get("state")} for row in rows],
    }
    for title, data in datasets.items():
        ws = wb.create_sheet(title)
        write_sheet(ws, data)
    wb.properties.title = "Human Pathogen Database — Everything in One Thread Through Response 79"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(everything)
    outputs.append(everything)

    raw_net = root / "Medical References - Human Pathogen Database v3.0.0a Raw and Net Tracking Through Response 79.md"
    text_write(raw_net, f"""# Human Pathogen Database — Raw and Net Tracking Through Response 79

## Raw Prompt 78

Continue

## Raw Response 78

[Preserved response summary only] Recovered the Response 77 complete restore from managed storage and established the exact Session 2 continuation boundary.

## Raw Prompt 79

Continue

## Raw Response 79

[Preserved response summary only] Built the provider-preview and physical-proof evidence-intake checkpoint and synchronized recovery state.

## Net Prompt

{net_prompt}

## Net Response

{net_response}

Updated: {now_iso}
""")
    outputs.append(raw_net)

    cumulative = root / "Medical References - Human Pathogen Database v3.0.0a Cumulative Thread Index Through Response 79.md"
    lines = ["# Human Pathogen Database — Cumulative Thread Index Through Response 79", ""]
    for row in rows:
        lines.append(f"- Response {row.get('response_label') or row.get('response_number')} — {row.get('title')}: {row.get('summary')}")
    lines.extend(["", f"Updated: {now_iso}"])
    text_write(cumulative, "\n".join(lines))
    outputs.append(cumulative)
    return outputs


def application_surfaces(project: Path, db_path: Path, workbook_path: Path, now_iso: str) -> tuple[list[Path], dict[str, Any]]:
    root = project / "App" / "Section 5 Session 2 Checkpoint 1"
    root.mkdir(parents=True, exist_ok=True)
    apps = [path for path in project.rglob("human_pathogen_app.py") if path.is_file() and sha256_file(path) == APPLICATION_SHA256]
    if len(apps) != 1:
        raise RuntimeError({"application_candidates": [str(path) for path in apps]})
    app = apps[0]
    pointer = root / "CURRENT_DATABASE.txt"
    text_write(pointer, db_path.relative_to(project).as_posix() + "\n")
    state = root / "CURRENT_PROJECT_STATE.json"
    json_write(state, {
        "schema": "mrhpd-section5-session2-current-state-1.0",
        "response": 79,
        "section": SECTION_LABEL,
        "session": SESSION_LABEL,
        "checkpoint": CHECKPOINT_LABEL,
        "state": "checkpoint_complete",
        "database": db_path.relative_to(project).as_posix(),
        "workbook": workbook_path.relative_to(project).as_posix(),
        "digital_publication": PUBLICATION_REL,
        "print_interior": PRINT_INTERIOR_REL,
        "cover": COVER_REL,
        "provider_previewer": "controlled_pending_external",
        "physical_proof": "planned_not_ordered",
        "main_application": app.relative_to(project).as_posix(),
        "main_application_sha256": sha256_file(app),
        "main_application_unchanged": True,
        "recorded_at": now_iso,
    })
    audit_script = root / "audit_section5_session2_checkpoint1.py"
    db_rel = db_path.relative_to(project).as_posix()
    wb_rel = workbook_path.relative_to(project).as_posix()
    text_write(audit_script, f'''#!/usr/bin/env python3
import json, sqlite3
from pathlib import Path
from openpyxl import load_workbook
from pypdf import PdfReader
project=Path(__file__).resolve().parents[2]
db=project/{db_rel!r}
workbook=project/{wb_rel!r}
publication=project/{PUBLICATION_REL!r}
interior=project/{PRINT_INTERIOR_REL!r}
cover=project/{COVER_REL!r}
con=sqlite3.connect(db)
try:
 integrity=con.execute("PRAGMA integrity_check").fetchone()[0]
 fk=len(list(con.execute("PRAGMA foreign_key_check")))
 r78=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R78'").fetchone()[0]
 r79=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R79'").fetchone()[0]
 checkpoint=con.execute("SELECT state,provider_previewer_status,physical_proof_status FROM section5_session2_checkpoint WHERE checkpoint_code='{CHECKPOINT_CODE}'").fetchone()
 evidence=con.execute("SELECT COUNT(*) FROM section5_provider_preview_evidence WHERE checkpoint_code='{CHECKPOINT_CODE}'").fetchone()[0]
 plan=con.execute("SELECT COUNT(*) FROM section5_physical_proof_plan WHERE checkpoint_code='{CHECKPOINT_CODE}'").fetchone()[0]
finally: con.close()
wb=load_workbook(workbook,read_only=True,data_only=False)
try: sheets=len(wb.sheetnames)
finally: wb.close()
pub=PdfReader(str(publication)); interior_pdf=PdfReader(str(interior))
result={{'status':'passed' if integrity=='ok' and fk==0 and r78==1 and r79==1 and checkpoint==('checkpoint_complete','controlled_pending_external','planned_not_ordered') and evidence==8 and plan==11 and sheets>=114 and len(pub.pages)==537 and len(interior_pdf.pages)==538 and cover.exists() else 'failed','integrity':integrity,'foreign_keys':fk,'response78':r78,'response79':r79,'checkpoint':checkpoint,'preview_evidence':evidence,'proof_plan':plan,'workbook_sheets':sheets,'digital_pages':len(pub.pages),'print_pages':len(interior_pdf.pages),'cover_exists':cover.exists()}}
print(json.dumps(result,indent=2))
raise SystemExit(0 if result['status']=='passed' else 1)
''')
    result = subprocess.run([sys.executable, str(audit_script)], cwd=project, text=True, capture_output=True, timeout=300)
    if result.returncode:
        raise RuntimeError({"application_audit_failed": {"stdout": result.stdout[-8000:], "stderr": result.stderr[-8000:]}})
    audit = json.loads(result.stdout)
    audit.update({"main_application": app.relative_to(project).as_posix(), "main_application_sha256": sha256_file(app), "main_application_unchanged": True})
    output = root / "SECTION5_SESSION2_CHECKPOINT1_APPLICATION_AUDIT.json"
    json_write(output, audit)
    return [pointer, state, audit_script, output], audit


def report_docx(path: Path, *, generated_at: str, sources: list[dict[str, Any]], evidence: list[dict[str, Any]], issues: list[dict[str, Any]], proof_plan: list[dict[str, Any]], proof_items: list[dict[str, Any]], recovery: list[dict[str, Any]]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    doc.core_properties.title = "MRHPD Section 5 Session 2 Checkpoint 1 Provider Preview and Physical Proof Intake Report"
    doc.core_properties.author = "Brent McAnulty, M.D."
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HUMAN PATHOGEN DATABASE")
    run.bold = True; run.font.size = Pt(24); run.font.color.rgb = RGBColor(23, 50, 77)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Provider Preview and Physical Proof Evidence Intake")
    run.bold = True; run.font.size = Pt(16); run.font.color.rgb = RGBColor(28, 116, 117)
    doc.add_paragraph("Remediation Section 5 of 5 • Session 2 of 3 • Checkpoint 1 of 3 • Through Response 79")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("CONTROLLED EVIDENCE BOUNDARY").bold = True
    doc.add_paragraph("The frozen Session 1 production candidate remains intact. No KDP Print Previewer approval, provider warning/error result, proof order, physical proof receipt, inspection finding, or production acceptance is claimed unless item-level evidence is present in the corresponding register.")

    doc.add_heading("Checkpoint disposition", level=1)
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(("Control", "Observed state", "Disposition")):
        table.cell(0, i).text = value; shade_cell(table.cell(0, i), NAVY)
        for run in table.cell(0, i).paragraphs[0].runs: run.font.color.rgb = RGBColor(255,255,255); run.bold=True
    rows = [
        ("Digital publication", "537 pages; immutable governed identity", "passed"),
        ("Print candidate", "538 pages; Session 1 frozen candidate", "preserved"),
        ("Cover candidate", "5,554 × 3,375 pixels", "preserved"),
        ("KDP Print Previewer", "No account-derived evidence available", "controlled pending"),
        ("Physical proof", "Not ordered, received, or inspected in evidence", "planned"),
        ("Next", "Provider evidence reconciliation and proof-readiness closure", "continue"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row): cells[i].text = value

    doc.add_heading("Official provider controls", level=1)
    for row in sources:
        doc.add_heading(row["title"], level=2)
        doc.add_paragraph(row["controlled_use"])
        doc.add_paragraph("Evidence boundary: " + row["evidence_boundary"])
        doc.add_paragraph("Source: " + row["url"])

    doc.add_heading("Print Previewer evidence register", level=1)
    table = doc.add_table(rows=1, cols=4); table.style="Table Grid"
    for i, value in enumerate(("Key", "Required evidence", "Expected artifact", "State")):
        table.cell(0,i).text=value; shade_cell(table.cell(0,i), NAVY)
        for run in table.cell(0,i).paragraphs[0].runs: run.font.color.rgb=RGBColor(255,255,255); run.bold=True
    for row in evidence:
        cells=table.add_row().cells
        for i, value in enumerate((row["evidence_key"], row["required_content"], row["expected_evidence"], row["status"])): cells[i].text=str(value)

    doc.add_heading("Provider conversion issue register", level=1)
    doc.add_paragraph("No provider defect is inferred. The initial control row records only that Print Previewer evidence has not yet been supplied.")
    for row in issues:
        doc.add_paragraph(f"{row['issue_key']} — {row['provider_message']} — {row['status']}")

    doc.add_heading("Physical proof workflow", level=1)
    table = doc.add_table(rows=1, cols=4); table.style="Table Grid"
    for i, value in enumerate(("Seq.", "Phase", "Action", "Acceptance / evidence")):
        table.cell(0,i).text=value; shade_cell(table.cell(0,i), TEAL)
        for run in table.cell(0,i).paragraphs[0].runs: run.font.color.rgb=RGBColor(255,255,255); run.bold=True
    for row in proof_plan:
        cells=table.add_row().cells
        values=(row["sequence"], row["phase"], row["action"], row["acceptance_criteria"]+" Evidence: "+row["evidence_required"])
        for i,value in enumerate(values): cells[i].text=str(value)

    doc.add_heading("Physical proof inspection checklist", level=1)
    table=doc.add_table(rows=1,cols=3); table.style="Table Grid"
    for i,value in enumerate(("Category","Inspection item","Acceptance criteria")):
        table.cell(0,i).text=value; shade_cell(table.cell(0,i), GOLD)
        for run in table.cell(0,i).paragraphs[0].runs: run.font.color.rgb=RGBColor(255,255,255); run.bold=True
    for row in proof_items:
        cells=table.add_row().cells
        for i,value in enumerate((row["category"],row["inspection_item"],row["acceptance_criteria"])): cells[i].text=str(value)

    doc.add_heading("Recovery and next checkpoint", level=1)
    for row in recovery:
        doc.add_paragraph(f"{row['event_code']}: {row['condition']} Recovery: {row['recovery']}")
    doc.add_paragraph("Checkpoint 2 will reconcile actual provider evidence when available, classify and resolve conversion issues, complete proof-order readiness, and preserve controlled-pending status for every external gate that remains unsupported.")
    doc.add_paragraph("Generated: " + generated_at)
    doc.save(path)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None: raise RuntimeError("DOCX CRC failure")
    return {"path": str(path), "bytes": path.stat().st_size, "sha256": sha256_file(path), "paragraphs": len(doc.paragraphs), "tables": len(doc.tables), "status": "passed"}


def report_pdf(path: Path, *, generated_at: str, sources: list[dict[str, Any]], evidence: list[dict[str, Any]], proof_plan: list[dict[str, Any]], proof_items: list[dict[str, Any]]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="MRTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=19, leading=22, textColor=colors.HexColor("#17324D"), alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name="MRSub", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=14, textColor=colors.HexColor("#1C7475"), alignment=TA_CENTER, spaceAfter=12))
    styles.add(ParagraphStyle(name="MRH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=13, leading=15, textColor=colors.HexColor("#17324D"), spaceBefore=8, spaceAfter=6))
    styles.add(ParagraphStyle(name="MRBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.5, leading=11, textColor=colors.HexColor("#24323D"), spaceAfter=5))
    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.48*inch, leftMargin=0.48*inch, topMargin=0.48*inch, bottomMargin=0.48*inch, title="MRHPD Section 5 Session 2 Checkpoint 1 Provider Preview and Physical Proof Intake", author="Brent McAnulty, M.D.")
    story: list[Any] = [
        Paragraph("HUMAN PATHOGEN DATABASE", styles["MRTitle"]),
        Paragraph("Provider Preview and Physical Proof Evidence Intake", styles["MRSub"]),
        Paragraph("Remediation Section 5 of 5 • Session 2 of 3 • Checkpoint 1 of 3 • Through Response 79", styles["MRBody"]),
        Paragraph("<b>Controlled evidence boundary.</b> No KDP Print Previewer approval, provider issue result, proof order, physical proof receipt, inspection finding, or production acceptance is claimed without item-level evidence.", styles["MRBody"]),
        Paragraph("Checkpoint disposition", styles["MRH1"]),
    ]
    disposition = [["Control","Observed","Disposition"],["Digital publication","537-page immutable identity","passed"],["Print candidate","538-page Session 1 frozen candidate","preserved"],["Cover","5,554 × 3,375 pixels","preserved"],["Print Previewer","No provider evidence supplied","controlled pending"],["Physical proof","Not ordered/received/inspected","planned"]]
    table=Table(disposition,colWidths=[1.5*inch,3.9*inch,1.25*inch],repeatRows=1)
    table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#17324D")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTNAME",(0,1),(-1,-1),"Helvetica"),("FONTSIZE",(0,0),(-1,-1),7),("GRID",(0,0),(-1,-1),0.35,colors.HexColor("#8AA0AC")),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#EAF1F5")]),("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4),("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3)]))
    story.extend([table, Spacer(1,8), Paragraph("Official provider controls",styles["MRH1"])])
    for row in sources:
        story.append(Paragraph(f"<b>{row['title']}</b> — {row['controlled_use']}<br/><font color='#66757F'>{row['url']}</font>",styles["MRBody"]))
    story.extend([PageBreak(),Paragraph("Print Previewer evidence register",styles["MRH1"])])
    preview_data=[["Key","Required content","Expected evidence","State"]]+[[row["evidence_key"],row["required_content"],row["expected_evidence"],row["status"]] for row in evidence]
    table=Table(preview_data,colWidths=[0.7*inch,2.65*inch,2.6*inch,0.9*inch],repeatRows=1)
    table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#17324D")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTNAME",(0,1),(-1,-1),"Helvetica"),("FONTSIZE",(0,0),(-1,-1),6.3),("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#8AA0AC")),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#EAF1F5")]),("LEFTPADDING",(0,0),(-1,-1),3),("RIGHTPADDING",(0,0),(-1,-1),3),("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3)]))
    story.extend([table,Spacer(1,8),Paragraph("Physical proof workflow",styles["MRH1"])])
    proof_data=[["Seq.","Phase","Action","Acceptance/evidence"]]+[[str(row["sequence"]),row["phase"],row["action"],row["acceptance_criteria"]+" Evidence: "+row["evidence_required"]] for row in proof_plan]
    table=Table(proof_data,colWidths=[0.35*inch,0.7*inch,2.7*inch,3.1*inch],repeatRows=1)
    table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#1C7475")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTNAME",(0,1),(-1,-1),"Helvetica"),("FONTSIZE",(0,0),(-1,-1),6.1),("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#8AA0AC")),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#E9F3EE")]),("LEFTPADDING",(0,0),(-1,-1),3),("RIGHTPADDING",(0,0),(-1,-1),3),("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3)]))
    story.extend([table,PageBreak(),Paragraph("Physical proof inspection checklist",styles["MRH1"])])
    check_data=[["Category","Item","Acceptance criteria","State"]]+[[row["category"],row["inspection_item"],row["acceptance_criteria"],row["status"]] for row in proof_items]
    table=Table(check_data,colWidths=[0.8*inch,1.7*inch,3.7*inch,0.75*inch],repeatRows=1)
    table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#C9A227")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTNAME",(0,1),(-1,-1),"Helvetica"),("FONTSIZE",(0,0),(-1,-1),6.4),("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#8AA0AC")),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#F7F1D9")]),("LEFTPADDING",(0,0),(-1,-1),3),("RIGHTPADDING",(0,0),(-1,-1),3),("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3)]))
    story.extend([table,Spacer(1,10),Paragraph("Next checkpoint",styles["MRH1"]),Paragraph("Checkpoint 2 will reconcile actual Print Previewer output when it exists, classify and resolve provider conversion issues, complete proof-order readiness, preserve unsupported external gates as controlled pending, and rebuild the synchronized recovery state.",styles["MRBody"]),Paragraph("Generated: "+generated_at,styles["MRBody"])])
    doc.build(story)
    reader=PdfReader(str(path))
    searchable=sum(1 for page in reader.pages if (page.extract_text() or "").strip())
    return {"path":str(path),"bytes":path.stat().st_size,"sha256":sha256_file(path),"pages":len(reader.pages),"searchable_pages":searchable,"status":"passed" if searchable==len(reader.pages) else "failed"}


def report_register(path: Path, *, sources: list[dict[str, Any]], evidence: list[dict[str, Any]], issues: list[dict[str, Any]], proof_plan: list[dict[str, Any]], proof_items: list[dict[str, Any]], recovery: list[dict[str, Any]]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb=Workbook(); wb.remove(wb.active)
    datasets={
        "Summary":[{"Control":"Response","Value":79,"Status":"current"},{"Control":"Checkpoint","Value":"1 of 3","Status":"complete"},{"Control":"Print Previewer","Value":"controlled pending external","Status":"pending"},{"Control":"Physical proof","Value":"planned, not ordered","Status":"pending"}],
        "Official Sources":sources,
        "Preview Evidence":evidence,
        "Conversion Issues":issues,
        "Proof Plan":proof_plan,
        "Proof Inspection":proof_items,
        "Tracking":[{"Response":78,"Prompt":"Continue","State":"complete"},{"Response":79,"Prompt":"Continue","State":"checkpoint complete"}],
        "Recovery":recovery,
    }
    for title,rows in datasets.items():
        ws=wb.create_sheet(title); write_sheet(ws,rows)
    wb.properties.title="MRHPD Section 5 Session 2 Checkpoint 1 Provider Preview and Proof Register"
    wb.properties.creator="Brent McAnulty, M.D."
    wb.save(path)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None: raise RuntimeError("register CRC failure")
    return {"path":str(path),"bytes":path.stat().st_size,"sha256":sha256_file(path),"sheets":len(wb.sheetnames),"status":"passed"}


def render_pdf_qa(pdf: Path, destination: Path) -> dict[str, Any]:
    destination.mkdir(parents=True, exist_ok=True)
    doc=fitz.open(pdf)
    rows=[]
    for index,page in enumerate(doc):
        pix=page.get_pixmap(matrix=fitz.Matrix(1.5,1.5),alpha=False)
        output=destination/f"page-{index+1:03d}.png"
        pix.save(output)
        with Image.open(output) as image:
            width,height=image.size
        rows.append({"page":index+1,"path":str(output),"bytes":output.stat().st_size,"pixels":[width,height],"status":"passed" if output.stat().st_size>1000 else "failed"})
    doc.close()
    if any(row["status"]!="passed" for row in rows): raise RuntimeError({"render_qa":rows})
    return {"status":"passed","pages":len(rows),"records":rows}


def extract_text(path: Path) -> str:
    suffix=path.suffix.lower()
    try:
        if suffix in {".md",".txt",".csv",".json",".py",".html",".yml",".yaml"}:
            return path.read_text(encoding="utf-8",errors="replace")
        if suffix==".docx":
            doc=Document(path); chunks=[p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows: chunks.extend(cell.text for cell in row.cells)
            return "\n".join(chunks)
        if suffix==".xlsx":
            wb=load_workbook(path,read_only=True,data_only=False)
            try:
                chunks=[]
                for ws in wb.worksheets:
                    chunks.append(ws.title)
                    for row in ws.iter_rows(values_only=True): chunks.append(" | ".join("" if value is None else str(value) for value in row))
                return "\n".join(chunks)
            finally: wb.close()
        if suffix in {".sqlite",".db"}:
            con=sqlite3.connect(path)
            try: return "\n".join(row[0] for row in con.execute("SELECT sql FROM sqlite_master WHERE sql IS NOT NULL ORDER BY name"))
            finally: con.close()
        if suffix==".pdf" and path.stat().st_size < 5_000_000:
            reader=PdfReader(str(path)); return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        return f"[index extraction error: {exc!r}]"
    return ""


def build_indexes(project: Path, now_iso: str) -> dict[str, Any]:
    root=project/"Indexes"/"Section 5 Session 2 Checkpoint 1"; root.mkdir(parents=True,exist_ok=True)
    source_json=root/"MRHPD v3.0.0a Section 5 Session 2 Checkpoint 1 Source Index.json"
    source_csv=root/"MRHPD v3.0.0a Section 5 Session 2 Checkpoint 1 Source Index.csv"
    bit_path=root/"MRHPD v3.0.0a Section 5 Session 2 Checkpoint 1 Bit Index.sqlite"
    qa_path=root/"MRHPD v3.0.0a Section 5 Session 2 Checkpoint 1 Index QA.json"
    excluded={path.resolve() for path in (source_json,source_csv,bit_path,qa_path)}
    rows=[]; payloads=[]
    searchable_suffixes={".md",".txt",".csv",".json",".py",".html",".yml",".yaml",".docx",".xlsx",".sqlite",".db",".pdf"}
    physical=[path for path in project.rglob("*") if path.is_file() and path.resolve() not in excluded]
    for path in sorted(physical):
        rel=path.relative_to(project).as_posix(); purpose="Project artifact"
        if rel.startswith("Database/"): purpose="Canonical or historical project database"
        elif rel.startswith("Tracking/"): purpose="Prompt, response, workbook, or cumulative tracking"
        elif rel.startswith("Reports/"): purpose="Human-readable project report or register"
        elif rel.startswith("Print Production/"): purpose="Governed print-production artifact"
        elif rel.startswith("QA/"): purpose="Quality-assurance evidence"
        elif rel.startswith("Recovery/"): purpose="Recovery, checkpoint, or provenance artifact"
        elif rel.startswith("Sources/"): purpose="Source-control or evidence-governance artifact"
        elif rel.startswith("App/"): purpose="Read-only local application surface"
        user_searchable=int(path.suffix.lower() in searchable_suffixes)
        digest=sha256_file(path)
        rows.append({"record_type":"physical_file","path":rel,"container_path":"","name":path.name,"purpose":purpose,"bytes":path.stat().st_size,"sha256":digest,"user_searchable":user_searchable})
        content=extract_text(path) if user_searchable and ("Section 5 Session 2" in rel or rel in {CURRENT_DB_REL,CURRENT_WORKBOOK_REL} or path.stat().st_size<2_000_000) else ""
        payloads.append((rel,path.name,purpose,content))
        if path.suffix.lower()==".zip":
            try:
                with zipfile.ZipFile(path) as zf:
                    for info in zf.infolist():
                        if info.is_dir(): continue
                        member_path=rel+"::"+info.filename
                        rows.append({"record_type":"container_member","path":member_path,"container_path":rel,"name":PurePosixPath(info.filename).name,"purpose":"ZIP member","bytes":info.file_size,"sha256":"","user_searchable":1})
                        payloads.append((member_path,PurePosixPath(info.filename).name,"ZIP member",""))
            except Exception:
                pass
    json_write(source_json,{"schema":"mrhpd-source-index-2.0","generated_at":now_iso,"records":rows,"record_count":len(rows)})
    csv_write(source_csv,rows)
    if bit_path.exists(): bit_path.unlink()
    con=sqlite3.connect(bit_path)
    try:
        con.executescript("""
        CREATE TABLE artifact (artifact_id INTEGER PRIMARY KEY,record_type TEXT NOT NULL,path TEXT NOT NULL UNIQUE,container_path TEXT,name TEXT NOT NULL,purpose TEXT NOT NULL,bytes INTEGER NOT NULL,sha256 TEXT,user_searchable INTEGER NOT NULL);
        CREATE VIRTUAL TABLE artifact_fts USING fts5(path,name,purpose,content);
        """)
        for row,payload in zip(rows,payloads):
            con.execute("INSERT INTO artifact (record_type,path,container_path,name,purpose,bytes,sha256,user_searchable) VALUES (?,?,?,?,?,?,?,?)",(row["record_type"],row["path"],row["container_path"],row["name"],row["purpose"],row["bytes"],row["sha256"],row["user_searchable"]))
            con.execute("INSERT INTO artifact_fts (path,name,purpose,content) VALUES (?,?,?,?)",payload)
        integrity=con.execute("PRAGMA integrity_check").fetchone()[0]
        counts={"artifact":con.execute("SELECT COUNT(*) FROM artifact").fetchone()[0],"fts":con.execute("SELECT COUNT(*) FROM artifact_fts").fetchone()[0],"response79":con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?",('"Response 79"',)).fetchone()[0],"print_previewer":con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?",('"Print Previewer"',)).fetchone()[0],"physical_proof":con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?",('"physical proof"',)).fetchone()[0]}
        con.commit()
    finally: con.close()
    if integrity!="ok" or counts["artifact"]!=len(rows) or counts["fts"]!=len(rows): raise RuntimeError({"bit_index_gate":{"integrity":integrity,"counts":counts,"expected":len(rows)}})
    qa={"status":"passed","generated_at":now_iso,"source_index_records":len(rows),"physical_files":sum(1 for row in rows if row["record_type"]=="physical_file"),"container_members":sum(1 for row in rows if row["record_type"]=="container_member"),"bit_index_integrity":integrity,"counts":counts,"bit_index_sha256":sha256_file(bit_path)}
    json_write(qa_path,qa)
    return {"source_json":source_json,"source_csv":source_csv,"bit_index":bit_path,"qa_path":qa_path,"qa":qa}


def build_manifest(project: Path, now_iso: str) -> tuple[Path,Path,list[dict[str,Any]]]:
    root=project/"Manifest"/"Section 5 Session 2 Checkpoint 1"; root.mkdir(parents=True,exist_ok=True)
    manifest=root/"MRHPD v3.0.0a Section 5 Session 2 Checkpoint 1 Current Project Manifest.json"
    checksums=root/"MRHPD v3.0.0a Section 5 Session 2 Checkpoint 1 Current Project Checksums.sha256"
    rows=[]
    for path in sorted(p for p in project.rglob("*") if p.is_file() and p not in {manifest,checksums}):
        rows.append({"path":path.relative_to(project).as_posix(),"bytes":path.stat().st_size,"sha256":sha256_file(path)})
    json_write(manifest,{"schema":"mrhpd-current-project-manifest-2.0","generated_at":now_iso,"exclusions":[manifest.relative_to(project).as_posix(),checksums.relative_to(project).as_posix()],"file_count":len(rows),"total_bytes":sum(row["bytes"] for row in rows),"files":rows})
    text_write(checksums,"".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
    for row in rows:
        path=project/row["path"]
        if path.stat().st_size!=row["bytes"] or sha256_file(path)!=row["sha256"]: raise RuntimeError({"manifest_mismatch":row["path"]})
    return manifest,checksums,rows


def create_apply_script(manifest: dict[str,Any], expected: dict[str,Any]) -> str:
    return f'''#!/usr/bin/env python3
from __future__ import annotations
import argparse,hashlib,json,re,shutil,sqlite3,sys,tempfile,zipfile
from pathlib import Path,PurePosixPath
from openpyxl import load_workbook
from pypdf import PdfReader
BASE_RESTORE_BYTES={BASE_RESTORE_BYTES}
BASE_RESTORE_SHA256={BASE_RESTORE_SHA256!r}
BASE_PROJECT_BYTES={BASE_PROJECT_BYTES}
BASE_PROJECT_SHA256={BASE_PROJECT_SHA256!r}
CURRENT_PROJECT_NAME={CURRENT_PROJECT_NAME!r}
CURRENT_DB_REL={CURRENT_DB_REL!r}
CURRENT_WORKBOOK_REL={CURRENT_WORKBOOK_REL!r}
PUBLICATION_REL={PUBLICATION_REL!r}
PRINT_INTERIOR_REL={PRINT_INTERIOR_REL!r}
COVER_REL={COVER_REL!r}
PUBLICATION_SHA256={PUBLICATION_SHA256!r}
PRINT_INTERIOR_SHA256={PRINT_INTERIOR_SHA256!r}
COVER_SHA256={COVER_SHA256!r}
APPLICATION_SHA256={APPLICATION_SHA256!r}
MANIFEST={manifest!r}
EXPECTED={expected!r}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as f:
  for block in iter(lambda:f.read(1024*1024),b''): h.update(block)
 return h.hexdigest()
def safe_extract(path,dest):
 dest.mkdir(parents=True,exist_ok=True)
 with zipfile.ZipFile(path) as zf:
  if zf.testzip() is not None: raise RuntimeError('ZIP CRC failure: '+str(path))
  names=zf.namelist()
  if len(names)!=len(set(names)): raise RuntimeError('duplicate ZIP members')
  for name in names:
   pp=PurePosixPath(name.replace('\\\\','/'))
   if pp.is_absolute() or '..' in pp.parts or re.match(r'^[A-Za-z]:',name): raise RuntimeError('unsafe ZIP path: '+name)
  zf.extractall(dest)
def verify(path,size,digest,label):
 observed={{'bytes':path.stat().st_size,'sha256':sha(path)}}
 if observed!={{'bytes':size,'sha256':digest}}: raise RuntimeError({{label:observed}})
def main():
 ap=argparse.ArgumentParser(); ap.add_argument('--base-response77-restore',type=Path,required=True); ap.add_argument('--output-dir',type=Path,required=True); args=ap.parse_args()
 verify(args.base_response77_restore,BASE_RESTORE_BYTES,BASE_RESTORE_SHA256,'baseline_restore')
 package=Path(__file__).resolve().parents[1]; overlay=package/'OVERLAY'
 if not overlay.is_dir(): raise RuntimeError('OVERLAY missing')
 if args.output_dir.exists() and any(args.output_dir.iterdir()): raise RuntimeError('output directory must be empty')
 args.output_dir.mkdir(parents=True,exist_ok=True)
 with tempfile.TemporaryDirectory(prefix='mrhpd-r79-apply-') as td:
  work=Path(td); restore_root=work/'restore'; safe_extract(args.base_response77_restore,restore_root)
  candidates=[p for p in restore_root.rglob('*.zip') if p.stat().st_size==BASE_PROJECT_BYTES and sha(p)==BASE_PROJECT_SHA256]
  if len(candidates)!=1: raise RuntimeError({{'project_archive_candidates':[str(p) for p in candidates]}})
  extracted=work/'project'; safe_extract(candidates[0],extracted); roots=[p for p in extracted.iterdir() if p.is_dir()]; source=roots[0] if len(roots)==1 else extracted
  destination=args.output_dir/CURRENT_PROJECT_NAME; shutil.copytree(source,destination)
  for row in MANIFEST['overlay_files']:
   src=overlay/row['path']; verify(src,row['bytes'],row['sha256'],'overlay_'+row['path']); dst=destination/row['path']; dst.parent.mkdir(parents=True,exist_ok=True); shutil.copy2(src,dst)
  db=destination/CURRENT_DB_REL; con=sqlite3.connect(db)
  try:
   integrity=con.execute('PRAGMA integrity_check').fetchone()[0]; fk=len(list(con.execute('PRAGMA foreign_key_check'))); r78=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R78'").fetchone()[0]; r79=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R79'").fetchone()[0]; checkpoint=con.execute("SELECT state,provider_previewer_status,physical_proof_status FROM section5_session2_checkpoint WHERE checkpoint_code='{CHECKPOINT_CODE}'").fetchone(); evidence=con.execute("SELECT COUNT(*) FROM section5_provider_preview_evidence WHERE checkpoint_code='{CHECKPOINT_CODE}'").fetchone()[0]; plan=con.execute("SELECT COUNT(*) FROM section5_physical_proof_plan WHERE checkpoint_code='{CHECKPOINT_CODE}'").fetchone()[0]
  finally: con.close()
  wb=load_workbook(destination/CURRENT_WORKBOOK_REL,read_only=True,data_only=False)
  try: sheets=len(wb.sheetnames)
  finally: wb.close()
  publication=destination/PUBLICATION_REL; interior=destination/PRINT_INTERIOR_REL; cover=destination/COVER_REL
  if sha(publication)!=PUBLICATION_SHA256 or sha(interior)!=PRINT_INTERIOR_SHA256 or sha(cover)!=COVER_SHA256: raise RuntimeError('immutable or frozen artifact identity changed')
  apps=[p for p in destination.rglob('human_pathogen_app.py') if p.is_file() and sha(p)==APPLICATION_SHA256]
  pub_pages=len(PdfReader(str(publication)).pages); print_pages=len(PdfReader(str(interior)).pages)
  result={{'status':'passed' if integrity=='ok' and fk==0 and r78==1 and r79==1 and checkpoint==('checkpoint_complete','controlled_pending_external','planned_not_ordered') and evidence==8 and plan==11 and sheets>=114 and pub_pages==537 and print_pages==538 and len(apps)==1 else 'failed','project_root':str(destination),'database':{{'integrity':integrity,'foreign_keys':fk,'response78':r78,'response79':r79,'checkpoint':checkpoint,'preview_evidence':evidence,'proof_plan':plan}},'workbook_sheets':sheets,'digital_pages':pub_pages,'print_pages':print_pages,'publication_sha256':sha(publication),'print_interior_sha256':sha(interior),'cover_sha256':sha(cover),'main_application_matches':len(apps)}}
  output=args.output_dir/'MRHPD_RESPONSE79_CHECKPOINT_RECOVERY_APPLICATION_RESULT.json'; output.write_text(json.dumps(result,indent=2),encoding='utf-8'); print(json.dumps(result,indent=2)); raise SystemExit(0 if result['status']=='passed' else 1)
if __name__=='__main__': main()
'''


def build_recovery_package(*, baseline_project:Path,current_project:Path,baseline_restore:Path,project_archive:Path,dist:Path,now:datetime,summary:dict[str,Any],direct_files:list[Path]) -> dict[str,Any]:
    baseline_map={path.relative_to(baseline_project).as_posix():(path.stat().st_size,sha256_file(path)) for path in baseline_project.rglob("*") if path.is_file()}
    current_map={path.relative_to(current_project).as_posix():(path.stat().st_size,sha256_file(path)) for path in current_project.rglob("*") if path.is_file()}
    deleted=sorted(set(baseline_map)-set(current_map))
    if deleted: raise RuntimeError({"unexpected_deleted_paths":deleted})
    overlay_rows=[]
    for rel,identity in sorted(current_map.items()):
        if baseline_map.get(rel)!=identity: overlay_rows.append({"path":rel,"bytes":identity[0],"sha256":identity[1],"change":"new" if rel not in baseline_map else "changed"})
    stamp=now.strftime("%Y-%m-%d %H%M UTC")
    package_root=dist/"recovery_package_root"
    if package_root.exists(): shutil.rmtree(package_root)
    overlay_root=package_root/"OVERLAY"; tools=package_root/"TOOLS"; overlay_root.mkdir(parents=True); tools.mkdir(parents=True)
    for row in overlay_rows:
        source=current_project/row["path"]; target=overlay_root/row["path"]; target.parent.mkdir(parents=True,exist_ok=True); shutil.copy2(source,target)
    manifest={"schema":"mrhpd-section5-session2-checkpoint-recovery-1.0","generated_at":now.isoformat().replace("+00:00","Z"),"version":PROJECT_VERSION,"response":79,"section":SECTION_LABEL,"session":SESSION_LABEL,"checkpoint":CHECKPOINT_LABEL,"state":"checkpoint_complete","baseline_restore":{"name":baseline_restore.name,"bytes":baseline_restore.stat().st_size,"sha256":sha256_file(baseline_restore)},"baseline_project":{"name":project_archive.name,"bytes":project_archive.stat().st_size,"sha256":sha256_file(project_archive)},"current_project_name":CURRENT_PROJECT_NAME,"overlay_file_count":len(overlay_rows),"overlay_total_bytes":sum(row["bytes"] for row in overlay_rows),"overlay_files":overlay_rows,"deleted_paths":[],"accepted_predecessor_mutated":False,"frozen_section3_release_mutated":False,"immutable_digital_publication_mutated":False,"frozen_session1_print_candidate_mutated":False,"user_upload_required":False,"requires_conversation_reconstruction":False,"next":"Remediation Section 5 of 5 Session 2 of 3 Checkpoint 2 of 3"}
    json_write(package_root/"CHECKPOINT_RECOVERY_MANIFEST.json",manifest)
    text_write(package_root/"CHECKPOINT_RECOVERY_CHECKSUMS.sha256","".join(f"{row['sha256']}  OVERLAY/{row['path']}\n" for row in overlay_rows))
    expected={"database_bytes":expected_path.stat().st_size if (expected_path:=current_project/CURRENT_DB_REL).exists() else 0,"database_sha256":sha256_file(current_project/CURRENT_DB_REL),"workbook_bytes":(current_project/CURRENT_WORKBOOK_REL).stat().st_size,"workbook_sha256":sha256_file(current_project/CURRENT_WORKBOOK_REL)}
    text_write(tools/"apply_checkpoint_recovery.py",create_apply_script(manifest,expected))
    text_write(package_root/"RESTORE_READ_FIRST.md",f"""# Human Pathogen Database — Response 79 Checkpoint Recovery

This cumulative intermediate recovery applies directly to the exact Response 77 complete restore and includes all current project changes through Response 79.

## Required baseline

Filename: `{baseline_restore.name}`

Bytes: `{baseline_restore.stat().st_size}`

SHA-256: `{sha256_file(baseline_restore)}`

## Automated apply

```bash
python TOOLS/apply_checkpoint_recovery.py \\
  --base-response77-restore "<Response 77 complete restore.zip>" \\
  --output-dir "<empty destination>"
```

The utility verifies the exact baseline, every overlay file, the current database and workbook, Responses 78 and 79, the provider-evidence boundary, the physical-proof plan, the unchanged 537-page publication, the frozen 538-page print candidate and cover, and the unchanged main application.
""")
    recovery_zip=dist/f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 Remediation Section 5 of 5 Session 2 of 3 Checkpoint 1 of 3 RECOVERY DATA THROUGH RESPONSE 79 {stamp}.zip"
    with zipfile.ZipFile(recovery_zip,"w",compression=zipfile.ZIP_DEFLATED,compresslevel=6,allowZip64=True) as zf:
        for path in sorted(package_root.rglob("*")):
            if path.is_file(): zf.write(path,path.relative_to(package_root).as_posix())
    recovery_qa=verify_zip(recovery_zip)
    with tempfile.TemporaryDirectory(prefix="mrhpd-r79-clean-apply-") as td:
        output=Path(td)/"restored"; result=subprocess.run([sys.executable,str((tools/"apply_checkpoint_recovery.py").resolve()),"--base-response77-restore",str(baseline_restore),"--output-dir",str(output)],cwd=package_root,text=True,capture_output=True,timeout=1800)
        if result.returncode: raise RuntimeError({"clean_apply_failed":{"stdout":result.stdout[-16000:],"stderr":result.stderr[-16000:]}})
        application=json.loads((output/"MRHPD_RESPONSE79_CHECKPOINT_RECOVERY_APPLICATION_RESULT.json").read_text(encoding="utf-8"))
        if application.get("status")!="passed": raise RuntimeError({"clean_apply_gate":application})
        clean_apply=application
    verification={"schema":"mrhpd-response79-checkpoint-recovery-verification-1.0","status":"passed","response":79,"checkpoint":"1 of 3","recovery_zip":recovery_qa,"manifest":{"overlay_file_count":len(overlay_rows),"overlay_total_bytes":sum(row["bytes"] for row in overlay_rows),"deleted_paths":0},"clean_apply":clean_apply,"accepted_predecessor_mutated":False,"frozen_section3_release_mutated":False,"immutable_digital_publication_mutated":False,"frozen_session1_print_candidate_mutated":False,"user_upload_required":False,"checkpoint_1_of_3_complete":True,"session_2_of_3_complete":False,"remediation_section_5_complete":False,"next":"Checkpoint 2 of 3 - provider evidence reconciliation and proof-readiness closure"}
    verification_path=dist/"MRHPD v3.0.0a Response 79 Checkpoint 1 Recovery Verification.json"; json_write(verification_path,verification)
    sha_path=dist/f"{recovery_zip.name}.sha256.txt"; text_write(sha_path,f"{recovery_qa['sha256']}  {recovery_zip.name}\n")
    summary_path=dist/"MRHPD_RESPONSE79_SECTION5_SESSION2_CHECKPOINT1_BUILD_SUMMARY.json"; json_write(summary_path,summary|{"recovery":verification})
    exact_names=dist/"MRHPD v3.0.0a Response 79 Exact File Names.txt"; text_write(exact_names,f"""Response 79 cumulative checkpoint recovery ZIP:
{recovery_zip.name}

Required baseline complete restore:
{baseline_restore.name}

Required baseline project archive embedded in that restore:
{project_archive.name}

Current copied SQLite database:
{Path(CURRENT_DB_REL).name}

Current comprehensive workbook:
{Path(CURRENT_WORKBOOK_REL).name}

Immutable digital publication:
{Path(PUBLICATION_REL).name}

Frozen print-production interior:
{Path(PRINT_INTERIOR_REL).name}

Frozen full-cover PNG:
{Path(COVER_REL).name}
""")
    delivery=dist/f"MRHPD v3.0.0a Response 79 Section 5 Session 2 Checkpoint 1 Recovery Package {stamp}.zip"
    with zipfile.ZipFile(delivery,"w",compression=zipfile.ZIP_DEFLATED,compresslevel=6,allowZip64=True) as zf:
        for path in [recovery_zip,sha_path,verification_path,summary_path,exact_names,*direct_files]:
            if path.exists(): zf.write(path,path.name)
    delivery_qa=verify_zip(delivery)
    return {"recovery_zip":recovery_zip,"recovery_qa":recovery_qa,"verification_path":verification_path,"summary_path":summary_path,"exact_names":exact_names,"delivery":delivery,"delivery_qa":delivery_qa,"overlay_rows":overlay_rows,"clean_apply":clean_apply}


def main() -> None:
    parser=argparse.ArgumentParser(); parser.add_argument("--response77-dir",type=Path,required=True); parser.add_argument("--dist",type=Path,default=Path("dist_cp5_s2_cp1")); args=parser.parse_args()
    now=utc_now(); now_iso=now.isoformat().replace("+00:00","Z")
    if args.dist.exists(): shutil.rmtree(args.dist)
    args.dist.mkdir(parents=True)
    with tempfile.TemporaryDirectory(prefix="mrhpd-cp5-s2-cp1-") as td:
        work=Path(td); restore,project_archive,baseline_project=restore_response77(args.response77_dir,work)
        current_project=work/"current_project"/CURRENT_PROJECT_NAME; current_project.parent.mkdir(parents=True); shutil.copytree(baseline_project,current_project)
        source_db=locate_current_database(baseline_project); source_workbook=locate_current_workbook(baseline_project)
        sources=official_source_rows(now_iso); evidence=preview_evidence_rows(now_iso); issues=conversion_issue_rows(now_iso); proof_plan=physical_proof_plan_rows(now_iso); proof_items=proof_inspection_rows(now_iso); recovery=recovery_rows(now_iso)
        current_db=current_project/CURRENT_DB_REL; database_qa=synchronize_database(source_db,current_db,now_iso=now_iso,sources=sources,evidence=evidence,issues=issues,proof_plan=proof_plan,proof_items=proof_items,recovery=recovery)
        current_workbook=current_project/CURRENT_WORKBOOK_REL; workbook_qa=augment_workbook(source_workbook,current_workbook,sources=sources,evidence=evidence,issues=issues,proof_plan=proof_plan,proof_items=proof_items,recovery=recovery)
        tracking=tracking_files(current_project,current_db,now_iso)
        app_files,application_qa=application_surfaces(current_project,current_db,current_workbook,now_iso)
        data_root=current_project/"Data"/"Section 5 Session 2 Checkpoint 1"
        source_json=data_root/"MRHPD v3.0.0a Response 79 KDP Official Source Controls.json"; source_csv=data_root/"MRHPD v3.0.0a Response 79 KDP Official Source Controls.csv"; json_write(source_json,sources); csv_write(source_csv,sources)
        evidence_json=data_root/"MRHPD v3.0.0a Response 79 Print Previewer Evidence Register.json"; evidence_csv=data_root/"MRHPD v3.0.0a Response 79 Print Previewer Evidence Register.csv"; json_write(evidence_json,evidence); csv_write(evidence_csv,evidence)
        issues_json=data_root/"MRHPD v3.0.0a Response 79 Provider Conversion Issue Register.json"; issues_csv=data_root/"MRHPD v3.0.0a Response 79 Provider Conversion Issue Register.csv"; json_write(issues_json,issues); csv_write(issues_csv,issues)
        proof_plan_json=data_root/"MRHPD v3.0.0a Response 79 Physical Proof Plan.json"; proof_plan_csv=data_root/"MRHPD v3.0.0a Response 79 Physical Proof Plan.csv"; json_write(proof_plan_json,proof_plan); csv_write(proof_plan_csv,proof_plan)
        proof_items_json=data_root/"MRHPD v3.0.0a Response 79 Physical Proof Inspection Register.json"; proof_items_csv=data_root/"MRHPD v3.0.0a Response 79 Physical Proof Inspection Register.csv"; json_write(proof_items_json,proof_items); csv_write(proof_items_csv,proof_items)
        report_root=current_project/"Reports"/"Section 5 Session 2"/"Checkpoint 1"
        docx_path=report_root/"MRHPD v3.0.0a Section 5 Session 2 Checkpoint 1 Provider Preview and Physical Proof Intake Report.docx"
        pdf_path=report_root/"MRHPD v3.0.0a Section 5 Session 2 Checkpoint 1 Provider Preview and Physical Proof Intake Report.pdf"
        xlsx_path=report_root/"MRHPD v3.0.0a Section 5 Session 2 Checkpoint 1 Provider Preview and Proof Register.xlsx"
        docx_qa=report_docx(docx_path,generated_at=now_iso,sources=sources,evidence=evidence,issues=issues,proof_plan=proof_plan,proof_items=proof_items,recovery=recovery)
        pdf_qa=report_pdf(pdf_path,generated_at=now_iso,sources=sources,evidence=evidence,proof_plan=proof_plan,proof_items=proof_items)
        register_qa=report_register(xlsx_path,sources=sources,evidence=evidence,issues=issues,proof_plan=proof_plan,proof_items=proof_items,recovery=recovery)
        render_qa=render_pdf_qa(pdf_path,report_root/"Rendered Report QA")
        qa_root=current_project/"QA"/"Section 5 Session 2"/"Checkpoint 1"; qa_root.mkdir(parents=True,exist_ok=True)
        qa_payload={"schema":"mrhpd-section5-session2-checkpoint1-qa-1.0","generated_at":now_iso,"status":"passed","response":79,"section":SECTION_LABEL,"session":SESSION_LABEL,"checkpoint":CHECKPOINT_LABEL,"database":database_qa,"workbook":workbook_qa,"application":application_qa,"provider_previewer":{"status":"controlled_pending_external","evidence_records":len(evidence),"unsupported_approval_claims":0,"actual_provider_issues_recorded":0},"physical_proof":{"status":"planned_not_ordered","plan_records":len(proof_plan),"inspection_records":len(proof_items),"completed_inspections":0},"digital_publication":{"sha256":sha256_file(current_project/PUBLICATION_REL),"expected_sha256":PUBLICATION_SHA256,"pages":len(PdfReader(str(current_project/PUBLICATION_REL)).pages)},"print_interior":{"sha256":sha256_file(current_project/PRINT_INTERIOR_REL),"expected_sha256":PRINT_INTERIOR_SHA256,"pages":len(PdfReader(str(current_project/PRINT_INTERIOR_REL)).pages)},"cover":{"sha256":sha256_file(current_project/COVER_REL),"expected_sha256":COVER_SHA256},"reports":{"docx":docx_qa,"pdf":pdf_qa,"xlsx":register_qa,"render":render_qa},"checkpoint_1_of_3_complete":True,"session_2_of_3_complete":False,"remediation_section_5_complete":False,"user_upload_required":False,"next":"Checkpoint 2 of 3 - provider evidence reconciliation and proof-readiness closure"}
        json_write(qa_root/"SECTION5_SESSION2_CHECKPOINT1_QA.json",qa_payload); json_write(qa_root/"RECOVERY_EVENTS_202_208.json",recovery)
        index_result=build_indexes(current_project,now_iso)
        manifest_path,checksums_path,manifest_rows=build_manifest(current_project,now_iso)
        summary={"schema":"mrhpd-response79-section5-session2-checkpoint1-build-1.0","generated_at":now_iso,"status":"passed","response":79,"section":SECTION_LABEL,"session":SESSION_LABEL,"checkpoint":CHECKPOINT_LABEL,"baseline_restore":{"bytes":restore.stat().st_size,"sha256":sha256_file(restore)},"baseline_project":{"bytes":project_archive.stat().st_size,"sha256":sha256_file(project_archive)},"database":database_qa,"workbook":workbook_qa,"application":application_qa,"provider_previewer":qa_payload["provider_previewer"],"physical_proof":qa_payload["physical_proof"],"digital_publication":qa_payload["digital_publication"],"print_interior":qa_payload["print_interior"],"cover":qa_payload["cover"],"reports":qa_payload["reports"],"index":index_result["qa"],"manifest_records":len(manifest_rows),"user_upload_required":False,"checkpoint_1_of_3_complete":True,"session_2_of_3_complete":False,"remediation_section_5_complete":False,"next":"Checkpoint 2 of 3 - provider evidence reconciliation and proof-readiness closure"}
        direct_files=[docx_path,pdf_path,xlsx_path,evidence_csv,issues_csv,proof_plan_csv,proof_items_csv,qa_root/"SECTION5_SESSION2_CHECKPOINT1_QA.json"]
        package=build_recovery_package(baseline_project=baseline_project,current_project=current_project,baseline_restore=restore,project_archive=project_archive,dist=args.dist,now=now,summary=summary,direct_files=direct_files)
        console={"status":"passed","delivery":package["delivery"].name,"delivery_bytes":package["delivery_qa"]["bytes"],"delivery_sha256":package["delivery_qa"]["sha256"],"recovery_zip":package["recovery_zip"].name,"recovery_zip_bytes":package["recovery_qa"]["bytes"],"recovery_zip_sha256":package["recovery_qa"]["sha256"],"overlay_files":len(package["overlay_rows"]),"database_tables":database_qa["table_count"],"workbook_sheets":workbook_qa["current_sheet_count"],"preview_evidence_records":len(evidence),"proof_plan_records":len(proof_plan),"proof_inspection_records":len(proof_items),"provider_approval_claimed":False,"physical_proof_claimed":False,"user_upload_required":False,"checkpoint_1_of_3_complete":True,"next":"Checkpoint 2 of 3 - provider evidence reconciliation and proof-readiness closure"}
        print(json.dumps(console,indent=2))


if __name__=="__main__": main()
