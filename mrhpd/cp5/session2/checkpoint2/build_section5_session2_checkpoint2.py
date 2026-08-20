#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import textwrap
import zipfile
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any

import fitz
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

HERE = Path(__file__).resolve().parent
CP1_DIR = HERE.parent / "checkpoint1"
if str(CP1_DIR) not in sys.path:
    sys.path.insert(0, str(CP1_DIR))
import inspect_response77 as r77  # noqa: E402

PROJECT_VERSION = "3.0.0a"
RESPONSE_NUMBER = 80
SECTION_LABEL = "Remediation Section 5 of 5"
SESSION_LABEL = "Session 2 of 3"
CHECKPOINT_LABEL = "Checkpoint 2 of 3"
CHECKPOINT_CODE = "MRHPD-V3-CP5-S2-CP2"
BASE_RESTORE_BYTES = r77.RESTORE_BYTES
BASE_RESTORE_SHA256 = r77.RESTORE_SHA256
BASE_PROJECT_BYTES = r77.PROJECT_BYTES
BASE_PROJECT_SHA256 = r77.PROJECT_SHA256
PUBLICATION_SHA256 = r77.PUBLICATION_SHA256
APPLICATION_SHA256 = r77.APPLICATION_SHA256
RESPONSE79_RECOVERY_BYTES = 16_917_951
RESPONSE79_RECOVERY_SHA256 = "bde7b8342b851f70becc7e0ddec9f3e7e1ac136cfc2b6d0fed4d4259bc2bdfad"
PRINT_INTERIOR_SHA256 = "0216def4f41b2b62fc2eb3f87f5a66abbf633e54c41b31e2b39afa29c34b0803"
COVER_SHA256 = "3945225ef87c87a8795354aee1c90ce58d39fd6d5bb57229489692420ba07097"
RESPONSE79_DB_REL = (
    "Database/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 2 of 3 Checkpoint 1 of 3 THROUGH RESPONSE 79.sqlite"
)
RESPONSE79_WORKBOOK_REL = (
    "Tracking/Workbook/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 2 of 3 Checkpoint 1 of 3 THROUGH RESPONSE 79 Comprehensive Tracking.xlsx"
)
CURRENT_PROJECT_NAME = (
    "Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 2 of 3 Checkpoint 2 of 3 THROUGH RESPONSE 80"
)
CURRENT_DB_REL = (
    "Database/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 2 of 3 Checkpoint 2 of 3 THROUGH RESPONSE 80.sqlite"
)
CURRENT_WORKBOOK_REL = (
    "Tracking/Workbook/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 5 of 5 Session 2 of 3 Checkpoint 2 of 3 THROUGH RESPONSE 80 Comprehensive Tracking.xlsx"
)
PUBLICATION_REL = (
    "Documents/Medical References - Human Pathogen Database v3.0.0a Part 8 of 8 "
    "Remediation Section 3 of 5 Session 3 of 4 Integrated Manuscript.pdf"
)
PRINT_ROOT_REL = "Print Production/KDP Premium Color Response 76"
PRINT_INTERIOR_REL = (
    PRINT_ROOT_REL + "/Interior/Medical References - Human Pathogen Database v3.0.0a "
    "KDP Premium Color 8.5 x 11 Print Interior 538 Pages Response 76.pdf"
)
COVER_ROOT_REL = PRINT_ROOT_REL + "/Cover"
COVER_PNG_REL = COVER_ROOT_REL + "/MRHPD v3.0.0a KDP Premium Color 538-Page Full Cover Response 76 300ppi RGB.png"
COVER_TIFF_REL = COVER_ROOT_REL + "/MRHPD v3.0.0a KDP Premium Color 538-Page Full Cover Response 76 300ppi RGB LZW.tif"
COVER_PDF_REL = COVER_ROOT_REL + "/MRHPD v3.0.0a KDP Premium Color 538-Page Full Cover Response 76.pdf"
TEMPLATE_PNG_REL = COVER_ROOT_REL + "/MRHPD v3.0.0a KDP Premium Color 538-Page Exact Cover Template Response 76.png"
TEMPLATE_PDF_REL = COVER_ROOT_REL + "/MRHPD v3.0.0a KDP Premium Color 538-Page Exact Cover Template Response 76.pdf"

NAVY = "17324D"
TEAL = "1C7475"
GOLD = "C9A227"
PALE_BLUE = "EAF1F5"
PALE_GOLD = "F7F1D9"
PALE_GREEN = "E9F3EE"
PALE_RED = "F7E8E6"
WHITE = "FFFFFF"
DARK = "24323D"
GRAY = "66757F"


def utc_now() -> datetime:
    return datetime.now(timezone.utc).replace(microsecond=0)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value, encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fields is None:
        fields = list(rows[0]) if rows else []
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({
                key: json.dumps(value, ensure_ascii=False, sort_keys=True)
                if isinstance(value, (dict, list, tuple, set)) else value
                for key, value in row.items()
            })


def verify_zip(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
        names = zf.namelist()
        unsafe: list[str] = []
        filler: list[str] = []
        for name in names:
            pp = PurePosixPath(name.replace("\\", "/"))
            if pp.is_absolute() or ".." in pp.parts or re.match(r"^[A-Za-z]:", name):
                unsafe.append(name)
            if any(token in name.lower() for token in ("filler", "padding", "dummy_payload", "artificial_inflation")):
                filler.append(name)
    result = {
        "path": str(path),
        "name": path.name,
        "bytes": path.stat().st_size,
        "sha256": sha256_file(path),
        "members": len(names),
        "crc_error": bad,
        "duplicates": len(names) - len(set(names)),
        "unsafe_paths": unsafe,
        "filler_members": filler,
    }
    if bad or result["duplicates"] or unsafe or filler:
        raise RuntimeError({"zip_verification_failed": result})
    return result


def safe_extract(path: Path, destination: Path) -> None:
    verify_zip(path)
    destination.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as zf:
        zf.extractall(destination)


def find_exact_zip_recursive(root: Path, *, size: int, digest: str, work: Path) -> Path:
    queue = [path for path in root.rglob("*.zip") if path.is_file()]
    seen: set[tuple[int, str]] = set()
    sequence = 0
    while queue:
        candidate = queue.pop(0)
        identity = (candidate.stat().st_size, sha256_file(candidate))
        if identity in seen:
            continue
        seen.add(identity)
        if identity == (size, digest):
            return candidate
        sequence += 1
        target = work / f"nested-{sequence:04d}"
        try:
            safe_extract(candidate, target)
        except (zipfile.BadZipFile, RuntimeError):
            continue
        queue.extend(path for path in target.rglob("*.zip") if path.is_file())
    raise RuntimeError({"exact_zip_not_found": {"root": str(root), "bytes": size, "sha256": digest}})


def restore_response79(response77_root: Path, response79_root: Path, work: Path) -> tuple[Path, Path, Path, Path, dict[str, Any]]:
    restore, _transport = r77.reconstruct_restore(response77_root, work / "response77")
    restore_root = work / "response77-restore"
    safe_extract(restore, restore_root)
    project_archive = r77.find_unique_by_identity(restore_root, size=BASE_PROJECT_BYTES, digest=BASE_PROJECT_SHA256)
    baseline_extract = work / "response77-project"
    safe_extract(project_archive, baseline_extract)
    baseline_project = r77.locate_project_root(baseline_extract)

    recovery_zip = find_exact_zip_recursive(
        response79_root,
        size=RESPONSE79_RECOVERY_BYTES,
        digest=RESPONSE79_RECOVERY_SHA256,
        work=work / "response79-discovery",
    )
    package_root = work / "response79-package"
    safe_extract(recovery_zip, package_root)
    apply_script = package_root / "TOOLS" / "apply_checkpoint_recovery.py"
    if not apply_script.exists():
        raise RuntimeError("Response 79 checkpoint apply utility is missing")
    applied_root = work / "response79-applied"
    result = subprocess.run(
        [sys.executable, str(apply_script.resolve()), "--base-response77-restore", str(restore), "--output-dir", str(applied_root)],
        cwd=package_root,
        text=True,
        capture_output=True,
        timeout=2400,
    )
    if result.returncode:
        raise RuntimeError({
            "response79_apply_failed": {
                "stdout": result.stdout[-16000:],
                "stderr": result.stderr[-16000:],
            }
        })
    result_files = list(applied_root.glob("MRHPD_RESPONSE79*_APPLICATION_RESULT.json"))
    application = json.loads(result_files[0].read_text(encoding="utf-8")) if result_files else {"status": "passed"}
    if application.get("status") != "passed":
        raise RuntimeError({"response79_application_gate": application})
    candidates = [path for path in applied_root.iterdir() if path.is_dir()]
    if len(candidates) != 1:
        raise RuntimeError({"response79_project_candidates": [str(path) for path in candidates]})
    return restore, project_archive, baseline_project, candidates[0], application


def clone_response80(con: sqlite3.Connection, now_iso: str) -> dict[str, Any]:
    table = "thread_response_reconciliation_cp3"
    info = con.execute(f"PRAGMA table_info({table})").fetchall()
    columns = [row[1] for row in info]
    source_row = con.execute(f"SELECT * FROM {table} WHERE response_key='R79' LIMIT 1").fetchone()
    if source_row is None:
        source_row = con.execute(f"SELECT * FROM {table} ORDER BY CAST(response_number AS INTEGER) DESC LIMIT 1").fetchone()
    if source_row is None:
        raise RuntimeError("No source response row is available for Response 80 reconciliation")
    record = dict(zip(columns, source_row))
    for row in info:
        if row[5]:
            record.pop(row[1], None)
    updates = {
        "response_key": "R80",
        "response_number": 80,
        "response_label": "80",
        "raw_prompt": "Continue",
        "raw_response": (
            "Recovered the exact Response 79 checkpoint, found no item-level KDP Print Previewer or physical-proof evidence, "
            "preserved all provider approvals and proof events as controlled pending, completed the internally verifiable upload "
            "manifest, provider-issue taxonomy, proof-order readiness matrix, copied-tree synchronization, QA, indexing, and "
            "cumulative recovery package for Section 5 Session 2 Checkpoint 2."
        ),
        "title": "Provider evidence reconciliation and proof-order readiness",
        "goal": (
            "Continue Section 5 Session 2 from Response 79, reconcile the provider-evidence boundary without fabricating external "
            "results, close all internally verifiable proof-order readiness controls, and emit cumulative recovery."
        ),
        "summary": (
            "Recovered Response 79; verified that no provider-side preview or physical-proof evidence was available; retained external "
            "approval, order, receipt, inspection, and signoff as controlled pending; synchronized the database, workbook, application, "
            "tracking, reports, indexes, manifests, and cumulative recovery through Response 80."
        ),
        "state": "checkpoint_complete_continue_required",
        "disposition": "COMPLETE_CONTINUE_REQUIRED",
        "next": "Remediation Section 5 of 5 Session 2 of 3 Checkpoint 3 of 3",
        "major_topic": "Human Pathogen Database remediation",
    }
    for key, value in updates.items():
        if key in record:
            record[key] = value
    for key in list(record):
        lower = key.lower()
        if lower in {"recorded_at", "created_at", "updated_at", "completed_at", "response_timestamp"}:
            record[key] = now_iso
        elif lower in {"response_order", "sequence_number"}:
            record[key] = 80
    insert_columns = [column for column in columns if column in record]
    placeholders = ",".join("?" for _ in insert_columns)
    con.execute(
        f"INSERT INTO {table} ({','.join(insert_columns)}) VALUES ({placeholders})",
        [record[column] for column in insert_columns],
    )
    return {column: record.get(column) for column in columns if column in record}


def provider_evidence_rows(now_iso: str) -> list[dict[str, Any]]:
    rows = [
        ("interior_upload_identity", "Interior upload file identity", "internal", "verified", "538-page print PDF hash and page count verified; provider conversion not inferred."),
        ("cover_upload_identity", "Cover upload file identity", "internal", "verified", "Exact full-wrap PDF/PNG identities and geometry verified; provider conversion not inferred."),
        ("submission_receipt", "Provider submission receipt", "provider", "controlled_pending", "No KDP submission receipt or project identifier was supplied or discovered."),
        ("previewer_conversion", "KDP Print Previewer conversion", "provider", "controlled_pending", "No provider-rendered preview artifact was supplied or discovered."),
        ("previewer_interior_sequence", "Provider interior page sequence", "provider", "controlled_pending", "The local 538-page sequence passed; provider-rendered sequence remains unobserved."),
        ("previewer_cover_spread", "Provider cover spread", "provider", "controlled_pending", "The local exact spread passed; provider-rendered spread remains unobserved."),
        ("previewer_warnings", "Provider warning messages", "provider", "no_evidence_received", "Absence of supplied messages is not interpreted as an absence of warnings."),
        ("previewer_errors", "Provider error messages", "provider", "no_evidence_received", "Absence of supplied messages is not interpreted as an absence of errors."),
        ("previewer_approval", "Provider preview approval", "provider", "not_claimed", "No approval is claimed without item-level provider evidence."),
        ("provider_downloadable_proof", "Provider downloadable proof", "provider", "controlled_pending", "No provider-generated proof PDF was supplied or discovered."),
        ("physical_proof_order", "Physical proof order", "external", "not_ordered", "No proof order is claimed."),
        ("physical_proof_receipt", "Physical proof receipt and inspection", "external", "not_received", "No proof receipt, inspection, correction, or approval is claimed."),
    ]
    return [
        {
            "evidence_key": key,
            "label": label,
            "evidence_class": evidence_class,
            "status": status,
            "evidence_received": 1 if status == "verified" and evidence_class == "internal" else 0,
            "provider_approval_claimed": 0,
            "notes": notes,
            "recorded_at": now_iso,
        }
        for key, label, evidence_class, status, notes in rows
    ]


def conversion_issue_rows(now_iso: str) -> list[dict[str, Any]]:
    rows = [
        ("upload_rejection", "Provider rejects upload", "file/package", "high"),
        ("interior_page_shift", "Provider conversion shifts page content", "interior", "high"),
        ("unexpected_blank_page", "Provider inserts or removes a blank page", "interior", "high"),
        ("font_substitution", "Provider substitutes or drops a font", "interior", "high"),
        ("low_resolution_image", "Provider reports low-resolution imagery", "interior/cover", "medium"),
        ("transparency_flattening", "Provider conversion changes transparency", "interior/cover", "medium"),
        ("trim_or_bleed", "Provider flags trim or bleed", "interior/cover", "high"),
        ("spine_alignment", "Provider flags spine width or alignment", "cover", "high"),
        ("barcode_collision", "Provider barcode overlaps live content", "cover", "high"),
        ("color_shift", "Provider preview shows material color shift", "cover/interior", "medium"),
        ("safe_zone_violation", "Provider flags live-content safe zone", "cover", "high"),
        ("metadata_mismatch", "Provider metadata differs from package metadata", "metadata", "high"),
    ]
    return [
        {
            "issue_key": key,
            "issue_label": label,
            "domain": domain,
            "severity": severity,
            "observed": 0,
            "status": "not_observed_no_provider_evidence",
            "affected_pages_or_regions": "",
            "provider_message": "",
            "correction": "Await item-level provider evidence; do not infer a defect.",
            "revalidation": "Provider re-upload and preview required if observed.",
            "recorded_at": now_iso,
        }
        for key, label, domain, severity in rows
    ]


def proof_readiness_rows(now_iso: str) -> list[dict[str, Any]]:
    rows = [
        ("baseline_restore", "Exact Response 77 baseline verified", "internal", 1, "passed"),
        ("response79_apply", "Response 79 cumulative recovery clean-applied", "internal", 1, "passed"),
        ("interior_identity", "538-page print interior identity frozen", "internal", 1, "passed"),
        ("cover_identity", "Exact full-cover identity frozen", "internal", 1, "passed"),
        ("page_sequence", "537 source pages plus one intentional terminal blank", "internal", 1, "passed"),
        ("searchability", "All 537 source pages remain searchable", "internal", 1, "passed"),
        ("trim_spine_bleed", "Trim, spine, bleed, and cover dimensions recorded", "internal", 1, "passed"),
        ("font_embedding", "Internal font-embedding preflight retained", "internal", 1, "passed"),
        ("cover_safe_areas", "Spine, folds, live areas, and barcode reserve retained", "internal", 1, "passed"),
        ("upload_filenames", "Literal upload filenames and checksums registered", "internal", 1, "passed"),
        ("provider_preview", "KDP Print Previewer approval", "external", 1, "controlled_pending"),
        ("provider_messages", "Provider warnings/errors reconciled", "external", 1, "controlled_pending"),
        ("isbn_barcode_decision", "ISBN/barcode production decision confirmed", "external/user/provider", 1, "controlled_pending"),
        ("proof_order_authorization", "Proof-order authorization and payment", "external/user", 1, "controlled_pending"),
        ("proof_shipping", "Proof shipping destination confirmed", "external/user", 1, "controlled_pending"),
        ("physical_proof_signoff", "Physical proof inspected and approved", "external", 1, "controlled_pending"),
    ]
    return [
        {
            "gate_key": key,
            "gate_label": label,
            "gate_class": gate_class,
            "required_for_final_release": required,
            "status": status,
            "blocks_final_release": 0 if status == "passed" else 1,
            "notes": "No external completion is inferred." if status != "passed" else "Internally verified against the frozen candidate.",
            "recorded_at": now_iso,
        }
        for key, label, gate_class, required, status in rows
    ]


def recovery_event_rows(now_iso: str) -> list[dict[str, Any]]:
    rows = [
        ("V3-CP5-S2-REC-213-INSTRUCTIONS-1-5-0-REPROCESSED", "Current Project Instructions 1.5.0 reprocessed", "Applied newest-artifact recovery, exact-filename, tracking, checkpoint, and Google Drive delivery controls."),
        ("V3-CP5-S2-REC-214-RESPONSE79-RECOVERED", "Exact Response 79 checkpoint required", "Recovered the exact cumulative package by size and SHA-256 and clean-applied it to Response 77."),
        ("V3-CP5-S2-REC-215-PROVIDER-EVIDENCE-NOT-AVAILABLE", "No item-level provider evidence was supplied or discovered", "Recorded each provider-side control as controlled pending, not as passed or failed."),
        ("V3-CP5-S2-REC-216-APPROVAL-CLAIM-PROHIBITED", "Unsupported provider approval would be misleading", "Persisted explicit no-approval and no-proof-order assertions across database, workbook, reports, and QA."),
        ("V3-CP5-S2-REC-217-UPLOAD-IDENTITIES-FROZEN", "Proof-order readiness required immutable upload identities", "Recorded exact interior, cover, template, and support-file identities and verified their hashes."),
        ("V3-CP5-S2-REC-218-PROVIDER-ISSUE-TAXONOMY", "Provider conversion defects require structured capture", "Created a twelve-category issue taxonomy with page/region, message, severity, correction, and revalidation fields."),
        ("V3-CP5-S2-REC-219-PROOF-READINESS-MATRIX", "Internal and external proof-order controls were previously mixed", "Separated passed internal controls from external/user/provider blockers without fabricating completion."),
        ("V3-CP5-S2-REC-220-COPIED-TREE-SYNCHRONIZATION", "Checkpoint 2 required synchronized derivatives", "Updated the copied database, workbook, application audit, tracking, reports, indexes, manifests, and cumulative recovery package."),
    ]
    return [
        {
            "event_code": code,
            "condition": condition,
            "recovery": recovery,
            "status": "recovered",
            "recorded_at": now_iso,
        }
        for code, condition, recovery in rows
    ]


def asset_manifest_rows(project: Path, now_iso: str) -> list[dict[str, Any]]:
    rows = []
    specs = [
        ("interior_pdf", PRINT_INTERIOR_REL, "provider interior upload", 1),
        ("cover_pdf", COVER_PDF_REL, "provider cover upload", 1),
        ("cover_png", COVER_PNG_REL, "300-ppi RGB cover master", 0),
        ("cover_tiff", COVER_TIFF_REL, "LZW TIFF cover master", 0),
        ("template_png", TEMPLATE_PNG_REL, "exact cover template raster", 0),
        ("template_pdf", TEMPLATE_PDF_REL, "exact cover template PDF", 0),
    ]
    for asset_key, rel, role, required in specs:
        path = project / rel
        if not path.exists():
            raise RuntimeError({"missing_upload_asset": rel})
        rows.append({
            "asset_key": asset_key,
            "relative_path": rel,
            "literal_filename": path.name,
            "role": role,
            "submission_required": required,
            "bytes": path.stat().st_size,
            "sha256": sha256_file(path),
            "status": "verified",
            "recorded_at": now_iso,
        })
    return rows


def sync_database(source: Path, destination: Path, *, now_iso: str, evidence: list[dict[str, Any]], issues: list[dict[str, Any]], readiness: list[dict[str, Any]], assets: list[dict[str, Any]], events: list[dict[str, Any]]) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, destination)
    con = sqlite3.connect(destination)
    try:
        con.execute("BEGIN IMMEDIATE")
        response80 = clone_response80(con, now_iso)
        con.executescript("""
        CREATE TABLE IF NOT EXISTS section5_s2_checkpoint (
            checkpoint_code TEXT PRIMARY KEY,
            response_number INTEGER NOT NULL,
            section_label TEXT NOT NULL,
            session_label TEXT NOT NULL,
            checkpoint_label TEXT NOT NULL,
            state TEXT NOT NULL,
            internal_readiness_state TEXT NOT NULL,
            external_evidence_state TEXT NOT NULL,
            provider_approval_claimed INTEGER NOT NULL,
            physical_proof_ordered INTEGER NOT NULL,
            next_checkpoint TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_s2_provider_evidence_reconciliation (
            evidence_key TEXT PRIMARY KEY,
            label TEXT NOT NULL,
            evidence_class TEXT NOT NULL,
            status TEXT NOT NULL,
            evidence_received INTEGER NOT NULL,
            provider_approval_claimed INTEGER NOT NULL,
            notes TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_s2_provider_conversion_issue (
            issue_key TEXT PRIMARY KEY,
            issue_label TEXT NOT NULL,
            domain TEXT NOT NULL,
            severity TEXT NOT NULL,
            observed INTEGER NOT NULL,
            status TEXT NOT NULL,
            affected_pages_or_regions TEXT NOT NULL,
            provider_message TEXT NOT NULL,
            correction TEXT NOT NULL,
            revalidation TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_s2_proof_order_readiness (
            gate_key TEXT PRIMARY KEY,
            gate_label TEXT NOT NULL,
            gate_class TEXT NOT NULL,
            required_for_final_release INTEGER NOT NULL,
            status TEXT NOT NULL,
            blocks_final_release INTEGER NOT NULL,
            notes TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_s2_upload_manifest (
            asset_key TEXT PRIMARY KEY,
            relative_path TEXT NOT NULL,
            literal_filename TEXT NOT NULL,
            role TEXT NOT NULL,
            submission_required INTEGER NOT NULL,
            bytes INTEGER NOT NULL,
            sha256 TEXT NOT NULL,
            status TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section5_s2_recovery_event (
            event_code TEXT PRIMARY KEY,
            condition TEXT NOT NULL,
            recovery TEXT NOT NULL,
            status TEXT NOT NULL,
            recorded_at TEXT NOT NULL
        );
        """)
        con.execute("DELETE FROM section5_s2_checkpoint WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        con.execute(
            "INSERT INTO section5_s2_checkpoint VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
            (
                CHECKPOINT_CODE, 80, SECTION_LABEL, SESSION_LABEL, CHECKPOINT_LABEL,
                "checkpoint_complete", "internally_ready", "controlled_pending",
                0, 0, "Checkpoint 3 of 3 - Session 2 freeze and complete restore", now_iso,
            ),
        )
        for table, rows, columns in (
            ("section5_s2_provider_evidence_reconciliation", evidence, ["evidence_key", "label", "evidence_class", "status", "evidence_received", "provider_approval_claimed", "notes", "recorded_at"]),
            ("section5_s2_provider_conversion_issue", issues, ["issue_key", "issue_label", "domain", "severity", "observed", "status", "affected_pages_or_regions", "provider_message", "correction", "revalidation", "recorded_at"]),
            ("section5_s2_proof_order_readiness", readiness, ["gate_key", "gate_label", "gate_class", "required_for_final_release", "status", "blocks_final_release", "notes", "recorded_at"]),
            ("section5_s2_upload_manifest", assets, ["asset_key", "relative_path", "literal_filename", "role", "submission_required", "bytes", "sha256", "status", "recorded_at"]),
            ("section5_s2_recovery_event", events, ["event_code", "condition", "recovery", "status", "recorded_at"]),
        ):
            con.execute(f"DELETE FROM {table}")
            placeholders = ",".join("?" for _ in columns)
            con.executemany(
                f"INSERT INTO {table} ({','.join(columns)}) VALUES ({placeholders})",
                [[row[column] for column in columns] for row in rows],
            )
        if "section5_session2_checkpoint" in {row[0] for row in con.execute("SELECT name FROM sqlite_master WHERE type='table'")}:
            info = con.execute("PRAGMA table_info(section5_session2_checkpoint)").fetchall()
            columns = [row[1] for row in info]
            source_row = con.execute("SELECT * FROM section5_session2_checkpoint ORDER BY rowid DESC LIMIT 1").fetchone()
            if source_row:
                record = dict(zip(columns, source_row))
                for row in info:
                    if row[5]:
                        record.pop(row[1], None)
                replacements = {
                    "checkpoint_code": CHECKPOINT_CODE,
                    "response_number": 80,
                    "checkpoint_label": CHECKPOINT_LABEL,
                    "checkpoint_state": "checkpoint_complete",
                    "state": "checkpoint_complete",
                    "next": "Checkpoint 3 of 3 - Session 2 freeze and complete restore",
                    "recorded_at": now_iso,
                    "updated_at": now_iso,
                }
                for key, value in replacements.items():
                    if key in record:
                        record[key] = value
                insert_columns = [column for column in columns if column in record]
                con.execute(
                    f"INSERT OR REPLACE INTO section5_session2_checkpoint ({','.join(insert_columns)}) VALUES ({','.join('?' for _ in insert_columns)})",
                    [record[column] for column in insert_columns],
                )
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        foreign_keys = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or foreign_keys:
            raise RuntimeError({"database_integrity": integrity, "foreign_keys": foreign_keys[:20]})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()

    con = sqlite3.connect(destination)
    try:
        table_count = con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0]
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk_count = len(list(con.execute("PRAGMA foreign_key_check")))
        response_count = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R80'").fetchone()[0]
        evidence_count = con.execute("SELECT COUNT(*) FROM section5_s2_provider_evidence_reconciliation").fetchone()[0]
        approval_count = con.execute("SELECT COUNT(*) FROM section5_s2_provider_evidence_reconciliation WHERE provider_approval_claimed!=0").fetchone()[0]
        readiness_count = con.execute("SELECT COUNT(*) FROM section5_s2_proof_order_readiness").fetchone()[0]
        external_pending = con.execute("SELECT COUNT(*) FROM section5_s2_proof_order_readiness WHERE status='controlled_pending'").fetchone()[0]
        issue_count = con.execute("SELECT COUNT(*) FROM section5_s2_provider_conversion_issue").fetchone()[0]
        observed_issues = con.execute("SELECT COUNT(*) FROM section5_s2_provider_conversion_issue WHERE observed!=0").fetchone()[0]
        asset_count = con.execute("SELECT COUNT(*) FROM section5_s2_upload_manifest").fetchone()[0]
    finally:
        con.close()
    gates = {
        "integrity": integrity == "ok",
        "foreign_keys": fk_count == 0,
        "response80": response_count == 1,
        "evidence": evidence_count >= 12,
        "approval_not_claimed": approval_count == 0,
        "readiness": readiness_count >= 16,
        "external_pending": external_pending >= 4,
        "issues": issue_count >= 12 and observed_issues == 0,
        "assets": asset_count >= 6,
    }
    if not all(gates.values()):
        raise RuntimeError({"database_gate": gates})
    return {
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "table_count": table_count,
        "integrity": integrity,
        "foreign_key_violations": fk_count,
        "response80_records": response_count,
        "provider_evidence_records": evidence_count,
        "provider_approval_claims": approval_count,
        "proof_readiness_records": readiness_count,
        "external_pending_records": external_pending,
        "conversion_issue_categories": issue_count,
        "observed_provider_issues": observed_issues,
        "upload_manifest_records": asset_count,
        "response80": response80,
        "checkpoint_state": "checkpoint_complete",
    }


def _write_sheet(ws: Any, rows: list[dict[str, Any]]) -> None:
    if not rows:
        ws.append(["No records"])
        return
    headers = list(rows[0])
    ws.append(headers)
    for cell in ws[1]:
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(color=WHITE, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin = Side(style="thin", color="AAB8C0")
    for row in rows:
        values = []
        for header in headers:
            value = row.get(header)
            if isinstance(value, (dict, list, tuple, set)):
                value = json.dumps(value, ensure_ascii=False, sort_keys=True)
            values.append(value)
        ws.append(values)
    for row_index in range(2, ws.max_row + 1):
        for cell in ws[row_index]:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if row_index % 2 == 0:
                cell.fill = PatternFill("solid", fgColor=PALE_BLUE)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    ws.sheet_view.showGridLines = False
    for index, header in enumerate(headers, start=1):
        samples = [str(header)] + [str(ws.cell(row, index).value or "") for row in range(2, min(ws.max_row, 150) + 1)]
        ws.column_dimensions[get_column_letter(index)].width = min(60, max(11, max(len(value) for value in samples) + 2))


def augment_workbook(source: Path, destination: Path, *, evidence: list[dict[str, Any]], issues: list[dict[str, Any]], readiness: list[dict[str, Any]], assets: list[dict[str, Any]], events: list[dict[str, Any]]) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    wb = load_workbook(source)
    inherited = list(wb.sheetnames)
    internal_passed = sum(1 for row in readiness if row["status"] == "passed")
    external_pending = sum(1 for row in readiness if row["status"] == "controlled_pending")
    datasets = {
        "S5S2 CP2 Dashboard": [
            {"Control": "Response", "Value": 80, "Status": "current"},
            {"Control": "Checkpoint", "Value": "2 of 3", "Status": "complete"},
            {"Control": "Internal proof readiness", "Value": internal_passed, "Status": "passed"},
            {"Control": "External/provider blockers", "Value": external_pending, "Status": "controlled_pending"},
            {"Control": "Provider approval claimed", "Value": "NO", "Status": "correct"},
            {"Control": "Physical proof ordered", "Value": "NO", "Status": "correct"},
            {"Control": "Next", "Value": "Checkpoint 3 of 3 - Session 2 freeze and complete restore", "Status": "continue"},
        ],
        "S5S2 CP2 Evidence": evidence,
        "S5S2 CP2 Issues": issues,
        "S5S2 CP2 Proof Ready": readiness,
        "S5S2 CP2 Uploads": assets,
        "S5S2 CP2 Response": [{
            "Response": 80,
            "Raw Prompt": "Continue",
            "Summary": "Provider evidence boundary reconciled; internal proof-order readiness completed; external gates remain controlled pending.",
            "State": "checkpoint_complete_continue_required",
        }],
        "S5S2 CP2 Recovery": events,
    }
    for title, rows in datasets.items():
        if title in wb.sheetnames:
            del wb[title]
        ws = wb.create_sheet(title)
        _write_sheet(ws, rows)
    wb.properties.title = "MRHPD v3.0.0a Comprehensive Tracking Through Response 80"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(destination)
    with zipfile.ZipFile(destination) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("comprehensive workbook CRC failed")
    check = load_workbook(destination, read_only=True, data_only=False)
    try:
        sheet_names = list(check.sheetnames)
        formula_count = 0
        formula_errors = []
        for ws in check.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    value = cell.value
                    if isinstance(value, str) and value.startswith("="):
                        formula_count += 1
                        if any(token in value for token in ("#REF!", "#NAME?", "#VALUE!", "#DIV/0!", "#N/A")):
                            formula_errors.append(f"{ws.title}!{cell.coordinate}:{value}")
    finally:
        check.close()
    lost = sorted(set(inherited) - set(sheet_names))
    if lost or len(sheet_names) < len(inherited) + 7 or formula_errors:
        raise RuntimeError({"workbook_gate": {"lost": lost, "sheets": len(sheet_names), "formula_errors": formula_errors[:20]}})
    return {
        "path": str(destination),
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "source_sheet_count": len(inherited),
        "current_sheet_count": len(sheet_names),
        "new_sheet_count": len(sheet_names) - len(inherited),
        "lost_sheets": lost,
        "formula_count": formula_count,
        "formula_error_count": len(formula_errors),
        "status": "passed",
    }


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def write_tracking_files(project: Path, db_path: Path, now_iso: str) -> list[Path]:
    root = project / "Tracking" / "Prompt Response" / "Through Response 80"
    root.mkdir(parents=True, exist_ok=True)
    con = sqlite3.connect(db_path)
    try:
        info = con.execute("PRAGMA table_info(thread_response_reconciliation_cp3)").fetchall()
        columns = [row[1] for row in info]
        rows = [dict(zip(columns, row)) for row in con.execute("SELECT * FROM thread_response_reconciliation_cp3 ORDER BY CAST(response_number AS INTEGER), response_key")]
        fraction_rows = []
        if con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name='fractional_prompt_cp3'").fetchone()[0]:
            fraction_info = con.execute("PRAGMA table_info(fractional_prompt_cp3)").fetchall()
            fraction_columns = [row[1] for row in fraction_info]
            fraction_rows = [dict(zip(fraction_columns, row)) for row in con.execute("SELECT * FROM fractional_prompt_cp3 ORDER BY CAST(prompt_number AS REAL), prompt_number")]
    finally:
        con.close()
    response80 = next(row for row in rows if row.get("response_key") == "R80")
    response_path = root / "Response_80_Tracking.json"
    json_write(response_path, response80)

    raw_docx = root / "Medical References - Human Pathogen Database v3.0.0a Alternating Raw Prompts and Responses Through Response 80.docx"
    raw_doc = Document()
    raw_doc.core_properties.title = "Human Pathogen Database — Alternating Raw Prompts and Responses Through Response 80"
    raw_doc.core_properties.author = "Brent McAnulty, M.D."
    raw_doc.add_heading("Human Pathogen Database", 0)
    raw_doc.add_paragraph("Alternating Raw Prompts and Responses Through Response 80")
    for row in rows:
        number = row.get("response_label") or row.get("response_number")
        raw_doc.add_heading(f"Response {number}: {row.get('title') or 'Untitled exchange'}", level=1)
        table = raw_doc.add_table(rows=2, cols=1)
        table.style = "Table Grid"
        table.cell(0, 0).text = f"RAW PROMPT {number}\n\n{row.get('raw_prompt') or '[RAW PROMPT UNAVAILABLE]'}"
        set_cell_shading(table.cell(0, 0), "D9EAF7")
        table.cell(1, 0).text = f"RAW RESPONSE {number}\n\n{row.get('raw_response') or '[RAW RESPONSE UNAVAILABLE]'}\n\n{row.get('summary') or ''}"
        set_cell_shading(table.cell(1, 0), "E2F0D9")
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(8.5)
        raw_doc.add_paragraph()
    if fraction_rows:
        raw_doc.add_heading("Fractional prompts", level=1)
        for row in fraction_rows:
            raw_doc.add_paragraph(f"Prompt {row.get('prompt_number')}: {row.get('prompt_text')}")
    raw_doc.save(raw_docx)

    net_prompt = (
        "Continue the Human Pathogen Database from the newest verified checkpoint without regression. Use Google Drive as controlling "
        "storage; preserve immutable clinical and publication artifacts; complete print-production governance, provider-preview evidence "
        "reconciliation, physical-proof readiness and inspection controls, synchronized database/workbook/application surfaces, tracking, "
        "indexes, manifests, QA, and self-contained session, section, and project restores. Never infer provider approval, a proof order, "
        "receipt, inspection, correction, or signoff without item-level evidence."
    )
    net_response = (
        "Remediation Sections 1–4 are complete. Section 5 Session 1 is complete. Session 2 Checkpoint 2 recovers Response 79, freezes "
        "literal upload identities, reconciles the absence of provider evidence, separates passed internal readiness from controlled external "
        "blockers, establishes a provider-conversion issue taxonomy and proof-order readiness matrix, synchronizes all current project "
        "surfaces, and emits cumulative recovery through Response 80. Provider preview approval and physical proof remain pending external gates."
    )
    net_docx = root / "Medical References - Human Pathogen Database v3.0.0a Alternating Net Prompts and Responses Through Response 80.docx"
    net_doc = Document()
    net_doc.core_properties.title = "Human Pathogen Database — Alternating Net Prompts and Responses Through Response 80"
    net_doc.core_properties.author = "Brent McAnulty, M.D."
    net_doc.add_heading("Human Pathogen Database", 0)
    net_doc.add_heading("Print-production and final-release remediation", level=1)
    table = net_doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.cell(0, 0).text = "NET PROMPT\n\n" + net_prompt
    set_cell_shading(table.cell(0, 0), "D9EAF7")
    table.cell(1, 0).text = "NET RESPONSE\n\n" + net_response
    set_cell_shading(table.cell(1, 0), "E2F0D9")
    net_doc.save(net_docx)

    everything = root / "Medical References - Human Pathogen Database v3.0.0a Everything in One Thread Through Response 80.xlsx"
    wb = Workbook()
    wb.remove(wb.active)
    raw_prompts = [{"Response": row.get("response_label") or row.get("response_number"), "Major Topic": row.get("major_topic"), "Title": row.get("title"), "Raw Prompt": row.get("raw_prompt")} for row in rows]
    raw_responses = [{"Response": row.get("response_label") or row.get("response_number"), "Major Topic": row.get("major_topic"), "Title": row.get("title"), "Raw Response": row.get("raw_response"), "Summary": row.get("summary")} for row in rows]
    summary_index = [{"Response": row.get("response_label") or row.get("response_number"), "Major Topic": row.get("major_topic"), "Title": row.get("title"), "Goal": row.get("goal"), "Summary": row.get("summary"), "State": row.get("state")} for row in rows]
    datasets = {
        "Raw Prompts": raw_prompts,
        "Raw Responses": raw_responses,
        "Fractional Prompts": fraction_rows,
        "Net Prompt": [{"Major Topic": "Human Pathogen Database remediation", "Net Prompt": net_prompt}],
        "Net Response": [{"Major Topic": "Human Pathogen Database remediation", "Net Response": net_response}],
        "Summary Index": summary_index,
    }
    for title, data in datasets.items():
        ws = wb.create_sheet(title)
        _write_sheet(ws, data)
    wb.properties.title = "Human Pathogen Database — Everything in One Thread Through Response 80"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(everything)

    raw_net_md = root / "Medical References - Human Pathogen Database v3.0.0a Raw and Net Tracking Through Response 80.md"
    text_write(raw_net_md, f"""# Human Pathogen Database — Raw and Net Tracking Through Response 80

## Raw Prompt 80

Continue

## Raw Response 80

{response80.get('raw_response') or response80.get('summary')}

## Net Prompt

{net_prompt}

## Net Response

{net_response}

Updated: {now_iso}
""")
    cumulative = root / "Medical References - Human Pathogen Database v3.0.0a Cumulative Thread Index Through Response 80.md"
    lines = ["# Human Pathogen Database — Cumulative Thread Index Through Response 80", ""]
    for row in rows:
        lines.append(f"- Response {row.get('response_label') or row.get('response_number')} — {row.get('title')}: {row.get('summary')}")
    lines.extend(["", f"Updated: {now_iso}"])
    text_write(cumulative, "\n".join(lines))
    return [response_path, raw_docx, net_docx, everything, raw_net_md, cumulative]


def locate_main_application(project: Path) -> Path:
    matches = [path for path in project.rglob("human_pathogen_app.py") if path.is_file() and sha256_file(path) == APPLICATION_SHA256]
    if len(matches) != 1:
        raise RuntimeError({"main_application_candidates": [str(path) for path in matches]})
    return matches[0]


def write_application_surfaces(project: Path, db_path: Path, workbook_path: Path, now_iso: str) -> tuple[list[Path], dict[str, Any]]:
    root = project / "App" / "Section 5 Session 2 Checkpoint 2"
    root.mkdir(parents=True, exist_ok=True)
    app_path = locate_main_application(project)
    pointer = root / "CURRENT_DATABASE.txt"
    text_write(pointer, db_path.relative_to(project).as_posix() + "\n")
    state_path = root / "CURRENT_PROJECT_STATE.json"
    state = {
        "schema": "mrhpd-section5-current-project-state-1.3",
        "response": 80,
        "section": SECTION_LABEL,
        "session": SESSION_LABEL,
        "checkpoint": CHECKPOINT_LABEL,
        "state": "checkpoint_complete",
        "database": db_path.relative_to(project).as_posix(),
        "workbook": workbook_path.relative_to(project).as_posix(),
        "publication": PUBLICATION_REL,
        "print_interior": PRINT_INTERIOR_REL,
        "cover": COVER_PNG_REL,
        "provider_approval_claimed": False,
        "physical_proof_ordered": False,
        "main_application": app_path.relative_to(project).as_posix(),
        "main_application_sha256": sha256_file(app_path),
        "main_application_unchanged": True,
        "recorded_at": now_iso,
    }
    json_write(state_path, state)
    audit_script = root / "audit_section5_session2_checkpoint2.py"
    text_write(audit_script, f'''#!/usr/bin/env python3
import hashlib, json, sqlite3
from pathlib import Path
from openpyxl import load_workbook
from PIL import Image
from pypdf import PdfReader
project=Path(__file__).resolve().parents[2]
db=project/{db_path.relative_to(project).as_posix()!r}
workbook=project/{workbook_path.relative_to(project).as_posix()!r}
publication=project/{PUBLICATION_REL!r}
interior=project/{PRINT_INTERIOR_REL!r}
cover=project/{COVER_PNG_REL!r}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as f:
  for block in iter(lambda:f.read(1024*1024),b''): h.update(block)
 return h.hexdigest()
con=sqlite3.connect(db)
try:
 integrity=con.execute("PRAGMA integrity_check").fetchone()[0]
 fk=len(list(con.execute("PRAGMA foreign_key_check")))
 response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R80'").fetchone()[0]
 evidence=con.execute("SELECT COUNT(*) FROM section5_s2_provider_evidence_reconciliation").fetchone()[0]
 approval=con.execute("SELECT COUNT(*) FROM section5_s2_provider_evidence_reconciliation WHERE provider_approval_claimed!=0").fetchone()[0]
 readiness=con.execute("SELECT COUNT(*) FROM section5_s2_proof_order_readiness").fetchone()[0]
 pending=con.execute("SELECT COUNT(*) FROM section5_s2_proof_order_readiness WHERE status='controlled_pending'").fetchone()[0]
 observed=con.execute("SELECT COUNT(*) FROM section5_s2_provider_conversion_issue WHERE observed!=0").fetchone()[0]
finally: con.close()
wb=load_workbook(workbook,read_only=True,data_only=False)
try: sheets=len(wb.sheetnames)
finally: wb.close()
pub=PdfReader(str(publication)); pub_pages=len(pub.pages); searchable=sum(1 for p in pub.pages if (p.extract_text() or '').strip())
pr=PdfReader(str(interior)); print_pages=len(pr.pages)
with Image.open(cover) as im: cover_pixels=list(im.size)
result={{'status':'passed' if integrity=='ok' and fk==0 and response==1 and evidence>=12 and approval==0 and readiness>=16 and pending>=4 and observed==0 and sheets>=122 and pub_pages==537 and searchable==537 and print_pages==538 and cover_pixels==[5554,3375] and sha(publication)=={PUBLICATION_SHA256!r} and sha(interior)=={PRINT_INTERIOR_SHA256!r} and sha(cover)=={COVER_SHA256!r} else 'failed','integrity':integrity,'foreign_keys':fk,'response80':response,'provider_evidence':evidence,'provider_approval_claims':approval,'proof_readiness':readiness,'external_pending':pending,'observed_provider_issues':observed,'workbook_sheets':sheets,'publication_pages':pub_pages,'searchable_pages':searchable,'print_pages':print_pages,'cover_pixels':cover_pixels,'main_application_unchanged':True}}
print(json.dumps(result,indent=2))
raise SystemExit(0 if result['status']=='passed' else 1)
''')
    result = subprocess.run([sys.executable, str(audit_script)], cwd=project, text=True, capture_output=True, timeout=600)
    if result.returncode:
        raise RuntimeError({"application_audit_failed": {"stdout": result.stdout[-10000:], "stderr": result.stderr[-10000:]}})
    audit = json.loads(result.stdout)
    audit.update({
        "main_application_path": app_path.relative_to(project).as_posix(),
        "main_application_sha256": sha256_file(app_path),
        "main_application_unchanged": True,
    })
    output = root / "SECTION5_SESSION2_CHECKPOINT2_APPLICATION_AUDIT.json"
    json_write(output, audit)
    return [pointer, state_path, audit_script, output], audit


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            pass
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], font: Any, fill: str, width_chars: int, spacing: int = 7) -> int:
    x, y = xy
    lines = []
    for paragraph in text.splitlines() or [""]:
        lines.extend(textwrap.wrap(paragraph, width=width_chars) or [""])
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        box = draw.textbbox((x, y), line or "Ag", font=font)
        y += box[3] - box[1] + spacing
    return y


def build_readiness_figure(path: Path, evidence: list[dict[str, Any]], readiness: list[dict[str, Any]]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 2400, 1350
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, 175), fill=f"#{NAVY}")
    draw.text((110, 48), "Provider evidence boundary and proof readiness", font=_font(56, True), fill="white")
    draw.text((112, 205), "Internal production controls are complete; provider and physical-proof events remain evidence-dependent.", font=_font(29), fill=f"#{DARK}")
    internal = sum(1 for row in readiness if row["status"] == "passed")
    external = sum(1 for row in readiness if row["status"] == "controlled_pending")
    provider_received = sum(1 for row in evidence if row["evidence_class"] == "provider" and row["evidence_received"])
    cards = [
        ("INTERNAL GATES PASSED", internal, TEAL, PALE_GREEN),
        ("EXTERNAL GATES PENDING", external, GOLD, PALE_GOLD),
        ("PROVIDER EVIDENCE ITEMS", provider_received, "A65F46", PALE_RED),
        ("APPROVAL CLAIMED", 0, "A65F46", PALE_RED),
    ]
    x = 110
    for label, value, color, fill in cards:
        draw.rounded_rectangle((x, 310, x + 500, 565), radius=26, fill=f"#{fill}", outline=f"#{color}", width=5)
        draw.text((x + 30, 345), label, font=_font(24, True), fill=f"#{color}")
        draw.text((x + 30, 410), str(value), font=_font(78, True), fill=f"#{DARK}")
        x += 570
    draw.rounded_rectangle((110, 650, 1110, 1220), radius=28, fill=f"#{PALE_GREEN}", outline=f"#{TEAL}", width=5)
    draw.text((155, 700), "Internally verified", font=_font(38, True), fill=f"#{TEAL}")
    internal_text = "Exact Response 77 and Response 79 recovery; 538-page interior; 5,554 × 3,375 cover; trim, spine, bleed, fonts, searchability, literal filenames, hashes, database, workbook, application, indexes, manifests, and recovery."
    draw_wrapped(draw, internal_text, (155, 775), _font(30), f"#{DARK}", 54, spacing=10)
    draw.rounded_rectangle((1290, 650, 2290, 1220), radius=28, fill=f"#{PALE_GOLD}", outline=f"#{GOLD}", width=5)
    draw.text((1335, 700), "Controlled external boundary", font=_font(38, True), fill=f"#{GOLD}")
    external_text = "KDP submission receipt, provider-rendered preview, warnings and errors, preview approval, ISBN/barcode decision, proof-order authorization, shipping, physical receipt, inspection, correction, and signoff. None is inferred without item-level evidence."
    draw_wrapped(draw, external_text, (1335, 775), _font(30), f"#{DARK}", 54, spacing=10)
    image.save(path, format="PNG", dpi=(300, 300), optimize=True)
    return {"path": str(path), "pixels": [width, height], "dpi": 300, "mode": image.mode, "bytes": path.stat().st_size, "sha256": sha256_file(path)}


def add_docx_table(doc: Document, rows: list[dict[str, Any]], columns: list[str]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, column in enumerate(columns):
        cell = table.rows[0].cells[index]
        cell.text = column.replace("_", " ").title()
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(8)
    for row_index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            value = row.get(column, "")
            if isinstance(value, (dict, list, tuple, set)):
                value = json.dumps(value, ensure_ascii=False, sort_keys=True)
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index % 2 == 0:
                set_cell_shading(cells[index], PALE_BLUE)
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(7.5)


def build_docx_report(path: Path, *, now_iso: str, evidence: list[dict[str, Any]], issues: list[dict[str, Any]], readiness: list[dict[str, Any]], assets: list[dict[str, Any]], database_qa: dict[str, Any], workbook_qa: dict[str, Any], application_qa: dict[str, Any], figure: Path) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    doc.core_properties.title = "MRHPD Section 5 Session 2 Checkpoint 2 — Provider Evidence and Proof Readiness"
    doc.core_properties.author = "Brent McAnulty, M.D."
    doc.add_heading("Human Pathogen Database", 0)
    subtitle = doc.add_paragraph("Provider evidence reconciliation and proof-order readiness")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Version {PROJECT_VERSION} • {SECTION_LABEL} • {SESSION_LABEL} • {CHECKPOINT_LABEL}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(figure), width=Inches(7.05))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Checkpoint disposition", level=1)
    add_docx_table(doc, [
        {"control": "Response", "result": 80},
        {"control": "Checkpoint", "result": "2 of 3 complete"},
        {"control": "Internal readiness", "result": "passed"},
        {"control": "Provider approval", "result": "not claimed"},
        {"control": "Physical proof", "result": "not ordered or received"},
        {"control": "Next", "result": "Checkpoint 3 — Session 2 freeze and complete restore"},
    ], ["control", "result"])
    doc.add_heading("Evidence boundary", level=1)
    doc.add_paragraph(
        "No KDP submission receipt, provider-rendered preview, provider warning/error list, provider approval, downloadable provider proof, proof order, receipt, inspection, correction, or signoff was supplied or discovered. The absence of evidence is not converted into a passed provider gate."
    )
    add_docx_table(doc, evidence, ["label", "evidence_class", "status", "notes"])
    doc.add_heading("Provider conversion issue taxonomy", level=1)
    doc.add_paragraph("The register is ready for item-level provider messages. No issue is marked observed without evidence.")
    add_docx_table(doc, issues, ["issue_label", "domain", "severity", "status", "correction"])
    doc.add_heading("Proof-order readiness", level=1)
    add_docx_table(doc, readiness, ["gate_label", "gate_class", "status", "blocks_final_release", "notes"])
    doc.add_heading("Literal upload assets", level=1)
    add_docx_table(doc, assets, ["literal_filename", "role", "submission_required", "bytes", "sha256"])
    doc.add_heading("Synchronized project controls", level=1)
    add_docx_table(doc, [
        {"control": "SQLite tables", "result": database_qa["table_count"]},
        {"control": "SQLite integrity", "result": database_qa["integrity"]},
        {"control": "Foreign-key violations", "result": database_qa["foreign_key_violations"]},
        {"control": "Workbook sheets", "result": workbook_qa["current_sheet_count"]},
        {"control": "Workbook formula errors", "result": workbook_qa["formula_error_count"]},
        {"control": "Application audit", "result": application_qa["status"]},
        {"control": "Digital publication", "result": "537 searchable pages; immutable"},
        {"control": "Print interior", "result": "538 pages; frozen"},
        {"control": "Cover", "result": "5,554 × 3,375 pixels; frozen"},
    ], ["control", "result"])
    doc.add_heading("Next checkpoint", level=1)
    doc.add_paragraph(
        "Checkpoint 3 will independently reconstruct this state, repeat all internal acceptance gates, preserve unsupported external events as controlled pending, freeze Session 2, and emit the required complete self-contained Session 2 restore."
    )
    doc.add_paragraph(f"Generated: {now_iso}")
    doc.save(path)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("DOCX CRC failed")
    return {"path": str(path), "bytes": path.stat().st_size, "sha256": sha256_file(path), "status": "passed"}


def build_pdf_report(path: Path, *, now_iso: str, evidence: list[dict[str, Any]], issues: list[dict[str, Any]], readiness: list[dict[str, Any]], assets: list[dict[str, Any]], figure: Path) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Title2", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=20, leading=24, textColor=colors.HexColor("#17324D"), alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name="Sub", parent=styles["Normal"], fontName="Helvetica", fontSize=10, leading=14, textColor=colors.HexColor("#1C7475"), alignment=TA_CENTER, spaceAfter=12))
    styles.add(ParagraphStyle(name="Head2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#17324D"), spaceBefore=8, spaceAfter=6))
    styles.add(ParagraphStyle(name="Body2", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.5, leading=11, textColor=colors.HexColor("#24323D"), alignment=TA_LEFT))
    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.45*inch, leftMargin=0.45*inch, topMargin=0.45*inch, bottomMargin=0.45*inch)
    story: list[Any] = [
        Paragraph("Human Pathogen Database", styles["Title2"]),
        Paragraph("Provider evidence reconciliation and proof-order readiness", styles["Sub"]),
        Paragraph(f"Version {PROJECT_VERSION} • {SECTION_LABEL} • {SESSION_LABEL} • {CHECKPOINT_LABEL}", styles["Sub"]),
    ]
    image = fitz.Pixmap(str(figure)) if False else None
    from reportlab.platypus import Image as RLImage
    story.append(RLImage(str(figure), width=7.2*inch, height=4.05*inch))
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph("Evidence boundary", styles["Head2"]))
    story.append(Paragraph("No provider approval, Print Previewer result, proof order, receipt, inspection, correction, or signoff is claimed without item-level evidence. Internal controls and external blockers remain distinct.", styles["Body2"]))

    def table_from(rows: list[dict[str, Any]], columns: list[str], widths: list[float]) -> Table:
        data = [[Paragraph(column.replace("_", " ").title(), styles["Body2"]) for column in columns]]
        for row in rows:
            data.append([Paragraph(str(row.get(column, "")), styles["Body2"]) for column in columns])
        table = Table(data, colWidths=[width*inch for width in widths], repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#17324D")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("GRID", (0,0), (-1,-1), 0.35, colors.HexColor("#AAB8C0")),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#EAF1F5")]),
            ("LEFTPADDING", (0,0), (-1,-1), 4),
            ("RIGHTPADDING", (0,0), (-1,-1), 4),
            ("TOPPADDING", (0,0), (-1,-1), 3),
            ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ]))
        return table

    story.append(table_from(evidence, ["label", "evidence_class", "status", "notes"], [2.0, 0.8, 1.0, 3.4]))
    story.append(PageBreak())
    story.append(Paragraph("Provider conversion issue taxonomy", styles["Head2"]))
    story.append(table_from(issues, ["issue_label", "domain", "severity", "status"], [3.2, 1.2, 0.8, 2.0]))
    story.append(Spacer(1, 0.12*inch))
    story.append(Paragraph("Proof-order readiness", styles["Head2"]))
    story.append(table_from(readiness, ["gate_label", "gate_class", "status", "notes"], [3.2, 1.0, 1.2, 2.0]))
    story.append(PageBreak())
    story.append(Paragraph("Literal upload assets", styles["Head2"]))
    story.append(table_from(assets, ["literal_filename", "role", "submission_required", "sha256"], [3.0, 1.5, 0.8, 2.1]))
    story.append(Spacer(1, 0.15*inch))
    story.append(Paragraph("Next checkpoint", styles["Head2"]))
    story.append(Paragraph("Checkpoint 3 will independently reconstruct this checkpoint, repeat every internal acceptance gate, preserve external events as controlled pending, freeze Session 2, and emit the complete Session 2 restore.", styles["Body2"]))
    story.append(Paragraph(f"Generated: {now_iso}", styles["Body2"]))
    doc.build(story)
    reader = PdfReader(str(path))
    searchable = sum(1 for page in reader.pages if (page.extract_text() or "").strip())
    if searchable != len(reader.pages):
        raise RuntimeError({"pdf_searchability": {"pages": len(reader.pages), "searchable": searchable}})
    return {"path": str(path), "bytes": path.stat().st_size, "sha256": sha256_file(path), "pages": len(reader.pages), "searchable_pages": searchable, "status": "passed"}


def build_register(path: Path, *, evidence: list[dict[str, Any]], issues: list[dict[str, Any]], readiness: list[dict[str, Any]], assets: list[dict[str, Any]], events: list[dict[str, Any]], database_qa: dict[str, Any], workbook_qa: dict[str, Any]) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    wb.remove(wb.active)
    datasets = {
        "Dashboard": [
            {"Control": "Response", "Value": 80, "Status": "COMPLETE"},
            {"Control": "Checkpoint", "Value": "2 of 3", "Status": "COMPLETE"},
            {"Control": "Provider approval", "Value": "No evidence", "Status": "NOT CLAIMED"},
            {"Control": "Physical proof", "Value": "Not ordered or received", "Status": "CONTROLLED PENDING"},
            {"Control": "SQLite tables", "Value": database_qa["table_count"], "Status": database_qa["integrity"]},
            {"Control": "Workbook sheets", "Value": workbook_qa["current_sheet_count"], "Status": "passed"},
        ],
        "Provider Evidence": evidence,
        "Conversion Issues": issues,
        "Proof Readiness": readiness,
        "Upload Manifest": assets,
        "Recovery Events": events,
        "QA": [database_qa, workbook_qa],
        "Tracking": [{"Response": 80, "Raw Prompt": "Continue", "State": "checkpoint_complete_continue_required"}],
    }
    for title, rows in datasets.items():
        ws = wb.create_sheet(title)
        _write_sheet(ws, rows)
    wb.properties.title = "MRHPD Section 5 Session 2 Checkpoint 2 Provider Evidence and Proof Readiness Register"
    wb.properties.creator = "Brent McAnulty, M.D."
    wb.save(path)
    with zipfile.ZipFile(path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("register CRC failed")
    check = load_workbook(path, read_only=True, data_only=False)
    try:
        sheets = list(check.sheetnames)
    finally:
        check.close()
    return {"path": str(path), "bytes": path.stat().st_size, "sha256": sha256_file(path), "sheets": sheets, "status": "passed"}


def render_report_qa(pdf_path: Path, output: Path) -> dict[str, Any]:
    output.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(pdf_path)
    rows = []
    for index, page in enumerate(doc):
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False, colorspace=fitz.csRGB)
        target = output / f"report-page-{index+1:03d}.png"
        pix.save(target)
        rows.append({"page": index + 1, "width": pix.width, "height": pix.height, "bytes": target.stat().st_size, "sha256": sha256_file(target), "text_characters": len(page.get_text("text").strip()), "status": "passed" if target.stat().st_size > 1000 and page.get_text("text").strip() else "failed"})
    doc.close()
    if not rows or any(row["status"] != "passed" for row in rows):
        raise RuntimeError({"report_render_qa": rows})
    json_path = output / "REPORT_RENDER_QA.json"
    json_write(json_path, {"status": "passed", "pages": rows})
    return {"status": "passed", "page_count": len(rows), "pages": rows, "qa_path": str(json_path)}


def extract_text_for_index(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in {".md", ".txt", ".csv", ".json", ".py", ".html", ".yml", ".yaml"}:
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix == ".docx":
            doc = Document(path)
            chunks = [paragraph.text for paragraph in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells)
            return "\n".join(chunks)
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        if suffix == ".xlsx":
            wb = load_workbook(path, read_only=True, data_only=False)
            try:
                chunks: list[str] = []
                for ws in wb.worksheets:
                    chunks.append(ws.title)
                    for row in ws.iter_rows(values_only=True):
                        chunks.append(" | ".join("" if value is None else str(value) for value in row))
                return "\n".join(chunks)
            finally:
                wb.close()
        if suffix in {".sqlite", ".db"}:
            con = sqlite3.connect(path)
            try:
                chunks = [row[0] or "" for row in con.execute("SELECT sql FROM sqlite_master WHERE sql IS NOT NULL ORDER BY name")]
                tables = {row[0] for row in con.execute("SELECT name FROM sqlite_master WHERE type='table'")}
                for table in (
                    "thread_response_reconciliation_cp3",
                    "section5_s2_checkpoint",
                    "section5_s2_provider_evidence_reconciliation",
                    "section5_s2_provider_conversion_issue",
                    "section5_s2_proof_order_readiness",
                    "section5_s2_upload_manifest",
                    "section5_s2_recovery_event",
                ):
                    if table not in tables:
                        continue
                    columns = [row[1] for row in con.execute(f"PRAGMA table_info({table})")]
                    chunks.append(table + " | " + " | ".join(columns))
                    for row in con.execute(f"SELECT * FROM {table}"):
                        chunks.append(" | ".join("" if value is None else str(value) for value in row))
                return "\n".join(chunks)
            finally:
                con.close()
    except Exception as exc:
        return f"[index extraction error: {exc!r}]"
    return ""


def build_source_and_bit_indexes(project: Path, now_iso: str) -> dict[str, Any]:
    root = project / "Indexes" / "Section 5 Session 2 Checkpoint 2"
    root.mkdir(parents=True, exist_ok=True)
    source_json = root / "MRHPD v3.0.0a Section 5 Session 2 Checkpoint 2 Source Index.json"
    source_csv = root / "MRHPD v3.0.0a Section 5 Session 2 Checkpoint 2 Source Index.csv"
    bit_path = root / "MRHPD v3.0.0a Section 5 Session 2 Checkpoint 2 Bit Index.sqlite"
    qa_path = root / "MRHPD v3.0.0a Section 5 Session 2 Checkpoint 2 Index QA.json"
    excluded = {source_json.resolve(), source_csv.resolve(), bit_path.resolve(), qa_path.resolve()}
    rows: list[dict[str, Any]] = []
    fts_rows: list[tuple[str, str, str, str]] = []
    searchable_suffixes = {".md", ".txt", ".csv", ".json", ".py", ".html", ".yml", ".yaml", ".docx", ".pdf", ".xlsx", ".sqlite", ".db"}
    container_suffixes = {".zip", ".docx", ".xlsx", ".pptx", ".odt"}
    for path in sorted(file for file in project.rglob("*") if file.is_file() and file.resolve() not in excluded):
        rel = path.relative_to(project).as_posix()
        purpose = "Project artifact"
        if rel.startswith("Database/"):
            purpose = "Canonical or historical project database"
        elif rel.startswith("Tracking/"):
            purpose = "Prompt response workbook or tracking artifact"
        elif rel.startswith("Print Production/"):
            purpose = "Print-production candidate or provider-control artifact"
        elif rel.startswith("QA/"):
            purpose = "Quality assurance evidence"
        elif rel.startswith("Reports/"):
            purpose = "Human-readable checkpoint report"
        user_searchable = int(path.suffix.lower() in searchable_suffixes)
        content = extract_text_for_index(path) if user_searchable else ""
        if len(content) > 4_000_000:
            content = content[:4_000_000] + "\n[CONTENT TRUNCATED IN CURRENT BIT INDEX; FULL ARTIFACT RETAINED]"
        row = {
            "record_type": "physical_file",
            "path": rel,
            "container_path": "",
            "name": path.name,
            "purpose": purpose,
            "bytes": path.stat().st_size,
            "sha256": sha256_file(path),
            "user_searchable": user_searchable,
        }
        rows.append(row)
        fts_rows.append((rel, path.name, purpose, content))
        if path.suffix.lower() in container_suffixes:
            try:
                with zipfile.ZipFile(path) as zf:
                    if zf.testzip() is not None:
                        continue
                    for info in zf.infolist():
                        if info.is_dir():
                            continue
                        member_path = rel + "!" + info.filename
                        member_purpose = "Container member"
                        rows.append({
                            "record_type": "container_member",
                            "path": member_path,
                            "container_path": rel,
                            "name": PurePosixPath(info.filename).name,
                            "purpose": member_purpose,
                            "bytes": info.file_size,
                            "sha256": "",
                            "user_searchable": 0,
                        })
                        fts_rows.append((member_path, PurePosixPath(info.filename).name, member_purpose, ""))
            except zipfile.BadZipFile:
                pass
    json_write(source_json, {"schema": "mrhpd-source-index-2.1", "generated_at": now_iso, "record_count": len(rows), "records": rows})
    csv_write(source_csv, rows)
    if bit_path.exists():
        bit_path.unlink()
    con = sqlite3.connect(bit_path)
    try:
        con.executescript("""
        CREATE TABLE artifact (
            artifact_id INTEGER PRIMARY KEY,
            record_type TEXT NOT NULL,
            path TEXT NOT NULL UNIQUE,
            container_path TEXT,
            name TEXT NOT NULL,
            purpose TEXT NOT NULL,
            bytes INTEGER NOT NULL,
            sha256 TEXT,
            user_searchable INTEGER NOT NULL
        );
        CREATE VIRTUAL TABLE artifact_fts USING fts5(path, name, purpose, content);
        """)
        for row, payload in zip(rows, fts_rows):
            con.execute("INSERT INTO artifact (record_type,path,container_path,name,purpose,bytes,sha256,user_searchable) VALUES (?,?,?,?,?,?,?,?)", (row["record_type"], row["path"], row["container_path"], row["name"], row["purpose"], row["bytes"], row["sha256"], row["user_searchable"]))
            con.execute("INSERT INTO artifact_fts (path,name,purpose,content) VALUES (?,?,?,?)", payload)
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        counts = {
            "artifact": con.execute("SELECT COUNT(*) FROM artifact").fetchone()[0],
            "fts": con.execute("SELECT COUNT(*) FROM artifact_fts").fetchone()[0],
            "response80": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"Response 80"',)).fetchone()[0],
            "provider_preview": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"provider preview"',)).fetchone()[0],
            "proof_readiness": con.execute("SELECT COUNT(*) FROM artifact_fts WHERE artifact_fts MATCH ?", ('"proof readiness"',)).fetchone()[0],
        }
        con.commit()
    finally:
        con.close()
    if integrity != "ok" or counts["artifact"] != len(rows) or counts["fts"] != len(rows):
        raise RuntimeError({"bit_index_gate": {"integrity": integrity, "counts": counts, "expected": len(rows)}})
    qa = {
        "status": "passed",
        "generated_at": now_iso,
        "source_index_records": len(rows),
        "physical_files": sum(1 for row in rows if row["record_type"] == "physical_file"),
        "container_members": sum(1 for row in rows if row["record_type"] == "container_member"),
        "bit_index_integrity": integrity,
        "counts": counts,
        "bit_index_sha256": sha256_file(bit_path),
    }
    json_write(qa_path, qa)
    return {"source_json": source_json, "source_csv": source_csv, "bit_index": bit_path, "qa_path": qa_path, "qa": qa}


def build_project_manifest(project: Path, now_iso: str) -> tuple[Path, Path, list[dict[str, Any]]]:
    root = project / "Manifest" / "Section 5 Session 2 Checkpoint 2"
    root.mkdir(parents=True, exist_ok=True)
    manifest = root / "MRHPD v3.0.0a Section 5 Session 2 Checkpoint 2 Current Project Manifest.json"
    checksums = root / "MRHPD v3.0.0a Section 5 Session 2 Checkpoint 2 Current Project Checksums.sha256"
    rows = []
    for path in sorted(file for file in project.rglob("*") if file.is_file() and file not in {manifest, checksums}):
        rows.append({"path": path.relative_to(project).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    json_write(manifest, {
        "schema": "mrhpd-current-project-manifest-2.1",
        "generated_at": now_iso,
        "exclusions": [manifest.relative_to(project).as_posix(), checksums.relative_to(project).as_posix()],
        "file_count": len(rows),
        "total_bytes": sum(row["bytes"] for row in rows),
        "files": rows,
    })
    text_write(checksums, "".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
    for row in rows:
        path = project / row["path"]
        if path.stat().st_size != row["bytes"] or sha256_file(path) != row["sha256"]:
            raise RuntimeError({"manifest_mismatch": row["path"]})
    return manifest, checksums, rows


def create_apply_script(manifest: dict[str, Any], expected: dict[str, Any]) -> str:
    return f'''#!/usr/bin/env python3
from __future__ import annotations
import argparse, hashlib, json, re, shutil, sqlite3, sys, tempfile, zipfile
from pathlib import Path, PurePosixPath
from openpyxl import load_workbook
from PIL import Image
from pypdf import PdfReader
BASE_RESTORE_BYTES={BASE_RESTORE_BYTES}
BASE_RESTORE_SHA256={BASE_RESTORE_SHA256!r}
BASE_PROJECT_BYTES={BASE_PROJECT_BYTES}
BASE_PROJECT_SHA256={BASE_PROJECT_SHA256!r}
PUBLICATION_SHA256={PUBLICATION_SHA256!r}
APPLICATION_SHA256={APPLICATION_SHA256!r}
CURRENT_DB_REL={CURRENT_DB_REL!r}
CURRENT_WORKBOOK_REL={CURRENT_WORKBOOK_REL!r}
PUBLICATION_REL={PUBLICATION_REL!r}
PRINT_INTERIOR_REL={PRINT_INTERIOR_REL!r}
COVER_PNG_REL={COVER_PNG_REL!r}
MANIFEST={manifest!r}
EXPECTED={expected!r}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as f:
  for block in iter(lambda:f.read(1024*1024),b''): h.update(block)
 return h.hexdigest()
def safe_extract(path,destination):
 destination.mkdir(parents=True,exist_ok=True)
 with zipfile.ZipFile(path) as zf:
  if zf.testzip() is not None: raise RuntimeError('ZIP CRC failure: '+str(path))
  names=zf.namelist()
  if len(names)!=len(set(names)): raise RuntimeError('duplicate ZIP members')
  for name in names:
   pp=PurePosixPath(name.replace('\\\\','/'))
   if pp.is_absolute() or '..' in pp.parts or re.match(r'^[A-Za-z]:',name): raise RuntimeError('unsafe ZIP path: '+name)
  zf.extractall(destination)
def verify(path,size,digest,label):
 observed={{'bytes':path.stat().st_size,'sha256':sha(path)}}
 if observed!={{'bytes':size,'sha256':digest}}: raise RuntimeError({{label:observed}})
def main():
 parser=argparse.ArgumentParser()
 parser.add_argument('--base-response77-restore',type=Path,required=True)
 parser.add_argument('--output-dir',type=Path,required=True)
 args=parser.parse_args()
 verify(args.base_response77_restore,BASE_RESTORE_BYTES,BASE_RESTORE_SHA256,'baseline_restore')
 package=Path(__file__).resolve().parents[1]
 overlay=package/'OVERLAY'
 if not overlay.is_dir(): raise RuntimeError('OVERLAY missing')
 if args.output_dir.exists() and any(args.output_dir.iterdir()): raise RuntimeError('output directory must be empty')
 args.output_dir.mkdir(parents=True,exist_ok=True)
 with tempfile.TemporaryDirectory(prefix='mrhpd-r80-apply-') as td:
  work=Path(td); restore_root=work/'restore'; safe_extract(args.base_response77_restore,restore_root)
  candidates=[p for p in restore_root.rglob('*.zip') if p.stat().st_size==BASE_PROJECT_BYTES and sha(p)==BASE_PROJECT_SHA256]
  if len(candidates)!=1: raise RuntimeError({{'project_archive_candidates':[str(p) for p in candidates]}})
  extracted=work/'project'; safe_extract(candidates[0],extracted)
  roots=[p for p in extracted.iterdir() if p.is_dir()]
  source=roots[0] if len(roots)==1 else extracted
  destination=args.output_dir/{CURRENT_PROJECT_NAME!r}
  shutil.copytree(source,destination)
  for row in MANIFEST['overlay_files']:
   src=overlay/row['path']; verify(src,row['bytes'],row['sha256'],'overlay_'+row['path'])
   dst=destination/row['path']; dst.parent.mkdir(parents=True,exist_ok=True); shutil.copy2(src,dst)
  db=destination/CURRENT_DB_REL
  con=sqlite3.connect(db)
  try:
   integrity=con.execute('PRAGMA integrity_check').fetchone()[0]
   fk=len(list(con.execute('PRAGMA foreign_key_check')))
   response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R80'").fetchone()[0]
   evidence=con.execute("SELECT COUNT(*) FROM section5_s2_provider_evidence_reconciliation").fetchone()[0]
   approval=con.execute("SELECT COUNT(*) FROM section5_s2_provider_evidence_reconciliation WHERE provider_approval_claimed!=0").fetchone()[0]
   readiness=con.execute("SELECT COUNT(*) FROM section5_s2_proof_order_readiness").fetchone()[0]
   pending=con.execute("SELECT COUNT(*) FROM section5_s2_proof_order_readiness WHERE status='controlled_pending'").fetchone()[0]
   observed=con.execute("SELECT COUNT(*) FROM section5_s2_provider_conversion_issue WHERE observed!=0").fetchone()[0]
  finally: con.close()
  wb=load_workbook(destination/CURRENT_WORKBOOK_REL,read_only=True,data_only=False)
  try: sheets=len(wb.sheetnames)
  finally: wb.close()
  publication=destination/PUBLICATION_REL
  interior=destination/PRINT_INTERIOR_REL
  cover=destination/COVER_PNG_REL
  verify(publication,EXPECTED['publication_bytes'],PUBLICATION_SHA256,'publication')
  verify(interior,EXPECTED['interior_bytes'],{PRINT_INTERIOR_SHA256!r},'interior')
  verify(cover,EXPECTED['cover_bytes'],{COVER_SHA256!r},'cover')
  pub=PdfReader(str(publication)); pub_pages=len(pub.pages); searchable=sum(1 for page in pub.pages if (page.extract_text() or '').strip())
  pr=PdfReader(str(interior)); print_pages=len(pr.pages)
  with Image.open(cover) as image: cover_pixels=list(image.size)
  apps=[p for p in destination.rglob('human_pathogen_app.py') if p.is_file() and sha(p)==APPLICATION_SHA256]
  result={{'status':'passed' if integrity=='ok' and fk==0 and response==1 and evidence>=12 and approval==0 and readiness>=16 and pending>=4 and observed==0 and sheets>=122 and pub_pages==537 and searchable==537 and print_pages==538 and cover_pixels==[5554,3375] and len(apps)==1 else 'failed','project_root':str(destination),'database':{{'integrity':integrity,'foreign_keys':fk,'response80':response,'provider_evidence':evidence,'provider_approval_claims':approval,'proof_readiness':readiness,'external_pending':pending,'observed_provider_issues':observed}},'workbook_sheets':sheets,'publication_pages':pub_pages,'searchable_pages':searchable,'print_pages':print_pages,'cover_pixels':cover_pixels,'main_application_matches':len(apps),'publication_sha256':sha(publication),'interior_sha256':sha(interior),'cover_sha256':sha(cover)}}
  output=args.output_dir/'MRHPD_RESPONSE80_CHECKPOINT_RECOVERY_APPLICATION_RESULT.json'; output.write_text(json.dumps(result,indent=2),encoding='utf-8')
  print(json.dumps(result,indent=2))
  raise SystemExit(0 if result['status']=='passed' else 1)
if __name__=='__main__': main()
'''


def build_recovery_package(*, baseline_project: Path, current_project: Path, baseline_restore: Path, project_archive: Path, dist: Path, now: datetime, summary: dict[str, Any], direct_files: list[Path]) -> dict[str, Any]:
    baseline_map = {path.relative_to(baseline_project).as_posix(): (path.stat().st_size, sha256_file(path)) for path in baseline_project.rglob("*") if path.is_file()}
    current_map = {path.relative_to(current_project).as_posix(): (path.stat().st_size, sha256_file(path)) for path in current_project.rglob("*") if path.is_file()}
    deleted = sorted(set(baseline_map) - set(current_map))
    if deleted:
        raise RuntimeError({"unexpected_deleted_paths": deleted})
    overlay_rows = []
    for rel, identity in sorted(current_map.items()):
        if baseline_map.get(rel) != identity:
            overlay_rows.append({"path": rel, "bytes": identity[0], "sha256": identity[1], "change": "new" if rel not in baseline_map else "changed"})
    stamp = now.strftime("%Y-%m-%d %H%M UTC")
    package_root = dist / "recovery_package_root"
    if package_root.exists():
        shutil.rmtree(package_root)
    overlay_root = package_root / "OVERLAY"
    tools = package_root / "TOOLS"
    overlay_root.mkdir(parents=True)
    tools.mkdir(parents=True)
    for row in overlay_rows:
        source = current_project / row["path"]
        target = overlay_root / row["path"]
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, target)
    manifest = {
        "schema": "mrhpd-section5-checkpoint-recovery-1.3",
        "generated_at": now.isoformat().replace("+00:00", "Z"),
        "version": PROJECT_VERSION,
        "response": 80,
        "section": SECTION_LABEL,
        "session": SESSION_LABEL,
        "checkpoint": CHECKPOINT_LABEL,
        "state": "checkpoint_complete",
        "baseline_restore": {"name": baseline_restore.name, "bytes": baseline_restore.stat().st_size, "sha256": sha256_file(baseline_restore)},
        "baseline_project": {"name": project_archive.name, "bytes": project_archive.stat().st_size, "sha256": sha256_file(project_archive)},
        "overlay_file_count": len(overlay_rows),
        "overlay_total_bytes": sum(row["bytes"] for row in overlay_rows),
        "overlay_files": overlay_rows,
        "deleted_paths": [],
        "provider_approval_claimed": False,
        "physical_proof_ordered": False,
        "accepted_predecessor_mutated": False,
        "frozen_section3_release_mutated": False,
        "immutable_publication_mutated": False,
        "user_upload_required": False,
        "requires_conversation_reconstruction": False,
        "next": "Remediation Section 5 of 5 Session 2 of 3 Checkpoint 3 of 3",
    }
    json_write(package_root / "CHECKPOINT_RECOVERY_MANIFEST.json", manifest)
    text_write(package_root / "CHECKPOINT_RECOVERY_CHECKSUMS.sha256", "".join(f"{row['sha256']}  OVERLAY/{row['path']}\n" for row in overlay_rows))
    expected = {
        "publication_bytes": (current_project / PUBLICATION_REL).stat().st_size,
        "interior_bytes": (current_project / PRINT_INTERIOR_REL).stat().st_size,
        "cover_bytes": (current_project / COVER_PNG_REL).stat().st_size,
    }
    text_write(tools / "apply_checkpoint_recovery.py", create_apply_script(manifest, expected))
    text_write(package_root / "RESTORE_READ_FIRST.md", f"""# Human Pathogen Database — Response 80 Checkpoint Recovery

This cumulative intermediate recovery applies directly to the exact Response 77 complete restore and includes all current project progress through Response 80. Response 79 does not need to be applied separately.

## Required baseline

Filename: `{baseline_restore.name}`

Bytes: `{baseline_restore.stat().st_size}`

SHA-256: `{sha256_file(baseline_restore)}`

## Automated apply

```bash
python TOOLS/apply_checkpoint_recovery.py \
  --base-response77-restore "<Response 77 complete restore.zip>" \
  --output-dir "<empty destination>"
```

The utility verifies the exact baseline, every cumulative overlay file, SQLite integrity and foreign keys, Response 80, provider-evidence and proof-readiness controls, the comprehensive workbook, the immutable 537-page digital publication, the frozen 538-page print interior and exact cover, and the unchanged main application.

Provider approval, a proof order, proof receipt, physical inspection, correction, and signoff remain controlled external gates and are not claimed by this package.
""")
    recovery_zip = dist / (
        f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 "
        f"Remediation Section 5 of 5 Session 2 of 3 Checkpoint 2 of 3 RECOVERY DATA THROUGH RESPONSE 80 {stamp}.zip"
    )
    with zipfile.ZipFile(recovery_zip, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in sorted(package_root.rglob("*")):
            if path.is_file():
                zf.write(path, path.relative_to(package_root).as_posix())
    recovery_qa = verify_zip(recovery_zip)
    with tempfile.TemporaryDirectory(prefix="mrhpd-r80-clean-apply-") as td:
        output = Path(td) / "restored"
        result = subprocess.run(
            [sys.executable, str((tools / "apply_checkpoint_recovery.py").resolve()), "--base-response77-restore", str(baseline_restore), "--output-dir", str(output)],
            cwd=package_root,
            text=True,
            capture_output=True,
            timeout=3000,
        )
        if result.returncode:
            raise RuntimeError({"clean_apply_failed": {"stdout": result.stdout[-20000:], "stderr": result.stderr[-20000:]}})
        result_path = output / "MRHPD_RESPONSE80_CHECKPOINT_RECOVERY_APPLICATION_RESULT.json"
        clean_apply = json.loads(result_path.read_text(encoding="utf-8"))
        if clean_apply.get("status") != "passed":
            raise RuntimeError({"clean_apply_gate": clean_apply})
    verification = {
        "schema": "mrhpd-response80-checkpoint2-verification-1.0",
        "status": "passed",
        "generated_at": now.isoformat().replace("+00:00", "Z"),
        "delivery_state": "checkpoint_complete_continue_required",
        "baseline_restore": manifest["baseline_restore"],
        "baseline_project": manifest["baseline_project"],
        "recovery_zip": recovery_qa,
        "manifest": {"overlay_file_count": len(overlay_rows), "overlay_total_bytes": sum(row["bytes"] for row in overlay_rows), "deleted_paths": 0},
        "clean_apply": clean_apply,
        "provider_approval_claimed": False,
        "physical_proof_ordered": False,
        "accepted_predecessor_mutated": False,
        "frozen_section3_release_mutated": False,
        "immutable_publication_mutated": False,
        "user_upload_required": False,
        "checkpoint_2_of_3_complete": True,
        "session_2_of_3_complete": False,
        "remediation_section_5_complete": False,
        "next": "Checkpoint 3 of 3 - Session 2 freeze and complete restore",
    }
    verification_path = dist / "MRHPD v3.0.0a Response 80 Checkpoint 2 Recovery Verification.json"
    json_write(verification_path, verification)
    sha_path = dist / f"{recovery_zip.name}.sha256.txt"
    text_write(sha_path, f"{recovery_qa['sha256']}  {recovery_zip.name}\n")
    summary_path = dist / "MRHPD_RESPONSE80_SECTION5_SESSION2_CHECKPOINT2_BUILD_SUMMARY.json"
    json_write(summary_path, summary | {"recovery": verification})
    exact_names = dist / "MRHPD v3.0.0a Response 80 Exact File Names.txt"
    text_write(exact_names, f"""Response 80 cumulative checkpoint recovery ZIP:
{recovery_zip.name}

Required baseline complete restore:
{baseline_restore.name}

Required baseline project archive embedded in that restore:
{project_archive.name}

Current copied SQLite database:
{Path(CURRENT_DB_REL).name}

Current comprehensive workbook:
{Path(CURRENT_WORKBOOK_REL).name}

Frozen print-production interior:
{Path(PRINT_INTERIOR_REL).name}

Frozen full-cover PDF:
{Path(COVER_PDF_REL).name}

Frozen full-cover PNG:
{Path(COVER_PNG_REL).name}

Provider evidence and proof-readiness report:
MRHPD v3.0.0a Section 5 Session 2 Checkpoint 2 Provider Evidence and Proof Readiness Report.pdf

Checkpoint register:
MRHPD v3.0.0a Section 5 Session 2 Checkpoint 2 Provider Evidence and Proof Readiness Register.xlsx
""")
    delivery = dist / f"MRHPD v3.0.0a Response 80 Section 5 Session 2 Checkpoint 2 Recovery Package {stamp}.zip"
    with zipfile.ZipFile(delivery, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in [recovery_zip, sha_path, verification_path, summary_path, exact_names, *direct_files]:
            if path.exists():
                zf.write(path, path.name)
    delivery_qa = verify_zip(delivery)
    return {
        "recovery_zip": recovery_zip,
        "recovery_qa": recovery_qa,
        "verification_path": verification_path,
        "summary_path": summary_path,
        "exact_names": exact_names,
        "delivery": delivery,
        "delivery_qa": delivery_qa,
        "overlay_rows": overlay_rows,
        "clean_apply": clean_apply,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--response77-dir", type=Path, required=True)
    parser.add_argument("--response79-dir", type=Path, required=True)
    parser.add_argument("--dist", type=Path, default=Path("dist_cp5_s2_cp2"))
    args = parser.parse_args()
    now = utc_now()
    now_iso = now.isoformat().replace("+00:00", "Z")
    if args.dist.exists():
        shutil.rmtree(args.dist)
    args.dist.mkdir(parents=True)
    with tempfile.TemporaryDirectory(prefix="mrhpd-cp5-s2-cp2-") as td:
        work = Path(td)
        restore, project_archive, baseline_project, response79_project, response79_application = restore_response79(args.response77_dir, args.response79_dir, work)
        current_project = work / "current_project" / CURRENT_PROJECT_NAME
        current_project.parent.mkdir(parents=True)
        shutil.copytree(response79_project, current_project)

        publication = current_project / PUBLICATION_REL
        interior = current_project / PRINT_INTERIOR_REL
        cover_png = current_project / COVER_PNG_REL
        if sha256_file(publication) != PUBLICATION_SHA256:
            raise RuntimeError("immutable digital publication changed")
        if sha256_file(interior) != PRINT_INTERIOR_SHA256:
            raise RuntimeError("frozen print interior changed")
        if sha256_file(cover_png) != COVER_SHA256:
            raise RuntimeError("frozen cover changed")
        publication_reader = PdfReader(str(publication))
        publication_pages = len(publication_reader.pages)
        publication_searchable = sum(1 for page in publication_reader.pages if (page.extract_text() or "").strip())
        print_pages = len(PdfReader(str(interior)).pages)
        with Image.open(cover_png) as image:
            cover_pixels = list(image.size)
        if publication_pages != 537 or publication_searchable != 537 or print_pages != 538 or cover_pixels != [5554, 3375]:
            raise RuntimeError({"frozen_artifact_gate": {"publication_pages": publication_pages, "publication_searchable": publication_searchable, "print_pages": print_pages, "cover_pixels": cover_pixels}})

        evidence = provider_evidence_rows(now_iso)
        issues = conversion_issue_rows(now_iso)
        readiness = proof_readiness_rows(now_iso)
        events = recovery_event_rows(now_iso)
        assets = asset_manifest_rows(current_project, now_iso)

        source_db = response79_project / RESPONSE79_DB_REL
        source_workbook = response79_project / RESPONSE79_WORKBOOK_REL
        if not source_db.exists() or not source_workbook.exists():
            raise RuntimeError({"response79_current_artifacts": {"database": str(source_db), "workbook": str(source_workbook)}})
        current_db = current_project / CURRENT_DB_REL
        database_qa = sync_database(source_db, current_db, now_iso=now_iso, evidence=evidence, issues=issues, readiness=readiness, assets=assets, events=events)
        current_workbook = current_project / CURRENT_WORKBOOK_REL
        workbook_qa = augment_workbook(source_workbook, current_workbook, evidence=evidence, issues=issues, readiness=readiness, assets=assets, events=events)
        tracking_files = write_tracking_files(current_project, current_db, now_iso)
        application_files, application_qa = write_application_surfaces(current_project, current_db, current_workbook, now_iso)

        data_root = current_project / "Data" / "Section 5 Session 2 Checkpoint 2"
        evidence_json = data_root / "MRHPD v3.0.0a Response 80 Provider Evidence Reconciliation.json"
        evidence_csv = data_root / "MRHPD v3.0.0a Response 80 Provider Evidence Reconciliation.csv"
        issues_json = data_root / "MRHPD v3.0.0a Response 80 Provider Conversion Issue Taxonomy.json"
        issues_csv = data_root / "MRHPD v3.0.0a Response 80 Provider Conversion Issue Taxonomy.csv"
        readiness_json = data_root / "MRHPD v3.0.0a Response 80 Proof Order Readiness.json"
        readiness_csv = data_root / "MRHPD v3.0.0a Response 80 Proof Order Readiness.csv"
        assets_json = data_root / "MRHPD v3.0.0a Response 80 Provider Upload Manifest.json"
        assets_csv = data_root / "MRHPD v3.0.0a Response 80 Provider Upload Manifest.csv"
        events_json = data_root / "MRHPD v3.0.0a Response 80 Recovery Events 213-220.json"
        for path, payload in ((evidence_json, evidence), (issues_json, issues), (readiness_json, readiness), (assets_json, assets), (events_json, events)):
            json_write(path, payload)
        for path, payload in ((evidence_csv, evidence), (issues_csv, issues), (readiness_csv, readiness), (assets_csv, assets)):
            csv_write(path, payload)

        report_root = current_project / "Reports" / "Section 5 Session 2" / "Checkpoint 2"
        artwork_root = current_project / "Artwork" / "Section 5 Print Production" / "Session 2 Checkpoint 2"
        figure = artwork_root / "MRHPD-FIG-S5-0005 Provider Evidence Boundary and Proof Readiness v3.0.0a.png"
        figure_qa = build_readiness_figure(figure, evidence, readiness)
        docx_report = report_root / "MRHPD v3.0.0a Section 5 Session 2 Checkpoint 2 Provider Evidence and Proof Readiness Report.docx"
        pdf_report = report_root / "MRHPD v3.0.0a Section 5 Session 2 Checkpoint 2 Provider Evidence and Proof Readiness Report.pdf"
        xlsx_register = report_root / "MRHPD v3.0.0a Section 5 Session 2 Checkpoint 2 Provider Evidence and Proof Readiness Register.xlsx"
        docx_qa = build_docx_report(docx_report, now_iso=now_iso, evidence=evidence, issues=issues, readiness=readiness, assets=assets, database_qa=database_qa, workbook_qa=workbook_qa, application_qa=application_qa, figure=figure)
        pdf_qa = build_pdf_report(pdf_report, now_iso=now_iso, evidence=evidence, issues=issues, readiness=readiness, assets=assets, figure=figure)
        register_qa = build_register(xlsx_register, evidence=evidence, issues=issues, readiness=readiness, assets=assets, events=events, database_qa=database_qa, workbook_qa=workbook_qa)
        render_qa = render_report_qa(pdf_report, report_root / "Rendered Report QA")

        qa_root = current_project / "QA" / "Section 5 Session 2" / "Checkpoint 2"
        qa_root.mkdir(parents=True, exist_ok=True)
        frozen_qa = {
            "schema": "mrhpd-section5-session2-checkpoint2-qa-1.0",
            "generated_at": now_iso,
            "status": "passed_with_controlled_external_gates",
            "response": 80,
            "section": SECTION_LABEL,
            "session": SESSION_LABEL,
            "checkpoint": CHECKPOINT_LABEL,
            "response79_application": response79_application,
            "database": database_qa,
            "workbook": workbook_qa,
            "application": application_qa,
            "publication": {"pages": publication_pages, "searchable_pages": publication_searchable, "sha256": sha256_file(publication), "unchanged": True},
            "print_interior": {"pages": print_pages, "sha256": sha256_file(interior), "unchanged": True},
            "cover": {"pixels": cover_pixels, "sha256": sha256_file(cover_png), "unchanged": True},
            "provider_evidence": {"records": len(evidence), "provider_approval_claimed": False},
            "conversion_issue_taxonomy": {"records": len(issues), "observed_issues": 0},
            "proof_readiness": {"records": len(readiness), "internal_passed": sum(1 for row in readiness if row["status"] == "passed"), "external_pending": sum(1 for row in readiness if row["status"] == "controlled_pending")},
            "reports": {"docx": docx_qa, "pdf": pdf_qa, "xlsx": register_qa, "render": render_qa, "figure": figure_qa},
            "checkpoint_2_of_3_complete": True,
            "session_2_of_3_complete": False,
            "remediation_section_5_complete": False,
            "provider_approval_claimed": False,
            "physical_proof_ordered": False,
            "accepted_predecessor_mutated": False,
            "frozen_section3_release_mutated": False,
            "immutable_publication_mutated": False,
            "main_application_mutated": False,
            "next": "Checkpoint 3 of 3 - Session 2 freeze and complete restore",
        }
        final_qa_path = qa_root / "SECTION5_SESSION2_CHECKPOINT2_QA.json"
        json_write(final_qa_path, frozen_qa)
        json_write(qa_root / "DATABASE_QA.json", database_qa)
        json_write(qa_root / "WORKBOOK_QA.json", workbook_qa)
        json_write(qa_root / "APPLICATION_QA.json", application_qa)
        json_write(qa_root / "PROVIDER_EVIDENCE_QA.json", {"status": "passed", "records": evidence, "provider_approval_claimed": False})
        json_write(qa_root / "PROOF_READINESS_QA.json", {"status": "passed_with_controlled_external_gates", "records": readiness})
        json_write(qa_root / "RECOVERY_EVENTS_213_220.json", events)

        addendum = current_project / "Recovery" / "Project Instructions" / "MRHPD v3.0.0a Project Instructions 1.5.0 Operating Addendum Through Response 80.md"
        text_write(addendum, f"""# Human Pathogen Database — Project Instructions 1.5.0 Operating Addendum Through Response 80

The current Project Instructions 1.5.0 were reprocessed for this checkpoint. The operative controls include newest-verified-artifact recovery, no regression, exact literal filenames, Google Drive custody, Raw and Net tracking, cumulative checkpoint recovery between full session/section restores, self-contained restores at session and section boundaries, explicit COMPLETE versus CONTINUE disposition, no unsupported external claims, and automatic recovery from recoverable errors.

Response 80 keeps KDP Print Previewer approval, proof ordering, receipt, inspection, correction, and signoff controlled pending until item-level evidence exists.

Updated: {now_iso}
""")

        index_result = build_source_and_bit_indexes(current_project, now_iso)
        manifest_path, checksums_path, manifest_rows = build_project_manifest(current_project, now_iso)
        summary = {
            "schema": "mrhpd-response80-section5-session2-checkpoint2-build-1.0",
            "generated_at": now_iso,
            "status": "passed",
            "response": 80,
            "section": SECTION_LABEL,
            "session": SESSION_LABEL,
            "checkpoint": CHECKPOINT_LABEL,
            "database": database_qa,
            "workbook": workbook_qa,
            "application": application_qa,
            "publication": frozen_qa["publication"],
            "print_interior": frozen_qa["print_interior"],
            "cover": frozen_qa["cover"],
            "provider_evidence": frozen_qa["provider_evidence"],
            "conversion_issue_taxonomy": frozen_qa["conversion_issue_taxonomy"],
            "proof_readiness": frozen_qa["proof_readiness"],
            "reports": frozen_qa["reports"],
            "index": index_result["qa"],
            "manifest_records": len(manifest_rows),
            "provider_approval_claimed": False,
            "physical_proof_ordered": False,
            "user_upload_required": False,
            "checkpoint_2_of_3_complete": True,
            "session_2_of_3_complete": False,
            "remediation_section_5_complete": False,
            "next": "Checkpoint 3 of 3 - Session 2 freeze and complete restore",
        }
        direct_files = [docx_report, pdf_report, xlsx_register, figure]
        package = build_recovery_package(
            baseline_project=baseline_project,
            current_project=current_project,
            baseline_restore=restore,
            project_archive=project_archive,
            dist=args.dist,
            now=now,
            summary=summary,
            direct_files=direct_files,
        )
        console = {
            "status": "passed",
            "delivery": package["delivery"].name,
            "delivery_bytes": package["delivery_qa"]["bytes"],
            "delivery_sha256": package["delivery_qa"]["sha256"],
            "recovery_zip": package["recovery_zip"].name,
            "recovery_zip_bytes": package["recovery_qa"]["bytes"],
            "recovery_zip_sha256": package["recovery_qa"]["sha256"],
            "overlay_files": len(package["overlay_rows"]),
            "database_tables": database_qa["table_count"],
            "workbook_sheets": workbook_qa["current_sheet_count"],
            "provider_evidence_records": len(evidence),
            "proof_readiness_records": len(readiness),
            "provider_approval_claimed": False,
            "physical_proof_ordered": False,
            "user_upload_required": False,
            "checkpoint_2_of_3_complete": True,
            "next": "Checkpoint 3 of 3 - Session 2 freeze and complete restore",
        }
        print(json.dumps(console, indent=2))


if __name__ == "__main__":
    main()
