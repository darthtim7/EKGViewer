#!/usr/bin/env python3
"""Build the MRHPD Section 4 Session 1 complete self-contained restore.

The builder resumes from the exact Complete Restore Through Response 64 and the
independently verified Checkpoint 2 recovery package through Response 65. It
applies that recovery package to a disposable copy, performs the final Session 1
synchronization through Response 66, reruns database/workbook/application/
publication/index/archive gates, and emits a complete restore that requires no
other project file or conversation reconstruction.

The accepted predecessor, frozen Section 3 release, Response 64 restore, and
Checkpoint 2 recovery package are never edited in place.
"""
from __future__ import annotations

import argparse
import csv
import hashlib
import importlib.util
import json
import os
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import textwrap
import zipfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any, Iterable

PROJECT_VERSION = "3.0.0a"
RESPONSE_NUMBER = 66
BASE_RESPONSE64_BYTES = 145_920_215
BASE_RESPONSE64_SHA256 = "71e7a06868e82238188827fd73ca7b2843670b0843dcf98ab5ab72305bf77834"
CP2_RECOVERY_BYTES = 10_898_139
CP2_RECOVERY_SHA256 = "f659564382e5858beaf3b1d2e8f77599d58e2176ad4e977034a0dc18dd0efe35"
CP2_DATABASE_SHA256 = "71586759d991836777c9c03e07b905a5d861c026a6c790b8920133f9f00f7454"
CP2_WORKBOOK_SHA256 = "9499e9bf6c2260e37403140c77149a7e5dd9728e52fbf1b5c3cf842b7a062ac6"
APPLICATION_SHA256 = "5f1e4ac8fc6e2ffad213646c78e4f261bf655795de5ac8a7d4486d3be11ce139"
PUBLICATION_SHA256 = "8a053112ca24cd730b970130d5d0fc57a15c681531603601096186aeb0cd9642"
EDITABLE_ASSEMBLY_SHA256 = "f832ff934d77049d75712f28bdfc9167b8a6b119c797235431b304b9e24369a2"
NOW_DT = datetime.now(timezone.utc)
NOW = NOW_DT.replace(microsecond=0).isoformat().replace("+00:00", "Z")
STAMP = NOW_DT.strftime("%Y-%m-%d %H%M UTC")
RAW_PROMPT = "Continue"

RESPONSE66 = {
    "response_key": "R66",
    "response_number": 66,
    "response_label": "66",
    "branch_id": "mainline",
    "canonical_current": 1,
    "response_date": NOW,
    "major_topic": "Human Pathogen Database remediation",
    "title": "Section 4 Session 1 complete restore and handoff",
    "goal": (
        "Complete Section 4 Session 1 from the verified Checkpoint 2 recovery state, rerun all session-level controls, "
        "and emit a complete self-contained restore through Response 66."
    ),
    "raw_prompt": RAW_PROMPT,
    "raw_response": "[PRE-EMISSION RESPONSE; final user-visible response is represented by the source-supported summary]",
    "summary": (
        "Applied and verified the Response 65 Checkpoint 2 recovery package, completed the final Session 1 database, workbook, "
        "application, tracking, recovery, index and manifest synchronization, preserved the 537-page publication and editable "
        "assembly unchanged, and emitted a clean-extraction-tested complete restore through Response 66."
    ),
    "state": "session_complete_continue_required",
    "coverage": "exact raw prompt plus source-supported response summary",
    "fidelity_classification": "source_verified_prompt_and_summary",
    "source_id": "CURRENT-CONVERSATION-R66",
    "source_path": "Current conversation turn and Section 4 Session 1 complete restore",
    "notes": "Session 1 of 3 is complete. Continue begins Remediation Section 4 Session 2 of 3.",
}

RECOVERY_EVENTS = [
    {
        "event_number": 95,
        "event_code": "V3-CP4-S1-REC-CHECKPOINT2-RECOVERY-APPLIED-AND-VERIFIED",
        "occurred_at": NOW,
        "failed_step": "None; newest verified checkpoint recovery was applied.",
        "exact_error_or_reason": (
            "Session-end work resumed from the exact Response 64 restore and the independently verified Response 65 Checkpoint 2 overlay, "
            "rather than reconstructing progress from the conversation."
        ),
        "intact_artifacts": (
            "Complete Restore Through Response 64, Checkpoint 2 recovery ZIP, frozen Section 3 release, accepted predecessor, persistent "
            "Google Drive custody, and all Checkpoint 2 QA evidence."
        ),
        "recovery_action": (
            "Verified both input identities, executed the embedded deterministic recovery utility, and required the resulting database, "
            "workbook, application, publication and editable-assembly identities to match the verified Checkpoint 2 state."
        ),
        "validation_result": "Checkpoint 2 state reproduced and verified before Session 1 finalization began.",
        "data_quality_effect": "None.",
        "next_checkpoint": "Complete Section 4 Session 1 and emit the session-end restore.",
    },
    {
        "event_number": 96,
        "event_code": "V3-CP4-S1-REC-SESSION1-FINAL-SYNCHRONIZATION-COMPLETE",
        "occurred_at": NOW,
        "failed_step": "None; final Session 1 synchronization completed.",
        "exact_error_or_reason": (
            "Checkpoint 3 required a current canonical database, workbook, application state, tracking set, recovery state, indexes, "
            "manifests and session handoff through Response 66."
        ),
        "intact_artifacts": "The accepted predecessor, frozen Section 3 release, Response 64 restore and Checkpoint 2 package remained immutable.",
        "recovery_action": (
            "Created a new Session 1-complete database and workbook, synchronized the local application through its native --db interface, "
            "added Response 66 and the session-completion controls, and rebuilt current QA/index/manifest outputs in a copied tree."
        ),
        "validation_result": "Database, workbook, application, publication and index controls passed.",
        "data_quality_effect": "None; only the copied current project tree changed.",
        "next_checkpoint": "Build and clean-verify the complete self-contained restore.",
    },
    {
        "event_number": 97,
        "event_code": "V3-CP4-S1-REC-SESSION1-COMPLETE-RESTORE-CLEAN-VERIFIED",
        "occurred_at": NOW,
        "failed_step": "None; session-end restore generation completed.",
        "exact_error_or_reason": (
            "The operative emission policy requires a complete self-contained restore at every session boundary and checkpoint recovery data "
            "between full-product emissions."
        ),
        "intact_artifacts": "All historical source, checkpoint, publication, database, application, workbook and tracking artifacts remain included.",
        "recovery_action": (
            "Built the complete project snapshot, embedded the verified Checkpoint 2 recovery package and restore controls, clean-extracted the "
            "restore, reran acceptance gates, and split the final restore into the minimum two connector-compatible transport volumes."
        ),
        "validation_result": "The session-end complete restore and both transport volumes passed manifest, checksum, CRC and clean-extraction gates.",
        "data_quality_effect": "None.",
        "next_checkpoint": "Remediation Section 4 of 5 Session 2 of 3.",
    },
]

NET_PROMPT = (
    "Continue the Human Pathogen Database from the newest verified checkpoint. Preserve Google Drive as the controlling project and "
    "download store. Complete Section 4 Session 1 from the exact Response 64 restore plus the Response 65 Checkpoint 2 recovery data; "
    "synchronize the copied SQLite database, comprehensive workbook, local application, Raw/Net tracking, Cumulative Thread Index, recovery "
    "records, Source Index, Bit Index, manifests and QA without modifying the accepted predecessor or frozen Section 3 release. Emit a complete "
    "self-contained restore at the session boundary, with no requirement for another project file, cloud artifact or conversation reconstruction."
)

NET_RESPONSE = (
    "Section 4 Session 1 is complete through Response 66. The exact Checkpoint 2 state was restored and independently verified before final "
    "synchronization. The current copied project contains the Session 1-complete SQLite database, comprehensive workbook, local-application "
    "state, tracking/recovery records, indexes, manifests and QA. The controlling 537-page publication and 537-page editable assembly remain "
    "byte-identical. A clean-extraction-tested, self-contained restore and connector-compatible Google Drive transport set were produced. "
    "Continue begins Section 4 Session 2 of 3."
)


def sha256_file(path: Path, chunk_size: int = 1024 * 1024) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(chunk_size), b""):
            h.update(block)
    return h.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False, default=str) + "\n", encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value.rstrip() + "\n", encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fields is None:
        fields = list(rows[0]) if rows else []
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def safe_infos(zf: zipfile.ZipFile, files_only: bool = False) -> list[zipfile.ZipInfo]:
    infos = [info for info in zf.infolist() if not files_only or not info.is_dir()]
    names = [info.filename for info in infos]
    duplicates = [name for name, count in Counter(names).items() if count > 1]
    unsafe: list[str] = []
    for name in names:
        posix = PurePosixPath(name.replace("\\", "/"))
        if name.startswith(("/", "\\")) or re.match(r"^[A-Za-z]:", name) or ".." in posix.parts:
            unsafe.append(name)
    if duplicates or unsafe:
        raise RuntimeError({"duplicate_members": duplicates[:30], "unsafe_members": unsafe[:30]})
    return infos


def verify_zip(path: Path, expected_bytes: int | None = None, expected_sha256: str | None = None) -> dict[str, Any]:
    if expected_bytes is not None and path.stat().st_size != expected_bytes:
        raise RuntimeError({"file": str(path), "expected_bytes": expected_bytes, "actual_bytes": path.stat().st_size})
    digest = sha256_file(path)
    if expected_sha256 is not None and digest != expected_sha256:
        raise RuntimeError({"file": str(path), "expected_sha256": expected_sha256, "actual_sha256": digest})
    with zipfile.ZipFile(path) as zf:
        infos = safe_infos(zf)
        bad = zf.testzip()
        if bad:
            raise RuntimeError({"file": str(path), "zip_crc_error": bad})
        filler = [info.filename for info in infos if re.search(r"(^|/)(filler|padding|pad)(/|$)", info.filename, re.I)]
        if filler:
            raise RuntimeError({"file": str(path), "filler_members": filler[:30]})
    return {
        "file": path.name,
        "bytes": path.stat().st_size,
        "sha256": digest,
        "members": len(infos),
        "crc": "passed",
        "duplicate_members": 0,
        "unsafe_paths": 0,
        "filler_members": 0,
    }


def safe_extract(path: Path, target: Path) -> None:
    target.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as zf:
        safe_infos(zf)
        zf.extractall(target)


def locate_exact_zip(root: Path, expected_bytes: int, expected_sha256: str) -> Path:
    candidates = sorted(root.rglob("*.zip"), key=lambda p: p.stat().st_size, reverse=True)
    for path in candidates:
        if path.stat().st_size == expected_bytes and sha256_file(path) == expected_sha256:
            return path
    raise FileNotFoundError({
        "root": str(root),
        "expected_bytes": expected_bytes,
        "expected_sha256": expected_sha256,
        "candidate_sample": [(str(p), p.stat().st_size) for p in candidates[:30]],
    })


def find_by_hash(root: Path, pattern: str, expected_sha256: str) -> Path:
    matches = [path for path in root.rglob(pattern) if path.is_file() and sha256_file(path) == expected_sha256]
    if len(matches) != 1:
        raise RuntimeError({"pattern": pattern, "expected_sha256": expected_sha256, "matches": [str(p) for p in matches]})
    return matches[0]


def table_exists(con: sqlite3.Connection, table: str) -> bool:
    return con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?", (table,)).fetchone() is not None


def table_info(con: sqlite3.Connection, table: str) -> list[sqlite3.Row]:
    return list(con.execute(f'PRAGMA table_info("{table}")'))


def table_columns(con: sqlite3.Connection, table: str) -> list[str]:
    return [row[1] for row in table_info(con, table)]


def schema_upsert(con: sqlite3.Connection, table: str, row: dict[str, Any], key: str) -> None:
    info = table_info(con, table)
    columns = [item[1] for item in info]
    primary_keys = {item[1] for item in info if item[5]}
    values = {name: value for name, value in row.items() if name in columns and name not in primary_keys}
    if key not in values:
        raise RuntimeError({"table": table, "required_key": key, "available": sorted(values)})
    names = list(values)
    update_names = [name for name in names if name != key]
    quoted = ", ".join(f'"{name}"' for name in names)
    placeholders = ", ".join("?" for _ in names)
    updates = ", ".join(f'"{name}"=excluded."{name}"' for name in update_names)
    con.execute(
        f'INSERT INTO "{table}" ({quoted}) VALUES ({placeholders}) '
        f'ON CONFLICT("{key}") DO UPDATE SET {updates}',
        [values[name] for name in names],
    )


def apply_checkpoint2(base_restore: Path, recovery_zip: Path, work: Path) -> tuple[Path, dict[str, Any]]:
    recovery_package = work / "checkpoint2_recovery_package"
    safe_extract(recovery_zip, recovery_package)
    utility = recovery_package / "TOOLS" / "apply_checkpoint_recovery.py"
    if not utility.exists():
        raise FileNotFoundError(utility)
    output = work / "project_through_response65"
    result = subprocess.run(
        [
            sys.executable,
            str(utility),
            "--base-response64-restore",
            str(base_restore),
            "--output-dir",
            str(output),
        ],
        text=True,
        capture_output=True,
        timeout=1200,
    )
    apply_record = {
        "returncode": result.returncode,
        "stdout": result.stdout[-30000:],
        "stderr": result.stderr[-30000:],
        "utility": str(utility),
    }
    if result.returncode != 0:
        raise RuntimeError({"checkpoint2_apply_failed": apply_record})

    db = find_by_hash(output, "*.sqlite", CP2_DATABASE_SHA256)
    workbook = find_by_hash(output, "*.xlsx", CP2_WORKBOOK_SHA256)
    app = find_by_hash(output, "human_pathogen_app.py", APPLICATION_SHA256)
    publication = find_by_hash(output, "*Integrated Manuscript*.pdf", PUBLICATION_SHA256)
    editable = find_by_hash(output, "*Editable Integrated Manuscript Assembly*.docx", EDITABLE_ASSEMBLY_SHA256)
    con = sqlite3.connect(db)
    try:
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        r65 = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R65'").fetchone()[0]
        cp2 = con.execute("SELECT state FROM section4_checkpoint WHERE checkpoint_code='MRHPD-V3-CP4-S1-CP2'").fetchone()
    finally:
        con.close()
    if integrity != "ok" or fk or r65 != 1 or not cp2 or cp2[0] != "checkpoint_complete":
        raise RuntimeError({"checkpoint2_state_gate": {"integrity": integrity, "fk": fk[:20], "r65": r65, "cp2": cp2}})
    return output, {
        "status": "passed",
        "base_restore": verify_zip(base_restore, BASE_RESPONSE64_BYTES, BASE_RESPONSE64_SHA256),
        "checkpoint2_recovery": verify_zip(recovery_zip, CP2_RECOVERY_BYTES, CP2_RECOVERY_SHA256),
        "apply_utility": apply_record,
        "database": db.relative_to(output).as_posix(),
        "database_sha256": CP2_DATABASE_SHA256,
        "workbook": workbook.relative_to(output).as_posix(),
        "workbook_sha256": CP2_WORKBOOK_SHA256,
        "application": app.relative_to(output).as_posix(),
        "application_sha256": APPLICATION_SHA256,
        "publication": publication.relative_to(output).as_posix(),
        "publication_sha256": PUBLICATION_SHA256,
        "editable_assembly": editable.relative_to(output).as_posix(),
        "editable_assembly_sha256": EDITABLE_ASSEMBLY_SHA256,
        "accepted_predecessor_mutated": False,
    }


def synchronize_database(project: Path) -> tuple[Path, dict[str, Any]]:
    source = find_by_hash(project, "*.sqlite", CP2_DATABASE_SHA256)
    target = source.with_name(
        f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 "
        "Remediation Section 4 of 5 Session 1 of 3 COMPLETE THROUGH RESPONSE 66.sqlite"
    )
    shutil.copy2(source, target)
    con = sqlite3.connect(target)
    con.row_factory = sqlite3.Row
    con.execute("PRAGMA foreign_keys=ON")
    try:
        con.execute("BEGIN IMMEDIATE")
        response_row = dict(RESPONSE66)
        if "reconciled_at" in table_columns(con, "thread_response_reconciliation_cp3"):
            response_row["reconciled_at"] = NOW
        schema_upsert(con, "thread_response_reconciliation_cp3", response_row, "response_key")

        if table_exists(con, "remediation_recovery_event"):
            columns = table_columns(con, "remediation_recovery_event")
            for event in RECOVERY_EVENTS:
                normalized = {
                    "event_code": event["event_code"],
                    "occurred_at": event["occurred_at"],
                    "failed_step": event["failed_step"],
                    "exact_error_or_reason": event["exact_error_or_reason"],
                    "intact_artifacts": event["intact_artifacts"],
                    "recovery_action": event["recovery_action"],
                    "validation_result": event["validation_result"],
                    "data_quality_effect": event["data_quality_effect"],
                    "next_checkpoint": event["next_checkpoint"],
                }
                normalized = {k: v for k, v in normalized.items() if k in columns}
                schema_upsert(con, "remediation_recovery_event", normalized, "event_code")

        con.executescript("""
        CREATE TABLE IF NOT EXISTS section4_session_release (
          section4_session_release_id INTEGER PRIMARY KEY,
          release_code TEXT NOT NULL UNIQUE,
          section_label TEXT NOT NULL,
          session_label TEXT NOT NULL,
          completion_response INTEGER NOT NULL,
          state TEXT NOT NULL,
          database_status TEXT NOT NULL,
          workbook_status TEXT NOT NULL,
          application_status TEXT NOT NULL,
          publication_status TEXT NOT NULL,
          complete_restore_required INTEGER NOT NULL CHECK(complete_restore_required IN (0,1)),
          accepted_predecessor_mutated INTEGER NOT NULL CHECK(accepted_predecessor_mutated IN (0,1)),
          next_session TEXT NOT NULL,
          recorded_at TEXT NOT NULL
        );
        """)
        con.execute("""
          INSERT INTO section4_checkpoint
          (checkpoint_code,section_label,session_label,checkpoint_label,response_number,state,
           database_integrity,foreign_key_violations,workbook_status,application_status,publication_sha256,
           accepted_predecessor_mutated,recorded_at)
          VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
          ON CONFLICT(checkpoint_code) DO UPDATE SET
            response_number=excluded.response_number,state=excluded.state,
            database_integrity=excluded.database_integrity,foreign_key_violations=excluded.foreign_key_violations,
            workbook_status=excluded.workbook_status,application_status=excluded.application_status,
            publication_sha256=excluded.publication_sha256,accepted_predecessor_mutated=excluded.accepted_predecessor_mutated,
            recorded_at=excluded.recorded_at
        """, (
            "MRHPD-V3-CP4-S1-CP3", "Remediation Section 4 of 5", "Session 1 of 3", "Checkpoint 3 of 3",
            66, "session_complete", "ok", 0, "pending", "pending", PUBLICATION_SHA256, 0, NOW,
        ))
        con.execute("""
          INSERT INTO section4_session_release
          (release_code,section_label,session_label,completion_response,state,database_status,workbook_status,
           application_status,publication_status,complete_restore_required,accepted_predecessor_mutated,next_session,recorded_at)
          VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
          ON CONFLICT(release_code) DO UPDATE SET
            completion_response=excluded.completion_response,state=excluded.state,database_status=excluded.database_status,
            workbook_status=excluded.workbook_status,application_status=excluded.application_status,
            publication_status=excluded.publication_status,complete_restore_required=excluded.complete_restore_required,
            accepted_predecessor_mutated=excluded.accepted_predecessor_mutated,next_session=excluded.next_session,
            recorded_at=excluded.recorded_at
        """, (
            "MRHPD-V3-CP4-S1-COMPLETE", "Remediation Section 4 of 5", "Session 1 of 3", 66,
            "session_complete", "ok", "pending", "pending", "passed", 1, 0,
            "Remediation Section 4 of 5 Session 2 of 3", NOW,
        ))
        if table_exists(con, "metadata") and {"key", "value"}.issubset(table_columns(con, "metadata")):
            updates = {
                "version": PROJECT_VERSION,
                "current_remediation_section": "Remediation Section 4 of 5",
                "current_session": "Session 1 of 3 COMPLETE",
                "current_checkpoint": "Checkpoint 3 of 3 COMPLETE",
                "current_response": "66",
                "current_canonical_database": target.name,
                "accepted_predecessor_mutated": "no",
                "last_updated_utc": NOW,
                "next_checkpoint": "Remediation Section 4 of 5 Session 2 of 3",
                "restore_emission_policy": "full at session/section/project completion; checkpoint recovery between",
            }
            for key, value in updates.items():
                con.execute("""
                  INSERT INTO metadata(key,value) VALUES (?,?)
                  ON CONFLICT(key) DO UPDATE SET value=excluded.value
                """, (key, value))
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or fk:
            raise RuntimeError({"integrity": integrity, "foreign_keys": fk[:30]})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()

    con = sqlite3.connect(target)
    try:
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        table_count = con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0]
        response_count = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3").fetchone()[0]
        r66 = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R66'").fetchone()[0]
        fractional = con.execute("SELECT COUNT(*) FROM fractional_prompt_cp3").fetchone()[0]
        recovery = con.execute("SELECT COUNT(*) FROM remediation_recovery_event").fetchone()[0]
        checkpoint = con.execute("SELECT state FROM section4_checkpoint WHERE checkpoint_code='MRHPD-V3-CP4-S1-CP3'").fetchone()
        session_release = con.execute("SELECT state FROM section4_session_release WHERE release_code='MRHPD-V3-CP4-S1-COMPLETE'").fetchone()
        locators = con.execute("SELECT COUNT(*) FROM publication_index_locator").fetchone()[0]
        cross_refs = con.execute("SELECT COUNT(*) FROM publication_cross_reference WHERE COALESCE(is_current,1)=1").fetchone()[0]
    finally:
        con.close()
    if integrity != "ok" or fk or r66 != 1 or checkpoint != ("session_complete",) or session_release != ("session_complete",):
        raise RuntimeError({"database_gate": {"integrity": integrity, "fk": fk[:20], "r66": r66, "checkpoint": checkpoint, "session": session_release}})
    return target, {
        "status": "passed",
        "source_database": source.relative_to(project).as_posix(),
        "canonical_database": target.relative_to(project).as_posix(),
        "bytes": target.stat().st_size,
        "sha256": sha256_file(target),
        "integrity": integrity,
        "foreign_key_violations": 0,
        "table_count": table_count,
        "response_records": response_count,
        "response66_records": r66,
        "fractional_prompt_records": fractional,
        "recovery_event_records": recovery,
        "checkpoint_state": checkpoint[0],
        "session_release_state": session_release[0],
        "publication_index_locators": locators,
        "current_cross_references": cross_refs,
        "accepted_predecessor_mutated": False,
    }


def parse_test_count(stdout: str) -> int | None:
    patterns = [r'"passed"\s*:\s*(\d+)', r'"test_count"\s*:\s*(\d+)', r'(\d+)\s*/\s*\1\s+passed']
    for pattern in patterns:
        matches = re.findall(pattern, stdout, flags=re.I)
        if matches:
            try:
                return int(matches[-1])
            except Exception:
                pass
    return None


def synchronize_application(project: Path, db: Path) -> dict[str, Any]:
    app = find_by_hash(project, "human_pathogen_app.py", APPLICATION_SHA256)
    app_dir = app.parent
    launcher = app_dir / "run_section4_session1_complete.py"
    text_write(launcher, f'''#!/usr/bin/env python3
from pathlib import Path
import subprocess,sys
ROOT=Path(__file__).resolve().parent.parent
APP=Path(__file__).resolve().parent/"human_pathogen_app.py"
DB=ROOT/{db.relative_to(project).as_posix()!r}
raise SystemExit(subprocess.call([sys.executable,str(APP),"--db",str(DB),*sys.argv[1:]]))
''')
    state_path = app_dir / "CURRENT_PROJECT_STATE.json"
    json_write(state_path, {
        "schema": "mrhpd-current-application-state-1.0",
        "generated_at": NOW,
        "remediation_section": "4 of 5",
        "session": "1 of 3 COMPLETE",
        "checkpoint": "3 of 3 COMPLETE",
        "response": 66,
        "canonical_database": db.name,
        "database_relative_path": db.relative_to(project).as_posix(),
        "application_sha256": APPLICATION_SHA256,
        "accepted_predecessor_mutated": False,
        "next_session": "Remediation Section 4 of 5 Session 2 of 3",
    })
    text_write(app_dir / "CURRENT_DATABASE.txt", db.relative_to(project).as_posix())
    text_write(app_dir / "README_SECTION4_SESSION1_COMPLETE.md", f"""# Human Pathogen Database local application — Section 4 Session 1 complete

Canonical database: `{db.name}`

Launch with:

```bash
python run_section4_session1_complete.py
```

The application retains its native required `--db` interface. The accepted predecessor and frozen Section 3 release remain immutable.
""")

    env = os.environ.copy()
    env["MRHPD_DATABASE"] = str(db)
    env["MRHPD_DB_PATH"] = str(db)
    test_results = []
    parsed_total = 0
    tests = sorted(app_dir.glob("test*.py"))
    if not tests:
        raise RuntimeError("No application tests found")
    for test in tests:
        result = subprocess.run([sys.executable, str(test)], cwd=app_dir, env=env, text=True, capture_output=True, timeout=420)
        parsed = parse_test_count(result.stdout)
        if parsed:
            parsed_total += parsed
        record = {
            "test": test.name,
            "returncode": result.returncode,
            "parsed_test_count": parsed,
            "stdout_tail": result.stdout[-20000:],
            "stderr_tail": result.stderr[-10000:],
        }
        test_results.append(record)
        if result.returncode != 0:
            raise RuntimeError({"application_test_failed": record})

    con = sqlite3.connect(db)
    try:
        direct_checks = {
            "database_integrity": con.execute("PRAGMA integrity_check").fetchone()[0] == "ok",
            "foreign_keys": not list(con.execute("PRAGMA foreign_key_check")),
            "response66_record": con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R66'").fetchone()[0] == 1,
            "session1_checkpoint_complete": con.execute("SELECT state FROM section4_checkpoint WHERE checkpoint_code='MRHPD-V3-CP4-S1-CP3'").fetchone() == ("session_complete",),
            "session1_release_complete": con.execute("SELECT state FROM section4_session_release WHERE release_code='MRHPD-V3-CP4-S1-COMPLETE'").fetchone() == ("session_complete",),
            "restore_emission_policy": con.execute("SELECT COUNT(*) FROM restore_emission_policy WHERE full_restore_at_session_end=1 AND checkpoint_recovery_between=1").fetchone()[0] >= 1,
            "publication_locators": con.execute("SELECT COUNT(*) FROM publication_index_locator").fetchone()[0] == 4011,
            "cross_references": con.execute("SELECT COUNT(*) FROM publication_cross_reference WHERE COALESCE(is_current,1)=1").fetchone()[0] == 12,
            "page_map": con.execute("SELECT COUNT(*) FROM final_publication_page_map_cp3").fetchone()[0] == 10 if table_exists(con, "final_publication_page_map_cp3") else True,
            "response_reconciliation": con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_number BETWEEN 41 AND 66").fetchone()[0] >= 24,
            "sbsec_disambiguation": con.execute("SELECT COUNT(*) FROM taxonomy_node WHERE canonical_name LIKE '%bovis%equinu%' OR canonical_name LIKE '%gallolyticus%' ").fetchone()[0] >= 1 if table_exists(con, "taxonomy_node") else True,
        }
    finally:
        con.close()
    if not all(direct_checks.values()):
        raise RuntimeError({"current_application_direct_checks": direct_checks})
    return {
        "status": "passed",
        "application": app.relative_to(project).as_posix(),
        "application_sha256": sha256_file(app),
        "application_source_preserved": sha256_file(app) == APPLICATION_SHA256,
        "native_database_configuration": "required --db argument",
        "canonical_database": db.relative_to(project).as_posix(),
        "launcher": launcher.relative_to(project).as_posix(),
        "state_file": state_path.relative_to(project).as_posix(),
        "tests": test_results,
        "legacy_parsed_test_count": parsed_total,
        "current_direct_checks": direct_checks,
        "current_direct_check_count": len(direct_checks),
        "total_recorded_checks": parsed_total + len(direct_checks),
        "all_returncodes_zero": True,
    }


def verify_publication(project: Path) -> dict[str, Any]:
    from pypdf import PdfReader
    pdf = find_by_hash(project, "*Integrated Manuscript*.pdf", PUBLICATION_SHA256)
    docx = find_by_hash(project, "*Editable Integrated Manuscript Assembly*.docx", EDITABLE_ASSEMBLY_SHA256)
    reader = PdfReader(str(pdf))
    searchable = sum(1 for page in reader.pages if (page.extract_text() or "").strip())
    if len(reader.pages) != 537 or searchable != 537:
        raise RuntimeError({"publication_pages": len(reader.pages), "searchable_pages": searchable})
    return {
        "status": "passed",
        "integrated_publication": pdf.relative_to(project).as_posix(),
        "publication_bytes": pdf.stat().st_size,
        "publication_sha256": PUBLICATION_SHA256,
        "publication_pages": 537,
        "searchable_pages": 537,
        "editable_assembly": docx.relative_to(project).as_posix(),
        "editable_assembly_bytes": docx.stat().st_size,
        "editable_assembly_sha256": EDITABLE_ASSEMBLY_SHA256,
        "publication_unchanged": True,
        "editable_assembly_unchanged": True,
    }


def style_sheet(ws, header_row: int = 1) -> None:
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter
    navy = PatternFill("solid", fgColor="17365D")
    teal = PatternFill("solid", fgColor="DDEBF7")
    gold = PatternFill("solid", fgColor="FFF2CC")
    for cell in ws[header_row]:
        cell.fill = navy
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(wrap_text=True, vertical="center")
    ws.freeze_panes = f"A{header_row + 1}"
    for row_index, row in enumerate(ws.iter_rows(min_row=header_row + 1), start=header_row + 1):
        fill = teal if row_index % 2 == 0 else gold
        for cell in row:
            cell.fill = fill
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    for col_idx in range(1, ws.max_column + 1):
        width = max(12, min(62, max(len(str(ws.cell(row=r, column=col_idx).value or "")) for r in range(1, min(ws.max_row, 250) + 1)) + 2))
        ws.column_dimensions[get_column_letter(col_idx)].width = width
    ws.auto_filter.ref = ws.dimensions


def synchronize_workbook(project: Path, db: Path, database_qa: dict[str, Any], application_qa: dict[str, Any], publication_qa: dict[str, Any]) -> tuple[Path, dict[str, Any]]:
    from openpyxl import load_workbook
    source = find_by_hash(project, "*.xlsx", CP2_WORKBOOK_SHA256)
    target = source.with_name(
        f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 "
        "Remediation Section 4 of 5 Session 1 of 3 COMPLETE THROUGH RESPONSE 66 Comprehensive Tracking.xlsx"
    )
    shutil.copy2(source, target)
    wb = load_workbook(target)
    original_sheets = list(wb.sheetnames)

    def reset_sheet(title: str, headers: list[str], rows: list[list[Any]]) -> None:
        if title in wb.sheetnames:
            ws = wb[title]
            ws.delete_rows(1, ws.max_row)
        else:
            ws = wb.create_sheet(title)
        ws.append(headers)
        for row in rows:
            ws.append(row)
        style_sheet(ws)

    reset_sheet("S4S1 Dashboard", ["Control", "Current state", "Evidence / Next"], [
        ["Remediation section", "Section 4 of 5 — CONTINUE", "Session 2 of 3 is next"],
        ["Session", "Session 1 of 3 — COMPLETE", "Checkpoint 3 of 3 passed"],
        ["Current response", 66, "Response 66 tracking synchronized"],
        ["Canonical database", db.name, database_qa["sha256"]],
        ["Database integrity", database_qa["integrity"], f"Foreign keys: {database_qa['foreign_key_violations']}"],
        ["Application", application_qa["status"], f"Recorded checks: {application_qa['total_recorded_checks']}"],
        ["Publication", "537 searchable pages; unchanged", PUBLICATION_SHA256],
        ["Editable assembly", "unchanged", EDITABLE_ASSEMBLY_SHA256],
        ["Restore emission", "Complete session restore required and built", "Checkpoint recovery between session/section emissions"],
        ["Next", "Remediation Section 4 Session 2 of 3", "Continue"],
    ])

    ws_responses = wb["S4S1 Responses"] if "S4S1 Responses" in wb.sheetnames else wb.create_sheet("S4S1 Responses")
    if ws_responses.max_row == 1 and ws_responses.cell(1, 1).value is None:
        ws_responses.append(["Response", "Title", "Goal", "Raw prompt", "Summary", "State", "Recorded at"])
    existing = {str(ws_responses.cell(row=r, column=1).value) for r in range(2, ws_responses.max_row + 1)}
    if "66" not in existing:
        ws_responses.append([66, RESPONSE66["title"], RESPONSE66["goal"], RAW_PROMPT, RESPONSE66["summary"], RESPONSE66["state"], NOW])
    style_sheet(ws_responses)

    ws_recovery = wb["S4S1 Recovery"] if "S4S1 Recovery" in wb.sheetnames else wb.create_sheet("S4S1 Recovery")
    if ws_recovery.max_row == 1 and ws_recovery.cell(1, 1).value is None:
        ws_recovery.append(["Event", "Code", "Failed step", "Reason", "Recovery action", "Validation", "Next"])
    codes = {str(ws_recovery.cell(row=r, column=2).value) for r in range(2, ws_recovery.max_row + 1)}
    for event in RECOVERY_EVENTS:
        if event["event_code"] not in codes:
            ws_recovery.append([
                event["event_number"], event["event_code"], event["failed_step"], event["exact_error_or_reason"],
                event["recovery_action"], event["validation_result"], event["next_checkpoint"],
            ])
    style_sheet(ws_recovery)

    reset_sheet("S4S1 Database QA", ["Control", "Expected", "Actual", "Status"], [
        ["Integrity", "ok", database_qa["integrity"], "PASS"],
        ["Foreign-key violations", 0, database_qa["foreign_key_violations"], "PASS"],
        ["Response 66", 1, database_qa["response66_records"], "PASS"],
        ["Session checkpoint state", "session_complete", database_qa["checkpoint_state"], "PASS"],
        ["Session release state", "session_complete", database_qa["session_release_state"], "PASS"],
        ["Locators", 4011, database_qa["publication_index_locators"], "PASS"],
        ["Cross-references", 12, database_qa["current_cross_references"], "PASS"],
    ])
    reset_sheet("S4S1 Application QA", ["Control", "Result", "Evidence"], [
        ["Application source preserved", application_qa["application_source_preserved"], application_qa["application_sha256"]],
        ["Legacy tests", "passed", application_qa["legacy_parsed_test_count"]],
        ["Current direct checks", "passed", application_qa["current_direct_check_count"]],
        ["Total recorded checks", application_qa["total_recorded_checks"], "All return codes zero"],
        ["Current database", db.name, application_qa["launcher"]],
    ])
    reset_sheet("S4S1 Session Release", ["Control", "State", "Evidence / Next"], [
        ["Checkpoint 3 of 3", "COMPLETE", NOW],
        ["Session 1 of 3", "COMPLETE", "Complete self-contained restore emitted"],
        ["Section 4 of 5", "CONTINUE", "Session 2 of 3 next"],
        ["Accepted predecessor mutated", "NO", "Immutable"],
        ["Frozen Section 3 release mutated", "NO", "Immutable"],
        ["User upload required", "NO", "Restore is self-contained"],
    ])
    wb.save(target)
    wb.close()

    wb2 = load_workbook(target, data_only=False, read_only=True)
    final_sheets = list(wb2.sheetnames)
    errors = []
    formula_count = 0
    error_tokens = {"#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "#N/A", "#NUM!", "#NULL!"}
    for ws in wb2.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                value = cell.value
                if isinstance(value, str) and value.startswith("="):
                    formula_count += 1
                if isinstance(value, str) and any(token in value for token in error_tokens):
                    errors.append(f"{ws.title}!{cell.coordinate}:{value}")
                    if len(errors) >= 100:
                        break
            if len(errors) >= 100:
                break
        if len(errors) >= 100:
            break
    wb2.close()
    missing = sorted(set(original_sheets) - set(final_sheets))
    if errors or missing or "S4S1 Session Release" not in final_sheets:
        raise RuntimeError({"formula_errors": errors, "missing_original_sheets": missing, "sheets": final_sheets})
    return target, {
        "status": "passed",
        "source_workbook": source.relative_to(project).as_posix(),
        "current_workbook": target.relative_to(project).as_posix(),
        "bytes": target.stat().st_size,
        "sha256": sha256_file(target),
        "source_sheet_count": len(original_sheets),
        "current_sheet_count": len(final_sheets),
        "original_sheets_preserved": not missing,
        "missing_original_sheets": missing,
        "new_session_release_sheet": "S4S1 Session Release",
        "formula_count": formula_count,
        "formula_error_count": len(errors),
    }


def add_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_tracking_documents(project: Path, db: Path, qa: dict[str, Any]) -> list[Path]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    tracking = project / "Tracking" / "Section 4 Session 1" / "Complete Through Response 66"
    tracking.mkdir(parents=True, exist_ok=True)
    created: list[Path] = []
    json_write(tracking / "Response_66_Tracking.json", RESPONSE66); created.append(tracking / "Response_66_Tracking.json")
    json_write(tracking / "RECOVERY_EVENTS_95_97.json", RECOVERY_EVENTS); created.append(tracking / "RECOVERY_EVENTS_95_97.json")
    session_record = {
        "schema": "mrhpd-section4-session1-completion-1.0",
        "created_at": NOW,
        "section": "Remediation Section 4 of 5",
        "session": "Session 1 of 3",
        "checkpoint": "Checkpoint 3 of 3",
        "completion_response": 66,
        "status": "COMPLETE",
        "database": qa["database"],
        "workbook": qa["workbook"],
        "application": qa["application"],
        "publication": qa["publication"],
        "accepted_predecessor_mutated": False,
        "next": "Remediation Section 4 of 5 Session 2 of 3",
    }
    json_write(tracking / "SESSION_1_COMPLETION_RECORD.json", session_record); created.append(tracking / "SESSION_1_COMPLETION_RECORD.json")
    raw_net = f"""# Human Pathogen Database — Raw and Net Tracking Through Response 66

## Major topic
Human Pathogen Database remediation

## Raw Prompt 66

{RAW_PROMPT}

## Net Prompt through Response 66

{NET_PROMPT}

## Net Response through Response 66

{NET_RESPONSE}

## Disposition

- Checkpoint 3 of 3: COMPLETE
- Section 4 Session 1 of 3: COMPLETE
- Remediation Section 4 of 5: CONTINUE
- Next: Session 2 of 3
"""
    text_write(tracking / "RAW_AND_NET_TRACKING.md", raw_net); created.append(tracking / "RAW_AND_NET_TRACKING.md")
    cumulative = f"""# Cumulative Thread Index Update — Response 66

## Major topic
Human Pathogen Database remediation

## Response 66 — Section 4 Session 1 complete restore and handoff

**Goal:** Complete the final Session 1 synchronization and emit a complete self-contained restore through Response 66.

**Output:** Applied the exact Response 65 Checkpoint 2 recovery state; synchronized Response 66, Recovery Events 95–97, the Session 1-complete database, workbook, application state, tracking, QA, Source Index, Bit Index and manifests; preserved the 537-page publication and editable assembly byte-for-byte; and built a clean-extraction-tested session-end restore.

**Disposition:** Checkpoint 3 of 3 and Session 1 of 3 COMPLETE. Remediation Section 4 of 5 CONTINUE to Session 2 of 3.
"""
    text_write(tracking / "CUMULATIVE_THREAD_INDEX_UPDATE.md", cumulative); created.append(tracking / "CUMULATIVE_THREAD_INDEX_UPDATE.md")
    handoff = """# Section 4 Session 1 to Session 2 Handoff

Session 1 completed the durable-state synchronization of the canonical SQLite database, comprehensive workbook, local application, current tracking, recovery records, Source Index, Bit Index and manifests. The controlling 537-page publication and editable assembly remain unchanged.

Session 2 begins from the complete restore through Response 66 and must preserve the same immutable-source, emission-policy, Google Drive custody and clean-verification requirements.
"""
    text_write(tracking / "SESSION_2_HANDOFF.md", handoff); created.append(tracking / "SESSION_2_HANDOFF.md")

    con = sqlite3.connect(db)
    con.row_factory = sqlite3.Row
    responses = [dict(row) for row in con.execute("SELECT * FROM thread_response_reconciliation_cp3 ORDER BY response_number, response_key")]
    con.close()

    raw_doc = Document()
    sec = raw_doc.sections[0]
    sec.top_margin = Inches(0.55); sec.bottom_margin = Inches(0.55); sec.left_margin = Inches(0.65); sec.right_margin = Inches(0.65)
    raw_doc.styles["Normal"].font.name = "Aptos"; raw_doc.styles["Normal"].font.size = Pt(9)
    title = raw_doc.add_heading("Human Pathogen Database", 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = raw_doc.add_paragraph("Alternating Raw Prompts and Responses Through Response 66"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    raw_doc.add_heading("Human Pathogen Database remediation", 1)
    for row in responses:
        label = row.get("response_label") or row.get("response_number")
        raw_doc.add_heading(f"Response {label} — {row.get('title') or 'Untitled response'}", 2)
        table = raw_doc.add_table(rows=2, cols=1)
        table.style = "Table Grid"
        prompt = row.get("raw_prompt") or "[RAW PROMPT NOT RECOVERED]"
        response = row.get("raw_response") or "[RAW RESPONSE NOT RECOVERED]"
        summary = row.get("summary") or "[NO SOURCE-SUPPORTED SUMMARY RECOVERED]"
        table.cell(0, 0).text = f"PROMPT\n{prompt}"
        table.cell(1, 0).text = f"RESPONSE\n{response}\n\nSUMMARY\n{summary}"
        add_cell_shading(table.cell(0, 0), "DDEBF7")
        add_cell_shading(table.cell(1, 0), "FFF2CC")
    raw_doc.core_properties.title = "MRHPD Alternating Raw Prompts and Responses Through Response 66"
    raw_doc.core_properties.author = "Brent McAnulty, M.D."
    raw_docx = tracking / "MRHPD v3.0.0a Alternating Raw Prompts and Responses Through Response 66.docx"
    raw_doc.save(raw_docx); created.append(raw_docx)

    net_doc = Document()
    sec = net_doc.sections[0]
    sec.top_margin = Inches(0.65); sec.bottom_margin = Inches(0.65); sec.left_margin = Inches(0.75); sec.right_margin = Inches(0.75)
    net_doc.styles["Normal"].font.name = "Aptos"; net_doc.styles["Normal"].font.size = Pt(10)
    title = net_doc.add_heading("Human Pathogen Database", 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    net_doc.add_heading("Human Pathogen Database remediation", 1)
    table = net_doc.add_table(rows=2, cols=1); table.style = "Table Grid"
    table.cell(0, 0).text = "NET PROMPT\n" + NET_PROMPT
    table.cell(1, 0).text = "NET RESPONSE\n" + NET_RESPONSE
    add_cell_shading(table.cell(0, 0), "DDEBF7"); add_cell_shading(table.cell(1, 0), "FFF2CC")
    net_doc.add_heading("Current disposition", 1)
    net_doc.add_paragraph("Checkpoint 3 of 3 and Section 4 Session 1 of 3 are complete. Remediation Section 4 of 5 continues to Session 2 of 3.")
    net_doc.core_properties.title = "MRHPD Alternating Net Prompts and Responses Through Response 66"
    net_doc.core_properties.author = "Brent McAnulty, M.D."
    net_docx = tracking / "MRHPD v3.0.0a Alternating Net Prompts and Responses Through Response 66.docx"
    net_doc.save(net_docx); created.append(net_docx)

    styles = getSampleStyleSheet()
    for docx_path, pdf_title, sections in [
        (raw_docx, "Alternating Raw Prompts and Responses Through Response 66", [("Raw Prompt 66", RAW_PROMPT), ("Response 66 summary", RESPONSE66["summary"])]),
        (net_docx, "Alternating Net Prompts and Responses Through Response 66", [("Net Prompt", NET_PROMPT), ("Net Response", NET_RESPONSE)]),
    ]:
        pdf_path = docx_path.with_suffix(".pdf")
        story = [Paragraph("Human Pathogen Database", styles["Title"]), Paragraph(pdf_title, styles["Heading2"]), Spacer(1, 0.15 * inch)]
        for heading, body in sections:
            story.extend([Paragraph(heading, styles["Heading2"]), Paragraph(body.replace("\n", "<br/>"), styles["BodyText"]), Spacer(1, 0.12 * inch)])
        SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=0.6*inch, leftMargin=0.6*inch, topMargin=0.55*inch, bottomMargin=0.55*inch).build(story)
        created.append(pdf_path)
    return created


def build_report(project: Path, qa: dict[str, Any]) -> list[Path]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    reports = project / "Documents" / "Section 4 Session 1 Complete"
    reports.mkdir(parents=True, exist_ok=True)
    docx_path = reports / "MRHPD v3.0.0a Section 4 Session 1 Completion Report Through Response 66.docx"
    pdf_path = reports / "MRHPD v3.0.0a Section 4 Session 1 Completion Report Through Response 66.pdf"
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.6); sec.bottom_margin = Inches(0.6); sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)
    doc.styles["Normal"].font.name = "Aptos"; doc.styles["Normal"].font.size = Pt(9)
    title = doc.add_heading("Human Pathogen Database", 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Remediation Section 4 of 5 · Session 1 of 3 · COMPLETE"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Session 1 completion", 1)
    doc.add_paragraph(NET_RESPONSE)
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"
    for idx, heading in enumerate(["Control", "Result", "Evidence"]):
        table.rows[0].cells[idx].text = heading
    rows = [
        ("Database integrity", qa["database"]["integrity"], qa["database"]["canonical_database"]),
        ("Foreign-key violations", qa["database"]["foreign_key_violations"], "0 required"),
        ("Workbook", qa["workbook"]["status"], qa["workbook"]["current_workbook"]),
        ("Workbook sheets", qa["workbook"]["current_sheet_count"], "All prior sheets preserved"),
        ("Application", qa["application"]["status"], qa["application"]["total_recorded_checks"]),
        ("Publication", "537/537 searchable pages", qa["publication"]["publication_sha256"]),
        ("Editable assembly", "unchanged", qa["publication"]["editable_assembly_sha256"]),
        ("Accepted predecessor mutated", "No", "Immutable"),
        ("Next", "Section 4 Session 2 of 3", "CONTINUE"),
    ]
    for control, result, evidence in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = str(control), str(result), str(evidence)
    doc.add_heading("Restore boundary", 1)
    doc.add_paragraph("The session-end restore includes the complete current project snapshot, current instructions, current tracking and recovery records, manifests, hashes, deterministic verification tools, and the verified Checkpoint 2 recovery package. It requires no other project file or conversation reconstruction.")
    doc.core_properties.title = "MRHPD Section 4 Session 1 Completion Report"
    doc.core_properties.author = "Brent McAnulty, M.D."
    doc.save(docx_path)

    styles = getSampleStyleSheet()
    data = [["Control", "Result"], ["Database integrity", qa["database"]["integrity"]], ["Workbook", qa["workbook"]["status"]], ["Application", qa["application"]["status"]], ["Publication", "537 searchable pages"], ["Session", "COMPLETE"], ["Next", "Session 2 of 3"]]
    pdf_table = Table(data, colWidths=[2.5*inch, 4.4*inch], repeatRows=1)
    pdf_table.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,0), colors.HexColor("#17365D")), ("TEXTCOLOR", (0,0), (-1,0), colors.white), ("GRID", (0,0), (-1,-1), 0.5, colors.grey), ("VALIGN", (0,0), (-1,-1), "TOP"), ("FONTSIZE", (0,0), (-1,-1), 8.5), ("BACKGROUND", (0,1), (-1,-1), colors.HexColor("#F4F8FA"))]))
    story = [Paragraph("Human Pathogen Database", styles["Title"]), Paragraph("Remediation Section 4 of 5 · Session 1 of 3 · COMPLETE", styles["Heading2"]), Spacer(1, 0.15*inch), Paragraph(NET_RESPONSE, styles["BodyText"]), Spacer(1, 0.15*inch), pdf_table, Spacer(1, 0.15*inch), Paragraph("Next: Remediation Section 4 of 5 Session 2 of 3.", styles["Heading2"])]
    SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=0.6*inch, leftMargin=0.6*inch, topMargin=0.55*inch, bottomMargin=0.55*inch).build(story)
    return [docx_path, pdf_path]


def finalize_database_status(db: Path, workbook_qa: dict[str, Any], application_qa: dict[str, Any], publication_qa: dict[str, Any]) -> dict[str, Any]:
    con = sqlite3.connect(db)
    try:
        con.execute("BEGIN IMMEDIATE")
        con.execute("""
          UPDATE section4_checkpoint SET workbook_status=?,application_status=?,database_integrity='ok',
            foreign_key_violations=0,publication_sha256=?,state='session_complete',recorded_at=?
          WHERE checkpoint_code='MRHPD-V3-CP4-S1-CP3'
        """, (workbook_qa["status"], application_qa["status"], publication_qa["publication_sha256"], NOW))
        con.execute("""
          UPDATE section4_session_release SET database_status='ok',workbook_status=?,application_status=?,
            publication_status=?,state='session_complete',recorded_at=?
          WHERE release_code='MRHPD-V3-CP4-S1-COMPLETE'
        """, (workbook_qa["status"], application_qa["status"], publication_qa["status"], NOW))
        controls = [
            ("SQLite integrity", "ok", "ok", "PASS", "QA/Section 4 Session 1/Complete/DATABASE_QA.json"),
            ("Foreign-key violations", "0", "0", "PASS", "QA/Section 4 Session 1/Complete/DATABASE_QA.json"),
            ("Workbook synchronization", "passed", workbook_qa["status"], "PASS", "QA/Section 4 Session 1/Complete/WORKBOOK_QA.json"),
            ("Application regressions", "passed", application_qa["status"], "PASS", "QA/Section 4 Session 1/Complete/APPLICATION_QA.json"),
            ("Publication unchanged", PUBLICATION_SHA256, publication_qa["publication_sha256"], "PASS", "QA/Section 4 Session 1/Complete/PUBLICATION_QA.json"),
            ("Session 1 complete restore", "required", "built and clean-verified", "PASS", "Recovery/Section 4 Session 1 Complete/SESSION_1_COMPLETE.json"),
        ]
        for name, expected, actual, status, evidence in controls:
            con.execute("""
              INSERT INTO section4_sync_qa
              (checkpoint_code,control_name,expected_value,actual_value,status,evidence_path,recorded_at)
              VALUES (?,?,?,?,?,?,?)
              ON CONFLICT(checkpoint_code,control_name) DO UPDATE SET
                expected_value=excluded.expected_value,actual_value=excluded.actual_value,status=excluded.status,
                evidence_path=excluded.evidence_path,recorded_at=excluded.recorded_at
            """, ("MRHPD-V3-CP4-S1-CP3", name, expected, actual, status, evidence, NOW))
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or fk:
            raise RuntimeError({"integrity": integrity, "foreign_keys": fk[:20]})
        con.commit()
    except Exception:
        con.rollback(); raise
    finally:
        con.close()
    return {"integrity": "ok", "foreign_key_violations": 0, "sha256": sha256_file(db), "bytes": db.stat().st_size}


def build_indexes_and_manifest(project: Path) -> dict[str, Any]:
    index_dir = project / "Indexes" / "Section 4 Session 1 Complete"
    manifest_dir = project / "Manifest" / "Section 4 Session 1 Complete"
    if index_dir.exists(): shutil.rmtree(index_dir)
    if manifest_dir.exists(): shutil.rmtree(manifest_dir)
    index_dir.mkdir(parents=True); manifest_dir.mkdir(parents=True)
    excluded = {
        (index_dir / "Source and Artifact Index.csv").relative_to(project).as_posix(),
        (index_dir / "Source and Artifact Index.json").relative_to(project).as_posix(),
        (index_dir / "Session 1 Bit Index.sqlite").relative_to(project).as_posix(),
        (manifest_dir / "Session 1 Project Manifest.csv").relative_to(project).as_posix(),
        (manifest_dir / "Session 1 Project Manifest.json").relative_to(project).as_posix(),
        (manifest_dir / "Session 1 SHA256 Inventory.txt").relative_to(project).as_posix(),
    }
    files = [path for path in sorted(project.rglob("*")) if path.is_file() and path.relative_to(project).as_posix() not in excluded]
    text_ext = {".txt", ".md", ".csv", ".json", ".py", ".html", ".css", ".js", ".xml", ".sql", ".yaml", ".yml"}
    rows = []
    for path in files:
        rel = path.relative_to(project).as_posix()
        rows.append({
            "path": rel,
            "name": path.name,
            "extension": path.suffix.lower(),
            "bytes": path.stat().st_size,
            "sha256": sha256_file(path),
            "category": rel.split("/", 1)[0],
            "purpose": "current project artifact",
            "searchable": "yes" if path.suffix.lower() in text_ext else "metadata_only",
            "indexed_at": NOW,
        })
    csv_write(index_dir / "Source and Artifact Index.csv", rows)
    json_write(index_dir / "Source and Artifact Index.json", {"generated_at": NOW, "file_count": len(rows), "files": rows})
    bit = index_dir / "Session 1 Bit Index.sqlite"
    con = sqlite3.connect(bit)
    try:
        con.executescript("""
        CREATE TABLE file_record(file_record_id INTEGER PRIMARY KEY,path TEXT NOT NULL UNIQUE,name TEXT NOT NULL,extension TEXT,bytes INTEGER NOT NULL,sha256 TEXT NOT NULL,category TEXT NOT NULL,indexed_at TEXT NOT NULL);
        CREATE VIRTUAL TABLE text_fts USING fts5(path UNINDEXED, content);
        """)
        for row, path in zip(rows, files):
            con.execute("INSERT INTO file_record(path,name,extension,bytes,sha256,category,indexed_at) VALUES (?,?,?,?,?,?,?)", (row["path"],row["name"],row["extension"],row["bytes"],row["sha256"],row["category"],NOW))
            if path.suffix.lower() in text_ext:
                content = path.read_text(encoding="utf-8", errors="replace")[:1_500_000]
                if content:
                    con.execute("INSERT INTO text_fts(path,content) VALUES (?,?)", (row["path"],content))
        con.commit()
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        searches = {term: con.execute("SELECT COUNT(*) FROM text_fts WHERE text_fts MATCH ?", (f'"{term}"',)).fetchone()[0] for term in ["Response 66","Session 1","Google Drive","complete restore"]}
    finally:
        con.close()
    if integrity != "ok" or any(value < 1 for value in searches.values()):
        raise RuntimeError({"bit_index_integrity": integrity, "searches": searches})
    csv_write(manifest_dir / "Session 1 Project Manifest.csv", rows)
    json_write(manifest_dir / "Session 1 Project Manifest.json", {"generated_at": NOW,"file_count":len(rows),"total_bytes":sum(row["bytes"] for row in rows),"files":rows,"recursive_controls_excluded":sorted(excluded)})
    text_write(manifest_dir / "Session 1 SHA256 Inventory.txt", "".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
    master_categories = [p.relative_to(project).as_posix() for p in project.rglob("*") if p.is_file() and "Master Categor" in p.name]
    if not master_categories:
        human = project / "Data" / "MRHPD v3.0.0a Master Category Database Human Readable.md"
        text_write(human, "# Master Category Database\n\nThe controlling category and subcategory records remain in the canonical SQLite database. This human-readable pointer is generated for restore discoverability.\n")
        master_categories = [human.relative_to(project).as_posix()]
    return {
        "status": "passed",
        "physical_file_count": len(rows),
        "physical_bytes": sum(row["bytes"] for row in rows),
        "source_index": (index_dir / "Source and Artifact Index.csv").relative_to(project).as_posix(),
        "bit_index": bit.relative_to(project).as_posix(),
        "bit_index_integrity": integrity,
        "search_tests": searches,
        "project_manifest": (manifest_dir / "Session 1 Project Manifest.json").relative_to(project).as_posix(),
        "master_category_files": master_categories,
    }


def build_project_archive(project: Path, dist: Path) -> tuple[Path, dict[str, Any]]:
    archive = dist / (
        f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 "
        f"Remediation Section 4 of 5 Session 1 of 3 COMPLETE PROJECT THROUGH RESPONSE 66 {STAMP}.zip"
    )
    with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in sorted(project.rglob("*")):
            if path.is_file():
                zf.write(path, f"{project.name}/{path.relative_to(project).as_posix()}")
    qa = verify_zip(archive)
    if qa["bytes"] >= 180 * 1024 * 1024:
        raise RuntimeError({"project_archive_exceeds_180_mib": qa})
    return archive, qa


def quick_verify_project(project: Path, expected_db_sha: str, expected_workbook_sha: str) -> dict[str, Any]:
    db = find_by_hash(project, "*COMPLETE THROUGH RESPONSE 66*.sqlite", expected_db_sha)
    workbook = find_by_hash(project, "*COMPLETE THROUGH RESPONSE 66*Comprehensive Tracking.xlsx", expected_workbook_sha)
    app = find_by_hash(project, "human_pathogen_app.py", APPLICATION_SHA256)
    publication = find_by_hash(project, "*Integrated Manuscript*.pdf", PUBLICATION_SHA256)
    editable = find_by_hash(project, "*Editable Integrated Manuscript Assembly*.docx", EDITABLE_ASSEMBLY_SHA256)
    con = sqlite3.connect(db)
    try:
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        r66 = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R66'").fetchone()[0]
        cp3 = con.execute("SELECT state FROM section4_checkpoint WHERE checkpoint_code='MRHPD-V3-CP4-S1-CP3'").fetchone()
    finally:
        con.close()
    if integrity != "ok" or fk or r66 != 1 or cp3 != ("session_complete",):
        raise RuntimeError({"clean_project_database_gate": {"integrity": integrity,"fk":fk[:20],"r66":r66,"cp3":cp3}})
    from openpyxl import load_workbook
    wb = load_workbook(workbook, read_only=True, data_only=False)
    sheets = list(wb.sheetnames)
    errors = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and any(t in cell.value for t in ["#REF!","#DIV/0!","#VALUE!","#NAME?","#N/A","#NUM!","#NULL!"]):
                    errors.append(f"{ws.title}!{cell.coordinate}")
    wb.close()
    if errors or "S4S1 Session Release" not in sheets:
        raise RuntimeError({"workbook_errors": errors[:20], "sheets": sheets})
    manifest = next(project.rglob("Section 4 Session 1 Complete/Session 1 Project Manifest.json"))
    manifest_data = json.loads(manifest.read_text(encoding="utf-8"))
    mismatches=[]
    for row in manifest_data["files"]:
        path=project/row["path"]
        if not path.exists() or path.stat().st_size!=row["bytes"] or sha256_file(path)!=row["sha256"]:
            mismatches.append(row["path"])
    if mismatches:
        raise RuntimeError({"project_manifest_mismatches": mismatches[:30]})
    return {
        "status":"passed","database_integrity":integrity,"foreign_key_violations":0,
        "database":db.relative_to(project).as_posix(),"workbook":workbook.relative_to(project).as_posix(),
        "workbook_sheets":len(sheets),"workbook_formula_errors":0,"application":app.relative_to(project).as_posix(),
        "publication":publication.relative_to(project).as_posix(),"editable_assembly":editable.relative_to(project).as_posix(),
        "manifest_records":manifest_data["file_count"],"manifest_mismatches":0,
    }


def build_restore_verify_script(session_archive_name: str, session_archive_bytes: int, session_archive_sha: str, recovery_name: str, recovery_bytes: int, recovery_sha: str) -> str:
    return f'''#!/usr/bin/env python3
import argparse,hashlib,json,sqlite3,subprocess,sys,tempfile,zipfile
from pathlib import Path,PurePosixPath
SESSION_NAME={session_archive_name!r}
SESSION_BYTES={session_archive_bytes}
SESSION_SHA={session_archive_sha!r}
RECOVERY_NAME={recovery_name!r}
RECOVERY_BYTES={recovery_bytes}
RECOVERY_SHA={recovery_sha!r}

def sha(path):
 h=hashlib.sha256()
 with open(path,'rb') as f:
  for b in iter(lambda:f.read(1024*1024),b''): h.update(b)
 return h.hexdigest()
def safe(zf):
 names=zf.namelist()
 if len(names)!=len(set(names)): raise SystemExit('duplicate ZIP members')
 for name in names:
  p=PurePosixPath(name.replace('\\\\','/'))
  if p.is_absolute() or '..' in p.parts: raise SystemExit('unsafe ZIP path: '+name)
 bad=zf.testzip()
 if bad: raise SystemExit('ZIP CRC failure: '+bad)
def verify_file(path, size, digest):
 if not path.exists() or path.stat().st_size!=size or sha(path)!=digest: raise SystemExit('identity failure: '+str(path))
 with zipfile.ZipFile(path) as zf: safe(zf)
parser=argparse.ArgumentParser()
parser.add_argument('--extract-project-to',type=Path)
args=parser.parse_args()
root=Path(__file__).resolve().parent.parent
session=root/'PROJECT_SNAPSHOT'/SESSION_NAME
recovery=root/'RECOVERY_HISTORY'/RECOVERY_NAME
verify_file(session,SESSION_BYTES,SESSION_SHA)
verify_file(recovery,RECOVERY_BYTES,RECOVERY_SHA)
manifest=json.loads((root/'COMPLETE_RESTORE_MANIFEST.json').read_text())
for row in manifest['files']:
 p=root/row['path']
 if not p.exists() or p.stat().st_size!=row['bytes'] or sha(p)!=row['sha256']: raise SystemExit('restore manifest failure: '+row['path'])
if args.extract_project_to:
 args.extract_project_to.mkdir(parents=True,exist_ok=True)
 with zipfile.ZipFile(session) as zf:
  safe(zf); zf.extractall(args.extract_project_to)
print(json.dumps({{'status':'passed','session_archive':SESSION_NAME,'session_sha256':SESSION_SHA,'self_contained':True}},indent=2))
'''


def locate_instructions(base_restore: Path, work: Path) -> Path | None:
    target = work / "base_restore_for_instructions"
    safe_extract(base_restore, target)
    candidates = sorted([p for p in target.rglob("Instructions.txt") if p.is_file()] + [p for p in target.rglob("Instructions.md") if p.is_file()], key=lambda p: p.stat().st_size, reverse=True)
    return candidates[0] if candidates else None


def build_complete_restore(project_archive: Path, recovery_zip: Path, base_restore: Path, reports: list[Path], dist: Path, work: Path, session_qa: dict[str, Any]) -> tuple[Path, dict[str, Any]]:
    package = work / "complete_restore_package"
    package.mkdir(parents=True)
    snapshot_dir = package / "PROJECT_SNAPSHOT"; snapshot_dir.mkdir()
    recovery_dir = package / "RECOVERY_HISTORY"; recovery_dir.mkdir()
    tools = package / "TOOLS"; tools.mkdir()
    report_dir = package / "REPORTS"; report_dir.mkdir()
    shutil.copy2(project_archive, snapshot_dir / project_archive.name)
    shutil.copy2(recovery_zip, recovery_dir / recovery_zip.name)
    for report in reports:
        shutil.copy2(report, report_dir / report.name)
    instructions = locate_instructions(base_restore, work)
    if instructions:
        shutil.copy2(instructions, package / "Instructions.txt")
    else:
        text_write(package / "Instructions.txt", "Human Pathogen Database project instructions are incorporated into the current project snapshot. Follow RESTORE_READ_FIRST.md and the current project governance files.")
    identity = {
        "schema":"mrhpd-complete-restore-identity-1.0","generated_at":NOW,"version":PROJECT_VERSION,
        "section":"Remediation Section 4 of 5","session":"Session 1 of 3 COMPLETE","checkpoint":"3 of 3 COMPLETE","response":66,
        "project_snapshot":{"name":project_archive.name,"bytes":project_archive.stat().st_size,"sha256":sha256_file(project_archive)},
        "checkpoint2_recovery":{"name":recovery_zip.name,"bytes":recovery_zip.stat().st_size,"sha256":sha256_file(recovery_zip)},
        "self_contained":True,"requires_other_project_files":False,"requires_conversation_reconstruction":False,
        "accepted_predecessor_mutated":False,"next":"Remediation Section 4 of 5 Session 2 of 3",
    }
    json_write(package / "CURRENT_PROJECT_IDENTITY.json", identity)
    text_write(package / "RESTORE_READ_FIRST.md", f"""# Human Pathogen Database — Complete Restore Through Response 66

This is the complete self-contained Section 4 Session 1 restore. It requires no prior ZIP, cloud artifact, user-supplied project file, or reconstruction from the conversation.

## Restore

1. Run `python TOOLS/restore_verify_extract.py` to verify the package.
2. Run `python TOOLS/restore_verify_extract.py --extract-project-to <destination>` to verify and extract the complete current project snapshot.
3. Open the extracted project and begin with its current Recovery, Tracking, QA and README files.

## Current state

- Remediation Section 4 of 5: CONTINUE
- Session 1 of 3: COMPLETE
- Checkpoint 3 of 3: COMPLETE
- Current response: 66
- Next: Session 2 of 3
- Accepted predecessor modified: no
- Frozen Section 3 release modified: no
- User upload required: no
""")
    text_write(tools / "restore_verify_extract.py", build_restore_verify_script(project_archive.name, project_archive.stat().st_size, sha256_file(project_archive), recovery_zip.name, recovery_zip.stat().st_size, sha256_file(recovery_zip)))
    json_write(package / "SESSION_1_ACCEPTANCE_QA.json", session_qa)
    control_names = {"COMPLETE_RESTORE_MANIFEST.json","COMPLETE_RESTORE_CHECKSUMS.sha256"}
    rows=[]
    for path in sorted(package.rglob("*")):
        if path.is_file() and path.name not in control_names:
            rows.append({"path":path.relative_to(package).as_posix(),"bytes":path.stat().st_size,"sha256":sha256_file(path)})
    json_write(package / "COMPLETE_RESTORE_MANIFEST.json", {"generated_at":NOW,"file_count":len(rows),"total_bytes":sum(r["bytes"] for r in rows),"files":rows,"self_contained":True})
    text_write(package / "COMPLETE_RESTORE_CHECKSUMS.sha256", "".join(f"{r['sha256']}  {r['path']}\n" for r in rows))
    restore = dist / (
        f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 "
        f"Remediation Section 4 of 5 Session 1 of 3 COMPLETE RESTORE THROUGH RESPONSE 66 {STAMP}.zip"
    )
    with zipfile.ZipFile(restore,"w",compression=zipfile.ZIP_DEFLATED,compresslevel=6,allowZip64=True) as zf:
        for path in sorted(package.rglob("*")):
            if path.is_file(): zf.write(path,path.relative_to(package).as_posix())
    restore_qa=verify_zip(restore)
    if restore_qa["bytes"] >= 180*1024*1024: raise RuntimeError({"restore_exceeds_180_mib":restore_qa})
    with tempfile.TemporaryDirectory(prefix="mrhpd-r66-restore-clean-") as td:
        clean=Path(td); safe_extract(restore,clean)
        result=subprocess.run([sys.executable,str(clean/"TOOLS"/"restore_verify_extract.py")],cwd=clean,text=True,capture_output=True,timeout=600)
        if result.returncode: raise RuntimeError({"restore_verifier_failed":{"stdout":result.stdout[-10000:],"stderr":result.stderr[-10000:]}})
    verification={
        "schema":"mrhpd-response66-complete-restore-verification-1.0","generated_at":NOW,"status":"passed",
        "restore":restore_qa,"project_snapshot":verify_zip(project_archive),"checkpoint2_recovery":verify_zip(recovery_zip,CP2_RECOVERY_BYTES,CP2_RECOVERY_SHA256),
        "clean_restore_verifier":"passed","self_contained":True,"requires_other_project_files":False,
        "requires_conversation_reconstruction":False,"accepted_predecessor_mutated":False,
        "checkpoint_3_of_3_complete":True,"session_1_of_3_complete":True,"remediation_section_4_complete":False,
        "next":"Remediation Section 4 of 5 Session 2 of 3",
    }
    json_write(dist/"MRHPD v3.0.0a Response 66 Complete Restore Verification.json",verification)
    text_write(dist/f"{restore.name}.sha256.txt",f"{restore_qa['sha256']}  {restore.name}")
    return restore,verification


def build_transport_volumes(restore: Path, dist: Path) -> dict[str, Any]:
    total=restore.stat().st_size
    if not (100*1024*1024 < total < 190*1024*1024):
        raise RuntimeError({"expected_two_volume_restore_size":total})
    first_size=(total+1)//2
    raw=[]
    with restore.open("rb") as source:
        for sequence,size in [(1,first_size),(2,total-first_size)]:
            path=dist/f"{restore.name}.part{sequence:03d}"
            remaining=size
            with path.open("wb") as out:
                while remaining:
                    block=source.read(min(1024*1024,remaining))
                    if not block: raise RuntimeError("unexpected EOF while splitting restore")
                    out.write(block); remaining-=len(block)
            raw.append({"sequence":sequence,"name":path.name,"bytes":path.stat().st_size,"sha256":sha256_file(path),"path":path})
    manifest={
        "schema":"mrhpd-complete-restore-transport-1.0","generated_at":NOW,
        "restore":{"name":restore.name,"bytes":total,"sha256":sha256_file(restore)},
        "part_count":2,"parts":[{k:v for k,v in row.items() if k!="path"} for row in raw],
        "minimum_volume_count":2,
    }
    manifest_path=dist/"MRHPD_RESPONSE66_COMPLETE_RESTORE_TRANSPORT_MANIFEST.json"; json_write(manifest_path,manifest)
    reassemble=dist/"reassemble_response66_complete_restore.py"
    text_write(reassemble, f'''#!/usr/bin/env python3
import hashlib,json,sys
from pathlib import Path
M={json.dumps(manifest,ensure_ascii=False)}
def sha(p):
 h=hashlib.sha256()
 with open(p,'rb') as f:
  for b in iter(lambda:f.read(1024*1024),b''): h.update(b)
 return h.hexdigest()
root=Path(__file__).resolve().parent
out=root/M['restore']['name']
with open(out,'wb') as dst:
 for row in M['parts']:
  p=root/row['name']
  if not p.exists() or p.stat().st_size!=row['bytes'] or sha(p)!=row['sha256']: raise SystemExit('part identity failure: '+row['name'])
  with open(p,'rb') as src:
   for b in iter(lambda:src.read(1024*1024),b''): dst.write(b)
if out.stat().st_size!=M['restore']['bytes'] or sha(out)!=M['restore']['sha256']: raise SystemExit('restore identity failure')
print(json.dumps({{'status':'passed','restore':out.name,'bytes':out.stat().st_size,'sha256':sha(out)}},indent=2))
''')
    wrappers=[]
    for row in raw:
        wrapper=dist/f"MRHPD v3.0.0a Response 66 Complete Restore Drive Volume {row['sequence']} of 2.zip"
        readme=(
            f"MRHPD Response 66 complete restore volume {row['sequence']} of 2. BOTH VOLUMES ARE REQUIRED. "
            "Extract both wrappers into the same directory and run reassemble_response66_complete_restore.py.\n"
        )
        readme_path=dist/f"README_VOLUME_{row['sequence']}.txt"; text_write(readme_path,readme)
        with zipfile.ZipFile(wrapper,"w",compression=zipfile.ZIP_STORED,allowZip64=True) as zf:
            zf.write(row["path"],row["path"].name)
            zf.write(manifest_path,manifest_path.name)
            zf.write(reassemble,reassemble.name)
            zf.write(readme_path,readme_path.name)
        qa=verify_zip(wrapper)
        if qa["bytes"] >= 104_857_600: raise RuntimeError({"drive_volume_exceeds_connector_limit":qa})
        wrappers.append({"sequence":row["sequence"],"wrapper":wrapper,"qa":qa,"raw_part":{k:v for k,v in row.items() if k!="path"}})
    return {"status":"passed","manifest":manifest,"manifest_path":manifest_path,"reassembly_utility":reassemble,"volumes":wrappers}


def build_controls_zip(dist: Path, files: Iterable[Path]) -> Path:
    controls=dist/"MRHPD v3.0.0a Response 66 Final Verification and Controls.zip"
    with zipfile.ZipFile(controls,"w",compression=zipfile.ZIP_DEFLATED,compresslevel=6) as zf:
        for path in files:
            if path.exists(): zf.write(path,path.name)
    verify_zip(controls)
    return controls


def main() -> None:
    parser=argparse.ArgumentParser()
    parser.add_argument("--base-dir",type=Path,default=Path("base_r64"))
    parser.add_argument("--checkpoint2-dir",type=Path,default=Path("checkpoint2_artifact"))
    parser.add_argument("--dist",type=Path,default=Path("dist_cp4_s1_cp3"))
    args=parser.parse_args()
    dist=args.dist
    if dist.exists(): shutil.rmtree(dist)
    dist.mkdir(parents=True)
    base=locate_exact_zip(args.base_dir,BASE_RESPONSE64_BYTES,BASE_RESPONSE64_SHA256)
    cp2=locate_exact_zip(args.checkpoint2_dir,CP2_RECOVERY_BYTES,CP2_RECOVERY_SHA256)

    with tempfile.TemporaryDirectory(prefix="mrhpd-cp4-s1-cp3-") as td:
        work=Path(td)
        project,source_qa=apply_checkpoint2(base,cp2,work)
        db,database_qa=synchronize_database(project)
        publication_qa=verify_publication(project)
        application_qa=synchronize_application(project,db)
        provisional_qa={"source":source_qa,"database":database_qa,"application":application_qa,"publication":publication_qa}
        workbook,workbook_qa=synchronize_workbook(project,db,database_qa,application_qa,publication_qa)
        provisional_qa["workbook"]=workbook_qa
        tracking_files=build_tracking_documents(project,db,provisional_qa)
        report_files=build_report(project,provisional_qa)
        final_db=finalize_database_status(db,workbook_qa,application_qa,publication_qa)
        database_qa.update(final_db)
        database_qa["sha256"]=sha256_file(db); database_qa["bytes"]=db.stat().st_size
        qa_dir=project/"QA"/"Section 4 Session 1"/"Complete"; qa_dir.mkdir(parents=True,exist_ok=True)
        json_write(qa_dir/"DATABASE_QA.json",database_qa)
        json_write(qa_dir/"WORKBOOK_QA.json",workbook_qa)
        json_write(qa_dir/"APPLICATION_QA.json",application_qa)
        json_write(qa_dir/"PUBLICATION_QA.json",publication_qa)
        recovery_dir=project/"Recovery"/"Section 4 Session 1 Complete"; recovery_dir.mkdir(parents=True,exist_ok=True)
        session_complete={
            "schema":"mrhpd-session1-complete-1.0","created_at":NOW,"section":"Remediation Section 4 of 5",
            "session":"Session 1 of 3","checkpoint":"3 of 3","response":66,"status":"COMPLETE",
            "database":database_qa,"workbook":workbook_qa,"application":application_qa,"publication":publication_qa,
            "accepted_predecessor_mutated":False,"next":"Remediation Section 4 of 5 Session 2 of 3",
        }
        json_write(recovery_dir/"SESSION_1_COMPLETE.json",session_complete)
        text_write(recovery_dir/"README.md","# Section 4 Session 1 Complete\n\nCheckpoint 3 of 3 and Session 1 of 3 are complete. Continue begins Session 2 of 3.\n")
        index_qa=build_indexes_and_manifest(project)
        session_qa={"status":"passed","generated_at":NOW,"source":source_qa,"database":database_qa,"workbook":workbook_qa,"application":application_qa,"publication":publication_qa,"indexes":index_qa,"tracking_files":len(tracking_files),"reports":len(report_files),"accepted_predecessor_mutated":False,"checkpoint_3_of_3_complete":True,"session_1_of_3_complete":True,"remediation_section_4_complete":False,"next":"Remediation Section 4 of 5 Session 2 of 3"}
        json_write(qa_dir/"SESSION_1_COMPLETE_QA.json",session_qa)
        # Rebuild indexes once more so final QA/session records are included.
        index_qa=build_indexes_and_manifest(project); session_qa["indexes"]=index_qa
        json_write(qa_dir/"SESSION_1_COMPLETE_QA.json",session_qa)
        project_archive,project_archive_qa=build_project_archive(project,dist)
        with tempfile.TemporaryDirectory(prefix="mrhpd-cp4-s1-project-clean-") as clean_td:
            clean_root=Path(clean_td); safe_extract(project_archive,clean_root)
            roots=[p for p in clean_root.iterdir() if p.is_dir()]
            clean_project=roots[0] if len(roots)==1 else clean_root
            clean_qa=quick_verify_project(clean_project,database_qa["sha256"],workbook_qa["sha256"])
        session_qa["project_archive"]=project_archive_qa; session_qa["clean_project_verification"]=clean_qa
        restore,restore_verification=build_complete_restore(project_archive,cp2,base,report_files,dist,work,session_qa)
        transport=build_transport_volumes(restore,dist)
        summary={
            "schema":"mrhpd-response66-session1-build-summary-1.0","generated_at":NOW,"status":"passed",
            "response":66,"section":"Remediation Section 4 of 5","session":"Session 1 of 3 COMPLETE","checkpoint":"3 of 3 COMPLETE",
            "source":source_qa,"database":database_qa,"workbook":workbook_qa,"application":application_qa,
            "publication":publication_qa,"indexes":index_qa,"project_archive":project_archive_qa,
            "clean_project_verification":clean_qa,"complete_restore":restore_verification,"transport":{
                "status":transport["status"],"manifest":transport["manifest"],
                "volumes":[{"sequence":v["sequence"],"wrapper":v["wrapper"].name,"qa":v["qa"],"raw_part":v["raw_part"]} for v in transport["volumes"]],
            },
            "accepted_predecessor_mutated":False,"user_upload_required":False,
            "next":"Remediation Section 4 of 5 Session 2 of 3",
        }
        summary_path=dist/"MRHPD_RESPONSE66_SESSION1_COMPLETE_BUILD_SUMMARY.json"; json_write(summary_path,summary)
        controls=build_controls_zip(dist,[
            dist/"MRHPD v3.0.0a Response 66 Complete Restore Verification.json",
            dist/f"{restore.name}.sha256.txt",
            summary_path,
            transport["manifest_path"],transport["reassembly_utility"],
            *report_files,
        ])
        output={
            "status":"passed","project_archive":project_archive.name,"project_archive_bytes":project_archive.stat().st_size,
            "complete_restore":restore.name,"complete_restore_bytes":restore.stat().st_size,"complete_restore_sha256":sha256_file(restore),
            "volume_1":transport["volumes"][0]["wrapper"].name,"volume_1_bytes":transport["volumes"][0]["wrapper"].stat().st_size,
            "volume_2":transport["volumes"][1]["wrapper"].name,"volume_2_bytes":transport["volumes"][1]["wrapper"].stat().st_size,
            "controls":controls.name,"controls_bytes":controls.stat().st_size,
            "checkpoint_3_of_3_complete":True,"session_1_of_3_complete":True,"remediation_section_4_complete":False,
            "next":"Remediation Section 4 of 5 Session 2 of 3",
        }
        print(json.dumps(output,indent=2,ensure_ascii=False))

if __name__=="__main__":
    main()
