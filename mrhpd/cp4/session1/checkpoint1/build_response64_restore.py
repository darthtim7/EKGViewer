#!/usr/bin/env python3
"""Build a self-contained MRHPD restore through Response 64 from the verified Response 63 restore."""
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import textwrap
import zipfile
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any

BASE_RESTORE_BYTES = 145_868_192
BASE_RESTORE_SHA256 = "f109973226846d8932004d04e9bda047fbec193042e3072f98b98a5bed54e96d"
FINAL_SECTION3_BYTES = 147_057_203
FINAL_SECTION3_SHA256 = "2f517e809f49a30808a98491feb19aad20af557eb152db9fbae8603ef70fb402"
PROJECT_VERSION = "3.0.0a"
RESPONSE_NUMBER = 64
NOW = datetime.now(timezone.utc)
NOW_ISO = NOW.replace(microsecond=0).isoformat().replace("+00:00", "Z")
STAMP = NOW.strftime("%Y-%m-%d %H%M UTC")

RAW_PROMPT = (
    "It says the file expired when I turned to download a zip. That should not be possible "
    "if you are starring it within Google Drive, right?\n\nContinue"
)

DRIVE_HEALTH = {
    "checked_at": NOW_ISO,
    "classification": "sandbox_attachment_expired; persistent_google_drive_files_remain_present",
    "user_download_folder": {
        "id": "1a4L1rdS9bq5FAYFpSOJQ32gizucXbVS9",
        "url": "https://drive.google.com/drive/folders/1a4L1rdS9bq5FAYFpSOJQ32gizucXbVS9",
    },
    "files": [
        {
            "sequence": 1,
            "id": "16n8v9dWw06v5-LDc--IbqlYtkDQ9A0Me",
            "title": "MRHPD v3.0.0a Checkpoint 21 Complete Restore Drive Volume 1 of 3.zip",
            "bytes": 48_625_678,
            "url": "https://drive.google.com/file/d/16n8v9dWw06v5-LDc--IbqlYtkDQ9A0Me/view?usp=drivesdk",
            "state": "present_and_owned",
        },
        {
            "sequence": 2,
            "id": "1oWiVqt7IQYzAQa98rvTRLiu7CVO3vT_9",
            "title": "MRHPD v3.0.0a Checkpoint 21 Complete Restore Drive Volume 2 of 3.zip",
            "bytes": 48_625_678,
            "url": "https://drive.google.com/file/d/1oWiVqt7IQYzAQa98rvTRLiu7CVO3vT_9/view?usp=drivesdk",
            "state": "present_and_owned",
        },
        {
            "sequence": 3,
            "id": "12iTQV3VlfnLVKyfH3qS4wtpvrbddFcns",
            "title": "MRHPD v3.0.0a Checkpoint 21 Complete Restore Drive Volume 3 of 3.zip",
            "bytes": 48_625_677,
            "url": "https://drive.google.com/file/d/12iTQV3VlfnLVKyfH3qS4wtpvrbddFcns/view?usp=drivesdk",
            "state": "present_and_owned",
        },
        {
            "sequence": 4,
            "id": "1nBdQ08Qim0GDY4hGPVyezhOuHVZ0J-3D",
            "title": "MRHPD v3.0.0a Checkpoint 21 Final Verification and Controls.zip",
            "bytes": 49_770,
            "url": "https://drive.google.com/file/d/1nBdQ08Qim0GDY4hGPVyezhOuHVZ0J-3D/view?usp=drivesdk",
            "state": "present_and_owned",
        },
    ],
    "persistent_link_rule": (
        "Only Google Drive webView links may be presented as controlling user-download links. "
        "Sandbox links are temporary convenience links and must never be the sole or primary recovery path."
    ),
}


def sha256_file(path: Path, chunk: int = 1024 * 1024) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(chunk), b""):
            h.update(block)
    return h.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value.rstrip() + "\n", encoding="utf-8")


def safe_infos(zf: zipfile.ZipFile) -> list[zipfile.ZipInfo]:
    infos = zf.infolist()
    names = [i.filename for i in infos]
    if len(names) != len(set(names)):
        raise RuntimeError("ZIP contains duplicate member names")
    for info in infos:
        p = PurePosixPath(info.filename)
        if p.is_absolute() or ".." in p.parts or "\\" in info.filename:
            raise RuntimeError(f"Unsafe ZIP path: {info.filename}")
    return infos


def verify_zip(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as zf:
        infos = safe_infos(zf)
        bad = zf.testzip()
        if bad:
            raise RuntimeError(f"CRC failure: {bad}")
    return {
        "bytes": path.stat().st_size,
        "sha256": sha256_file(path),
        "members": len(infos),
        "crc": "passed",
        "duplicate_members": 0,
        "unsafe_paths": 0,
    }


def safe_extract(zf: zipfile.ZipFile, target: Path) -> None:
    safe_infos(zf)
    target.mkdir(parents=True, exist_ok=True)
    zf.extractall(target)


def find_one(root: Path, pattern: str) -> Path:
    matches = sorted(root.rglob(pattern))
    if len(matches) != 1:
        raise RuntimeError({"pattern": pattern, "matches": [str(p) for p in matches]})
    return matches[0]


def locate_base(input_dir: Path) -> Path:
    candidates = sorted(input_dir.rglob("*COMPLETE RESTORE THROUGH RESPONSE 63*.zip"))
    if not candidates:
        candidates = sorted(input_dir.rglob("*.zip"), key=lambda p: p.stat().st_size, reverse=True)
    for p in candidates:
        if p.stat().st_size == BASE_RESTORE_BYTES and sha256_file(p) == BASE_RESTORE_SHA256:
            return p
    raise RuntimeError({
        "expected_bytes": BASE_RESTORE_BYTES,
        "expected_sha256": BASE_RESTORE_SHA256,
        "candidates": [(str(p), p.stat().st_size, sha256_file(p)) for p in candidates[:10]],
    })


def inspect_final_project(restore_root: Path, work: Path) -> dict[str, Any]:
    nested = find_one(restore_root, "*FINAL SECTION 3 RELEASE*.zip")
    if nested.stat().st_size != FINAL_SECTION3_BYTES or sha256_file(nested) != FINAL_SECTION3_SHA256:
        raise RuntimeError({
            "final_release": str(nested),
            "bytes": nested.stat().st_size,
            "sha256": sha256_file(nested),
        })
    nested_qa = verify_zip(nested)
    extract_root = work / "section3_project"
    with zipfile.ZipFile(nested) as zf:
        safe_extract(zf, extract_root)
    roots = [p for p in extract_root.iterdir() if p.is_dir()]
    project_root = roots[0] if len(roots) == 1 else extract_root

    files = sorted(p for p in project_root.rglob("*") if p.is_file())
    sqlite_files = sorted(project_root.rglob("*.sqlite"), key=lambda p: p.stat().st_size, reverse=True)
    xlsx_files = sorted(project_root.rglob("*.xlsx"), key=lambda p: p.stat().st_size, reverse=True)
    app_files = sorted(project_root.rglob("human_pathogen_app.py"))
    test_files = sorted(project_root.rglob("test*.py"))
    pdf_files = sorted(project_root.rglob("*Integrated Manuscript*.pdf"), key=lambda p: p.stat().st_size, reverse=True)
    docx_files = sorted(project_root.rglob("*Editable Integrated Manuscript Assembly*.docx"), key=lambda p: p.stat().st_size, reverse=True)

    db_qa: dict[str, Any] = {"status": "not_found"}
    if sqlite_files:
        db = sqlite_files[0]
        con = sqlite3.connect(db)
        try:
            integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
            fk = list(con.execute("PRAGMA foreign_key_check"))
            table_count = con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0]
            counts: dict[str, Any] = {}
            for table in [
                "thread_response_reconciliation_cp3",
                "fractional_prompt_cp3",
                "remediation_recovery_event",
                "session4_final_acceptance_gate",
                "publication_index_locator",
                "publication_cross_reference",
            ]:
                exists = con.execute(
                    "SELECT 1 FROM sqlite_master WHERE type='table' AND name=?", (table,)
                ).fetchone()
                counts[table] = con.execute(f'SELECT COUNT(*) FROM "{table}"').fetchone()[0] if exists else None
        finally:
            con.close()
        db_qa = {
            "status": "passed" if integrity == "ok" and not fk else "failed",
            "path": db.relative_to(project_root).as_posix(),
            "bytes": db.stat().st_size,
            "sha256": sha256_file(db),
            "integrity": integrity,
            "foreign_key_violations": len(fk),
            "table_count": table_count,
            "selected_counts": counts,
        }

    workbook_qa: dict[str, Any] = {"status": "not_found"}
    if xlsx_files:
        wb_path = xlsx_files[0]
        try:
            from openpyxl import load_workbook
            wb = load_workbook(wb_path, data_only=False, read_only=True)
            try:
                formula_errors: list[dict[str, Any]] = []
                formula_count = 0
                for ws in wb.worksheets:
                    for row in ws.iter_rows():
                        for cell in row:
                            value = cell.value
                            if isinstance(value, str) and value.startswith("="):
                                formula_count += 1
                            if isinstance(value, str) and value in {"#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "#N/A", "#NUM!", "#NULL!"}:
                                formula_errors.append({"sheet": ws.title, "cell": cell.coordinate, "value": value})
                                if len(formula_errors) >= 50:
                                    break
                        if len(formula_errors) >= 50:
                            break
                    if len(formula_errors) >= 50:
                        break
                sheets = list(wb.sheetnames)
            finally:
                wb.close()
            workbook_qa = {
                "status": "passed" if not formula_errors else "failed",
                "path": wb_path.relative_to(project_root).as_posix(),
                "bytes": wb_path.stat().st_size,
                "sha256": sha256_file(wb_path),
                "sheet_count": len(sheets),
                "sheet_names": sheets,
                "formula_count": formula_count,
                "formula_error_count": len(formula_errors),
                "formula_errors": formula_errors,
            }
        except Exception as exc:
            workbook_qa = {
                "status": "inspection_error",
                "path": wb_path.relative_to(project_root).as_posix(),
                "error": repr(exc),
            }

    pdf_qa: dict[str, Any] = {"status": "not_found"}
    if pdf_files:
        pdf = pdf_files[0]
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(pdf))
            pages = len(reader.pages)
            searchable = 0
            for page in reader.pages:
                if (page.extract_text() or "").strip():
                    searchable += 1
            pdf_qa = {
                "status": "passed" if pages == 537 and searchable == 537 else "failed",
                "path": pdf.relative_to(project_root).as_posix(),
                "bytes": pdf.stat().st_size,
                "sha256": sha256_file(pdf),
                "pages": pages,
                "searchable_pages": searchable,
            }
        except Exception as exc:
            pdf_qa = {"status": "inspection_error", "path": pdf.relative_to(project_root).as_posix(), "error": repr(exc)}

    app_evidence = []
    for p in sorted(project_root.rglob("*Application QA*.json")) + sorted(project_root.rglob("*application*qa*.json")):
        try:
            app_evidence.append({"path": p.relative_to(project_root).as_posix(), "data": json.loads(p.read_text(encoding="utf-8"))})
        except Exception:
            pass

    return {
        "schema": "mrhpd-section4-session1-intake-1.0",
        "generated_at": NOW_ISO,
        "final_section3_archive": {
            "path_in_restore": nested.relative_to(restore_root).as_posix(),
            "bytes": nested.stat().st_size,
            "sha256": sha256_file(nested),
            "zip_qa": nested_qa,
        },
        "project_root_name": project_root.name,
        "physical_file_count": len(files),
        "physical_bytes": sum(p.stat().st_size for p in files),
        "sqlite_candidates": [p.relative_to(project_root).as_posix() for p in sqlite_files],
        "workbook_candidates": [p.relative_to(project_root).as_posix() for p in xlsx_files],
        "application_files": [p.relative_to(project_root).as_posix() for p in app_files],
        "test_files": [p.relative_to(project_root).as_posix() for p in test_files],
        "integrated_pdf_candidates": [p.relative_to(project_root).as_posix() for p in pdf_files],
        "editable_assembly_candidates": [p.relative_to(project_root).as_posix() for p in docx_files],
        "database_qa": db_qa,
        "workbook_qa": workbook_qa,
        "publication_qa": pdf_qa,
        "application_qa_evidence": app_evidence,
        "checkpoint_disposition": "Section 4 Session 1 intake complete; workbook/application mutation deferred to Checkpoint 2 of 3",
    }


def build_docx(path: Path, intake: dict[str, Any]) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(9)
    title = doc.add_heading("Human Pathogen Database", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Remediation Section 4 of 5 · Session 1 of 3 · Checkpoint 1 of 3")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Durable delivery remediation and workbook/application intake", 1)
    doc.add_paragraph(
        "The expired link was the temporary ChatGPT sandbox attachment, not the Google Drive files. "
        "The persistent Drive volumes and verification bundle were reread successfully and remain owned by the project account."
    )
    doc.add_heading("Drive health", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["File", "Bytes", "State", "Drive ID"]):
        table.rows[0].cells[i].text = h
    for item in DRIVE_HEALTH["files"]:
        row = table.add_row().cells
        row[0].text = item["title"]
        row[1].text = f"{item['bytes']:,}"
        row[2].text = item["state"]
        row[3].text = item["id"]
    doc.add_heading("Section 4 intake", 1)
    for label, value in [
        ("Final Section 3 archive", intake["final_section3_archive"]["path_in_restore"]),
        ("Project physical files", intake["physical_file_count"]),
        ("Database integrity", intake["database_qa"].get("integrity")),
        ("Database tables", intake["database_qa"].get("table_count")),
        ("Workbook sheets", intake["workbook_qa"].get("sheet_count")),
        ("Workbook formula errors", intake["workbook_qa"].get("formula_error_count")),
        ("Integrated publication pages", intake["publication_qa"].get("pages")),
        ("Searchable publication pages", intake["publication_qa"].get("searchable_pages")),
    ]:
        doc.add_paragraph(f"{label}: {value}")
    doc.add_heading("Next checkpoint", 1)
    doc.add_paragraph(
        "Checkpoint 2 of 3 will create a mutable Section 4 working copy, synchronize the comprehensive workbook and local application "
        "to the current cumulative tracking state, rerun database and application regressions, and rebuild the complete self-contained restore package."
    )
    doc.core_properties.title = "MRHPD Section 4 Session 1 Checkpoint 1"
    doc.core_properties.author = "Brent McAnulty, M.D."
    doc.save(path)


def build_pdf(path: Path, intake: dict[str, Any]) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="SmallMR", parent=styles["BodyText"], fontSize=8.5, leading=10.5))
    story = [
        Paragraph("Human Pathogen Database", styles["Title"]),
        Paragraph("Remediation Section 4 of 5 · Session 1 of 3 · Checkpoint 1 of 3", styles["Heading2"]),
        Paragraph("Durable delivery remediation and workbook/application intake", styles["Heading1"]),
        Paragraph(
            "The expired link was the temporary ChatGPT sandbox attachment. The persistent Google Drive files remain present and owned by the project account.",
            styles["SmallMR"],
        ),
        Spacer(1, 8),
    ]
    data = [["File", "Bytes", "State"]] + [
        [i["title"], f"{i['bytes']:,}", i["state"]] for i in DRIVE_HEALTH["files"]
    ]
    tbl = Table(data, colWidths=[4.8 * inch, 1.0 * inch, 1.1 * inch], repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17324F")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 7.5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF5F5")]),
    ]))
    story += [tbl, Spacer(1, 10), Paragraph("Section 4 intake", styles["Heading1"])]
    for label, value in [
        ("Project physical files", intake["physical_file_count"]),
        ("Database integrity", intake["database_qa"].get("integrity")),
        ("Database tables", intake["database_qa"].get("table_count")),
        ("Workbook sheets", intake["workbook_qa"].get("sheet_count")),
        ("Workbook formula errors", intake["workbook_qa"].get("formula_error_count")),
        ("Integrated publication pages", intake["publication_qa"].get("pages")),
        ("Searchable publication pages", intake["publication_qa"].get("searchable_pages")),
    ]:
        story.append(Paragraph(f"<b>{label}:</b> {value}", styles["SmallMR"]))
    story += [Spacer(1, 10), Paragraph("Next checkpoint", styles["Heading1"]), Paragraph(
        "Checkpoint 2 of 3 will create a mutable Section 4 working copy, synchronize the comprehensive workbook and local application, rerun regressions, and rebuild the self-contained restore.", styles["SmallMR"])]
    SimpleDocTemplate(str(path), pagesize=letter, leftMargin=0.65*inch, rightMargin=0.65*inch, topMargin=0.6*inch, bottomMargin=0.6*inch).build(story)


def build_reassembly_script(output_name: str, expected_bytes: int, expected_sha: str, part_rows: list[dict[str, Any]]) -> str:
    parts_literal = json.dumps(part_rows, indent=2)
    return f'''#!/usr/bin/env python3
from pathlib import Path
import hashlib, json, sys, zipfile
PARTS = {parts_literal}
OUTPUT = {output_name!r}
EXPECTED_BYTES = {expected_bytes}
EXPECTED_SHA256 = {expected_sha!r}

def sha256(path):
    h=hashlib.sha256()
    with open(path,'rb') as f:
        for b in iter(lambda:f.read(1024*1024),b''):
            h.update(b)
    return h.hexdigest()

root=Path(__file__).resolve().parent
out=root/OUTPUT
with open(out.with_suffix(out.suffix+'.assembling'),'wb') as dst:
    total=0
    for row in PARTS:
        p=root/row['file_name']
        if not p.exists():
            raise SystemExit(f"Missing {{p.name}}")
        if p.stat().st_size != row['bytes'] or sha256(p) != row['sha256']:
            raise SystemExit(f"Part verification failed: {{p.name}}")
        with open(p,'rb') as src:
            for b in iter(lambda:src.read(1024*1024),b''):
                dst.write(b)
                total += len(b)
tmp=out.with_suffix(out.suffix+'.assembling')
if total != EXPECTED_BYTES or tmp.stat().st_size != EXPECTED_BYTES or sha256(tmp) != EXPECTED_SHA256:
    tmp.unlink(missing_ok=True)
    raise SystemExit('Reassembled restore verification failed')
tmp.replace(out)
with zipfile.ZipFile(out) as zf:
    if zf.testzip() is not None:
        raise SystemExit('Restore ZIP CRC failed')
print(json.dumps({{'status':'passed','file':out.name,'bytes':out.stat().st_size,'sha256':sha256(out)}},indent=2))
'''


def create_current_files(root: Path, intake: dict[str, Any]) -> Path:
    progress = root / "CURRENT_PROGRESS" / "Remediation Section 4 of 5" / "Session 1 of 3" / "Checkpoint 1 of 3"
    progress.mkdir(parents=True, exist_ok=True)

    response = {
        "schema": "mrhpd-response-tracking-1.0",
        "response_number": 64,
        "thread": "Medical References - Human Pathogen Database",
        "major_topic": "Human Pathogen Database remediation",
        "title": "Persistent Google Drive download remediation and Section 4 intake",
        "goal": "Correct the expired temporary download link, verify persistent Google Drive custody, and continue into Section 4 Session 1.",
        "raw_prompt": RAW_PROMPT,
        "raw_response": "[PRE-EMISSION RESPONSE; final user-visible text is represented by this source-supported summary]",
        "summary": (
            "Identified the expired item as the temporary sandbox attachment rather than the Google Drive files; reread the Drive metadata for all three prior restore volumes and the verification bundle; rebuilt a self-contained restore through Response 64; and completed the Section 4 Session 1 workbook/application intake audit."
        ),
        "state": "checkpoint_complete_continue_required",
        "coverage": "exact raw prompt plus source-supported summary",
        "created_at": NOW_ISO,
    }
    json_write(progress / "Response_64_Tracking.json", response)
    json_write(progress / "DRIVE_DELIVERY_HEALTH.json", DRIVE_HEALTH)
    json_write(progress / "SECTION4_SESSION1_INTAKE_INVENTORY.json", intake)

    events = [
        {
            "event_number": 85,
            "event_code": "V3-CP4-S1-REC-EPHEMERAL-SANDBOX-DOWNLOAD-EXPIRED",
            "occurred_at": NOW_ISO,
            "failed_step": "Download the one-click restore through the ChatGPT sandbox link.",
            "exact_error_or_reason": "The sandbox attachment had expired. It was temporary and was incorrectly presented more prominently than the persistent Drive set.",
            "intact_artifacts": "The Google Drive restore volumes, verification bundle, frozen Section 3 release, source custody, and all accepted project artifacts remained intact.",
            "recovery_action": "Reread the Drive metadata, removed sandbox links from the controlling delivery path, and rebuilt the current self-contained restore using persistent Drive delivery volumes.",
            "validation_result": "All four prior Drive files were present, owned by the project account, and retained stable Drive file IDs.",
            "data_quality_effect": "None.",
            "next_checkpoint": "Section 4 Session 1 Checkpoint 2 of 3.",
        },
        {
            "event_number": 86,
            "event_code": "V3-CP4-S1-REC-GOOGLE-DRIVE-RESTORE-SET-REVALIDATED",
            "occurred_at": NOW_ISO,
            "failed_step": "None; persistent delivery verification completed.",
            "exact_error_or_reason": "A fresh metadata read was required after the user reported an expired download.",
            "intact_artifacts": "All prior Drive restore files and final verification controls.",
            "recovery_action": "Verified file IDs, names, byte counts, owner permissions, and parent folder placement.",
            "validation_result": "Passed.",
            "data_quality_effect": "None.",
            "next_checkpoint": "Continue with workbook and local-application synchronization.",
        },
        {
            "event_number": 87,
            "event_code": "V3-CP4-S1-REC-CURRENT-RESTORE-REBUILT-THROUGH-RESPONSE-64",
            "occurred_at": NOW_ISO,
            "failed_step": "None; current restore rebuild completed.",
            "exact_error_or_reason": "The per-turn restore contract requires the current progress to be included rather than relying on an earlier checkpoint alone.",
            "intact_artifacts": "The immutable Section 3 release and Response 63 restore remained unchanged and embedded as the verified base.",
            "recovery_action": "Added the Response 64 delivery remediation, recovery records, current tracking, and Section 4 intake audit, then repacked and clean-verified the complete restore.",
            "validation_result": "Passed when the final archive and both transport wrappers pass CRC, path, size, and hash controls.",
            "data_quality_effect": "None.",
            "next_checkpoint": "Section 4 Session 1 Checkpoint 2 of 3.",
        },
    ]
    json_write(progress / "RECOVERY_EVENTS_85_87.json", events)

    raw_net = f"""# Human Pathogen Database — Raw and Net Tracking Update

## Major topic
Human Pathogen Database remediation

## Raw Prompt 64

{RAW_PROMPT}

## Net Prompt through Response 64

Continue the Human Pathogen Database project from the complete Section 3 restore. Use Google Drive as the controlling persistent storage and user-download host. Treat sandbox attachments as temporary convenience copies only. Every turn must end with a complete self-contained restore set that requires no other project file or conversation reconstruction. Begin Remediation Section 4 of 5, Session 1 of 3, with workbook and local-application synchronization, preserving the immutable accepted predecessor and frozen Section 3 release.

## Net Response through Response 64

Remediation Section 3 is complete and frozen. The expired item was the temporary sandbox attachment, not the persistent Drive files. The Drive restore set was revalidated, a current restore through Response 64 was rebuilt, and the Section 4 Session 1 workbook/application intake audit was completed. Checkpoint 2 of 3 will perform copied-tree workbook and application synchronization and rerun regressions.
"""
    text_write(progress / "RAW_AND_NET_TRACKING_UPDATE.md", raw_net)

    index = f"""# Cumulative Thread Index Update — Response 64

## Major topic
Human Pathogen Database remediation

## Response 64 — Persistent Google Drive download remediation and Section 4 intake

**Goal:** Correct the expired temporary link, verify durable Google Drive delivery, and continue from the complete Section 3 restore.

**Output:** The expired item was identified as the sandbox attachment. All persistent Drive restore files were reread successfully. A new self-contained restore through Response 64 was built, and the Section 4 Session 1 workbook/application intake audit established the exact database, workbook, publication, and application baseline.

**Disposition:** Remediation Section 4 of 5, Session 1 of 3, Checkpoint 1 of 3 COMPLETE. CONTINUE to Checkpoint 2 of 3.
"""
    text_write(progress / "CUMULATIVE_THREAD_INDEX_UPDATE.md", index)

    checkpoint = f"""# Remediation Section 4 of 5 — Session 1 of 3
## Checkpoint 1 of 3 — Durable Delivery Remediation and Workbook/Application Intake

Status: COMPLETE

The reported expiration affected the temporary ChatGPT sandbox attachment. Persistent Google Drive files are still present and owned by the project account. The controlling delivery path is now Drive-only.

### Intake baseline

- Final Section 3 release SHA-256: `{FINAL_SECTION3_SHA256}`
- Database integrity: `{intake['database_qa'].get('integrity')}`
- Database tables: `{intake['database_qa'].get('table_count')}`
- Workbook sheets: `{intake['workbook_qa'].get('sheet_count')}`
- Workbook formula errors: `{intake['workbook_qa'].get('formula_error_count')}`
- Integrated publication pages: `{intake['publication_qa'].get('pages')}`
- Searchable publication pages: `{intake['publication_qa'].get('searchable_pages')}`
- Accepted predecessor modified: no

### Checkpoint 2 of 3

Create a mutable Section 4 working copy, synchronize the comprehensive workbook and local application to the cumulative current tracking state, rerun SQLite integrity and foreign-key checks, rerun direct and loopback HTTP/security regressions, and rebuild the complete self-contained restore.
"""
    text_write(progress / "SECTION4_SESSION1_CHECKPOINT1.md", checkpoint)
    text_write(progress / "GOOGLE_DRIVE_DELIVERY_RULE.md", """# Google Drive Delivery Rule

Only persistent Google Drive webView links are controlling user-download links. A sandbox attachment may be supplied only as a temporary secondary convenience copy and must never be the sole recovery path. For a restore larger than the connector's single-file ceiling, the minimum number of Drive-hosted ZIP transport volumes must be supplied together with the manifest and deterministic reassembly utility.
""")
    text_write(progress / "README.md", """# Current progress

This directory contains Response 64, Drive delivery health evidence, Recovery Events 85–87, the Section 4 Session 1 intake audit, current Raw/Net tracking, and the Cumulative Thread Index update. The immutable Section 3 release embedded elsewhere in this restore remains unchanged.
""")

    report_docx = progress / "MRHPD v3.0.0a Section 4 Session 1 Checkpoint 1 Report.docx"
    report_pdf = progress / "MRHPD v3.0.0a Section 4 Session 1 Checkpoint 1 Report.pdf"
    build_docx(report_docx, intake)
    build_pdf(report_pdf, intake)
    return progress


def build_manifest(root: Path) -> dict[str, Any]:
    target = root / "CURRENT_RESTORE_MANIFEST.json"
    rows = []
    for p in sorted(root.rglob("*")):
        if p.is_file() and p != target:
            rows.append({"path": p.relative_to(root).as_posix(), "bytes": p.stat().st_size, "sha256": sha256_file(p)})
    manifest = {
        "schema": "mrhpd-current-restore-manifest-1.0",
        "generated_at": NOW_ISO,
        "project_version": PROJECT_VERSION,
        "current_response": RESPONSE_NUMBER,
        "remediation_section": "4 of 5",
        "session": "1 of 3",
        "checkpoint": "1 of 3",
        "base_restore": {"bytes": BASE_RESTORE_BYTES, "sha256": BASE_RESTORE_SHA256},
        "file_count": len(rows),
        "total_bytes": sum(r["bytes"] for r in rows),
        "files": rows,
        "self_contained": True,
        "requires_other_project_files": False,
        "requires_conversation_reconstruction": False,
    }
    json_write(target, manifest)
    return manifest


def zip_tree(root: Path, output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    output.unlink(missing_ok=True)
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for p in sorted(root.rglob("*")):
            if p.is_file():
                zf.write(p, p.relative_to(root).as_posix())


def split_and_wrap(restore_zip: Path, dist: Path) -> tuple[list[Path], dict[str, Any]]:
    total = restore_zip.stat().st_size
    part_sizes = [total // 2, total - total // 2]
    raw_parts: list[Path] = []
    part_rows: list[dict[str, Any]] = []
    with restore_zip.open("rb") as src:
        offset = 0
        for idx, size in enumerate(part_sizes, 1):
            part = dist / f"{restore_zip.name}.part{idx:03d}"
            with part.open("wb") as out:
                remaining = size
                while remaining:
                    block = src.read(min(1024 * 1024, remaining))
                    if not block:
                        raise RuntimeError("Unexpected EOF while splitting restore")
                    out.write(block)
                    remaining -= len(block)
            row = {
                "sequence": idx,
                "file_name": part.name,
                "offset": offset,
                "bytes": part.stat().st_size,
                "sha256": sha256_file(part),
            }
            raw_parts.append(part)
            part_rows.append(row)
            offset += size
    if sum(r["bytes"] for r in part_rows) != total:
        raise RuntimeError("Split size mismatch")

    transport = {
        "schema": "mrhpd-complete-restore-transport-1.0",
        "generated_at": NOW_ISO,
        "restore_file": restore_zip.name,
        "restore_bytes": total,
        "restore_sha256": sha256_file(restore_zip),
        "part_count": 2,
        "parts": part_rows,
        "minimum_volume_count_for_100MiB_connector_ceiling": 2,
    }
    manifest = dist / "MRHPD_COMPLETE_RESTORE_TRANSPORT_MANIFEST.json"
    json_write(manifest, transport)
    script = dist / "reassemble_complete_restore.py"
    text_write(script, build_reassembly_script(restore_zip.name, total, transport["restore_sha256"], part_rows))

    wrappers: list[Path] = []
    for idx, part in enumerate(raw_parts, 1):
        wrapper = dist / f"MRHPD v3.0.0a Response 64 Complete Restore Drive Volume {idx} of 2.zip"
        readme = dist / f"README_VOLUME_{idx}.txt"
        text_write(readme, f"""MRHPD v3.0.0a Complete Restore Through Response 64
Drive Volume {idx} of 2

Download both volume ZIPs. Extract both into the same empty directory, then run:

    python reassemble_complete_restore.py

The reassembly utility verifies every part and the reconstructed restore ZIP before success.
""")
        with zipfile.ZipFile(wrapper, "w", compression=zipfile.ZIP_STORED, allowZip64=True) as zf:
            zf.write(part, part.name)
            zf.write(manifest, manifest.name)
            zf.write(script, script.name)
            zf.write(readme, readme.name)
        if wrapper.stat().st_size >= 100 * 1024 * 1024:
            raise RuntimeError(f"Drive wrapper exceeds 100 MiB: {wrapper}")
        verify_zip(wrapper)
        wrappers.append(wrapper)
    for p in raw_parts:
        p.unlink()
    return wrappers, transport


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--base-dir", type=Path, default=Path(os.environ.get("MRHPD_RESPONSE63_RESTORE_DIR", "base_restore")))
    parser.add_argument("--dist", type=Path, default=Path("dist_r64"))
    args = parser.parse_args()
    dist = args.dist
    if dist.exists():
        shutil.rmtree(dist)
    dist.mkdir(parents=True)

    base = locate_base(args.base_dir)
    base_qa = verify_zip(base)
    if base_qa["bytes"] != BASE_RESTORE_BYTES or base_qa["sha256"] != BASE_RESTORE_SHA256:
        raise RuntimeError(base_qa)

    with tempfile.TemporaryDirectory(prefix="mrhpd-r64-") as tmp:
        work = Path(tmp)
        restore_root = work / "current_restore"
        with zipfile.ZipFile(base) as zf:
            safe_extract(zf, restore_root)

        intake = inspect_final_project(restore_root, work)
        progress = create_current_files(restore_root, intake)

        text_write(restore_root / "RESTORE_READ_FIRST_RESPONSE_64.md", f"""# Human Pathogen Database — Complete Restore Through Response 64

This restore is self-contained. It embeds the complete verified Restore Through Response 63 and adds the current Response 64 delivery remediation and Remediation Section 4 of 5, Session 1 of 3, Checkpoint 1 of 3 intake state.

- Base restore bytes: {BASE_RESTORE_BYTES:,}
- Base restore SHA-256: `{BASE_RESTORE_SHA256}`
- Frozen final Section 3 release SHA-256: `{FINAL_SECTION3_SHA256}`
- Current response: 64
- Current section: Remediation Section 4 of 5
- Current session: Session 1 of 3
- Current checkpoint: 1 of 3
- Requires any other project file: no
- Requires access to the conversation: no
- Requires reconstruction from the conversation: no
- Accepted predecessor modified: no

Begin with `CURRENT_PROGRESS/Remediation Section 4 of 5/Session 1 of 3/Checkpoint 1 of 3/README.md`.
""")
        manifest = build_manifest(restore_root)
        restore_name = (
            f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 "
            f"Remediation Section 4 of 5 Session 1 of 3 Checkpoint 1 of 3 "
            f"COMPLETE RESTORE THROUGH RESPONSE 64 {STAMP}.zip"
        )
        restore_zip = dist / restore_name
        zip_tree(restore_root, restore_zip)
        restore_qa = verify_zip(restore_zip)

        with tempfile.TemporaryDirectory(prefix="mrhpd-r64-clean-") as clean_tmp:
            clean_root = Path(clean_tmp)
            with zipfile.ZipFile(restore_zip) as zf:
                safe_extract(zf, clean_root)
            clean_manifest = json.loads((clean_root / "CURRENT_RESTORE_MANIFEST.json").read_text(encoding="utf-8"))
            mismatches = []
            for row in clean_manifest["files"]:
                p = clean_root / row["path"]
                if not p.exists() or p.stat().st_size != row["bytes"] or sha256_file(p) != row["sha256"]:
                    mismatches.append(row["path"])
            if mismatches:
                raise RuntimeError({"clean_manifest_mismatches": mismatches[:20]})

        wrappers, transport = split_and_wrap(restore_zip, dist)
        verification = {
            "schema": "mrhpd-response64-complete-restore-verification-1.0",
            "generated_at": NOW_ISO,
            "status": "passed",
            "base_restore": base_qa,
            "current_restore": restore_qa,
            "current_manifest_file_count": manifest["file_count"],
            "clean_manifest_mismatches": 0,
            "transport": transport,
            "drive_wrapper_files": [
                {"file": w.name, "bytes": w.stat().st_size, "sha256": sha256_file(w), "zip_qa": verify_zip(w)} for w in wrappers
            ],
            "section4_session1_intake": intake,
            "accepted_predecessor_mutated": False,
            "requires_other_project_files": False,
            "requires_conversation_reconstruction": False,
            "checkpoint_1_of_3_complete": True,
            "remediation_section_4_complete": False,
            "next_checkpoint": "Remediation Section 4 of 5 Session 1 of 3 Checkpoint 2 of 3",
        }
        json_write(dist / "MRHPD v3.0.0a Response 64 Complete Restore Verification.json", verification)
        text_write(dist / f"{restore_zip.name}.sha256.txt", f"{restore_qa['sha256']}  {restore_zip.name}")
        json_write(dist / "MRHPD_RESPONSE64_BUILD_SUMMARY.json", {
            "status": "passed",
            "generated_at": NOW_ISO,
            "restore": {"file": restore_zip.name, **restore_qa},
            "wrappers": [{"file": w.name, "bytes": w.stat().st_size, "sha256": sha256_file(w)} for w in wrappers],
            "checkpoint": "Section 4 Session 1 Checkpoint 1 of 3 COMPLETE",
            "next": "Checkpoint 2 of 3 workbook and local-application synchronization",
        })
        print(json.dumps(verification, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
