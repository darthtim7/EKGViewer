#!/usr/bin/env python3
"""Build MRHPD Section 4 Session 2 Checkpoint 2 recovery through Response 68.

This builder reconstructs the exact complete restore through Response 66,
applies the independently verified Response 67 Checkpoint 1 recovery, copies
that current state, and then performs detailed field-level, query-level,
source-governance, cross-artifact drift, workbook, application, index,
manifest, tracking, report, and recovery synchronization.

The Response 66 restore, frozen Section 3 release, accepted predecessor,
537-page publication, editable assembly, and main application source are never
modified in place.
"""
from __future__ import annotations

import argparse
import csv
import hashlib
import importlib.util
import json
import os
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import textwrap
import zipfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any, Iterable

CP1_PATH = Path(__file__).resolve().parents[1] / "checkpoint1" / "build_checkpoint1_recovery.py"
spec = importlib.util.spec_from_file_location("mrhpd_cp4_s2_cp1", CP1_PATH)
if spec is None or spec.loader is None:
    raise RuntimeError(f"Unable to load Checkpoint 1 builder: {CP1_PATH}")
cp1 = importlib.util.module_from_spec(spec)
spec.loader.exec_module(cp1)

PROJECT_VERSION = "3.0.0a"
RESPONSE_NUMBER = 68
NOW_DT = datetime.now(timezone.utc)
NOW = NOW_DT.replace(microsecond=0).isoformat().replace("+00:00", "Z")
STAMP = NOW_DT.strftime("%Y-%m-%d %H%M UTC")
RAW_PROMPT = "Continue"
BASE_RESTORE_BYTES = cp1.BASE_RESTORE_BYTES
BASE_RESTORE_SHA256 = cp1.BASE_RESTORE_SHA256
BASE_PROJECT_BYTES = cp1.BASE_PROJECT_BYTES
BASE_PROJECT_SHA256 = cp1.BASE_PROJECT_SHA256
PUBLICATION_SHA256 = cp1.PUBLICATION_SHA256
EDITABLE_ASSEMBLY_SHA256 = cp1.EDITABLE_ASSEMBLY_SHA256
APPLICATION_SHA256 = cp1.APPLICATION_SHA256

RESPONSE68 = {
    "response_key": "R68",
    "response_number": 68,
    "response_label": "68",
    "branch_id": "mainline",
    "canonical_current": 1,
    "response_date": NOW,
    "major_topic": "Human Pathogen Database remediation",
    "title": "Section 4 Session 2 detailed field, query, evidence, and drift-reconciliation checkpoint",
    "goal": (
        "Continue from the verified Response 67 checkpoint; expand capability parity into detailed field- and query-level coverage; "
        "reconcile source-governance evidence across SQLite, workbook, application, reports, and indexes; resolve cross-artifact drift; "
        "extend current-database search and disambiguation audits; and emit a cleanly applicable checkpoint-recovery package."
    ),
    "raw_prompt": RAW_PROMPT,
    "raw_response": "[PRE-EMISSION RESPONSE; final user-visible response is represented by the source-supported summary]",
    "summary": (
        "Restored the complete Response 66 baseline, clean-applied the verified Response 67 recovery, created a copied Response 68 "
        "working tree, audited every logical SQLite table and field, executed controlled search/disambiguation and semantic queries, "
        "reconciled evidence-source governance, resolved expected versus prohibited drift, synchronized the workbook and read-only "
        "application audit surfaces, reran publication/index/manifest/package gates, and emitted clean-verified Checkpoint 2 recovery data."
    ),
    "state": "checkpoint_complete_continue_required",
    "coverage": "exact raw prompt plus source-supported response summary",
    "fidelity_classification": "source_verified_prompt_and_summary",
    "source_id": "CURRENT-CONVERSATION-R68",
    "source_path": "Current conversation turn and Section 4 Session 2 Checkpoint 2 recovery package",
    "notes": "Checkpoint 2 of 3 is complete. Checkpoint 3 ends Session 2 and emits a complete self-contained restore.",
}

NET_PROMPT = (
    "Continue the Human Pathogen Database from the newest verified Section 4 Session 2 checkpoint. Preserve Google Drive as the "
    "controlling storage and download host. Use a copied project tree; add Response 68 and recovery history; expand cross-artifact "
    "capability governance into detailed field-level and query-level coverage; reconcile evidence and source-governance metadata; "
    "detect and resolve drift against the frozen Checkpoint 1 baseline; extend current-database application, search, alias, historical-name, "
    "and disambiguation regressions; preserve the accepted predecessor, frozen Section 3 release, 537-page publication, editable assembly, "
    "and main application source; rebuild the workbook, tracking, reports, Source Index, Bit Index, manifests, checksums, and QA; and emit "
    "complete intermediate checkpoint-recovery data tied directly to the exact Response 66 complete restore."
)

NET_RESPONSE = (
    "Section 4 Session 2 Checkpoint 2 is complete through Response 68. The exact Response 66 restore and Response 67 recovery were "
    "verified before work began. The copied current project now includes comprehensive logical-table and field coverage, controlled term "
    "and semantic query regressions, source-governance audits, expected/prohibited drift resolution, expanded workbook controls, and a "
    "read-only current-database application audit. The 537-page publication, editable assembly, and main application source remain "
    "byte-identical. The recovery package applies directly to the Response 66 restore and clean-verifies the current state. Continue "
    "proceeds to Checkpoint 3 of 3, which emits the Session 2 complete restore."
)

SEARCH_TERMS = [
    ("streptococcus", "Streptococcus"),
    ("strep", "Strep"),
    ("strep_bovis", "Strep bovis"),
    ("bovis", "bovis"),
    ("sbsec", "SBSEC"),
    ("gas", "GAS"),
    ("gbs", "GBS"),
    ("pneumococcus", "pneumococcus"),
    ("sanguis", "sanguis"),
    ("gallolyticus", "gallolyticus"),
    ("anginosus", "anginosus"),
    ("crypto", "Crypto"),
    ("asymptomatic_bacteriuria", "asymptomatic bacteriuria"),
    ("endocarditis", "endocarditis"),
    ("meningitis", "meningitis"),
    ("antibiogram", "antibiogram"),
]


def qident(value: str) -> str:
    return '"' + value.replace('"', '""') + '"'


def json_write(path: Path, value: Any) -> None:
    cp1.json_write(path, value)


def text_write(path: Path, value: str) -> None:
    cp1.text_write(path, value)


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            normalized = {}
            for key in fields:
                value = row.get(key)
                if isinstance(value, (list, dict)):
                    value = json.dumps(value, ensure_ascii=False)
                normalized[key] = value
            writer.writerow(normalized)


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def find_one(root: Path, pattern: str) -> Path:
    return cp1.find_one(root, pattern)


def verify_zip(path: Path, expected_bytes: int | None = None, expected_sha256: str | None = None) -> dict[str, Any]:
    return cp1.verify_zip(path, expected_bytes, expected_sha256)


def safe_extract(path: Path, target: Path) -> None:
    cp1.safe_extract(path, target)


def sha256_file(path: Path) -> str:
    return cp1.sha256_file(path)


def reconstruct_response67(base_restore: Path, response67_input: Path, work: Path) -> tuple[Path, dict[str, Any]]:
    wrappers = sorted(response67_input.rglob("*Response 67 Checkpoint 1 Recovery Delivery.zip"))
    if len(wrappers) != 1:
        raise RuntimeError({"response67_delivery_wrappers": [str(path) for path in wrappers]})
    wrapper = wrappers[0]
    wrapper_qa = verify_zip(wrapper)
    wrapper_root = work / "response67_delivery"
    safe_extract(wrapper, wrapper_root)
    recovery_matches = sorted(wrapper_root.rglob("*RECOVERY DATA THROUGH RESPONSE 67*.zip"))
    if len(recovery_matches) != 1:
        raise RuntimeError({"response67_recovery_matches": [str(path) for path in recovery_matches]})
    recovery = recovery_matches[0]
    recovery_qa = verify_zip(recovery)
    recovery_root = work / "response67_recovery"
    safe_extract(recovery, recovery_root)
    utility = recovery_root / "TOOLS" / "apply_checkpoint_recovery.py"
    if not utility.exists():
        raise FileNotFoundError(utility)
    output = work / "response67_project"
    result = subprocess.run(
        [sys.executable, str(utility), "--base-response66-restore", str(base_restore), "--output-dir", str(output)],
        text=True,
        capture_output=True,
        timeout=2400,
    )
    if result.returncode != 0:
        raise RuntimeError({"response67_clean_apply_failed": {"stdout": result.stdout[-20000:], "stderr": result.stderr[-20000:]}})
    return output, {
        "status": "passed",
        "delivery_wrapper": wrapper_qa,
        "checkpoint_recovery": recovery_qa,
        "apply_stdout": result.stdout[-12000:],
        "response67_project": str(output),
    }


def logical_tables(con: sqlite3.Connection) -> list[tuple[str, str]]:
    rows = list(con.execute("SELECT name, COALESCE(sql,'') FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name"))
    names = {row[0] for row in rows}
    shadow_suffixes = ("_data", "_idx", "_content", "_docsize", "_config")
    logical: list[tuple[str, str]] = []
    for name, sql in rows:
        if any(name.endswith(suffix) and name[: -len(suffix)] in names for suffix in shadow_suffixes):
            continue
        logical.append((name, sql or ""))
    return logical


def field_level_coverage(con: sqlite3.Connection) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    table_failures: list[dict[str, Any]] = []
    for table, sql in logical_tables(con):
        try:
            row_count = int(con.execute(f"SELECT COUNT(*) FROM {qident(table)}").fetchone()[0])
        except Exception as exc:
            table_failures.append({"table": table, "error": repr(exc)})
            continue
        info = list(con.execute(f"PRAGMA table_info({qident(table)})"))
        for cid, field, declared_type, not_null, default_value, primary_key in info:
            field_type = (declared_type or "").upper()
            try:
                null_count = int(con.execute(f"SELECT COUNT(*) FROM {qident(table)} WHERE {qident(field)} IS NULL").fetchone()[0])
                blank_count = 0
                if any(token in field_type for token in ("CHAR", "TEXT", "CLOB")) or field_type == "":
                    blank_count = int(
                        con.execute(
                            f"SELECT COUNT(*) FROM {qident(table)} WHERE {qident(field)} IS NOT NULL AND TRIM(CAST({qident(field)} AS TEXT))=''"
                        ).fetchone()[0]
                    )
                distinct_count: int | None
                if "BLOB" in field_type:
                    distinct_count = None
                else:
                    distinct_count = int(con.execute(f"SELECT COUNT(DISTINCT {qident(field)}) FROM {qident(table)}").fetchone()[0])
                critical_name = bool(re.search(r"(^id$|_id$|_key$|_code$|^name$|_name$|^title$|^status$|_status$)", field, re.I))
                required = bool(not_null or primary_key)
                failed = (required and null_count > 0) or (required and critical_name and blank_count > 0)
                status = "failed" if failed else "passed"
                notes = []
                if sql.upper().startswith("CREATE VIRTUAL TABLE"):
                    notes.append("virtual table")
                if required:
                    notes.append("required field")
                elif null_count or blank_count:
                    notes.append("nullable gaps documented")
                rows.append(
                    {
                        "table_name": table,
                        "field_name": field,
                        "declared_type": declared_type or "",
                        "row_count": row_count,
                        "not_null_declared": int(bool(not_null)),
                        "primary_key": int(bool(primary_key)),
                        "null_count": null_count,
                        "blank_count": blank_count,
                        "distinct_count": distinct_count,
                        "status": status,
                        "notes": "; ".join(notes),
                    }
                )
            except Exception as exc:
                rows.append(
                    {
                        "table_name": table,
                        "field_name": field,
                        "declared_type": declared_type or "",
                        "row_count": row_count,
                        "not_null_declared": int(bool(not_null)),
                        "primary_key": int(bool(primary_key)),
                        "null_count": None,
                        "blank_count": None,
                        "distinct_count": None,
                        "status": "failed",
                        "notes": f"audit error: {exc!r}",
                    }
                )
    failed_rows = [row for row in rows if row["status"] != "passed"]
    summary = {
        "status": "passed" if not failed_rows and not table_failures else "failed",
        "logical_tables": len({row["table_name"] for row in rows}),
        "audited_fields": len(rows),
        "passed_fields": sum(1 for row in rows if row["status"] == "passed"),
        "failed_fields": len(failed_rows),
        "table_failures": table_failures,
        "required_fields": sum(1 for row in rows if row["not_null_declared"] or row["primary_key"]),
        "documented_nullable_gaps": sum(1 for row in rows if (row.get("null_count") or 0) + (row.get("blank_count") or 0) > 0),
    }
    if summary["status"] != "passed":
        raise RuntimeError({"field_level_coverage": summary, "failed_samples": failed_rows[:30]})
    return rows, summary


def text_columns(con: sqlite3.Connection, table: str) -> list[str]:
    columns: list[str] = []
    for row in con.execute(f"PRAGMA table_info({qident(table)})"):
        name = row[1]
        declared = (row[2] or "").upper()
        if any(token in declared for token in ("CHAR", "TEXT", "CLOB")) or declared == "":
            columns.append(name)
    return columns


def search_term(con: sqlite3.Connection, term: str) -> tuple[int, list[dict[str, Any]]]:
    candidate_tokens = (
        "search", "resolver", "alias", "taxonomy", "organism", "syndrome", "clinical", "disease", "manifestation",
        "treatment", "evidence", "source", "publication", "glossary", "master_category", "reconciliation",
    )
    total = 0
    samples: list[dict[str, Any]] = []
    for table, _sql in logical_tables(con):
        if not any(token in table.lower() for token in candidate_tokens):
            continue
        columns = text_columns(con, table)[:14]
        if not columns:
            continue
        where = " OR ".join(f"LOWER(CAST({qident(column)} AS TEXT)) LIKE LOWER(?)" for column in columns)
        params = [f"%{term}%"] * len(columns)
        try:
            count = int(con.execute(f"SELECT COUNT(*) FROM {qident(table)} WHERE {where}", params).fetchone()[0])
        except Exception:
            continue
        if count <= 0:
            continue
        total += count
        if len(samples) < 12:
            select_cols = columns[:6]
            query = f"SELECT {', '.join(qident(column) for column in select_cols)} FROM {qident(table)} WHERE {where} LIMIT 2"
            try:
                values = con.execute(query, params).fetchall()
                for value in values:
                    samples.append({"table": table, "values": [str(item)[:250] if item is not None else None for item in value]})
            except Exception:
                samples.append({"table": table, "values": [f"{count} matching rows"]})
        if total >= 500:
            break
    return total, samples


def semantic_control(con: sqlite3.Connection, key: str) -> tuple[int, int, dict[str, Any]]:
    if key == "locator_count":
        value = int(con.execute("SELECT COUNT(*) FROM publication_index_locator").fetchone()[0])
        return value, 4011, {"operator": "eq"}
    if key == "cross_reference_count":
        value = int(con.execute("SELECT COUNT(*) FROM publication_cross_reference").fetchone()[0])
        return value, 12, {"operator": "gte"}
    if key == "page_map_count":
        table = "final_publication_page_map" if cp1.table_exists(con, "final_publication_page_map") else "publication_page_map"
        value = int(con.execute(f"SELECT COUNT(*) FROM {qident(table)}").fetchone()[0])
        return value, 10, {"operator": "eq", "table": table}
    if key == "capability_pass_count":
        value = int(con.execute("SELECT COUNT(*) FROM section4_session2_capability WHERE status='passed'").fetchone()[0])
        total = int(con.execute("SELECT COUNT(*) FROM section4_session2_capability").fetchone()[0])
        return value, total, {"operator": "eq", "total": total}
    if key == "response68":
        value = int(con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R68'").fetchone()[0])
        return value, 1, {"operator": "eq"}
    raise KeyError(key)


def query_level_coverage(con: sqlite3.Connection) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for key, term in SEARCH_TERMS:
        count, samples = search_term(con, term)
        rows.append(
            {
                "query_key": key,
                "query_text": term,
                "query_type": "cross_table_text_search",
                "result_count": count,
                "expected_minimum": 1,
                "status": "passed" if count >= 1 else "failed",
                "sample_json": samples,
                "notes": "Case-insensitive search across governed resolver, alias, taxonomy, clinical, evidence, publication, and tracking tables.",
            }
        )
    for key in ("locator_count", "cross_reference_count", "page_map_count", "capability_pass_count", "response68"):
        actual, expected, details = semantic_control(con, key)
        operator = details["operator"]
        passed = actual == expected if operator == "eq" else actual >= expected
        rows.append(
            {
                "query_key": key,
                "query_text": key,
                "query_type": "semantic_control",
                "result_count": actual,
                "expected_minimum": expected,
                "status": "passed" if passed else "failed",
                "sample_json": details,
                "notes": f"Operator: {operator}",
            }
        )
    failed = [row for row in rows if row["status"] != "passed"]
    summary = {
        "status": "passed" if not failed else "failed",
        "queries": len(rows),
        "passed": len(rows) - len(failed),
        "failed": len(failed),
        "search_queries": len(SEARCH_TERMS),
        "semantic_controls": len(rows) - len(SEARCH_TERMS),
        "failed_queries": [row["query_key"] for row in failed],
    }
    if failed:
        raise RuntimeError({"query_level_coverage": summary, "failed_rows": failed})
    return rows, summary


def first_column(columns: list[str], candidates: Iterable[str], contains: Iterable[str] = ()) -> str | None:
    lower = {column.lower(): column for column in columns}
    for candidate in candidates:
        if candidate.lower() in lower:
            return lower[candidate.lower()]
    for column in columns:
        if any(token.lower() in column.lower() for token in contains):
            return column
    return None


def source_governance_audit(con: sqlite3.Connection) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for table, _sql in logical_tables(con):
        lower = table.lower()
        if not ("evidence" in lower or "source" in lower):
            continue
        if any(token in lower for token in ("source_index", "artifact_index", "bit_index")):
            continue
        try:
            row_count = int(con.execute(f"SELECT COUNT(*) FROM {qident(table)}").fetchone()[0])
        except Exception as exc:
            rows.append({"audit_key": table, "source_table": table, "row_count": None, "status": "failed", "details_json": {"error": repr(exc)}})
            continue
        columns = cp1.table_columns(con, table)
        url_col = first_column(columns, ("url", "source_url", "plaintext_url"), ("url",))
        year_col = first_column(columns, ("year", "publication_year", "source_year"), ("year",))
        title_col = first_column(columns, ("title", "source_title", "name"), ("title",))
        authority_col = first_column(columns, ("authority", "source_type", "source_family", "publisher", "organization"), ("authority", "publisher", "organization", "family", "type"))
        missing_url = None
        nonempty_url = None
        latest_year = None
        missing_title = None
        missing_authority = None
        try:
            if url_col:
                missing_url = int(con.execute(f"SELECT COUNT(*) FROM {qident(table)} WHERE {qident(url_col)} IS NULL OR TRIM(CAST({qident(url_col)} AS TEXT))='' ").fetchone()[0])
                nonempty_url = row_count - missing_url
            if title_col:
                missing_title = int(con.execute(f"SELECT COUNT(*) FROM {qident(table)} WHERE {qident(title_col)} IS NULL OR TRIM(CAST({qident(title_col)} AS TEXT))='' ").fetchone()[0])
            if authority_col:
                missing_authority = int(con.execute(f"SELECT COUNT(*) FROM {qident(table)} WHERE {qident(authority_col)} IS NULL OR TRIM(CAST({qident(authority_col)} AS TEXT))='' ").fetchone()[0])
            if year_col:
                values = [value[0] for value in con.execute(f"SELECT {qident(year_col)} FROM {qident(table)} WHERE {qident(year_col)} IS NOT NULL")]
                years = []
                for value in values:
                    match = re.search(r"(19|20)\d{2}", str(value))
                    if match:
                        years.append(int(match.group(0)))
                latest_year = max(years) if years else None
        except Exception as exc:
            rows.append({"audit_key": table, "source_table": table, "row_count": row_count, "status": "failed", "details_json": {"error": repr(exc)}})
            continue
        linkage_columns = [column for column in columns if column.lower().endswith("_id") or "source" in column.lower() or "evidence" in column.lower()]
        schema_supported = bool(title_col or url_col or authority_col or linkage_columns)
        status = "passed" if schema_supported else "failed"
        rows.append(
            {
                "audit_key": table,
                "source_table": table,
                "row_count": row_count,
                "url_column": url_col,
                "missing_url_count": missing_url,
                "nonempty_url_count": nonempty_url,
                "year_column": year_col,
                "latest_year": latest_year,
                "title_column": title_col,
                "missing_title_count": missing_title,
                "authority_column": authority_col,
                "missing_authority_count": missing_authority,
                "status": status,
                "details_json": {"columns": columns, "linkage_columns": linkage_columns},
            }
        )
    failed = [row for row in rows if row["status"] != "passed"]
    summary = {
        "status": "passed" if rows and not failed else "failed",
        "audited_tables": len(rows),
        "passed_tables": len(rows) - len(failed),
        "failed_tables": len(failed),
        "tables_with_url": sum(1 for row in rows if row.get("url_column")),
        "total_nonempty_urls": sum(int(row.get("nonempty_url_count") or 0) for row in rows),
        "latest_year": max((row.get("latest_year") or 0 for row in rows), default=0) or None,
        "failed_table_names": [row["source_table"] for row in failed],
    }
    if summary["status"] != "passed":
        raise RuntimeError({"source_governance": summary, "failed_rows": failed})
    return rows, summary


def baseline_artifacts(project: Path) -> dict[str, Path]:
    db = cp1.find_canonical_database(project)
    workbook, _ = cp1.find_workbook(project)
    app = cp1.find_by_hash(project, "human_pathogen_app.py", APPLICATION_SHA256)
    publication = cp1.find_by_hash(project, "*Integrated Manuscript*.pdf", PUBLICATION_SHA256)
    editable = cp1.find_by_hash(project, "*Editable Integrated Manuscript Assembly*.docx", EDITABLE_ASSEMBLY_SHA256)
    return {"database": db, "workbook": workbook, "application": app, "publication": publication, "editable_assembly": editable}


def critical_drift_rows(baseline: dict[str, Path], current: dict[str, Path], baseline_root: Path, current_root: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for key in ("application", "publication", "editable_assembly", "database", "workbook"):
        before = baseline[key]
        after = current[key]
        before_hash = sha256_file(before)
        after_hash = sha256_file(after)
        changed = before_hash != after_hash or before.stat().st_size != after.stat().st_size
        prohibited = key in {"application", "publication", "editable_assembly"}
        expected = key in {"database", "workbook"}
        if prohibited:
            status = "passed" if not changed else "failed"
            drift_class = "prohibited_drift_absent" if not changed else "prohibited_drift_detected"
            resolution = "No mutation permitted; byte identity required."
        else:
            status = "passed" if changed else "failed"
            drift_class = "expected_governed_mutation" if changed else "expected_mutation_missing"
            resolution = "Copied derivative updated and independently verified."
        rows.append(
            {
                "artifact_key": key,
                "baseline_path": before.relative_to(baseline_root).as_posix(),
                "baseline_bytes": before.stat().st_size,
                "baseline_sha256": before_hash,
                "current_path": after.relative_to(current_root).as_posix(),
                "current_bytes": after.stat().st_size,
                "current_sha256": after_hash,
                "drift_class": drift_class,
                "resolution": resolution,
                "status": status,
            }
        )
    failed = [row for row in rows if row["status"] != "passed"]
    if failed:
        raise RuntimeError({"critical_drift_failure": failed})
    return rows


def project_diff(baseline: Path, current: Path) -> dict[str, Any]:
    base = {path.relative_to(baseline).as_posix(): (path.stat().st_size, sha256_file(path)) for path in baseline.rglob("*") if path.is_file()}
    now = {path.relative_to(current).as_posix(): (path.stat().st_size, sha256_file(path)) for path in current.rglob("*") if path.is_file()}
    added = sorted(set(now) - set(base))
    deleted = sorted(set(base) - set(now))
    modified = sorted(path for path in set(base) & set(now) if base[path] != now[path])
    prohibited_patterns = ("Integrated Manuscript", "Editable Integrated Manuscript Assembly", "App/human_pathogen_app.py")
    prohibited = [path for path in modified + deleted if any(pattern in path for pattern in prohibited_patterns)]
    status = "passed" if not deleted and not prohibited else "failed"
    result = {
        "status": status,
        "baseline_files": len(base),
        "current_files": len(now),
        "added_count": len(added),
        "modified_count": len(modified),
        "deleted_count": len(deleted),
        "prohibited_drift_count": len(prohibited),
        "added_paths": added,
        "modified_paths": modified,
        "deleted_paths": deleted,
        "prohibited_paths": prohibited,
    }
    if status != "passed":
        raise RuntimeError({"project_diff": result})
    return result


def recovery_events() -> list[dict[str, Any]]:
    events = [
        ("V3-CP4-S2-REC-CHECKPOINT1-CLEAN-RESTORE-PASSED", "None; Checkpoint 1 restoration completed.", "The exact Response 66 restore and Response 67 recovery were verified and clean-applied before Checkpoint 2 mutation.", "Restored the current project through Response 67 in an isolated working directory and verified database, workbook, application, and publication controls."),
        ("V3-CP4-S2-REC-FIELD-LEVEL-COVERAGE-COMPLETE", "None; logical-table and field audit completed.", "Checkpoint 2 required explicit field-level coverage rather than table-presence assumptions.", "Audited every logical SQLite table and every declared field, including required-field null/blank controls, distinctness, type, and primary-key metadata."),
        ("V3-CP4-S2-REC-QUERY-DISAMBIGUATION-COVERAGE-COMPLETE", "None; query coverage completed.", "Search, alias, historical-name, and semantic controls required reproducible current-database evidence.", "Executed controlled searches for Streptococcus, Strep bovis, bovis, SBSEC, GAS, GBS, pneumococcus, sanguis, gallolyticus, anginosus, Crypto, asymptomatic bacteriuria, endocarditis, meningitis, and antibiogram, plus locator, cross-reference, page-map, capability, and response controls."),
        ("V3-CP4-S2-REC-SOURCE-GOVERNANCE-RECONCILED", "None; source-governance reconciliation completed.", "Evidence/source tables vary in URL, year, title, authority, and linkage schema and required an explicit cross-table audit.", "Recorded table counts, source linkage, URL completeness, year coverage, authority/title fields, documented permissible gaps, and synchronized the results to SQLite, workbook, reports, QA, and indexes."),
        ("V3-CP4-S2-REC-CROSS-ARTIFACT-DRIFT-RESOLVED", "None; expected and prohibited drift classified.", "The database and workbook were expected to change while the publication, editable assembly, and main application source were required to remain immutable.", "Verified expected derivative mutation, prohibited drift absence, no deletions, and no unapproved publication or application-source change."),
        ("V3-CP4-S2-REC-APPLICATION-WORKBOOK-QUERY-AUDIT-PASSED", "None; application and workbook synchronization completed.", "Current database query and governance evidence needed read-only application and workbook surfaces.", "Added a checkpoint-specific audit and launcher, preserved the main application source, retained all prior workbook sheets, added detailed audit sheets, reran legacy tests, and verified all current controls."),
        ("V3-CP4-S2-REC-CHECKPOINT2-RECOVERY-CLEAN-VERIFIED", "None; Checkpoint 2 recovery package completed.", "Intermediate turns require complete recovery data tied directly to the last session-end full restore.", "Built the cumulative Response 66-to-68 overlay, deterministic apply utility, reports, manifests, checksums, and clean-applied the package before emission."),
    ]
    output = []
    for code, failed_step, reason, action in events:
        output.append(
            {
                "event_code": code,
                "occurred_at": NOW,
                "failed_step": failed_step,
                "exact_error_or_reason": reason,
                "intact_artifacts": "Accepted predecessor, frozen Section 3 release, complete Response 66 restore, Response 67 recovery, 537-page publication, editable assembly, main application source, prior database/workbook, and Google Drive custody remained intact.",
                "recovery_action": action,
                "validation_result": "Passed. No accepted or frozen source artifact was modified in place.",
                "data_quality_effect": "Governance, traceability, and regression coverage improved; no clinical claim was altered.",
                "next_checkpoint": "Remediation Section 4 of 5 Session 2 of 3 Checkpoint 3 of 3 — complete Session 2 restore.",
            }
        )
    return output


def synchronize_database(
    project: Path,
    source_db: Path,
    field_rows: list[dict[str, Any]],
    query_rows: list[dict[str, Any]],
    governance_rows: list[dict[str, Any]],
    events: list[dict[str, Any]],
) -> Path:
    target = source_db.with_name(
        f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 "
        "Remediation Section 4 of 5 Session 2 of 3 Checkpoint 2 of 3.sqlite"
    )
    shutil.copy2(source_db, target)
    con = sqlite3.connect(target)
    con.execute("PRAGMA foreign_keys=ON")
    try:
        con.execute("BEGIN IMMEDIATE")
        response_row = dict(RESPONSE68)
        if "reconciled_at" in cp1.table_columns(con, "thread_response_reconciliation_cp3"):
            response_row["reconciled_at"] = NOW
        cp1.schema_upsert(con, "thread_response_reconciliation_cp3", response_row, "response_key")
        if cp1.table_exists(con, "remediation_recovery_event"):
            columns = cp1.table_columns(con, "remediation_recovery_event")
            for event in events:
                cp1.schema_upsert(con, "remediation_recovery_event", {key: value for key, value in event.items() if key in columns}, "event_code")
        con.executescript(
            """
            CREATE TABLE IF NOT EXISTS section4_session2_field_coverage (
              section4_session2_field_coverage_id INTEGER PRIMARY KEY,
              checkpoint_code TEXT NOT NULL,
              table_name TEXT NOT NULL,
              field_name TEXT NOT NULL,
              declared_type TEXT,
              row_count INTEGER NOT NULL,
              not_null_declared INTEGER NOT NULL CHECK(not_null_declared IN (0,1)),
              primary_key INTEGER NOT NULL CHECK(primary_key IN (0,1)),
              null_count INTEGER,
              blank_count INTEGER,
              distinct_count INTEGER,
              status TEXT NOT NULL,
              notes TEXT,
              checked_at TEXT NOT NULL,
              UNIQUE(checkpoint_code, table_name, field_name)
            );
            CREATE TABLE IF NOT EXISTS section4_session2_query_coverage (
              section4_session2_query_coverage_id INTEGER PRIMARY KEY,
              checkpoint_code TEXT NOT NULL,
              query_key TEXT NOT NULL,
              query_text TEXT NOT NULL,
              query_type TEXT NOT NULL,
              result_count INTEGER NOT NULL,
              expected_minimum INTEGER NOT NULL,
              status TEXT NOT NULL,
              sample_json TEXT NOT NULL,
              notes TEXT,
              checked_at TEXT NOT NULL,
              UNIQUE(checkpoint_code, query_key)
            );
            CREATE TABLE IF NOT EXISTS section4_session2_source_governance (
              section4_session2_source_governance_id INTEGER PRIMARY KEY,
              checkpoint_code TEXT NOT NULL,
              audit_key TEXT NOT NULL,
              source_table TEXT NOT NULL,
              row_count INTEGER,
              url_column TEXT,
              missing_url_count INTEGER,
              nonempty_url_count INTEGER,
              year_column TEXT,
              latest_year INTEGER,
              title_column TEXT,
              missing_title_count INTEGER,
              authority_column TEXT,
              missing_authority_count INTEGER,
              status TEXT NOT NULL,
              details_json TEXT NOT NULL,
              checked_at TEXT NOT NULL,
              UNIQUE(checkpoint_code, audit_key)
            );
            CREATE TABLE IF NOT EXISTS section4_session2_drift_resolution (
              section4_session2_drift_resolution_id INTEGER PRIMARY KEY,
              checkpoint_code TEXT NOT NULL,
              artifact_key TEXT NOT NULL,
              baseline_path TEXT NOT NULL,
              baseline_bytes INTEGER NOT NULL,
              baseline_sha256 TEXT NOT NULL,
              current_path TEXT NOT NULL,
              current_bytes INTEGER NOT NULL,
              current_sha256 TEXT NOT NULL,
              drift_class TEXT NOT NULL,
              resolution TEXT NOT NULL,
              status TEXT NOT NULL,
              checked_at TEXT NOT NULL,
              UNIQUE(checkpoint_code, artifact_key)
            );
            """
        )
        checkpoint_code = "MRHPD-V3-CP4-S2-CP2"
        con.execute("DELETE FROM section4_session2_field_coverage WHERE checkpoint_code=?", (checkpoint_code,))
        for row in field_rows:
            con.execute(
                """
                INSERT INTO section4_session2_field_coverage
                (checkpoint_code,table_name,field_name,declared_type,row_count,not_null_declared,primary_key,null_count,blank_count,distinct_count,status,notes,checked_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (
                    checkpoint_code, row["table_name"], row["field_name"], row["declared_type"], row["row_count"],
                    row["not_null_declared"], row["primary_key"], row["null_count"], row["blank_count"], row["distinct_count"],
                    row["status"], row["notes"], NOW,
                ),
            )
        con.execute("DELETE FROM section4_session2_query_coverage WHERE checkpoint_code=?", (checkpoint_code,))
        for row in query_rows:
            con.execute(
                """
                INSERT INTO section4_session2_query_coverage
                (checkpoint_code,query_key,query_text,query_type,result_count,expected_minimum,status,sample_json,notes,checked_at)
                VALUES (?,?,?,?,?,?,?,?,?,?)
                """,
                (
                    checkpoint_code, row["query_key"], row["query_text"], row["query_type"], row["result_count"],
                    row["expected_minimum"], row["status"], json.dumps(row["sample_json"], ensure_ascii=False), row["notes"], NOW,
                ),
            )
        con.execute("DELETE FROM section4_session2_source_governance WHERE checkpoint_code=?", (checkpoint_code,))
        for row in governance_rows:
            con.execute(
                """
                INSERT INTO section4_session2_source_governance
                (checkpoint_code,audit_key,source_table,row_count,url_column,missing_url_count,nonempty_url_count,year_column,latest_year,title_column,missing_title_count,authority_column,missing_authority_count,status,details_json,checked_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (
                    checkpoint_code, row["audit_key"], row["source_table"], row.get("row_count"), row.get("url_column"),
                    row.get("missing_url_count"), row.get("nonempty_url_count"), row.get("year_column"), row.get("latest_year"),
                    row.get("title_column"), row.get("missing_title_count"), row.get("authority_column"), row.get("missing_authority_count"),
                    row["status"], json.dumps(row.get("details_json", {}), ensure_ascii=False), NOW,
                ),
            )
        con.execute(
            """
            INSERT INTO section4_session2_checkpoint
            (checkpoint_code,section_label,session_label,checkpoint_label,response_number,state,database_status,workbook_status,
             application_status,publication_status,capability_status,evidence_status,drift_status,accepted_predecessor_mutated,
             next_checkpoint,recorded_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            ON CONFLICT(checkpoint_code) DO UPDATE SET
              response_number=excluded.response_number,state=excluded.state,database_status=excluded.database_status,
              workbook_status=excluded.workbook_status,application_status=excluded.application_status,
              publication_status=excluded.publication_status,capability_status=excluded.capability_status,
              evidence_status=excluded.evidence_status,drift_status=excluded.drift_status,
              accepted_predecessor_mutated=excluded.accepted_predecessor_mutated,
              next_checkpoint=excluded.next_checkpoint,recorded_at=excluded.recorded_at
            """,
            (
                checkpoint_code, "Remediation Section 4 of 5", "Session 2 of 3", "Checkpoint 2 of 3", 68,
                "checkpoint_complete", "ok", "pending", "pending", "passed", "passed", "passed", "pending", 0,
                "Remediation Section 4 of 5 Session 2 of 3 Checkpoint 3 of 3 — complete Session 2 restore", NOW,
            ),
        )
        if cp1.table_exists(con, "section4_checkpoint"):
            con.execute(
                """
                INSERT INTO section4_checkpoint
                (checkpoint_code,section_label,session_label,checkpoint_label,response_number,state,database_integrity,
                 foreign_key_violations,workbook_status,application_status,publication_sha256,accepted_predecessor_mutated,recorded_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
                ON CONFLICT(checkpoint_code) DO UPDATE SET
                  response_number=excluded.response_number,state=excluded.state,database_integrity=excluded.database_integrity,
                  foreign_key_violations=excluded.foreign_key_violations,workbook_status=excluded.workbook_status,
                  application_status=excluded.application_status,publication_sha256=excluded.publication_sha256,
                  accepted_predecessor_mutated=excluded.accepted_predecessor_mutated,recorded_at=excluded.recorded_at
                """,
                (checkpoint_code, "Remediation Section 4 of 5", "Session 2 of 3", "Checkpoint 2 of 3", 68, "checkpoint_complete", "ok", 0, "pending", "pending", PUBLICATION_SHA256, 0, NOW),
            )
        if cp1.table_exists(con, "metadata") and {"key", "value"}.issubset(cp1.table_columns(con, "metadata")):
            updates = {
                "version": PROJECT_VERSION,
                "current_remediation_section": "Remediation Section 4 of 5",
                "current_session": "Session 2 of 3",
                "current_checkpoint": "Checkpoint 2 of 3 COMPLETE",
                "current_response": "68",
                "current_canonical_database": target.name,
                "accepted_predecessor_mutated": "no",
                "last_updated_utc": NOW,
                "next_checkpoint": "Remediation Section 4 of 5 Session 2 of 3 Checkpoint 3 of 3 — complete Session 2 restore",
            }
            for key, value in updates.items():
                con.execute("INSERT INTO metadata(key,value) VALUES (?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value", (key, value))
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or fk:
            raise RuntimeError({"integrity": integrity, "foreign_keys": fk[:30]})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()
    return target


def insert_drift_rows(db: Path, rows: list[dict[str, Any]]) -> None:
    con = sqlite3.connect(db)
    try:
        con.execute("BEGIN IMMEDIATE")
        checkpoint_code = "MRHPD-V3-CP4-S2-CP2"
        con.execute("DELETE FROM section4_session2_drift_resolution WHERE checkpoint_code=?", (checkpoint_code,))
        for row in rows:
            con.execute(
                """
                INSERT INTO section4_session2_drift_resolution
                (checkpoint_code,artifact_key,baseline_path,baseline_bytes,baseline_sha256,current_path,current_bytes,current_sha256,drift_class,resolution,status,checked_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (
                    checkpoint_code, row["artifact_key"], row["baseline_path"], row["baseline_bytes"], row["baseline_sha256"],
                    row["current_path"], row["current_bytes"], row["current_sha256"], row["drift_class"], row["resolution"], row["status"], NOW,
                ),
            )
        con.execute("UPDATE section4_session2_checkpoint SET drift_status='passed',recorded_at=? WHERE checkpoint_code=?", (NOW, checkpoint_code))
        if con.execute("PRAGMA integrity_check").fetchone()[0] != "ok" or list(con.execute("PRAGMA foreign_key_check")):
            raise RuntimeError("Database failed after drift insert")
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()


def database_qa(db: Path, project: Path) -> dict[str, Any]:
    con = sqlite3.connect(db)
    try:
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        table_count = int(con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0])
        response68 = int(con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R68'").fetchone()[0])
        field_total = int(con.execute("SELECT COUNT(*) FROM section4_session2_field_coverage WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2'").fetchone()[0])
        field_failed = int(con.execute("SELECT COUNT(*) FROM section4_session2_field_coverage WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2' AND status!='passed'").fetchone()[0])
        query_total = int(con.execute("SELECT COUNT(*) FROM section4_session2_query_coverage WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2'").fetchone()[0])
        query_failed = int(con.execute("SELECT COUNT(*) FROM section4_session2_query_coverage WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2' AND status!='passed'").fetchone()[0])
        governance_total = int(con.execute("SELECT COUNT(*) FROM section4_session2_source_governance WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2'").fetchone()[0])
        governance_failed = int(con.execute("SELECT COUNT(*) FROM section4_session2_source_governance WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2' AND status!='passed'").fetchone()[0])
        drift_total = int(con.execute("SELECT COUNT(*) FROM section4_session2_drift_resolution WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2'").fetchone()[0])
        drift_failed = int(con.execute("SELECT COUNT(*) FROM section4_session2_drift_resolution WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2' AND status!='passed'").fetchone()[0])
        checkpoint = con.execute("SELECT state,workbook_status,application_status,publication_status FROM section4_session2_checkpoint WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2'").fetchone()
        locators = int(con.execute("SELECT COUNT(*) FROM publication_index_locator").fetchone()[0])
        cross_refs = int(con.execute("SELECT COUNT(*) FROM publication_cross_reference").fetchone()[0])
        response_records = int(con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3").fetchone()[0])
        fractional_records = int(con.execute("SELECT COUNT(*) FROM fractional_prompt_cp3").fetchone()[0])
        recovery_records = int(con.execute("SELECT COUNT(*) FROM remediation_recovery_event").fetchone()[0])
    finally:
        con.close()
    qa = {
        "status": "passed",
        "canonical_database": db.relative_to(project).as_posix(),
        "bytes": db.stat().st_size,
        "sha256": sha256_file(db),
        "integrity": integrity,
        "foreign_key_violations": len(fk),
        "table_count": table_count,
        "response_records": response_records,
        "response68_records": response68,
        "fractional_prompt_records": fractional_records,
        "recovery_event_records": recovery_records,
        "field_coverage_records": field_total,
        "field_coverage_failures": field_failed,
        "query_coverage_records": query_total,
        "query_coverage_failures": query_failed,
        "source_governance_records": governance_total,
        "source_governance_failures": governance_failed,
        "drift_resolution_records": drift_total,
        "drift_resolution_failures": drift_failed,
        "checkpoint_state": checkpoint[0] if checkpoint else None,
        "workbook_status": checkpoint[1] if checkpoint else None,
        "application_status": checkpoint[2] if checkpoint else None,
        "publication_status": checkpoint[3] if checkpoint else None,
        "publication_index_locators": locators,
        "current_cross_references": cross_refs,
        "accepted_predecessor_mutated": False,
    }
    required = [
        integrity == "ok", not fk, response68 == 1, field_total > 0, field_failed == 0, query_total >= len(SEARCH_TERMS) + 5,
        query_failed == 0, governance_total > 0, governance_failed == 0, drift_total >= 5, drift_failed == 0,
        checkpoint is not None and checkpoint[0] == "checkpoint_complete", locators == 4011, cross_refs >= 12,
    ]
    if not all(required):
        qa["status"] = "failed"
        raise RuntimeError({"database_qa": qa})
    return qa


def application_audit_source() -> str:
    return '''#!/usr/bin/env python3
+from __future__ import annotations
+import argparse,hashlib,json,re,sqlite3,subprocess,sys
+from pathlib import Path
+
+def sha(path):
+    h=hashlib.sha256()
+    with open(path,'rb') as f:
+        for block in iter(lambda:f.read(1024*1024),b''): h.update(block)
+    return h.hexdigest()
+
+p=argparse.ArgumentParser(description="MRHPD Section 4 Session 2 Checkpoint 2 current-database audit")
+p.add_argument('--db',type=Path,required=True)
+p.add_argument('--app',type=Path,required=True)
+p.add_argument('--output',type=Path)
+a=p.parse_args()
+con=sqlite3.connect(a.db)
+try:
+    checks={
+        'integrity':con.execute('PRAGMA integrity_check').fetchone()[0]=='ok',
+        'foreign_keys':not list(con.execute('PRAGMA foreign_key_check')),
+        'response68':con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R68'").fetchone()[0]==1,
+        'field_coverage':con.execute("SELECT COUNT(*) FROM section4_session2_field_coverage WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2' AND status!='passed'").fetchone()[0]==0 and con.execute("SELECT COUNT(*) FROM section4_session2_field_coverage WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2'").fetchone()[0]>0,
+        'query_coverage':con.execute("SELECT COUNT(*) FROM section4_session2_query_coverage WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2' AND status!='passed'").fetchone()[0]==0 and con.execute("SELECT COUNT(*) FROM section4_session2_query_coverage WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2'").fetchone()[0]>=21,
+        'source_governance':con.execute("SELECT COUNT(*) FROM section4_session2_source_governance WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2' AND status!='passed'").fetchone()[0]==0,
+        'drift_resolution':con.execute("SELECT COUNT(*) FROM section4_session2_drift_resolution WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2' AND status!='passed'").fetchone()[0]==0,
+        'locators':con.execute('SELECT COUNT(*) FROM publication_index_locator').fetchone()[0]==4011,
+        'cross_references':con.execute('SELECT COUNT(*) FROM publication_cross_reference').fetchone()[0]>=12,
+        'search_terms':con.execute("SELECT COUNT(*) FROM section4_session2_query_coverage WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2' AND query_type='cross_table_text_search' AND status='passed'").fetchone()[0]>=16,
+    }
+finally:
+    con.close()
+source=a.app.read_text(encoding='utf-8',errors='replace')
+checks['native_db_argument']='--db' in source
+help_run=subprocess.run([sys.executable,str(a.app),'--help'],text=True,capture_output=True,timeout=60)
+checks['application_help']=help_run.returncode==0
+result={'schema':'mrhpd-s4s2-cp2-application-audit-1.0','status':'passed' if all(checks.values()) else 'failed','database':str(a.db),'database_sha256':sha(a.db),'application':str(a.app),'application_sha256':sha(a.app),'checks':checks,'help_stdout_tail':help_run.stdout[-3000:],'help_stderr_tail':help_run.stderr[-3000:]}
+if a.output:
+    a.output.parent.mkdir(parents=True,exist_ok=True)
+    a.output.write_text(json.dumps(result,indent=2)+'\n',encoding='utf-8')
+print(json.dumps(result,indent=2))
+raise SystemExit(0 if result['status']=='passed' else 1)
+'''.replace("\n+", "\n")


def synchronize_application(project: Path, db: Path) -> tuple[dict[str, Path], dict[str, Any]]:
    app = cp1.find_by_hash(project, "human_pathogen_app.py", APPLICATION_SHA256)
    app_dir = app.parent
    audit = app_dir / "section4_session2_checkpoint2_audit.py"
    text_write(audit, application_audit_source())
    launcher = app_dir / "run_section4_session2_checkpoint2.py"
    text_write(
        launcher,
        f'''#!/usr/bin/env python3
from pathlib import Path
import subprocess,sys
ROOT=Path(__file__).resolve().parent.parent
APP=Path(__file__).resolve().parent/"human_pathogen_app.py"
AUDIT=Path(__file__).resolve().parent/"section4_session2_checkpoint2_audit.py"
DB=ROOT/{db.relative_to(project).as_posix()!r}
if len(sys.argv)>1 and sys.argv[1]=='--audit':
    raise SystemExit(subprocess.call([sys.executable,str(AUDIT),'--db',str(DB),'--app',str(APP),*sys.argv[2:]]))
raise SystemExit(subprocess.call([sys.executable,str(APP),'--db',str(DB),*sys.argv[1:]]))
''',
    )
    state = app_dir / "CURRENT_PROJECT_STATE.json"
    json_write(
        state,
        {
            "schema": "mrhpd-current-application-state-1.0",
            "generated_at": NOW,
            "remediation_section": "4 of 5",
            "session": "2 of 3",
            "checkpoint": "2 of 3 COMPLETE",
            "response": 68,
            "canonical_database": db.name,
            "database_relative_path": db.relative_to(project).as_posix(),
            "database_sha256": sha256_file(db),
            "application_sha256": APPLICATION_SHA256,
            "audit_utility": audit.name,
            "accepted_predecessor_mutated": False,
            "next_checkpoint": "Checkpoint 3 of 3 — complete Session 2 restore",
        },
    )
    text_write(app_dir / "CURRENT_DATABASE.txt", db.relative_to(project).as_posix())
    readme = app_dir / "README_SECTION4_SESSION2_CHECKPOINT2.md"
    text_write(
        readme,
        f"""# Human Pathogen Database local application — Section 4 Session 2 Checkpoint 2

Canonical database: `{db.name}`

- Run `python run_section4_session2_checkpoint2.py --audit` for detailed field, query, governance, and drift verification.
- Pass ordinary application arguments after `python run_section4_session2_checkpoint2.py` to launch the read-only application with the current database.
- The main application source remains byte-identical and receives the current database through its native `--db` interface.
""",
    )
    env = os.environ.copy()
    env["MRHPD_DATABASE"] = str(db)
    env["MRHPD_DB_PATH"] = str(db)
    tests = []
    for test in sorted(app_dir.glob("test*.py")):
        result = subprocess.run([sys.executable, str(test)], cwd=app_dir, env=env, text=True, capture_output=True, timeout=600)
        record = {"test": test.name, "returncode": result.returncode, "stdout_tail": result.stdout[-20000:], "stderr_tail": result.stderr[-12000:]}
        tests.append(record)
        if result.returncode != 0:
            raise RuntimeError({"application_test_failed": record})
    audit_output = project / "QA" / "Section 4 Session 2" / "Checkpoint 2" / "APPLICATION_CHECKPOINT2_AUDIT.json"
    result = subprocess.run([sys.executable, str(audit), "--db", str(db), "--app", str(app), "--output", str(audit_output)], cwd=app_dir, text=True, capture_output=True, timeout=300)
    if result.returncode != 0:
        raise RuntimeError({"checkpoint2_application_audit_failed": {"stdout": result.stdout[-20000:], "stderr": result.stderr[-12000:]}})
    qa = {
        "status": "passed",
        "main_application": app.relative_to(project).as_posix(),
        "main_application_sha256": sha256_file(app),
        "main_application_unchanged": sha256_file(app) == APPLICATION_SHA256,
        "canonical_database": db.relative_to(project).as_posix(),
        "legacy_test_files": len(tests),
        "legacy_tests": tests,
        "checkpoint2_audit": audit_output.relative_to(project).as_posix(),
        "checkpoint2_audit_stdout": result.stdout[-15000:],
        "launcher": launcher.relative_to(project).as_posix(),
        "current_state": state.relative_to(project).as_posix(),
    }
    return {"application": app, "audit": audit, "launcher": launcher, "state": state, "readme": readme, "audit_output": audit_output}, qa


def autosize(ws, max_width: int = 70) -> None:
    from openpyxl.styles import Alignment
    from openpyxl.utils import get_column_letter
    for index in range(1, ws.max_column + 1):
        values = [str(ws.cell(row=row, column=index).value or "") for row in range(1, min(ws.max_row, 500) + 1)]
        ws.column_dimensions[get_column_letter(index)].width = max(11, min(max_width, max((len(value) for value in values), default=9) + 2))
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")


def synchronize_workbook(
    project: Path,
    source_workbook: Path,
    original_sheets: list[str],
    field_rows: list[dict[str, Any]],
    query_rows: list[dict[str, Any]],
    governance_rows: list[dict[str, Any]],
    drift_rows: list[dict[str, Any]],
    summaries: dict[str, Any],
) -> tuple[Path, dict[str, Any]]:
    from openpyxl import load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    target = source_workbook.with_name(
        f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 "
        "Remediation Section 4 of 5 Session 2 of 3 Checkpoint 2 of 3 Detailed Coverage.xlsx"
    )
    shutil.copy2(source_workbook, target)
    wb = load_workbook(target)
    managed = ["S4S2 Field Coverage", "S4S2 Query Coverage", "S4S2 Governance", "S4S2 Drift Resolve", "S4S2 CP2 QA"]
    for name in managed:
        if name in wb.sheetnames:
            del wb[name]
    navy, teal, white = "17324F", "167D86", "FFFFFF"
    if "S4S2 Dashboard" in wb.sheetnames:
        ws = wb["S4S2 Dashboard"]
        ws.append(["Checkpoint 2 field records", ">0", summaries["field"]["audited_fields"], "PASS", "section4_session2_field_coverage"])
        ws.append(["Checkpoint 2 query records", len(query_rows), summaries["query"]["queries"], "PASS", "section4_session2_query_coverage"])
        ws.append(["Checkpoint 2 governance records", len(governance_rows), summaries["governance"]["audited_tables"], "PASS", "section4_session2_source_governance"])
        ws.append(["Checkpoint 2 prohibited drift", 0, summaries["project_diff"]["prohibited_drift_count"], "PASS", "DRIFT_FILE_DIFF.json"])
        autosize(ws)

    datasets = [
        ("S4S2 Field Coverage", ["table_name", "field_name", "declared_type", "row_count", "not_null_declared", "primary_key", "null_count", "blank_count", "distinct_count", "status", "notes"], field_rows),
        ("S4S2 Query Coverage", ["query_key", "query_text", "query_type", "result_count", "expected_minimum", "status", "sample_json", "notes"], query_rows),
        ("S4S2 Governance", ["audit_key", "source_table", "row_count", "url_column", "missing_url_count", "nonempty_url_count", "year_column", "latest_year", "title_column", "missing_title_count", "authority_column", "missing_authority_count", "status", "details_json"], governance_rows),
        ("S4S2 Drift Resolve", ["artifact_key", "baseline_path", "baseline_bytes", "baseline_sha256", "current_path", "current_bytes", "current_sha256", "drift_class", "resolution", "status"], drift_rows),
    ]
    for name, headers, rows in datasets:
        ws = wb.create_sheet(name)
        ws.append(headers)
        for row in rows:
            ws.append([json.dumps(row.get(header), ensure_ascii=False) if isinstance(row.get(header), (dict, list)) else row.get(header) for header in headers])
        for cell in ws[1]:
            cell.font = Font(bold=True, color=white)
            cell.fill = PatternFill("solid", fgColor=navy)
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        autosize(ws)
    qa_ws = wb.create_sheet("S4S2 CP2 QA")
    qa_ws.append(["Area", "Control", "Value", "Status"])
    for area, summary in summaries.items():
        if isinstance(summary, dict):
            for key, value in summary.items():
                qa_ws.append([area, key, json.dumps(value, ensure_ascii=False) if isinstance(value, (list, dict)) else value, "PASS" if key == "status" and value == "passed" else ""])
    for cell in qa_ws[1]:
        cell.font = Font(bold=True, color=white)
        cell.fill = PatternFill("solid", fgColor=teal)
    qa_ws.freeze_panes = "A2"
    qa_ws.auto_filter.ref = qa_ws.dimensions
    autosize(qa_ws)
    wb.save(target)

    wb2 = load_workbook(target, read_only=True, data_only=False)
    final_sheets = list(wb2.sheetnames)
    formula_count = 0
    formula_errors: list[str] = []
    tokens = ("#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "#N/A", "#NUM!", "#NULL!")
    for ws in wb2.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                value = cell.value
                if isinstance(value, str) and value.startswith("="):
                    formula_count += 1
                if isinstance(value, str) and any(token in value for token in tokens):
                    formula_errors.append(f"{ws.title}!{cell.coordinate}:{value}")
    wb2.close()
    missing_original = sorted(set(original_sheets) - set(final_sheets))
    missing_managed = sorted(set(managed) - set(final_sheets))
    qa = {
        "status": "passed" if not missing_original and not missing_managed and not formula_errors else "failed",
        "source_workbook": source_workbook.relative_to(project).as_posix(),
        "current_workbook": target.relative_to(project).as_posix(),
        "bytes": target.stat().st_size,
        "sha256": sha256_file(target),
        "source_sheet_count": len(original_sheets),
        "current_sheet_count": len(final_sheets),
        "original_sheets_preserved": not missing_original,
        "missing_original_sheets": missing_original,
        "managed_sheets": managed,
        "missing_managed_sheets": missing_managed,
        "formula_count": formula_count,
        "formula_error_count": len(formula_errors),
        "formula_errors": formula_errors,
    }
    if qa["status"] != "passed":
        raise RuntimeError({"workbook_qa": qa})
    return target, qa


def finalize_database(db: Path, workbook_qa: dict[str, Any], application_qa: dict[str, Any], publication_qa: dict[str, Any]) -> None:
    con = sqlite3.connect(db)
    try:
        con.execute("BEGIN IMMEDIATE")
        con.execute(
            "UPDATE section4_session2_checkpoint SET workbook_status=?,application_status=?,publication_status=?,drift_status='passed',recorded_at=? WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2'",
            (workbook_qa["status"], application_qa["status"], publication_qa["status"], NOW),
        )
        if cp1.table_exists(con, "section4_checkpoint"):
            con.execute(
                "UPDATE section4_checkpoint SET workbook_status=?,application_status=?,recorded_at=? WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2'",
                (workbook_qa["status"], application_qa["status"], NOW),
            )
        if con.execute("PRAGMA integrity_check").fetchone()[0] != "ok" or list(con.execute("PRAGMA foreign_key_check")):
            raise RuntimeError("Database failed during final status synchronization")
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()


def build_tracking(project: Path, db: Path) -> list[Path]:
    from docx import Document
    from docx.shared import Inches, RGBColor

    root = project / "Tracking" / "Section 4 Session 2" / "Checkpoint 2"
    root.mkdir(parents=True, exist_ok=True)
    response = root / "Response_68_Tracking.json"
    json_write(response, RESPONSE68)
    raw_net = root / "RAW_AND_NET_TRACKING.md"
    text_write(
        raw_net,
        f"""# Human Pathogen Database — Raw and Net Tracking Through Response 68

## Major Topic: Human Pathogen Database remediation

### Raw Prompt 68

```text
{RAW_PROMPT}
```

### Raw Response 68

The final user-visible response is represented by the source-supported summary below. Complete current files and QA evidence are in the checkpoint-recovery package.

**Summary:** {RESPONSE68['summary']}

### Net Prompt

{NET_PROMPT}

### Net Response

{NET_RESPONSE}

## Current disposition

- Remediation Section 4 of 5: CONTINUE
- Session 2 of 3: CONTINUE
- Checkpoint 2 of 3: COMPLETE
- Next: Checkpoint 3 of 3 — complete Session 2 restore
""",
    )
    cumulative = root / "CUMULATIVE_THREAD_INDEX_UPDATE.md"
    text_write(
        cumulative,
        f"""# Cumulative Thread Index Update — Response 68

## Human Pathogen Database remediation

### Response 68 — Detailed field, query, evidence, and drift-reconciliation checkpoint

**Goal:** {RESPONSE68['goal']}

**Output:** {RESPONSE68['summary']}

**Disposition:** Checkpoint 2 of 3 COMPLETE; Section 4 Session 2 CONTINUE; Checkpoint 3 of 3 next with the complete Session 2 restore.
""",
    )
    con = sqlite3.connect(db)
    con.row_factory = sqlite3.Row
    rows = [dict(row) for row in con.execute("SELECT * FROM thread_response_reconciliation_cp3 ORDER BY response_number, response_key")]
    con.close()

    def make_doc(path: Path, raw: bool) -> None:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        doc.add_heading("Human Pathogen Database", 0)
        doc.add_paragraph("Alternating Raw Prompts/Responses" if raw else "Alternating Net Prompt/Response")
        doc.add_heading("Human Pathogen Database remediation", level=1)
        if raw:
            for row in rows:
                number = row.get("response_label") or row.get("response_number")
                doc.add_heading(f"Response {number}: {(row.get('title') or 'Untitled')[:120]}", level=2)
                prompt = doc.add_paragraph().add_run("Prompt\n" + (row.get("raw_prompt") or "(UNRECOVERED OR NOT STORED)"))
                prompt.font.color.rgb = RGBColor(31, 78, 121)
                response_text = row.get("raw_response") or "(RAW RESPONSE NOT RECOVERED)"
                if row.get("summary"):
                    response_text += "\n\nSummary: " + row["summary"]
                response_run = doc.add_paragraph().add_run("Response\n" + response_text)
                response_run.font.color.rgb = RGBColor(46, 125, 50)
        else:
            doc.add_heading("Net Prompt", level=2)
            run = doc.add_paragraph().add_run(NET_PROMPT)
            run.font.color.rgb = RGBColor(31, 78, 121)
            doc.add_heading("Net Response", level=2)
            run = doc.add_paragraph().add_run(NET_RESPONSE)
            run.font.color.rgb = RGBColor(46, 125, 50)
        doc.save(path)

    raw_docx = root / "Alternating Raw Prompts and Responses Through Response 68.docx"
    net_docx = root / "Alternating Net Prompts and Responses Through Response 68.docx"
    make_doc(raw_docx, True)
    make_doc(net_docx, False)
    return [response, raw_net, cumulative, raw_docx, net_docx]


def build_reports(
    project: Path,
    field_rows: list[dict[str, Any]],
    query_rows: list[dict[str, Any]],
    governance_rows: list[dict[str, Any]],
    drift_rows: list[dict[str, Any]],
    summaries: dict[str, Any],
) -> list[Path]:
    from docx import Document
    from docx.shared import Inches
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from pypdf import PdfReader
    import fitz

    report_dir = project / "Reports" / "Section 4 Session 2" / "Checkpoint 2"
    report_dir.mkdir(parents=True, exist_ok=True)
    stem = f"MRHPD v{PROJECT_VERSION} Section 4 Session 2 Checkpoint 2 Detailed Coverage and Drift Reconciliation"
    docx_path = report_dir / f"{stem}.docx"
    pdf_path = report_dir / f"{stem}.pdf"
    xlsx_path = report_dir / f"{stem} Register.xlsx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    doc.add_heading("Human Pathogen Database", 0)
    doc.add_paragraph("Section 4 Session 2 — Detailed Coverage and Drift Reconciliation")
    doc.add_paragraph(f"Checkpoint 2 of 3 • Response 68 • {NOW}")
    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(RESPONSE68["summary"])
    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = "Table Grid"
    for index, heading in enumerate(["Area", "Result", "Status"]):
        summary_table.rows[0].cells[index].text = heading
    for area, result in [
        ("Logical tables", summaries["field"]["logical_tables"]),
        ("Audited fields", summaries["field"]["audited_fields"]),
        ("Controlled queries", summaries["query"]["queries"]),
        ("Source-governance tables", summaries["governance"]["audited_tables"]),
        ("Critical drift records", len(drift_rows)),
        ("Prohibited drift", summaries["project_diff"]["prohibited_drift_count"]),
    ]:
        cells = summary_table.add_row().cells
        cells[0].text = str(area)
        cells[1].text = str(result)
        cells[2].text = "PASS"
    doc.add_heading("Field-level coverage", level=1)
    doc.add_paragraph("Every logical SQLite table and declared field was audited for type, row count, required-field nulls, blank required text, distinct values, and key constraints. Nullable gaps are documented rather than misclassified as failures.")
    field_table = doc.add_table(rows=1, cols=6)
    field_table.style = "Table Grid"
    for index, heading in enumerate(["Table", "Field", "Type", "Rows", "Null/blank", "Status"]):
        field_table.rows[0].cells[index].text = heading
    for row in field_rows[:120]:
        cells = field_table.add_row().cells
        values = [row["table_name"], row["field_name"], row["declared_type"], row["row_count"], f"{row['null_count']}/{row['blank_count']}", row["status"].upper()]
        for index, value in enumerate(values):
            cells[index].text = str(value)
    doc.add_heading("Query and disambiguation coverage", level=1)
    query_table = doc.add_table(rows=1, cols=5)
    query_table.style = "Table Grid"
    for index, heading in enumerate(["Query", "Type", "Actual", "Expected", "Status"]):
        query_table.rows[0].cells[index].text = heading
    for row in query_rows:
        cells = query_table.add_row().cells
        for index, value in enumerate([row["query_text"], row["query_type"], row["result_count"], row["expected_minimum"], row["status"].upper()]):
            cells[index].text = str(value)
    doc.add_heading("Source governance", level=1)
    governance_table = doc.add_table(rows=1, cols=6)
    governance_table.style = "Table Grid"
    for index, heading in enumerate(["Table", "Rows", "URL field", "Nonempty URLs", "Latest year", "Status"]):
        governance_table.rows[0].cells[index].text = heading
    for row in governance_rows:
        cells = governance_table.add_row().cells
        for index, value in enumerate([row["source_table"], row.get("row_count"), row.get("url_column"), row.get("nonempty_url_count"), row.get("latest_year"), row["status"].upper()]):
            cells[index].text = "" if value is None else str(value)
    doc.add_heading("Drift reconciliation", level=1)
    drift_table = doc.add_table(rows=1, cols=5)
    drift_table.style = "Table Grid"
    for index, heading in enumerate(["Artifact", "Class", "Before", "After", "Status"]):
        drift_table.rows[0].cells[index].text = heading
    for row in drift_rows:
        cells = drift_table.add_row().cells
        for index, value in enumerate([row["artifact_key"], row["drift_class"], row["baseline_sha256"][:16], row["current_sha256"][:16], row["status"].upper()]):
            cells[index].text = str(value)
    doc.add_heading("Next checkpoint", level=1)
    doc.add_paragraph("Checkpoint 3 of 3 will complete Session 2, freeze the synchronized current package, rerun complete database/workbook/application/publication/index/archive controls, and emit the next full self-contained restore.")
    doc.save(docx_path)

    styles = getSampleStyleSheet()
    story = [
        Paragraph("Human Pathogen Database", styles["Title"]),
        Paragraph("Section 4 Session 2 — Detailed Coverage and Drift Reconciliation", styles["Heading1"]),
        Paragraph(f"Checkpoint 2 of 3 • Response 68 • {NOW}", styles["Normal"]),
        Spacer(1, 0.18 * inch),
        Paragraph("Executive summary", styles["Heading1"]),
        Paragraph(RESPONSE68["summary"], styles["BodyText"]),
    ]
    summary_data = [["Area", "Result", "Status"], ["Logical tables", summaries["field"]["logical_tables"], "PASS"], ["Audited fields", summaries["field"]["audited_fields"], "PASS"], ["Queries", summaries["query"]["queries"], "PASS"], ["Governance tables", summaries["governance"]["audited_tables"], "PASS"], ["Prohibited drift", 0, "PASS"]]
    table = Table(summary_data, colWidths=[2.3 * inch, 3.4 * inch, 0.8 * inch], repeatRows=1)
    table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17324F")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.4, colors.grey), ("FONTSIZE", (0, 0), (-1, -1), 8)]))
    story += [Spacer(1, 0.15 * inch), table, PageBreak(), Paragraph("Field-level coverage", styles["Heading1"])]
    field_data = [["Table", "Field", "Rows", "Null", "Blank", "Status"]]
    for row in field_rows[:180]:
        field_data.append([row["table_name"][:36], row["field_name"][:32], row["row_count"], row["null_count"], row["blank_count"], row["status"].upper()])
    field_pdf = Table(field_data, colWidths=[1.7 * inch, 1.55 * inch, 0.7 * inch, 0.65 * inch, 0.65 * inch, 0.7 * inch], repeatRows=1)
    field_pdf.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#167D86")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.25, colors.grey), ("FONTSIZE", (0, 0), (-1, -1), 5.5), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story += [field_pdf, PageBreak(), Paragraph("Query and disambiguation coverage", styles["Heading1"])]
    query_data = [["Query", "Type", "Actual", "Expected", "Status"]] + [[row["query_text"], row["query_type"], row["result_count"], row["expected_minimum"], row["status"].upper()] for row in query_rows]
    query_pdf = Table(query_data, colWidths=[2.1 * inch, 2.1 * inch, 0.8 * inch, 0.8 * inch, 0.7 * inch], repeatRows=1)
    query_pdf.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17324F")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.35, colors.grey), ("FONTSIZE", (0, 0), (-1, -1), 7)]))
    story += [query_pdf, PageBreak(), Paragraph("Source governance", styles["Heading1"])]
    governance_data = [["Table", "Rows", "URL", "Nonempty", "Year", "Status"]] + [[row["source_table"], row.get("row_count"), row.get("url_column") or "", row.get("nonempty_url_count") if row.get("nonempty_url_count") is not None else "", row.get("latest_year") or "", row["status"].upper()] for row in governance_rows]
    gov_pdf = Table(governance_data, colWidths=[2.3 * inch, 0.65 * inch, 1.25 * inch, 0.8 * inch, 0.65 * inch, 0.7 * inch], repeatRows=1)
    gov_pdf.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#167D86")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.3, colors.grey), ("FONTSIZE", (0, 0), (-1, -1), 6.5)]))
    story += [gov_pdf, PageBreak(), Paragraph("Drift reconciliation and next checkpoint", styles["Heading1"])]
    drift_data = [["Artifact", "Class", "Status"]] + [[row["artifact_key"], row["drift_class"], row["status"].upper()] for row in drift_rows]
    drift_pdf = Table(drift_data, colWidths=[1.8 * inch, 4.2 * inch, 0.8 * inch], repeatRows=1)
    drift_pdf.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17324F")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.4, colors.grey), ("FONTSIZE", (0, 0), (-1, -1), 8)]))
    story += [drift_pdf, Spacer(1, 0.2 * inch), Paragraph("Checkpoint 3 of 3 will complete Session 2 and emit the next full self-contained restore.", styles["BodyText"])]
    SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=32, leftMargin=32, topMargin=32, bottomMargin=32, title=stem, author="Brent McAnulty, M.D.").build(story)

    wb = Workbook()
    wb.remove(wb.active)
    datasets = {
        "Summary": [["Area", "Value"], ["Status", "passed"], ["Response", 68], ["Checkpoint", "2 of 3 COMPLETE"], ["Session", "2 of 3 CONTINUE"], ["Logical tables", summaries["field"]["logical_tables"]], ["Fields", summaries["field"]["audited_fields"]], ["Queries", summaries["query"]["queries"]], ["Governance tables", summaries["governance"]["audited_tables"]], ["Prohibited drift", summaries["project_diff"]["prohibited_drift_count"]]],
        "Field Coverage": [["Table", "Field", "Type", "Rows", "Not null", "PK", "Null", "Blank", "Distinct", "Status", "Notes"]] + [[row["table_name"], row["field_name"], row["declared_type"], row["row_count"], row["not_null_declared"], row["primary_key"], row["null_count"], row["blank_count"], row["distinct_count"], row["status"], row["notes"]] for row in field_rows],
        "Query Coverage": [["Key", "Query", "Type", "Actual", "Expected", "Status", "Samples", "Notes"]] + [[row["query_key"], row["query_text"], row["query_type"], row["result_count"], row["expected_minimum"], row["status"], json.dumps(row["sample_json"], ensure_ascii=False), row["notes"]] for row in query_rows],
        "Governance": [["Table", "Rows", "URL", "Missing URLs", "Nonempty URLs", "Year", "Latest", "Title", "Missing Titles", "Authority", "Missing Authority", "Status"]] + [[row["source_table"], row.get("row_count"), row.get("url_column"), row.get("missing_url_count"), row.get("nonempty_url_count"), row.get("year_column"), row.get("latest_year"), row.get("title_column"), row.get("missing_title_count"), row.get("authority_column"), row.get("missing_authority_count"), row["status"]] for row in governance_rows],
        "Drift": [["Artifact", "Baseline path", "Baseline SHA", "Current path", "Current SHA", "Class", "Resolution", "Status"]] + [[row["artifact_key"], row["baseline_path"], row["baseline_sha256"], row["current_path"], row["current_sha256"], row["drift_class"], row["resolution"], row["status"]] for row in drift_rows],
    }
    for name, data in datasets.items():
        ws = wb.create_sheet(name)
        for row in data:
            ws.append(row)
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="17324F")
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        autosize(ws)
    wb.save(xlsx_path)

    with zipfile.ZipFile(docx_path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("DOCX CRC failed")
    with zipfile.ZipFile(xlsx_path) as zf:
        if zf.testzip() is not None:
            raise RuntimeError("XLSX CRC failed")
    reader = PdfReader(str(pdf_path))
    text_chars = sum(len(page.extract_text() or "") for page in reader.pages)
    if len(reader.pages) < 5 or text_chars < 3000:
        raise RuntimeError({"pdf_validation": {"pages": len(reader.pages), "text_chars": text_chars}})
    render_dir = report_dir / "PDF Render QA"
    render_dir.mkdir(exist_ok=True)
    rendered = []
    pdf_doc = fitz.open(pdf_path)
    for index, page in enumerate(pdf_doc):
        pix = page.get_pixmap(matrix=fitz.Matrix(1.25, 1.25), alpha=False)
        path = render_dir / f"page-{index + 1:03d}.png"
        pix.save(path)
        if path.stat().st_size < 5000 or pix.width < 700 or pix.height < 900:
            raise RuntimeError({"render_failure": {"page": index + 1, "bytes": path.stat().st_size, "width": pix.width, "height": pix.height}})
        rendered.append({"page": index + 1, "path": path.relative_to(project).as_posix(), "bytes": path.stat().st_size, "width": pix.width, "height": pix.height})
    pdf_doc.close()
    report_qa = {
        "status": "passed",
        "docx": {"path": docx_path.relative_to(project).as_posix(), "bytes": docx_path.stat().st_size, "sha256": sha256_file(docx_path)},
        "pdf": {"path": pdf_path.relative_to(project).as_posix(), "bytes": pdf_path.stat().st_size, "sha256": sha256_file(pdf_path), "pages": len(reader.pages), "text_chars": text_chars, "rendered_pages": len(rendered)},
        "xlsx": {"path": xlsx_path.relative_to(project).as_posix(), "bytes": xlsx_path.stat().st_size, "sha256": sha256_file(xlsx_path), "sheets": list(datasets)},
        "rendered": rendered,
    }
    json_write(report_dir / "REPORT_QA.json", report_qa)
    return [docx_path, pdf_path, xlsx_path, report_dir / "REPORT_QA.json"] + [project / row["path"] for row in rendered]


def build_indexes(project: Path) -> dict[str, Any]:
    index_dir = project / "Indexes" / "Section 4 Session 2 Checkpoint 2"
    if index_dir.exists():
        shutil.rmtree(index_dir)
    index_dir.mkdir(parents=True)
    excluded_prefixes = {index_dir.relative_to(project).as_posix(), "Manifest/Section 4 Session 2 Checkpoint 2"}
    rows = []
    for path in sorted(project.rglob("*")):
        if not path.is_file():
            continue
        rel = path.relative_to(project).as_posix()
        if any(rel.startswith(prefix) for prefix in excluded_prefixes):
            continue
        rows.append({"path": rel, "bytes": path.stat().st_size, "sha256": sha256_file(path), "extension": path.suffix.lower(), "purpose": cp1.infer_purpose(path)})
    json_path = index_dir / "Source and Artifact Index.json"
    csv_path = index_dir / "Source and Artifact Index.csv"
    json_write(json_path, {"schema": "mrhpd-source-index-1.0", "generated_at": NOW, "file_count": len(rows), "total_bytes": sum(row["bytes"] for row in rows), "files": rows})
    csv_write(csv_path, rows, ["path", "bytes", "sha256", "extension", "purpose"])
    bit = index_dir / "Package Bit Index.sqlite"
    con = sqlite3.connect(bit)
    con.executescript(
        """
        CREATE TABLE file_index(file_id INTEGER PRIMARY KEY,path TEXT NOT NULL UNIQUE,bytes INTEGER NOT NULL,sha256 TEXT NOT NULL,extension TEXT,purpose TEXT,text_preview TEXT);
        CREATE TABLE container_member(container_member_id INTEGER PRIMARY KEY,container_path TEXT NOT NULL,member_path TEXT NOT NULL,member_bytes INTEGER,compressed_bytes INTEGER,member_crc INTEGER,UNIQUE(container_path,member_path));
        CREATE INDEX idx_file_extension ON file_index(extension);
        CREATE INDEX idx_container_member_path ON container_member(container_path,member_path);
        """
    )
    for row in rows:
        path = project / row["path"]
        preview = ""
        if path.suffix.lower() in {".txt", ".md", ".json", ".csv", ".py", ".html", ".yml", ".yaml"} and path.stat().st_size <= 4_000_000:
            preview = path.read_text(encoding="utf-8", errors="replace")[:25000]
        con.execute("INSERT INTO file_index(path,bytes,sha256,extension,purpose,text_preview) VALUES (?,?,?,?,?,?)", (row["path"], row["bytes"], row["sha256"], row["extension"], row["purpose"], preview))
        if path.suffix.lower() == ".zip":
            try:
                with zipfile.ZipFile(path) as zf:
                    cp1.safe_infos(zf)
                    for info in zf.infolist():
                        if not info.is_dir():
                            con.execute("INSERT OR REPLACE INTO container_member(container_path,member_path,member_bytes,compressed_bytes,member_crc) VALUES (?,?,?,?,?)", (row["path"], info.filename, info.file_size, info.compress_size, info.CRC))
            except Exception:
                pass
    con.commit()
    integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
    file_rows = int(con.execute("SELECT COUNT(*) FROM file_index").fetchone()[0])
    searches = {}
    for term in ["Response 68", "field coverage", "query coverage", "source governance", "Strep bovis", "Google Drive"]:
        searches[term] = int(con.execute("SELECT COUNT(*) FROM file_index WHERE path LIKE ? OR text_preview LIKE ?", (f"%{term}%", f"%{term}%")).fetchone()[0])
    member_rows = int(con.execute("SELECT COUNT(*) FROM container_member").fetchone()[0])
    con.close()
    if integrity != "ok" or file_rows != len(rows) or searches["Response 68"] < 1 or searches["field coverage"] < 1:
        raise RuntimeError({"bit_index": {"integrity": integrity, "file_rows": file_rows, "expected": len(rows), "searches": searches}})
    qa = {"status": "passed", "source_index_files": len(rows), "source_index_total_bytes": sum(row["bytes"] for row in rows), "bit_index_integrity": integrity, "bit_index_file_rows": file_rows, "container_member_rows": member_rows, "searches": searches, "source_json": json_path.relative_to(project).as_posix(), "source_csv": csv_path.relative_to(project).as_posix(), "bit_index": bit.relative_to(project).as_posix()}
    json_write(index_dir / "INDEX_QA.json", qa)
    return qa


def build_manifest(project: Path) -> dict[str, Any]:
    manifest_dir = project / "Manifest" / "Section 4 Session 2 Checkpoint 2"
    if manifest_dir.exists():
        shutil.rmtree(manifest_dir)
    manifest_dir.mkdir(parents=True)
    excluded = manifest_dir.relative_to(project).as_posix()
    rows = []
    for path in sorted(project.rglob("*")):
        if not path.is_file():
            continue
        rel = path.relative_to(project).as_posix()
        if rel.startswith(excluded):
            continue
        rows.append({"path": rel, "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    json_path = manifest_dir / "Package Manifest.json"
    csv_path = manifest_dir / "Package Manifest.csv"
    checksums = manifest_dir / "SHA256 Inventory.txt"
    json_write(json_path, {"schema": "mrhpd-package-manifest-1.0", "generated_at": NOW, "file_count": len(rows), "total_bytes": sum(row["bytes"] for row in rows), "files": rows})
    csv_write(csv_path, rows, ["path", "bytes", "sha256"])
    text_write(checksums, "".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
    mismatches = []
    for row in rows:
        path = project / row["path"]
        if not path.exists() or path.stat().st_size != row["bytes"] or sha256_file(path) != row["sha256"]:
            mismatches.append(row["path"])
    if mismatches:
        raise RuntimeError({"manifest_mismatches": mismatches[:50]})
    qa = {"status": "passed", "manifest_records": len(rows), "total_bytes": sum(row["bytes"] for row in rows), "mismatches": 0, "json": json_path.relative_to(project).as_posix(), "csv": csv_path.relative_to(project).as_posix(), "checksums": checksums.relative_to(project).as_posix()}
    json_write(manifest_dir / "MANIFEST_QA.json", qa)
    return qa


def build_apply_utility(manifest: dict[str, Any]) -> str:
    return f'''#!/usr/bin/env python3
import argparse,hashlib,json,re,shutil,sqlite3,subprocess,sys,tempfile,zipfile
from pathlib import Path,PurePosixPath
M={json.dumps(manifest,ensure_ascii=False)}
def sha(path):
    h=hashlib.sha256()
    with open(path,'rb') as f:
        for block in iter(lambda:f.read(1024*1024),b''): h.update(block)
    return h.hexdigest()
def safe(zf):
    names=zf.namelist()
    if len(names)!=len(set(names)): raise SystemExit('duplicate ZIP members')
    for name in names:
        p=PurePosixPath(name.replace('\\\\','/'))
        if name.startswith(('/', '\\\\')) or re.match(r'^[A-Za-z]:',name) or '..' in p.parts: raise SystemExit('unsafe ZIP path: '+name)
    bad=zf.testzip()
    if bad: raise SystemExit('ZIP CRC failure: '+bad)
def verify(path,size,digest):
    if not path.exists() or path.stat().st_size!=size or sha(path)!=digest: raise SystemExit('identity failure: '+str(path))
    with zipfile.ZipFile(path) as zf: safe(zf)
def extract(path,target):
    target.mkdir(parents=True,exist_ok=True)
    with zipfile.ZipFile(path) as zf: safe(zf); zf.extractall(target)
def find_snapshot(root):
    matches=[]
    for path in root.rglob('*.zip'):
        if path.stat().st_size==M['base_project']['bytes'] and sha(path)==M['base_project']['sha256']: matches.append(path)
    if len(matches)!=1: raise SystemExit('base project snapshot not found exactly once')
    return matches[0]
p=argparse.ArgumentParser()
p.add_argument('--base-response66-restore',type=Path,required=True)
p.add_argument('--output-dir',type=Path,required=True)
a=p.parse_args()
verify(a.base_response66_restore,M['base_restore']['bytes'],M['base_restore']['sha256'])
root=Path(__file__).resolve().parent.parent
with tempfile.TemporaryDirectory(prefix='mrhpd-r68-apply-') as td:
    td=Path(td); restore=td/'restore'; extract(a.base_response66_restore,restore); snapshot=find_snapshot(restore); verify(snapshot,M['base_project']['bytes'],M['base_project']['sha256'])
    project=td/'project'; extract(snapshot,project); roots=[x for x in project.iterdir() if x.is_dir()]; files=[x for x in project.iterdir() if x.is_file()]; source=roots[0] if len(roots)==1 and not files else project
    if a.output_dir.exists() and any(a.output_dir.iterdir()): raise SystemExit('output directory must be empty')
    a.output_dir.mkdir(parents=True,exist_ok=True)
    for item in source.iterdir():
        destination=a.output_dir/item.name
        shutil.copytree(item,destination) if item.is_dir() else shutil.copy2(item,destination)
    for rel in M['deleted_paths']:
        target=a.output_dir/rel
        if target.is_dir(): shutil.rmtree(target)
        else: target.unlink(missing_ok=True)
    for row in M['overlay_files']:
        src=root/'OVERLAY'/row['path']; dst=a.output_dir/row['path']; dst.parent.mkdir(parents=True,exist_ok=True); shutil.copy2(src,dst)
        if dst.stat().st_size!=row['bytes'] or sha(dst)!=row['sha256']: raise SystemExit('overlay identity failure: '+row['path'])
    for key,row in M['critical_files'].items():
        target=a.output_dir/row['path']
        if not target.exists() or target.stat().st_size!=row['bytes'] or sha(target)!=row['sha256']: raise SystemExit('critical identity failure: '+key)
    db=a.output_dir/M['critical_files']['database']['path']; con=sqlite3.connect(db)
    try:
        if con.execute('PRAGMA integrity_check').fetchone()[0]!='ok': raise SystemExit('SQLite integrity failure')
        if list(con.execute('PRAGMA foreign_key_check')): raise SystemExit('foreign-key failure')
        if con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R68'").fetchone()[0]!=1: raise SystemExit('Response 68 missing')
        if con.execute("SELECT COUNT(*) FROM section4_session2_field_coverage WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2' AND status!='passed'").fetchone()[0]!=0: raise SystemExit('field coverage failure')
        if con.execute("SELECT COUNT(*) FROM section4_session2_query_coverage WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2' AND status!='passed'").fetchone()[0]!=0: raise SystemExit('query coverage failure')
        if con.execute("SELECT COUNT(*) FROM section4_session2_source_governance WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2' AND status!='passed'").fetchone()[0]!=0: raise SystemExit('source governance failure')
        if con.execute("SELECT COUNT(*) FROM section4_session2_drift_resolution WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP2' AND status!='passed'").fetchone()[0]!=0: raise SystemExit('drift resolution failure')
    finally:
        con.close()
    audit=a.output_dir/M['critical_files']['application_audit']['path']; app=a.output_dir/M['critical_files']['application']['path']
    result=subprocess.run([sys.executable,str(audit),'--db',str(db),'--app',str(app)],text=True,capture_output=True)
    if result.returncode: raise SystemExit('application checkpoint audit failed: '+result.stderr[-3000:])
print(json.dumps({{'status':'passed','output_dir':str(a.output_dir),'database_sha256':sha(db),'response':68}},indent=2))
'''


def build_recovery_package(
    base_restore: Path,
    base_project_archive: Path,
    immutable_response66: Path,
    current: Path,
    critical: dict[str, Path],
    reports: list[Path],
    checkpoint_qa: dict[str, Any],
    dist: Path,
    work: Path,
) -> tuple[Path, dict[str, Any]]:
    package = work / "checkpoint2_recovery_package"
    overlay = package / "OVERLAY"
    tools = package / "TOOLS"
    overlay.mkdir(parents=True)
    tools.mkdir(parents=True)
    changed, deleted = cp1.compare_overlay(immutable_response66, current, overlay)
    critical_rows = {key: {"path": path.relative_to(current).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path)} for key, path in critical.items()}
    manifest = {
        "schema": "mrhpd-checkpoint-recovery-manifest-2.1",
        "generated_at": NOW,
        "version": PROJECT_VERSION,
        "section": "Remediation Section 4 of 5",
        "session": "Session 2 of 3",
        "checkpoint": "2 of 3",
        "response": 68,
        "base_restore": {"name": base_restore.name, "bytes": BASE_RESTORE_BYTES, "sha256": BASE_RESTORE_SHA256},
        "base_project": {"name": base_project_archive.name, "bytes": BASE_PROJECT_BYTES, "sha256": BASE_PROJECT_SHA256},
        "overlay_file_count": len(changed),
        "overlay_total_bytes": sum(row["bytes"] for row in changed),
        "overlay_files": changed,
        "deleted_paths": deleted,
        "critical_files": critical_rows,
        "accepted_predecessor_mutated": False,
        "requires_checkpoint1_package": False,
        "requires_conversation_reconstruction": False,
        "next_checkpoint": "Checkpoint 3 of 3 — complete Session 2 restore",
    }
    json_write(package / "RECOVERY_MANIFEST.json", manifest)
    json_write(package / "BASELINE_IDENTITY.json", {"response66_restore": manifest["base_restore"], "response66_project": manifest["base_project"], "accepted_predecessor_mutated": False})
    text_write(
        package / "RESTORE_READ_FIRST.md",
        f"""# Human Pathogen Database — Checkpoint Recovery Through Response 68

This is intermediate Checkpoint 2 recovery data for Remediation Section 4 Session 2. It applies directly to the exact complete restore through Response 66; the earlier Response 67 checkpoint package is not required.

## Required baseline

- Bytes: {BASE_RESTORE_BYTES}
- SHA-256: `{BASE_RESTORE_SHA256}`

## Apply

```bash
python TOOLS/apply_checkpoint_recovery.py \
  --base-response66-restore "<Complete Restore Through Response 66.zip>" \
  --output-dir "<restored project through Response 68>"
```

The utility verifies the baseline, project snapshot, every overlay file, SQLite integrity and foreign keys, Response 68, field coverage, query coverage, source governance, drift resolution, critical hashes, and the read-only application audit.

## Current state

- Section 4 of 5: CONTINUE
- Session 2 of 3: CONTINUE
- Checkpoint 2 of 3: COMPLETE
- Next: Checkpoint 3 of 3 — complete Session 2 restore
- User upload required: no
- Conversation reconstruction required: no
""",
    )
    text_write(tools / "apply_checkpoint_recovery.py", build_apply_utility(manifest))
    json_write(package / "CHECKPOINT_2_QA.json", checkpoint_qa)
    for report in reports:
        destination = package / "REPORTS" / report.name
        destination.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(report, destination)
    controls = {"CHECKPOINT_RECOVERY_CHECKSUMS.sha256"}
    rows = []
    for path in sorted(package.rglob("*")):
        if path.is_file() and path.name not in controls:
            rows.append({"path": path.relative_to(package).as_posix(), "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    text_write(package / "CHECKPOINT_RECOVERY_CHECKSUMS.sha256", "".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
    recovery = dist / (
        f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 Remediation Section 4 of 5 "
        f"Session 2 of 3 Checkpoint 2 of 3 RECOVERY DATA THROUGH RESPONSE 68 {STAMP}.zip"
    )
    with zipfile.ZipFile(recovery, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6, allowZip64=True) as zf:
        for path in sorted(package.rglob("*")):
            if path.is_file():
                zf.write(path, path.relative_to(package).as_posix())
    recovery_qa = verify_zip(recovery)
    if recovery_qa["bytes"] >= 104_857_600:
        raise RuntimeError({"recovery_exceeds_drive_connector_limit": recovery_qa})
    with tempfile.TemporaryDirectory(prefix="mrhpd-r68-clean-apply-") as td:
        extracted = Path(td) / "recovery"
        safe_extract(recovery, extracted)
        output = Path(td) / "applied"
        result = subprocess.run(
            [sys.executable, str(extracted / "TOOLS" / "apply_checkpoint_recovery.py"), "--base-response66-restore", str(base_restore), "--output-dir", str(output)],
            text=True,
            capture_output=True,
            timeout=2400,
        )
        if result.returncode != 0:
            raise RuntimeError({"clean_apply_failed": {"stdout": result.stdout[-25000:], "stderr": result.stderr[-25000:]}})
        clean_db = output / critical_rows["database"]["path"]
        clean_qa = database_qa(clean_db, output)
    verification = {
        "schema": "mrhpd-response68-checkpoint2-recovery-verification-1.0",
        "generated_at": NOW,
        "status": "passed",
        "recovery": recovery_qa,
        "base_restore": verify_zip(base_restore, BASE_RESTORE_BYTES, BASE_RESTORE_SHA256),
        "base_project": verify_zip(base_project_archive, BASE_PROJECT_BYTES, BASE_PROJECT_SHA256),
        "overlay_files": len(changed),
        "overlay_bytes": sum(row["bytes"] for row in changed),
        "deleted_paths": deleted,
        "clean_apply": "passed",
        "clean_database": clean_qa,
        "accepted_predecessor_mutated": False,
        "checkpoint_2_of_3_complete": True,
        "session_2_of_3_complete": False,
        "remediation_section_4_complete": False,
        "next": "Checkpoint 3 of 3 — complete Session 2 restore",
    }
    verification_path = dist / "MRHPD v3.0.0a Response 68 Checkpoint 2 Recovery Verification.json"
    json_write(verification_path, verification)
    text_write(dist / f"{recovery.name}.sha256.txt", f"{recovery_qa['sha256']}  {recovery.name}")
    return recovery, verification


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--response66-dir", type=Path, default=Path("response66_artifacts"))
    parser.add_argument("--response67-dir", type=Path, default=Path("response67_artifact"))
    parser.add_argument("--dist", type=Path, default=Path("dist_cp4_s2_cp2"))
    args = parser.parse_args()
    if args.dist.exists():
        shutil.rmtree(args.dist)
    args.dist.mkdir(parents=True)

    with tempfile.TemporaryDirectory(prefix="mrhpd-cp4-s2-cp2-") as td:
        work = Path(td)
        base_restore, restore_qa = cp1.reconstruct_response66(args.response66_dir, work)
        immutable_response66, _unused_mutable, source_qa = cp1.extract_base_project(base_restore, work)
        base_project_archive = find_one(work / "response66_restore", "*COMPLETE PROJECT THROUGH RESPONSE 66*.zip")
        verify_zip(base_project_archive, BASE_PROJECT_BYTES, BASE_PROJECT_SHA256)
        response67_project, response67_qa = reconstruct_response67(base_restore, args.response67_dir, work)
        immutable_response67 = work / "immutable_response67_project"
        shutil.copytree(response67_project, immutable_response67)
        current = work / "mutable_section4_session2_checkpoint2" / response67_project.name
        current.parent.mkdir(parents=True)
        shutil.copytree(response67_project, current)

        baseline = baseline_artifacts(immutable_response67)
        source_db = cp1.find_canonical_database(current)
        source_workbook, original_sheets = cp1.find_workbook(current)
        events = recovery_events()

        con = sqlite3.connect(source_db)
        field_rows, field_summary = field_level_coverage(con)
        governance_rows, governance_summary = source_governance_audit(con)
        con.close()

        db = synchronize_database(current, source_db, field_rows, [], governance_rows, events)
        con = sqlite3.connect(db)
        query_rows, query_summary = query_level_coverage(con)
        con.close()
        # Query rows require Response 68 and the copied database, so insert them after initial synchronization.
        con = sqlite3.connect(db)
        try:
            con.execute("BEGIN IMMEDIATE")
            checkpoint_code = "MRHPD-V3-CP4-S2-CP2"
            con.execute("DELETE FROM section4_session2_query_coverage WHERE checkpoint_code=?", (checkpoint_code,))
            for row in query_rows:
                con.execute(
                    """
                    INSERT INTO section4_session2_query_coverage
                    (checkpoint_code,query_key,query_text,query_type,result_count,expected_minimum,status,sample_json,notes,checked_at)
                    VALUES (?,?,?,?,?,?,?,?,?,?)
                    """,
                    (checkpoint_code, row["query_key"], row["query_text"], row["query_type"], row["result_count"], row["expected_minimum"], row["status"], json.dumps(row["sample_json"], ensure_ascii=False), row["notes"], NOW),
                )
            con.commit()
        except Exception:
            con.rollback()
            raise
        finally:
            con.close()

        publication_qa = cp1.verify_publication(current)
        app_paths, application_qa = synchronize_application(current, db)
        preliminary = {"field": field_summary, "query": query_summary, "governance": governance_summary, "project_diff": {"status": "pending"}}
        # Create a provisional workbook before drift classification so expected mutation is observable.
        provisional_drift = []
        workbook, workbook_qa = synchronize_workbook(current, source_workbook, original_sheets, field_rows, query_rows, governance_rows, provisional_drift, preliminary)
        current_artifacts = {"database": db, "workbook": workbook, "application": app_paths["application"], "publication": cp1.find_by_hash(current, "*Integrated Manuscript*.pdf", PUBLICATION_SHA256), "editable_assembly": cp1.find_by_hash(current, "*Editable Integrated Manuscript Assembly*.docx", EDITABLE_ASSEMBLY_SHA256)}
        drift_rows = critical_drift_rows(baseline, current_artifacts, immutable_response67, current)
        insert_drift_rows(db, drift_rows)
        project_diff_qa = project_diff(immutable_response67, current)
        summaries = {"field": field_summary, "query": query_summary, "governance": governance_summary, "project_diff": project_diff_qa}
        # Rebuild workbook with finalized drift records.
        workbook, workbook_qa = synchronize_workbook(current, source_workbook, original_sheets, field_rows, query_rows, governance_rows, drift_rows, summaries)
        current_artifacts["workbook"] = workbook
        # Refresh expected database/workbook drift after final copied derivatives are frozen.
        drift_rows = critical_drift_rows(baseline, current_artifacts, immutable_response67, current)
        insert_drift_rows(db, drift_rows)
        finalize_database(db, workbook_qa, application_qa, publication_qa)
        db_qa = database_qa(db, current)
        # Refresh current application state and rerun current audit against final database identity.
        state = app_paths["state"]
        state_data = load_json(state)
        state_data["database_sha256"] = db_qa["sha256"]
        state_data["workbook_sha256"] = workbook_qa["sha256"]
        json_write(state, state_data)
        audit_result = subprocess.run([sys.executable, str(app_paths["audit"]), "--db", str(db), "--app", str(app_paths["application"]), "--output", str(app_paths["audit_output"])], text=True, capture_output=True, timeout=300)
        if audit_result.returncode != 0:
            raise RuntimeError({"final_application_audit_failed": {"stdout": audit_result.stdout[-20000:], "stderr": audit_result.stderr[-12000:]}})
        application_qa["final_audit_stdout"] = audit_result.stdout[-15000:]
        application_qa["status"] = "passed"

        tracking_files = build_tracking(current, db)
        qa_dir = current / "QA" / "Section 4 Session 2" / "Checkpoint 2"
        qa_dir.mkdir(parents=True, exist_ok=True)
        json_write(qa_dir / "FIELD_LEVEL_COVERAGE.json", field_rows)
        csv_write(qa_dir / "FIELD_LEVEL_COVERAGE.csv", field_rows, ["table_name", "field_name", "declared_type", "row_count", "not_null_declared", "primary_key", "null_count", "blank_count", "distinct_count", "status", "notes"])
        json_write(qa_dir / "QUERY_LEVEL_COVERAGE.json", query_rows)
        csv_write(qa_dir / "QUERY_LEVEL_COVERAGE.csv", query_rows, ["query_key", "query_text", "query_type", "result_count", "expected_minimum", "status", "sample_json", "notes"])
        json_write(qa_dir / "SOURCE_GOVERNANCE_AUDIT.json", governance_rows)
        csv_write(qa_dir / "SOURCE_GOVERNANCE_AUDIT.csv", governance_rows, ["audit_key", "source_table", "row_count", "url_column", "missing_url_count", "nonempty_url_count", "year_column", "latest_year", "title_column", "missing_title_count", "authority_column", "missing_authority_count", "status", "details_json"])
        json_write(qa_dir / "DRIFT_RESOLUTION.json", drift_rows)
        csv_write(qa_dir / "DRIFT_RESOLUTION.csv", drift_rows, ["artifact_key", "baseline_path", "baseline_bytes", "baseline_sha256", "current_path", "current_bytes", "current_sha256", "drift_class", "resolution", "status"])
        json_write(qa_dir / "DRIFT_FILE_DIFF.json", project_diff_qa)
        json_write(qa_dir / "DATABASE_QA.json", db_qa)
        json_write(qa_dir / "WORKBOOK_QA.json", workbook_qa)
        json_write(qa_dir / "APPLICATION_QA.json", application_qa)
        json_write(qa_dir / "PUBLICATION_QA.json", publication_qa)

        report_files = build_reports(current, field_rows, query_rows, governance_rows, drift_rows, summaries)
        recovery_dir = current / "Recovery" / "Section 4 Session 2 Checkpoint 2"
        recovery_dir.mkdir(parents=True, exist_ok=True)
        json_write(recovery_dir / "RECOVERY_EVENTS_106_112.json", events)
        json_write(
            recovery_dir / "CHECKPOINT_STATE.json",
            {
                "schema": "mrhpd-section4-session2-checkpoint2-1.0",
                "created_at": NOW,
                "section": "Remediation Section 4 of 5",
                "session": "Session 2 of 3",
                "checkpoint": "2 of 3",
                "response": 68,
                "status": "COMPLETE",
                "database": db_qa,
                "workbook": workbook_qa,
                "application": application_qa,
                "publication": publication_qa,
                "field_coverage": field_summary,
                "query_coverage": query_summary,
                "source_governance": governance_summary,
                "drift": project_diff_qa,
                "accepted_predecessor_mutated": False,
                "next": "Checkpoint 3 of 3 — complete Session 2 restore",
            },
        )
        text_write(
            current / "README_SECTION4_SESSION2_CHECKPOINT2.md",
            """# Human Pathogen Database — Section 4 Session 2 Checkpoint 2

This copied current project state is synchronized through Response 68. It includes detailed logical-table and field coverage, controlled query and disambiguation regressions, source-governance reconciliation, expected/prohibited drift resolution, synchronized workbook and read-only application audits, tracking, indexes, manifests, reports, QA, and recovery controls.

- Checkpoint 2 of 3: COMPLETE
- Session 2 of 3: CONTINUE
- Next: Checkpoint 3 of 3 — complete Session 2 restore
- Accepted predecessor modified: no
- Frozen Section 3 release modified: no
""",
        )

        index_qa = build_indexes(current)
        manifest_qa = build_manifest(current)
        checkpoint_qa = {
            "schema": "mrhpd-section4-session2-checkpoint2-qa-1.0",
            "generated_at": NOW,
            "status": "passed",
            "response66_restore": restore_qa,
            "response66_source": source_qa,
            "response67_recovery": response67_qa,
            "database": db_qa,
            "workbook": workbook_qa,
            "application": application_qa,
            "publication": publication_qa,
            "field_coverage": field_summary,
            "query_coverage": query_summary,
            "source_governance": governance_summary,
            "critical_drift": drift_rows,
            "project_diff": project_diff_qa,
            "indexes": index_qa,
            "manifest": manifest_qa,
            "tracking_files": len(tracking_files),
            "report_files": len(report_files),
            "accepted_predecessor_mutated": False,
            "checkpoint_2_of_3_complete": True,
            "session_2_of_3_complete": False,
            "next": "Checkpoint 3 of 3 — complete Session 2 restore",
        }
        json_write(qa_dir / "CHECKPOINT_2_COMPLETE_QA.json", checkpoint_qa)
        index_qa = build_indexes(current)
        manifest_qa = build_manifest(current)
        checkpoint_qa["indexes"] = index_qa
        checkpoint_qa["manifest"] = manifest_qa
        json_write(qa_dir / "CHECKPOINT_2_COMPLETE_QA.json", checkpoint_qa)
        manifest_qa = build_manifest(current)
        checkpoint_qa["manifest"] = manifest_qa

        critical = {
            "database": db,
            "workbook": workbook,
            "application": app_paths["application"],
            "application_audit": app_paths["audit"],
            "publication": current_artifacts["publication"],
            "editable_assembly": current_artifacts["editable_assembly"],
            "checkpoint_qa": qa_dir / "CHECKPOINT_2_COMPLETE_QA.json",
        }
        recovery, verification = build_recovery_package(base_restore, base_project_archive, immutable_response66, current, critical, report_files, checkpoint_qa, args.dist, work)
        summary = {
            "schema": "mrhpd-response68-checkpoint2-build-summary-1.0",
            "generated_at": NOW,
            "status": "passed",
            "response": 68,
            "section": "Remediation Section 4 of 5",
            "session": "Session 2 of 3 CONTINUE",
            "checkpoint": "2 of 3 COMPLETE",
            "database": db_qa,
            "workbook": workbook_qa,
            "application": application_qa,
            "publication": publication_qa,
            "field_coverage": field_summary,
            "query_coverage": query_summary,
            "source_governance": governance_summary,
            "project_diff": project_diff_qa,
            "indexes": index_qa,
            "manifest": manifest_qa,
            "recovery": verification,
            "accepted_predecessor_mutated": False,
            "user_upload_required": False,
            "next": "Remediation Section 4 of 5 Session 2 of 3 Checkpoint 3 of 3 — complete Session 2 restore",
        }
        json_write(args.dist / "MRHPD_RESPONSE68_CHECKPOINT2_BUILD_SUMMARY.json", summary)
        print(
            json.dumps(
                {
                    "status": "passed",
                    "recovery_zip": recovery.name,
                    "recovery_bytes": recovery.stat().st_size,
                    "recovery_sha256": sha256_file(recovery),
                    "database_tables": db_qa["table_count"],
                    "workbook_sheets": workbook_qa["current_sheet_count"],
                    "logical_tables": field_summary["logical_tables"],
                    "audited_fields": field_summary["audited_fields"],
                    "queries": query_summary["queries"],
                    "source_governance_tables": governance_summary["audited_tables"],
                    "next": summary["next"],
                },
                indent=2,
            )
        )


if __name__ == "__main__":
    main()
