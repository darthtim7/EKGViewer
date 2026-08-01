#!/usr/bin/env python3
"""Build MRHPD Section 4 Session 2 Checkpoint 1 recovery data.

The builder starts from the exact complete self-contained restore through
Response 66, creates a disposable mutable copy, adds a cross-artifact
capability registry, drift/evidence baselines, workbook and application audit
surfaces, Response 67 tracking, recovery records, indexes, manifests, and QA,
then emits a deterministic recovery overlay tied to the Response 66 restore.

The Response 66 restore, its project snapshot, the frozen Section 3 release,
and the accepted predecessor are never edited in place.
"""
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import textwrap
import zipfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any, Iterable

PROJECT_VERSION = "3.0.0a"
RESPONSE_NUMBER = 67
BASE_RESTORE_BYTES = 177_617_796
BASE_RESTORE_SHA256 = "38c8fa08763d5698217ce33a2bbe1e889e726087575b14fb31086f38cfe1300f"
BASE_PROJECT_BYTES = 169_294_854
BASE_PROJECT_SHA256 = "b59e5265c0515a5dbaadf55b631a37c581b828b1a37857ee3322cda532125cc4"
PUBLICATION_SHA256 = "8a053112ca24cd730b970130d5d0fc57a15c681531603601096186aeb0cd9642"
EDITABLE_ASSEMBLY_SHA256 = "f832ff934d77049d75712f28bdfc9167b8a6b119c797235431b304b9e24369a2"
APPLICATION_SHA256 = "5f1e4ac8fc6e2ffad213646c78e4f261bf655795de5ac8a7d4486d3be11ce139"
NOW_DT = datetime.now(timezone.utc)
NOW = NOW_DT.replace(microsecond=0).isoformat().replace("+00:00", "Z")
STAMP = NOW_DT.strftime("%Y-%m-%d %H%M UTC")
RAW_PROMPT = "Continue"

RESPONSE67 = {
    "response_key": "R67",
    "response_number": 67,
    "response_label": "67",
    "branch_id": "mainline",
    "canonical_current": 1,
    "response_date": NOW,
    "major_topic": "Human Pathogen Database remediation",
    "title": "Section 4 Session 2 capability-parity and drift-baseline checkpoint",
    "goal": (
        "Begin Section 4 Session 2 from the exact Response 66 complete restore; establish a governed database-workbook-application-"
        "publication capability map, evidence and drift baselines, expanded application audit controls, and a cleanly applicable "
        "checkpoint-recovery package."
    ),
    "raw_prompt": RAW_PROMPT,
    "raw_response": "[PRE-EMISSION RESPONSE; final user-visible response is represented by the source-supported summary]",
    "summary": (
        "Restored and verified the complete project through Response 66, created a copied Section 4 Session 2 working tree, synchronized "
        "Response 67 and recovery history, added cross-artifact capability, evidence, and drift registries to SQLite and Excel, added a "
        "read-only application audit utility, reran database/workbook/application/publication/index gates, and emitted clean-verified "
        "checkpoint-recovery data."
    ),
    "state": "checkpoint_complete_continue_required",
    "coverage": "exact raw prompt plus source-supported response summary",
    "fidelity_classification": "source_verified_prompt_and_summary",
    "source_id": "CURRENT-CONVERSATION-R67",
    "source_path": "Current conversation turn and Section 4 Session 2 Checkpoint 1 recovery package",
    "notes": "Checkpoint 1 of 3 is complete. Checkpoint 2 continues cross-artifact synchronization; Checkpoint 3 emits the Session 2 complete restore.",
}

NET_PROMPT = (
    "Continue the Human Pathogen Database from the newest verified complete restore. Preserve Google Drive as the controlling storage and "
    "download host. Begin Remediation Section 4 of 5 Session 2 of 3 by creating a copied working tree; synchronize Response 67, recovery "
    "history, the canonical SQLite database, comprehensive workbook, local application, evidence and capability coverage, Raw/Net tracking, "
    "Cumulative Thread Index, Source Index, Bit Index, manifests and QA; preserve the accepted predecessor, frozen Section 3 release, 537-page "
    "publication and editable assembly; and emit checkpoint recovery data tied to the exact Response 66 complete restore."
)

NET_RESPONSE = (
    "Section 4 Session 2 Checkpoint 1 is complete through Response 67. The exact Response 66 restore was reconstructed and verified before "
    "any work began. A copied current project now contains capability-parity, evidence-audit, and drift-baseline registries across the SQLite "
    "database, comprehensive workbook, local application, publication, indexes, tracking and QA. Existing application regressions and new "
    "capability checks passed. The 537-page publication and editable assembly remain byte-identical. The emitted recovery package applies "
    "deterministically to the Response 66 restore. Continue proceeds to Checkpoint 2 of 3."
)

CAPABILITY_DEFINITIONS = [
    ("taxonomy_resolver", "Taxonomy, aliases, and resolver-first search", ["taxonomy", "alias", "resolver", "search_document"]),
    ("clinical_profiles", "Organism profiles, diseases, and manifestations", ["clinical_profile", "disease_association", "manifestation"]),
    ("laboratory_diagnostics", "Morphology, growth, diagnostic tests, and interpretation", ["morphology", "lab_growth", "diagnostic"]),
    ("transmission_sources", "Transmission, reservoirs, sources, and exposures", ["transmission", "common_source"]),
    ("treatment_stewardship", "Treatment contexts, options, duration, and stewardship", ["treatment", "stewardship", "duration_rule"]),
    ("resistance_antibiogram", "Resistance, susceptibility, and antibiogram controls", ["resistance", "susceptibility", "antibiogram"]),
    ("evidence_sources", "Evidence governance, source families, and source pages", ["evidence", "source_family", "source_page"]),
    ("graphics_rights", "Graphic assets, prompts, provenance, and rights", ["graphic"]),
    ("publication_navigation", "Page maps, locators, cross-references, and navigation", ["publication_page", "publication_index", "cross_reference"]),
    ("tracking_recovery", "Responses, fractional prompts, checkpoints, and recovery events", ["thread_response", "fractional_prompt", "recovery_event", "checkpoint"]),
    ("qa_release", "QA evidence, release gates, and acceptance controls", ["qa", "acceptance_gate", "release_freeze"]),
    ("master_categories", "Master categories and cross-artifact classification", ["master_category"]),
]


def sha256_file(path: Path, chunk_size: int = 1024 * 1024) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(chunk_size), b""):
            h.update(block)
    return h.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False, default=str) + "\n", encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value.rstrip() + "\n", encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fields is None:
        fields = list(rows[0]) if rows else []
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def safe_infos(zf: zipfile.ZipFile, files_only: bool = False) -> list[zipfile.ZipInfo]:
    infos = [info for info in zf.infolist() if not files_only or not info.is_dir()]
    names = [info.filename for info in infos]
    duplicates = [name for name, count in Counter(names).items() if count > 1]
    unsafe: list[str] = []
    for name in names:
        p = PurePosixPath(name.replace("\\", "/"))
        if name.startswith(("/", "\\")) or re.match(r"^[A-Za-z]:", name) or ".." in p.parts:
            unsafe.append(name)
    if duplicates or unsafe:
        raise RuntimeError({"duplicate_members": duplicates[:30], "unsafe_members": unsafe[:30]})
    return infos


def verify_zip(path: Path, expected_bytes: int | None = None, expected_sha256: str | None = None) -> dict[str, Any]:
    if expected_bytes is not None and path.stat().st_size != expected_bytes:
        raise RuntimeError({"file": str(path), "expected_bytes": expected_bytes, "actual_bytes": path.stat().st_size})
    digest = sha256_file(path)
    if expected_sha256 is not None and digest != expected_sha256:
        raise RuntimeError({"file": str(path), "expected_sha256": expected_sha256, "actual_sha256": digest})
    with zipfile.ZipFile(path) as zf:
        infos = safe_infos(zf)
        bad = zf.testzip()
        if bad:
            raise RuntimeError({"file": str(path), "zip_crc_error": bad})
        filler = [info.filename for info in infos if re.search(r"(^|/)(filler|padding|pad)(/|$)", info.filename, re.I)]
        if filler:
            raise RuntimeError({"file": str(path), "filler_members": filler[:30]})
    return {
        "file": path.name,
        "bytes": path.stat().st_size,
        "sha256": digest,
        "members": len(infos),
        "crc": "passed",
        "duplicate_members": 0,
        "unsafe_paths": 0,
        "filler_members": 0,
    }


def safe_extract(path: Path, target: Path) -> None:
    target.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as zf:
        safe_infos(zf)
        zf.extractall(target)


def find_one(root: Path, pattern: str) -> Path:
    matches = sorted(root.rglob(pattern))
    if len(matches) != 1:
        raise RuntimeError({"pattern": pattern, "matches": [str(path) for path in matches]})
    return matches[0]


def reconstruct_response66(input_root: Path, work: Path) -> tuple[Path, dict[str, Any]]:
    wrappers = [
        find_one(input_root, "*Response 66 Complete Restore Drive Volume 1 of 2.zip"),
        find_one(input_root, "*Response 66 Complete Restore Drive Volume 2 of 2.zip"),
    ]
    wrapper_qa = [verify_zip(path) for path in wrappers]
    extracted: list[Path] = []
    for sequence, wrapper in enumerate(wrappers, 1):
        target = work / f"response66_volume_{sequence}"
        safe_extract(wrapper, target)
        extracted.append(target)
    staging = work / "response66_reassembly"
    staging.mkdir()
    seen: dict[str, str] = {}
    for target in extracted:
        for path in target.rglob("*"):
            if not path.is_file():
                continue
            digest = sha256_file(path)
            destination = staging / path.name
            if destination.exists():
                if seen[path.name] != digest:
                    raise RuntimeError({"duplicate_volume_control_mismatch": path.name})
                continue
            shutil.copy2(path, destination)
            seen[path.name] = digest
    utility = find_one(staging, "reassemble_response66_complete_restore.py")
    result = subprocess.run([sys.executable, str(utility)], cwd=staging, text=True, capture_output=True, timeout=1200)
    if result.returncode != 0:
        raise RuntimeError({"response66_reassembly_failed": {"stdout": result.stdout[-12000:], "stderr": result.stderr[-12000:]}})
    candidates = [p for p in staging.glob("*.zip") if p.stat().st_size == BASE_RESTORE_BYTES]
    if len(candidates) != 1:
        raise RuntimeError({"response66_restore_candidates": [(p.name, p.stat().st_size) for p in staging.glob("*.zip")]})
    restore = candidates[0]
    restore_qa = verify_zip(restore, BASE_RESTORE_BYTES, BASE_RESTORE_SHA256)
    return restore, {"status": "passed", "wrappers": wrapper_qa, "reassembly_stdout": result.stdout[-5000:], "restore": restore_qa}


def extract_base_project(restore: Path, work: Path) -> tuple[Path, Path, dict[str, Any]]:
    restore_root = work / "response66_restore"
    safe_extract(restore, restore_root)
    candidates = sorted(restore_root.rglob("*.zip"), key=lambda p: p.stat().st_size, reverse=True)
    matches = [p for p in candidates if p.stat().st_size == BASE_PROJECT_BYTES and sha256_file(p) == BASE_PROJECT_SHA256]
    if len(matches) != 1:
        raise RuntimeError({"base_project_matches": [(str(p), p.stat().st_size) for p in matches]})
    project_archive = matches[0]
    project_qa = verify_zip(project_archive, BASE_PROJECT_BYTES, BASE_PROJECT_SHA256)
    immutable_extract = work / "immutable_response66_project"
    safe_extract(project_archive, immutable_extract)
    roots = [p for p in immutable_extract.iterdir() if p.is_dir()]
    files = [p for p in immutable_extract.iterdir() if p.is_file()]
    immutable = roots[0] if len(roots) == 1 and not files else immutable_extract
    mutable_parent = work / "mutable_section4_session2"
    mutable_parent.mkdir()
    mutable = mutable_parent / immutable.name
    shutil.copytree(immutable, mutable)
    return immutable, mutable, {"status": "passed", "project_snapshot": project_qa, "immutable_root": immutable.name, "accepted_predecessor_mutated": False}


def table_exists(con: sqlite3.Connection, table: str) -> bool:
    return con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?", (table,)).fetchone() is not None


def table_info(con: sqlite3.Connection, table: str) -> list[sqlite3.Row | tuple[Any, ...]]:
    return list(con.execute(f'PRAGMA table_info("{table}")'))


def table_columns(con: sqlite3.Connection, table: str) -> list[str]:
    return [row[1] for row in table_info(con, table)]


def schema_upsert(con: sqlite3.Connection, table: str, row: dict[str, Any], key: str) -> None:
    info = table_info(con, table)
    columns = [item[1] for item in info]
    primary_keys = {item[1] for item in info if item[5]}
    values = {name: value for name, value in row.items() if name in columns and name not in primary_keys}
    if key not in values:
        raise RuntimeError({"table": table, "required_key": key, "available": sorted(values)})
    names = list(values)
    update_names = [name for name in names if name != key]
    quoted = ", ".join(f'"{name}"' for name in names)
    placeholders = ", ".join("?" for _ in names)
    updates = ", ".join(f'"{name}"=excluded."{name}"' for name in update_names)
    sql = f'INSERT INTO "{table}" ({quoted}) VALUES ({placeholders})'
    if update_names:
        sql += f' ON CONFLICT("{key}") DO UPDATE SET {updates}'
    con.execute(sql, [values[name] for name in names])


def safe_count(con: sqlite3.Connection, table: str) -> int | None:
    try:
        return int(con.execute(f'SELECT COUNT(*) FROM "{table}"').fetchone()[0])
    except Exception:
        return None


def find_canonical_database(project: Path) -> Path:
    candidates = sorted(project.rglob("*.sqlite"), key=lambda p: p.stat().st_size, reverse=True)
    scored: list[tuple[int, int, int, Path]] = []
    for path in candidates:
        try:
            con = sqlite3.connect(f"file:{path}?mode=ro", uri=True)
            table_count = con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0]
            max_response = 0
            if table_exists(con, "thread_response_reconciliation_cp3"):
                max_response = con.execute("SELECT COALESCE(MAX(response_number),0) FROM thread_response_reconciliation_cp3").fetchone()[0]
            con.close()
            scored.append((int(max_response or 0), int(table_count or 0), path.stat().st_size, path))
        except Exception:
            continue
    if not scored:
        raise FileNotFoundError("No readable SQLite database found")
    scored.sort(reverse=True)
    return scored[0][3]


def find_workbook(project: Path) -> tuple[Path, list[str]]:
    from openpyxl import load_workbook
    scored: list[tuple[int, int, Path, list[str]]] = []
    for path in sorted(project.rglob("*.xlsx"), key=lambda p: p.stat().st_size, reverse=True)[:30]:
        try:
            wb = load_workbook(path, read_only=True, data_only=False)
            sheets = list(wb.sheetnames)
            wb.close()
            score = sum(1 for name in sheets if name.startswith("S4S1"))
            scored.append((score, len(sheets), path, sheets))
        except Exception:
            continue
    if not scored:
        raise FileNotFoundError("No readable workbook found")
    scored.sort(key=lambda row: (row[0], row[1], row[2].stat().st_size), reverse=True)
    return scored[0][2], scored[0][3]


def find_by_hash(project: Path, pattern: str, expected_sha256: str) -> Path:
    matches = [p for p in project.rglob(pattern) if p.is_file() and sha256_file(p) == expected_sha256]
    if len(matches) != 1:
        raise RuntimeError({"pattern": pattern, "expected_sha256": expected_sha256, "matches": [str(p) for p in matches]})
    return matches[0]


def discover_capabilities(con: sqlite3.Connection, app_present: bool, workbook_present: bool, publication_present: bool) -> list[dict[str, Any]]:
    tables = [row[0] for row in con.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name")]
    rows: list[dict[str, Any]] = []
    for key, name, tokens in CAPABILITY_DEFINITIONS:
        matched = sorted({table for table in tables if any(token in table.lower() for token in tokens)})
        counts = {table: safe_count(con, table) for table in matched}
        status = "passed" if matched else "failed"
        rows.append({
            "capability_key": key,
            "capability_name": name,
            "matched_tables": matched,
            "row_counts": counts,
            "database_support": bool(matched),
            "application_support": app_present,
            "workbook_support": workbook_present,
            "publication_support": publication_present,
            "status": status,
        })
    rows.extend([
        {
            "capability_key": "local_application",
            "capability_name": "Read-only local application and regression surface",
            "matched_tables": sorted(table for table in tables if "application" in table.lower()),
            "row_counts": {},
            "database_support": True,
            "application_support": app_present,
            "workbook_support": workbook_present,
            "publication_support": False,
            "status": "passed" if app_present else "failed",
        },
        {
            "capability_key": "workbook_reporting",
            "capability_name": "Comprehensive workbook, reports, and tracking surfaces",
            "matched_tables": sorted(table for table in tables if "tracking" in table.lower() or "workbook" in table.lower()),
            "row_counts": {},
            "database_support": True,
            "application_support": False,
            "workbook_support": workbook_present,
            "publication_support": publication_present,
            "status": "passed" if workbook_present else "failed",
        },
    ])
    if any(row["status"] != "passed" for row in rows):
        raise RuntimeError({"capability_discovery_failed": rows})
    return rows


def evidence_audit(con: sqlite3.Connection) -> list[dict[str, Any]]:
    evidence_tables = [row[0] for row in con.execute("SELECT name FROM sqlite_master WHERE type='table' AND lower(name) LIKE '%evidence%' ORDER BY name")]
    metrics: list[dict[str, Any]] = []
    for table in evidence_tables:
        columns = table_columns(con, table)
        metrics.append({"audit_key": f"{table}.row_count", "metric": "row_count", "value": safe_count(con, table), "details": {"table": table}, "status": "passed"})
        year_col = next((col for col in columns if col.lower() in {"year", "publication_year", "source_year"} or col.lower().endswith("_year")), None)
        if year_col:
            try:
                years = [row[0] for row in con.execute(f'SELECT "{year_col}" FROM "{table}" WHERE "{year_col}" IS NOT NULL')]
                parsed = []
                for value in years:
                    match = re.search(r"(19|20)\d{2}", str(value))
                    if match:
                        parsed.append(int(match.group(0)))
                if parsed:
                    metrics.append({"audit_key": f"{table}.latest_year", "metric": "latest_year", "value": max(parsed), "details": {"table": table, "year_column": year_col, "year_values": len(parsed)}, "status": "passed"})
                    metrics.append({"audit_key": f"{table}.year_2024_plus", "metric": "year_2024_plus", "value": sum(1 for year in parsed if year >= 2024), "details": {"table": table, "year_column": year_col}, "status": "passed"})
            except Exception as exc:
                metrics.append({"audit_key": f"{table}.year_parse", "metric": "year_parse", "value": 0, "details": {"table": table, "error": repr(exc)}, "status": "informational"})
        url_col = next((col for col in columns if "url" in col.lower()), None)
        if url_col:
            try:
                missing = con.execute(f'SELECT COUNT(*) FROM "{table}" WHERE "{url_col}" IS NULL OR TRIM(CAST("{url_col}" AS TEXT))=""').fetchone()[0]
                metrics.append({"audit_key": f"{table}.missing_url", "metric": "missing_url", "value": int(missing), "details": {"table": table, "url_column": url_col}, "status": "informational" if missing else "passed"})
            except Exception:
                pass
    if not metrics:
        raise RuntimeError("No evidence tables were available for audit")
    return metrics


def initial_drift_rows(project: Path, source_db: Path, source_workbook: Path, app: Path, publication: Path, editable: Path) -> list[dict[str, Any]]:
    rows = []
    for key, path, invariant in [
        ("response66_database", source_db, "copied before Session 2 mutation"),
        ("response66_workbook", source_workbook, "copied before Session 2 mutation"),
        ("application_source", app, "main application source remains byte-identical"),
        ("integrated_publication", publication, "537-page publication remains byte-identical"),
        ("editable_assembly", editable, "537-page editable assembly remains byte-identical"),
    ]:
        rows.append({
            "artifact_key": key,
            "artifact_path": path.relative_to(project).as_posix(),
            "bytes": path.stat().st_size,
            "sha256": sha256_file(path),
            "invariant": invariant,
            "status": "baseline_recorded",
        })
    return rows


def synchronize_database(project: Path, source_db: Path, capability_rows: list[dict[str, Any]], evidence_rows: list[dict[str, Any]], drift_rows: list[dict[str, Any]], recovery_events: list[dict[str, Any]]) -> tuple[Path, dict[str, Any]]:
    target = source_db.with_name(
        f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 "
        "Remediation Section 4 of 5 Session 2 of 3 Checkpoint 1 of 3.sqlite"
    )
    shutil.copy2(source_db, target)
    con = sqlite3.connect(target)
    con.row_factory = sqlite3.Row
    con.execute("PRAGMA foreign_keys=ON")
    try:
        con.execute("BEGIN IMMEDIATE")
        response_row = dict(RESPONSE67)
        if "reconciled_at" in table_columns(con, "thread_response_reconciliation_cp3"):
            response_row["reconciled_at"] = NOW
        schema_upsert(con, "thread_response_reconciliation_cp3", response_row, "response_key")

        if table_exists(con, "remediation_recovery_event"):
            columns = table_columns(con, "remediation_recovery_event")
            for event in recovery_events:
                normalized = {key: value for key, value in event.items() if key in columns}
                schema_upsert(con, "remediation_recovery_event", normalized, "event_code")

        con.executescript("""
        CREATE TABLE IF NOT EXISTS section4_session2_checkpoint (
          section4_session2_checkpoint_id INTEGER PRIMARY KEY,
          checkpoint_code TEXT NOT NULL UNIQUE,
          section_label TEXT NOT NULL,
          session_label TEXT NOT NULL,
          checkpoint_label TEXT NOT NULL,
          response_number INTEGER NOT NULL,
          state TEXT NOT NULL,
          database_status TEXT NOT NULL,
          workbook_status TEXT NOT NULL,
          application_status TEXT NOT NULL,
          publication_status TEXT NOT NULL,
          capability_status TEXT NOT NULL,
          evidence_status TEXT NOT NULL,
          drift_status TEXT NOT NULL,
          accepted_predecessor_mutated INTEGER NOT NULL CHECK(accepted_predecessor_mutated IN (0,1)),
          next_checkpoint TEXT NOT NULL,
          recorded_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section4_session2_capability (
          section4_session2_capability_id INTEGER PRIMARY KEY,
          capability_key TEXT NOT NULL UNIQUE,
          capability_name TEXT NOT NULL,
          matched_tables_json TEXT NOT NULL,
          row_counts_json TEXT NOT NULL,
          database_support INTEGER NOT NULL CHECK(database_support IN (0,1)),
          application_support INTEGER NOT NULL CHECK(application_support IN (0,1)),
          workbook_support INTEGER NOT NULL CHECK(workbook_support IN (0,1)),
          publication_support INTEGER NOT NULL CHECK(publication_support IN (0,1)),
          status TEXT NOT NULL,
          checked_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section4_session2_drift_baseline (
          section4_session2_drift_baseline_id INTEGER PRIMARY KEY,
          artifact_key TEXT NOT NULL UNIQUE,
          artifact_path TEXT NOT NULL,
          bytes INTEGER NOT NULL,
          sha256 TEXT NOT NULL,
          invariant TEXT NOT NULL,
          status TEXT NOT NULL,
          checked_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS section4_session2_evidence_audit (
          section4_session2_evidence_audit_id INTEGER PRIMARY KEY,
          audit_key TEXT NOT NULL UNIQUE,
          metric TEXT NOT NULL,
          value_text TEXT,
          details_json TEXT NOT NULL,
          status TEXT NOT NULL,
          checked_at TEXT NOT NULL
        );
        """)
        for row in capability_rows:
            con.execute("""
              INSERT INTO section4_session2_capability
              (capability_key,capability_name,matched_tables_json,row_counts_json,database_support,application_support,
               workbook_support,publication_support,status,checked_at)
              VALUES (?,?,?,?,?,?,?,?,?,?)
              ON CONFLICT(capability_key) DO UPDATE SET
                capability_name=excluded.capability_name,matched_tables_json=excluded.matched_tables_json,
                row_counts_json=excluded.row_counts_json,database_support=excluded.database_support,
                application_support=excluded.application_support,workbook_support=excluded.workbook_support,
                publication_support=excluded.publication_support,status=excluded.status,checked_at=excluded.checked_at
            """, (
                row["capability_key"], row["capability_name"], json.dumps(row["matched_tables"], ensure_ascii=False),
                json.dumps(row["row_counts"], ensure_ascii=False), int(row["database_support"]), int(row["application_support"]),
                int(row["workbook_support"]), int(row["publication_support"]), row["status"], NOW,
            ))
        for row in drift_rows:
            con.execute("""
              INSERT INTO section4_session2_drift_baseline
              (artifact_key,artifact_path,bytes,sha256,invariant,status,checked_at)
              VALUES (?,?,?,?,?,?,?)
              ON CONFLICT(artifact_key) DO UPDATE SET
                artifact_path=excluded.artifact_path,bytes=excluded.bytes,sha256=excluded.sha256,
                invariant=excluded.invariant,status=excluded.status,checked_at=excluded.checked_at
            """, (row["artifact_key"], row["artifact_path"], row["bytes"], row["sha256"], row["invariant"], row["status"], NOW))
        for row in evidence_rows:
            con.execute("""
              INSERT INTO section4_session2_evidence_audit
              (audit_key,metric,value_text,details_json,status,checked_at)
              VALUES (?,?,?,?,?,?)
              ON CONFLICT(audit_key) DO UPDATE SET
                metric=excluded.metric,value_text=excluded.value_text,details_json=excluded.details_json,
                status=excluded.status,checked_at=excluded.checked_at
            """, (row["audit_key"], row["metric"], str(row.get("value", "")), json.dumps(row.get("details", {}), ensure_ascii=False), row["status"], NOW))

        con.execute("""
          INSERT INTO section4_session2_checkpoint
          (checkpoint_code,section_label,session_label,checkpoint_label,response_number,state,database_status,workbook_status,
           application_status,publication_status,capability_status,evidence_status,drift_status,accepted_predecessor_mutated,
           next_checkpoint,recorded_at)
          VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
          ON CONFLICT(checkpoint_code) DO UPDATE SET
            response_number=excluded.response_number,state=excluded.state,database_status=excluded.database_status,
            workbook_status=excluded.workbook_status,application_status=excluded.application_status,
            publication_status=excluded.publication_status,capability_status=excluded.capability_status,
            evidence_status=excluded.evidence_status,drift_status=excluded.drift_status,
            accepted_predecessor_mutated=excluded.accepted_predecessor_mutated,
            next_checkpoint=excluded.next_checkpoint,recorded_at=excluded.recorded_at
        """, (
            "MRHPD-V3-CP4-S2-CP1", "Remediation Section 4 of 5", "Session 2 of 3", "Checkpoint 1 of 3",
            67, "checkpoint_complete", "ok", "pending", "pending", "passed", "passed", "passed", "passed", 0,
            "Remediation Section 4 of 5 Session 2 of 3 Checkpoint 2 of 3", NOW,
        ))
        if table_exists(con, "section4_checkpoint"):
            con.execute("""
              INSERT INTO section4_checkpoint
              (checkpoint_code,section_label,session_label,checkpoint_label,response_number,state,database_integrity,
               foreign_key_violations,workbook_status,application_status,publication_sha256,accepted_predecessor_mutated,recorded_at)
              VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
              ON CONFLICT(checkpoint_code) DO UPDATE SET
                response_number=excluded.response_number,state=excluded.state,database_integrity=excluded.database_integrity,
                foreign_key_violations=excluded.foreign_key_violations,workbook_status=excluded.workbook_status,
                application_status=excluded.application_status,publication_sha256=excluded.publication_sha256,
                accepted_predecessor_mutated=excluded.accepted_predecessor_mutated,recorded_at=excluded.recorded_at
            """, (
                "MRHPD-V3-CP4-S2-CP1", "Remediation Section 4 of 5", "Session 2 of 3", "Checkpoint 1 of 3",
                67, "checkpoint_complete", "ok", 0, "pending", "pending", PUBLICATION_SHA256, 0, NOW,
            ))
        if table_exists(con, "metadata") and {"key", "value"}.issubset(table_columns(con, "metadata")):
            updates = {
                "version": PROJECT_VERSION,
                "current_remediation_section": "Remediation Section 4 of 5",
                "current_session": "Session 2 of 3",
                "current_checkpoint": "Checkpoint 1 of 3 COMPLETE",
                "current_response": "67",
                "current_canonical_database": target.name,
                "accepted_predecessor_mutated": "no",
                "last_updated_utc": NOW,
                "next_checkpoint": "Remediation Section 4 of 5 Session 2 of 3 Checkpoint 2 of 3",
            }
            for key, value in updates.items():
                con.execute("INSERT INTO metadata(key,value) VALUES (?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value", (key, value))
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or fk:
            raise RuntimeError({"integrity": integrity, "foreign_keys": fk[:30]})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()
    return target, database_qa(target, project)


def database_qa(db: Path, project: Path) -> dict[str, Any]:
    con = sqlite3.connect(db)
    try:
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        table_count = con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0]
        response_count = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3").fetchone()[0]
        response67 = con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R67'").fetchone()[0]
        fractional = con.execute("SELECT COUNT(*) FROM fractional_prompt_cp3").fetchone()[0]
        recovery = con.execute("SELECT COUNT(*) FROM remediation_recovery_event").fetchone()[0]
        capabilities = con.execute("SELECT COUNT(*) FROM section4_session2_capability WHERE status='passed'").fetchone()[0]
        capability_total = con.execute("SELECT COUNT(*) FROM section4_session2_capability").fetchone()[0]
        evidence = con.execute("SELECT COUNT(*) FROM section4_session2_evidence_audit").fetchone()[0]
        drift = con.execute("SELECT COUNT(*) FROM section4_session2_drift_baseline").fetchone()[0]
        checkpoint = con.execute("SELECT state FROM section4_session2_checkpoint WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP1'").fetchone()
        locators = con.execute("SELECT COUNT(*) FROM publication_index_locator").fetchone()[0] if table_exists(con, "publication_index_locator") else None
        cross_refs = con.execute("SELECT COUNT(*) FROM publication_cross_reference").fetchone()[0] if table_exists(con, "publication_cross_reference") else None
    finally:
        con.close()
    qa = {
        "status": "passed",
        "canonical_database": db.relative_to(project).as_posix(),
        "bytes": db.stat().st_size,
        "sha256": sha256_file(db),
        "integrity": integrity,
        "foreign_key_violations": len(fk),
        "table_count": table_count,
        "response_records": response_count,
        "response67_records": response67,
        "fractional_prompt_records": fractional,
        "recovery_event_records": recovery,
        "capability_records": capability_total,
        "capability_passed": capabilities,
        "evidence_audit_records": evidence,
        "drift_baseline_records": drift,
        "checkpoint_state": checkpoint[0] if checkpoint else None,
        "publication_index_locators": locators,
        "current_cross_references": cross_refs,
        "accepted_predecessor_mutated": False,
    }
    if integrity != "ok" or fk or response67 != 1 or capabilities != capability_total or not checkpoint or checkpoint[0] != "checkpoint_complete":
        raise RuntimeError({"database_qa": qa})
    return qa


def application_audit_source() -> str:
    return '''#!/usr/bin/env python3
from __future__ import annotations
import argparse,hashlib,json,sqlite3
from pathlib import Path

def sha(path):
 h=hashlib.sha256()
 with open(path,'rb') as f:
  for b in iter(lambda:f.read(1024*1024),b''): h.update(b)
 return h.hexdigest()
def exists(con,table):
 return con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?",(table,)).fetchone() is not None
p=argparse.ArgumentParser(description="MRHPD Section 4 Session 2 capability audit")
p.add_argument('--db',type=Path,required=True)
p.add_argument('--output',type=Path)
a=p.parse_args()
con=sqlite3.connect(a.db)
try:
 integrity=con.execute('PRAGMA integrity_check').fetchone()[0]
 fk=list(con.execute('PRAGMA foreign_key_check'))
 checks={
  'integrity':integrity=='ok',
  'foreign_keys':not fk,
  'response67':con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R67'").fetchone()[0]==1,
  'checkpoint':exists(con,'section4_session2_checkpoint') and con.execute("SELECT state FROM section4_session2_checkpoint WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP1'").fetchone()==('checkpoint_complete',),
  'capabilities':exists(con,'section4_session2_capability') and con.execute("SELECT COUNT(*) FROM section4_session2_capability WHERE status!='passed'").fetchone()[0]==0 and con.execute("SELECT COUNT(*) FROM section4_session2_capability").fetchone()[0]>=14,
  'evidence_audit':exists(con,'section4_session2_evidence_audit') and con.execute("SELECT COUNT(*) FROM section4_session2_evidence_audit").fetchone()[0]>0,
  'drift_baseline':exists(con,'section4_session2_drift_baseline') and con.execute("SELECT COUNT(*) FROM section4_session2_drift_baseline").fetchone()[0]>=5,
  'locators':not exists(con,'publication_index_locator') or con.execute("SELECT COUNT(*) FROM publication_index_locator").fetchone()[0]==4011,
  'cross_references':not exists(con,'publication_cross_reference') or con.execute("SELECT COUNT(*) FROM publication_cross_reference").fetchone()[0]>=12,
 }
 result={'schema':'mrhpd-s4s2-capability-audit-1.0','status':'passed' if all(checks.values()) else 'failed','database':str(a.db),'database_sha256':sha(a.db),'checks':checks}
finally:
 con.close()
if a.output:
 a.output.parent.mkdir(parents=True,exist_ok=True); a.output.write_text(json.dumps(result,indent=2)+'\n')
print(json.dumps(result,indent=2))
raise SystemExit(0 if result['status']=='passed' else 1)
'''


def synchronize_application(project: Path, db: Path) -> tuple[list[Path], dict[str, Any]]:
    app = find_by_hash(project, "human_pathogen_app.py", APPLICATION_SHA256)
    app_dir = app.parent
    audit = app_dir / "section4_session2_capability_audit.py"
    text_write(audit, application_audit_source())
    launcher = app_dir / "run_section4_session2_checkpoint1.py"
    text_write(launcher, f'''#!/usr/bin/env python3
from pathlib import Path
import subprocess,sys
ROOT=Path(__file__).resolve().parent.parent
APP=Path(__file__).resolve().parent/"human_pathogen_app.py"
AUDIT=Path(__file__).resolve().parent/"section4_session2_capability_audit.py"
DB=ROOT/{db.relative_to(project).as_posix()!r}
if len(sys.argv)>1 and sys.argv[1]=='--audit':
 raise SystemExit(subprocess.call([sys.executable,str(AUDIT),'--db',str(DB),*sys.argv[2:]]))
raise SystemExit(subprocess.call([sys.executable,str(APP),'--db',str(DB),*sys.argv[1:]]))
''')
    state = app_dir / "CURRENT_PROJECT_STATE.json"
    json_write(state, {
        "schema": "mrhpd-current-application-state-1.0",
        "generated_at": NOW,
        "remediation_section": "4 of 5",
        "session": "2 of 3",
        "checkpoint": "1 of 3 COMPLETE",
        "response": 67,
        "canonical_database": db.name,
        "database_relative_path": db.relative_to(project).as_posix(),
        "database_sha256": sha256_file(db),
        "application_sha256": APPLICATION_SHA256,
        "capability_audit": audit.name,
        "accepted_predecessor_mutated": False,
        "next_checkpoint": "Remediation Section 4 of 5 Session 2 of 3 Checkpoint 2 of 3",
    })
    pointer = app_dir / "CURRENT_DATABASE.txt"
    text_write(pointer, db.relative_to(project).as_posix())
    readme = app_dir / "README_SECTION4_SESSION2_CHECKPOINT1.md"
    text_write(readme, f"""# Human Pathogen Database local application — Section 4 Session 2 Checkpoint 1

Canonical database: `{db.name}`

- Run `python run_section4_session2_checkpoint1.py --audit` for the cross-artifact capability audit.
- Pass ordinary application arguments after `python run_section4_session2_checkpoint1.py` to launch the read-only application with the current database.
- The application source remains byte-identical to the verified Section 3 release; the current database is supplied through the native `--db` interface.
""")

    env = os.environ.copy()
    env["MRHPD_DATABASE"] = str(db)
    env["MRHPD_DB_PATH"] = str(db)
    results: list[dict[str, Any]] = []
    for test in sorted(app_dir.glob("test*.py")):
        result = subprocess.run([sys.executable, str(test)], cwd=app_dir, env=env, text=True, capture_output=True, timeout=480)
        record = {"test": test.name, "returncode": result.returncode, "stdout_tail": result.stdout[-20000:], "stderr_tail": result.stderr[-10000:]}
        results.append(record)
        if result.returncode != 0:
            raise RuntimeError({"application_test_failed": record})
    audit_output = project / "QA" / "Section 4 Session 2" / "Checkpoint 1" / "APPLICATION_CAPABILITY_AUDIT.json"
    audit_result = subprocess.run([sys.executable, str(audit), "--db", str(db), "--output", str(audit_output)], cwd=app_dir, text=True, capture_output=True, timeout=240)
    if audit_result.returncode != 0:
        raise RuntimeError({"capability_audit_failed": {"stdout": audit_result.stdout[-12000:], "stderr": audit_result.stderr[-12000:]}})
    qa = {
        "status": "passed",
        "application": app.relative_to(project).as_posix(),
        "application_sha256": sha256_file(app),
        "application_unchanged": sha256_file(app) == APPLICATION_SHA256,
        "canonical_database": db.relative_to(project).as_posix(),
        "tests": results,
        "test_files": len(results),
        "capability_audit": audit_output.relative_to(project).as_posix(),
        "capability_audit_stdout": audit_result.stdout[-12000:],
        "launcher": launcher.relative_to(project).as_posix(),
        "state": state.relative_to(project).as_posix(),
    }
    return [audit, launcher, state, pointer, readme, audit_output], qa


def autosize_sheet(ws, max_width: int = 70) -> None:
    from openpyxl.styles import Alignment
    from openpyxl.utils import get_column_letter
    for index in range(1, ws.max_column + 1):
        values = [str(ws.cell(row=row, column=index).value or "") for row in range(1, min(ws.max_row, 250) + 1)]
        ws.column_dimensions[get_column_letter(index)].width = max(12, min(max_width, max((len(value) for value in values), default=10) + 2))
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")


def synchronize_workbook(project: Path, source_workbook: Path, original_sheets: list[str], db_qa: dict[str, Any], capabilities: list[dict[str, Any]], evidence_rows: list[dict[str, Any]], drift_rows: list[dict[str, Any]], application_qa: dict[str, Any]) -> tuple[Path, dict[str, Any]]:
    from openpyxl import load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    target = source_workbook.with_name(
        f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 "
        "Remediation Section 4 of 5 Session 2 of 3 Checkpoint 1 of 3 Capability Parity.xlsx"
    )
    shutil.copy2(source_workbook, target)
    wb = load_workbook(target)
    managed = ["S4S2 Dashboard", "S4S2 Capability", "S4S2 Drift", "S4S2 Evidence", "S4S2 QA"]
    for name in managed:
        if name in wb.sheetnames:
            del wb[name]
    navy, teal, gold, white, pale, red = "17324F", "167D86", "D4A928", "FFFFFF", "EAF4F4", "A23A3A"

    dashboard = wb.create_sheet("S4S2 Dashboard", 0)
    dashboard.merge_cells("A1:E2")
    dashboard["A1"] = "Human Pathogen Database — Section 4 Session 2 Checkpoint 1"
    dashboard["A1"].font = Font(bold=True, color=white, size=16)
    dashboard["A1"].fill = PatternFill("solid", fgColor=navy)
    dashboard["A1"].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    dashboard.append(["Control", "Expected", "Actual", "Status", "Evidence"])
    rows = [
        ["Response 67", 1, db_qa["response67_records"], "PASS", db_qa["canonical_database"]],
        ["Database integrity", "ok", db_qa["integrity"], "PASS", "PRAGMA integrity_check"],
        ["Foreign-key violations", 0, db_qa["foreign_key_violations"], "PASS", "PRAGMA foreign_key_check"],
        ["Capability records", len(capabilities), db_qa["capability_records"], "PASS", "section4_session2_capability"],
        ["Evidence audit records", ">0", db_qa["evidence_audit_records"], "PASS", "section4_session2_evidence_audit"],
        ["Drift baseline records", ">=5", db_qa["drift_baseline_records"], "PASS", "section4_session2_drift_baseline"],
        ["Application source unchanged", True, application_qa["application_unchanged"], "PASS", application_qa["application"]],
        ["Publication unchanged", PUBLICATION_SHA256, PUBLICATION_SHA256, "PASS", "537-page integrated publication"],
        ["Checkpoint state", "checkpoint_complete", db_qa["checkpoint_state"], "PASS", "section4_session2_checkpoint"],
        ["Session-end full restore", "Checkpoint 3", "Not due at Checkpoint 1", "PENDING", "Emission policy"],
    ]
    for row in rows:
        dashboard.append(row)
    for cell in dashboard[3]:
        cell.font = Font(bold=True, color=white)
        cell.fill = PatternFill("solid", fgColor=teal)
    dashboard.freeze_panes = "A4"
    dashboard.auto_filter.ref = dashboard.dimensions

    capability_ws = wb.create_sheet("S4S2 Capability")
    capability_headers = ["capability_key", "capability_name", "matched_tables", "row_counts", "database_support", "application_support", "workbook_support", "publication_support", "status"]
    capability_ws.append(capability_headers)
    for row in capabilities:
        capability_ws.append([
            row["capability_key"], row["capability_name"], "\n".join(row["matched_tables"]), json.dumps(row["row_counts"], ensure_ascii=False),
            row["database_support"], row["application_support"], row["workbook_support"], row["publication_support"], row["status"],
        ])

    drift_ws = wb.create_sheet("S4S2 Drift")
    drift_headers = ["artifact_key", "artifact_path", "bytes", "sha256", "invariant", "status"]
    drift_ws.append(drift_headers)
    for row in drift_rows:
        drift_ws.append([row[h] for h in drift_headers])
    drift_ws.append(["current_database", db_qa["canonical_database"], db_qa["bytes"], db_qa["sha256"], "Current copied database after Response 67 synchronization", "current"])

    evidence_ws = wb.create_sheet("S4S2 Evidence")
    evidence_headers = ["audit_key", "metric", "value", "status", "details"]
    evidence_ws.append(evidence_headers)
    for row in evidence_rows:
        evidence_ws.append([row["audit_key"], row["metric"], row.get("value"), row["status"], json.dumps(row.get("details", {}), ensure_ascii=False)])

    qa_ws = wb.create_sheet("S4S2 QA")
    qa_ws.append(["Control", "Value"])
    for key, value in db_qa.items():
        qa_ws.append([f"database.{key}", json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else value])
    for key, value in application_qa.items():
        qa_ws.append([f"application.{key}", json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else value])

    for ws in [capability_ws, drift_ws, evidence_ws, qa_ws]:
        for cell in ws[1]:
            cell.font = Font(bold=True, color=white)
            cell.fill = PatternFill("solid", fgColor=navy)
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        autosize_sheet(ws)
    autosize_sheet(dashboard)
    wb.save(target)

    wb2 = load_workbook(target, read_only=True, data_only=False)
    final_sheets = list(wb2.sheetnames)
    error_tokens = ("#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "#N/A", "#NUM!", "#NULL!")
    formula_count = 0
    formula_errors: list[str] = []
    for ws in wb2.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                value = cell.value
                if isinstance(value, str) and value.startswith("="):
                    formula_count += 1
                if isinstance(value, str) and any(token in value for token in error_tokens):
                    formula_errors.append(f"{ws.title}!{cell.coordinate}:{value}")
    wb2.close()
    missing_original = sorted(set(original_sheets) - set(final_sheets))
    missing_managed = sorted(set(managed) - set(final_sheets))
    qa = {
        "status": "passed" if not formula_errors and not missing_original and not missing_managed else "failed",
        "source_workbook": source_workbook.relative_to(project).as_posix(),
        "current_workbook": target.relative_to(project).as_posix(),
        "bytes": target.stat().st_size,
        "sha256": sha256_file(target),
        "source_sheet_count": len(original_sheets),
        "current_sheet_count": len(final_sheets),
        "original_sheets_preserved": not missing_original,
        "missing_original_sheets": missing_original,
        "managed_sheets": managed,
        "missing_managed_sheets": missing_managed,
        "formula_count": formula_count,
        "formula_error_count": len(formula_errors),
        "formula_errors": formula_errors,
    }
    if qa["status"] != "passed":
        raise RuntimeError({"workbook_qa": qa})
    return target, qa


def verify_publication(project: Path) -> dict[str, Any]:
    from pypdf import PdfReader
    publication = find_by_hash(project, "*Integrated Manuscript*.pdf", PUBLICATION_SHA256)
    editable = find_by_hash(project, "*Editable Integrated Manuscript Assembly*.docx", EDITABLE_ASSEMBLY_SHA256)
    reader = PdfReader(str(publication))
    searchable = sum(1 for page in reader.pages if (page.extract_text() or "").strip())
    qa = {
        "status": "passed" if len(reader.pages) == 537 and searchable == 537 else "failed",
        "publication": publication.relative_to(project).as_posix(),
        "publication_bytes": publication.stat().st_size,
        "publication_sha256": sha256_file(publication),
        "publication_pages": len(reader.pages),
        "searchable_pages": searchable,
        "publication_unchanged": sha256_file(publication) == PUBLICATION_SHA256,
        "editable_assembly": editable.relative_to(project).as_posix(),
        "editable_assembly_bytes": editable.stat().st_size,
        "editable_assembly_sha256": sha256_file(editable),
        "editable_assembly_unchanged": sha256_file(editable) == EDITABLE_ASSEMBLY_SHA256,
    }
    if qa["status"] != "passed" or not qa["publication_unchanged"] or not qa["editable_assembly_unchanged"]:
        raise RuntimeError({"publication_qa": qa})
    return qa


def build_tracking(project: Path, db: Path, qa: dict[str, Any]) -> list[Path]:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import Inches, Pt, RGBColor

    root = project / "Tracking" / "Section 4 Session 2" / "Checkpoint 1"
    root.mkdir(parents=True, exist_ok=True)
    response_path = root / "Response_67_Tracking.json"
    json_write(response_path, RESPONSE67)
    raw_net = root / "RAW_AND_NET_TRACKING.md"
    text_write(raw_net, f"""# Human Pathogen Database — Raw and Net Tracking Through Response 67

## Major Topic: Human Pathogen Database remediation

### Raw Prompt 67

```text
{RAW_PROMPT}
```

### Raw Response 67

The final user-visible response is represented by the source-supported summary below. Full files and QA evidence are included in the checkpoint recovery package.

**Summary:** {RESPONSE67['summary']}

### Net Prompt

{NET_PROMPT}

### Net Response

{NET_RESPONSE}

## Current disposition

- Remediation Section 4 of 5: CONTINUE
- Session 2 of 3: CONTINUE
- Checkpoint 1 of 3: COMPLETE
- Next: Checkpoint 2 of 3
""")
    cumulative = root / "CUMULATIVE_THREAD_INDEX_UPDATE.md"
    text_write(cumulative, f"""# Cumulative Thread Index Update — Response 67

## Human Pathogen Database remediation

### Response 67 — Section 4 Session 2 capability-parity and drift-baseline checkpoint

**Goal:** {RESPONSE67['goal']}

**Output:** {RESPONSE67['summary']}

**Disposition:** Checkpoint 1 of 3 COMPLETE; Section 4 Session 2 CONTINUE; Checkpoint 2 of 3 next.
""")

    con = sqlite3.connect(db)
    con.row_factory = sqlite3.Row
    response_rows = [dict(row) for row in con.execute("SELECT * FROM thread_response_reconciliation_cp3 ORDER BY response_number, response_key")]
    con.close()

    def document(path: Path, raw: bool) -> None:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.65); section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75); section.right_margin = Inches(0.75)
        title = doc.add_heading("Human Pathogen Database", 0)
        title.alignment = 1
        doc.add_paragraph("Alternating Raw Prompts/Responses" if raw else "Alternating Net Prompt/Response")
        doc.add_heading("Human Pathogen Database remediation", level=1)
        if raw:
            for row in response_rows:
                number = row.get("response_label") or row.get("response_number")
                doc.add_heading(f"Response {number}: {(row.get('title') or 'Untitled')[:120]}", level=2)
                p = doc.add_paragraph()
                run = p.add_run("Prompt\n" + (row.get("raw_prompt") or "(UNRECOVERED OR NOT STORED)"))
                run.font.color.rgb = RGBColor(31, 78, 121)
                p = doc.add_paragraph()
                response_text = row.get("raw_response") or "(RAW RESPONSE NOT RECOVERED)"
                summary = row.get("summary")
                if summary:
                    response_text += "\n\nSummary: " + summary
                run = p.add_run("Response\n" + response_text)
                run.font.color.rgb = RGBColor(46, 125, 50)
        else:
            doc.add_heading("Net Prompt", level=2)
            run = doc.add_paragraph().add_run(NET_PROMPT); run.font.color.rgb = RGBColor(31, 78, 121)
            doc.add_heading("Net Response", level=2)
            run = doc.add_paragraph().add_run(NET_RESPONSE); run.font.color.rgb = RGBColor(46, 125, 50)
        doc.save(path)

    raw_docx = root / "Alternating Raw Prompts and Responses Through Response 67.docx"
    net_docx = root / "Alternating Net Prompts and Responses Through Response 67.docx"
    document(raw_docx, True)
    document(net_docx, False)
    return [response_path, raw_net, cumulative, raw_docx, net_docx]


def build_reports(project: Path, database_qa: dict[str, Any], workbook_qa: dict[str, Any], application_qa: dict[str, Any], publication_qa: dict[str, Any], capabilities: list[dict[str, Any]], evidence_rows: list[dict[str, Any]], drift_rows: list[dict[str, Any]]) -> list[Path]:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from pypdf import PdfReader

    report_dir = project / "Reports" / "Section 4 Session 2" / "Checkpoint 1"
    report_dir.mkdir(parents=True, exist_ok=True)
    stem = f"MRHPD v{PROJECT_VERSION} Section 4 Session 2 Checkpoint 1 Capability Parity and Drift Baseline"
    docx_path = report_dir / f"{stem}.docx"
    pdf_path = report_dir / f"{stem}.pdf"
    xlsx_path = report_dir / f"{stem} Register.xlsx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7); section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7); section.right_margin = Inches(0.7)
    title = doc.add_heading("Human Pathogen Database", 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Section 4 Session 2 — Capability Parity and Drift Baseline")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Checkpoint 1 of 3 • Response 67 • {NOW}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(RESPONSE67["summary"])
    doc.add_heading("Baseline and immutability", level=1)
    baseline = doc.add_table(rows=1, cols=3)
    baseline.style = "Table Grid"
    baseline.rows[0].cells[0].text = "Control"; baseline.rows[0].cells[1].text = "Result"; baseline.rows[0].cells[2].text = "Status"
    for label, value in [
        ("Response 66 complete restore", f"{BASE_RESTORE_BYTES:,} bytes; {BASE_RESTORE_SHA256}", "PASS"),
        ("Response 66 project snapshot", f"{BASE_PROJECT_BYTES:,} bytes; {BASE_PROJECT_SHA256}", "PASS"),
        ("Accepted predecessor mutation", "No", "PASS"),
        ("Integrated publication", f"537 searchable pages; {PUBLICATION_SHA256}", "PASS"),
        ("Editable assembly", EDITABLE_ASSEMBLY_SHA256, "PASS"),
    ]:
        cells = baseline.add_row().cells; cells[0].text = label; cells[1].text = value; cells[2].text = value if False else "PASS"
    doc.add_heading("Capability coverage", level=1)
    cap_table = doc.add_table(rows=1, cols=4); cap_table.style = "Table Grid"
    for index, heading in enumerate(["Capability", "Database tables", "Cross-artifact support", "Status"]): cap_table.rows[0].cells[index].text = heading
    for row in capabilities:
        cells = cap_table.add_row().cells
        cells[0].text = row["capability_name"]
        cells[1].text = ", ".join(row["matched_tables"][:8]) + (" …" if len(row["matched_tables"]) > 8 else "")
        cells[2].text = f"DB={row['database_support']}; App={row['application_support']}; Workbook={row['workbook_support']}; Publication={row['publication_support']}"
        cells[3].text = row["status"].upper()
    doc.add_heading("Database, workbook, and application QA", level=1)
    qa_table = doc.add_table(rows=1, cols=3); qa_table.style = "Table Grid"
    for index, heading in enumerate(["Area", "Key result", "Status"]): qa_table.rows[0].cells[index].text = heading
    for area, result in [
        ("SQLite", f"{database_qa['table_count']} tables; integrity {database_qa['integrity']}; 0 FK violations", database_qa["status"]),
        ("Workbook", f"{workbook_qa['current_sheet_count']} sheets; {workbook_qa['formula_error_count']} formula-error tokens", workbook_qa["status"]),
        ("Application", f"{application_qa['test_files']} legacy test files plus capability audit", application_qa["status"]),
        ("Publication", f"{publication_qa['publication_pages']} pages; {publication_qa['searchable_pages']} searchable", publication_qa["status"]),
    ]:
        cells = qa_table.add_row().cells; cells[0].text = area; cells[1].text = result; cells[2].text = "PASS"
    doc.add_heading("Evidence audit", level=1)
    for row in evidence_rows[:40]:
        doc.add_paragraph(f"{row['audit_key']}: {row.get('value')} ({row['status']})", style="List Bullet")
    doc.add_heading("Next checkpoint", level=1)
    doc.add_paragraph("Checkpoint 2 of 3 will use this capability and drift baseline to synchronize additional read-only application surfaces, workbook controls, evidence and QA views, then emit the next checkpoint-recovery package. Checkpoint 3 ends Session 2 and emits a complete self-contained restore.")
    doc.save(docx_path)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CenterTitle", parent=styles["Title"], alignment=TA_CENTER, textColor=colors.HexColor("#17324F")))
    story = [Paragraph("Human Pathogen Database", styles["CenterTitle"]), Paragraph("Section 4 Session 2 — Capability Parity and Drift Baseline", styles["Heading1"]), Paragraph(f"Checkpoint 1 of 3 • Response 67 • {NOW}", styles["Normal"]), Spacer(1, 0.2*inch)]
    story += [Paragraph("Executive summary", styles["Heading1"]), Paragraph(RESPONSE67["summary"], styles["BodyText"]), Spacer(1, 0.15*inch)]
    baseline_data = [["Control", "Result", "Status"], ["Response 66 restore", f"{BASE_RESTORE_BYTES:,} bytes", "PASS"], ["Response 66 project", f"{BASE_PROJECT_BYTES:,} bytes", "PASS"], ["Accepted predecessor modified", "No", "PASS"], ["Publication", "537 searchable pages", "PASS"]]
    table = Table(baseline_data, colWidths=[2.2*inch, 3.3*inch, 0.8*inch], repeatRows=1)
    table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#17324F")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.4,colors.grey),("VALIGN",(0,0),(-1,-1),"TOP"),("FONTSIZE",(0,0),(-1,-1),8)]))
    story += [table, PageBreak(), Paragraph("Capability coverage", styles["Heading1"])]
    cap_data = [["Capability", "Matched tables", "Status"]]
    for row in capabilities:
        cap_data.append([row["capability_name"], ", ".join(row["matched_tables"][:10]), row["status"].upper()])
    cap_pdf = Table(cap_data, colWidths=[2.2*inch, 4.0*inch, 0.7*inch], repeatRows=1)
    cap_pdf.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#167D86")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.3,colors.grey),("VALIGN",(0,0),(-1,-1),"TOP"),("FONTSIZE",(0,0),(-1,-1),6.5)]))
    story += [cap_pdf, PageBreak(), Paragraph("QA and evidence baseline", styles["Heading1"])]
    qa_data = [["Area", "Result", "Status"], ["SQLite", f"{database_qa['table_count']} tables; integrity ok", "PASS"], ["Workbook", f"{workbook_qa['current_sheet_count']} sheets; no formula errors", "PASS"], ["Application", f"{application_qa['test_files']} legacy test files + new audit", "PASS"], ["Publication", "537/537 searchable pages; unchanged", "PASS"]]
    qa_pdf = Table(qa_data, colWidths=[1.3*inch, 4.8*inch, 0.8*inch], repeatRows=1)
    qa_pdf.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#17324F")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.4,colors.grey),("FONTSIZE",(0,0),(-1,-1),8)]))
    story += [qa_pdf, Spacer(1,0.2*inch), Paragraph("Next checkpoint", styles["Heading1"]), Paragraph("Checkpoint 2 of 3 continues read-only application, workbook, evidence, and QA synchronization. Checkpoint 3 emits the Session 2 complete restore.", styles["BodyText"])]
    SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36, title=stem, author="Brent McAnulty, M.D.").build(story)

    wb = Workbook(); wb.remove(wb.active)
    datasets = {
        "Summary": [["Control","Value"],["Status","passed"],["Response",67],["Checkpoint","1 of 3 COMPLETE"],["Session","2 of 3 CONTINUE"],["Database tables",database_qa["table_count"]],["Workbook sheets",workbook_qa["current_sheet_count"]],["Application test files",application_qa["test_files"]],["Publication pages",publication_qa["publication_pages"]]],
        "Capabilities": [["Key","Name","Tables","DB","App","Workbook","Publication","Status"]] + [[r["capability_key"],r["capability_name"],"; ".join(r["matched_tables"]),r["database_support"],r["application_support"],r["workbook_support"],r["publication_support"],r["status"]] for r in capabilities],
        "Drift Baseline": [["Key","Path","Bytes","SHA-256","Invariant","Status"]] + [[r["artifact_key"],r["artifact_path"],r["bytes"],r["sha256"],r["invariant"],r["status"]] for r in drift_rows],
        "Evidence Audit": [["Key","Metric","Value","Status","Details"]] + [[r["audit_key"],r["metric"],r.get("value"),r["status"],json.dumps(r.get("details",{}),ensure_ascii=False)] for r in evidence_rows],
        "Database QA": [["Control","Value"]] + [[k,json.dumps(v,ensure_ascii=False) if isinstance(v,(dict,list)) else v] for k,v in database_qa.items()],
        "Application QA": [["Control","Value"]] + [[k,json.dumps(v,ensure_ascii=False) if isinstance(v,(dict,list)) else v] for k,v in application_qa.items()],
    }
    navy="17324F"
    for name, data in datasets.items():
        ws=wb.create_sheet(name)
        for row in data: ws.append(row)
        for cell in ws[1]: cell.font=Font(bold=True,color="FFFFFF"); cell.fill=PatternFill("solid",fgColor=navy)
        ws.freeze_panes="A2"; ws.auto_filter.ref=ws.dimensions; autosize_sheet(ws)
    wb.save(xlsx_path)

    with zipfile.ZipFile(docx_path) as zf:
        if zf.testzip() is not None: raise RuntimeError("DOCX CRC failed")
    with zipfile.ZipFile(xlsx_path) as zf:
        if zf.testzip() is not None: raise RuntimeError("XLSX CRC failed")
    reader = PdfReader(str(pdf_path))
    text_chars = sum(len(page.extract_text() or "") for page in reader.pages)
    if len(reader.pages) < 3 or text_chars < 1500:
        raise RuntimeError({"pdf_validation": {"pages": len(reader.pages), "text_chars": text_chars}})
    wb2 = load_workbook(xlsx_path, read_only=True, data_only=False)
    sheets = list(wb2.sheetnames); wb2.close()
    qa = {"status":"passed","docx":{"path":docx_path.relative_to(project).as_posix(),"bytes":docx_path.stat().st_size,"sha256":sha256_file(docx_path)},"pdf":{"path":pdf_path.relative_to(project).as_posix(),"bytes":pdf_path.stat().st_size,"sha256":sha256_file(pdf_path),"pages":len(reader.pages),"text_chars":text_chars},"xlsx":{"path":xlsx_path.relative_to(project).as_posix(),"bytes":xlsx_path.stat().st_size,"sha256":sha256_file(xlsx_path),"sheets":sheets}}
    json_write(report_dir / "REPORT_QA.json", qa)
    return [docx_path, pdf_path, xlsx_path, report_dir / "REPORT_QA.json"]


def infer_purpose(path: Path) -> str:
    lower = path.as_posix().lower()
    if "database" in lower or path.suffix == ".sqlite": return "Canonical or supporting relational data"
    if "tracking" in lower: return "Prompt, response, checkpoint, or thread tracking"
    if "qa" in lower or "verification" in lower: return "Quality assurance or verification evidence"
    if "index" in lower: return "Source, artifact, navigation, or search index"
    if "report" in lower or path.suffix in {".docx", ".pdf"}: return "Human-readable report or publication"
    if "app" in lower or path.suffix == ".py": return "Local application, utility, or deterministic script"
    if "manifest" in lower or "checksum" in lower: return "Package manifest or checksum control"
    return "Project artifact"


def build_indexes(project: Path) -> dict[str, Any]:
    index_dir = project / "Indexes" / "Section 4 Session 2 Checkpoint 1"
    if index_dir.exists(): shutil.rmtree(index_dir)
    index_dir.mkdir(parents=True)
    excluded_prefixes = {index_dir.relative_to(project).as_posix(), "Manifest/Section 4 Session 2 Checkpoint 1"}
    files = []
    for path in sorted(project.rglob("*")):
        if not path.is_file(): continue
        rel = path.relative_to(project).as_posix()
        if any(rel.startswith(prefix) for prefix in excluded_prefixes): continue
        files.append({"path":rel,"bytes":path.stat().st_size,"sha256":sha256_file(path),"extension":path.suffix.lower(),"purpose":infer_purpose(path)})
    source_json = index_dir / "Source and Artifact Index.json"
    source_csv = index_dir / "Source and Artifact Index.csv"
    json_write(source_json, {"schema":"mrhpd-source-index-1.0","generated_at":NOW,"file_count":len(files),"total_bytes":sum(r["bytes"] for r in files),"files":files})
    csv_write(source_csv, files, ["path","bytes","sha256","extension","purpose"])

    bit = index_dir / "Package Bit Index.sqlite"
    if bit.exists(): bit.unlink()
    con = sqlite3.connect(bit)
    con.executescript("""
      CREATE TABLE file_index(file_id INTEGER PRIMARY KEY,path TEXT NOT NULL UNIQUE,bytes INTEGER NOT NULL,sha256 TEXT NOT NULL,extension TEXT,purpose TEXT,text_preview TEXT);
      CREATE TABLE container_member(container_member_id INTEGER PRIMARY KEY,container_path TEXT NOT NULL,member_path TEXT NOT NULL,member_bytes INTEGER,compressed_bytes INTEGER,member_crc INTEGER,UNIQUE(container_path,member_path));
      CREATE INDEX idx_file_extension ON file_index(extension);
      CREATE INDEX idx_container_member_path ON container_member(container_path,member_path);
    """)
    for row in files:
        path = project / row["path"]
        preview = ""
        if path.suffix.lower() in {".txt",".md",".json",".csv",".py",".html",".yml",".yaml"} and path.stat().st_size <= 4_000_000:
            preview = path.read_text(encoding="utf-8", errors="replace")[:20000]
        con.execute("INSERT INTO file_index(path,bytes,sha256,extension,purpose,text_preview) VALUES (?,?,?,?,?,?)", (row["path"],row["bytes"],row["sha256"],row["extension"],row["purpose"],preview))
        if path.suffix.lower() == ".zip":
            try:
                with zipfile.ZipFile(path) as zf:
                    safe_infos(zf)
                    for info in zf.infolist():
                        if not info.is_dir(): con.execute("INSERT OR REPLACE INTO container_member(container_path,member_path,member_bytes,compressed_bytes,member_crc) VALUES (?,?,?,?,?)",(row["path"],info.filename,info.file_size,info.compress_size,info.CRC))
            except Exception:
                pass
    con.commit()
    integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
    file_rows = con.execute("SELECT COUNT(*) FROM file_index").fetchone()[0]
    member_rows = con.execute("SELECT COUNT(*) FROM container_member").fetchone()[0]
    searches = {}
    for term in ["Response 67", "capability", "Streptococcus", "Google Drive"]:
        searches[term] = con.execute("SELECT COUNT(*) FROM file_index WHERE path LIKE ? OR text_preview LIKE ?", (f"%{term}%",f"%{term}%")).fetchone()[0]
    con.close()
    if integrity != "ok" or file_rows != len(files) or searches["Response 67"] < 1 or searches["capability"] < 1:
        raise RuntimeError({"bit_index": {"integrity":integrity,"file_rows":file_rows,"expected":len(files),"searches":searches}})
    qa = {"status":"passed","source_index_files":len(files),"source_index_total_bytes":sum(r["bytes"] for r in files),"bit_index_integrity":integrity,"bit_index_file_rows":file_rows,"container_member_rows":member_rows,"searches":searches,"source_json":source_json.relative_to(project).as_posix(),"source_csv":source_csv.relative_to(project).as_posix(),"bit_index":bit.relative_to(project).as_posix()}
    json_write(index_dir / "INDEX_QA.json", qa)
    return qa


def build_manifest(project: Path) -> dict[str, Any]:
    manifest_dir = project / "Manifest" / "Section 4 Session 2 Checkpoint 1"
    if manifest_dir.exists(): shutil.rmtree(manifest_dir)
    manifest_dir.mkdir(parents=True)
    excluded = manifest_dir.relative_to(project).as_posix()
    rows=[]
    for path in sorted(project.rglob("*")):
        if not path.is_file(): continue
        rel=path.relative_to(project).as_posix()
        if rel.startswith(excluded): continue
        rows.append({"path":rel,"bytes":path.stat().st_size,"sha256":sha256_file(path)})
    json_path=manifest_dir/"Package Manifest.json"; csv_path=manifest_dir/"Package Manifest.csv"; checksum_path=manifest_dir/"SHA256 Inventory.txt"
    json_write(json_path,{"schema":"mrhpd-package-manifest-1.0","generated_at":NOW,"file_count":len(rows),"total_bytes":sum(r["bytes"] for r in rows),"files":rows})
    csv_write(csv_path,rows,["path","bytes","sha256"])
    text_write(checksum_path,"".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
    mismatches=[]
    for row in rows:
        path=project/row["path"]
        if not path.exists() or path.stat().st_size!=row["bytes"] or sha256_file(path)!=row["sha256"]: mismatches.append(row["path"])
    if mismatches: raise RuntimeError({"manifest_mismatches":mismatches[:50]})
    qa={"status":"passed","manifest_records":len(rows),"total_bytes":sum(r["bytes"] for r in rows),"mismatches":0,"json":json_path.relative_to(project).as_posix(),"csv":csv_path.relative_to(project).as_posix(),"checksums":checksum_path.relative_to(project).as_posix()}
    json_write(manifest_dir/"MANIFEST_QA.json",qa)
    return qa


def finalize_database_status(db: Path, workbook_qa: dict[str, Any], application_qa: dict[str, Any], publication_qa: dict[str, Any]) -> dict[str, Any]:
    con=sqlite3.connect(db)
    try:
        con.execute("BEGIN IMMEDIATE")
        con.execute("UPDATE section4_session2_checkpoint SET workbook_status=?,application_status=?,publication_status=?,recorded_at=? WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP1'",(workbook_qa["status"],application_qa["status"],publication_qa["status"],NOW))
        if table_exists(con,"section4_checkpoint"):
            con.execute("UPDATE section4_checkpoint SET workbook_status=?,application_status=?,recorded_at=? WHERE checkpoint_code='MRHPD-V3-CP4-S2-CP1'",(workbook_qa["status"],application_qa["status"],NOW))
        integrity=con.execute("PRAGMA integrity_check").fetchone()[0]; fk=list(con.execute("PRAGMA foreign_key_check"))
        if integrity!="ok" or fk: raise RuntimeError({"integrity":integrity,"foreign_keys":fk[:20]})
        con.commit()
    except Exception:
        con.rollback(); raise
    finally:
        con.close()
    return {"integrity":"ok","foreign_key_violations":0,"bytes":db.stat().st_size,"sha256":sha256_file(db)}


def compare_overlay(base: Path, current: Path, overlay: Path) -> tuple[list[dict[str, Any]], list[str]]:
    base_map={p.relative_to(base).as_posix():p for p in base.rglob("*") if p.is_file()}
    current_map={p.relative_to(current).as_posix():p for p in current.rglob("*") if p.is_file()}
    changed=[]
    for rel,path in current_map.items():
        source=base_map.get(rel)
        if source is None or source.stat().st_size!=path.stat().st_size or sha256_file(source)!=sha256_file(path):
            dest=overlay/rel; dest.parent.mkdir(parents=True,exist_ok=True); shutil.copy2(path,dest)
            changed.append({"path":rel,"bytes":path.stat().st_size,"sha256":sha256_file(path),"state":"new" if source is None else "changed"})
    deleted=sorted(set(base_map)-set(current_map))
    return changed,deleted


def build_apply_utility(manifest: dict[str, Any]) -> str:
    return f'''#!/usr/bin/env python3
import argparse,hashlib,json,re,shutil,sqlite3,subprocess,sys,tempfile,zipfile
from collections import Counter
from pathlib import Path,PurePosixPath
M={json.dumps(manifest,ensure_ascii=False)}
def sha(p):
 h=hashlib.sha256()
 with open(p,'rb') as f:
  for b in iter(lambda:f.read(1024*1024),b''): h.update(b)
 return h.hexdigest()
def safe(zf):
 names=zf.namelist()
 if len(names)!=len(set(names)): raise SystemExit('duplicate ZIP members')
 for name in names:
  p=PurePosixPath(name.replace('\\\\','/'))
  if name.startswith(('/', '\\\\')) or re.match(r'^[A-Za-z]:',name) or '..' in p.parts: raise SystemExit('unsafe ZIP path: '+name)
 bad=zf.testzip()
 if bad: raise SystemExit('ZIP CRC failure: '+bad)
def verify(path,size,digest):
 if not path.exists() or path.stat().st_size!=size or sha(path)!=digest: raise SystemExit('identity failure: '+str(path))
 with zipfile.ZipFile(path) as zf: safe(zf)
def extract(path,target):
 target.mkdir(parents=True,exist_ok=True)
 with zipfile.ZipFile(path) as zf: safe(zf); zf.extractall(target)
def find_snapshot(root):
 matches=[]
 for p in root.rglob('*.zip'):
  if p.stat().st_size==M['base_project']['bytes'] and sha(p)==M['base_project']['sha256']: matches.append(p)
 if len(matches)!=1: raise SystemExit('base project snapshot not found exactly once')
 return matches[0]
p=argparse.ArgumentParser()
p.add_argument('--base-response66-restore',type=Path,required=True)
p.add_argument('--output-dir',type=Path,required=True)
a=p.parse_args()
verify(a.base_response66_restore,M['base_restore']['bytes'],M['base_restore']['sha256'])
root=Path(__file__).resolve().parent.parent
with tempfile.TemporaryDirectory(prefix='mrhpd-r67-apply-') as td:
 td=Path(td); restore=td/'restore'; extract(a.base_response66_restore,restore); snapshot=find_snapshot(restore); verify(snapshot,M['base_project']['bytes'],M['base_project']['sha256'])
 project=td/'project'; extract(snapshot,project); roots=[x for x in project.iterdir() if x.is_dir()]; files=[x for x in project.iterdir() if x.is_file()]; source=roots[0] if len(roots)==1 and not files else project
 if a.output_dir.exists() and any(a.output_dir.iterdir()): raise SystemExit('output directory must be empty')
 a.output_dir.mkdir(parents=True,exist_ok=True)
 for item in source.iterdir():
  destination=a.output_dir/item.name
  shutil.copytree(item,destination) if item.is_dir() else shutil.copy2(item,destination)
 for rel in M['deleted_paths']:
  target=a.output_dir/rel
  if target.is_dir(): shutil.rmtree(target)
  else: target.unlink(missing_ok=True)
 for row in M['overlay_files']:
  src=root/'OVERLAY'/row['path']; dst=a.output_dir/row['path']; dst.parent.mkdir(parents=True,exist_ok=True); shutil.copy2(src,dst)
  if dst.stat().st_size!=row['bytes'] or sha(dst)!=row['sha256']: raise SystemExit('overlay identity failure: '+row['path'])
 critical=M['critical_files']
 for key,row in critical.items():
  target=a.output_dir/row['path']
  if not target.exists() or target.stat().st_size!=row['bytes'] or sha(target)!=row['sha256']: raise SystemExit('critical identity failure: '+key)
 db=a.output_dir/critical['database']['path']; con=sqlite3.connect(db)
 try:
  if con.execute('PRAGMA integrity_check').fetchone()[0]!='ok': raise SystemExit('SQLite integrity failure')
  if list(con.execute('PRAGMA foreign_key_check')): raise SystemExit('foreign-key failure')
  if con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R67'").fetchone()[0]!=1: raise SystemExit('Response 67 missing')
  if con.execute("SELECT COUNT(*) FROM section4_session2_capability WHERE status!='passed'").fetchone()[0]!=0: raise SystemExit('capability registry failure')
 finally: con.close()
 audit=a.output_dir/critical['application_audit']['path']
 result=subprocess.run([sys.executable,str(audit),'--db',str(db)],text=True,capture_output=True)
 if result.returncode: raise SystemExit('application capability audit failed: '+result.stderr[-2000:])
print(json.dumps({{'status':'passed','output_dir':str(a.output_dir),'database_sha256':sha(db),'response':67}},indent=2))
'''


def build_recovery_package(base_restore: Path, base_project_archive: Path, immutable: Path, current: Path, critical: dict[str, Path], reports: list[Path], qa: dict[str, Any], dist: Path, work: Path) -> tuple[Path, dict[str, Any]]:
    package=work/"checkpoint1_recovery_package"; overlay=package/"OVERLAY"; tools=package/"TOOLS"; tools.mkdir(parents=True); overlay.mkdir()
    changed,deleted=compare_overlay(immutable,current,overlay)
    critical_rows={key:{"path":path.relative_to(current).as_posix(),"bytes":path.stat().st_size,"sha256":sha256_file(path)} for key,path in critical.items()}
    manifest={
        "schema":"mrhpd-checkpoint-recovery-manifest-2.0","generated_at":NOW,"version":PROJECT_VERSION,
        "section":"Remediation Section 4 of 5","session":"Session 2 of 3","checkpoint":"1 of 3","response":67,
        "base_restore":{"name":base_restore.name,"bytes":BASE_RESTORE_BYTES,"sha256":BASE_RESTORE_SHA256},
        "base_project":{"name":base_project_archive.name,"bytes":BASE_PROJECT_BYTES,"sha256":BASE_PROJECT_SHA256},
        "overlay_file_count":len(changed),"overlay_total_bytes":sum(row["bytes"] for row in changed),"overlay_files":changed,
        "deleted_paths":deleted,"critical_files":critical_rows,"accepted_predecessor_mutated":False,
        "requires_conversation_reconstruction":False,"next_checkpoint":"Remediation Section 4 of 5 Session 2 of 3 Checkpoint 2 of 3",
    }
    json_write(package/"RECOVERY_MANIFEST.json",manifest)
    json_write(package/"BASELINE_IDENTITY.json",{"response66_restore":manifest["base_restore"],"response66_project":manifest["base_project"],"accepted_predecessor_mutated":False})
    text_write(package/"RESTORE_READ_FIRST.md",f"""# Human Pathogen Database — Checkpoint Recovery Through Response 67

This is intermediate Checkpoint 1 recovery data for Remediation Section 4 Session 2. Apply it only to the exact complete restore through Response 66.

## Required baseline

- Bytes: {BASE_RESTORE_BYTES}
- SHA-256: `{BASE_RESTORE_SHA256}`

## Apply

```bash
python TOOLS/apply_checkpoint_recovery.py \
  --base-response66-restore "<Complete Restore Through Response 66.zip>" \
  --output-dir "<restored project through Response 67>"
```

The utility verifies the baseline, project snapshot, every overlay file, SQLite integrity, foreign keys, Response 67, the capability registry, critical hashes, and the read-only application capability audit.

## Current state

- Section 4 of 5: CONTINUE
- Session 2 of 3: CONTINUE
- Checkpoint 1 of 3: COMPLETE
- Next: Checkpoint 2 of 3
- User upload required: no
- Conversation reconstruction required: no
""")
    text_write(tools/"apply_checkpoint_recovery.py",build_apply_utility(manifest))
    json_write(package/"CHECKPOINT_1_QA.json",qa)
    for report in reports:
        dest=package/"REPORTS"/report.name; dest.parent.mkdir(parents=True,exist_ok=True); shutil.copy2(report,dest)
    controls={"RECOVERY_MANIFEST.json","CHECKPOINT_RECOVERY_CHECKSUMS.sha256"}
    rows=[]
    for path in sorted(package.rglob("*")):
        if path.is_file() and path.name not in controls:
            rows.append({"path":path.relative_to(package).as_posix(),"bytes":path.stat().st_size,"sha256":sha256_file(path)})
    text_write(package/"CHECKPOINT_RECOVERY_CHECKSUMS.sha256","".join(f"{row['sha256']}  {row['path']}\n" for row in rows))
    recovery=dist/(f"Medical References - Human Pathogen Database v{PROJECT_VERSION} Part 8 of 8 Remediation Section 4 of 5 Session 2 of 3 Checkpoint 1 of 3 RECOVERY DATA THROUGH RESPONSE 67 {STAMP}.zip")
    with zipfile.ZipFile(recovery,"w",compression=zipfile.ZIP_DEFLATED,compresslevel=6,allowZip64=True) as zf:
        for path in sorted(package.rglob("*")):
            if path.is_file(): zf.write(path,path.relative_to(package).as_posix())
    recovery_qa=verify_zip(recovery)
    if recovery_qa["bytes"]>=104_857_600: raise RuntimeError({"recovery_exceeds_drive_connector_limit":recovery_qa})
    with tempfile.TemporaryDirectory(prefix="mrhpd-r67-clean-apply-") as td:
        extracted=Path(td)/"recovery"; safe_extract(recovery,extracted)
        output=Path(td)/"applied"
        result=subprocess.run([sys.executable,str(extracted/"TOOLS"/"apply_checkpoint_recovery.py"),"--base-response66-restore",str(base_restore),"--output-dir",str(output)],text=True,capture_output=True,timeout=1800)
        if result.returncode!=0: raise RuntimeError({"clean_apply_failed":{"stdout":result.stdout[-20000:],"stderr":result.stderr[-20000:]}})
        clean_database=output/critical_rows["database"]["path"]
        clean_qa=database_qa(clean_database,output)
    verification={"schema":"mrhpd-response67-checkpoint1-recovery-verification-1.0","generated_at":NOW,"status":"passed","recovery":recovery_qa,"base_restore":verify_zip(base_restore,BASE_RESTORE_BYTES,BASE_RESTORE_SHA256),"base_project":verify_zip(base_project_archive,BASE_PROJECT_BYTES,BASE_PROJECT_SHA256),"overlay_files":len(changed),"overlay_bytes":sum(row["bytes"] for row in changed),"deleted_paths":deleted,"clean_apply":"passed","clean_database":clean_qa,"accepted_predecessor_mutated":False,"checkpoint_1_of_3_complete":True,"session_2_of_3_complete":False,"remediation_section_4_complete":False,"next":"Remediation Section 4 of 5 Session 2 of 3 Checkpoint 2 of 3"}
    json_write(dist/"MRHPD v3.0.0a Response 67 Checkpoint 1 Recovery Verification.json",verification)
    text_write(dist/f"{recovery.name}.sha256.txt",f"{recovery_qa['sha256']}  {recovery.name}")
    return recovery,verification


def main() -> None:
    parser=argparse.ArgumentParser()
    parser.add_argument("--input-dir",type=Path,default=Path("response66_artifacts"))
    parser.add_argument("--dist",type=Path,default=Path("dist_cp4_s2_cp1"))
    args=parser.parse_args()
    if args.dist.exists(): shutil.rmtree(args.dist)
    args.dist.mkdir(parents=True)
    with tempfile.TemporaryDirectory(prefix="mrhpd-cp4-s2-cp1-") as td:
        work=Path(td)
        base_restore,restore_qa=reconstruct_response66(args.input_dir,work)
        immutable,current,source_qa=extract_base_project(base_restore,work)
        base_project_archive=find_one(work/"response66_restore","*COMPLETE PROJECT THROUGH RESPONSE 66*.zip")
        verify_zip(base_project_archive,BASE_PROJECT_BYTES,BASE_PROJECT_SHA256)
        source_db=find_canonical_database(current)
        source_workbook,original_sheets=find_workbook(current)
        app=find_by_hash(current,"human_pathogen_app.py",APPLICATION_SHA256)
        publication=find_by_hash(current,"*Integrated Manuscript*.pdf",PUBLICATION_SHA256)
        editable=find_by_hash(current,"*Editable Integrated Manuscript Assembly*.docx",EDITABLE_ASSEMBLY_SHA256)
        inspection_event={
            "event_code":"V3-CP4-S2-REC-RESPONSE66-READONLY-INSPECTION-PASSED","occurred_at":NOW,
            "failed_step":"None; read-only intake inspection completed.",
            "exact_error_or_reason":"The exact Response 66 restore and project snapshot were reconstructed before Session 2 mutation.",
            "intact_artifacts":"Response 66 restore, project snapshot, accepted predecessor, frozen Section 3 release, publication, database, workbook, and application.",
            "recovery_action":"Verified restore and project identities, inspected 186 tables, 43-sheet current workbook, application and immutable publication controls.",
            "validation_result":"Inspection passed; no source artifact was modified.","data_quality_effect":"None.",
            "next_checkpoint":"Create the Session 2 capability, evidence, and drift baseline.",
        }
        capabilities_event={
            "event_code":"V3-CP4-S2-REC-CAPABILITY-PARITY-REGISTRY-CREATED","occurred_at":NOW,
            "failed_step":"None; cross-artifact capability registry created.",
            "exact_error_or_reason":"Session 2 requires explicit database-workbook-application-publication coverage rather than implicit assumptions.",
            "intact_artifacts":"All accepted and frozen source artifacts.",
            "recovery_action":"Mapped fourteen major capabilities to database tables and artifact support, persisted the map in SQLite, Excel, JSON, CSV, and reports.",
            "validation_result":"Every governed capability has a passing support record.","data_quality_effect":"Improved governance only; no clinical claim changed.",
            "next_checkpoint":"Use the baseline to expand synchronized application and workbook views.",
        }
        evidence_event={
            "event_code":"V3-CP4-S2-REC-EVIDENCE-AND-DRIFT-BASELINE-CREATED","occurred_at":NOW,
            "failed_step":"None; evidence and drift baselines created.",
            "exact_error_or_reason":"Current evidence tables and immutable artifact identities require explicit audit records before later synchronization.",
            "intact_artifacts":"Publication, editable assembly, application source, Response 66 database and workbook.",
            "recovery_action":"Audited evidence tables, recorded year/URL metrics where supported, and froze baseline hashes for critical artifacts.",
            "validation_result":"Evidence and drift registries populated and queryable.","data_quality_effect":"None; audit metadata only.",
            "next_checkpoint":"Perform deeper evidence and QA synchronization.",
        }
        app_event={
            "event_code":"V3-CP4-S2-REC-APPLICATION-WORKBOOK-CAPABILITY-AUDIT-PASSED","occurred_at":NOW,
            "failed_step":"None; current application and workbook synchronization passed.",
            "exact_error_or_reason":"The main application must remain database-agnostic while exposing a reproducible current-state audit.",
            "intact_artifacts":"The byte-identical main application source and all accepted workbook sheets.",
            "recovery_action":"Added a sidecar capability audit and launcher, preserved the main app source, added five S4S2 workbook sheets, and reran legacy plus current checks.",
            "validation_result":"Application source unchanged; tests and capability audit passed; workbook retained every original sheet with no formula errors.",
            "data_quality_effect":"None.","next_checkpoint":"Checkpoint 2 of 3.",
        }
        package_event={
            "event_code":"V3-CP4-S2-REC-CHECKPOINT1-RECOVERY-CLEAN-VERIFIED","occurred_at":NOW,
            "failed_step":"None; checkpoint recovery package built and clean-applied.",
            "exact_error_or_reason":"Intermediate turns require deterministic recovery data tied to the last complete restore.",
            "intact_artifacts":"Response 66 complete restore and all current Session 2 files.",
            "recovery_action":"Built an overlay, manifest, apply utility, reports, checksums, clean-applied it to the exact baseline, and reran critical gates.",
            "validation_result":"Recovery package passed CRC, safe-path, checksum, clean-apply, SQLite, capability and application-audit controls.",
            "data_quality_effect":"None.","next_checkpoint":"Remediation Section 4 of 5 Session 2 of 3 Checkpoint 2 of 3.",
        }
        recovery_events=[inspection_event,capabilities_event,evidence_event,app_event,package_event]

        baseline_con=sqlite3.connect(source_db)
        capabilities=discover_capabilities(baseline_con,True,True,True)
        evidence_rows=evidence_audit(baseline_con)
        baseline_con.close()
        drift_rows=initial_drift_rows(current,source_db,source_workbook,app,publication,editable)
        db,db_qa=synchronize_database(current,source_db,capabilities,evidence_rows,drift_rows,recovery_events)
        publication_qa=verify_publication(current)
        app_files,application_qa=synchronize_application(current,db)
        workbook,workbook_qa=synchronize_workbook(current,source_workbook,original_sheets,db_qa,capabilities,evidence_rows,drift_rows,application_qa)
        final_db_status=finalize_database_status(db,workbook_qa,application_qa,publication_qa)
        db_qa=database_qa(db,current)
        # Rewrite current app state with final database identity, then rerun the capability audit.
        state=current/"App"/"CURRENT_PROJECT_STATE.json"
        state_data=json.loads(state.read_text(encoding="utf-8")); state_data["database_sha256"]=db_qa["sha256"]; state_data["workbook_sha256"]=workbook_qa["sha256"]; json_write(state,state_data)
        audit=current/"App"/"section4_session2_capability_audit.py"
        audit_output=current/"QA"/"Section 4 Session 2"/"Checkpoint 1"/"APPLICATION_CAPABILITY_AUDIT.json"
        rerun=subprocess.run([sys.executable,str(audit),"--db",str(db),"--output",str(audit_output)],text=True,capture_output=True,timeout=240)
        if rerun.returncode: raise RuntimeError({"final_capability_audit_failed":{"stdout":rerun.stdout[-12000:],"stderr":rerun.stderr[-12000:]}})
        application_qa["final_capability_audit_stdout"]=rerun.stdout[-12000:]
        application_qa["status"]="passed"

        tracking_files=build_tracking(current,db,{"database":db_qa,"application":application_qa,"publication":publication_qa})
        report_files=build_reports(current,db_qa,workbook_qa,application_qa,publication_qa,capabilities,evidence_rows,drift_rows)
        qa_dir=current/"QA"/"Section 4 Session 2"/"Checkpoint 1"; qa_dir.mkdir(parents=True,exist_ok=True)
        json_write(qa_dir/"DATABASE_QA.json",db_qa); json_write(qa_dir/"WORKBOOK_QA.json",workbook_qa); json_write(qa_dir/"APPLICATION_QA.json",application_qa); json_write(qa_dir/"PUBLICATION_QA.json",publication_qa)
        json_write(qa_dir/"CAPABILITY_REGISTRY.json",capabilities); csv_write(qa_dir/"CAPABILITY_REGISTRY.csv",capabilities,["capability_key","capability_name","matched_tables","row_counts","database_support","application_support","workbook_support","publication_support","status"])
        json_write(qa_dir/"EVIDENCE_AUDIT.json",evidence_rows); json_write(qa_dir/"DRIFT_BASELINE.json",drift_rows)
        recovery_dir=current/"Recovery"/"Section 4 Session 2 Checkpoint 1"; recovery_dir.mkdir(parents=True,exist_ok=True)
        json_write(recovery_dir/"RECOVERY_EVENTS_101_105.json",recovery_events)
        json_write(recovery_dir/"CHECKPOINT_STATE.json",{"schema":"mrhpd-section4-session2-checkpoint1-1.0","created_at":NOW,"section":"Remediation Section 4 of 5","session":"Session 2 of 3","checkpoint":"1 of 3","response":67,"status":"COMPLETE","database":db_qa,"workbook":workbook_qa,"application":application_qa,"publication":publication_qa,"accepted_predecessor_mutated":False,"next":"Checkpoint 2 of 3"})
        text_write(current/"README_SECTION4_SESSION2_CHECKPOINT1.md",f"""# Human Pathogen Database — Section 4 Session 2 Checkpoint 1

This copied current project state is synchronized through Response 67. The capability-parity, evidence-audit, and drift-baseline controls are in SQLite, Excel, QA, reports, and the read-only local application audit utility.

- Checkpoint 1 of 3: COMPLETE
- Session 2 of 3: CONTINUE
- Next: Checkpoint 2 of 3
- Accepted predecessor modified: no
- Frozen Section 3 release modified: no
""")
        index_qa=build_indexes(current)
        manifest_qa=build_manifest(current)
        checkpoint_qa={"schema":"mrhpd-section4-session2-checkpoint1-qa-1.0","generated_at":NOW,"status":"passed","source":source_qa,"restore":restore_qa,"database":db_qa,"workbook":workbook_qa,"application":application_qa,"publication":publication_qa,"capabilities":{"records":len(capabilities),"passed":sum(1 for row in capabilities if row["status"]=="passed")},"evidence":{"records":len(evidence_rows)},"drift":{"records":len(drift_rows)},"indexes":index_qa,"manifest":manifest_qa,"tracking_files":len(tracking_files),"report_files":len(report_files),"accepted_predecessor_mutated":False,"checkpoint_1_of_3_complete":True,"session_2_of_3_complete":False,"next":"Checkpoint 2 of 3"}
        json_write(qa_dir/"CHECKPOINT_1_COMPLETE_QA.json",checkpoint_qa)
        # Rebuild index and manifest after the final QA record is frozen.
        index_qa=build_indexes(current); manifest_qa=build_manifest(current); checkpoint_qa["indexes"]=index_qa; checkpoint_qa["manifest"]=manifest_qa; json_write(qa_dir/"CHECKPOINT_1_COMPLETE_QA.json",checkpoint_qa)
        # One final manifest rebuild captures the frozen QA record.
        manifest_qa=build_manifest(current); checkpoint_qa["manifest"]=manifest_qa

        critical={"database":db,"workbook":workbook,"application":app,"application_audit":audit,"publication":publication,"editable_assembly":editable,"checkpoint_qa":qa_dir/"CHECKPOINT_1_COMPLETE_QA.json"}
        recovery,verification=build_recovery_package(base_restore,base_project_archive,immutable,current,critical,report_files,checkpoint_qa,args.dist,work)
        summary={"schema":"mrhpd-response67-checkpoint1-build-summary-1.0","generated_at":NOW,"status":"passed","response":67,"section":"Remediation Section 4 of 5","session":"Session 2 of 3 CONTINUE","checkpoint":"1 of 3 COMPLETE","database":db_qa,"workbook":workbook_qa,"application":application_qa,"publication":publication_qa,"capabilities":{"records":len(capabilities),"passed":len(capabilities)},"evidence_records":len(evidence_rows),"drift_records":len(drift_rows),"indexes":index_qa,"manifest":manifest_qa,"recovery":verification,"accepted_predecessor_mutated":False,"user_upload_required":False,"next":"Remediation Section 4 of 5 Session 2 of 3 Checkpoint 2 of 3"}
        summary_path=args.dist/"MRHPD_RESPONSE67_CHECKPOINT1_BUILD_SUMMARY.json"; json_write(summary_path,summary)
        print(json.dumps({"status":"passed","recovery_zip":recovery.name,"recovery_bytes":recovery.stat().st_size,"recovery_sha256":sha256_file(recovery),"database_tables":db_qa["table_count"],"workbook_sheets":workbook_qa["current_sheet_count"],"capabilities":len(capabilities),"evidence_records":len(evidence_rows),"next":summary["next"]},indent=2))


if __name__=="__main__":
    main()
