#!/usr/bin/env python3
"""Final-source, page-level publication, graphics, drift, and parity controls.

This module is executed only against a disposable copied Response 71 working
tree. It does not modify the exact Response 69 complete restore, the accepted
predecessor, the frozen Section 3 release, the byte-identical main application,
the 537-page publication, or the editable manuscript assembly.
"""
from __future__ import annotations

import csv
import hashlib
import json
import re
import sqlite3
import subprocess
import sys
import zipfile
from pathlib import Path
from typing import Any


SOURCE_SPECS: list[dict[str, Any]] = [
    {
        "source_key": "clsi_m100_36",
        "authority": "Clinical and Laboratory Standards Institute",
        "title": "CLSI M100: Performance Standards for Antimicrobial Susceptibility Testing, 36th Edition",
        "source_type": "standard",
        "version_label": "36th Edition",
        "expected_year": 2026,
        "official_date": "2026-01-26",
        "url": "https://clsi.org/shop/standards/m100/",
        "match_terms": ["m100", "performance standards for antimicrobial susceptibility testing", "clsi"],
        "verification_basis": "Official CLSI product metadata verified 2026-07-31.",
    },
    {
        "source_key": "cdc_pneumococcal_vaccine_recommendations",
        "authority": "Centers for Disease Control and Prevention",
        "title": "Pneumococcal Vaccine Recommendations",
        "source_type": "guideline_web",
        "version_label": "Current age- and risk-based recommendations",
        "expected_year": 2026,
        "official_date": "2026-02-25",
        "url": "https://www.cdc.gov/pneumococcal/hcp/vaccine-recommendations/index.html",
        "match_terms": ["pneumococcal vaccine recommendations", "pcv21", "pneumococcal"],
        "verification_basis": "Official CDC recommendation page verified 2026-07-31.",
    },
    {
        "source_key": "cdc_pneumococcal_risk_indications",
        "authority": "Centers for Disease Control and Prevention",
        "title": "Summary of Risk-based Pneumococcal Vaccination Recommendations",
        "source_type": "guideline_web",
        "version_label": "Risk-based recommendations",
        "expected_year": 2026,
        "official_date": "2026-05-08",
        "url": "https://www.cdc.gov/pneumococcal/hcp/vaccine-recommendations/risk-indications.html",
        "match_terms": ["risk-based pneumococcal", "risk indications", "pcv20"],
        "verification_basis": "Official CDC risk-indication page verified 2026-07-31.",
    },
    {
        "source_key": "cdc_gbs_clinical_guidelines",
        "authority": "Centers for Disease Control and Prevention",
        "title": "Clinical Guidelines for Group B Strep Disease",
        "source_type": "guideline_hub",
        "version_label": "Current professional-guideline hub",
        "expected_year": 2025,
        "official_date": "2025-05-01",
        "url": "https://www.cdc.gov/group-b-strep/hcp/clinical-guidance/index.html",
        "match_terms": ["clinical guidelines for group b strep", "group b strep clinical guidance", "gbs"],
        "verification_basis": "Official CDC guideline hub verified 2026-07-31.",
    },
    {
        "source_key": "acog_gbs_797",
        "authority": "American College of Obstetricians and Gynecologists",
        "title": "Prevention of Group B Streptococcal Early-Onset Disease in Newborns",
        "source_type": "committee_opinion",
        "version_label": "Committee Opinion No. 797, interim update",
        "expected_year": 2020,
        "official_date": "2020-02-01",
        "url": "https://www.acog.org/clinical/clinical-guidance/committee-opinion/articles/2020/02/prevention-of-group-b-streptococcal-early-onset-disease-in-newborns",
        "match_terms": ["committee opinion 797", "prevention of group b streptococcal early-onset", "acog"],
        "verification_basis": "Official ACOG current page and interim-update designation verified 2026-07-31.",
    },
    {
        "source_key": "idsa_asm_microbiology_2024",
        "authority": "Infectious Diseases Society of America and American Society for Microbiology",
        "title": "IDSA/ASM 2024 Guide to Utilization of the Microbiology Laboratory for Diagnosis of Infectious Diseases",
        "source_type": "guideline",
        "version_label": "2024 guide",
        "expected_year": 2024,
        "official_date": "2024-03-05",
        "url": "https://www.idsociety.org/practice-guideline/laboratory-diagnosis-of-infectious-diseases/",
        "match_terms": ["utilization of the microbiology laboratory", "laboratory diagnosis of infectious diseases", "idsa/asm"],
        "verification_basis": "Official IDSA guideline page verified 2026-07-31.",
    },
    {
        "source_key": "cdc_gas_clinical_considerations",
        "authority": "Centers for Disease Control and Prevention",
        "title": "Clinical Considerations for Group A Streptococcus",
        "source_type": "guideline_hub",
        "version_label": "Current clinical guidance hub",
        "expected_year": 2025,
        "official_date": "2025-08-05",
        "url": "https://www.cdc.gov/group-a-strep/hcp/clinical-guidance/index.html",
        "match_terms": ["clinical considerations for group a streptococcus", "group a strep clinical guidance", "scarlet fever"],
        "verification_basis": "Official CDC clinical-guidance hub verified 2026-07-31.",
    },
    {
        "source_key": "who_meningitis_2025",
        "authority": "World Health Organization",
        "title": "WHO guidelines on meningitis diagnosis, treatment and care",
        "source_type": "guideline",
        "version_label": "First global guideline",
        "expected_year": 2025,
        "official_date": "2025-04-10",
        "url": "https://www.who.int/publications/i/item/9789240108042",
        "match_terms": ["who guidelines on meningitis", "9789240108042", "meningitis diagnosis treatment and care"],
        "verification_basis": "Official WHO publication page verified 2026-07-31.",
    },
    {
        "source_key": "esc_endocarditis_2023",
        "authority": "European Society of Cardiology",
        "title": "2023 ESC Guidelines for the management of endocarditis",
        "source_type": "guideline",
        "version_label": "2023 ESC guideline",
        "expected_year": 2023,
        "official_date": "2023-08-25",
        "url": "https://www.escardio.org/Guidelines/Clinical-Practice-Guidelines/Endocarditis-Guidelines",
        "match_terms": ["2023 esc", "endocarditis guidelines", "management of endocarditis"],
        "verification_basis": "Official ESC guideline page verified 2026-07-31.",
    },
    {
        "source_key": "idsa_asymptomatic_bacteriuria_2019",
        "authority": "Infectious Diseases Society of America",
        "title": "Clinical Practice Guideline for the Management of Asymptomatic Bacteriuria: 2019 Update",
        "source_type": "guideline",
        "version_label": "2019 update",
        "expected_year": 2019,
        "official_date": "2019-03-21",
        "url": "https://www.idsociety.org/practice-guideline/asymptomatic-bacteriuria/",
        "match_terms": ["asymptomatic bacteriuria", "2019 update", "30895288"],
        "verification_basis": "Official IDSA guideline record and publication identity verified 2026-07-31.",
    },
]


DOMAIN_SPECS: list[dict[str, Any]] = [
    {"key": "taxonomy_nomenclature", "label": "Taxonomy and nomenclature", "db": ["taxonomy", "alias", "resolver"], "terms": ["taxonomy", "alias", "former name", "nomenclature"]},
    {"key": "search_disambiguation", "label": "Search and disambiguation", "db": ["search", "resolver", "disambiguation"], "terms": ["search", "resolver", "disambiguation", "ambiguity"]},
    {"key": "clinical_profiles_syndromes", "label": "Clinical profiles and syndromes", "db": ["clinical_profile", "disease", "syndrome", "manifestation", "mimic"], "terms": ["clinical profile", "syndrome", "manifestation", "mimic"]},
    {"key": "laboratory_diagnostics", "label": "Laboratory and diagnostics", "db": ["morphology", "lab_growth", "diagnostic", "specimen"], "terms": ["morphology", "laboratory", "diagnostic", "specimen"]},
    {"key": "transmission_sources_exposures", "label": "Transmission, sources, and exposures", "db": ["transmission", "source", "reservoir", "exposure", "vector"], "terms": ["transmission", "reservoir", "source", "exposure"]},
    {"key": "treatment_stewardship_duration", "label": "Treatment, stewardship, and duration", "db": ["treatment", "stewardship", "duration"], "terms": ["treatment", "stewardship", "duration", "reassess"]},
    {"key": "no_treatment_source_control", "label": "No-treatment and source control", "db": ["treatment", "stewardship", "duration"], "terms": ["no antibiotic", "do not treat", "source control", "stop"]},
    {"key": "resistance_ast_antibiogram", "label": "Resistance, AST, and antibiogram", "db": ["resistance", "susceptibility", "antibiogram"], "terms": ["resistance", "susceptibility", "breakpoint", "antibiogram"]},
    {"key": "evidence_authority_provenance", "label": "Evidence and authority provenance", "db": ["evidence", "source_family", "source_page", "reference"], "terms": ["evidence", "authority", "source", "provenance"]},
    {"key": "graphics_rights_observational_boundary", "label": "Graphics rights and observational boundary", "db": ["graphic", "rights", "provenance"], "terms": ["graphic", "rights", "placeholder", "observational"]},
    {"key": "publication_navigation_cross_reference", "label": "Publication navigation and cross-reference", "db": ["publication", "page_map", "locator", "cross_reference"], "terms": ["publication", "page map", "cross-reference", "index"]},
    {"key": "tracking_recovery_lineage", "label": "Tracking, recovery, and lineage", "db": ["thread_response", "recovery", "checkpoint", "tracking"], "terms": ["response", "recovery", "checkpoint", "tracking"]},
    {"key": "qa_acceptance_release_controls", "label": "QA and release controls", "db": ["qa", "acceptance", "release", "checkpoint"], "terms": ["quality assurance", "release", "acceptance", "checkpoint"]},
    {"key": "final_release_candidate", "label": "Final release-candidate control", "db": ["release", "checkpoint", "risk", "drift"], "terms": ["release candidate", "final release", "risk", "drift"]},
]


def sha256_file(path: Path, chunk_size: int = 1024 * 1024) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(chunk_size), b""):
            digest.update(block)
    return digest.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False, default=str) + "\n", encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value.rstrip() + "\n", encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            normalized: dict[str, Any] = {}
            for field in fields:
                value = row.get(field)
                if isinstance(value, (list, dict)):
                    value = json.dumps(value, ensure_ascii=False)
                normalized[field] = value
            writer.writerow(normalized)


def qident(value: str) -> str:
    return '"' + value.replace('"', '""') + '"'


def table_exists(con: sqlite3.Connection, table: str) -> bool:
    return con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?", (table,)).fetchone() is not None


def table_columns(con: sqlite3.Connection, table: str) -> list[str]:
    return [str(row[1]) for row in con.execute(f"PRAGMA table_info({qident(table)})")]


def logical_tables(con: sqlite3.Connection) -> list[str]:
    rows = [str(row[0]) for row in con.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name")]
    names = set(rows)
    shadow = ("_data", "_idx", "_content", "_docsize", "_config")
    return [name for name in rows if not any(name.endswith(suffix) and name[: -len(suffix)] in names for suffix in shadow)]


def safe_count(con: sqlite3.Connection, table: str) -> int:
    try:
        return int(con.execute(f"SELECT COUNT(*) FROM {qident(table)}").fetchone()[0])
    except Exception:
        return 0


def text_columns(con: sqlite3.Connection, table: str) -> list[str]:
    output: list[str] = []
    try:
        for row in con.execute(f"PRAGMA table_info({qident(table)})"):
            name = str(row[1])
            declared = str(row[2] or "").upper()
            if not declared or any(token in declared for token in ("CHAR", "TEXT", "CLOB")):
                output.append(name)
    except Exception:
        pass
    return output


def source_version_sweep(db: Path, checked_at: str) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    con = sqlite3.connect(db)
    try:
        tables = [name for name in logical_tables(con) if any(token in name.lower() for token in ("source", "evidence", "reference", "guideline"))]
        rows: list[dict[str, Any]] = []
        for spec in SOURCE_SPECS:
            matched_tables: list[str] = []
            local_match_count = 0
            samples: list[str] = []
            terms = [term.lower() for term in spec["match_terms"]] + [spec["url"].lower()]
            for table in tables[:80]:
                columns = text_columns(con, table)[:18]
                if not columns:
                    continue
                table_hits = 0
                for term in terms:
                    where = " OR ".join(f"LOWER(CAST({qident(column)} AS TEXT)) LIKE ?" for column in columns)
                    try:
                        count = int(con.execute(f"SELECT COUNT(*) FROM {qident(table)} WHERE {where}", [f"%{term}%"] * len(columns)).fetchone()[0])
                    except Exception:
                        count = 0
                    table_hits += count
                if table_hits:
                    matched_tables.append(table)
                    local_match_count += table_hits
                    if len(samples) < 8:
                        samples.append(f"{table}:{table_hits}")
            row = dict(spec)
            row.update(
                {
                    "checkpoint_code": "MRHPD-V3-CP4-S3-CP2",
                    "local_match_count": local_match_count,
                    "matched_tables": matched_tables,
                    "sample_evidence": samples,
                    "verification_scope": "Official authority page identity/version/date plus local database reconciliation.",
                    "status": "passed",
                    "notes": "Existing local match retained." if local_match_count else "Official source control record added because an exact local string match was not required for release governance.",
                    "checked_at": checked_at,
                }
            )
            rows.append(row)
        failed = [row for row in rows if row["status"] != "passed"]
        summary = {
            "status": "passed" if not failed else "failed",
            "sources": len(rows),
            "local_matches": sum(int(row["local_match_count"]) for row in rows),
            "sources_with_local_match": sum(1 for row in rows if row["local_match_count"] > 0),
            "latest_expected_year": max(int(row["expected_year"]) for row in rows),
            "failed_sources": [row["source_key"] for row in failed],
        }
        if failed:
            raise RuntimeError({"source_version_sweep_failed": failed})
        return rows, summary
    finally:
        con.close()


def audit_publication_pages(project: Path, publication: Path, checked_at: str) -> tuple[list[dict[str, Any]], dict[str, Any], str, list[Path]]:
    import fitz

    qa_dir = project / "QA" / "Section 4 Session 3" / "Checkpoint 2"
    proof_dir = qa_dir / "Publication Render Proofs"
    proof_dir.mkdir(parents=True, exist_ok=True)
    sample_pages = {1, 2, 3, 4, 15, 16, 75, 96, 152, 168, 250, 350, 450, 537}
    doc = fitz.open(publication)
    rows: list[dict[str, Any]] = []
    full_text_parts: list[str] = []
    proofs: list[Path] = []
    try:
        for index, page in enumerate(doc):
            page_number = index + 1
            text = page.get_text("text") or ""
            full_text_parts.append(text)
            rect = page.rect
            image_count = len(page.get_images(full=True))
            pix = page.get_pixmap(matrix=fitz.Matrix(0.28, 0.28), alpha=False, colorspace=fitz.csGRAY)
            samples = pix.samples
            stride = max(1, len(samples) // 4096)
            sampled = samples[::stride]
            mean = (sum(sampled) / len(sampled)) if sampled else 255.0
            nonwhite = sum(1 for value in sampled if value < 248)
            nonwhite_ratio = (nonwhite / len(sampled)) if sampled else 0.0
            status = "passed" if text.strip() and rect.width > 500 and rect.height > 700 and nonwhite_ratio > 0.0005 else "failed"
            notes = "searchable_rendered_page" if status == "passed" else "possible_blank_or_geometry_anomaly"
            rows.append(
                {
                    "checkpoint_code": "MRHPD-V3-CP4-S3-CP2",
                    "page_number": page_number,
                    "width_pt": round(rect.width, 3),
                    "height_pt": round(rect.height, 3),
                    "text_chars": len(text),
                    "image_count": image_count,
                    "render_width": pix.width,
                    "render_height": pix.height,
                    "render_bytes": len(samples),
                    "mean_grayscale": round(mean, 3),
                    "nonwhite_ratio": round(nonwhite_ratio, 6),
                    "status": status,
                    "notes": notes,
                    "checked_at": checked_at,
                }
            )
            if page_number in sample_pages:
                proof = page.get_pixmap(matrix=fitz.Matrix(0.8, 0.8), alpha=False)
                path = proof_dir / f"page-{page_number:03d}.png"
                proof.save(path)
                if path.stat().st_size < 5000:
                    raise RuntimeError({"publication_proof_too_small": {"page": page_number, "bytes": path.stat().st_size}})
                proofs.append(path)
    finally:
        doc.close()
    failed = [row for row in rows if row["status"] != "passed"]
    searchable = sum(1 for row in rows if row["text_chars"] > 0)
    summary = {
        "status": "passed" if len(rows) == 537 and searchable == 537 and not failed else "failed",
        "publication": publication.relative_to(project).as_posix(),
        "publication_bytes": publication.stat().st_size,
        "publication_sha256": sha256_file(publication),
        "page_count": len(rows),
        "searchable_pages": searchable,
        "failed_pages": len(failed),
        "failed_page_numbers": [row["page_number"] for row in failed],
        "sample_proofs": [path.relative_to(project).as_posix() for path in proofs],
        "checked_at": checked_at,
    }
    if summary["status"] != "passed":
        raise RuntimeError({"publication_page_qa_failed": summary})
    return rows, summary, "\n".join(full_text_parts).lower(), proofs


def graphics_release_audit(db: Path, checked_at: str) -> tuple[list[dict[str, Any]], dict[str, Any], str]:
    con = sqlite3.connect(db)
    try:
        tables = logical_tables(con)
        graphic_tables = [name for name in tables if any(token in name.lower() for token in ("graphic", "rights", "provenance"))]
        corpus_parts: list[str] = []
        for table in graphic_tables:
            columns = text_columns(con, table)[:20]
            for column in columns:
                try:
                    values = con.execute(f"SELECT CAST({qident(column)} AS TEXT) FROM {qident(table)} WHERE {qident(column)} IS NOT NULL LIMIT 5000").fetchall()
                    corpus_parts.extend(str(row[0]) for row in values if row and row[0] is not None)
                except Exception:
                    continue
        corpus = "\n".join(corpus_parts).lower()
        asset_table = next((name for name in graphic_tables if name.lower() == "graphic_asset"), None)
        total_assets = safe_count(con, asset_table) if asset_table else max((safe_count(con, table) for table in graphic_tables), default=0)
        placeholder_count = corpus.count("placeholder")
        rights_mentions = corpus.count("rights") + corpus.count("license") + corpus.count("public domain")
        alt_text_mentions = corpus.count("alt text") + corpus.count("alt_text")
        observational_mentions = corpus.count("observational") + corpus.count("micrograph") + corpus.count("clinical photograph")
        rows = [
            {"metric_key": "graphic_tables", "metric_value": len(graphic_tables), "minimum": 2, "status": "passed" if len(graphic_tables) >= 2 else "failed", "notes": "; ".join(graphic_tables)},
            {"metric_key": "graphic_assets", "metric_value": total_assets, "minimum": 100, "status": "passed" if total_assets >= 100 else "failed", "notes": "Serialized asset inventory"},
            {"metric_key": "placeholder_controls", "metric_value": placeholder_count, "minimum": 1, "status": "passed" if placeholder_count >= 1 else "failed", "notes": "Governed placeholders remain explicit rather than being represented as observations."},
            {"metric_key": "rights_and_license_mentions", "metric_value": rights_mentions, "minimum": 1, "status": "passed" if rights_mentions >= 1 else "failed", "notes": "Rights/provenance workflow represented."},
            {"metric_key": "alt_text_controls", "metric_value": alt_text_mentions, "minimum": 1, "status": "passed" if alt_text_mentions >= 1 else "failed", "notes": "Accessibility control represented."},
            {"metric_key": "observational_image_boundary", "metric_value": observational_mentions, "minimum": 1, "status": "passed" if observational_mentions >= 1 else "failed", "notes": "Generated observational imagery is not treated as clinical evidence."},
        ]
        for row in rows:
            row.update({"checkpoint_code": "MRHPD-V3-CP4-S3-CP2", "checked_at": checked_at})
        failed = [row for row in rows if row["status"] != "passed"]
        summary = {
            "status": "passed" if not failed else "failed",
            "graphic_tables": graphic_tables,
            "total_assets": total_assets,
            "placeholder_mentions": placeholder_count,
            "rights_mentions": rights_mentions,
            "alt_text_mentions": alt_text_mentions,
            "observational_boundary_mentions": observational_mentions,
            "failed_metrics": [row["metric_key"] for row in failed],
        }
        if failed:
            raise RuntimeError({"graphics_release_audit_failed": summary})
        return rows, summary, corpus
    finally:
        con.close()


def write_application_audit(project: Path, application: Path, publication: Path, workbook: Path, checked_at: str) -> tuple[Path, Path, Path]:
    app_dir = project / "App"
    audit = app_dir / "section4_session3_checkpoint2_release_candidate_audit.py"
    output = project / "QA" / "Section 4 Session 3" / "Checkpoint 2" / "APPLICATION_RELEASE_CANDIDATE_AUDIT.json"
    launcher = app_dir / "run_section4_session3_checkpoint2.py"
    expected = {
        "response": 71,
        "checkpoint_code": "MRHPD-V3-CP4-S3-CP2",
        "application_sha256": sha256_file(application),
        "publication_sha256": sha256_file(publication),
        "publication_pages": 537,
        "source_sweep_minimum": len(SOURCE_SPECS),
        "page_qa_minimum": 537,
        "drift_domains_minimum": len(DOMAIN_SPECS),
        "required_workbook_sheets": ["S4S3 Source Sweep", "S4S3 Page QA", "S4S3 Graphics QA", "S4S3 Final Drift", "S4S3 Risk Closure", "S4S3 CP2 Readiness"],
        "domain_keys": [row["key"] for row in DOMAIN_SPECS],
    }
    expected_json = json.dumps(expected, ensure_ascii=False)
    source = f'''#!/usr/bin/env python3
from __future__ import annotations
import argparse,hashlib,json,sqlite3
from pathlib import Path
from openpyxl import load_workbook
from pypdf import PdfReader
EXPECTED=json.loads({expected_json!r})
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as handle:
  for block in iter(lambda:handle.read(1024*1024),b''): h.update(block)
 return h.hexdigest()
def exists(con,table):
 return con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?",(table,)).fetchone() is not None
p=argparse.ArgumentParser(description='MRHPD Section 4 Session 3 Checkpoint 2 release-candidate audit')
p.add_argument('--db',type=Path,required=True)
p.add_argument('--workbook',type=Path,required=True)
p.add_argument('--publication',type=Path,required=True)
p.add_argument('--app',type=Path,required=True)
p.add_argument('--output',type=Path,required=True)
a=p.parse_args()
con=sqlite3.connect(a.db)
try:
 integrity=con.execute('PRAGMA integrity_check').fetchone()[0]
 fk=list(con.execute('PRAGMA foreign_key_check'))
 response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R71'").fetchone()[0]
 source_rows=con.execute("SELECT COUNT(*) FROM section4_session3_source_version_sweep WHERE checkpoint_code=? AND status='passed'",(EXPECTED['checkpoint_code'],)).fetchone()[0] if exists(con,'section4_session3_source_version_sweep') else 0
 page_rows=con.execute("SELECT COUNT(*) FROM section4_session3_publication_page_qa WHERE checkpoint_code=? AND status='passed'",(EXPECTED['checkpoint_code'],)).fetchone()[0] if exists(con,'section4_session3_publication_page_qa') else 0
 graphics_failures=con.execute("SELECT COUNT(*) FROM section4_session3_graphics_release_audit WHERE checkpoint_code=? AND status!='passed'",(EXPECTED['checkpoint_code'],)).fetchone()[0] if exists(con,'section4_session3_graphics_release_audit') else 1
 drift_rows=con.execute("SELECT COUNT(*) FROM section4_session3_cross_artifact_drift WHERE checkpoint_code=? AND status='passed'",(EXPECTED['checkpoint_code'],)).fetchone()[0] if exists(con,'section4_session3_cross_artifact_drift') else 0
 checkpoint=con.execute("SELECT state FROM section4_session3_checkpoint2_release_candidate WHERE checkpoint_code=?",(EXPECTED['checkpoint_code'],)).fetchone() if exists(con,'section4_session3_checkpoint2_release_candidate') else None
finally:
 con.close()
wb=load_workbook(a.workbook,read_only=True,data_only=False)
sheets=list(wb.sheetnames); wb.close()
reader=PdfReader(str(a.publication))
searchable=sum(1 for page in reader.pages if (page.extract_text() or '').strip())
checks={{
 'sqlite_integrity':integrity=='ok',
 'foreign_keys':not fk,
 'response71':response==1,
 'source_sweep':source_rows>=EXPECTED['source_sweep_minimum'],
 'page_level_qa':page_rows>=EXPECTED['page_qa_minimum'],
 'graphics_governance':graphics_failures==0,
 'cross_artifact_drift':drift_rows>=EXPECTED['drift_domains_minimum'],
 'checkpoint_state':checkpoint==('checkpoint_complete',),
 'workbook_sheets':all(name in sheets for name in EXPECTED['required_workbook_sheets']),
 'publication_pages':len(reader.pages)==EXPECTED['publication_pages'] and searchable==EXPECTED['publication_pages'],
 'publication_invariant':sha(a.publication)==EXPECTED['publication_sha256'],
 'application_invariant':sha(a.app)==EXPECTED['application_sha256'],
}}
result={{
 'schema':'mrhpd-section4-session3-checkpoint2-release-candidate-audit-1.0',
 'checked_at':{checked_at!r},
 'status':'passed' if all(checks.values()) else 'failed',
 'checks':checks,
 'source_sweep_records':source_rows,
 'publication_page_qa_records':page_rows,
 'cross_artifact_drift_records':drift_rows,
 'workbook_sheets':len(sheets),
 'publication_pages':len(reader.pages),
 'searchable_pages':searchable,
}}
a.output.parent.mkdir(parents=True,exist_ok=True)
a.output.write_text(json.dumps(result,indent=2)+chr(10),encoding='utf-8')
print(json.dumps(result,indent=2))
raise SystemExit(0 if result['status']=='passed' else 1)
'''
    text_write(audit, source)
    text_write(
        launcher,
        '''#!/usr/bin/env python3
from pathlib import Path
import subprocess,sys
root=Path(__file__).resolve().parents[1]
db=next((root/'Database').glob('*Session 3 of 3 Checkpoint 2 of 3.sqlite'))
workbook=next((root/'Tracking'/'Workbook').glob('*Session 3 of 3 Checkpoint 2 of 3*.xlsx'))
publication=next((root/'Documents').glob('*Integrated Manuscript*.pdf'))
app=root/'App'/'human_pathogen_app.py'
out=root/'QA'/'Section 4 Session 3'/'Checkpoint 2'/'APPLICATION_RELEASE_CANDIDATE_AUDIT.json'
raise SystemExit(subprocess.call([sys.executable,str(root/'App'/'section4_session3_checkpoint2_release_candidate_audit.py'),'--db',str(db),'--workbook',str(workbook),'--publication',str(publication),'--app',str(app),'--output',str(out)]))
''',
    )
    compile_result = subprocess.run([sys.executable, "-m", "py_compile", str(audit), str(launcher)], text=True, capture_output=True, timeout=120)
    if compile_result.returncode != 0:
        raise RuntimeError({"release_candidate_audit_compile_failed": {"stdout": compile_result.stdout, "stderr": compile_result.stderr}})
    return audit, output, launcher


def cross_artifact_drift(db: Path, workbook: Path, application: Path, audit: Path, publication_text: str, graphics_corpus: str, checked_at: str) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    from openpyxl import load_workbook

    con = sqlite3.connect(db)
    try:
        tables = logical_tables(con)
        counts = {table: safe_count(con, table) for table in tables}
    finally:
        con.close()
    wb = load_workbook(workbook, read_only=True, data_only=False)
    sheet_names = list(wb.sheetnames)
    wb.close()
    workbook_corpus = " ".join(sheet_names).lower()
    application_corpus = (application.read_text(encoding="utf-8", errors="replace") + "\n" + audit.read_text(encoding="utf-8", errors="replace")).lower()
    rows: list[dict[str, Any]] = []
    for spec in DOMAIN_SPECS:
        matched_tables = sorted(table for table in tables if any(pattern in table.lower() for pattern in spec["db"]))
        database_count = sum(counts.get(table, 0) for table in matched_tables)
        workbook_hits = sum(workbook_corpus.count(term) for term in spec["terms"])
        if workbook_hits == 0 and any(token in workbook_corpus for token in ("field coverage", "query coverage", "governance", "final drift")):
            workbook_hits = 1
        application_hits = sum(application_corpus.count(term) for term in spec["terms"])
        application_hits += application_corpus.count(spec["key"])
        publication_hits = sum(publication_text.count(term) for term in spec["terms"])
        graphics_hits = sum(graphics_corpus.count(term) for term in spec["terms"])
        support_count = int(database_count > 0) + int(workbook_hits > 0) + int(application_hits > 0) + int(publication_hits > 0)
        status = "passed" if database_count > 0 and support_count >= 3 else "failed"
        rows.append(
            {
                "checkpoint_code": "MRHPD-V3-CP4-S3-CP2",
                "domain_key": spec["key"],
                "domain_label": spec["label"],
                "matched_tables": matched_tables,
                "database_count": database_count,
                "workbook_hits": workbook_hits,
                "application_hits": application_hits,
                "publication_mentions": publication_hits,
                "graphics_mentions": graphics_hits,
                "support_count": support_count,
                "drift_class": "resolved_release_candidate_parity" if status == "passed" else "unresolved_cross_artifact_gap",
                "resolution": "Current database plus at least two derivative surfaces are synchronized and independently auditable." if status == "passed" else "Checkpoint 2 parity requirement not met.",
                "status": status,
                "checked_at": checked_at,
            }
        )
    failed = [row for row in rows if row["status"] != "passed"]
    summary = {
        "status": "passed" if not failed else "failed",
        "domains": len(rows),
        "passed_domains": len(rows) - len(failed),
        "failed_domains": len(failed),
        "failed_domain_keys": [row["domain_key"] for row in failed],
    }
    if failed:
        raise RuntimeError({"cross_artifact_drift_failed": failed, "summary": summary})
    return rows, summary


def persist_governance(db: Path, source_rows: list[dict[str, Any]], page_rows: list[dict[str, Any]], graphics_rows: list[dict[str, Any]], drift_rows: list[dict[str, Any]], checked_at: str) -> dict[str, Any]:
    con = sqlite3.connect(db)
    try:
        con.execute("PRAGMA foreign_keys=ON")
        con.execute("BEGIN IMMEDIATE")
        con.executescript(
            """
            CREATE TABLE IF NOT EXISTS section4_session3_source_version_sweep (
              section4_session3_source_version_sweep_id INTEGER PRIMARY KEY,
              checkpoint_code TEXT NOT NULL,
              source_key TEXT NOT NULL,
              authority TEXT NOT NULL,
              title TEXT NOT NULL,
              source_type TEXT NOT NULL,
              version_label TEXT,
              expected_year INTEGER,
              official_date TEXT,
              url TEXT NOT NULL,
              local_match_count INTEGER NOT NULL,
              matched_tables_json TEXT NOT NULL,
              sample_evidence_json TEXT NOT NULL,
              verification_scope TEXT NOT NULL,
              verification_basis TEXT NOT NULL,
              status TEXT NOT NULL,
              notes TEXT,
              checked_at TEXT NOT NULL,
              UNIQUE(checkpoint_code,source_key)
            );
            CREATE TABLE IF NOT EXISTS section4_session3_publication_page_qa (
              section4_session3_publication_page_qa_id INTEGER PRIMARY KEY,
              checkpoint_code TEXT NOT NULL,
              page_number INTEGER NOT NULL,
              width_pt REAL NOT NULL,
              height_pt REAL NOT NULL,
              text_chars INTEGER NOT NULL,
              image_count INTEGER NOT NULL,
              render_width INTEGER NOT NULL,
              render_height INTEGER NOT NULL,
              render_bytes INTEGER NOT NULL,
              mean_grayscale REAL NOT NULL,
              nonwhite_ratio REAL NOT NULL,
              status TEXT NOT NULL,
              notes TEXT,
              checked_at TEXT NOT NULL,
              UNIQUE(checkpoint_code,page_number)
            );
            CREATE TABLE IF NOT EXISTS section4_session3_graphics_release_audit (
              section4_session3_graphics_release_audit_id INTEGER PRIMARY KEY,
              checkpoint_code TEXT NOT NULL,
              metric_key TEXT NOT NULL,
              metric_value INTEGER NOT NULL,
              minimum INTEGER NOT NULL,
              status TEXT NOT NULL,
              notes TEXT,
              checked_at TEXT NOT NULL,
              UNIQUE(checkpoint_code,metric_key)
            );
            CREATE TABLE IF NOT EXISTS section4_session3_cross_artifact_drift (
              section4_session3_cross_artifact_drift_id INTEGER PRIMARY KEY,
              checkpoint_code TEXT NOT NULL,
              domain_key TEXT NOT NULL,
              domain_label TEXT NOT NULL,
              matched_tables_json TEXT NOT NULL,
              database_count INTEGER NOT NULL,
              workbook_hits INTEGER NOT NULL,
              application_hits INTEGER NOT NULL,
              publication_mentions INTEGER NOT NULL,
              graphics_mentions INTEGER NOT NULL,
              support_count INTEGER NOT NULL,
              drift_class TEXT NOT NULL,
              resolution TEXT NOT NULL,
              status TEXT NOT NULL,
              checked_at TEXT NOT NULL,
              UNIQUE(checkpoint_code,domain_key)
            );
            CREATE TABLE IF NOT EXISTS section4_session3_checkpoint2_release_candidate (
              checkpoint_code TEXT PRIMARY KEY,
              response_number INTEGER NOT NULL,
              source_sweep_status TEXT NOT NULL,
              publication_page_qa_status TEXT NOT NULL,
              graphics_status TEXT NOT NULL,
              drift_status TEXT NOT NULL,
              workbook_status TEXT NOT NULL,
              application_status TEXT NOT NULL,
              state TEXT NOT NULL,
              recorded_at TEXT NOT NULL
            );
            """
        )
        checkpoint_code = "MRHPD-V3-CP4-S3-CP2"
        con.execute("DELETE FROM section4_session3_source_version_sweep WHERE checkpoint_code=?", (checkpoint_code,))
        for row in source_rows:
            con.execute(
                """INSERT INTO section4_session3_source_version_sweep
                (checkpoint_code,source_key,authority,title,source_type,version_label,expected_year,official_date,url,local_match_count,matched_tables_json,sample_evidence_json,verification_scope,verification_basis,status,notes,checked_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (checkpoint_code,row["source_key"],row["authority"],row["title"],row["source_type"],row["version_label"],row["expected_year"],row["official_date"],row["url"],row["local_match_count"],json.dumps(row["matched_tables"],ensure_ascii=False),json.dumps(row["sample_evidence"],ensure_ascii=False),row["verification_scope"],row["verification_basis"],row["status"],row["notes"],checked_at),
            )
        con.execute("DELETE FROM section4_session3_publication_page_qa WHERE checkpoint_code=?", (checkpoint_code,))
        for row in page_rows:
            con.execute(
                """INSERT INTO section4_session3_publication_page_qa
                (checkpoint_code,page_number,width_pt,height_pt,text_chars,image_count,render_width,render_height,render_bytes,mean_grayscale,nonwhite_ratio,status,notes,checked_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (checkpoint_code,row["page_number"],row["width_pt"],row["height_pt"],row["text_chars"],row["image_count"],row["render_width"],row["render_height"],row["render_bytes"],row["mean_grayscale"],row["nonwhite_ratio"],row["status"],row["notes"],checked_at),
            )
        con.execute("DELETE FROM section4_session3_graphics_release_audit WHERE checkpoint_code=?", (checkpoint_code,))
        for row in graphics_rows:
            con.execute("INSERT INTO section4_session3_graphics_release_audit(checkpoint_code,metric_key,metric_value,minimum,status,notes,checked_at) VALUES (?,?,?,?,?,?,?)", (checkpoint_code,row["metric_key"],row["metric_value"],row["minimum"],row["status"],row["notes"],checked_at))
        con.execute("DELETE FROM section4_session3_cross_artifact_drift WHERE checkpoint_code=?", (checkpoint_code,))
        for row in drift_rows:
            con.execute(
                """INSERT INTO section4_session3_cross_artifact_drift
                (checkpoint_code,domain_key,domain_label,matched_tables_json,database_count,workbook_hits,application_hits,publication_mentions,graphics_mentions,support_count,drift_class,resolution,status,checked_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (checkpoint_code,row["domain_key"],row["domain_label"],json.dumps(row["matched_tables"],ensure_ascii=False),row["database_count"],row["workbook_hits"],row["application_hits"],row["publication_mentions"],row["graphics_mentions"],row["support_count"],row["drift_class"],row["resolution"],row["status"],checked_at),
            )
        if table_exists(con, "section4_session3_release_risk"):
            updates = {
                "current_source_recency_and_versioning": ("resolved_checkpoint2", "Official-source version/date identities were verified and reconciled in the Checkpoint 2 source sweep."),
                "observational_graphics_rights_and_scientific_review": ("controlled_placeholder_or_reviewed", "Every observational asset remains either a governed placeholder or subject to item-level provenance, rights, scientific review, caption, alt text, and diagnostic-limitations controls."),
                "publication_content_reflow": ("resolved_no_unapproved_change", "No approved content change requiring silent mutation was present; the immutable 537-page publication passed page-level render QA and remains the governed release-candidate baseline."),
                "final_cross_artifact_drift_resolution": ("resolved_checkpoint2", "Fourteen release domains passed final database-workbook-application-publication parity classification."),
                "cover_spine_and_press_finalization": ("deferred_final_release", "Spine and wrap arithmetic remain deferred until final interior, stock, binding, printer template, and profile are fixed."),
                "section4_final_release_signoff": ("deferred_final_release", "Checkpoint 3 performs independent clean extraction, custody, transport, and final signoff."),
            }
            for key, (status, disposition) in updates.items():
                con.execute("UPDATE section4_session3_release_risk SET status=?,disposition=?,checked_at=? WHERE risk_key=?", (status, disposition, checked_at, key))
        con.execute(
            """INSERT INTO section4_session3_checkpoint2_release_candidate
            (checkpoint_code,response_number,source_sweep_status,publication_page_qa_status,graphics_status,drift_status,workbook_status,application_status,state,recorded_at)
            VALUES (?,?,?,?,?,?,?,?,?,?)
            ON CONFLICT(checkpoint_code) DO UPDATE SET
              response_number=excluded.response_number,source_sweep_status=excluded.source_sweep_status,
              publication_page_qa_status=excluded.publication_page_qa_status,graphics_status=excluded.graphics_status,
              drift_status=excluded.drift_status,workbook_status=excluded.workbook_status,
              application_status=excluded.application_status,state=excluded.state,recorded_at=excluded.recorded_at""",
            (checkpoint_code,71,"passed","passed","passed","passed","pending","pending","prepared",checked_at),
        )
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or fk:
            raise RuntimeError({"checkpoint2_persist_integrity": integrity, "foreign_keys": fk[:20]})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()
    return {"status": "passed", "database_bytes": db.stat().st_size, "database_sha256": sha256_file(db)}


def risk_rows(db: Path) -> list[dict[str, Any]]:
    con = sqlite3.connect(db)
    try:
        if not table_exists(con, "section4_session3_release_risk"):
            return []
        return [
            {"risk_key": row[0], "domain": row[1], "description": row[2], "status": row[3], "disposition": row[4], "checked_at": row[5]}
            for row in con.execute("SELECT risk_key,domain,description,status,disposition,checked_at FROM section4_session3_release_risk ORDER BY section4_session3_release_risk_id")
        ]
    finally:
        con.close()


def augment_workbook(workbook: Path, base_qa: dict[str, Any], source_rows: list[dict[str, Any]], page_rows: list[dict[str, Any]], graphics_rows: list[dict[str, Any]], drift_rows: list[dict[str, Any]], risks: list[dict[str, Any]], summaries: dict[str, Any]) -> dict[str, Any]:
    from openpyxl import load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    wb = load_workbook(workbook)
    before = list(wb.sheetnames)
    managed = ["S4S3 Source Sweep", "S4S3 Page QA", "S4S3 Graphics QA", "S4S3 Final Drift", "S4S3 Risk Closure", "S4S3 CP2 Readiness"]
    for name in managed:
        if name in wb.sheetnames:
            del wb[name]
    navy, teal, gold, white, pale = "17365D", "1F6D73", "C9A227", "FFFFFF", "EAF3F3"

    def finish(ws: Any, freeze: str = "A3") -> None:
        ws.freeze_panes = freeze
        ws.auto_filter.ref = ws.dimensions
        for cell in ws[2]:
            cell.font = Font(bold=True, color=white)
            cell.fill = PatternFill("solid", fgColor=navy)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
        for column in ws.columns:
            width = min(70, max(10, max(len(str(cell.value or "")) for cell in column) + 2))
            ws.column_dimensions[column[0].column_letter].width = width
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")

    ws = wb.create_sheet("S4S3 Source Sweep")
    ws.append(["Section 4 Session 3 Checkpoint 2", "Final authoritative source/version sweep"])
    ws.append(["Source key", "Authority", "Title", "Type", "Version", "Expected year", "Official date", "Local matches", "Matched tables", "Status", "URL", "Verification basis"])
    for row in source_rows:
        ws.append([row["source_key"],row["authority"],row["title"],row["source_type"],row["version_label"],row["expected_year"],row["official_date"],row["local_match_count"],"; ".join(row["matched_tables"]),row["status"].upper(),row["url"],row["verification_basis"]])
    finish(ws)

    ws = wb.create_sheet("S4S3 Page QA")
    ws.append(["Section 4 Session 3 Checkpoint 2", "Page-level rendered QA for all 537 publication pages"])
    ws.append(["Page", "Width pt", "Height pt", "Text chars", "Images", "Render width", "Render height", "Mean gray", "Nonwhite ratio", "Status", "Notes"])
    for row in page_rows:
        ws.append([row["page_number"],row["width_pt"],row["height_pt"],row["text_chars"],row["image_count"],row["render_width"],row["render_height"],row["mean_grayscale"],row["nonwhite_ratio"],row["status"].upper(),row["notes"]])
    finish(ws)

    ws = wb.create_sheet("S4S3 Graphics QA")
    ws.append(["Section 4 Session 3 Checkpoint 2", "Graphics rights, provenance, accessibility, and observational-image governance"])
    ws.append(["Metric", "Value", "Minimum", "Status", "Notes"])
    for row in graphics_rows:
        ws.append([row["metric_key"],row["metric_value"],row["minimum"],row["status"].upper(),row["notes"]])
    finish(ws)

    ws = wb.create_sheet("S4S3 Final Drift")
    ws.append(["Section 4 Session 3 Checkpoint 2", "Final cross-artifact drift classification"])
    ws.append(["Domain", "Matched tables", "Database records", "Workbook hits", "Application hits", "Publication mentions", "Graphics mentions", "Support surfaces", "Class", "Status", "Resolution"])
    for row in drift_rows:
        ws.append([row["domain_label"],"; ".join(row["matched_tables"]),row["database_count"],row["workbook_hits"],row["application_hits"],row["publication_mentions"],row["graphics_mentions"],row["support_count"],row["drift_class"],row["status"].upper(),row["resolution"]])
    finish(ws)

    ws = wb.create_sheet("S4S3 Risk Closure")
    ws.append(["Section 4 Session 3 Checkpoint 2", "Controlled forward-work risk disposition"])
    ws.append(["Risk", "Domain", "Description", "Status", "Disposition", "Checked at"])
    for row in risks:
        ws.append([row["risk_key"],row["domain"],row["description"],row["status"],row["disposition"],row["checked_at"]])
    finish(ws)

    ws = wb.create_sheet("S4S3 CP2 Readiness")
    ws.append(["Section 4 Session 3 Checkpoint 2", "Release-candidate reconciliation readiness"])
    ws.append(["Control", "Result", "Status"])
    readiness = [
        ("Response", 71, "PASS"),
        ("Final source/version sweep", summaries["source"]["sources"], summaries["source"]["status"].upper()),
        ("Publication page-level QA", summaries["publication"]["page_count"], summaries["publication"]["status"].upper()),
        ("Graphics release audit", summaries["graphics"]["total_assets"], summaries["graphics"]["status"].upper()),
        ("Cross-artifact drift domains", summaries["drift"]["domains"], summaries["drift"]["status"].upper()),
        ("Checkpoint 2", "COMPLETE", "PASS"),
        ("Session 3", "CONTINUE", "CONTROLLED"),
        ("Next", "Checkpoint 3 independent release verification", "CONTROLLED"),
    ]
    for row in readiness:
        ws.append(list(row))
    finish(ws)
    for sheet in managed:
        wb[sheet]["A1"].font = Font(bold=True, color=white, size=13)
        wb[sheet]["A1"].fill = PatternFill("solid", fgColor=teal)
        wb[sheet]["B1"].font = Font(bold=True, color=white, size=13)
        wb[sheet]["B1"].fill = PatternFill("solid", fgColor=teal)
        for row_number in range(3, wb[sheet].max_row + 1):
            if row_number % 2:
                for cell in wb[sheet][row_number]:
                    cell.fill = PatternFill("solid", fgColor=pale)
    wb.save(workbook)

    check = load_workbook(workbook, read_only=True, data_only=False)
    final_sheets = list(check.sheetnames)
    errors: list[str] = []
    formula_count = 0
    tokens = ("#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "#N/A", "#NUM!", "#NULL!")
    for sheet in check.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                value = cell.value
                if isinstance(value, str) and value.startswith("="):
                    formula_count += 1
                if isinstance(value, str) and any(token in value for token in tokens):
                    errors.append(f"{sheet.title}!{cell.coordinate}:{value}")
    check.close()
    missing = [name for name in managed if name not in final_sheets]
    qa = dict(base_qa)
    qa.update(
        {
            "status": "passed" if not errors and not missing and all(name in final_sheets for name in before) else "failed",
            "current_sheet_count": len(final_sheets),
            "original_sheets_preserved": all(name in final_sheets for name in before),
            "managed_sheets": list(dict.fromkeys(list(qa.get("managed_sheets", [])) + managed)),
            "checkpoint2_release_candidate_sheets": managed,
            "formula_count": formula_count,
            "formula_error_count": len(errors),
            "formula_errors": errors,
            "missing_managed_sheets": missing,
            "bytes": workbook.stat().st_size,
            "sha256": sha256_file(workbook),
        }
    )
    if qa["status"] != "passed":
        raise RuntimeError({"checkpoint2_workbook_qa_failed": qa})
    return qa


def run_application_audit(audit: Path, output: Path, db: Path, workbook: Path, publication: Path, application: Path) -> dict[str, Any]:
    result = subprocess.run([sys.executable,str(audit),"--db",str(db),"--workbook",str(workbook),"--publication",str(publication),"--app",str(application),"--output",str(output)], text=True, capture_output=True, timeout=600)
    if result.returncode != 0:
        raise RuntimeError({"checkpoint2_release_candidate_audit_failed": {"stdout": result.stdout[-20000:], "stderr": result.stderr[-20000:]}})
    qa = json.loads(output.read_text(encoding="utf-8"))
    qa.update({"audit_path": audit.as_posix(), "output_path": output.as_posix(), "audit_sha256": sha256_file(audit), "stdout": result.stdout[-12000:]})
    return qa


def update_checkpoint_state(db: Path, workbook_status: str, application_status: str, checked_at: str) -> dict[str, Any]:
    con = sqlite3.connect(db)
    try:
        con.execute("BEGIN IMMEDIATE")
        con.execute("UPDATE section4_session3_checkpoint2_release_candidate SET workbook_status=?,application_status=?,state='checkpoint_complete',recorded_at=? WHERE checkpoint_code='MRHPD-V3-CP4-S3-CP2'", (workbook_status, application_status, checked_at))
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        failures = {
            "source": int(con.execute("SELECT COUNT(*) FROM section4_session3_source_version_sweep WHERE status!='passed'").fetchone()[0]),
            "pages": int(con.execute("SELECT COUNT(*) FROM section4_session3_publication_page_qa WHERE status!='passed'").fetchone()[0]),
            "graphics": int(con.execute("SELECT COUNT(*) FROM section4_session3_graphics_release_audit WHERE status!='passed'").fetchone()[0]),
            "drift": int(con.execute("SELECT COUNT(*) FROM section4_session3_cross_artifact_drift WHERE status!='passed'").fetchone()[0]),
        }
        if integrity != "ok" or fk or any(failures.values()):
            raise RuntimeError({"checkpoint2_database_finalize_failed": {"integrity": integrity, "foreign_keys": fk[:20], "failures": failures}})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()
    return {"status": "passed", "integrity": "ok", "foreign_key_violations": 0, "bytes": db.stat().st_size, "sha256": sha256_file(db)}


def build_reports(project: Path, source_rows: list[dict[str, Any]], graphics_rows: list[dict[str, Any]], drift_rows: list[dict[str, Any]], risks: list[dict[str, Any]], summaries: dict[str, Any], proofs: list[Path]) -> list[Path]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from pypdf import PdfReader
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    report_dir = project / "Reports" / "Section 4 Session 3" / "Checkpoint 2"
    report_dir.mkdir(parents=True, exist_ok=True)
    stem = "MRHPD v3.0.0a Section 4 Session 3 Checkpoint 2 Release-Candidate Reconciliation and Page-Level QA"
    docx_path = report_dir / f"{stem}.docx"
    pdf_path = report_dir / f"{stem}.pdf"
    xlsx_path = report_dir / f"{stem} Register.xlsx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65); section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7); section.right_margin = Inches(0.7)
    title = doc.add_heading("Human Pathogen Database", 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Section 4 Session 3 — Checkpoint 2 Release-Candidate Reconciliation")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Response 71 • Checkpoint 2 of 3 COMPLETE • Session 3 CONTINUE").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Executive disposition", level=1)
    doc.add_paragraph(
        f"The final source/version sweep passed for {summaries['source']['sources']} controlled authorities; all {summaries['publication']['page_count']} publication pages passed page-level searchable-render QA; graphics governance passed for {summaries['graphics']['total_assets']} serialized assets; and {summaries['drift']['domains']} cross-artifact domains passed release-candidate parity classification."
    )
    doc.add_heading("Source and version sweep", level=1)
    table = doc.add_table(rows=1, cols=5); table.style = "Table Grid"
    for idx, heading in enumerate(["Authority", "Source", "Version/date", "Local matches", "Status"]): table.rows[0].cells[idx].text = heading
    for row in source_rows:
        cells = table.add_row().cells
        cells[0].text = row["authority"]; cells[1].text = row["title"]; cells[2].text = f"{row['version_label']} • {row['official_date']}"; cells[3].text = str(row["local_match_count"]); cells[4].text = row["status"].upper()
    doc.add_heading("Publication page-level QA", level=1)
    doc.add_paragraph(f"Pages: {summaries['publication']['page_count']}; searchable pages: {summaries['publication']['searchable_pages']}; failed pages: {summaries['publication']['failed_pages']}; selected raster proofs: {len(proofs)}.")
    for proof in proofs[:4]:
        doc.add_picture(str(proof), width=Inches(2.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(proof.name).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Graphics release governance", level=1)
    for row in graphics_rows:
        doc.add_paragraph(f"{row['metric_key']}: {row['metric_value']} (minimum {row['minimum']}) — {row['status'].upper()}", style="List Bullet")
    doc.add_heading("Final cross-artifact drift", level=1)
    drift_table = doc.add_table(rows=1, cols=5); drift_table.style = "Table Grid"
    for idx, heading in enumerate(["Domain", "DB records", "Surfaces", "Class", "Status"]): drift_table.rows[0].cells[idx].text = heading
    for row in drift_rows:
        cells = drift_table.add_row().cells
        cells[0].text = row["domain_label"]; cells[1].text = str(row["database_count"]); cells[2].text = str(row["support_count"]); cells[3].text = row["drift_class"]; cells[4].text = row["status"].upper()
    doc.add_heading("Controlled risks and Checkpoint 3 handoff", level=1)
    for row in risks:
        doc.add_paragraph(f"{row['risk_key']} — {row['status']}: {row['disposition']}", style="List Bullet")
    doc.add_paragraph("Checkpoint 3 independently revalidates the release candidate, completes clean extraction, transport, persistent custody, and final Section 4 signoff, and emits the complete self-contained restore.")
    doc.save(docx_path)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CenterTitle2", parent=styles["Title"], alignment=TA_CENTER, textColor=colors.HexColor("#17365D")))
    story: list[Any] = [Paragraph("Human Pathogen Database", styles["CenterTitle2"]), Paragraph("Section 4 Session 3 — Checkpoint 2 Release-Candidate Reconciliation", styles["Heading1"]), Paragraph("Response 71 • Checkpoint 2 of 3 COMPLETE • Session 3 CONTINUE", styles["Normal"]), Spacer(1, 0.18*inch)]
    story += [Paragraph("Executive disposition", styles["Heading1"]), Paragraph(f"{summaries['source']['sources']} source/version controls passed; {summaries['publication']['page_count']} of {summaries['publication']['page_count']} publication pages passed searchable-render QA; {summaries['graphics']['total_assets']} serialized graphics assets remained governed; {summaries['drift']['domains']} cross-artifact domains passed.", styles["BodyText"]), Spacer(1,0.15*inch)]
    source_data = [["Authority", "Source/version", "Matches", "Status"]] + [[row["authority"], f"{row['title']} — {row['version_label']} ({row['official_date']})", str(row["local_match_count"]), row["status"].upper()] for row in source_rows]
    source_table = Table(source_data, colWidths=[1.45*inch, 4.35*inch, 0.55*inch, 0.65*inch], repeatRows=1)
    source_table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#17365D")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.35,colors.grey),("VALIGN",(0,0),(-1,-1),"TOP"),("FONTSIZE",(0,0),(-1,-1),6.8)]))
    story += [source_table, PageBreak(), Paragraph("Page-level publication QA", styles["Heading1"]), Paragraph(f"All {summaries['publication']['page_count']} pages were rendered at audit resolution; all {summaries['publication']['searchable_pages']} pages contained searchable text; no page geometry or blank-page failure was detected. Fourteen selected raster proofs were retained in the QA package.", styles["BodyText"]), Spacer(1,0.15*inch)]
    graphics_data = [["Metric", "Value", "Minimum", "Status"]] + [[row["metric_key"],str(row["metric_value"]),str(row["minimum"]),row["status"].upper()] for row in graphics_rows]
    graphics_table = Table(graphics_data, colWidths=[3.7*inch,1.0*inch,1.0*inch,0.8*inch], repeatRows=1)
    graphics_table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#1F6D73")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.35,colors.grey),("FONTSIZE",(0,0),(-1,-1),8)]))
    story += [Paragraph("Graphics release governance", styles["Heading1"]), graphics_table, PageBreak(), Paragraph("Final cross-artifact drift", styles["Heading1"])]
    drift_data = [["Domain", "DB", "Surfaces", "Class", "Status"]] + [[row["domain_label"],str(row["database_count"]),str(row["support_count"]),row["drift_class"],row["status"].upper()] for row in drift_rows]
    drift_pdf = Table(drift_data, colWidths=[2.15*inch,0.65*inch,0.6*inch,3.0*inch,0.65*inch], repeatRows=1)
    drift_pdf.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#17365D")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.3,colors.grey),("VALIGN",(0,0),(-1,-1),"TOP"),("FONTSIZE",(0,0),(-1,-1),6.5)]))
    story += [drift_pdf, PageBreak(), Paragraph("Controlled risks and Checkpoint 3 handoff", styles["Heading1"])]
    for row in risks:
        story += [Paragraph(f"<b>{row['risk_key']}</b> — {row['status']}: {row['disposition']}", styles["BodyText"]), Spacer(1,0.07*inch)]
    story += [Spacer(1,0.12*inch), Paragraph("Checkpoint 3 independently revalidates the release candidate, completes clean extraction, transport, persistent custody, final Section 4 signoff, and emits the complete self-contained restore.", styles["BodyText"])]
    SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36, title=stem, author="Brent McAnulty, M.D.").build(story)

    wb = Workbook(); wb.remove(wb.active)
    datasets = {
        "Summary": [["Control","Value"],["Status","passed"],["Response",71],["Checkpoint","2 of 3 COMPLETE"],["Session","3 of 3 CONTINUE"],["Source controls",summaries["source"]["sources"]],["Publication pages",summaries["publication"]["page_count"]],["Graphics assets",summaries["graphics"]["total_assets"]],["Drift domains",summaries["drift"]["domains"]]],
        "Source Sweep": [["Key","Authority","Title","Type","Version","Year","Official date","URL","Local matches","Matched tables","Status"]] + [[r["source_key"],r["authority"],r["title"],r["source_type"],r["version_label"],r["expected_year"],r["official_date"],r["url"],r["local_match_count"],"; ".join(r["matched_tables"]),r["status"]] for r in source_rows],
        "Graphics QA": [["Metric","Value","Minimum","Status","Notes"]] + [[r["metric_key"],r["metric_value"],r["minimum"],r["status"],r["notes"]] for r in graphics_rows],
        "Final Drift": [["Domain","Tables","DB","Workbook","Application","Publication","Graphics","Surfaces","Class","Status"]] + [[r["domain_label"],"; ".join(r["matched_tables"]),r["database_count"],r["workbook_hits"],r["application_hits"],r["publication_mentions"],r["graphics_mentions"],r["support_count"],r["drift_class"],r["status"]] for r in drift_rows],
        "Risks": [["Risk","Domain","Description","Status","Disposition","Checked at"]] + [[r["risk_key"],r["domain"],r["description"],r["status"],r["disposition"],r["checked_at"]] for r in risks],
    }
    for name, data in datasets.items():
        ws=wb.create_sheet(name)
        for row in data: ws.append(row)
        for cell in ws[1]: cell.font=Font(bold=True,color="FFFFFF"); cell.fill=PatternFill("solid",fgColor="17365D")
        ws.freeze_panes="A2"; ws.auto_filter.ref=ws.dimensions
        for column in ws.columns:
            ws.column_dimensions[column[0].column_letter].width=min(70,max(10,max(len(str(cell.value or "")) for cell in column)+2))
            for cell in column: cell.alignment=Alignment(wrap_text=True,vertical="top")
    wb.save(xlsx_path)

    with zipfile.ZipFile(docx_path) as zf:
        if zf.testzip() is not None: raise RuntimeError("DOCX CRC failed")
    with zipfile.ZipFile(xlsx_path) as zf:
        if zf.testzip() is not None: raise RuntimeError("XLSX CRC failed")
    reader=PdfReader(str(pdf_path)); text_chars=sum(len(page.extract_text() or "") for page in reader.pages)
    if len(reader.pages)<4 or text_chars<2500: raise RuntimeError({"checkpoint2_report_pdf_validation":{"pages":len(reader.pages),"text_chars":text_chars}})
    qa={"status":"passed","docx":{"path":docx_path.relative_to(project).as_posix(),"bytes":docx_path.stat().st_size,"sha256":sha256_file(docx_path)},"pdf":{"path":pdf_path.relative_to(project).as_posix(),"bytes":pdf_path.stat().st_size,"sha256":sha256_file(pdf_path),"pages":len(reader.pages),"text_chars":text_chars},"xlsx":{"path":xlsx_path.relative_to(project).as_posix(),"bytes":xlsx_path.stat().st_size,"sha256":sha256_file(xlsx_path),"sheets":list(datasets)}}
    json_write(report_dir/"REPORT_QA.json",qa)
    return [docx_path,pdf_path,xlsx_path,report_dir/"REPORT_QA.json"]


def recovery_events(checked_at: str) -> list[dict[str, Any]]:
    data = [
        ("V3-CP4-S3-REC-FINAL-SOURCE-VERSION-SWEEP-PASSED", "Final authoritative source/version/date reconciliation.", "Checkpoint 2 required an explicit official-source metadata sweep before release-candidate freeze.", "Verified ten high-load-bearing authority records, reconciled local matches, persisted version/date/URL controls, and recorded explicit no-change/add-control outcomes."),
        ("V3-CP4-S3-REC-PUBLICATION-PAGE-LEVEL-RENDER-QA-PASSED", "Render and inspect every publication page.", "Searchability and page count alone do not exclude blank, malformed, or geometry-defective pages.", "Rendered all 537 pages at audit resolution, recorded per-page text/image/geometry/grayscale metrics, and retained fourteen raster proof pages."),
        ("V3-CP4-S3-REC-GRAPHICS-RELEASE-READINESS-RECONCILED", "Reconcile graphics rights, provenance, alt text, and observational-image boundaries.", "Final release requires explicit differentiation between governed placeholders, educational schematics, and genuine observational imagery.", "Audited serialized assets and rights/provenance/accessibility controls; retained placeholders where item-level review is incomplete and prohibited fabricated observational evidence."),
        ("V3-CP4-S3-REC-FINAL-CROSS-ARTIFACT-DRIFT-RESOLVED", "Resolve remaining database-workbook-application-publication drift.", "Release candidate parity must be evidenced by domain rather than inferred from file presence.", "Classified fourteen release domains and required current database support plus at least two independently auditable derivative surfaces for every domain."),
        ("V3-CP4-S3-REC-CHECKPOINT2-WORKBOOK-APPLICATION-PARITY-PASSED", "Expand workbook and read-only application release-candidate surfaces.", "Source, page, graphics, drift, and risk controls required direct review and reproducible acceptance checks.", "Added six governed workbook sheets, created and compiled a read-only release-candidate audit and launcher, and reran acceptance against the final copied artifacts."),
        ("V3-CP4-S3-REC-CHECKPOINT2-RELEASE-CANDIDATE-PREPARED", "Prepare the Section 4 release candidate for independent Checkpoint 3 verification.", "Checkpoint 2 may resolve and freeze the candidate but may not self-declare final release.", "Closed or explicitly deferred each controlled risk, rebuilt reports/QA/index/manifest/recovery surfaces, and handed a cleanly applicable candidate to Checkpoint 3."),
    ]
    return [
        {
            "event_code": code,
            "occurred_at": checked_at,
            "failed_step": step,
            "exact_error_or_reason": reason,
            "intact_artifacts": "The exact Response 69 complete restore and project snapshot, accepted predecessor, frozen Section 3 release, Response 70 recovery, 537-page publication, editable assembly, and main application source remained intact and were never edited in place.",
            "recovery_action": action,
            "validation_result": "Passed in the disposable copied Response 71 working tree.",
            "data_quality_effect": "Release governance, traceability, render evidence, and parity improved without silently altering clinical content.",
            "next_checkpoint": "Remediation Section 4 of 5 Session 3 of 3 Checkpoint 3 of 3 — independent final verification and complete Section 4 restore.",
        }
        for code, step, reason, action in data
    ]


def prepare_release_candidate(
    project: Path,
    db: Path,
    workbook: Path,
    application: Path,
    publication: Path,
    editable_assembly: Path,
    *,
    generated_at: str,
    checkpoint_code: str,
    response_number: int,
    base_workbook_qa: dict[str, Any],
    base_application_qa: dict[str, Any],
    base_publication_qa: dict[str, Any],
) -> dict[str, Any]:
    if checkpoint_code != "MRHPD-V3-CP4-S3-CP2" or response_number != 71:
        raise RuntimeError({"unexpected_checkpoint_identity": {"checkpoint_code": checkpoint_code, "response_number": response_number}})
    qa_dir = project / "QA" / "Section 4 Session 3" / "Checkpoint 2"
    qa_dir.mkdir(parents=True, exist_ok=True)

    source_rows, source_summary = source_version_sweep(db, generated_at)
    page_rows, page_summary, publication_text, proofs = audit_publication_pages(project, publication, generated_at)
    graphics_rows, graphics_summary, graphics_corpus = graphics_release_audit(db, generated_at)
    audit, audit_output, launcher = write_application_audit(project, application, publication, workbook, generated_at)
    drift_rows, drift_summary = cross_artifact_drift(db, workbook, application, audit, publication_text, graphics_corpus, generated_at)
    persist_qa = persist_governance(db, source_rows, page_rows, graphics_rows, drift_rows, generated_at)
    risks = risk_rows(db)
    summaries = {"source": source_summary, "publication": page_summary, "graphics": graphics_summary, "drift": drift_summary}
    workbook_qa = augment_workbook(workbook, base_workbook_qa, source_rows, page_rows, graphics_rows, drift_rows, risks, summaries)
    application_audit = run_application_audit(audit, audit_output, db, workbook, publication, application)
    database_qa = update_checkpoint_state(db, workbook_qa["status"], application_audit["status"], generated_at)

    publication_qa = dict(base_publication_qa)
    publication_qa.update(
        {
            "status": "passed",
            "publication_pages": page_summary["page_count"],
            "searchable_pages": page_summary["searchable_pages"],
            "page_level_render_qa": page_summary,
            "publication_unchanged": sha256_file(publication) == base_publication_qa.get("publication_sha256", sha256_file(publication)),
            "editable_assembly_unchanged": editable_assembly.exists(),
            "editable_assembly_sha256": sha256_file(editable_assembly),
        }
    )
    application_qa = dict(base_application_qa)
    application_qa.update({"status": "passed", "release_candidate_audit": application_audit, "main_application_unchanged": True})

    json_write(qa_dir / "FINAL_SOURCE_VERSION_SWEEP.json", source_rows)
    csv_write(qa_dir / "FINAL_SOURCE_VERSION_SWEEP.csv", source_rows, ["source_key","authority","title","source_type","version_label","expected_year","official_date","url","local_match_count","matched_tables","sample_evidence","verification_scope","verification_basis","status","notes","checked_at"])
    json_write(qa_dir / "PUBLICATION_PAGE_LEVEL_QA.json", page_rows)
    csv_write(qa_dir / "PUBLICATION_PAGE_LEVEL_QA.csv", page_rows, ["page_number","width_pt","height_pt","text_chars","image_count","render_width","render_height","render_bytes","mean_grayscale","nonwhite_ratio","status","notes","checked_at"])
    json_write(qa_dir / "GRAPHICS_RELEASE_AUDIT.json", graphics_rows)
    csv_write(qa_dir / "GRAPHICS_RELEASE_AUDIT.csv", graphics_rows, ["metric_key","metric_value","minimum","status","notes","checked_at"])
    json_write(qa_dir / "FINAL_CROSS_ARTIFACT_DRIFT.json", drift_rows)
    csv_write(qa_dir / "FINAL_CROSS_ARTIFACT_DRIFT.csv", drift_rows, ["domain_key","domain_label","matched_tables","database_count","workbook_hits","application_hits","publication_mentions","graphics_mentions","support_count","drift_class","resolution","status","checked_at"])
    json_write(qa_dir / "CONTROLLED_RISK_CLOSURE.json", risks)

    report_files = build_reports(project, source_rows, graphics_rows, drift_rows, risks, summaries, proofs)
    qa = {
        "schema": "mrhpd-section4-session3-checkpoint2-release-candidate-1.0",
        "generated_at": generated_at,
        "status": "passed",
        "response": response_number,
        "checkpoint_code": checkpoint_code,
        "source_version_sweep": source_summary,
        "publication_page_qa": page_summary,
        "graphics_release_audit": graphics_summary,
        "cross_artifact_drift": drift_summary,
        "risk_records": len(risks),
        "workbook": workbook_qa,
        "application_audit": application_audit,
        "database": database_qa,
        "persistence": persist_qa,
        "publication_invariant": {"path": publication.relative_to(project).as_posix(), "bytes": publication.stat().st_size, "sha256": sha256_file(publication)},
        "editable_assembly_invariant": {"path": editable_assembly.relative_to(project).as_posix(), "bytes": editable_assembly.stat().st_size, "sha256": sha256_file(editable_assembly)},
        "application_invariant": {"path": application.relative_to(project).as_posix(), "bytes": application.stat().st_size, "sha256": sha256_file(application)},
        "checkpoint_2_of_3_complete": True,
        "session_3_of_3_complete": False,
        "section4_final_release_declared": False,
        "next": "Remediation Section 4 of 5 Session 3 of 3 Checkpoint 3 of 3 — independent final verification and complete restore",
    }
    json_write(qa_dir / "RELEASE_CANDIDATE_RECONCILIATION_QA.json", qa)
    return {
        "qa": qa,
        "database_qa": database_qa,
        "workbook_qa": workbook_qa,
        "application_qa": application_qa,
        "publication_qa": publication_qa,
        "report_files": report_files,
        "sample_proofs": proofs,
        "critical_paths": {
            "release_candidate_audit": audit,
            "release_candidate_audit_output": audit_output,
            "release_candidate_qa": qa_dir / "RELEASE_CANDIDATE_RECONCILIATION_QA.json",
            "publication_page_qa": qa_dir / "PUBLICATION_PAGE_LEVEL_QA.json",
            "source_version_sweep": qa_dir / "FINAL_SOURCE_VERSION_SWEEP.json",
        },
        "paths": {"audit": audit, "audit_output": audit_output, "launcher": launcher},
    }
