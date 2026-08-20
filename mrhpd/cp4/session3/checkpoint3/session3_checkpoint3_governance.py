#!/usr/bin/env python3
"""Independent final verification and Section 4 release governance.

This module runs only against the disposable project reconstructed from the
exact Response 69 complete restore and the verified Response 71 Checkpoint 2
recovery package. It independently revalidates the release candidate, records
Checkpoint 3, closes Session 3 and Remediation Section 4, and emits final QA and
release reports without editing accepted or frozen source artifacts in place.
"""
from __future__ import annotations

import csv
import hashlib
import json
import sqlite3
import subprocess
import sys
import zipfile
from pathlib import Path
from typing import Any

APPLICATION_SHA256 = "5f1e4ac8fc6e2ffad213646c78e4f261bf655795de5ac8a7d4486d3be11ce139"
PUBLICATION_SHA256 = "8a053112ca24cd730b970130d5d0fc57a15c681531603601096186aeb0cd9642"
EDITABLE_ASSEMBLY_SHA256 = "f832ff934d77049d75712f28bdfc9167b8a6b119c797235431b304b9e24369a2"
CHECKPOINT_CODE = "MRHPD-V3-CP4-S3-CP3"
CP2_CODE = "MRHPD-V3-CP4-S3-CP2"
FINAL_RELEASE_CODE = "MRHPD-V3-CP4-COMPLETE"
NEXT_SECTION = "Remediation Section 5 of 5"


def sha256_file(path: Path, chunk_size: int = 1024 * 1024) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(chunk_size), b""):
            digest.update(block)
    return digest.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False, default=str) + "\n", encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value.rstrip() + "\n", encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            normalized: dict[str, Any] = {}
            for field in fields:
                value = row.get(field)
                if isinstance(value, (dict, list)):
                    value = json.dumps(value, ensure_ascii=False)
                normalized[field] = value
            writer.writerow(normalized)


def qident(value: str) -> str:
    return '"' + value.replace('"', '""') + '"'


def table_exists(con: sqlite3.Connection, table: str) -> bool:
    return con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?", (table,)).fetchone() is not None


def table_columns(con: sqlite3.Connection, table: str) -> list[str]:
    return [str(row[1]) for row in con.execute(f"PRAGMA table_info({qident(table)})")]


def logical_tables(con: sqlite3.Connection) -> list[str]:
    rows = [str(row[0]) for row in con.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name")]
    names = set(rows)
    shadow = ("_data", "_idx", "_content", "_docsize", "_config")
    return [name for name in rows if not any(name.endswith(suffix) and name[: -len(suffix)] in names for suffix in shadow)]


def safe_scalar(con: sqlite3.Connection, sql: str, params: tuple[Any, ...] = (), default: Any = 0) -> Any:
    try:
        row = con.execute(sql, params).fetchone()
        return row[0] if row else default
    except Exception:
        return default


def safe_count(con: sqlite3.Connection, table: str, where: str = "", params: tuple[Any, ...] = ()) -> int:
    if not table_exists(con, table):
        return 0
    suffix = f" WHERE {where}" if where else ""
    return int(safe_scalar(con, f"SELECT COUNT(*) FROM {qident(table)}{suffix}", params, 0) or 0)


def locate_by_hash(root: Path, pattern: str, expected_sha256: str) -> Path:
    matches = [path for path in root.rglob(pattern) if path.is_file() and sha256_file(path) == expected_sha256]
    if len(matches) != 1:
        raise RuntimeError({"pattern": pattern, "expected_sha256": expected_sha256, "matches": [str(path) for path in matches]})
    return matches[0]


def checkpoint_state(con: sqlite3.Connection, code: str) -> str | None:
    for table in ("section4_session3_checkpoint", "section4_checkpoint"):
        if table_exists(con, table) and "checkpoint_code" in table_columns(con, table) and "state" in table_columns(con, table):
            row = con.execute(f"SELECT state FROM {qident(table)} WHERE checkpoint_code=?", (code,)).fetchone()
            if row:
                return str(row[0])
    return None


def audit_candidate_database(db: Path, generated_at: str) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    con = sqlite3.connect(db)
    try:
        integrity = str(con.execute("PRAGMA integrity_check").fetchone()[0])
        foreign_keys = list(con.execute("PRAGMA foreign_key_check"))
        physical_tables = int(con.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'").fetchone()[0])
        logical = logical_tables(con)
        response71 = safe_count(con, "thread_response_reconciliation_cp3", "response_key='R71'")
        response70 = safe_count(con, "thread_response_reconciliation_cp3", "response_key='R70'")
        cp1_state = checkpoint_state(con, "MRHPD-V3-CP4-S3-CP1")
        cp2_state = checkpoint_state(con, CP2_CODE)
        field_rows = safe_count(con, "section4_session3_field_coverage")
        field_failures = safe_count(con, "section4_session3_field_coverage", "status!='passed'")
        query_rows = safe_count(con, "section4_session3_query_coverage")
        query_failures = safe_count(con, "section4_session3_query_coverage", "status!='passed'")
        source_governance_rows = safe_count(con, "section4_session3_source_governance")
        source_governance_failures = safe_count(con, "section4_session3_source_governance", "status!='passed'")
        baseline_drift_rows = safe_count(con, "section4_session3_drift_resolution")
        baseline_drift_failures = safe_count(con, "section4_session3_drift_resolution", "status!='passed'")
        release_governance_rows = safe_count(con, "section4_session3_release_governance")
        release_governance_failures = safe_count(con, "section4_session3_release_governance", "status!='passed'")
        risk_rows = safe_count(con, "section4_session3_release_risk")
        source_sweep_rows = safe_count(con, "section4_session3_source_version_sweep", "checkpoint_code=?", (CP2_CODE,))
        source_sweep_failures = safe_count(con, "section4_session3_source_version_sweep", "checkpoint_code=? AND status!='passed'", (CP2_CODE,))
        page_rows = safe_count(con, "section4_session3_publication_page_qa", "checkpoint_code=?", (CP2_CODE,))
        page_failures = safe_count(con, "section4_session3_publication_page_qa", "checkpoint_code=? AND status!='passed'", (CP2_CODE,))
        graphics_rows = safe_count(con, "section4_session3_graphics_release_audit", "checkpoint_code=?", (CP2_CODE,))
        graphics_failures = safe_count(con, "section4_session3_graphics_release_audit", "checkpoint_code=? AND status!='passed'", (CP2_CODE,))
        graphic_assets = int(
            safe_scalar(
                con,
                "SELECT metric_value FROM section4_session3_graphics_release_audit WHERE checkpoint_code=? AND metric_key='graphic_assets'",
                (CP2_CODE,),
                0,
            )
            or 0
        )
        final_drift_rows = safe_count(con, "section4_session3_cross_artifact_drift", "checkpoint_code=?", (CP2_CODE,))
        final_drift_failures = safe_count(con, "section4_session3_cross_artifact_drift", "checkpoint_code=? AND status!='passed'", (CP2_CODE,))
        cp2_candidate_state = safe_scalar(
            con,
            "SELECT state FROM section4_session3_checkpoint2_release_candidate WHERE checkpoint_code=?",
            (CP2_CODE,),
            None,
        ) if table_exists(con, "section4_session3_checkpoint2_release_candidate") else None
    finally:
        con.close()

    controls = [
        {"control_key": "sqlite_integrity", "expected": "ok", "observed": integrity, "status": "passed" if integrity == "ok" else "failed"},
        {"control_key": "foreign_keys", "expected": 0, "observed": len(foreign_keys), "status": "passed" if not foreign_keys else "failed"},
        {"control_key": "physical_table_inventory", "expected": ">=209", "observed": physical_tables, "status": "passed" if physical_tables >= 209 else "failed"},
        {"control_key": "logical_table_inventory", "expected": ">=191", "observed": len(logical), "status": "passed" if len(logical) >= 191 else "failed"},
        {"control_key": "response70_lineage", "expected": 1, "observed": response70, "status": "passed" if response70 == 1 else "failed"},
        {"control_key": "response71_lineage", "expected": 1, "observed": response71, "status": "passed" if response71 == 1 else "failed"},
        {"control_key": "checkpoint1_state", "expected": "checkpoint_complete", "observed": cp1_state, "status": "passed" if cp1_state == "checkpoint_complete" else "failed"},
        {"control_key": "checkpoint2_state", "expected": "checkpoint_complete", "observed": cp2_state, "status": "passed" if cp2_state == "checkpoint_complete" else "failed"},
        {"control_key": "field_coverage", "expected": ">0 and 0 failures", "observed": {"rows": field_rows, "failures": field_failures}, "status": "passed" if field_rows > 0 and field_failures == 0 else "failed"},
        {"control_key": "query_coverage", "expected": ">=21 and 0 failures", "observed": {"rows": query_rows, "failures": query_failures}, "status": "passed" if query_rows >= 21 and query_failures == 0 else "failed"},
        {"control_key": "source_governance", "expected": ">0 and 0 failures", "observed": {"rows": source_governance_rows, "failures": source_governance_failures}, "status": "passed" if source_governance_rows > 0 and source_governance_failures == 0 else "failed"},
        {"control_key": "baseline_drift", "expected": ">0 and 0 failures", "observed": {"rows": baseline_drift_rows, "failures": baseline_drift_failures}, "status": "passed" if baseline_drift_rows > 0 and baseline_drift_failures == 0 else "failed"},
        {"control_key": "release_governance", "expected": ">=17 and 0 failures", "observed": {"rows": release_governance_rows, "failures": release_governance_failures}, "status": "passed" if release_governance_rows >= 17 and release_governance_failures == 0 else "failed"},
        {"control_key": "controlled_risk_register", "expected": ">=6", "observed": risk_rows, "status": "passed" if risk_rows >= 6 else "failed"},
        {"control_key": "final_source_version_sweep", "expected": ">=10 and 0 failures", "observed": {"rows": source_sweep_rows, "failures": source_sweep_failures}, "status": "passed" if source_sweep_rows >= 10 and source_sweep_failures == 0 else "failed"},
        {"control_key": "checkpoint2_publication_page_qa", "expected": "537 and 0 failures", "observed": {"rows": page_rows, "failures": page_failures}, "status": "passed" if page_rows == 537 and page_failures == 0 else "failed"},
        {"control_key": "graphics_release_governance", "expected": ">=6 metrics, >=336 assets, 0 failures", "observed": {"metrics": graphics_rows, "assets": graphic_assets, "failures": graphics_failures}, "status": "passed" if graphics_rows >= 6 and graphic_assets >= 336 and graphics_failures == 0 else "failed"},
        {"control_key": "final_cross_artifact_drift", "expected": "14 and 0 failures", "observed": {"rows": final_drift_rows, "failures": final_drift_failures}, "status": "passed" if final_drift_rows == 14 and final_drift_failures == 0 else "failed"},
        {"control_key": "checkpoint2_release_candidate_state", "expected": "checkpoint_complete", "observed": cp2_candidate_state, "status": "passed" if cp2_candidate_state == "checkpoint_complete" else "failed"},
    ]
    for row in controls:
        row.update({"checkpoint_code": CHECKPOINT_CODE, "checked_at": generated_at})
    failed = [row for row in controls if row["status"] != "passed"]
    summary = {
        "status": "passed" if not failed else "failed",
        "control_count": len(controls),
        "failed_controls": [row["control_key"] for row in failed],
        "physical_tables": physical_tables,
        "logical_tables": len(logical),
        "field_records": field_rows,
        "query_records": query_rows,
        "source_governance_records": source_governance_rows,
        "source_version_controls": source_sweep_rows,
        "publication_page_records": page_rows,
        "graphics_assets": graphic_assets,
        "drift_domains": final_drift_rows,
        "checked_at": generated_at,
    }
    if failed:
        raise RuntimeError({"candidate_database_acceptance_failed": failed, "summary": summary})
    return controls, summary


def audit_publication_pages(project: Path, publication: Path, generated_at: str) -> tuple[list[dict[str, Any]], dict[str, Any], list[Path]]:
    import fitz
    from pypdf import PdfReader

    qa_dir = project / "QA" / "Section 4 Session 3" / "Checkpoint 3"
    proof_dir = qa_dir / "Independent Publication Proofs"
    proof_dir.mkdir(parents=True, exist_ok=True)
    sample_pages = {1, 16, 208, 419, 455, 537}
    doc = fitz.open(publication)
    secondary = PdfReader(str(publication))
    rows: list[dict[str, Any]] = []
    proofs: list[Path] = []
    try:
        for index, page in enumerate(doc):
            page_number = index + 1
            primary_text = page.get_text("text") or ""
            secondary_text = secondary.pages[index].extract_text() or ""
            searchable_text = primary_text if primary_text.strip() else secondary_text
            rect = page.rect
            pix = page.get_pixmap(matrix=fitz.Matrix(0.20, 0.20), alpha=False, colorspace=fitz.csGRAY)
            samples = pix.samples
            mean = (sum(samples) / len(samples)) if samples else 255.0
            nonwhite = sum(1 for value in samples if value < 248)
            nonwhite_ratio = (nonwhite / len(samples)) if samples else 0.0
            image_count = len(page.get_images(full=True))
            drawing_count = len(page.get_drawings())
            block_count = len(page.get_text("blocks"))
            geometry_valid = min(rect.width, rect.height) > 500 and max(rect.width, rect.height) > 700
            render_valid = bool(samples) and pix.width > 80 and pix.height > 80
            object_valid = (image_count + drawing_count + block_count) > 0
            passed = bool(searchable_text.strip()) and geometry_valid and render_valid and object_valid
            if not passed:
                notes = "possible_searchability_render_geometry_or_object_failure"
            elif rect.width > rect.height:
                notes = "searchable_landscape_page_independently_rendered"
            elif not primary_text.strip() and secondary_text.strip():
                notes = "searchable_page_validated_by_secondary_extractor"
            elif nonwhite_ratio < 0.0005:
                notes = "searchable_sparse_or_low_ink_page_independently_rendered"
            else:
                notes = "searchable_rendered_page"
            rows.append(
                {
                    "checkpoint_code": CHECKPOINT_CODE,
                    "page_number": page_number,
                    "width_pt": round(rect.width, 3),
                    "height_pt": round(rect.height, 3),
                    "primary_text_chars": len(primary_text),
                    "secondary_text_chars": len(secondary_text),
                    "searchable_text_chars": len(searchable_text),
                    "image_count": image_count,
                    "drawing_count": drawing_count,
                    "block_count": block_count,
                    "render_width": pix.width,
                    "render_height": pix.height,
                    "render_bytes": len(samples),
                    "mean_grayscale": round(mean, 3),
                    "nonwhite_ratio": round(nonwhite_ratio, 6),
                    "status": "passed" if passed else "failed",
                    "notes": notes,
                    "checked_at": generated_at,
                }
            )
            if page_number in sample_pages:
                proof = page.get_pixmap(matrix=fitz.Matrix(0.75, 0.75), alpha=False)
                path = proof_dir / f"page-{page_number:03d}.png"
                proof.save(path)
                if path.stat().st_size < 5000:
                    raise RuntimeError({"independent_publication_proof_too_small": {"page": page_number, "bytes": path.stat().st_size}})
                proofs.append(path)
    finally:
        doc.close()

    failed = [row for row in rows if row["status"] != "passed"]
    searchable = sum(1 for row in rows if row["searchable_text_chars"] > 0)
    summary = {
        "status": "passed" if len(rows) == 537 and searchable == 537 and not failed else "failed",
        "publication": publication.relative_to(project).as_posix(),
        "publication_bytes": publication.stat().st_size,
        "publication_sha256": sha256_file(publication),
        "page_count": len(rows),
        "searchable_pages": searchable,
        "failed_pages": len(failed),
        "failed_page_numbers": [row["page_number"] for row in failed],
        "landscape_pages": sum(1 for row in rows if row["width_pt"] > row["height_pt"]),
        "secondary_extractor_pages": sum(1 for row in rows if row["primary_text_chars"] == 0 and row["secondary_text_chars"] > 0),
        "sample_proofs": [path.relative_to(project).as_posix() for path in proofs],
        "checked_at": generated_at,
    }
    if summary["status"] != "passed":
        raise RuntimeError({"independent_publication_page_qa_failed": summary})
    return rows, summary, proofs


def workbook_qa(path: Path, inherited_sheets: list[str], managed_sheets: list[str]) -> dict[str, Any]:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=False)
    sheets = list(wb.sheetnames)
    formula_count = 0
    errors: list[str] = []
    tokens = ("#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "#N/A", "#NUM!", "#NULL!")
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                value = cell.value
                if isinstance(value, str) and value.startswith("="):
                    formula_count += 1
                if isinstance(value, str) and any(token in value for token in tokens):
                    errors.append(f"{sheet.title}!{cell.coordinate}:{value}")
    wb.close()
    missing_inherited = [name for name in inherited_sheets if name not in sheets]
    missing_managed = [name for name in managed_sheets if name not in sheets]
    result = {
        "status": "passed" if not errors and not missing_inherited and not missing_managed else "failed",
        "current_sheet_count": len(sheets),
        "original_sheets_preserved": not missing_inherited,
        "missing_inherited_sheets": missing_inherited,
        "managed_sheets": managed_sheets,
        "missing_managed_sheets": missing_managed,
        "formula_count": formula_count,
        "formula_error_count": len(errors),
        "formula_errors": errors,
        "bytes": path.stat().st_size,
        "sha256": sha256_file(path),
    }
    if result["status"] != "passed":
        raise RuntimeError({"section4_final_workbook_qa_failed": result})
    return result


def augment_workbook(workbook: Path, controls: list[dict[str, Any]], database_summary: dict[str, Any], page_summary: dict[str, Any], *, final_state: str) -> dict[str, Any]:
    from openpyxl import load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    wb = load_workbook(workbook)
    inherited = list(wb.sheetnames)
    managed = ["S4S3 CP3 Final QA", "Section 4 Final Release"]
    for name in managed:
        if name in wb.sheetnames:
            del wb[name]

    navy, teal, gold, white, pale = "17365D", "1F6D73", "C9A227", "FFFFFF", "EAF3F3"

    def finish(ws: Any, header_row: int = 2) -> None:
        ws.freeze_panes = f"A{header_row + 1}"
        ws.auto_filter.ref = ws.dimensions
        for cell in ws[header_row]:
            cell.font = Font(bold=True, color=white)
            cell.fill = PatternFill("solid", fgColor=navy)
        for row_number in range(header_row + 1, ws.max_row + 1):
            if row_number % 2:
                for cell in ws[row_number]:
                    cell.fill = PatternFill("solid", fgColor=pale)
        for column in ws.columns:
            width = min(72, max(10, max(len(str(cell.value or "")) for cell in column) + 2))
            ws.column_dimensions[column[0].column_letter].width = width
            for cell in column:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
        ws["A1"].font = Font(bold=True, color=white, size=13)
        ws["A1"].fill = PatternFill("solid", fgColor=teal)
        ws["B1"].font = Font(bold=True, color=white, size=13)
        ws["B1"].fill = PatternFill("solid", fgColor=teal)

    ws = wb.create_sheet("S4S3 CP3 Final QA")
    ws.append(["Section 4 Session 3 Checkpoint 3", "Independent final acceptance matrix"])
    ws.append(["Control", "Expected", "Observed", "Status", "Checked at"])
    for row in controls:
        ws.append([row["control_key"], str(row["expected"]), json.dumps(row["observed"], ensure_ascii=False) if isinstance(row["observed"], (dict, list)) else row["observed"], row["status"].upper(), row["checked_at"]])
    finish(ws)

    ws = wb.create_sheet("Section 4 Final Release")
    ws.append(["Human Pathogen Database", "Remediation Section 4 of 5 final release"])
    ws.append(["Control", "Value"])
    rows = [
        ("Release code", FINAL_RELEASE_CODE),
        ("Response", 72),
        ("Session", "Session 3 of 3 COMPLETE"),
        ("Checkpoint", "Checkpoint 3 of 3 COMPLETE"),
        ("Section state", final_state),
        ("Database physical tables", database_summary["physical_tables"]),
        ("Database logical tables", database_summary["logical_tables"]),
        ("Source/version controls", database_summary["source_version_controls"]),
        ("Publication pages independently audited", page_summary["page_count"]),
        ("Searchable pages", page_summary["searchable_pages"]),
        ("Graphics assets governed", database_summary["graphics_assets"]),
        ("Cross-artifact drift domains", database_summary["drift_domains"]),
        ("Accepted predecessor mutated", "no"),
        ("Frozen Section 3 release mutated", "no"),
        ("Cover/spine finalization", "deferred to final print-production controls in Section 5"),
        ("Next", NEXT_SECTION),
    ]
    for row in rows:
        ws.append(list(row))
    finish(ws)
    for cell in ws[2]:
        cell.fill = PatternFill("solid", fgColor=gold)
        cell.font = Font(bold=True, color="000000")

    wb.save(workbook)
    return workbook_qa(workbook, inherited, managed)


def persist_candidate(db: Path, controls: list[dict[str, Any]], page_rows: list[dict[str, Any]], workbook_status: str, generated_at: str) -> dict[str, Any]:
    con = sqlite3.connect(db)
    try:
        con.execute("PRAGMA foreign_keys=ON")
        con.execute("BEGIN IMMEDIATE")
        con.executescript(
            """
            CREATE TABLE IF NOT EXISTS section4_session3_checkpoint3_acceptance (
              section4_session3_checkpoint3_acceptance_id INTEGER PRIMARY KEY,
              checkpoint_code TEXT NOT NULL,
              control_key TEXT NOT NULL,
              expected_json TEXT NOT NULL,
              observed_json TEXT NOT NULL,
              status TEXT NOT NULL,
              checked_at TEXT NOT NULL,
              UNIQUE(checkpoint_code,control_key)
            );
            CREATE TABLE IF NOT EXISTS section4_session3_checkpoint3_page_audit (
              section4_session3_checkpoint3_page_audit_id INTEGER PRIMARY KEY,
              checkpoint_code TEXT NOT NULL,
              page_number INTEGER NOT NULL,
              width_pt REAL NOT NULL,
              height_pt REAL NOT NULL,
              primary_text_chars INTEGER NOT NULL,
              secondary_text_chars INTEGER NOT NULL,
              searchable_text_chars INTEGER NOT NULL,
              image_count INTEGER NOT NULL,
              drawing_count INTEGER NOT NULL,
              block_count INTEGER NOT NULL,
              render_width INTEGER NOT NULL,
              render_height INTEGER NOT NULL,
              render_bytes INTEGER NOT NULL,
              mean_grayscale REAL NOT NULL,
              nonwhite_ratio REAL NOT NULL,
              status TEXT NOT NULL,
              notes TEXT,
              checked_at TEXT NOT NULL,
              UNIQUE(checkpoint_code,page_number)
            );
            CREATE TABLE IF NOT EXISTS section4_final_release (
              release_code TEXT PRIMARY KEY,
              response_number INTEGER NOT NULL,
              section_label TEXT NOT NULL,
              session_label TEXT NOT NULL,
              checkpoint_label TEXT NOT NULL,
              state TEXT NOT NULL,
              database_status TEXT NOT NULL,
              source_version_status TEXT NOT NULL,
              publication_page_status TEXT NOT NULL,
              graphics_status TEXT NOT NULL,
              drift_status TEXT NOT NULL,
              workbook_status TEXT NOT NULL,
              application_status TEXT NOT NULL,
              tracking_status TEXT NOT NULL,
              index_manifest_status TEXT NOT NULL,
              cover_spine_status TEXT NOT NULL,
              accepted_predecessor_mutated INTEGER NOT NULL,
              frozen_section3_mutated INTEGER NOT NULL,
              next_section TEXT NOT NULL,
              recorded_at TEXT NOT NULL
            );
            """
        )
        con.execute("DELETE FROM section4_session3_checkpoint3_acceptance WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        for row in controls:
            con.execute(
                "INSERT INTO section4_session3_checkpoint3_acceptance(checkpoint_code,control_key,expected_json,observed_json,status,checked_at) VALUES (?,?,?,?,?,?)",
                (CHECKPOINT_CODE,row["control_key"],json.dumps(row["expected"],ensure_ascii=False),json.dumps(row["observed"],ensure_ascii=False),row["status"],generated_at),
            )
        con.execute("DELETE FROM section4_session3_checkpoint3_page_audit WHERE checkpoint_code=?", (CHECKPOINT_CODE,))
        for row in page_rows:
            con.execute(
                """INSERT INTO section4_session3_checkpoint3_page_audit
                (checkpoint_code,page_number,width_pt,height_pt,primary_text_chars,secondary_text_chars,searchable_text_chars,image_count,drawing_count,block_count,render_width,render_height,render_bytes,mean_grayscale,nonwhite_ratio,status,notes,checked_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (CHECKPOINT_CODE,row["page_number"],row["width_pt"],row["height_pt"],row["primary_text_chars"],row["secondary_text_chars"],row["searchable_text_chars"],row["image_count"],row["drawing_count"],row["block_count"],row["render_width"],row["render_height"],row["render_bytes"],row["mean_grayscale"],row["nonwhite_ratio"],row["status"],row["notes"],generated_at),
            )
        con.execute(
            """INSERT INTO section4_final_release
            (release_code,response_number,section_label,session_label,checkpoint_label,state,database_status,source_version_status,
             publication_page_status,graphics_status,drift_status,workbook_status,application_status,tracking_status,index_manifest_status,
             cover_spine_status,accepted_predecessor_mutated,frozen_section3_mutated,next_section,recorded_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            ON CONFLICT(release_code) DO UPDATE SET
              response_number=excluded.response_number,section_label=excluded.section_label,session_label=excluded.session_label,
              checkpoint_label=excluded.checkpoint_label,state=excluded.state,database_status=excluded.database_status,
              source_version_status=excluded.source_version_status,publication_page_status=excluded.publication_page_status,
              graphics_status=excluded.graphics_status,drift_status=excluded.drift_status,workbook_status=excluded.workbook_status,
              application_status=excluded.application_status,tracking_status=excluded.tracking_status,index_manifest_status=excluded.index_manifest_status,
              cover_spine_status=excluded.cover_spine_status,accepted_predecessor_mutated=excluded.accepted_predecessor_mutated,
              frozen_section3_mutated=excluded.frozen_section3_mutated,next_section=excluded.next_section,recorded_at=excluded.recorded_at""",
            (FINAL_RELEASE_CODE,72,"Remediation Section 4 of 5","Session 3 of 3","Checkpoint 3 of 3","candidate_finalization_pending","ok","passed","passed","passed","passed",workbook_status,"pending_final_audit","current_and_final_rebuild_gated","post_database_final_rebuild_external_gate","deferred_section5_print_production",0,0,NEXT_SECTION,generated_at),
        )
        if table_exists(con, "section4_session3_release_risk"):
            con.execute("UPDATE section4_session3_release_risk SET status='validated_checkpoint3_pending_final_state',disposition='Independently revalidated at Checkpoint 3; final section-state synchronization is pending.',checked_at=? WHERE risk_key!='cover_spine_and_press_finalization'", (generated_at,))
            con.execute("UPDATE section4_session3_release_risk SET status='deferred_section5',disposition='Final cover-wrap and spine arithmetic remain governed by final interior, stock, binding, printer template, and color-profile decisions in Section 5.',checked_at=? WHERE risk_key='cover_spine_and_press_finalization'", (generated_at,))
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or fk:
            raise RuntimeError({"checkpoint3_candidate_persistence_failed": {"integrity": integrity, "foreign_keys": fk[:20]}})
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()
    return {"status": "passed", "bytes": db.stat().st_size, "sha256": sha256_file(db)}


def prepare_final_verification(
    project: Path,
    db: Path,
    workbook: Path,
    database_qa: dict[str, Any],
    workbook_base_qa: dict[str, Any],
    application_base_qa: dict[str, Any],
    publication_base_qa: dict[str, Any],
    *,
    generated_at: str,
    response_number: int,
) -> dict[str, Any]:
    if response_number != 72:
        raise RuntimeError({"unexpected_response": response_number})
    application = locate_by_hash(project, "human_pathogen_app.py", APPLICATION_SHA256)
    publication = locate_by_hash(project, "*Integrated Manuscript*.pdf", PUBLICATION_SHA256)
    editable = locate_by_hash(project, "*Editable Integrated Manuscript Assembly*.docx", EDITABLE_ASSEMBLY_SHA256)
    controls, database_summary = audit_candidate_database(db, generated_at)
    page_rows, page_summary, proofs = audit_publication_pages(project, publication, generated_at)
    controls.append({"checkpoint_code":CHECKPOINT_CODE,"control_key":"independent_publication_render","expected":"537 searchable rendered pages","observed":{"pages":page_summary["page_count"],"searchable":page_summary["searchable_pages"],"failures":page_summary["failed_pages"]},"status":"passed" if page_summary["status"]=="passed" else "failed","checked_at":generated_at})
    controls.append({"checkpoint_code":CHECKPOINT_CODE,"control_key":"immutable_application_source","expected":APPLICATION_SHA256,"observed":sha256_file(application),"status":"passed" if sha256_file(application)==APPLICATION_SHA256 else "failed","checked_at":generated_at})
    controls.append({"checkpoint_code":CHECKPOINT_CODE,"control_key":"immutable_publication","expected":PUBLICATION_SHA256,"observed":sha256_file(publication),"status":"passed" if sha256_file(publication)==PUBLICATION_SHA256 else "failed","checked_at":generated_at})
    controls.append({"checkpoint_code":CHECKPOINT_CODE,"control_key":"immutable_editable_assembly","expected":EDITABLE_ASSEMBLY_SHA256,"observed":sha256_file(editable),"status":"passed" if sha256_file(editable)==EDITABLE_ASSEMBLY_SHA256 else "failed","checked_at":generated_at})
    artifact_files = {
        "Tracking": len([p for p in (project / "Tracking").rglob("*") if p.is_file()]) if (project / "Tracking").exists() else 0,
        "Recovery": len([p for p in (project / "Recovery").rglob("*") if p.is_file()]) if (project / "Recovery").exists() else 0,
        "Indexes": len([p for p in (project / "Indexes").rglob("*") if p.is_file()]) if (project / "Indexes").exists() else 0,
        "Manifest": len([p for p in (project / "Manifest").rglob("*") if p.is_file()]) if (project / "Manifest").exists() else 0,
    }
    controls.append({"checkpoint_code":CHECKPOINT_CODE,"control_key":"tracking_recovery_index_manifest_surfaces","expected":"all nonempty","observed":artifact_files,"status":"passed" if all(value>0 for value in artifact_files.values()) else "failed","checked_at":generated_at})
    failed = [row for row in controls if row["status"] != "passed"]
    if failed:
        raise RuntimeError({"checkpoint3_preparation_controls_failed": failed})

    workbook_candidate_qa = augment_workbook(workbook, controls, database_summary, page_summary, final_state="candidate_finalization_pending")
    persistence = persist_candidate(db, controls, page_rows, workbook_candidate_qa["status"], generated_at)

    qa_dir = project / "QA" / "Section 4 Session 3" / "Checkpoint 3"
    qa_dir.mkdir(parents=True, exist_ok=True)
    json_write(qa_dir / "INDEPENDENT_ACCEPTANCE_CONTROLS.json", controls)
    csv_write(qa_dir / "INDEPENDENT_ACCEPTANCE_CONTROLS.csv", controls, ["control_key","expected","observed","status","checked_at"])
    json_write(qa_dir / "INDEPENDENT_PUBLICATION_PAGE_QA.json", page_rows)
    csv_write(qa_dir / "INDEPENDENT_PUBLICATION_PAGE_QA.csv", page_rows, ["page_number","width_pt","height_pt","primary_text_chars","secondary_text_chars","searchable_text_chars","image_count","drawing_count","block_count","render_width","render_height","render_bytes","mean_grayscale","nonwhite_ratio","status","notes","checked_at"])
    candidate_qa = {
        "schema": "mrhpd-section4-session3-checkpoint3-candidate-verification-1.0",
        "generated_at": generated_at,
        "status": "passed",
        "response": response_number,
        "checkpoint_code": CHECKPOINT_CODE,
        "database": database_summary,
        "controls": {"count": len(controls), "failed": 0},
        "publication": page_summary,
        "workbook": workbook_candidate_qa,
        "application_invariant": {"path":application.relative_to(project).as_posix(),"bytes":application.stat().st_size,"sha256":sha256_file(application)},
        "publication_invariant": {"path":publication.relative_to(project).as_posix(),"bytes":publication.stat().st_size,"sha256":sha256_file(publication)},
        "editable_assembly_invariant": {"path":editable.relative_to(project).as_posix(),"bytes":editable.stat().st_size,"sha256":sha256_file(editable)},
        "persistence": persistence,
        "accepted_predecessor_mutated": False,
        "frozen_section3_release_mutated": False,
        "checkpoint_3_prepared": True,
        "section4_final_release_declared": False,
    }
    json_write(qa_dir / "CHECKPOINT3_CANDIDATE_VERIFICATION_QA.json", candidate_qa)
    application_qa = dict(application_base_qa)
    application_qa.update({"status":"passed","main_application_unchanged":True,"checkpoint3_candidate_controls":"passed"})
    publication_qa = dict(publication_base_qa)
    publication_qa.update({"status":"passed","publication_pages":537,"searchable_pages":537,"publication_unchanged":True,"editable_assembly_unchanged":True,"independent_page_qa":page_summary})
    merged_workbook_qa = dict(workbook_base_qa)
    merged_workbook_qa.update(workbook_candidate_qa)
    return {
        "qa": candidate_qa,
        "controls": controls,
        "database_summary": database_summary,
        "page_rows": page_rows,
        "page_summary": page_summary,
        "proofs": proofs,
        "workbook_qa": merged_workbook_qa,
        "application_qa": application_qa,
        "publication_qa": publication_qa,
        "paths": {"application":application,"publication":publication,"editable":editable},
    }


def generic_final_state(db: Path) -> dict[str, Any]:
    con = sqlite3.connect(db)
    try:
        cp3 = checkpoint_state(con, CHECKPOINT_CODE)
        session_release = safe_scalar(con, "SELECT state FROM section4_session_release WHERE release_code='MRHPD-V3-CP4-S3-COMPLETE'", (), None) if table_exists(con, "section4_session_release") else None
        response72 = safe_count(con, "thread_response_reconciliation_cp3", "response_key='R72'")
        return {"checkpoint_state": cp3, "session_release_state": session_release, "response72_records": response72}
    finally:
        con.close()


def write_final_application_audit(project: Path, db: Path, workbook: Path, application: Path, publication: Path, editable: Path, generated_at: str) -> tuple[Path, Path, dict[str, Any]]:
    app_dir = project / "App"
    qa_dir = project / "QA" / "Section 4 Session 3" / "Checkpoint 3"
    audit = app_dir / "section4_final_release_audit.py"
    output = qa_dir / "SECTION4_FINAL_APPLICATION_AUDIT.json"
    expected = {
        "application_sha256": sha256_file(application),
        "publication_sha256": sha256_file(publication),
        "editable_sha256": sha256_file(editable),
        "workbook_sha256": sha256_file(workbook),
        "required_sheets": ["S4S3 Source Sweep","S4S3 Page QA","S4S3 Graphics QA","S4S3 Final Drift","S4S3 Risk Closure","S4S3 CP2 Readiness","S4S3 CP3 Final QA","Section 4 Final Release"],
    }
    expected_json = json.dumps(expected, ensure_ascii=False)
    source = f'''#!/usr/bin/env python3
from __future__ import annotations
import argparse,hashlib,json,sqlite3
from pathlib import Path
from openpyxl import load_workbook
from pypdf import PdfReader
EXPECTED=json.loads({expected_json!r})
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as handle:
  for block in iter(lambda:handle.read(1024*1024),b''): h.update(block)
 return h.hexdigest()
def exists(con,table):
 return con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?",(table,)).fetchone() is not None
p=argparse.ArgumentParser(description='MRHPD Section 4 final release audit')
p.add_argument('--db',type=Path,required=True); p.add_argument('--workbook',type=Path,required=True)
p.add_argument('--publication',type=Path,required=True); p.add_argument('--editable',type=Path,required=True)
p.add_argument('--app',type=Path,required=True); p.add_argument('--output',type=Path,required=True)
a=p.parse_args(); con=sqlite3.connect(a.db)
try:
 integrity=con.execute('PRAGMA integrity_check').fetchone()[0]; fk=list(con.execute('PRAGMA foreign_key_check'))
 response72=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R72'").fetchone()[0]
 cp3=con.execute("SELECT state FROM section4_checkpoint WHERE checkpoint_code='MRHPD-V3-CP4-S3-CP3'").fetchone() if exists(con,'section4_checkpoint') else None
 session_release=con.execute("SELECT state FROM section4_session_release WHERE release_code='MRHPD-V3-CP4-S3-COMPLETE'").fetchone() if exists(con,'section4_session_release') else None
 final_release=con.execute("SELECT state FROM section4_final_release WHERE release_code='MRHPD-V3-CP4-COMPLETE'").fetchone() if exists(con,'section4_final_release') else None
 controls=con.execute("SELECT COUNT(*) FROM section4_session3_checkpoint3_acceptance WHERE checkpoint_code='MRHPD-V3-CP4-S3-CP3' AND status='passed'").fetchone()[0] if exists(con,'section4_session3_checkpoint3_acceptance') else 0
 control_failures=con.execute("SELECT COUNT(*) FROM section4_session3_checkpoint3_acceptance WHERE checkpoint_code='MRHPD-V3-CP4-S3-CP3' AND status!='passed'").fetchone()[0] if exists(con,'section4_session3_checkpoint3_acceptance') else 1
 pages=con.execute("SELECT COUNT(*) FROM section4_session3_checkpoint3_page_audit WHERE checkpoint_code='MRHPD-V3-CP4-S3-CP3' AND status='passed'").fetchone()[0] if exists(con,'section4_session3_checkpoint3_page_audit') else 0
finally: con.close()
wb=load_workbook(a.workbook,read_only=True,data_only=False); sheets=list(wb.sheetnames); wb.close()
reader=PdfReader(str(a.publication)); searchable=sum(1 for page in reader.pages if (page.extract_text() or '').strip())
checks={{
 'sqlite_integrity':integrity=='ok','foreign_keys':not fk,'response72':response72==1,
 'checkpoint3_state':cp3==('session_complete',),'session3_release_state':session_release==('session_complete',),
 'section4_release_state':final_release==('section_complete',),'acceptance_controls':controls>=23 and control_failures==0,
 'independent_page_records':pages==537,'workbook_sheets':all(name in sheets for name in EXPECTED['required_sheets']),
 'workbook_invariant':sha(a.workbook)==EXPECTED['workbook_sha256'],'publication_pages':len(reader.pages)==537 and searchable==537,
 'publication_invariant':sha(a.publication)==EXPECTED['publication_sha256'],'editable_invariant':sha(a.editable)==EXPECTED['editable_sha256'],
 'application_invariant':sha(a.app)==EXPECTED['application_sha256'],
}}
result={{'schema':'mrhpd-section4-final-application-audit-1.0','checked_at':{generated_at!r},'status':'passed' if all(checks.values()) else 'failed','checks':checks,'workbook_sheets':len(sheets),'publication_pages':len(reader.pages),'searchable_pages':searchable,'acceptance_controls':controls,'independent_page_records':pages}}
a.output.parent.mkdir(parents=True,exist_ok=True); a.output.write_text(json.dumps(result,indent=2)+chr(10),encoding='utf-8')
print(json.dumps(result,indent=2)); raise SystemExit(0 if result['status']=='passed' else 1)
'''
    text_write(audit, source)
    compile_result = subprocess.run([sys.executable, "-m", "py_compile", str(audit)], text=True, capture_output=True, timeout=120)
    if compile_result.returncode != 0:
        raise RuntimeError({"section4_final_audit_compile_failed": {"stdout": compile_result.stdout, "stderr": compile_result.stderr}})
    result = subprocess.run([sys.executable,str(audit),"--db",str(db),"--workbook",str(workbook),"--publication",str(publication),"--editable",str(editable),"--app",str(application),"--output",str(output)], text=True, capture_output=True, timeout=900)
    if result.returncode != 0:
        raise RuntimeError({"section4_final_application_audit_failed": {"stdout": result.stdout[-20000:], "stderr": result.stderr[-20000:]}})
    qa = json.loads(output.read_text(encoding="utf-8"))
    qa.update({"audit_path":audit.relative_to(project).as_posix(),"output_path":output.relative_to(project).as_posix(),"audit_sha256":sha256_file(audit),"stdout":result.stdout[-12000:]})
    return audit, output, qa


def build_reports(project: Path, final_qa: dict[str, Any], controls: list[dict[str, Any]], proofs: list[Path]) -> list[Path]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from pypdf import PdfReader
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    report_dir = project / "Reports" / "Section 4 Session 3" / "Checkpoint 3"
    report_dir.mkdir(parents=True, exist_ok=True)
    stem = "MRHPD v3.0.0a Section 4 Final Release and Independent Verification"
    docx_path = report_dir / f"{stem}.docx"
    pdf_path = report_dir / f"{stem}.pdf"
    xlsx_path = report_dir / f"{stem} Register.xlsx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65); section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7); section.right_margin = Inches(0.7)
    title = doc.add_heading("Human Pathogen Database", 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Remediation Section 4 of 5 — Final Release and Independent Verification"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Response 72 • Session 3 of 3 COMPLETE • Checkpoint 3 of 3 COMPLETE").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Final disposition", level=1)
    doc.add_paragraph("The exact Response 71 release candidate was reconstructed from the governed Response 69 baseline and independently revalidated. Remediation Section 4 is complete. The accepted predecessor, frozen Section 3 release, 537-page publication, editable assembly, and main application source remain unchanged.")
    doc.add_heading("Acceptance matrix", level=1)
    table = doc.add_table(rows=1, cols=4); table.style = "Table Grid"
    for idx, heading in enumerate(["Control", "Expected", "Observed", "Status"]): table.rows[0].cells[idx].text = heading
    for row in controls:
        cells = table.add_row().cells
        cells[0].text = row["control_key"]; cells[1].text = str(row["expected"]); cells[2].text = json.dumps(row["observed"],ensure_ascii=False) if isinstance(row["observed"],(dict,list)) else str(row["observed"]); cells[3].text = row["status"].upper()
    doc.add_heading("Publication page-level independent audit", level=1)
    pub = final_qa["publication"]
    doc.add_paragraph(f"Pages audited: {pub['page_count']}; searchable pages: {pub['searchable_pages']}; failed pages: {pub['failed_pages']}; landscape pages: {pub['landscape_pages']}; secondary-extractor pages: {pub['secondary_extractor_pages']}.")
    for proof in proofs:
        doc.add_picture(str(proof), width=Inches(2.75))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(proof.name).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Release boundary", level=1)
    doc.add_paragraph("Section 4 closes clinical-reference synchronization and release governance. Final cover-wrap/spine arithmetic, printer-template, stock, binding, and color-profile decisions remain explicitly deferred to Section 5 print-production controls.")
    doc.save(docx_path)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CenterTitleFinal", parent=styles["Title"], alignment=TA_CENTER, textColor=colors.HexColor("#17365D")))
    story: list[Any] = [Paragraph("Human Pathogen Database", styles["CenterTitleFinal"]), Paragraph("Remediation Section 4 of 5 — Final Release", styles["Heading1"]), Paragraph("Response 72 • Session 3 of 3 COMPLETE • Checkpoint 3 of 3 COMPLETE", styles["Normal"]), Spacer(1,0.15*inch)]
    story += [Paragraph("Final disposition", styles["Heading1"]), Paragraph("The exact Response 71 release candidate was reconstructed from the governed Response 69 baseline and independently revalidated. Remediation Section 4 is complete; immutable publication and application artifacts remain unchanged.", styles["BodyText"]), Spacer(1,0.12*inch)]
    summary_data = [["Metric","Result"],["Acceptance controls",str(final_qa["controls"]["count"])],["Database tables",str(final_qa["database"]["physical_tables"])],["Source/version controls",str(final_qa["database"]["source_version_controls"])],["Publication pages",str(final_qa["publication"]["page_count"])],["Graphics assets",str(final_qa["database"]["graphics_assets"])],["Drift domains",str(final_qa["database"]["drift_domains"])],["Application audit",final_qa["application_audit"]["status"].upper()],["Section state","COMPLETE"]]
    summary_table = Table(summary_data, colWidths=[3.2*inch,3.7*inch])
    summary_table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#17365D")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.35,colors.grey),("FONTSIZE",(0,0),(-1,-1),9)]))
    story += [summary_table, PageBreak(), Paragraph("Independent acceptance controls", styles["Heading1"])]
    control_data = [["Control","Status","Observed"]] + [[row["control_key"],row["status"].upper(),str(row["observed"])[:120]] for row in controls]
    control_table = Table(control_data, colWidths=[2.55*inch,0.8*inch,3.65*inch], repeatRows=1)
    control_table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#1F6D73")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),0.3,colors.grey),("VALIGN",(0,0),(-1,-1),"TOP"),("FONTSIZE",(0,0),(-1,-1),6.8)]))
    story += [control_table, PageBreak(), Paragraph("Publication audit and release boundary", styles["Heading1"]), Paragraph(f"All {pub['page_count']} pages passed independent searchable-render QA. Final cover and spine production remain deferred until Section 5 fixes the final interior, stock, binding, printer template, and color profile.", styles["BodyText"])]
    SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36, title=stem, author="Brent McAnulty, M.D.").build(story)

    wb = Workbook(); wb.remove(wb.active)
    datasets = {
        "Summary": [["Control","Value"],["Status","passed"],["Response",72],["Section","4 of 5 COMPLETE"],["Session","3 of 3 COMPLETE"],["Checkpoint","3 of 3 COMPLETE"],["Database tables",final_qa["database"]["physical_tables"]],["Publication pages",final_qa["publication"]["page_count"]],["Graphics assets",final_qa["database"]["graphics_assets"]],["Drift domains",final_qa["database"]["drift_domains"]]],
        "Acceptance Controls": [["Control","Expected","Observed","Status","Checked at"]] + [[r["control_key"],str(r["expected"]),json.dumps(r["observed"],ensure_ascii=False) if isinstance(r["observed"],(dict,list)) else str(r["observed"]),r["status"],r["checked_at"]] for r in controls],
        "Application Audit": [["Check","Passed"]] + [[key,value] for key,value in final_qa["application_audit"]["checks"].items()],
        "Release Boundary": [["Item","Disposition"],["Section 4","complete"],["Cover/spine","deferred to Section 5 print-production controls"],["Accepted predecessor mutated","no"],["Frozen Section 3 release mutated","no"],["Next",NEXT_SECTION]],
    }
    for name,data in datasets.items():
        ws=wb.create_sheet(name)
        for row in data: ws.append(row)
        for cell in ws[1]: cell.font=Font(bold=True,color="FFFFFF"); cell.fill=PatternFill("solid",fgColor="17365D")
        ws.freeze_panes="A2"; ws.auto_filter.ref=ws.dimensions
        for column in ws.columns:
            ws.column_dimensions[column[0].column_letter].width=min(72,max(10,max(len(str(cell.value or "")) for cell in column)+2))
            for cell in column: cell.alignment=Alignment(wrap_text=True,vertical="top")
    wb.save(xlsx_path)

    with zipfile.ZipFile(docx_path) as zf:
        if zf.testzip() is not None: raise RuntimeError("Final DOCX CRC failed")
    with zipfile.ZipFile(xlsx_path) as zf:
        if zf.testzip() is not None: raise RuntimeError("Final XLSX CRC failed")
    reader = PdfReader(str(pdf_path)); text_chars = sum(len(page.extract_text() or "") for page in reader.pages)
    if len(reader.pages) < 3 or text_chars < 2200:
        raise RuntimeError({"final_report_pdf_validation": {"pages":len(reader.pages),"text_chars":text_chars}})
    qa = {"status":"passed","docx":{"path":docx_path.relative_to(project).as_posix(),"bytes":docx_path.stat().st_size,"sha256":sha256_file(docx_path)},"pdf":{"path":pdf_path.relative_to(project).as_posix(),"bytes":pdf_path.stat().st_size,"sha256":sha256_file(pdf_path),"pages":len(reader.pages),"text_chars":text_chars},"xlsx":{"path":xlsx_path.relative_to(project).as_posix(),"bytes":xlsx_path.stat().st_size,"sha256":sha256_file(xlsx_path),"sheets":list(datasets)}}
    json_write(report_dir / "REPORT_QA.json", qa)
    return [docx_path,pdf_path,xlsx_path,report_dir/"REPORT_QA.json"]


def complete_section4_release(
    project: Path,
    db: Path,
    workbook: Path,
    database_qa: dict[str, Any],
    workbook_base_qa: dict[str, Any],
    application_base_qa: dict[str, Any],
    publication_base_qa: dict[str, Any],
    *,
    candidate: dict[str, Any],
    generated_at: str,
) -> dict[str, Any]:
    states = generic_final_state(db)
    if states["checkpoint_state"] != "session_complete" or states["session_release_state"] != "session_complete" or states["response72_records"] != 1:
        raise RuntimeError({"generic_final_state_failed": states})

    con = sqlite3.connect(db)
    try:
        con.execute("BEGIN IMMEDIATE")
        con.execute("UPDATE section4_final_release SET state='section_complete',database_status='ok',source_version_status='passed',publication_page_status='passed',graphics_status='passed',drift_status='passed',workbook_status='pending_final_save',application_status='pending_final_audit',tracking_status='current_and_final_rebuild_gated',index_manifest_status='post_database_final_rebuild_external_gate',recorded_at=? WHERE release_code=?", (generated_at,FINAL_RELEASE_CODE))
        if table_exists(con, "section4_session3_release_risk"):
            con.execute("UPDATE section4_session3_release_risk SET status='resolved_checkpoint3',disposition='Independent Checkpoint 3 acceptance, clean extraction, final state, and complete-restore gates passed.',checked_at=? WHERE risk_key='section4_final_release_signoff'", (generated_at,))
            con.execute("UPDATE section4_session3_release_risk SET status='resolved_checkpoint3',disposition='Independently revalidated and included in the final Section 4 release.',checked_at=? WHERE risk_key NOT IN ('section4_final_release_signoff','cover_spine_and_press_finalization')", (generated_at,))
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or fk:
            raise RuntimeError({"section4_state_update_failed": {"integrity":integrity,"foreign_keys":fk[:20]}})
        con.commit()
    except Exception:
        con.rollback(); raise
    finally:
        con.close()

    final_workbook_qa = augment_workbook(workbook, candidate["controls"], candidate["database_summary"], candidate["page_summary"], final_state="section_complete")
    application = candidate["paths"]["application"]
    publication = candidate["paths"]["publication"]
    editable = candidate["paths"]["editable"]
    audit_path, audit_output, application_audit = write_final_application_audit(project,db,workbook,application,publication,editable,generated_at)

    con = sqlite3.connect(db)
    try:
        con.execute("BEGIN IMMEDIATE")
        con.execute("UPDATE section4_final_release SET workbook_status=?,application_status=?,recorded_at=? WHERE release_code=?", (final_workbook_qa["status"],application_audit["status"],generated_at,FINAL_RELEASE_CODE))
        final_state = con.execute("SELECT state FROM section4_final_release WHERE release_code=?", (FINAL_RELEASE_CODE,)).fetchone()
        acceptance_failures = con.execute("SELECT COUNT(*) FROM section4_session3_checkpoint3_acceptance WHERE checkpoint_code=? AND status!='passed'", (CHECKPOINT_CODE,)).fetchone()[0]
        page_failures = con.execute("SELECT COUNT(*) FROM section4_session3_checkpoint3_page_audit WHERE checkpoint_code=? AND status!='passed'", (CHECKPOINT_CODE,)).fetchone()[0]
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        if integrity != "ok" or fk or final_state != ("section_complete",) or acceptance_failures or page_failures:
            raise RuntimeError({"section4_final_database_gate_failed":{"integrity":integrity,"foreign_keys":fk[:20],"state":final_state,"acceptance_failures":acceptance_failures,"page_failures":page_failures}})
        con.commit()
    except Exception:
        con.rollback(); raise
    finally:
        con.close()

    application_qa = dict(application_base_qa)
    application_qa.update({"status":"passed","main_application_unchanged":True,"section4_final_audit":application_audit,"section4_final_audit_path":audit_path.relative_to(project).as_posix()})
    workbook_qa_merged = dict(workbook_base_qa); workbook_qa_merged.update(final_workbook_qa)
    publication_qa = dict(publication_base_qa)
    publication_qa.update({"status":"passed","publication_pages":537,"searchable_pages":537,"publication_unchanged":True,"editable_assembly_unchanged":True,"independent_page_qa":candidate["page_summary"]})
    final_database_qa = dict(database_qa)
    final_database_qa.update({"status":"passed","integrity":"ok","foreign_key_violations":0,"checkpoint_state":"session_complete","session_release_state":"session_complete","section4_release_state":"section_complete","response72_records":1,"bytes":db.stat().st_size,"sha256":sha256_file(db)})

    final_qa = {
        "schema": "mrhpd-section4-final-release-1.0",
        "generated_at": generated_at,
        "status": "passed",
        "response": 72,
        "section": "Remediation Section 4 of 5 COMPLETE",
        "session": "Session 3 of 3 COMPLETE",
        "checkpoint": "Checkpoint 3 of 3 COMPLETE",
        "database": candidate["database_summary"],
        "controls": {"count":len(candidate["controls"]),"failed":0},
        "publication": candidate["page_summary"],
        "workbook": workbook_qa_merged,
        "application_audit": application_audit,
        "generic_final_state": states,
        "database_final_state": final_database_qa,
        "application_invariant": {"path":application.relative_to(project).as_posix(),"bytes":application.stat().st_size,"sha256":sha256_file(application)},
        "publication_invariant": {"path":publication.relative_to(project).as_posix(),"bytes":publication.stat().st_size,"sha256":sha256_file(publication)},
        "editable_assembly_invariant": {"path":editable.relative_to(project).as_posix(),"bytes":editable.stat().st_size,"sha256":sha256_file(editable)},
        "accepted_predecessor_mutated": False,
        "frozen_section3_release_mutated": False,
        "checkpoint_3_of_3_complete": True,
        "session_3_of_3_complete": True,
        "remediation_section_4_complete": True,
        "cover_spine_finalization": "deferred_section5_print_production",
        "next": NEXT_SECTION,
    }
    report_files = build_reports(project, final_qa, candidate["controls"], candidate["proofs"])
    qa_dir = project / "QA" / "Section 4 Session 3" / "Checkpoint 3"
    json_write(qa_dir / "SECTION4_FINAL_RELEASE_QA.json", final_qa)
    text_write(qa_dir / "SECTION4_FINAL_RELEASE_README.md", "# Remediation Section 4 of 5 — COMPLETE\n\nSession 3 of 3 and Checkpoint 3 of 3 are complete through Response 72. Independent candidate reconstruction, database, workbook, application, publication, graphics, source, drift, tracking, recovery, index, manifest, archive, and clean-extraction gates are represented in the final package. Cover-wrap and spine production remain deferred to Remediation Section 5 of 5.\n")
    return {
        "qa": final_qa,
        "database_qa": final_database_qa,
        "workbook_qa": workbook_qa_merged,
        "application_qa": application_qa,
        "publication_qa": publication_qa,
        "report_files": report_files,
        "critical_paths": {"final_application_audit":audit_path,"final_application_audit_output":audit_output,"final_release_qa":qa_dir/"SECTION4_FINAL_RELEASE_QA.json"},
    }
