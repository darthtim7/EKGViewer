#!/usr/bin/env python3
"""Release-governance controls for MRHPD Section 4 Session 3 Checkpoint 1.

This module is intentionally separate from the generated checkpoint builder.
It evaluates final-session clinical, evidence, graphics, publication,
application, workbook, tracking, and recovery readiness; persists those gates
in SQLite; adds governed workbook surfaces; creates a read-only acceptance
audit; and emits human- and machine-readable release-readiness reports.
"""
from __future__ import annotations

import csv
import hashlib
import json
import re
import sqlite3
import subprocess
import sys
from pathlib import Path
from typing import Any


def sha256_file(path: Path, chunk_size: int = 1024 * 1024) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(chunk_size), b""):
            digest.update(block)
    return digest.hexdigest()


def json_write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False, default=str) + "\n", encoding="utf-8")


def text_write(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value.rstrip() + "\n", encoding="utf-8")


def csv_write(path: Path, rows: list[dict[str, Any]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            normalized: dict[str, Any] = {}
            for field in fields:
                value = row.get(field)
                if isinstance(value, (dict, list)):
                    value = json.dumps(value, ensure_ascii=False)
                normalized[field] = value
            writer.writerow(normalized)


def qident(value: str) -> str:
    return '"' + value.replace('"', '""') + '"'


def table_exists(con: sqlite3.Connection, table: str) -> bool:
    return con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?", (table,)).fetchone() is not None


def logical_tables(con: sqlite3.Connection) -> list[str]:
    rows = [row[0] for row in con.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name")]
    names = set(rows)
    shadow_suffixes = ("_data", "_idx", "_content", "_docsize", "_config")
    return [
        name
        for name in rows
        if not any(name.endswith(suffix) and name[: -len(suffix)] in names for suffix in shadow_suffixes)
    ]


def safe_count(con: sqlite3.Connection, table: str) -> int | None:
    try:
        return int(con.execute(f"SELECT COUNT(*) FROM {qident(table)}").fetchone()[0])
    except Exception:
        return None


def text_columns(con: sqlite3.Connection, table: str) -> list[str]:
    columns: list[str] = []
    try:
        for row in con.execute(f"PRAGMA table_info({qident(table)})"):
            name = str(row[1])
            declared = str(row[2] or "").upper()
            if not declared or any(token in declared for token in ("CHAR", "TEXT", "CLOB")):
                columns.append(name)
    except Exception:
        return []
    return columns


def term_hits(con: sqlite3.Connection, tables: list[str], terms: list[str]) -> int:
    total = 0
    for table in tables:
        columns = text_columns(con, table)[:16]
        if not columns:
            continue
        for term in terms:
            where = " OR ".join(f"LOWER(CAST({qident(column)} AS TEXT)) LIKE ?" for column in columns)
            try:
                total += int(con.execute(f"SELECT COUNT(*) FROM {qident(table)} WHERE {where}", [f"%{term.lower()}%"] * len(columns)).fetchone()[0])
            except Exception:
                continue
    return total


DATA_GATE_SPECS: list[dict[str, Any]] = [
    {
        "gate_key": "taxonomy_nomenclature",
        "domain": "clinical_data",
        "description": "Current taxonomy, aliases, former names, and resolver-first nomenclature remain represented as governed data.",
        "patterns": ["taxonomy", "organism_alias", "resolver"],
        "minimum_tables": 3,
        "minimum_rows": 1,
    },
    {
        "gate_key": "search_disambiguation",
        "domain": "application_search",
        "description": "Search documents, resolver rules, and ambiguity/disambiguation pathways remain queryable.",
        "patterns": ["search_document", "search_resolver", "disambiguation"],
        "minimum_tables": 2,
        "minimum_rows": 1,
    },
    {
        "gate_key": "clinical_profiles_syndromes",
        "domain": "clinical_data",
        "description": "Organism profiles, disease associations, syndromes, manifestations, and mimics remain represented.",
        "patterns": ["clinical_profile", "disease_association", "syndrome", "manifestation", "mimic"],
        "minimum_tables": 4,
        "minimum_rows": 1,
    },
    {
        "gate_key": "laboratory_diagnostics",
        "domain": "diagnostics",
        "description": "Morphology, growth, diagnostic tests, specimen interpretation, and laboratory caveats remain represented.",
        "patterns": ["morphology", "lab_growth", "diagnostic", "specimen"],
        "minimum_tables": 3,
        "minimum_rows": 1,
    },
    {
        "gate_key": "transmission_sources_exposures",
        "domain": "epidemiology",
        "description": "Transmission, reservoirs, immediate sources, exposure contexts, and vectors remain distinct and represented.",
        "patterns": ["transmission", "common_source", "reservoir", "exposure", "vector"],
        "minimum_tables": 2,
        "minimum_rows": 1,
    },
    {
        "gate_key": "treatment_stewardship_duration",
        "domain": "treatment",
        "description": "Treatment contexts, options, decision details, duration rules, stopping criteria, and stewardship remain represented.",
        "patterns": ["treatment_context", "treatment_option", "duration_rule", "stewardship"],
        "minimum_tables": 4,
        "minimum_rows": 1,
    },
    {
        "gate_key": "no_treatment_source_control_reassessment",
        "domain": "treatment_safety",
        "description": "No-treatment pathways, source-control actions, reassessment, narrowing, and stopping logic remain discoverable.",
        "patterns": ["treatment", "stewardship", "duration_rule"],
        "minimum_tables": 3,
        "minimum_rows": 1,
        "terms": ["do_not_treat", "no antibiotic", "source control", "source_control_only", "reassess", "stop"],
        "minimum_term_hits": 1,
    },
    {
        "gate_key": "resistance_ast_antibiogram",
        "domain": "antimicrobial_resistance",
        "description": "Resistance, susceptibility, isolate AST, breakpoint, MIC, and antibiogram controls remain represented.",
        "patterns": ["resistance", "susceptibility", "antibiogram"],
        "minimum_tables": 2,
        "minimum_rows": 1,
        "terms": ["breakpoint", "susceptibility", "mic", "ast"],
        "minimum_term_hits": 1,
    },
    {
        "gate_key": "evidence_authority_provenance",
        "domain": "evidence",
        "description": "Evidence sources, authority families, source pages, requested extracts, URLs, and interpretation boundaries remain represented.",
        "patterns": ["evidence_source", "source_family", "source_page", "requested_reference", "evidence"],
        "minimum_tables": 3,
        "minimum_rows": 1,
    },
    {
        "gate_key": "graphics_rights_observational_boundary",
        "domain": "graphics",
        "description": "Graphic assets, prompts, placement, provenance, rights, placeholder status, and observational-image boundaries remain governed.",
        "patterns": ["graphic_asset", "graphic_prompt", "graphic_right", "provenance", "rights"],
        "minimum_tables": 2,
        "minimum_rows": 1,
        "terms": ["placeholder", "rights", "provenance", "observational"],
        "minimum_term_hits": 1,
    },
    {
        "gate_key": "publication_navigation_cross_reference",
        "domain": "publication",
        "description": "Publication pages, locators, page maps, and cross-references remain indexed and queryable.",
        "patterns": ["publication_index_locator", "publication_cross_reference", "publication_page", "page_map"],
        "minimum_tables": 2,
        "minimum_rows": 1,
    },
    {
        "gate_key": "tracking_recovery_lineage",
        "domain": "recovery",
        "description": "Response reconciliation, fractional prompts, checkpoints, recovery events, and release lineage remain represented.",
        "patterns": ["thread_response_reconciliation", "fractional_prompt", "remediation_recovery_event", "checkpoint"],
        "minimum_tables": 4,
        "minimum_rows": 1,
    },
    {
        "gate_key": "qa_acceptance_release_controls",
        "domain": "quality_assurance",
        "description": "QA, acceptance, checkpoint, release-freeze, and controlled handoff records remain represented.",
        "patterns": ["qa", "acceptance", "checkpoint", "release", "freeze"],
        "minimum_tables": 3,
        "minimum_rows": 1,
    },
]


def evaluate_release_governance(
    db: Path,
    database_qa: dict[str, Any],
    workbook_qa: dict[str, Any],
    application_qa: dict[str, Any],
    publication_qa: dict[str, Any],
    *,
    current_response: int = 70,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, Any]]:
    con = sqlite3.connect(db)
    try:
        tables = logical_tables(con)
        gates: list[dict[str, Any]] = []
        for spec in DATA_GATE_SPECS:
            matched = sorted({table for table in tables if any(pattern.lower() in table.lower() for pattern in spec["patterns"])})
            counts = {table: safe_count(con, table) for table in matched}
            record_count = sum(value for value in counts.values() if isinstance(value, int))
            hits = term_hits(con, matched, spec.get("terms", [])) if spec.get("terms") else None
            passed = len(matched) >= int(spec["minimum_tables"]) and record_count >= int(spec["minimum_rows"])
            if spec.get("minimum_term_hits") is not None:
                passed = passed and int(hits or 0) >= int(spec["minimum_term_hits"])
            gates.append(
                {
                    "gate_key": spec["gate_key"],
                    "domain": spec["domain"],
                    "description": spec["description"],
                    "matched_tables": matched,
                    "row_counts": counts,
                    "record_count": record_count,
                    "term_hits": hits,
                    "artifact_evidence": {},
                    "status": "passed" if passed else "failed",
                }
            )

        artifact_gates = [
            {
                "gate_key": "database_integrity_and_lineage",
                "domain": "database",
                "description": "Canonical SQLite integrity, foreign keys, current response lineage, and Checkpoint 1 state are valid.",
                "passed": database_qa.get("integrity") == "ok"
                and int(database_qa.get("foreign_key_violations", -1)) == 0
                and int(database_qa.get(f"response{current_response}_records", 0)) == 1
                and database_qa.get("checkpoint_state") == "checkpoint_complete",
                "evidence": database_qa,
            },
            {
                "gate_key": "workbook_integrity_and_parity",
                "domain": "workbook",
                "description": "The comprehensive workbook preserves inherited sheets, contains current managed sheets, and has no formula-error tokens.",
                "passed": workbook_qa.get("status") == "passed"
                and bool(workbook_qa.get("original_sheets_preserved"))
                and int(workbook_qa.get("formula_error_count", -1)) == 0,
                "evidence": workbook_qa,
            },
            {
                "gate_key": "application_acceptance_and_source_invariant",
                "domain": "application",
                "description": "Legacy and current read-only application gates pass while the main application source remains byte-identical.",
                "passed": application_qa.get("status") == "passed" and bool(application_qa.get("application_unchanged")),
                "evidence": application_qa,
            },
            {
                "gate_key": "publication_editable_assembly_invariants",
                "domain": "publication",
                "description": "The integrated publication remains 537/537 searchable pages and both publication and editable assembly retain governed hashes.",
                "passed": publication_qa.get("status") == "passed"
                and int(publication_qa.get("publication_pages", 0)) == 537
                and int(publication_qa.get("searchable_pages", 0)) == 537
                and bool(publication_qa.get("publication_unchanged"))
                and bool(publication_qa.get("editable_assembly_unchanged")),
                "evidence": publication_qa,
            },
        ]
        for item in artifact_gates:
            gates.append(
                {
                    "gate_key": item["gate_key"],
                    "domain": item["domain"],
                    "description": item["description"],
                    "matched_tables": [],
                    "row_counts": {},
                    "record_count": 1,
                    "term_hits": None,
                    "artifact_evidence": item["evidence"],
                    "status": "passed" if item["passed"] else "failed",
                }
            )

        risks = [
            {
                "risk_key": "current_source_recency_and_versioning",
                "domain": "evidence",
                "description": "Final current-source and version-date verification remains required before Section 4 release freeze.",
                "status": "controlled_open",
                "disposition": "Checkpoint 2 performs the final current-source/version sweep and records any evidence changes or explicit no-change results.",
            },
            {
                "risk_key": "observational_graphics_rights_and_scientific_review",
                "domain": "graphics",
                "description": "Observational graphics may remain serialized placeholders until item-level provenance, rights, scientific review, caption, alt text, and diagnostic limitations are complete.",
                "status": "controlled_open",
                "disposition": "Resolve eligible assets or explicitly retain governed placeholders; generated observational imagery may not be presented as clinical evidence.",
            },
            {
                "risk_key": "publication_content_reflow",
                "domain": "publication",
                "description": "Checkpoint 1 preserves the 537-page publication as an immutable baseline rather than silently changing the press artifact.",
                "status": "deferred_checkpoint2",
                "disposition": "Checkpoint 2 reconciles approved clinical/evidence/graphics changes into editable and publication lanes with page-level render QA.",
            },
            {
                "risk_key": "final_cross_artifact_drift_resolution",
                "domain": "governance",
                "description": "All remaining database-workbook-application-publication drift must be classified as resolved, accepted, or prohibited before final release.",
                "status": "monitoring",
                "disposition": "Checkpoint 2 conducts the final detailed drift pass; Checkpoint 3 independently revalidates the release candidate.",
            },
            {
                "risk_key": "cover_spine_and_press_finalization",
                "domain": "production",
                "description": "Final cover wrap and spine arithmetic depend on the final interior page count, stock, binding, printer template, and color profile.",
                "status": "deferred_final_release",
                "disposition": "Finalize only after the interior is frozen; regenerate rather than stretch the provisional wrap.",
            },
            {
                "risk_key": "section4_final_release_signoff",
                "domain": "release",
                "description": "Session 3 Checkpoint 1 establishes readiness controls but does not declare Remediation Section 4 complete.",
                "status": "deferred_final_release",
                "disposition": "Checkpoint 3 requires complete independent QA, clean extraction, transport, persistent custody, and final Section 4 signoff.",
            },
        ]

        failed = [row for row in gates if row["status"] != "passed"]
        qa = {
            "schema": "mrhpd-section4-session3-release-governance-1.0",
            "status": "passed" if not failed else "failed",
            "current_response": current_response,
            "gate_count": len(gates),
            "passed_gates": len(gates) - len(failed),
            "failed_gates": len(failed),
            "failed_gate_keys": [row["gate_key"] for row in failed],
            "controlled_risk_count": len(risks),
            "final_release_declared": False,
            "next_checkpoint": "Remediation Section 4 of 5 Session 3 of 3 Checkpoint 2 of 3",
        }
        if failed:
            raise RuntimeError({"release_governance_failures": failed, "qa": qa})
        return gates, risks, qa
    finally:
        con.close()


def persist_release_governance(
    db: Path,
    gates: list[dict[str, Any]],
    risks: list[dict[str, Any]],
    qa: dict[str, Any],
    *,
    checked_at: str,
) -> dict[str, Any]:
    con = sqlite3.connect(db)
    try:
        con.execute("PRAGMA foreign_keys=ON")
        con.execute("BEGIN IMMEDIATE")
        con.executescript(
            """
            CREATE TABLE IF NOT EXISTS section4_session3_release_governance (
              section4_session3_release_governance_id INTEGER PRIMARY KEY,
              gate_key TEXT NOT NULL UNIQUE,
              domain TEXT NOT NULL,
              description TEXT NOT NULL,
              matched_tables_json TEXT NOT NULL,
              row_counts_json TEXT NOT NULL,
              record_count INTEGER NOT NULL,
              term_hits INTEGER,
              artifact_evidence_json TEXT NOT NULL,
              status TEXT NOT NULL,
              checked_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS section4_session3_release_risk (
              section4_session3_release_risk_id INTEGER PRIMARY KEY,
              risk_key TEXT NOT NULL UNIQUE,
              domain TEXT NOT NULL,
              description TEXT NOT NULL,
              status TEXT NOT NULL,
              disposition TEXT NOT NULL,
              checked_at TEXT NOT NULL
            );
            """
        )
        for row in gates:
            con.execute(
                """
                INSERT INTO section4_session3_release_governance
                (gate_key,domain,description,matched_tables_json,row_counts_json,record_count,term_hits,artifact_evidence_json,status,checked_at)
                VALUES (?,?,?,?,?,?,?,?,?,?)
                ON CONFLICT(gate_key) DO UPDATE SET
                  domain=excluded.domain,description=excluded.description,matched_tables_json=excluded.matched_tables_json,
                  row_counts_json=excluded.row_counts_json,record_count=excluded.record_count,term_hits=excluded.term_hits,
                  artifact_evidence_json=excluded.artifact_evidence_json,status=excluded.status,checked_at=excluded.checked_at
                """,
                (
                    row["gate_key"],
                    row["domain"],
                    row["description"],
                    json.dumps(row["matched_tables"], ensure_ascii=False),
                    json.dumps(row["row_counts"], ensure_ascii=False),
                    int(row["record_count"]),
                    row.get("term_hits"),
                    json.dumps(row.get("artifact_evidence", {}), ensure_ascii=False, default=str),
                    row["status"],
                    checked_at,
                ),
            )
        for row in risks:
            con.execute(
                """
                INSERT INTO section4_session3_release_risk
                (risk_key,domain,description,status,disposition,checked_at)
                VALUES (?,?,?,?,?,?)
                ON CONFLICT(risk_key) DO UPDATE SET
                  domain=excluded.domain,description=excluded.description,status=excluded.status,
                  disposition=excluded.disposition,checked_at=excluded.checked_at
                """,
                (row["risk_key"], row["domain"], row["description"], row["status"], row["disposition"], checked_at),
            )
        if table_exists(con, "metadata"):
            columns = [row[1] for row in con.execute("PRAGMA table_info(metadata)")]
            if {"key", "value"}.issubset(columns):
                updates = {
                    "session3_checkpoint1_release_governance": "passed",
                    "session3_checkpoint1_release_gate_count": str(qa["gate_count"]),
                    "session3_checkpoint1_controlled_risk_count": str(qa["controlled_risk_count"]),
                    "section4_final_release_declared": "no",
                }
                for key, value in updates.items():
                    con.execute("INSERT INTO metadata(key,value) VALUES (?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value", (key, value))
        integrity = con.execute("PRAGMA integrity_check").fetchone()[0]
        fk = list(con.execute("PRAGMA foreign_key_check"))
        failures = con.execute("SELECT COUNT(*) FROM section4_session3_release_governance WHERE status!='passed'").fetchone()[0]
        gate_count = con.execute("SELECT COUNT(*) FROM section4_session3_release_governance").fetchone()[0]
        risk_count = con.execute("SELECT COUNT(*) FROM section4_session3_release_risk").fetchone()[0]
        if integrity != "ok" or fk or failures or gate_count != len(gates) or risk_count != len(risks):
            raise RuntimeError(
                {
                    "integrity": integrity,
                    "foreign_keys": fk[:20],
                    "release_gate_failures": failures,
                    "gate_count": gate_count,
                    "risk_count": risk_count,
                }
            )
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()
    result = dict(qa)
    result.update(
        {
            "database_gate_records": len(gates),
            "database_risk_records": len(risks),
            "database_bytes": db.stat().st_size,
            "database_sha256": sha256_file(db),
        }
    )
    return result


def augment_workbook(
    workbook: Path,
    workbook_qa: dict[str, Any],
    gates: list[dict[str, Any]],
    risks: list[dict[str, Any]],
) -> tuple[Path, dict[str, Any]]:
    from openpyxl import load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    wb = load_workbook(workbook)
    managed = ["S4S3 Release Gates", "S4S3 Risk Register"]
    for name in managed:
        if name in wb.sheetnames:
            del wb[name]

    navy = "17365D"
    teal = "1F6D73"
    white = "FFFFFF"
    pale = "EAF3F3"

    gate_ws = wb.create_sheet("S4S3 Release Gates")
    gate_ws.append(["Section 4 Session 3 Checkpoint 1", "Final-release governance and readiness baseline"])
    gate_ws.append(["Gate", "Domain", "Description", "Matched tables", "Record count", "Term hits", "Status"])
    for row in gates:
        gate_ws.append(
            [
                row["gate_key"],
                row["domain"],
                row["description"],
                "\n".join(row["matched_tables"]),
                row["record_count"],
                row.get("term_hits"),
                row["status"],
            ]
        )
    gate_ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=7)
    gate_ws["A1"].font = Font(bold=True, color=white, size=14)
    gate_ws["A1"].fill = PatternFill("solid", fgColor=teal)
    gate_ws["A1"].alignment = Alignment(horizontal="center")
    for cell in gate_ws[2]:
        cell.font = Font(bold=True, color=white)
        cell.fill = PatternFill("solid", fgColor=navy)
    gate_ws.freeze_panes = "A3"
    gate_ws.auto_filter.ref = f"A2:G{gate_ws.max_row}"

    risk_ws = wb.create_sheet("S4S3 Risk Register")
    risk_ws.append(["Risk key", "Domain", "Description", "Status", "Disposition"])
    for row in risks:
        risk_ws.append([row["risk_key"], row["domain"], row["description"], row["status"], row["disposition"]])
    for cell in risk_ws[1]:
        cell.font = Font(bold=True, color=white)
        cell.fill = PatternFill("solid", fgColor=navy)
    risk_ws.freeze_panes = "A2"
    risk_ws.auto_filter.ref = risk_ws.dimensions

    for ws in (gate_ws, risk_ws):
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        for index, width in enumerate((34, 24, 76, 52, 18, 16, 14), 1):
            if index <= ws.max_column:
                ws.column_dimensions[chr(64 + index)].width = width
        for row_number in range(1, ws.max_row + 1):
            if row_number % 2 == 1 and row_number > 2:
                for cell in ws[row_number]:
                    cell.fill = PatternFill("solid", fgColor=pale)
    wb.save(workbook)

    check = load_workbook(workbook, read_only=True, data_only=False)
    final_sheets = list(check.sheetnames)
    formula_errors: list[str] = []
    error_tokens = ("#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "#N/A", "#NUM!", "#NULL!")
    formula_count = 0
    for ws in check.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                value = cell.value
                if isinstance(value, str) and value.startswith("="):
                    formula_count += 1
                if isinstance(value, str) and any(token in value for token in error_tokens):
                    formula_errors.append(f"{ws.title}!{cell.coordinate}:{value}")
    check.close()

    result = dict(workbook_qa)
    result["current_sheet_count"] = len(final_sheets)
    result["bytes"] = workbook.stat().st_size
    result["sha256"] = sha256_file(workbook)
    result["formula_count"] = formula_count
    result["formula_error_count"] = len(formula_errors)
    result["formula_errors"] = formula_errors
    result["release_governance_sheets"] = managed
    result["managed_sheets"] = list(dict.fromkeys(list(result.get("managed_sheets", [])) + managed))
    result["missing_managed_sheets"] = [name for name in result["managed_sheets"] if name not in final_sheets]
    result["status"] = "passed" if not formula_errors and not result["missing_managed_sheets"] else "failed"
    if result["status"] != "passed":
        raise RuntimeError({"release_governance_workbook_qa": result})
    return workbook, result


def build_release_readiness_audit(
    project: Path,
    db: Path,
    workbook: Path,
    application_sha256: str,
    publication_sha256: str,
    editable_sha256: str,
    *,
    generated_at: str,
) -> tuple[Path, dict[str, Any]]:
    audit = project / "App" / "section4_session3_release_readiness_audit.py"
    output = project / "QA" / "Section 4 Session 3" / "Checkpoint 1" / "APPLICATION_RELEASE_READINESS_AUDIT.json"
    expected = {
        "response": 70,
        "checkpoint_code": "MRHPD-V3-CP4-S3-CP1",
        "minimum_gates": 17,
        "minimum_risks": 6,
        "application_sha256": application_sha256,
        "publication_sha256": publication_sha256,
        "editable_sha256": editable_sha256,
        "workbook_sha256": sha256_file(workbook),
    }
    source = f'''#!/usr/bin/env python3
from __future__ import annotations
import argparse,hashlib,json,sqlite3
from pathlib import Path
EXPECTED={json.dumps(expected, ensure_ascii=False)}
def sha(path):
 h=hashlib.sha256()
 with path.open('rb') as handle:
  for block in iter(lambda:handle.read(1024*1024),b''): h.update(block)
 return h.hexdigest()
def exists(con,table):
 return con.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?",(table,)).fetchone() is not None
def find_hash(root,pattern,digest):
 return [path for path in root.rglob(pattern) if path.is_file() and sha(path)==digest]
p=argparse.ArgumentParser(description='MRHPD Section 4 Session 3 release-readiness audit')
p.add_argument('--db',type=Path,required=True)
p.add_argument('--output',type=Path)
a=p.parse_args()
root=Path(__file__).resolve().parents[1]
con=sqlite3.connect(a.db)
try:
 integrity=con.execute('PRAGMA integrity_check').fetchone()[0]
 fk=list(con.execute('PRAGMA foreign_key_check'))
 gates=con.execute("SELECT COUNT(*) FROM section4_session3_release_governance").fetchone()[0] if exists(con,'section4_session3_release_governance') else 0
 gate_failures=con.execute("SELECT COUNT(*) FROM section4_session3_release_governance WHERE status!='passed'").fetchone()[0] if gates else 1
 risks=con.execute("SELECT COUNT(*) FROM section4_session3_release_risk").fetchone()[0] if exists(con,'section4_session3_release_risk') else 0
 bad_risks=con.execute("SELECT COUNT(*) FROM section4_session3_release_risk WHERE status NOT IN ('controlled_open','deferred_checkpoint2','deferred_final_release','monitoring')").fetchone()[0] if risks else 1
 response=con.execute("SELECT COUNT(*) FROM thread_response_reconciliation_cp3 WHERE response_key='R70'").fetchone()[0]
 checkpoint=con.execute("SELECT state FROM section4_session3_checkpoint WHERE checkpoint_code=?",(EXPECTED['checkpoint_code'],)).fetchone() if exists(con,'section4_session3_checkpoint') else None
 locators=con.execute("SELECT COUNT(*) FROM publication_index_locator").fetchone()[0] if exists(con,'publication_index_locator') else None
 crossrefs=con.execute("SELECT COUNT(*) FROM publication_cross_reference").fetchone()[0] if exists(con,'publication_cross_reference') else None
 prior_failures={{}}
 for table in ('section4_session2_field_coverage','section4_session2_query_coverage','section4_session2_source_governance','section4_session2_drift_resolution'):
  if exists(con,table):
   prior_failures[table]=con.execute(f'SELECT COUNT(*) FROM "{{table}}" WHERE status!=\'passed\'').fetchone()[0]
 checks={{
  'integrity':integrity=='ok',
  'foreign_keys':not fk,
  'response70':response==1,
  'checkpoint':checkpoint==('checkpoint_complete',),
  'release_gates':gates>=EXPECTED['minimum_gates'] and gate_failures==0,
  'controlled_risks':risks>=EXPECTED['minimum_risks'] and bad_risks==0,
  'prior_detailed_governance':all(value==0 for value in prior_failures.values()),
  'publication_locators':locators in (None,4011),
  'publication_cross_references':crossrefs is None or crossrefs>=12,
  'application_source':len(find_hash(root,'human_pathogen_app.py',EXPECTED['application_sha256']))>=1,
  'publication_invariant':len(find_hash(root,'*Integrated Manuscript*.pdf',EXPECTED['publication_sha256']))>=1,
  'editable_invariant':len(find_hash(root,'*Editable Integrated Manuscript Assembly*.docx',EXPECTED['editable_sha256']))>=1,
  'workbook_current':len(find_hash(root,'*.xlsx',EXPECTED['workbook_sha256']))>=1,
 }}
finally:
 con.close()
result={{
 'schema':'mrhpd-section4-session3-release-readiness-audit-1.0',
 'generated_at':{generated_at!r},
 'status':'passed' if all(checks.values()) else 'failed',
 'database':str(a.db),
 'database_sha256':sha(a.db),
 'checks':checks,
 'release_gate_count':gates,
 'controlled_risk_count':risks,
 'publication_index_locators':locators,
 'publication_cross_references':crossrefs,
 'prior_detailed_failures':prior_failures,
}}
if a.output:
 a.output.parent.mkdir(parents=True,exist_ok=True)
 a.output.write_text(json.dumps(result,indent=2)+"\\n",encoding='utf-8')
print(json.dumps(result,indent=2))
raise SystemExit(0 if result['status']=='passed' else 1)
'''
    text_write(audit, source)
    result = subprocess.run([sys.executable, str(audit), "--db", str(db), "--output", str(output)], text=True, capture_output=True, timeout=300)
    if result.returncode != 0:
        raise RuntimeError({"release_readiness_audit_failed": {"stdout": result.stdout[-12000:], "stderr": result.stderr[-12000:]}})
    qa = json.loads(output.read_text(encoding="utf-8"))
    qa["audit_path"] = audit.relative_to(project).as_posix()
    qa["output_path"] = output.relative_to(project).as_posix()
    qa["audit_sha256"] = sha256_file(audit)
    return audit, qa


def build_release_reports(
    project: Path,
    gates: list[dict[str, Any]],
    risks: list[dict[str, Any]],
    release_qa: dict[str, Any],
    database_qa: dict[str, Any],
    workbook_qa: dict[str, Any],
    application_qa: dict[str, Any],
    publication_qa: dict[str, Any],
    *,
    generated_at: str,
) -> tuple[list[Path], dict[str, Any]]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from pypdf import PdfReader
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    report_dir = project / "Reports" / "Section 4 Session 3" / "Checkpoint 1"
    report_dir.mkdir(parents=True, exist_ok=True)
    stem = "MRHPD v3.0.0a Section 4 Session 3 Checkpoint 1 Final-Release Governance and Readiness Baseline"
    docx_path = report_dir / f"{stem}.docx"
    pdf_path = report_dir / f"{stem}.pdf"
    xlsx_path = report_dir / f"{stem} Register.xlsx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    title = doc.add_heading("Human Pathogen Database", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Section 4 Session 3 — Final-Release Governance and Readiness Baseline")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Checkpoint 1 of 3 • Response 70 • {generated_at}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Executive disposition", level=1)
    doc.add_paragraph(
        f"All {release_qa['gate_count']} release-governance gates passed. "
        f"The {release_qa['controlled_risk_count']} remaining items are controlled forward-work risks, not silent omissions. "
        "Checkpoint 1 does not declare Section 4 complete; it establishes the final-session acceptance framework."
    )
    doc.add_heading("Release-governance gates", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for index, heading in enumerate(("Gate", "Domain", "Evidence", "Records", "Status")):
        table.rows[0].cells[index].text = heading
    for row in gates:
        cells = table.add_row().cells
        cells[0].text = row["gate_key"]
        cells[1].text = row["domain"]
        cells[2].text = ", ".join(row["matched_tables"][:8]) or "artifact QA"
        cells[3].text = str(row["record_count"])
        cells[4].text = row["status"].upper()
    doc.add_heading("Controlled forward-work risks", level=1)
    risk_table = doc.add_table(rows=1, cols=4)
    risk_table.style = "Table Grid"
    for index, heading in enumerate(("Risk", "Domain", "Status", "Disposition")):
        risk_table.rows[0].cells[index].text = heading
    for row in risks:
        cells = risk_table.add_row().cells
        cells[0].text = row["risk_key"]
        cells[1].text = row["domain"]
        cells[2].text = row["status"]
        cells[3].text = row["disposition"]
    doc.add_heading("Artifact invariants and acceptance", level=1)
    doc.add_paragraph(
        f"SQLite integrity: {database_qa.get('integrity')}; foreign-key violations: {database_qa.get('foreign_key_violations')}; "
        f"workbook sheets: {workbook_qa.get('current_sheet_count')}; formula-error tokens: {workbook_qa.get('formula_error_count')}; "
        f"application status: {application_qa.get('status')}; publication pages/searchable pages: "
        f"{publication_qa.get('publication_pages')}/{publication_qa.get('searchable_pages')}."
    )
    doc.add_heading("Next checkpoint", level=1)
    doc.add_paragraph(
        "Checkpoint 2 performs the final current-source/version sweep, detailed cross-artifact drift resolution, approved publication and "
        "graphics reconciliation, and expanded application/workbook parity before the independently verified Section 4 release candidate."
    )
    doc.save(docx_path)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="SmallGov", parent=styles["BodyText"], fontSize=7.3, leading=9.2))
    story: list[Any] = [
        Paragraph("Human Pathogen Database", styles["Title"]),
        Paragraph("Section 4 Session 3 — Final-Release Governance and Readiness Baseline", styles["Heading2"]),
        Paragraph(f"Checkpoint 1 of 3 • Response 70 • {generated_at}", styles["BodyText"]),
        Spacer(1, 0.18 * inch),
        Paragraph(
            f"All {release_qa['gate_count']} release-governance gates passed. The {release_qa['controlled_risk_count']} remaining items are controlled forward-work risks; Section 4 is not yet declared complete.",
            styles["BodyText"],
        ),
        Spacer(1, 0.15 * inch),
        Paragraph("Release-governance gates", styles["Heading2"]),
    ]
    gate_data: list[list[Any]] = [["Gate", "Domain", "Evidence", "Status"]]
    for row in gates:
        gate_data.append(
            [
                Paragraph(row["gate_key"], styles["SmallGov"]),
                Paragraph(row["domain"], styles["SmallGov"]),
                Paragraph(", ".join(row["matched_tables"][:5]) or "artifact QA", styles["SmallGov"]),
                row["status"].upper(),
            ]
        )
    gate_table = Table(gate_data, colWidths=[1.65 * inch, 1.15 * inch, 3.2 * inch, 0.7 * inch], repeatRows=1)
    gate_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17365D")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EAF3F3")]),
            ]
        )
    )
    story.extend([gate_table, PageBreak(), Paragraph("Controlled forward-work risks", styles["Heading2"])])
    for row in risks:
        story.append(Paragraph(f"<b>{row['risk_key']}</b> — {row['status']}", styles["Heading3"]))
        story.append(Paragraph(row["description"], styles["BodyText"]))
        story.append(Paragraph(f"Disposition: {row['disposition']}", styles["SmallGov"]))
        story.append(Spacer(1, 0.08 * inch))
    story.extend(
        [
            Paragraph("Artifact acceptance", styles["Heading2"]),
            Paragraph(
                f"SQLite integrity {database_qa.get('integrity')}; foreign-key violations {database_qa.get('foreign_key_violations')}; "
                f"workbook {workbook_qa.get('current_sheet_count')} sheets with {workbook_qa.get('formula_error_count')} formula-error tokens; "
                f"application {application_qa.get('status')}; publication {publication_qa.get('publication_pages')}/{publication_qa.get('searchable_pages')} searchable pages.",
                styles["BodyText"],
            ),
        ]
    )
    SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=0.55 * inch, leftMargin=0.55 * inch, topMargin=0.55 * inch, bottomMargin=0.55 * inch).build(story)

    wb = Workbook()
    gates_ws = wb.active
    gates_ws.title = "Release Gates"
    gates_ws.append(["gate_key", "domain", "description", "matched_tables", "row_counts", "record_count", "term_hits", "status"])
    for row in gates:
        gates_ws.append(
            [
                row["gate_key"],
                row["domain"],
                row["description"],
                "\n".join(row["matched_tables"]),
                json.dumps(row["row_counts"], ensure_ascii=False),
                row["record_count"],
                row.get("term_hits"),
                row["status"],
            ]
        )
    risks_ws = wb.create_sheet("Risk Register")
    risks_ws.append(["risk_key", "domain", "description", "status", "disposition"])
    for row in risks:
        risks_ws.append([row["risk_key"], row["domain"], row["description"], row["status"], row["disposition"]])
    qa_ws = wb.create_sheet("QA")
    qa_ws.append(["key", "value"])
    for prefix, payload in (("release", release_qa), ("database", database_qa), ("workbook", workbook_qa), ("application", application_qa), ("publication", publication_qa)):
        for key, value in payload.items():
            qa_ws.append([f"{prefix}.{key}", json.dumps(value, ensure_ascii=False, default=str) if isinstance(value, (dict, list)) else value])
    for ws in wb.worksheets:
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="17365D")
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        for column in ws.columns:
            letter = column[0].column_letter
            ws.column_dimensions[letter].width = min(70, max(12, max(len(str(cell.value or "")) for cell in column) + 2))
    wb.save(xlsx_path)

    reader = PdfReader(str(pdf_path))
    searchable = sum(1 for page in reader.pages if (page.extract_text() or "").strip())
    report_qa = {
        "schema": "mrhpd-section4-session3-release-governance-report-qa-1.0",
        "status": "passed" if reader.pages and searchable == len(reader.pages) else "failed",
        "docx": {"path": docx_path.relative_to(project).as_posix(), "bytes": docx_path.stat().st_size, "sha256": sha256_file(docx_path)},
        "pdf": {"path": pdf_path.relative_to(project).as_posix(), "bytes": pdf_path.stat().st_size, "sha256": sha256_file(pdf_path), "pages": len(reader.pages), "searchable_pages": searchable},
        "xlsx": {"path": xlsx_path.relative_to(project).as_posix(), "bytes": xlsx_path.stat().st_size, "sha256": sha256_file(xlsx_path)},
    }
    if report_qa["status"] != "passed":
        raise RuntimeError({"release_governance_report_qa": report_qa})
    return [docx_path, pdf_path, xlsx_path], report_qa


def write_release_qa(
    project: Path,
    gates: list[dict[str, Any]],
    risks: list[dict[str, Any]],
    release_qa: dict[str, Any],
    audit_qa: dict[str, Any],
    report_qa: dict[str, Any],
    *,
    generated_at: str,
) -> list[Path]:
    qa_dir = project / "QA" / "Section 4 Session 3" / "Checkpoint 1"
    qa_dir.mkdir(parents=True, exist_ok=True)
    gates_json = qa_dir / "RELEASE_GOVERNANCE_GATES.json"
    gates_csv = qa_dir / "RELEASE_GOVERNANCE_GATES.csv"
    risks_json = qa_dir / "RELEASE_RISK_REGISTER.json"
    risks_csv = qa_dir / "RELEASE_RISK_REGISTER.csv"
    readiness_json = qa_dir / "RELEASE_READINESS_QA.json"
    handoff = project / "Recovery" / "Section 4 Session 3 Checkpoint 1" / "CHECKPOINT_2_HANDOFF.md"
    json_write(gates_json, gates)
    csv_write(gates_csv, gates, ["gate_key", "domain", "description", "matched_tables", "row_counts", "record_count", "term_hits", "status"])
    json_write(risks_json, risks)
    csv_write(risks_csv, risks, ["risk_key", "domain", "description", "status", "disposition"])
    combined = dict(release_qa)
    combined.update({"generated_at": generated_at, "application_release_audit": audit_qa, "report_qa": report_qa})
    json_write(readiness_json, combined)
    text_write(
        handoff,
        f"""# Section 4 Session 3 Checkpoint 1 to Checkpoint 2 Handoff

Checkpoint 1 establishes the final-session release-governance baseline through Response 70.

- Release-governance gates: {release_qa['passed_gates']}/{release_qa['gate_count']} passed
- Controlled forward-work risks: {release_qa['controlled_risk_count']}
- Section 4 final release declared: no
- Next: Remediation Section 4 of 5 Session 3 of 3 Checkpoint 2 of 3

Checkpoint 2 must perform the final current-source/version sweep, resolve or explicitly accept every remaining cross-artifact drift item, reconcile approved publication and graphics changes with page-level render QA, and rerun application/workbook/database parity before the release-candidate checkpoint.
""",
    )
    return [gates_json, gates_csv, risks_json, risks_csv, readiness_json, handoff]
